import sys, os, time, re
sys.stdout.reconfigure(encoding='utf-8')
sys.path.insert(0, r'd:\UNIGO\scripts')
from docx import Document
from fix_borders_and_a1 import make_borderless_table, clean_a1_duplication

target = r'd:\UNIGO\KHBD_Tin_học\Lớp_7\Tuần_03\KHBD_Tin_hoc_Lớp_7_Bai02_Bai_2_Phan_mem_may_tinh.docx'
doc = Document(target)

# 1. Borderless table 0
if doc.tables:
    make_borderless_table(doc.tables[0])
    for r in doc.tables[0].rows:
        for c in r.cells:
            for p in c.paragraphs:
                for run in p.runs:
                    run.text = clean_a1_duplication(run.text)

# 2. Fix A1 in paragraphs
for p in doc.paragraphs:
    for run in p.runs:
        run.text = clean_a1_duplication(run.text)

# 3. Clean duplicated header block
if len(doc.paragraphs) >= 4:
    p0_text = doc.paragraphs[0].text.strip()
    p2_text = doc.paragraphs[2].text.strip()
    if p0_text.startswith("TÊN BÀI DẠY: BÀI 2") and p2_text == "Tiết theo PPCT: 1":
        print("Removing duplicated header paragraphs...")
        p1_elem = doc.paragraphs[1]._p
        p2_elem = doc.paragraphs[2]._p
        p3_elem = doc.paragraphs[3]._p
        p1_elem.getparent().remove(p1_elem)
        p2_elem.getparent().remove(p2_elem)
        p3_elem.getparent().remove(p3_elem)

try:
    doc.save(target)
    print("✅ ĐÃ LƯU TRỰC TIẾP THÀNH CÔNG VÀO FILE CHÍNH!")
except PermissionError:
    alt = target.replace(".docx", "_fixed.docx")
    doc.save(alt)
    print(f"⚠️ File đang mở trong Word, đã lưu bản sửa vào:\n   {alt}")
