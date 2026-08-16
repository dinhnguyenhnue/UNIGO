"""
generate_khbd_tin_hoc_full.py
=============================
Tự động tạo toàn bộ KHBD môn Tin học (Tiền Tiểu học -> Lớp 9)
theo đúng chuẩn PL4 (CV 5512), Thông tư 02/2025 (CV 3456 - 6 Miền NLS),
căn lề, thụt đầu dòng EMU, bảng 3 cột cho từng hoạt động dạy học,
và cấu trúc thư mục phân bổ theo TUẦN để đồng bộ với Lịch báo giảng.
"""

import os
import re
import sys
import shutil
from datetime import date, timedelta
import docx
from docx import Document
from docx.shared import Pt, Cm, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

sys.stdout.reconfigure(encoding='utf-8')

# ============================================================
# CONFIG
# ============================================================
TEMPLATE_PATH = r"D:\UNIGO\Hệ thống mẫu văn bản\PL4-Khung kế hoạch bài dạy (THCS).docx"
PPCT_DOC_PATH = r"D:\UNIGO\Phân phối chương trình\Tin học\Kế hoạch dạy học môn Tin học 2026-2027.docx"
OUT_BASE_DIR = r"D:\UNIGO\KHBD_Tin_học"

# Ngày đầu Tuần 01 (Thứ Hai 03/08/2026)
TUAN_01_START = date(2026, 8, 3)

# Indent EMU values
INDENT_0 = 0           # Tiêu đề section (I., II., III.)
INDENT_1 = 180340      # Mục con cấp 1 (1., 2., 3.)
INDENT_2 = 360045      # Mục con cấp 2 (2.1., 2.2., 2.3.)
INDENT_BULLET = 540000 # Bullet content (- NLa..., - Miền I...)

FONT_NAME = "Times New Roman"
FONT_SIZE = Pt(13)
LINE_SPACING = 1.15
SPACE_AFTER = Emu(38100)  # 3pt


# ============================================================
# HELPER FUNCTIONS
# ============================================================
def set_font(run, bold=False, italic=False, size=FONT_SIZE):
    run.font.name = FONT_NAME
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)


def add_paragraph(doc, text, bold=False, italic=False, first_indent=None,
                  left_indent=None, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_after=SPACE_AFTER, space_before=0):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = space_after
    pf.space_before = Emu(space_before) if isinstance(space_before, int) else space_before

    if first_indent is not None:
        pf.first_line_indent = Emu(first_indent)
    if left_indent is not None:
        pf.left_indent = Emu(left_indent)

    if text:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic)
    return p


def add_multi_run_paragraph(doc, runs_data, first_indent=None, left_indent=None,
                            alignment=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = SPACE_AFTER
    pf.space_before = Emu(0)

    if first_indent is not None:
        pf.first_line_indent = Emu(first_indent)
    if left_indent is not None:
        pf.left_indent = Emu(left_indent)

    for text, bold, italic in runs_data:
        run = p.add_run(text)
        set_font(run, bold=bold, italic=italic)
    return p


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
        '</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def set_table_no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="nil"/>'
        '  <w:left w:val="nil"/>'
        '  <w:bottom w:val="nil"/>'
        '  <w:right w:val="nil"/>'
        '  <w:insideH w:val="nil"/>'
        '  <w:insideV w:val="nil"/>'
        '</w:tblBorders>'
    )
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(borders)


def set_cell_text(cell, text, bold=False, italic=False, alignment=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = LINE_SPACING
    pf.space_after = Emu(25400)  # 2pt
    pf.space_before = Emu(0)
    
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            p = cell.add_paragraph()
            p.alignment = alignment
            pf = p.paragraph_format
            pf.line_spacing = LINE_SPACING
            pf.space_after = Emu(25400)
            pf.space_before = Emu(0)
        run = p.add_run(line)
        set_font(run, bold=bold, italic=italic)


def add_3col_table(doc, buoc_data, is_last=False):
    buoc_labels = [
        'Bước 1:\nChuyển giao\nnhiệm vụ',
        'Bước 2:\nHọc sinh tiếp nhận\nnhiệm vụ',
        'Bước 3:\nBáo cáo kết quả\nhoạt động',
        'Bước 4:\nĐánh giá kết quả\nthực hiện nhiệm vụ'
    ]
    if is_last:
        buoc_labels[3] = 'Bước 4:\nGiáo viên nhắc nhở\nnhiệm vụ về nhà'

    table = doc.add_table(rows=1, cols=3)
    set_table_borders(table)

    # Header row
    hdr = table.rows[0]
    for i, label in enumerate(['Bước', 'Hoạt động của GV', 'Hoạt động của HS']):
        set_cell_text(hdr.cells[i], label, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)

    # Data rows
    for i, (gv, hs) in enumerate(buoc_data):
        row = table.add_row()
        set_cell_text(row.cells[0], buoc_labels[i], bold=True, italic=True,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(row.cells[1], gv)
        set_cell_text(row.cells[2], hs)

    for row in table.rows:
        row.cells[0].width = Cm(3.2)
        row.cells[1].width = Cm(6.4)
        row.cells[2].width = Cm(6.4)

    return table


def sanitize_filename(name):
    import unicodedata
    name = unicodedata.normalize('NFD', name).encode('ascii', 'ignore').decode('utf-8')
    name = re.sub(r'[^a-zA-Z0-9_\-]', '_', name)
    name = re.sub(r'_+', '_', name).strip('_')
    return name


def get_dates_for_week(tuan_so):
    start = TUAN_01_START + timedelta(weeks=tuan_so - 1)
    ngay_soan = start.strftime('%d/%m/%Y')
    ngay_day = (start + timedelta(days=2)).strftime('%d/%m/%Y')
    return ngay_soan, ngay_day


# ============================================================
# DOCUMENT BUILDER
# ============================================================
def build_khbd_tin_hoc(grade_str, lesson_title, lesson_idx, tiet_ppct, yccd, tuan_so):
    """Tạo 1 file KHBD Tin học chuẩn theo đúng cấu trúc PL4 & NLS 6 miền."""
    doc = Document(TEMPLATE_PATH)

    # Clear body while keeping sectPr (preserves headers/footers)
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'sectPr':
            body.remove(child)

    ngay_soan, ngay_day = get_dates_for_week(tuan_so)
    is_thcs = grade_str in ['Lớp 6', 'Lớp 7', 'Lớp 8', 'Lớp 9']
    to_chuyen_mon = "Tổ chuyên môn THCS" if is_thcs else "Tổ chuyên môn Tiểu học"
    bac_nls = "Bậc 2" if is_thcs else "Bậc 1"

    # 1. Info Table (2x2) - NO BORDER
    tbl_info = doc.add_table(rows=2, cols=2)
    set_table_no_borders(tbl_info)
    set_cell_text(tbl_info.rows[0].cells[0], 'Trường: TH&THCS UNIGO', bold=True)
    set_cell_text(tbl_info.rows[0].cells[1], f'Ngày soạn: {ngay_soan}', bold=True)
    set_cell_text(tbl_info.rows[1].cells[0], f'GV: Đậu Đình Nguyên\nTổ: {to_chuyen_mon}', bold=False)
    set_cell_text(tbl_info.rows[1].cells[1], f'Ngày dạy: {ngay_day}\nLớp: {grade_str}', bold=False)

    # 2. Title Block
    add_paragraph(doc, '', bold=False)
    add_paragraph(doc, f'TÊN BÀI DẠY: {lesson_title.upper()}', bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=INDENT_0)
    add_paragraph(doc, f'Môn học: Tin học     Lớp: {grade_str}     Thời lượng: 1 tiết (45 phút)', bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=INDENT_0)
    add_paragraph(doc, f'Tiết theo PPCT: {tiet_ppct}', bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=INDENT_0)
    add_paragraph(doc, f'Tên tiết: {lesson_title}', bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=INDENT_0)

    # 3. I. Mục tiêu
    add_paragraph(doc, 'I. Mục tiêu', bold=True, first_indent=INDENT_0)

    # 1. Kiến thức
    add_paragraph(doc, '1. Kiến thức:', bold=True, first_indent=INDENT_1)
    if yccd and len(yccd.strip()) > 10:
        yccd_clean = yccd.strip().lstrip('-').strip()
        add_paragraph(doc, f'- Sự hiểu biết về nội dung: {yccd_clean}.',
                      first_indent=0, left_indent=INDENT_BULLET)
    else:
        add_paragraph(doc, f'- Sự hiểu biết về các khái niệm, quy trình và kỹ năng thực hành trong bài {lesson_title}.',
                      first_indent=0, left_indent=INDENT_BULLET)
    add_paragraph(doc, f'- Khả năng nhận diện và vận dụng kiến thức bài học để thực hành thao tác trên máy tính/thiết bị số.',
                  first_indent=0, left_indent=INDENT_BULLET)

    # 2. Năng lực
    add_paragraph(doc, '2. Năng lực:', bold=True, first_indent=INDENT_1)

    # 2.1. Năng lực đặc thù (Tin học)
    add_paragraph(doc, '2.1. Năng lực đặc thù (Tin học):', bold=True, first_indent=INDENT_2)
    add_paragraph(doc,
        f'- NLa (Sử dụng và quản lí các phương tiện ICT): Nhận diện, phân tích và thao tác đúng quy trình các công cụ, phần mềm trong bài {lesson_title}. (Đạt được thông qua Hoạt động 2)',
        first_indent=0, left_indent=INDENT_BULLET)
    add_paragraph(doc,
        f'- NLc (Giải quyết vấn đề với sự hỗ trợ của ICT): Vận dụng kiến thức, kỹ năng đã học để hoàn thành nhiệm vụ thực hành và xử lý tình huống thực tế. (Đạt được thông qua Hoạt động 3)',
        first_indent=0, left_indent=INDENT_BULLET)

    # 2.2. Năng lực số (Thông tư 02/2025 – CV 3456)
    add_paragraph(doc, '2.2. Năng lực số (Thông tư 02/2025 – CV 3456):', bold=True, first_indent=INDENT_2)
    add_paragraph(doc,
        f'- Miền I. Khai thác dữ liệu và thông tin (thành tố 1.1. Duyệt, tìm kiếm và lọc dữ liệu, thông tin và nội dung số – {bac_nls}): Khai thác thông tin số, dữ liệu đa phương tiện an toàn phục vụ bài học. (Đạt được thông qua Hoạt động 2)',
        first_indent=0, left_indent=INDENT_BULLET)
    add_paragraph(doc,
        f'- Miền IV. An toàn (thành tố 4.1. Bảo vệ thiết bị – {bac_nls}): Tuân thủ quy tắc an toàn về điện, bảo vệ mắt và giữ vệ sinh thiết bị phòng máy. (Đạt được thông qua Hoạt động 3, Hoạt động 4)',
        first_indent=0, left_indent=INDENT_BULLET)

    # 2.3. Năng lực chung
    add_paragraph(doc, '2.3. Năng lực chung:', bold=True, first_indent=INDENT_2)
    add_paragraph(doc,
        f'- Tự chủ và tự học: Chủ động đọc tài liệu SGK, quan sát thao tác mẫu và tự thực hiện nhiệm vụ học tập. (Đạt được thông qua Hoạt động 1, Hoạt động 2)',
        first_indent=0, left_indent=INDENT_BULLET)
    add_paragraph(doc,
        f'- Giao tiếp và hợp tác: Thảo luận nhóm hiệu quả, chia sẻ và hỗ trợ bạn trong giờ thực hành máy tính. (Đạt được thông qua Hoạt động 2, Hoạt động 3)',
        first_indent=0, left_indent=INDENT_BULLET)
    add_paragraph(doc,
        f'- Giải quyết vấn đề và sáng tạo: Phát hiện và sửa lỗi trong quá trình thực hành, đề xuất ý tưởng ứng dụng mới. (Đạt được thông qua Hoạt động 4)',
        first_indent=0, left_indent=INDENT_BULLET)

    # 3. Phẩm chất
    add_paragraph(doc, '3. Phẩm chất:', bold=True, first_indent=INDENT_1)
    add_paragraph(doc,
        f'- Chăm chỉ: Tích cực tham gia các hoạt động tìm hiểu bài học, hăng hái phát biểu và hoàn thành bài tập thực hành. (Thông qua Hoạt động 1, Hoạt động 2)',
        first_indent=0, left_indent=INDENT_BULLET)
    add_paragraph(doc,
        f'- Trung thực: Đánh giá đúng kết quả học tập của bản thân và bạn bè, tôn trọng sản phẩm số của người khác. (Thông qua Hoạt động 3)',
        first_indent=0, left_indent=INDENT_BULLET)
    add_paragraph(doc,
        f'- Trách nhiệm: Giữ gìn an toàn thiết bị phòng máy, có ý thức bảo vệ tài sản công cộng và dữ liệu số. (Thông qua Hoạt động 3, Hoạt động 4)',
        first_indent=0, left_indent=INDENT_BULLET)

    # 4. II. Thiết bị dạy học và học liệu
    add_paragraph(doc, 'II. Thiết bị dạy học và học liệu:', bold=True, first_indent=INDENT_0)
    add_paragraph(doc, '1. Thiết bị:', bold=True, first_indent=INDENT_1)
    add_paragraph(doc, '- Máy tính giáo viên có kết nối Internet và máy chiếu.',
                  first_indent=0, left_indent=INDENT_BULLET)
    add_paragraph(doc, '- Phòng máy tính cho học sinh thực hành.',
                  first_indent=0, left_indent=INDENT_BULLET)

    add_paragraph(doc, '2. Học liệu:', bold=True, first_indent=INDENT_1)
    add_paragraph(doc, f'- SGK Tin học {grade_str}, bài giảng điện tử tương tác.',
                  first_indent=0, left_indent=INDENT_BULLET)
    add_paragraph(doc, '- Phiếu học tập, tệp bài tập thực hành mẫu.',
                  first_indent=0, left_indent=INDENT_BULLET)

    # 5. III. Tiến trình dạy học
    add_paragraph(doc, 'III. Tiến trình dạy học', bold=True, first_indent=INDENT_0)

    # HĐ 1. Khởi động
    add_paragraph(doc, '1. Hoạt động 1. Khởi động (Xác định vấn đề/nhiệm vụ học tập/Mở đầu) (7 phút)',
                  bold=True, first_indent=INDENT_1)
    add_multi_run_paragraph(doc, [
        ('a) Mục tiêu: ', True, False),
        (f'Kích hoạt hiểu biết nền, tạo hứng thú và kết nối học sinh vào bài học {lesson_title}.', False, True)
    ], first_indent=INDENT_1)
    add_multi_run_paragraph(doc, [
        ('b) Nội dung: ', True, False),
        (f'GV đặt câu hỏi tình huống hoặc chiếu hình ảnh/video thực tế liên quan đến {lesson_title}.', False, True)
    ], first_indent=INDENT_1)
    add_multi_run_paragraph(doc, [
        ('c) Sản phẩm: ', True, False),
        ('HS đưa ra câu trả lời ban đầu, xác định được nhiệm vụ trọng tâm cần tìm hiểu.', False, True)
    ], first_indent=INDENT_1)
    add_multi_run_paragraph(doc, [
        ('d) Tổ chức thực hiện:', True, False),
    ], first_indent=INDENT_1)

    add_3col_table(doc, [
        (f'GV chiếu hình ảnh/video/tình huống thực tế liên quan đến {lesson_title}. Nêu câu hỏi gợi mở cho cả lớp.',
         'HS tập trung quan sát màn chiếu, lắng nghe câu hỏi của GV.'),
        ('GV bao quát lớp, khuyến khích HS suy nghĩ cá nhân hoặc trao đổi nhanh với bạn bên cạnh.',
         'HS suy nghĩ câu trả lời, thảo luận nhanh nhóm đôi.'),
        ('GV gọi đại diện 2-3 HS giơ tay phát biểu ý kiến.',
         'Đại diện HS nêu câu trả lời, các bạn khác lắng nghe và nhận xét.'),
        ('GV nhận xét, chuẩn hóa ý kiến và dẫn dắt giới thiệu vào bài học mới.',
         'HS lắng nghe, mở SGK và vở ghi chuẩn bị nội dung bài mới.')
    ])

    add_paragraph(doc, '', bold=False)

    # HĐ 2. Hình thành kiến thức mới
    add_paragraph(doc, '2. Hoạt động 2. Hình thành kiến thức mới/giải quyết vấn đề (18 phút)',
                  bold=True, first_indent=INDENT_1)
    add_multi_run_paragraph(doc, [
        ('a) Mục tiêu: ', True, False),
        (f'HS nắm vững các khái niệm, quy trình và thao tác kỹ thuật của bài {lesson_title}.', False, True)
    ], first_indent=INDENT_1)
    add_multi_run_paragraph(doc, [
        ('b) Nội dung: ', True, False),
        (f'Tìm hiểu kiến thức trong SGK, quan sát GV làm mẫu thao tác và phân tích các bước thực hiện.', False, True)
    ], first_indent=INDENT_1)
    add_multi_run_paragraph(doc, [
        ('c) Sản phẩm: ', True, False),
        ('Ghi nhớ kiến thức chuẩn, hoàn thành phiếu học tập hoặc thao tác mẫu đúng quy trình.', False, True)
    ], first_indent=INDENT_1)
    add_multi_run_paragraph(doc, [
        ('d) Tổ chức thực hiện:', True, False),
    ], first_indent=INDENT_1)

    add_3col_table(doc, [
        (f'GV giao nhiệm vụ: Yêu cầu HS đọc thông tin SGK và theo dõi GV làm mẫu thao tác {lesson_title} trên máy chiếu.',
         'HS tiếp nhận nhiệm vụ, mở SGK trang tương ứng và quan sát màn chiếu.'),
        ('GV hướng dẫn từng bước chi tiết, nhấn mạnh các lưu ý và lỗi thường gặp khi thao tác.',
         'HS theo dõi kỹ các bước, ghi chép từ khóa chính và thứ tự các bước vào vở.'),
        ('GV gọi 1-2 HS đại diện lên máy GV thực hiện lại thao tác mẫu hoặc trả lời câu hỏi kiểm tra.',
         'Đại diện HS thực hiện thao tác/phát biểu; cả lớp quan sát và nhận xét.'),
        ('GV chuẩn hóa kiến thức lý thuyết và chốt quy trình thao tác chuẩn xác.',
         'HS hoàn thiện ghi nhớ kiến thức chuẩn và chuẩn bị thực hành trên máy.')
    ])

    add_paragraph(doc, '', bold=False)

    # HĐ 3. Luyện tập
    add_paragraph(doc, '3. Hoạt động 3. Luyện tập (12 phút)',
                  bold=True, first_indent=INDENT_1)
    add_multi_run_paragraph(doc, [
        ('a) Mục tiêu: ', True, False),
        (f'HS rèn luyện kỹ năng thực hành, củng cố và vận dụng kiến thức bài {lesson_title} trên máy tính.', False, True)
    ], first_indent=INDENT_1)
    add_multi_run_paragraph(doc, [
        ('b) Nội dung: ', True, False),
        (f'Thực hành cá nhân/nhóm đôi trên máy tính theo các bài tập luyện tập trong SGK/phiếu bài tập.', False, True)
    ], first_indent=INDENT_1)
    add_multi_run_paragraph(doc, [
        ('c) Sản phẩm: ', True, False),
        ('Sản phẩm thực hành trên phần mềm/tệp kết quả đạt yêu cầu kỹ thuật.', False, True)
    ], first_indent=INDENT_1)
    add_multi_run_paragraph(doc, [
        ('d) Tổ chức thực hiện:', True, False),
    ], first_indent=INDENT_1)

    add_3col_table(doc, [
        (f'GV giao bài tập thực hành luyện tập bài {lesson_title} trên máy tính cho từng HS/nhóm đôi.',
         'HS mở phần mềm/tệp thực hành trên máy tính và bắt đầu làm bài.'),
        ('GV di chuyển quan sát, hỗ trợ kịp thời các HS còn lúng túng hoặc gặp sự cố kỹ thuật.',
         'HS tự giác thực hành cá nhân, chủ động trao đổi nhóm đôi khi gặp khó khăn.'),
        ('GV kiểm tra nhanh sản phẩm tại máy của một số HS, cho HS đổi chéo kiểm tra bài của nhau.',
         'HS đối chiếu kết quả thực hành với bạn bên cạnh, báo cáo kết quả hoàn thành.'),
        ('GV nhận xét, đánh giá kết quả thực hành của cả lớp, tuyên dương HS làm tốt.',
         'HS hoàn thiện sản phẩm, lưu tệp theo đúng hướng dẫn của GV.')
    ])

    add_paragraph(doc, '', bold=False)

    # HĐ 4. Vận dụng
    add_paragraph(doc, '4. Hoạt động 4. Vận dụng (8 phút)',
                  bold=True, first_indent=INDENT_1)
    add_multi_run_paragraph(doc, [
        ('a) Mục tiêu: ', True, False),
        (f'Vận dụng kiến thức, kỹ năng đã học vào giải quyết bài toán/tình huống thực tế cuộc sống.', False, True)
    ], first_indent=INDENT_1)
    add_multi_run_paragraph(doc, [
        ('b) Nội dung: ', True, False),
        (f'GV đặt câu hỏi mở rộng, yêu cầu liên hệ thực tiễn hoặc giao bài tập về nhà sáng tạo.', False, True)
    ], first_indent=INDENT_1)
    add_multi_run_paragraph(doc, [
        ('c) Sản phẩm: ', True, False),
        ('Ý kiến đề xuất giải pháp, sản phẩm ứng dụng sáng tạo và phòng máy được tắt an toàn.', False, True)
    ], first_indent=INDENT_1)
    add_multi_run_paragraph(doc, [
        ('d) Tổ chức thực hiện:', True, False),
    ], first_indent=INDENT_1)

    add_3col_table(doc, [
        (f'GV nêu câu hỏi vận dụng: "Em có thể ứng dụng kiến thức bài {lesson_title} như thế nào trong học tập và cuộc sống?"',
         'HS tiếp nhận câu hỏi, suy nghĩ cá nhân liên hệ thực tế.'),
        ('GV dành 2-3 phút cho HS suy nghĩ, gợi ý các hướng liên hệ gần gũi.',
         'HS suy nghĩ, thảo luận nhanh hoặc ghi chép ý tưởng vào vở.'),
        ('GV mời 1-2 HS chia sẻ ý tưởng vận dụng trước lớp.',
         'HS phát biểu ý tưởng sáng tạo; các bạn khác nhận xét, bổ sung.'),
        ('GV tổng kết tiết học, giao nhiệm vụ chuẩn bị bài sau; hướng dẫn HS lưu bài, tắt máy và xếp gọn phòng học.',
         'HS ghi nhận nhiệm vụ về nhà, thực hiện Shutdown máy tính an toàn, xếp bàn ghế ngay ngắn.')
    ], is_last=True)

    # 6. Phần cuối: Rút kinh nghiệm + Bảng ký tên (NO BORDER)
    add_paragraph(doc, '', bold=False)
    add_paragraph(doc, 'RÚT KINH NGHIỆM SAU BÀI DẠY:', bold=True, first_indent=INDENT_0)
    add_paragraph(doc, '...........................................................................................................................', first_indent=INDENT_0)
    add_paragraph(doc, '...........................................................................................................................', first_indent=INDENT_0)

    add_paragraph(doc, '', bold=False)

    tbl_sign = doc.add_table(rows=3, cols=3)
    set_table_no_borders(tbl_sign)
    for i, txt in enumerate(['DUYỆT CỦA BGH', 'DUYỆT CỦA TỔ CM', 'NGƯỜI SOẠN']):
        set_cell_text(tbl_sign.rows[0].cells[i], txt, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for i in range(3):
        set_cell_text(tbl_sign.rows[1].cells[i], '(Ký, ghi rõ họ tên)', italic=True,
                      alignment=WD_ALIGN_PARAGRAPH.CENTER)
    for i in range(3):
        tbl_sign.rows[2].cells[i].text = '\n\n\n'
    set_cell_text(tbl_sign.rows[2].cells[2], '\n\n\nĐậu Đình Nguyên', bold=True,
                  alignment=WD_ALIGN_PARAGRAPH.CENTER)

    return doc


# ============================================================
# GRADE 9 LESSONS (SGK Kết nối tri thức 9)
# ============================================================
def get_grade_9_lessons():
    return [
        ("Bài 1. Thế giới kĩ thuật số", 1, "Sự nhận biết sự phát triển của máy tính và các thiết bị số thông minh làm thay đổi xã hội loài người. Phân tích tác động hai mặt của thế giới kỹ thuật số."),
        ("Bài 2. Thông tin trong giải quyết vấn đề", 2, "Giải thích được tầm quan trọng của thông tin trong việc đưa ra quyết định giải quyết vấn đề thực tế."),
        ("Bài 3. Mạng máy tính", 3, "Nêu được khái niệm mạng máy tính, thành phần mạng và lợi ích của việc kết nối mạng."),
        ("Bài 4. Internet", 4, "Trình bày được đặc điểm của Internet, các dịch vụ phổ biến trên Internet và cách truy cập thông tin."),
        ("Ôn tập Đánh giá định kỳ 1", 5, "Hệ thống hóa kiến thức chuẩn bị Đánh giá định kỳ 1."),
        ("Đánh giá định kỳ 1", 6, "Kiểm tra đánh giá kết quả học tập môn Tin học."),
        ("Bài 5. Văn hoá ứng xử trên mạng", 7, "Nhận biết các quy tắc văn hóa, đạo đức và pháp luật khi giao tiếp, chia sẻ thông tin trên không gian mạng."),
        ("Bài 6. Sử dụng phần mềm bảng tính nâng cao", 8, "Sử dụng các hàm địa chỉ tương đối/tuyệt đối, hàm điều kiện IF và xử lý dữ liệu phức tạp trong bảng tính."),
        ("Bài 7. Sắp xếp và lọc dữ liệu nâng cao", 9, "Thực hiện các thao tác sắp xếp dữ liệu theo nhiều tiêu chí và lọc dữ liệu tùy biến."),
        ("Bài 8. Trình bày dữ liệu bằng biểu đồ nâng cao", 10, "Lựa chọn và tạo các dạng biểu đồ thích hợp để trực quan hóa dữ liệu thống kê."),
        ("Ôn tập Đánh giá định kỳ 2", 11, "Hệ thống hóa kiến thức chuẩn bị Đánh giá định kỳ 2."),
        ("Đánh giá định kỳ 2", 12, "Kiểm tra đánh giá kết quả học tập môn Tin học."),
        ("Bài 9. Làm quen với phần mềm tạo trang web", 13, "Nhận biết cấu trúc trang web, tạo và chỉnh sửa nội dung trang web đơn giản."),
        ("Bài 10. Chèn hình ảnh và tạo liên kết cho trang web", 14, "Thực hiện thao tác chèn hình ảnh, tạo siêu liên kết (hyperlink) giữa các trang web."),
        ("Bài 11. Thuật toán và lập trình", 15, "Mô tả thuật toán bằng sơ đồ khối, viết chương trình giải quyết bài toán thực tế."),
        ("Bài 12. Câu lệnh rẽ nhánh và lặp trong lập trình", 16, "Sử dụng thành thạo cấu trúc rẽ nhánh và cấu trúc lặp trong ngôn ngữ lập trình."),
        ("Ôn tập Đánh giá định kỳ 3", 17, "Hệ thống hóa kiến thức chuẩn bị Đánh giá định kỳ 3."),
        ("Đánh giá định kỳ 3", 18, "Kiểm tra đánh giá kết quả học tập môn Tin học."),
        ("Bài 13. Dự án học tập: Tạo sản phẩm kĩ thuật số", 19, "Làm việc nhóm lập kế hoạch, thiết kế và hoàn thiện sản phẩm kỹ thuật số phục vụ học tập."),
        ("Bài 14. Tin học và định hướng nghề nghiệp", 20, "Tìm hiểu các ngành nghề trong lĩnh vực CNTT và ứng dụng tin học trong xã hội hiện đại."),
        ("Ôn tập Đánh giá định kỳ 4", 21, "Hệ thống hóa kiến thức chuẩn bị Đánh giá định kỳ 4."),
        ("Đánh giá định kỳ 4", 22, "Kiểm tra đánh giá kết quả học tập môn Tin học cả năm.")
    ]


# ============================================================
# MAIN EXECUTION
# ============================================================
def main():
    print("==================================================")
    print(" BẮT ĐẦU TÁI TẠO TOÀN BỘ KHBD TIN HỌC (THEO TUẦN)")
    print("==================================================")

    # 1. Dọn dẹp thư mục KHBD_Tin_học cũ (bỏ qua file đang mở)
    if os.path.exists(OUT_BASE_DIR):
        for root, dirs, files in os.walk(OUT_BASE_DIR, topdown=False):
            for f in files:
                try:
                    os.remove(os.path.join(root, f))
                except Exception:
                    pass
            for d in dirs:
                try:
                    os.rmdir(os.path.join(root, d))
                except Exception:
                    pass
        print("  [+] Đã dọn dẹp thư mục KHBD_Tin_học cũ.")

    os.makedirs(OUT_BASE_DIR, exist_ok=True)

    # 2. Parse PPCT Document
    doc_ppct = Document(PPCT_DOC_PATH)
    tables = doc_ppct.tables

    grade_configs = [
        ("Tiền_tiểu_học", "Tiền tiểu học", 3, False),
        ("Lớp_1", "Lớp 1", 4, False),
        ("Lớp_2", "Lớp 2", 5, False),
        ("Lớp_3", "Lớp 3", 6, False),
        ("Lớp_4", "Lớp 4", 7, False),
        ("Lớp_5", "Lớp 5", 8, True),
        ("Lớp_6", "Lớp 6", 9, True),
        ("Lớp_7", "Lớp 7", 10, True),
        ("Lớp_8", "Lớp 8", 11, True),
    ]

    total_created = 0

    for folder_prefix, grade_str, t_idx, is_rotation in grade_configs:
        print(f"\n---> Đang xử lý {grade_str} (Table {t_idx})...")
        t = tables[t_idx]
        lesson_count = 0

        for row_idx, row in enumerate(t.rows[1:], start=1):
            cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
            stt = cells[0]
            title = cells[1]
            so_tiet = cells[2]
            ppct_str = cells[3]
            yccd = cells[4] if len(cells) > 4 else title

            if not title or title.startswith('Chủ đề'):
                continue

            try:
                tiet_ppct = int(ppct_str)
            except ValueError:
                tiet_ppct = lesson_count + 1

            lesson_count += 1

            # Tính Tuần theo đúng logic Lịch báo giảng:
            if is_rotation:
                # Lớp 5, 6, 7, 8: Tuần 1 = Tiết 0; Tuần chẵn 2, 4, 6... = Tin học (2 tiết/tuần)
                if tiet_ppct == 0:
                    tuan_so = 1
                else:
                    tuan_so = 2 * ((tiet_ppct - 1) // 2 + 1)
            else:
                # Tiền TH & Lớp 1-4: Tuần 1 = Tiết 0; Tuần 2+ = ppct + 1
                tuan_so = tiet_ppct + 1 if tiet_ppct >= 1 else 1

            safe_title = sanitize_filename(title)
            tuan_folder = f"Tuần_{tuan_so:02d}"

            out_dir = os.path.join(OUT_BASE_DIR, folder_prefix, tuan_folder)
            os.makedirs(out_dir, exist_ok=True)

            filename = f"KHBD_Tin_hoc_{folder_prefix}_Tiet{tiet_ppct:02d}_{safe_title}.docx"
            out_file = os.path.join(out_dir, filename)

            doc = build_khbd_tin_hoc(grade_str, title, lesson_count, tiet_ppct, yccd, tuan_so)

            try:
                doc.save(out_file)
                total_created += 1
                print(f"  [+] Đã tạo: {folder_prefix} -> {tuan_folder} -> {filename}")
            except Exception as e:
                print(f"  [!] Lỗi khi lưu {out_file}: {e}")

    # Process Lớp 9
    print(f"\n---> Đang xử lý Lớp 9 (SGK Kết nối tri thức 9)...")
    grade_9_lessons = get_grade_9_lessons()
    lesson_count = 0
    for title, ppct, yccd in grade_9_lessons:
        lesson_count += 1
        tuan_so = 2 * ((ppct - 1) // 2 + 1) if ppct >= 1 else 1
        safe_title = sanitize_filename(title)
        tuan_folder = f"Tuần_{tuan_so:02d}"

        out_dir = os.path.join(OUT_BASE_DIR, "Lớp_9", tuan_folder)
        os.makedirs(out_dir, exist_ok=True)

        filename = f"KHBD_Tin_hoc_Lop_9_Tiet{ppct:02d}_{safe_title}.docx"
        out_file = os.path.join(out_dir, filename)

        doc = build_khbd_tin_hoc("Lớp 9", title, lesson_count, ppct, yccd, tuan_so)
        try:
            doc.save(out_file)
            total_created += 1
            print(f"  [+] Đã tạo: Lớp_9 -> {tuan_folder} -> {filename}")
        except Exception as e:
            print(f"  [!] Lỗi khi lưu {out_file}: {e}")

    print(f"\n==========================================")
    print(f" HOÀN THÀNH TẠO {total_created} FILE KHBD TIN HỌC CHUẨN (THEO TUẦN)!")
    print(f"==========================================")


if __name__ == '__main__':
    main()
