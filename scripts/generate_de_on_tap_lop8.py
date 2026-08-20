# -*- coding: utf-8 -*-
"""
Tạo file Đề Ôn Tập Đánh Giá Định Kỳ 1 Môn Tin Học Lớp 8 (.docx)
Tuân thủ quy chuẩn UNIGO:
  - Header / Footer giữ nguyên form chuẩn UNIGO
  - Font Times New Roman 13pt
  - Cấu trúc: Ma trận kiến thức -> Đề ôn tập (Trắc nghiệm + Tự luận) -> Đáp án & Biểu điểm
"""
import sys, io, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

TEMPLATE_DOCX = r'D:\UNIGO\Hệ thống mẫu văn bản\PL4-Khung kế hoạch bài dạy (THCS).docx'
OUT_DOCX = r'D:\UNIGO\KHBD_Tin_học\Lớp_8\Tuần_04\De_on_tap_DGDK1_Lop_8.docx'

def set_cell_borders(cell, top="single", bottom="single", left="single", right="single", sz="4", color="000000"):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = parse_xml(f'''
        <w:tcBorders {nsdecls("w")}>
            <w:top w:val="{top}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="{left}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="{bottom}" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:right w:val="{right}" w:sz="{sz}" w:space="0" w:color="{color}"/>
        </w:tcBorders>
    ''')
    tcPr.append(borders)

def set_font(run, name="Times New Roman", size_pt=13, bold=False, italic=False, color_rgb=(0,0,0)):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{name}" w:hAnsi="{name}" w:cs="{name}"/>')
    rPr.append(rFonts)

def format_para(p, text="", align=WD_ALIGN_PARAGRAPH.LEFT, bold=False, italic=False, size_pt=13, space_before=0, space_after=4, color_rgb=(0,0,0)):
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    if text:
        run = p.add_run(text)
        set_font(run, size_pt=size_pt, bold=bold, italic=italic, color_rgb=color_rgb)
        return run
    return None

def generate_de_on_tap():
    print("[+] Tạo file Đề ôn tập ĐGĐK 1 Lớp 8...")
    doc = Document(TEMPLATE_DOCX)
    
    # Keep header / sectPr, clean body
    body = doc.element.body
    for child in list(body):
        if not child.tag.endswith('sectPr'):
            body.remove(child)

    # ─── 1. TIÊU ĐỀ TRƯỜNG & ĐỀ THI ───
    tbl_hdr = doc.add_table(rows=1, cols=2)
    tbl_hdr.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell_l, cell_r = tbl_hdr.rows[0].cells
    cell_l.width = Inches(3.2)
    cell_r.width = Inches(3.8)

    p1 = cell_l.paragraphs[0]
    format_para(p1, "TRƯỜNG TIỂU HỌC & THCS UNIGO", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=12)
    p2 = cell_l.add_paragraph()
    format_para(p2, "TỔ KHOA HỌC TỰ NHIÊN - TIN HỌC", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=12)

    p3 = cell_r.paragraphs[0]
    format_para(p3, "ĐỀ CƯƠNG & BÀI TẬP ÔN TẬP ĐÁNH GIÁ ĐỊNH KỲ 1", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=12)
    p4 = cell_r.add_paragraph()
    format_para(p4, "NĂM HỌC: 2026 - 2027 | MÔN: TIN HỌC 8", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=12)
    p5 = cell_r.add_paragraph()
    format_para(p5, "(Thời gian làm bài: 45 phút)", align=WD_ALIGN_PARAGRAPH.CENTER, italic=True, size_pt=11)

    for cell in [cell_l, cell_r]:
        set_cell_borders(cell, top="none", bottom="none", left="none", right="none")

    p_div = doc.add_paragraph()
    format_para(p_div, "─────────────────────────────────────────────────────────────", align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=10, color_rgb=(100,100,100))

    # ─── 2. PHẠM VI KIẾN THỨC TRỌNG TÂM ───
    p_sec1 = doc.add_paragraph()
    format_para(p_sec1, "A. MA TRẬN PHẠM VI KIẾN THỨC ÔN TẬP", bold=True, size_pt=13, space_before=6)
    
    bullets_k = [
        "1. Lược sử công cụ tính toán và các thế hệ máy tính điện tử (Thế hệ 1 đến thế hệ 5).",
        "2. Thông tin trong môi trường số: đặc điểm, tính chất, tác động tích cực và tiêu cực.",
        "3. Tiêu chí đánh giá độ tin cậy của thông tin trên môi trường mạng Internet.",
        "4. Kỹ năng tìm kiếm, trích dẫn, khai thác và sử dụng thông tin số đúng quy định pháp luật và bản quyền."
    ]
    for b in bullets_k:
        p = doc.add_paragraph()
        format_para(p, f"  • {b}", size_pt=12)

    # ─── 3. PHẦN I: TRẮC NGHIỆM KHÁCH QUAN (6.0 ĐIỂM - 12 CÂU) ───
    p_sec2 = doc.add_paragraph()
    format_para(p_sec2, "B. ĐỀ ÔN TẬP TRẮC NGHIỆM KHÁCH QUAN (12 CÂU - 6,0 ĐIỂM)", bold=True, size_pt=13, space_before=10)

    mcq_questions = [
        ("Câu 1: Máy tính điện tử thế hệ thứ nhất (1945 - 1955) sử dụng linh kiện điện tử nào sau đây?",
         "A. Đèn điện tử chân không", "B. Bóng bán dẫn (Transistor)", "C. Mạch tích hợp (IC)", "D. Mạch tích hợp cỡ lớn (VLSI)", "A"),
        ("Câu 2: Chiếc máy tính điện tử đầu tiên trên thế giới có tên gọi là gì?",
         "A. Apple I", "B. ENIAC", "C. IBM PC", "D. Pascaline", "B"),
        ("Câu 3: Đâu KHÔNG phải là đặc điểm của thông tin trong môi trường số?",
         "A. Khối lượng khổng lồ và tăng nhanh", "B. Đa dạng các dạng thông tin", "C. Hoàn toàn chính xác 100% không cần kiểm chứng", "D. Lan truyền nhanh chóng trên toàn cầu", "C"),
        ("Câu 4: Khi tiếp nhận một thông tin trên mạng xã hội, bước đầu tiên em cần làm để đánh giá độ tin cậy là gì?",
         "A. Nhấn nút Chia sẻ (Share) ngay cho bạn bè", "B. Kiểm tra tác giả và nguồn gốc xuất bản của thông tin", "C. Bấm Like và bình luận cảm xúc", "D. Tin tưởng tuyệt đối vì có nhiều lượt xem", "B"),
        ("Câu 5: Tên miền của các cơ quan chính phủ hoặc tổ chức giáo dục chính thống thường có phần mở rộng là gì?",
         "A. .gov hoặc .edu", "B. .xyz hoặc .club", "C. .tk hoặc .top", "D. .free hoặc .info", "A"),
        ("Câu 6: Hành vi nào sau đây vi phạm bản quyền và đạo đức khi sử dụng thông tin trên mạng?",
         "A. Ghi rõ nguồn tác giả khi trích dẫn tài liệu", "B. Tải tranh vẽ của người khác rồi tự nhận là tác phẩm của mình", "C. Sử dụng tài liệu học tập có ghi rõ bản quyền Creative Commons", "D. Chia sẻ đường link bài báo gốc cho mọi người cùng đọc", "B"),
        ("Câu 7: Máy tính thế hệ thứ tư (từ năm 1971 đến nay) gắn liền với sự phát triển của linh kiện nào?",
         "A. Vi xử lý (Microprocessor) tích hợp mật độ cao", "B. Đèn chân không", "C. Rơ le điện từ", "D. Ống tia âm cực", "A"),
        ("Câu 8: Thông tin nào sau đây có nguy cơ là tin giả (Fake news)?",
         "A. Thông báo lịch thi từ Cổng thông tin Bộ Giáo dục và Đào tạo", "B. Bài viết giật gân, không rõ tác giả, kêu gọi chia sẻ khẩn cấp trên mạng", "C. Bản tin thời sự phát sóng trực tiếp trên Đài Truyền hình Việt Nam", "D. Dự báo thời tiết từ Trung tâm Dự báo Khí tượng Thủy văn Quốc gia", "B"),
        ("Câu 9: Khi tìm kiếm thông tin bằng Google, để tìm chính xác một cụm từ ta nên đặt cụm từ đó trong ký hiệu nào?",
         "A. Cặp dấu ngoặc đơn ( )", "B. Cặp dấu ngoặc kép \" \"", "C. Cặp dấu ngoặc vuông [ ]", "D. Cặp dấu gạch chéo / /", "B"),
        ("Câu 10: Tác động tiêu cực của việc lạm dụng thông tin số và mạng xã hội là gì?",
         "A. Tăng khả năng tự học trực tuyến", "B. Nguy cơ nghiện mạng, suy giảm thị lực và tiếp xúc thông tin xấu độc", "C. Mở rộng giao lưu kết nối bạn bè quốc tế", "D. Tiếp cận nhanh chóng kho tàng tri thức nhân loại", "B"),
        ("Câu 11: Khi sử dụng bài viết, hình ảnh từ Internet vào bài thuyết trình của mình, em cần làm gì để tôn trọng tác giả?",
         "A. Đổi tên tác giả thành tên của mình", "B. Ghi chú nguồn gốc tác giả và đường liên kết tham khảo", "C. Xóa chữ ký tác giả trên hình ảnh", "D. Không cần làm gì vì thông tin trên mạng là miễn phí", "B"),
        ("Câu 12: Bộ não xử lý thông tin của máy tính điện tử hiện đại là thiết bị nào?",
         "A. Bàn phím", "B. Bộ vi xử lý trung tâm (CPU)", "C. Màn hình", "D. Chuột máy tính", "B")
    ]

    for q_text, opt_a, opt_b, opt_c, opt_d, _ in mcq_questions:
        p_q = doc.add_paragraph()
        format_para(p_q, q_text, bold=True, size_pt=12, space_before=4)
        p_opts = doc.add_paragraph()
        format_para(p_opts, f"    {opt_a}         {opt_b}\n    {opt_c}         {opt_d}", size_pt=12)

    # ─── 4. PHẦN II: TỰ LUẬN & BÀI TẬP TÌNH HUỐNG (4.0 ĐIỂM - 2 CÂU) ───
    p_sec3 = doc.add_paragraph()
    format_para(p_sec3, "C. PHẦN TỰ LUẬN VÀ BÀI TẬP TÌNH HUỐNG (4,0 ĐIỂM)", bold=True, size_pt=13, space_before=10)

    p_t1 = doc.add_paragraph()
    format_para(p_t1, "Câu 1 (2,0 điểm):", bold=True, size_pt=12)
    format_para(doc.add_paragraph(), "  Em hãy nêu 4 tiêu chí cơ bản để đánh giá độ tin cậy của một nguồn thông tin trên Internet. Lấy ví dụ minh họa cho một tiêu chí cụ thể.", size_pt=12)

    p_t2 = doc.add_paragraph()
    format_para(p_t2, "Câu 2 (2,0 điểm) - Tình huống thực tế:", bold=True, size_pt=12)
    format_para(doc.add_paragraph(), "  Bạn Nam đang làm bài báo cáo Tin học về chủ đề 'Bảo vệ môi trường số'. Nam tìm thấy một bài viết rất hay trên mạng nhưng không có tên tác giả và trên một trang web lạ. Theo em:\n  a) Nam có nên sao chép nguyên văn nội dung đó vào bài của mình không? Vì sao?\n  b) Em hãy đưa ra lời khuyên giúp Nam tìm kiếm và kiểm chứng nguồn tài liệu chính thống, tin cậy.", size_pt=12)

    # ─── 5. ĐÁP ÁN & HƯỚNG DẪN CHẤM ───
    p_sec4 = doc.add_paragraph()
    format_para(p_sec4, "D. ĐÁP ÁN VÀ BIỂU ĐIỂM CHI TIẾT", bold=True, size_pt=13, space_before=10)

    # Bảng đáp án trắc nghiệm
    tbl_ans = doc.add_table(rows=2, cols=12)
    tbl_ans.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = [f"C{i+1}" for i in range(12)]
    answers = [ans for _, _, _, _, _, ans in mcq_questions]

    for col_idx, h in enumerate(headers):
        cell = tbl_ans.rows[0].cells[col_idx]
        cell.width = Inches(0.55)
        p = cell.paragraphs[0]
        format_para(p, h, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=11)
        set_cell_borders(cell)

    for col_idx, a in enumerate(answers):
        cell = tbl_ans.rows[1].cells[col_idx]
        cell.width = Inches(0.55)
        p = cell.paragraphs[0]
        format_para(p, a, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=11, color_rgb=(180,0,0))
        set_cell_borders(cell)

    p_note = doc.add_paragraph()
    format_para(p_note, "*(Mỗi câu trắc nghiệm đúng được 0,5 điểm x 12 = 6,0 điểm)*", italic=True, size_pt=11, space_before=4)

    # Hướng dẫn chấm tự luận
    p_tl_hd = doc.add_paragraph()
    format_para(p_tl_hd, "Hướng dẫn chấm Tự luận:", bold=True, size_pt=12, space_before=6)
    tl_guides = [
        "Câu 1 (2,0đ): Nêu đúng 4 tiêu chí (Nguồn gốc/Tác giả; Tính cập nhật; Mục đích bài viết; Kiểm chứng chéo) (1,5đ). Lấy ví dụ đúng (0,5đ).",
        "Câu 2 (2,0đ): a) Khẳng định không nên sao chép nguyên văn vì vi phạm bản quyền và rủi ro tin sai lệch (1,0đ). b) Khuyên kiểm tra nguồn uy tín (.gov, .edu, báo chí) và trích dẫn hợp lệ (1,0đ)."
    ]
    for g in tl_guides:
        p = doc.add_paragraph()
        format_para(p, f"  • {g}", size_pt=11)

    # Ký duyệt
    tbl_sign = doc.add_table(rows=1, cols=3)
    tbl_sign.alignment = WD_TABLE_ALIGNMENT.CENTER
    c1, c2, c3 = tbl_sign.rows[0].cells
    c1.width = Inches(2.3); c2.width = Inches(2.3); c3.width = Inches(2.3)
    for c in [c1, c2, c3]:
        set_cell_borders(c, "none", "none", "none", "none")

    format_para(c1.paragraphs[0], "DUYỆT BGH", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=11)
    format_para(c2.paragraphs[0], "DUYỆT TỔ CHUYÊN MÔN", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=11)
    format_para(c3.paragraphs[0], "GIÁO VIÊN SOẠN ĐỀ", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size_pt=11)

    os.makedirs(os.path.dirname(OUT_DOCX), exist_ok=True)
    doc.save(OUT_DOCX)
    print(f"    [OK] Đã lưu file đề ôn tập: {OUT_DOCX}")

if __name__ == '__main__':
    generate_de_on_tap()
