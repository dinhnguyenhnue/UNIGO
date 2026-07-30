import os
import sys
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding='utf-8')

TMPL_FILE = r'D:\UNIGO\Hệ thống mẫu văn bản\Lịch báo giảng.docx'
OUT_DIR = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Lịch báo giảng'

OUT_LBG = os.path.join(OUT_DIR, 'Lịch báo giảng.docx')
OUT_LBG_T1 = os.path.join(OUT_DIR, 'Lịch báo giảng - Tuần 01.docx')

os.makedirs(OUT_DIR, exist_ok=True)

# Timetable data for Teacher Đậu Đình Nguyên (Tin học & Robotics) from "TKB toàn trường CHECK - lần 2.1.xlsx"
SCHEDULE = [
    # Thứ Hai (03/08)
    (0, 4, 'sang', 'Tin học', '1A1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (0, 1, 'chieu', 'Tin học', '7A1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    
    # Thứ Ba (04/08)
    (1, 3, 'sang', 'Tin học', '5C1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (1, 4, 'sang', 'Robotics', '5C1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
    (1, 1, 'chieu', 'Tin học', '2A1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (1, 2, 'chieu', 'Tin học', '1C1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (1, 3, 'chieu', 'Tin học', '7A1', '2', 'Bài 1. Thiết bị vào - ra', 'Phòng Tin học'),
    (1, 4, 'chieu', 'Robotics', '7A1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
    
    # Thứ Tư (05/08)
    (2, 2, 'sang', 'Robotics', '3A1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
    (2, 4, 'sang', 'Tin học', '3C1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (2, 3, 'chieu', 'Tin học', '4C1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (2, 4, 'chieu', 'Robotics', '2A1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
    
    # Thứ Năm (06/08)
    (3, 1, 'sang', 'Tin học', 'TT3', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (3, 2, 'sang', 'Robotics', '1A1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
    (3, 4, 'sang', 'Robotics', '2C1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
    (3, 1, 'chieu', 'Tin học', '3A1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (3, 2, 'chieu', 'Tin học', 'TTH 1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (3, 3, 'chieu', 'Tin học', 'TTH2', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (3, 4, 'chieu', 'Robotics', '1C1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
    
    # Thứ Sáu (07/08)
    (4, 2, 'sang', 'Robotics', '3C1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
    (4, 3, 'sang', 'Tin học', '2C1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (4, 4, 'sang', 'Tin học', '6A1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (4, 5, 'sang', 'Robotics', '6A1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
    (4, 1, 'chieu', 'Tin học', '8A1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (4, 2, 'chieu', 'Robotics', '8A1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
    (4, 3, 'chieu', 'Robotics', '4C1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
]

DAYS_DATES = [
    ("Hai (03/08)", 0),
    ("Ba (04/08)", 1),
    ("Tư (05/08)", 2),
    ("Năm (06/08)", 3),
    ("Sáu (07/08)", 4)
]

def set_table_borders(table):
    """Sets full black grid borders."""
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        for child in list(tblPr[0]):
            if child.tag.endswith('tblBorders'):
                tblPr[0].remove(child)
        borders_xml = parse_xml(
            r'<w:tblBorders %s>'
            r'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        tblPr[0].append(borders_xml)

def set_font_all(doc):
    """Sets Times New Roman font for paragraphs and tables without destroying layout."""
    for p in doc.paragraphs:
        p.style.font.name = 'Times New Roman'
        for r in p.runs:
            r.font.name = 'Times New Roman'
            
    for i, t in enumerate(doc.tables):
        if i != 0: # Do not overwrite Table 0 borders
            set_table_borders(t)
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.style.font.name = 'Times New Roman'
                    for r in p.runs:
                        r.font.name = 'Times New Roman'

def update_cover_page(table_0):
    """Updates Teacher Name and Class List on the cover page (Table 0) while keeping all layout intact."""
    cell0 = table_0.rows[0].cells[0]
    for p in cell0.paragraphs:
        txt = p.text.strip()
        if "Họ và tên giáo viên" in txt:
            p.text = "Họ và tên giáo viên  : Đậu Đình Nguyên"
            for r in p.runs:
                r.font.name = 'Times New Roman'
        elif txt.startswith("Lớp:"):
            p.text = "Lớp: Tiền tiểu học, Lớp 1, 2, 3, 4, 5, 6, 7, 8"
            for r in p.runs:
                r.font.name = 'Times New Roman'

def populate_lbg_table_clean(table, session_type):
    """Rebuilds LBG table with ONLY active teaching rows (deleting empty period rows)."""
    for row in list(table.rows)[1:]:
        table._element.remove(row._element)
        
    hdr_row = table.rows[0]
    for i, c in enumerate(hdr_row.cells):
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)
            
    session_slots = [item for item in SCHEDULE if item[2] == session_type]
    session_slots.sort(key=lambda x: (x[0], x[1]))
    
    for day_idx in range(5):
        day_name, _ = DAYS_DATES[day_idx]
        day_items = [item for item in session_slots if item[0] == day_idx]
        
        if not day_items:
            continue
            
        for idx, item in enumerate(day_items):
            _, period, _, mon, lop, ppct, ten_bai, do_dung = item
            row = table.add_row()
            values = [day_name, str(period), str(ppct), mon, lop, ten_bai, do_dung]
            aligns = [
                WD_ALIGN_PARAGRAPH.CENTER, # Thứ/ngày
                WD_ALIGN_PARAGRAPH.CENTER, # Tiết TKB
                WD_ALIGN_PARAGRAPH.CENTER, # Tiết PPCT
                WD_ALIGN_PARAGRAPH.CENTER, # Môn
                WD_ALIGN_PARAGRAPH.CENTER, # Lớp
                WD_ALIGN_PARAGRAPH.LEFT,   # Tên bài
                WD_ALIGN_PARAGRAPH.LEFT    # Đồ dùng
            ]
            
            for c_idx in range(7):
                cell = row.cells[c_idx]
                cell.text = values[c_idx]
                p = cell.paragraphs[0]
                p.alignment = aligns[c_idx]
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(13)
                    
        if len(day_items) > 1:
            start_row_idx = len(table.rows) - len(day_items)
            end_row_idx = len(table.rows) - 1
            first_cell = table.rows[start_row_idx].cells[0]
            for r_i in range(start_row_idx + 1, end_row_idx + 1):
                first_cell.merge(table.rows[r_i].cells[0])
            first_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
    set_table_borders(table)

def process_lbg():
    print(f"Reading template {TMPL_FILE}...")
    doc = docx.Document(TMPL_FILE)
    
    if len(doc.tables) > 0:
        print("Updating cover page (Table 0)...")
        update_cover_page(doc.tables[0])
        
    for p in doc.paragraphs:
        txt = p.text.strip()
        if "TUẦN 01" in txt or "TUẦN 1" in txt or "từ ngày" in txt:
            p.text = "TUẦN 01 (từ ngày 03/08/2026 đến ngày 07/08/2026)"
        elif "Ngày ........" in txt or "Ngày ....." in txt or "tháng........" in txt:
            p.text = "Ngày 01 tháng 08 năm 2026"
        elif "Buổi" in txt and "Tuần" in txt:
            if "Sáng" in txt or "SÁNG" in txt:
                p.text = "Buổi: Sáng    Tuần: 01 (Từ ngày 03/08/2026 đến ngày 07/08/2026)"
            elif "Chiều" in txt or "CHIỀU" in txt:
                p.text = "Buổi: Chiều    Tuần: 01 (Từ ngày 03/08/2026 đến ngày 07/08/2026)"

    print("Populating Table 1 (Sáng) - active rows only...")
    populate_lbg_table_clean(doc.tables[1], 'sang')
    
    print("Populating Table 3 (Chiều) - active rows only...")
    populate_lbg_table_clean(doc.tables[3], 'chieu')
    
    set_font_all(doc)
    
    for path in [OUT_LBG, OUT_LBG_T1]:
        try:
            doc.save(path)
            print(f"Successfully saved {path}")
        except PermissionError:
            print(f"PERMISSION ERROR: Could not save to {path}.")

if __name__ == '__main__':
    process_lbg()
