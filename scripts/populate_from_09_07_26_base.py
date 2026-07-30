import docx
import os
import sys
import copy
import math
import re
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding='utf-8')

BASE_09_07_26 = r'D:\UNIGO\Hệ thống mẫu văn bản\09.07.26. Kế hoạch tổ chuyên môn (THCS).docx'
if not os.path.exists(BASE_09_07_26):
    BASE_09_07_26 = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch tổ chuyên môn\09.07.26. Kế hoạch tổ chuyên môn (THCS).docx'

DIR_TOCM = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch tổ chuyên môn'
DIR_MAU = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm'

TARGET_1 = os.path.join(DIR_TOCM, '30.07.26. Kế hoạch tổ chuyên môn (THCS).docx')
TARGET_2 = os.path.join(DIR_TOCM, '09.07.26. Kế hoạch tổ chuyên môn (THCS).docx')
TARGET_3 = os.path.join(DIR_TOCM, 'Kế hoạch tổ chuyên môn (THCS).docx')

def clean_xml_element(elem):
    for bm in elem.xpath('.//w:bookmarkStart | .//w:bookmarkEnd'):
        bm.getparent().remove(bm)
    for node in elem.xpath('.//*[@r:id]'):
        for attr in list(node.attrib):
            if 'id' in attr.lower() and attr != 'id':
                del node.attrib[attr]
    return elem

def set_table_borders(table):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        for child in list(tblPr[0]):
            if child.tag.endswith('tblBorders'):
                tblPr[0].remove(child)
        borders_xml = parse_xml(
            r'<w:tblBorders %s>'
            r'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        tblPr[0].append(borders_xml)

def format_bullet_text(text):
    if not text:
        return ""
    text = text.replace("; -", "\n-").replace(". -", ".\n-")
    text = re.sub(r'(\S)\s*[-–—]\s+', r'\1\n- ', text)
    lines = [line.strip() for line in text.split('\n') if line.strip()]
    formatted_lines = []
    for line in lines:
        if line.startswith("-") or line.startswith("•"):
            if not line.startswith("- "):
                line = "- " + line.lstrip("-•").strip()
        formatted_lines.append(line)
    return "\n".join(formatted_lines)

def add_p_after(doc, anchor_elem, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=13):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(font_size)
    r.bold = bold
    r.italic = italic
    
    p_elem = clean_xml_element(p._element)
    p_elem.getparent().remove(p_elem)
    anchor_elem.addnext(p_elem)
    return p_elem

def add_table_7col_after(doc, anchor_elem, cleaned_rows, periods_per_week=1):
    table = doc.add_table(rows=0, cols=7)
    set_table_borders(table)
    
    col_widths = [Cm(0.8), Cm(3.8), Cm(1.0), Cm(1.2), Cm(1.2), Cm(3.5), Cm(6.5)]
    
    # Header row
    hdr_row = table.add_row()
    headers = ['TT', 'Bài/chủ đề', 'Tổng số tiết', 'Tuần', 'Tiết theo PPCT', 'Nội dung', 'Mục tiêu bài học']
    for idx, text in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.width = col_widths[idx]
        cell.text = text
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)

    period_seq = 0
    lesson_tt = 1
    
    for r_type, data in cleaned_rows:
        if r_type == 'topic':
            row = table.add_row()
            cell_top = row.cells[0]
            for c_i in range(1, 7):
                cell_top.merge(row.cells[c_i])
            cell_top.text = str(data)
            p = cell_top.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(13)
        elif r_type == 'lesson':
            stt_orig, bai, so_tiet_str, _, tiet_ppct_orig, noi_dung, yccd = data
            
            try:
                n_tiet = int(so_tiet_str)
            except ValueError:
                n_tiet = 1
                
            period_seq += n_tiet
            week_num = min(35, math.ceil(period_seq / max(1, periods_per_week)))
            
            formatted_noi_dung = format_bullet_text(noi_dung)
            formatted_yccd = format_bullet_text(yccd)
            
            row = table.add_row()
            row_data = [
                str(lesson_tt),
                bai,
                str(n_tiet),
                f"{week_num}",
                str(period_seq),
                formatted_noi_dung,
                formatted_yccd
            ]
            lesson_tt += 1
            
            for idx in range(7):
                cell = row.cells[idx]
                cell.width = col_widths[idx]
                cell.text = str(row_data[idx])
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in [0, 2, 3, 4] else WD_ALIGN_PARAGRAPH.JUSTIFY
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(13)

    for row in table.rows:
        for c_i in range(min(7, len(row.cells))):
            row.cells[c_i].width = col_widths[c_i]

    t_elem = clean_xml_element(table._element)
    t_elem.getparent().remove(t_elem)
    anchor_elem.addnext(t_elem)
    return t_elem

def add_table_5col_after(doc, anchor_elem, grade_name):
    table = doc.add_table(rows=0, cols=5)
    set_table_borders(table)
    
    col_widths = [Cm(0.8), Cm(1.5), Cm(5.0), Cm(6.5), Cm(4.2)]
    
    hdr_row = table.add_row()
    headers = ['TT', 'Lớp', 'Bài kiểm tra', 'Nội dung', 'Hình thức']
    for idx, text in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.width = col_widths[idx]
        cell.text = text
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)
            
    assessments = [
        ('1', grade_name, 'Đánh giá định kỳ 1 (Giữa HK1)', 'Nội dung Học kỳ 1 (Tuần 1 - 9)', 'Viết trên giấy / Thực hành'),
        ('2', grade_name, 'Đánh giá định kỳ 2 (Cuối HK1)', 'Nội dung Học kỳ 1 (Tuần 1 - 18)', 'Viết trên giấy / Thực hành'),
        ('3', grade_name, 'Đánh giá định kỳ 3 (Giữa HK2)', 'Nội dung Học kỳ 2 (Tuần 19 - 27)', 'Viết trên giấy / Thực hành'),
        ('4', grade_name, 'Đánh giá định kỳ 4 (Cuối HK2)', 'Nội dung Học kỳ 2 & Tổng hợp cả năm', 'Viết trên giấy / Thực hành'),
    ]
    
    for item in assessments:
        row = table.add_row()
        for idx in range(5):
            cell = row.cells[idx]
            cell.width = col_widths[idx]
            cell.text = item[idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in [0, 1, 4] else WD_ALIGN_PARAGRAPH.JUSTIFY
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(13)
                
    for row in table.rows:
        for c_i in range(min(5, len(row.cells))):
            row.cells[c_i].width = col_widths[c_i]

    t_elem = clean_xml_element(table._element)
    t_elem.getparent().remove(t_elem)
    anchor_elem.addnext(t_elem)
    return t_elem

def parse_lessons_from_source(doc_src, table_idx):
    cleaned_rows = []
    if not doc_src or table_idx >= len(doc_src.tables):
        return cleaned_rows

    t = doc_src.tables[table_idx]
    current_topic = ""

    for r_idx, row in enumerate(t.rows[1:]):
        vals = [c.text.strip().replace('\n', ' ') for c in row.cells]
        row_txt = " ".join(vals).upper()

        if any(k in row_txt for k in ['HỌC KỲ I:', 'HỌC KỲ II:', 'HỌC KÌ I:', 'HỌC KÌ II:', 'HỌC KỲ 1', 'HỌC KỲ 2', 'HỌC KÌ 1', 'HỌC KÌ 2', 'CẢ NĂM:']):
            continue
        if 'THIẾT BỊ' in row_txt or 'PHÒNG HỌC' in row_txt or 'CÒI' in row_txt or 'BẢNG TƯƠNG TÁC' in row_txt:
            continue

        if len(vals) == 4 and ('CHUYÊN ĐỀ' in vals[0].upper() or 'CÁC SỐ' in vals[0].upper()):
            topic = vals[0]
            noi_dung = vals[2]
            yccd = vals[3]
            if topic != current_topic:
                current_topic = topic
                cleaned_rows.append(('topic', current_topic))
            cleaned_rows.append(('lesson', [str(r_idx + 1), noi_dung, '1', '', '', noi_dung, yccd]))
            continue

        if vals[0].startswith("Bài 1:") or vals[0].startswith("Bài 2:") or "Bài 1:" in vals[0] or "Chủ đề" in vals[0] or "Chương" in vals[0] or (vals[0] and not vals[0].isdigit() and len(vals[0]) > 5):
            if len(set(vals)) == 1 or not vals[2].isdigit():
                cleaned_rows.append(('topic', vals[0]))
                continue

        if vals[0] and vals[1]:
            stt = vals[0]
            bai = vals[1]
            so_tiet = vals[2] if len(vals) > 2 else '1'
            tuan = ''
            tiet_ppct = vals[2] if len(vals) > 2 else ''
            noi_dung = bai
            yccd = vals[3] if len(vals) > 3 else ''
            if len(vals) >= 5:
                tiet_ppct = vals[2]
                yccd = vals[4] if len(vals) > 4 else vals[3]

            cleaned_rows.append(('lesson', [stt, bai, so_tiet, tuan, tiet_ppct, noi_dung, yccd]))

    return cleaned_rows

def add_subject_section_inplace(doc, anchor_elem, s_idx, s_name, has_data, source_files_map, periods_per_week=1):
    curr = anchor_elem
    curr = add_p_after(doc, curr, f"{s_idx}. Môn: {s_name}", bold=True, font_size=14)

    for g_num in range(6, 9):
        g_name = f"{s_name} {g_num}"
        curr = add_p_after(doc, curr, f"{s_idx}.{g_num-5}. Kế hoạch dạy học môn {g_name}", bold=True, font_size=13)
        curr = add_p_after(doc, curr, f"{s_idx}.{g_num-5}.1. Kế hoạch dạy học chính khoá", bold=True)
        curr = add_p_after(doc, curr, "Cả năm: 35 Tiết")
        curr = add_p_after(doc, curr, "Học kì 1: 18 Tiết")
        curr = add_p_after(doc, curr, "Học kì 2: 17 Tiết")
        curr = add_p_after(doc, curr, "Số điểm kiểm tra thường xuyên: 4/kỳ")
        curr = add_p_after(doc, curr, "Số điểm kiểm tra định kỳ:  2/kỳ")
        curr = add_p_after(doc, curr, "")

        fpath, t_idx = source_files_map.get(g_num, (None, None))
        if has_data and fpath and os.path.exists(fpath) and t_idx is not None:
            doc_src = docx.Document(fpath)
            cleaned_rows = parse_lessons_from_source(doc_src, t_idx)
            curr = add_table_7col_after(doc, curr, cleaned_rows, periods_per_week=periods_per_week)
        else:
            curr = add_table_7col_after(doc, curr, [], periods_per_week=periods_per_week)
            curr = add_p_after(doc, curr, f"(Ghi chú: Mục môn {s_name} - Đã ghi nhận form mẫu và sẽ bổ sung chi tiết Khung PPCT sau)", italic=True)

        curr = add_p_after(doc, curr, "")
        curr = add_p_after(doc, curr, f"{s_idx}.{g_num-5}.2. Kế hoạch dạy học tăng cường", bold=True)
        curr = add_p_after(doc, curr, ".....")
        curr = add_p_after(doc, curr, "")
        curr = add_p_after(doc, curr, f"{s_idx}.{g_num-5}.3. Kế hoạch kiểm tra đánh giá", bold=True)
        curr = add_table_5col_after(doc, curr, f"Lớp {g_num}A1")
        curr = add_p_after(doc, curr, "")

    return curr

def main():
    print(f"Loading base template file from {BASE_09_07_26}...")
    doc = docx.Document(BASE_09_07_26)

    p_vi = None
    p_vii = None

    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt.startswith("VI. KẾ HOẠCH GIẢNG DẠY"):
            p_vi = p
        elif txt.startswith("VII. KẾ HOẠCH CỤ THỂ") or txt.startswith("VIII. Nhiệm vụ"):
            if not p_vii:
                p_vii = p

    print(f"p_vi: {p_vi.text if p_vi else None}")
    print(f"p_vii: {p_vii.text if p_vii else None}")

    # Remove old elements between p_vi and p_vii
    print("Deleting old section VI content in-place...")
    curr_elem = p_vi._element.getnext()
    while curr_elem is not None and curr_elem != p_vii._element:
        nxt = curr_elem.getnext()
        curr_elem.getparent().remove(curr_elem)
        curr_elem = nxt

    anchor = p_vi._element

    # Source files paths in D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch tổ chuyên môn
    f_van6_lsdl = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch dạy học môn Văn 6 _ Lịch sử và Địa lý 6,7,8.docx')
    f_van78_gddp = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch dạy học môn Văn lớp 7,8_Giáo dục địa phương lớp 6,7,8 (THCS)-PPCT.docx')
    f_toan = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch môn Toán THCS (Kế hoạch cá nhân).docx')
    f_anh67 = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch dạy học môn Tiếng Anh 6,7 (THCS)-PPCT.docx')
    f_anh8 = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch dạy học môn Tiếng anh 8_ Cô Quỳnh.docx')
    f_khtn = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch dạy học môn KHTN6,7,8.26-27.docx')
    f_amnhac = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch dạy học môn ÂM NHẠC THCS.docx')
    f_mythuat = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch dạy học môn Mĩ Thuật (THCS)-PPCT.docx')
    f_gdtc = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch dạy học môn GDTC THCS 26-27.docx')
    f_hdtn = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch môn HĐTN THCS. Bùi Thị Phương Anh(Kế hoạch cá nhân).docx')

    f_tin = os.path.join(DIR_MAU, 'Kế hoạch dạy học môn Tin học (THCS) - 2026 - 2027.docx')
    f_robotics = os.path.join(DIR_MAU, 'Kế hoạch dạy học môn Robotics (THCS) - 2026 - 2027.docx')

    print("Populating 14 subjects (Setting HĐTN name to 'Hoạt động trải nghiệm' with PPCT tables 6, 7, 8)...")

    # 1. Ngữ văn (Văn 6 from f_van6_lsdl T7; Văn 7 from f_van78_gddp T5; Văn 8 from f_van78_gddp T8)
    map_van = {6: (f_van6_lsdl, 7), 7: (f_van78_gddp, 5), 8: (f_van78_gddp, 8)}
    anchor = add_subject_section_inplace(doc, anchor, 1, "Ngữ văn", True, map_van, periods_per_week=4)

    # 2. Toán (Toán 6: T5, Toán 7: T6, Toán 8: T8 in f_toan)
    map_toan = {6: (f_toan, 5), 7: (f_toan, 6), 8: (f_toan, 8)}
    anchor = add_subject_section_inplace(doc, anchor, 2, "Toán", True, map_toan, periods_per_week=4)

    # 3. Tiếng Anh (Anh 6: T3, Anh 7: T4 in f_anh67; Anh 8: T7 in f_anh8)
    map_anh = {6: (f_anh67, 3), 7: (f_anh67, 4), 8: (f_anh8, 7)}
    anchor = add_subject_section_inplace(doc, anchor, 3, "Tiếng Anh", True, map_anh, periods_per_week=3)

    # 4. Tiếng Trung (Placeholder)
    anchor = add_subject_section_inplace(doc, anchor, 4, "Tiếng Trung", False, {}, periods_per_week=2)

    # 5. Khoa học tự nhiên (KHTN 6: T8, KHTN 7: T9, KHTN 8: T11 in f_khtn)
    map_khtn = {6: (f_khtn, 8), 7: (f_khtn, 9), 8: (f_khtn, 11)}
    anchor = add_subject_section_inplace(doc, anchor, 5, "Khoa học tự nhiên", True, map_khtn, periods_per_week=4)

    # 6. Lịch sử và Địa lý (LS-ĐL 6: T11, LS-ĐL 7: T13, LS-ĐL 8: T15 in f_van6_lsdl)
    map_lsdl = {6: (f_van6_lsdl, 11), 7: (f_van6_lsdl, 13), 8: (f_van6_lsdl, 15)}
    anchor = add_subject_section_inplace(doc, anchor, 6, "Lịch sử và Địa lý", True, map_lsdl, periods_per_week=3)

    # 7. Tin học (Tin 6: T3, Tin 7: T4, Tin 8: T5 in f_tin)
    map_tin = {6: (f_tin, 3), 7: (f_tin, 4), 8: (f_tin, 5)}
    anchor = add_subject_section_inplace(doc, anchor, 7, "Tin học", True, map_tin, periods_per_week=1)

    # 8. Robotics (Rob 6: T3, Rob 7: T4, Rob 8: T5 in f_robotics)
    map_rob = {6: (f_robotics, 3), 7: (f_robotics, 4), 8: (f_robotics, 5)}
    anchor = add_subject_section_inplace(doc, anchor, 8, "Robotics", True, map_rob, periods_per_week=1)

    # 9. Giáo dục công dân (GDCD 6: T17, GDCD 7: T18, GDCD 8: T19 in f_van6_lsdl)
    map_gdcd = {6: (f_van6_lsdl, 17), 7: (f_van6_lsdl, 18), 8: (f_van6_lsdl, 19)}
    anchor = add_subject_section_inplace(doc, anchor, 9, "Giáo dục công dân", True, map_gdcd, periods_per_week=1)

    # 10. Giáo dục địa phương (GDĐP 6: T13, GDĐP 7: T14, GDĐP 8: T15 in f_van78_gddp)
    map_gddp = {6: (f_van78_gddp, 13), 7: (f_van78_gddp, 14), 8: (f_van78_gddp, 15)}
    anchor = add_subject_section_inplace(doc, anchor, 10, "Giáo dục địa phương", True, map_gddp, periods_per_week=1)

    # 11. Hoạt động trải nghiệm (HĐTN 6: T6, HĐTN 7: T7, HĐTN 8: T8 in f_hdtn)
    map_hdtn = {6: (f_hdtn, 6), 7: (f_hdtn, 7), 8: (f_hdtn, 8)}
    anchor = add_subject_section_inplace(doc, anchor, 11, "Hoạt động trải nghiệm", True, map_hdtn, periods_per_week=3)

    # 12. Âm nhạc (Âm nhạc 6: T3, Âm nhạc 7: T4, Âm nhạc 8: T5 in f_amnhac)
    map_amnhac = {6: (f_amnhac, 3), 7: (f_amnhac, 4), 8: (f_amnhac, 5)}
    anchor = add_subject_section_inplace(doc, anchor, 12, "Âm nhạc", True, map_amnhac, periods_per_week=1)

    # 13. Mỹ thuật (Mỹ thuật 6: T3, Mỹ thuật 7: T4, Mỹ thuật 8: T5 in f_mythuat)
    map_myt = {6: (f_mythuat, 3), 7: (f_mythuat, 4), 8: (f_mythuat, 5)}
    anchor = add_subject_section_inplace(doc, anchor, 13, "Mỹ thuật", True, map_myt, periods_per_week=1)

    # 14. Giáo dục thể chất (GDTC 6: T15, GDTC 7: T16, GDTC 8: T17 in f_gdtc)
    map_gdtc = {6: (f_gdtc, 15), 7: (f_gdtc, 16), 8: (f_gdtc, 17)}
    anchor = add_subject_section_inplace(doc, anchor, 14, "Giáo dục thể chất", True, map_gdtc, periods_per_week=2)

    # Format font Times New Roman 13pt
    print("Applying Times New Roman 13pt and clean borders...")
    for p in doc.paragraphs:
        p.style.font.name = 'Times New Roman'
        for r in p.runs:
            r.font.name = 'Times New Roman'

    for t in doc.tables:
        set_table_borders(t)
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.style.font.name = 'Times New Roman'
                    for r in p.runs:
                        r.font.name = 'Times New Roman'

    for target in [TARGET_1, TARGET_2, TARGET_3]:
        try:
            doc.save(target)
            print(f"Successfully saved document with HĐTN to {target}")
        except PermissionError:
            print(f"Permission error saving to {target}")

if __name__ == '__main__':
    main()
