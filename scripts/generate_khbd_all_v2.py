"""
generate_khbd_all_v2.py
========================
Tự động tạo toàn bộ KHBD Tin học & Robotics (TTH → Lớp 8) theo đúng luật:
- Bảng 2 cột: HOẠT ĐỘNG CỦA GV – HS | KẾT QUẢ CẦN ĐẠT
- Table info 3×2 (THCS) — NO BORDER
- Kiến thức dùng Danh từ trực tiếp (CẤM "Sự hiểu biết", "Khả năng")
- NLS Bậc đúng theo khối (L1-3=Bậc1, L4-5=Bậc2, L6-7=Bậc3, L8-9=Bậc4)
- Phân luồng TH / THCS riêng biệt
- Ngày soạn/dạy đúng mapping LBG
"""

import os
import re
import sys
import json
import copy
import shutil
from datetime import date, timedelta
from docx import Document
from docx.shared import Pt, Cm, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

sys.stdout.reconfigure(encoding='utf-8')

# ============================================================
# CONFIG
# ============================================================
TPL_TH = r"D:\UNIGO\Hệ thống mẫu văn bản\PL4-Khung kế hoạch bài dạy (THCS).docx"  # TH template not on disk, use THCS template + TH margins/format
TPL_THCS = r"D:\UNIGO\Hệ thống mẫu văn bản\PL4-Khung kế hoạch bài dạy (THCS).docx"

PPCT_TIN = r"D:\UNIGO\Phân phối chương trình\Tin học\Kế hoạch dạy học môn Tin học 2026-2027.docx"
PPCT_ROB = r"D:\UNIGO\Phân phối chương trình\Robotics\Kế hoạch dạy học môn Robotics 2026-2027.docx"

OUT_TIN = r"D:\UNIGO\KHBD_Tin_học"
OUT_ROB = r"D:\UNIGO\KHBD_Robotics"

NLS_JSON = r"D:\UNIGO\.agents\skills\tao-khbd\references\cv3456_full_data.json"

TUAN_01_START = date(2026, 8, 3)  # Thứ Hai đầu tiên năm học

FONT_NAME = "Times New Roman"
FONT_SIZE_PT = 13

# EMU indent constants
INDENT_0 = 0
INDENT_1 = 180340      # ~0.5cm
INDENT_2 = 360045      # ~1.0cm
INDENT_BULLET = 540000  # ~1.5cm
INDENT_TH_1 = 457200   # ~1.27cm (TH cấp 1)
INDENT_TH_2 = 450215   # ~1.25cm (TH cấp 2)

# Lịch dạy - Tin học
TIN_SCHEDULE = {
    'TTH': (3, 'TT3'),
    '1': (0, '1A1'), '2': (1, '2A1'), '3': (3, '3A1'),
    '4': (2, '4C1'), '5': (1, '5C1'),
    '6': (4, '6A1'), '7': (1, '7A1'), '8': (4, '8A1'),
}
# Lịch dạy - Robotics
ROB_SCHEDULE = {
    '1': (3, '1A1'), '2': (2, '2A1'), '3': (2, '3A1'),
    '4': (4, '4C1'), '5': (1, '5C1'),
    '6': (4, '6A1'), '7': (1, '7A1'), '8': (4, '8A1'),
}

# ============================================================
# HELPER FUNCTIONS
# ============================================================
def afont(run, bold=False, italic=False, size_pt=FONT_SIZE_PT):
    """Apply Times New Roman font to a run."""
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)


def clean_body(doc):
    """Remove all body elements except sectPr (preserves headers/footers)."""
    for child in list(doc.element.body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'sectPr':
            doc.element.body.remove(child)


def add_p(doc, text="", bold=False, italic=False, first_indent=None,
          left_indent=None, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
          space_after_pt=3, space_before_pt=0, line_spacing=1.15, size_pt=FONT_SIZE_PT):
    """Add a formatted paragraph."""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = Pt(space_after_pt)
    pf.space_before = Pt(space_before_pt)
    if first_indent is not None:
        pf.first_line_indent = Emu(first_indent)
    if left_indent is not None:
        pf.left_indent = Emu(left_indent)
    if text:
        run = p.add_run(text)
        afont(run, bold=bold, italic=italic, size_pt=size_pt)
    return p


def add_multi_run(doc, runs_data, first_indent=None, left_indent=None,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.15):
    """Add paragraph with multiple runs [(text, bold, italic), ...]."""
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = Pt(3)
    pf.space_before = Pt(0)
    if first_indent is not None:
        pf.first_line_indent = Emu(first_indent)
    if left_indent is not None:
        pf.left_indent = Emu(left_indent)
    for text, bold, italic in runs_data:
        run = p.add_run(text)
        afont(run, bold=bold, italic=italic)
    return p


def set_borders(table, val="single", sz="4", color="000000"):
    """Set table borders via XML."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    borders_el = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), val)
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        borders_el.append(b)
    tblPr.append(borders_el)


def set_no_borders(table):
    """Set table to have no borders."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    borders_el = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'nil')
        borders_el.append(b)
    tblPr.append(borders_el)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, v in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def fill_cell(cell, text, bold=False, italic=False,
              align=WD_ALIGN_PARAGRAPH.LEFT, size_pt=FONT_SIZE_PT,
              space_after=2, line_spacing=1.15):
    """Fill a cell with text, supporting \\n for multiple paragraphs."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)
    set_cell_margins(cell)
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            p = cell.add_paragraph()
            p.alignment = align
            p.paragraph_format.line_spacing = line_spacing
            p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(line)
        afont(run, bold=bold, italic=italic, size_pt=size_pt)


def fill_cell_rich(cell, lines_data, align=WD_ALIGN_PARAGRAPH.LEFT,
                   size_pt=FONT_SIZE_PT, line_spacing=1.15):
    """Fill cell with rich content: [(text, bold, italic), ...]."""
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(2)
    set_cell_margins(cell)
    for i, (text, bold, italic) in enumerate(lines_data):
        if i > 0:
            p = cell.add_paragraph()
            p.alignment = align
            p.paragraph_format.line_spacing = line_spacing
            p.paragraph_format.space_after = Pt(2)
        run = p.add_run(text)
        afont(run, bold=bold, italic=italic, size_pt=size_pt)


def sanitize(name):
    import unicodedata
    name = unicodedata.normalize('NFD', name).encode('ascii', 'ignore').decode('utf-8')
    name = re.sub(r'[^a-zA-Z0-9_\-]', '_', name)
    name = re.sub(r'_+', '_', name).strip('_')
    return name


def compute_dates(tuan_so, day_of_week):
    week_start = TUAN_01_START + timedelta(weeks=tuan_so - 1)
    ngay_day = week_start + timedelta(days=day_of_week)
    ngay_soan = week_start - timedelta(days=2)  # Saturday tuần trước
    return ngay_soan.strftime('%d/%m/%Y'), ngay_day.strftime('%d/%m/%Y')


def get_bac_nls(grade):
    if grade <= 3:
        return 1
    elif grade <= 5:
        return 2
    elif grade <= 7:
        return 3
    else:
        return 4


def get_nls_level_key(grade):
    if grade <= 3:
        return "L1-3"
    elif grade <= 5:
        return "L4-5"
    elif grade <= 7:
        return "L6-7"
    else:
        return "L8-9"


def get_kit_name(grade):
    if grade <= 2:
        return "OLLO Kinder"
    elif grade <= 4:
        return "OLLO Initiate"
    else:
        return "OLLO Excel 1"


def is_thcs(grade):
    return grade >= 6


def get_to_chuyen_mon(grade):
    return "Tổ chuyên môn THCS" if grade >= 6 else "Tổ chuyên môn Tiểu học"


def yccd_to_noun(item):
    """Convert a single YCCD item from verb phrase to noun phrase."""
    import re
    item = item.strip().strip('-').strip()
    if not item:
        return ''
    # Remove leading verb patterns to convert to noun phrase
    verb_patterns = [
        r'^Nhận biết được\s+', r'^Phân biệt được\s+', r'^Nêu được\s+',
        r'^Giải thích được\s+', r'^Biết\s+', r'^Hiểu được\s+',
        r'^Trình bày được\s+', r'^Mô tả được\s+', r'^Vận dụng được\s+',
        r'^Thực hiện được\s+', r'^Sử dụng được\s+', r'^Xác định được\s+',
        r'^Liệt kê được\s+', r'^So sánh được\s+', r'^Phân tích được\s+',
        r'^Đánh giá được\s+', r'^Tạo được\s+', r'^Viết được\s+',
        r'^Lập được\s+', r'^Thiết kế được\s+', r'^Lắp ráp được\s+',
        r'^Kể tên được\s+', r'^Nhận diện được\s+',
    ]
    for pat in verb_patterns:
        item = re.sub(pat, '', item, count=1)
    # Also remove banned noun patterns
    banned = ['Sự hiểu biết về ', 'Khả năng nhận diện ', 'Khả năng phân tích ',
              'Khả năng vận dụng ', 'Sự nhận biết ']
    for b in banned:
        item = item.replace(b, '')
    # Capitalize first letter
    if item:
        item = item[0].upper() + item[1:]
    # Remove trailing period if doubled
    item = item.rstrip('.')
    return item


def parse_yccd_bullets(yccd_raw):
    """Parse raw YCCD text from PPCT into list of noun-phrase bullets."""
    if not yccd_raw or len(yccd_raw.strip()) < 5:
        return []
    # Split on "- " pattern (may have spaces before dash)
    import re
    items = re.split(r'\s*-\s+', yccd_raw.strip())
    result = []
    for item in items:
        item = item.strip()
        if not item:
            continue
        noun = yccd_to_noun(item)
        if noun and len(noun) > 3:
            result.append(noun)
    return result


def parse_yccd_for_th(yccd_raw):
    """Parse raw YCCD text from PPCT/SGK into clean bullet points for Tiểu học."""
    if not yccd_raw or len(yccd_raw.strip()) < 3:
        return []
    text = yccd_raw.strip()
    items = []
    if ' - ' in text or text.startswith('- '):
        parts = re.split(r'\s*-\s+', text)
        items = [p.strip() for p in parts if p.strip()]
    elif ';' in text:
        parts = text.split(';')
        items = [p.strip() for p in parts if p.strip()]
    elif '\n' in text:
        parts = text.split('\n')
        items = [p.strip().lstrip('-').strip() for p in parts if p.strip()]
    else:
        items = [text]

    cleaned = []
    for it in items:
        it = it.lstrip('-').strip()
        if it:
            it = it[0].upper() + it[1:]
            if not it.endswith('.'):
                it += '.'
            cleaned.append(it)
    return cleaned


def format_th_title(raw_title, tiet_ppct):
    """Format title for Tiểu học cleanly without redundant 'BÀI: Bài'."""
    t = raw_title.strip()
    m = re.match(r'^(?:Bài|BÀI)\s*(\d+)[\.:\s]*(.*)$', t)
    if m:
        bai_num, bai_name = m.group(1), m.group(2).strip()
        return f"BÀI {bai_num}. {bai_name.upper()} (Tiết: {tiet_ppct} theo PPCT)"
    if t.lower().startswith('bài'):
        clean = re.sub(r'^(?:Bài|BÀI)[\s:]*', '', t).strip()
        return f"BÀI: {clean.upper()} (Tiết: {tiet_ppct} theo PPCT)"
    if any(t.lower().startswith(x) for x in ['tiết 0', 'ôn tập', 'đánh giá', 'tổng kết']):
        return f"{t.upper()} (Tiết: {tiet_ppct} theo PPCT)"
    return f"BÀI: {t.upper()} (Tiết: {tiet_ppct} theo PPCT)"


# ============================================================
# TIỂU HỌC BUILDER (TTH + Lớp 1-5)
# ============================================================
def build_khbd_th(mon_hoc, grade, grade_label, ten_lop, title, tiet_ppct,
                  yccd, tuan_so, day_of_week, kit_name=None):
    """Build KHBD for Tiểu học (TTH, Lớp 1-5)."""
    doc = Document(TPL_TH)
    clean_body(doc)

    # Margins: L3cm, R2cm, T2cm, B2cm
    for sec in doc.sections:
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(2.0)
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)

    ngay_soan, ngay_day = compute_dates(tuan_so, day_of_week)
    ls = 1.5  # TH line spacing
    to_cm = get_to_chuyen_mon(grade)

    # ---- Table info 2×2 — NO BORDER (đúng template PL4) ----
    tbl_info = doc.add_table(rows=2, cols=2)
    set_no_borders(tbl_info)
    fill_cell(tbl_info.rows[0].cells[0], 'Trường Tiểu học và THCS UNIGO', bold=True, line_spacing=ls)
    fill_cell(tbl_info.rows[0].cells[1], f'Họ tên giáo viên: Đậu Đình Nguyên', line_spacing=ls)
    fill_cell(tbl_info.rows[1].cells[0], f'Tổ {to_cm.replace("Tổ chuyên môn ", "")}', line_spacing=ls)
    fill_cell(tbl_info.rows[1].cells[1], f'Ngày soạn: {ngay_soan}\nNgày dạy: {ngay_day}\nLớp: {ten_lop}', line_spacing=ls)

    # ---- Tiêu đề bài học (KHÔNG lặp lại Thứ ngày hay Tên GV vì đã có ở Table 0) ----
    mon_display = mon_hoc.upper()
    add_p(doc, f'KẾ HOẠCH DẠY HỌC MÔN {mon_display}', bold=True,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls, size_pt=14)

    if mon_hoc == "Robotics" and kit_name:
        add_p(doc, f'CHỦ ĐIỂM: BỘ THIẾT BỊ: {kit_name.upper()}', bold=True,
              alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)

    title_formatted = format_th_title(title, tiet_ppct)
    add_p(doc, title_formatted, bold=True,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)
    add_p(doc, '', line_spacing=ls)

    # ---- I. YÊU CẦU CẦN ĐẠT ----
    add_p(doc, 'I. YÊU CẦU CẦN ĐẠT:', bold=True, first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, '- Sau bài học này em sẽ:', first_indent=INDENT_TH_1, line_spacing=ls)

    # Yêu cầu cần đạt của học sinh trích xuất từ SGK / PPCT
    yccd_items = parse_yccd_for_th(yccd)
    for y_item in yccd_items:
        add_p(doc, f'+ {y_item}', left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=1.33)

    # Phẩm chất TRƯỚC cho TH
    add_p(doc, '1. Phát triển phẩm chất', bold=True, first_indent=INDENT_TH_1, line_spacing=ls)

    pham_chat_items = [
        f'- Chăm chỉ: Hăng hái tham gia hoạt động tìm hiểu bài học, kiên trì thực hành theo hướng dẫn. (Đạt được thông qua Hoạt động 2, Hoạt động 3)',
        f'- Trách nhiệm: Giữ gìn cẩn thận thiết bị học tập, thu dọn gọn gàng sau giờ học. (Đạt được thông qua Hoạt động 4)',
        f'- Trung thực: Tự giác làm bài, tôn trọng kết quả của bản thân và bạn bè. (Đạt được thông qua Hoạt động 3)',
    ]
    for item in pham_chat_items:
        add_p(doc, item, left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=1.33)

    # Năng lực SAU cho TH
    add_p(doc, '2. Phát triển năng lực', bold=True, first_indent=INDENT_TH_1, line_spacing=ls)

    bac = get_bac_nls(grade)

    # 2.1. Năng lực đặc thù
    add_p(doc, f'2.1. Năng lực đặc thù ({mon_hoc}):', bold=True,
          first_indent=INDENT_TH_2, line_spacing=1.33)
    if mon_hoc == "Tin học":
        add_p(doc, f'- NLa (Sử dụng và quản lí các phương tiện ICT): Nhận diện và thao tác đúng các công cụ, phần mềm trong bài {title}. (Đạt được thông qua Hoạt động 2)',
              left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=1.33)
        add_p(doc, f'- NLd (Ứng dụng ICT trong học và tự học): Sử dụng phần mềm tạo sản phẩm số phục vụ học tập. (Đạt được thông qua Hoạt động 3)',
              left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=1.33)
    else:
        add_p(doc, f'- NLa (Sử dụng và quản lí các phương tiện ICT): Nhận biết tên gọi, hình dạng, chức năng các chi tiết bộ kit {kit_name}. (Đạt được thông qua Hoạt động 2)',
              left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=1.33)
        add_p(doc, f'- NLd (Ứng dụng ICT trong học và tự học): Thực hiện lắp ráp đúng quy trình mô hình {title}. (Đạt được thông qua Hoạt động 3)',
              left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=1.33)

    # 2.2. Năng lực số
    add_p(doc, '2.2. Năng lực số (Thông tư 02/2025 – CV 3456):', bold=True,
          first_indent=INDENT_TH_2, line_spacing=1.33)
    add_p(doc, f'- Miền V. Giải quyết vấn đề (thành tố 5.3. Sử dụng sáng tạo công nghệ số – Bậc {bac}): Sử dụng công cụ số để giải quyết vấn đề đơn giản trong bài học. (Đạt được thông qua Hoạt động 2, Hoạt động 3)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=1.33)

    # 2.3. Năng lực chung
    add_p(doc, '2.3. Năng lực chung:', bold=True, first_indent=INDENT_TH_2, line_spacing=1.33)
    add_p(doc, '- Tự chủ và tự học: Chủ động quan sát hướng dẫn và tự thực hiện nhiệm vụ. (Đạt được thông qua Hoạt động 2)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=1.33)
    add_p(doc, '- Giao tiếp và hợp tác: Thảo luận nhóm, chia sẻ và hỗ trợ bạn trong giờ thực hành. (Đạt được thông qua Hoạt động 3)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=1.33)

    # ---- II. ĐỒ DÙNG DẠY HỌC ----
    add_p(doc, 'II. ĐỒ DÙNG DẠY HỌC :', bold=True, first_indent=INDENT_0, line_spacing=ls)
    if mon_hoc == "Tin học":
        add_p(doc, '1. Giáo viên: Máy tính kết nối Internet, máy chiếu, bài giảng điện tử, phiếu học tập.',
              first_indent=INDENT_TH_1, line_spacing=ls)
        add_p(doc, '2. Học sinh: Máy tính thực hành, SGK Tin học, vở ghi.',
              first_indent=INDENT_TH_1, line_spacing=ls)
    else:
        add_p(doc, f'1. Giáo viên: Bộ Kit Robotics {kit_name} mẫu, máy chiếu, slide hướng dẫn lắp ráp {title}, phiếu hướng dẫn thực hành.',
              first_indent=INDENT_TH_1, line_spacing=ls)
        add_p(doc, f'2. Học sinh: Bộ Kit Robotics {kit_name} theo nhóm, dụng cụ tháo chốt, vở ghi.',
              first_indent=INDENT_TH_1, line_spacing=ls)

    # ---- III. PHƯƠNG PHÁP, KĨ THUẬT DẠY HỌC ----
    add_p(doc, 'III. PHƯƠNG PHÁP, KĨ THUẬT DẠY HỌC', bold=True,
          first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, '- Các phương pháp: Vấn đáp, thực hành, hoạt động nhóm, giải quyết vấn đề.',
          first_indent=INDENT_TH_1, line_spacing=ls)
    add_p(doc, '- Kĩ thuật: Đặt câu hỏi, trình bày 1 phút, động não, chia sẻ nhóm đôi.',
          first_indent=INDENT_TH_1, line_spacing=ls)

    # ---- IV. CÁC HOẠT ĐỘNG DẠY - HỌC CHỦ YẾU ----
    add_p(doc, 'IV. CÁC HOẠT ĐỘNG DẠY - HỌC CHỦ YẾU:', bold=True,
          first_indent=INDENT_0, line_spacing=ls)

    # Bảng 2 cột với hàng gộp (gridSpan) cho TH
    table = doc.add_table(rows=1, cols=2)
    set_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr = table.rows[0].cells
    fill_cell(hdr[0], 'HOẠT ĐỘNG CỦA GV – HS', bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)
    fill_cell(hdr[1], 'KẾT QUẢ CẦN ĐẠT', bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.5)

    # TH Activities data
    th_activities = _get_th_activities(mon_hoc, title, kit_name)
    for act_title, gv_hs_text, ket_qua_text in th_activities:
        # Merged header row
        row_hdr = table.add_row()
        c0 = row_hdr.cells[0]
        c1 = row_hdr.cells[1]
        tcPr = c0._tc.get_or_add_tcPr()
        gs = OxmlElement('w:gridSpan')
        gs.set(qn('w:val'), '2')
        tcPr.append(gs)
        row_hdr._tr.remove(c1._tc)
        fill_cell(c0, act_title, bold=True, size_pt=13, line_spacing=1.5)

        # Content row
        row_cnt = table.add_row()
        fill_cell(row_cnt.cells[0], gv_hs_text, line_spacing=1.5)
        fill_cell(row_cnt.cells[1], ket_qua_text, line_spacing=1.5)

    # ---- V. ĐIỀU CHỈNH - BỔ SUNG ----
    add_p(doc, '', line_spacing=ls)
    add_p(doc, 'V. ĐIỀU CHỈNH - BỔ SUNG SAU TIẾT DẠY :', bold=True,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)
    add_p(doc, '(GV ghi những nội dung mà mình đã bổ sung ngoài KHBD đã lên...)', bold=True,
          italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)
    add_p(doc, '.........................................................................................................................',
          line_spacing=ls)
    add_p(doc, '', line_spacing=ls)

    # Bảng chữ ký 3×3 — NO BORDER (chuẩn AGENTS.md)
    tbl_sign = doc.add_table(rows=3, cols=3)
    set_no_borders(tbl_sign)
    for i, txt in enumerate(['DUYỆT CỦA BGH', 'DUYỆT CỦA TỔ CM', 'NGƯỜI SOẠN']):
        fill_cell(tbl_sign.rows[0].cells[i], txt, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)
    for i in range(3):
        fill_cell(tbl_sign.rows[1].cells[i], '(Ký, ghi rõ họ tên)',
                  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)
    for i in range(3):
        fill_cell(tbl_sign.rows[2].cells[i], '\n\n\n',
                  align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)
    fill_cell(tbl_sign.rows[2].cells[2], '\n\n\nĐậu Đình Nguyên', bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)

    return doc


def _get_th_activities(mon_hoc, title, kit_name):
    """Return list of (act_title, gv_hs_text, ket_qua_text) for TH."""
    if mon_hoc == "Tin học":
        return [
            (f'1. Hoạt động MỞ ĐẦU (5 phút)\n*Mục tiêu: Kích hoạt hứng thú, tạo tình huống kết nối vào bài học {title}.',
             f'GV: Chiếu hình ảnh/video liên quan đến {title}. Đặt câu hỏi gợi mở.\n- GV gọi 2-3 HS trả lời.\n- GV nhận xét, dẫn dắt vào bài mới.\nHS: Quan sát, suy nghĩ, trả lời câu hỏi.\n- Thảo luận nhanh nhóm đôi.\n- Mở SGK chuẩn bị bài.',
             f'HS xác định được vấn đề bài học.\nHS hứng thú với nội dung mới.'),
            (f'2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI (15 phút)\n*Mục tiêu: HS nắm được các khái niệm, quy trình và thao tác kỹ thuật của bài {title}.',
             f'GV: Hướng dẫn HS đọc SGK, trình chiếu thao tác mẫu.\n- Nhấn mạnh các lưu ý và lỗi thường gặp.\n- Gọi HS lên thực hiện lại thao tác mẫu.\nHS: Theo dõi hướng dẫn, ghi chép kiến thức.\n- Đặt câu hỏi khi chưa hiểu.\n- Thực hiện thao tác mẫu khi được gọi.',
             f'HS ghi nhớ kiến thức chuẩn.\nHS nắm được quy trình thao tác.'),
            (f'3. HĐ LUYỆN TẬP-THỰC HÀNH (10 phút)\n*Mục tiêu: HS rèn kỹ năng thực hành trên máy tính.',
             f'GV: Giao bài tập thực hành trên máy tính.\n- Quan sát, hỗ trợ HS gặp khó khăn.\n- Kiểm tra sản phẩm tại máy.\nHS: Thực hành cá nhân trên máy tính.\n- Trao đổi nhóm đôi khi gặp khó.\n- Đối chiếu kết quả với bạn.',
             f'Sản phẩm thực hành đạt yêu cầu.\nHS tự tin thao tác trên máy.'),
            (f'4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
             f'GV: Đặt câu hỏi liên hệ thực tế.\n- Giao nhiệm vụ về nhà.\n- Hướng dẫn tắt máy, xếp gọn phòng học.\nHS: Suy nghĩ ý tưởng ứng dụng.\n- Ghi nhận nhiệm vụ về nhà.\n- Tắt máy an toàn, xếp bàn ghế.',
             f'HS liên hệ được kiến thức với thực tế.\nPhòng máy được thu dọn gọn gàng.'),
        ]
    else:  # Robotics
        return [
            (f'1. Hoạt động MỞ ĐẦU (5 phút)\n*Mục tiêu: Kích hoạt hứng thú, tạo tình huống kết nối vào bài học {title}.',
             f'GV: Chiếu video/hình ảnh liên quan đến {title}. Đặt câu hỏi gợi mở.\n- Mời 2-3 HS trả lời.\n- Chốt kiến thức, dẫn dắt vào bài.\nHS: Quan sát màn chiếu, suy nghĩ.\n- Thảo luận nhanh nhóm đôi.\n- Hăng hái giơ tay trả lời.',
             f'HS nhận diện được vấn đề kỹ thuật.\nHS sẵn sàng bước vào hoạt động chính.'),
            (f'2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI (15 phút)\n*Mục tiêu: Tìm hiểu linh kiện và nguyên lý lắp ráp mô hình {title}.',
             f'GV: Giới thiệu linh kiện, trình chiếu sơ đồ lắp ráp.\n- Hướng dẫn cách nhặt đúng linh kiện.\n- Kiểm tra khay linh kiện từng nhóm.\nHS: Theo dõi slide hướng dẫn.\n- Nhặt đúng số lượng linh kiện.\n- Ghi nhớ thứ tự các bước ghép.',
             f'HS nhận biết tên gọi linh kiện.\nHS nắm được quy trình lắp ráp.'),
            (f'3. HĐ LUYỆN TẬP-THỰC HÀNH (10 phút)\n*Mục tiêu: Thực hành lắp ráp và vận hành chạy thử robot {title}.',
             f'GV: Giao nhiệm vụ lắp ráp theo nhóm.\n- Quan sát, hỗ trợ nhóm gặp khó.\n- Cho các nhóm bật nguồn chạy thử.\nHS: Phân công nhiệm vụ trong nhóm.\n- Lắp ráp cẩn thận theo sơ đồ.\n- Bật robot chạy thử, quan sát.',
             f'Mô hình robot hoàn thiện.\nRobot hoạt động đúng yêu cầu.'),
            (f'4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
             f'GV: Đặt câu hỏi cải tiến tính năng robot.\n- Hướng dẫn tháo dỡ và kiểm kê linh kiện.\n- Nhận xét tiết học, dặn dò bài sau.\nHS: Suy nghĩ ý tưởng nâng cấp.\n- Tháo robot, xếp linh kiện vào hộp.\n- Đóng hộp Kit, cất gọn gàng.',
             f'HS đề xuất ý tưởng cải tiến.\nBộ Kit được sắp xếp ngăn nắp.'),
        ]


# ============================================================
# THCS BUILDER (Lớp 6-8)
# ============================================================
def build_khbd_thcs(mon_hoc, grade, grade_label, ten_lop, title, tiet_ppct,
                    yccd, tuan_so, day_of_week, kit_name=None):
    """Build KHBD for THCS (Lớp 6-8)."""
    doc = Document(TPL_THCS)
    clean_body(doc)

    # Margins: L2.54cm, R1.27cm, T1.27cm, B1.27cm
    for sec in doc.sections:
        sec.left_margin = Cm(2.54)
        sec.right_margin = Cm(1.27)
        sec.top_margin = Cm(1.27)
        sec.bottom_margin = Cm(1.27)

    ngay_soan, ngay_day = compute_dates(tuan_so, day_of_week)
    ls = 1.15  # THCS line spacing
    bac = get_bac_nls(grade)
    to_cm = get_to_chuyen_mon(grade)

    # ---- Table info 2×2 — NO BORDER (đúng template PL4) ----
    tbl_info = doc.add_table(rows=2, cols=2)
    set_no_borders(tbl_info)
    fill_cell(tbl_info.rows[0].cells[0], 'Trường Tiểu học và THCS UNIGO', bold=True)
    fill_cell(tbl_info.rows[0].cells[1], f'Họ tên giáo viên: Đậu Đình Nguyên')
    fill_cell(tbl_info.rows[1].cells[0], f'Tổ {to_cm.replace("Tổ chuyên môn ", "")}')
    fill_cell(tbl_info.rows[1].cells[1], f'Ngày soạn: {ngay_soan}\nNgày dạy: {ngay_day}\nLớp: {ten_lop}')

    # ---- Title Block ----
    add_p(doc, '', line_spacing=ls)
    add_p(doc, f'TÊN BÀI DẠY: {title.upper()}', bold=True,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, first_indent=INDENT_0, line_spacing=ls)

    if mon_hoc == "Robotics":
        add_p(doc, 'Môn học: Robotics',
              bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)
    else:
        add_p(doc, 'Môn học: Tin học',
              bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)
    add_p(doc, 'Thời lượng: 1 tiết (45 phút)',
          bold=True, italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)
    add_p(doc, f'Tiết theo PPCT: {tiet_ppct}', bold=True,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)

    # ---- I. Mục tiêu (Kiến thức → Năng lực → Phẩm chất cho THCS) ----
    add_p(doc, 'I. Mục tiêu', bold=True, first_indent=INDENT_0, line_spacing=ls)

    # 1. Kiến thức — Tách YCCD từ PPCT thành từng dòng danh từ
    add_p(doc, '1. Kiến thức:', bold=True, first_indent=INDENT_1, line_spacing=ls)
    yccd_bullets = parse_yccd_bullets(yccd)
    if yccd_bullets:
        for bullet in yccd_bullets:
            add_p(doc, f'- {bullet}.',
                  left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
    else:
        add_p(doc, f'- Các khái niệm, quy trình và kỹ năng thực hành trong bài {title}.',
              left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)

    # 2. Năng lực
    add_p(doc, '2. Năng lực:', bold=True, first_indent=INDENT_1, line_spacing=ls)

    # 2.1. Năng lực đặc thù
    add_p(doc, f'2.1. Năng lực đặc thù ({mon_hoc}):', bold=True,
          first_indent=INDENT_2, line_spacing=ls)
    if mon_hoc == "Tin học":
        add_p(doc, f'- NLa (Sử dụng và quản lí các phương tiện ICT): Thao tác đúng quy trình các công cụ, phần mềm trong bài {title}. (Đạt được thông qua Hoạt động 2)',
              left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
        add_p(doc, f'- NLc (Giải quyết vấn đề với sự hỗ trợ của ICT): Hoàn thành nhiệm vụ thực hành và xử lý tình huống thực tế. (Đạt được thông qua Hoạt động 3)',
              left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
    else:
        add_p(doc, f'- NLa (Sử dụng và quản lí các phương tiện ICT): Thiết kế mô phỏng 3D, thao tác lắp ráp chuẩn xác các khớp nối. (Đạt được thông qua Hoạt động 2, Hoạt động 3)',
              left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
        add_p(doc, f'- NLd (Ứng dụng ICT trong học và tự học): Nạp code lập trình và hiệu chỉnh robot {title}. (Đạt được thông qua Hoạt động 3)',
              left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)

    # 2.2. Năng lực số
    add_p(doc, '2.2. Năng lực số (Thông tư 02/2025 – CV 3456):', bold=True,
          first_indent=INDENT_2, line_spacing=ls)
    add_p(doc, f'- Miền I. Khai thác dữ liệu và thông tin (thành tố 1.1. Duyệt, tìm kiếm và lọc dữ liệu, thông tin và nội dung số – Bậc {bac}): Khai thác thông tin số phục vụ bài học. (Đạt được thông qua Hoạt động 2)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, f'- Miền IV. An toàn (thành tố 4.1. Bảo vệ thiết bị và nội dung số – Bậc {bac}): Tuân thủ quy tắc an toàn thiết bị và bảo vệ dữ liệu. (Đạt được thông qua Hoạt động 3, Hoạt động 4)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)

    # 2.3. Năng lực chung
    add_p(doc, '2.3. Năng lực chung:', bold=True, first_indent=INDENT_2, line_spacing=ls)
    add_p(doc, '- Tự chủ và tự học: Chủ động đọc tài liệu SGK, quan sát thao tác mẫu và tự thực hiện nhiệm vụ học tập. (Đạt được thông qua Hoạt động 1, Hoạt động 2)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, '- Giao tiếp và hợp tác: Thảo luận nhóm hiệu quả, chia sẻ và hỗ trợ bạn trong giờ thực hành. (Đạt được thông qua Hoạt động 2, Hoạt động 3)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, '- Giải quyết vấn đề và sáng tạo: Phát hiện và sửa lỗi trong quá trình thực hành, đề xuất ý tưởng ứng dụng mới. (Đạt được thông qua Hoạt động 4)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)

    # 3. Phẩm chất
    add_p(doc, '3. Phẩm chất:', bold=True, first_indent=INDENT_1, line_spacing=ls)
    add_p(doc, '- Chăm chỉ: Tích cực tham gia hoạt động tìm hiểu và hoàn thành bài tập thực hành. (Thông qua Hoạt động 1, Hoạt động 2)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, '- Trung thực: Đánh giá đúng kết quả học tập, tôn trọng sản phẩm của người khác. (Thông qua Hoạt động 3)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, '- Trách nhiệm: Giữ gìn an toàn thiết bị, bảo vệ tài sản công cộng và dữ liệu số. (Thông qua Hoạt động 3, Hoạt động 4)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)

    # ---- II. Thiết bị dạy học và học liệu ----
    add_p(doc, 'II. Thiết bị dạy học và học liệu:', bold=True,
          first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, '1. Thiết bị:', bold=True, first_indent=INDENT_1, line_spacing=ls)
    if mon_hoc == "Tin học":
        add_p(doc, '- Máy tính giáo viên có kết nối Internet và máy chiếu.',
              left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
        add_p(doc, '- Phòng máy tính cho học sinh thực hành.',
              left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
    else:
        add_p(doc, f'- Bộ Kit Robotics {kit_name}, máy tính GV, máy chiếu.',
              left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
        add_p(doc, '- Phần mềm nạp lập trình, máy chiếu bài giảng 3D.',
              left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)

    add_p(doc, '2. Học liệu:', bold=True, first_indent=INDENT_1, line_spacing=ls)
    if mon_hoc == "Tin học":
        add_p(doc, f'- SGK Tin học {grade_label}, bài giảng điện tử tương tác.',
              left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
    else:
        add_p(doc, f'- Phiếu học tập thực hành, sơ đồ nguyên lý cơ khí mô hình {title}.',
              left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, '- Phiếu học tập, tệp bài tập thực hành mẫu.',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)

    # ---- III. Tiến trình dạy học ----
    add_p(doc, 'III. Tiến trình dạy học', bold=True, first_indent=INDENT_0, line_spacing=ls)

    thcs_activities = _get_thcs_activities(mon_hoc, title, kit_name)

    for idx, (act_title, muc_tieu, noi_dung, san_pham, gv_hs_col, ket_qua_col) in enumerate(thcs_activities, 1):
        add_p(doc, act_title, bold=True, first_indent=INDENT_1, line_spacing=ls)
        add_multi_run(doc, [
            ('a) Mục tiêu: ', True, False), (muc_tieu, False, True)
        ], first_indent=INDENT_1, line_spacing=ls)
        add_multi_run(doc, [
            ('b) Nội dung: ', True, False), (noi_dung, False, True)
        ], first_indent=INDENT_1, line_spacing=ls)
        add_multi_run(doc, [
            ('c) Sản phẩm: ', True, False), (san_pham, False, True)
        ], first_indent=INDENT_1, line_spacing=ls)
        add_multi_run(doc, [
            ('d) Tổ chức thực hiện:', True, False),
        ], first_indent=INDENT_1, line_spacing=ls)

        # Bảng 2 cột: HOẠT ĐỘNG CỦA GV – HS | KẾT QUẢ CẦN ĐẠT
        tbl = doc.add_table(rows=1, cols=2)
        set_borders(tbl)
        # Set column widths ~55% / ~45%
        for cell in tbl.columns[0].cells:
            cell.width = Cm(9.0)
        for cell in tbl.columns[1].cells:
            cell.width = Cm(7.0)

        # Header
        fill_cell(tbl.rows[0].cells[0], 'HOẠT ĐỘNG CỦA GV – HS', bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(tbl.rows[0].cells[1], 'KẾT QUẢ CẦN ĐẠT', bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

        # Content row
        row = tbl.add_row()
        fill_cell(row.cells[0], gv_hs_col)
        fill_cell(row.cells[1], ket_qua_col)

        add_p(doc, '', line_spacing=ls)

    # ---- Rút kinh nghiệm + Bảng ký tên ----
    add_p(doc, 'RÚT KINH NGHIỆM SAU BÀI DẠY:', bold=True,
          first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, '...........................................................................................................................', line_spacing=ls)
    add_p(doc, '...........................................................................................................................', line_spacing=ls)
    add_p(doc, '', line_spacing=ls)

    # Bảng chữ ký 3×3 — NO BORDER
    tbl_sign = doc.add_table(rows=3, cols=3)
    set_no_borders(tbl_sign)
    for i, txt in enumerate(['DUYỆT CỦA BGH', 'DUYỆT CỦA TỔ CM', 'NGƯỜI SOẠN']):
        fill_cell(tbl_sign.rows[0].cells[i], txt, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    for i in range(3):
        fill_cell(tbl_sign.rows[1].cells[i], '(Ký, ghi rõ họ tên)',
                  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i in range(3):
        fill_cell(tbl_sign.rows[2].cells[i], '\n\n\n',
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell(tbl_sign.rows[2].cells[2], '\n\n\nĐậu Đình Nguyên', bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER)

    return doc


def _get_thcs_activities(mon_hoc, title, kit_name):
    """Return THCS activities: [(act_title, mt, nd, sp, gv_hs_col, ket_qua_col), ...]."""
    buoc1 = 'Bước 1: Chuyển giao nhiệm vụ học tập'
    buoc2 = 'Bước 2: Học sinh tiếp nhận nhiệm vụ học tập'
    buoc3 = 'Bước 3: Báo cáo kết quả hoạt động'
    buoc4 = 'Bước 4: Đánh giá kết quả thực hiện nhiệm vụ'
    buoc4_last = 'Bước 4: Giáo viên nhắc nhở nhiệm vụ về nhà'

    if mon_hoc == "Tin học":
        return [
            ('1. Hoạt động 1. Khởi động (Xác định vấn đề/nhiệm vụ học tập/Mở đầu) (7 phút)',
             f'Kích hoạt hiểu biết nền, tạo hứng thú và kết nối HS vào bài học {title}.',
             f'GV đặt câu hỏi tình huống hoặc chiếu hình ảnh/video thực tế liên quan đến {title}.',
             'HS đưa ra câu trả lời ban đầu, xác định được nhiệm vụ trọng tâm.',
             f'{buoc1}:\n- GV chiếu hình ảnh/video tình huống liên quan đến {title}.\n- GV nêu câu hỏi gợi mở cho cả lớp.\n{buoc2}:\n- HS tập trung quan sát, lắng nghe câu hỏi.\n- HS suy nghĩ cá nhân, thảo luận nhanh nhóm đôi.\n{buoc3}:\n- 2-3 HS giơ tay phát biểu ý kiến.\n- Các bạn khác lắng nghe, nhận xét.\n{buoc4}:\n- GV nhận xét, chuẩn hóa và dẫn dắt vào bài mới.',
             f'HS nhận diện được vấn đề bài học.\nHS hứng thú, sẵn sàng tìm hiểu nội dung mới.'),

            ('2. Hoạt động 2. Hình thành kiến thức mới/giải quyết vấn đề (18 phút)',
             f'HS nắm vững các khái niệm, quy trình và thao tác kỹ thuật của bài {title}.',
             'Tìm hiểu kiến thức trong SGK, quan sát GV làm mẫu thao tác.',
             'Ghi nhớ kiến thức chuẩn, hoàn thành phiếu học tập hoặc thao tác mẫu đúng quy trình.',
             f'{buoc1}:\n- GV giao nhiệm vụ: đọc SGK và theo dõi thao tác mẫu trên máy chiếu.\n{buoc2}:\n- HS mở SGK trang tương ứng và quan sát màn chiếu.\n- HS ghi chép từ khóa chính vào vở.\n{buoc3}:\n- 1-2 HS lên máy GV thực hiện lại thao tác mẫu.\n- Cả lớp quan sát, nhận xét.\n{buoc4}:\n- GV chuẩn hóa kiến thức lý thuyết và chốt quy trình thao tác.',
             f'Kiến thức lý thuyết chuẩn.\nQuy trình thao tác đúng.'),

            ('3. Hoạt động 3. Luyện tập (12 phút)',
             f'HS rèn luyện kỹ năng thực hành trên máy tính.',
             'Thực hành cá nhân/nhóm đôi trên máy tính theo bài tập SGK/phiếu bài tập.',
             'Sản phẩm thực hành trên phần mềm đạt yêu cầu kỹ thuật.',
             f'{buoc1}:\n- GV giao bài tập thực hành bài {title} trên máy tính.\n{buoc2}:\n- HS mở phần mềm/tệp thực hành, bắt đầu làm bài.\n- HS trao đổi nhóm đôi khi gặp khó khăn.\n{buoc3}:\n- GV kiểm tra sản phẩm tại máy.\n- HS đổi chéo kiểm tra bài của nhau.\n{buoc4}:\n- GV nhận xét, đánh giá kết quả thực hành.\n- Tuyên dương HS làm tốt.',
             f'Sản phẩm thực hành hoàn thành.\nHS tự tin thao tác trên máy tính.'),

            ('4. Hoạt động 4. Vận dụng (8 phút)',
             'Vận dụng kiến thức, kỹ năng đã học vào tình huống thực tế cuộc sống.',
             'GV đặt câu hỏi mở rộng, yêu cầu liên hệ thực tiễn.',
             'Ý kiến đề xuất giải pháp, phòng máy được tắt an toàn.',
             f'{buoc1}:\n- GV nêu câu hỏi: "Em có thể ứng dụng kiến thức bài {title} như thế nào trong cuộc sống?"\n{buoc2}:\n- HS suy nghĩ, thảo luận nhanh hoặc ghi chép ý tưởng.\n{buoc3}:\n- 1-2 HS chia sẻ ý tưởng vận dụng trước lớp.\n{buoc4_last}:\n- GV tổng kết tiết học, giao nhiệm vụ chuẩn bị bài sau.\n- Hướng dẫn HS lưu bài, tắt máy và xếp gọn phòng học.',
             f'HS liên hệ được kiến thức với thực tế.\nPhòng máy thu dọn gọn gàng.'),
        ]
    else:  # Robotics THCS
        return [
            ('1. Hoạt động 1. Khởi động (Xác định vấn đề/nhiệm vụ học tập/Mở đầu) (7 phút)',
             f'Kích hoạt tư duy kỹ thuật và kết nối tình huống vào bài học {title}.',
             f'GV chiếu hình ảnh/video thực tế về {title} và đặt câu hỏi phân tích cơ chế.',
             'Câu trả lời của HS xác định được vấn đề kỹ thuật cần giải quyết.',
             f'{buoc1}:\n- GV trình chiếu clip/hình ảnh thực tế liên quan đến {title}.\n- Yêu cầu HS quan sát và phân tích.\n{buoc2}:\n- HS tập trung quan sát màn chiếu.\n- HS thảo luận nhanh theo nhóm bàn.\n{buoc3}:\n- 2 đại diện HS trả lời câu hỏi khởi động.\n- Nêu giả thuyết và nguyên lý vận hành ban đầu.\n{buoc4}:\n- GV nhận xét, chuẩn hóa và dẫn dắt vào bài mới.',
             f'HS nhận diện được vấn đề kỹ thuật.\nHS sẵn sàng bước vào hoạt động chính.'),

            ('2. Hoạt động 2. Hình thành kiến thức mới/giải quyết vấn đề (18 phút)',
             f'Nghiên cứu sơ đồ thiết kế và tìm hiểu linh kiện cho mô hình {title}.',
             'GV yêu cầu HS tìm hiểu linh kiện khung, chốt, động cơ và cơ cấu truyền động.',
             'Sơ đồ khối nguyên lý cơ khí và danh mục linh kiện chính xác.',
             f'{buoc1}:\n- GV phát phiếu hướng dẫn và trình chiếu sơ đồ 2D/3D của {title}.\n{buoc2}:\n- HS tiếp nhận phiếu, quan sát các góc ghép linh kiện.\n- HS đối chiếu danh mục, nhặt đúng số lượng chốt và khung nối.\n{buoc3}:\n- GV kiểm tra khay linh kiện từng nhóm tại bàn.\n- Đại diện nhóm giơ khay linh kiện để kiểm tra.\n{buoc4}:\n- GV chốt quy trình các bước lắp ráp mô hình robot.',
             f'HS nắm được quy trình lắp ráp.\nDanh mục linh kiện đầy đủ, chính xác.'),

            ('3. Hoạt động 3. Luyện tập (12 phút)',
             f'Thực hành lắp ráp, lập trình và chạy thử mô hình robot {title}.',
             'GV yêu cầu các nhóm tiến hành lắp ráp phần cứng và thử nghiệm vận hành.',
             f'Mô hình robot {title} hoàn thiện, hoạt động chính xác.',
             f'{buoc1}:\n- GV giao nhiệm vụ cho từng nhóm lắp ráp mô hình {title}.\n{buoc2}:\n- HS phân công nhiệm vụ (xem sơ đồ, lắp khung, chọn chốt).\n- HS tiến hành lắp ráp cẩn thận từng chi tiết.\n{buoc3}:\n- GV yêu cầu các nhóm cấp nguồn/nạp code và cho robot vận hành thử.\n- HS đặt robot lên bàn thử nghiệm, quan sát hoạt động.\n{buoc4}:\n- GV đánh giá sản phẩm thực hành, chấm điểm tiêu chí kỹ thuật.',
             f'Mô hình robot hoàn thiện.\nRobot hoạt động đúng yêu cầu.'),

            ('4. Hoạt động 4. Vận dụng (8 phút)',
             'Vận dụng sáng tạo, nâng cấp tính năng robot và dọn dẹp.',
             'GV đặt yêu cầu cải tiến tối ưu hoặc thêm chi tiết trang trí.',
             'Báo cáo ý tưởng cải tiến mô hình và bộ Kit được sắp xếp ngăn nắp.',
             f'{buoc1}:\n- GV đặt câu hỏi: "Em có thể cải tiến cấu trúc nào để {title} hoạt động tốt hơn?"\n{buoc2}:\n- HS suy nghĩ ý tưởng sáng tạo mở rộng.\n- HS nhẹ nhàng tháo chốt, phân loại linh kiện về đúng vị trí.\n{buoc3}:\n- 1-2 HS trình bày ý tưởng nâng cấp robot.\n{buoc4_last}:\n- GV dặn dò nhiệm vụ chuẩn bị cho bài học tiếp theo.\n- HS thu dọn bàn học, nộp lại hộp Kit Robotics.',
             f'HS đề xuất ý tưởng cải tiến.\nBộ Kit được sắp xếp ngăn nắp.'),
        ]


# ============================================================
# MAIN EXECUTION
# ============================================================
def process_mon(mon_hoc, ppct_path, out_base_dir, schedule, grade_configs):
    """Process one subject (Tin học or Robotics)."""
    print(f"\n{'='*60}")
    print(f" ĐANG TẠO KHBD {mon_hoc.upper()}")
    print(f"{'='*60}")

    # Clean output dir
    if os.path.exists(out_base_dir):
        shutil.rmtree(out_base_dir)
        print(f"  [+] Đã xóa sạch thư mục {out_base_dir}")
    os.makedirs(out_base_dir, exist_ok=True)

    # Parse PPCT
    doc_ppct = Document(ppct_path)
    tables = doc_ppct.tables
    created = 0

    for folder_name, grade_label, grade_num, t_idx, is_rotation in grade_configs:
        print(f"\n---> Đang xử lý {grade_label} (Table {t_idx})...")

        # Get schedule info
        grade_key = folder_name.replace('Lớp_', '').replace('Tiền_tiểu_học', 'TTH')
        if grade_key not in schedule:
            print(f"  [!] Không tìm thấy lịch dạy cho {grade_key}, bỏ qua.")
            continue

        day_of_week, ten_lop = schedule[grade_key]
        kit_name = get_kit_name(grade_num) if mon_hoc == "Robotics" else None

        t = tables[t_idx]
        for row_idx, row in enumerate(t.rows[1:], start=1):
            cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
            title = cells[1] if len(cells) > 1 else ''
            tiet_ppct_str = cells[3] if len(cells) > 3 else str(row_idx)
            yccd = cells[4] if len(cells) > 4 else title

            if not title or title.startswith('Chủ đề'):
                continue

            try:
                tiet_ppct = int(tiet_ppct_str)
            except ValueError:
                tiet_ppct = row_idx

            # Tính Tuần
            if is_rotation:
                if mon_hoc == "Tin học":
                    tuan_so = 1 if tiet_ppct == 0 else 2 * ((tiet_ppct - 1) // 2 + 1)
                else:  # Robotics rotation
                    tuan_so = 1 if tiet_ppct == 0 else 2 * ((tiet_ppct - 1) // 2) + 3
            else:
                tuan_so = 1 if tiet_ppct == 0 else tiet_ppct + 1

            safe_title = sanitize(title)
            tuan_folder = f"Tuần_{tuan_so:02d}"
            out_dir = os.path.join(out_base_dir, folder_name, tuan_folder)
            os.makedirs(out_dir, exist_ok=True)

            prefix = "Tin_hoc" if mon_hoc == "Tin học" else "Robotics"
            filename = f"KHBD_{prefix}_{folder_name}_Tiet{tiet_ppct:02d}_{safe_title}.docx"
            out_file = os.path.join(out_dir, filename)

            # Build document
            if grade_num >= 6:
                doc = build_khbd_thcs(mon_hoc, grade_num, grade_label, ten_lop,
                                       title, tiet_ppct, yccd, tuan_so, day_of_week, kit_name)
            else:
                doc = build_khbd_th(mon_hoc, grade_num, grade_label, ten_lop,
                                     title, tiet_ppct, yccd, tuan_so, day_of_week, kit_name)

            try:
                doc.save(out_file)
                created += 1
                print(f"  [+] {folder_name} -> {tuan_folder} -> {filename}")
            except PermissionError:
                alt = out_file.replace('.docx', '_new.docx')
                doc.save(alt)
                created += 1
                print(f"  [!] File bị khóa, đã lưu: {alt}")
            except Exception as e:
                print(f"  [!] Lỗi: {e}")

    return created


def main():
    print("=" * 60)
    print(" TÁI TẠO TOÀN BỘ KHBD TIN HỌC & ROBOTICS (V2 - CHUẨN LUẬT)")
    print("=" * 60)

    # ---- TIN HỌC ----
    tin_configs = [
        # (folder, label, grade_num, table_index, is_rotation)
        ("Tiền_tiểu_học", "Tiền tiểu học", 0, 3, False),
        ("Lớp_1", "Lớp 1", 1, 4, False),
        ("Lớp_2", "Lớp 2", 2, 5, False),
        ("Lớp_3", "Lớp 3", 3, 6, False),
        ("Lớp_4", "Lớp 4", 4, 7, False),
        ("Lớp_5", "Lớp 5", 5, 8, True),
        ("Lớp_6", "Lớp 6", 6, 9, True),
        ("Lớp_7", "Lớp 7", 7, 10, True),
        ("Lớp_8", "Lớp 8", 8, 11, True),
    ]
    tin_count = process_mon("Tin học", PPCT_TIN, OUT_TIN, TIN_SCHEDULE, tin_configs)

    # ---- ROBOTICS ----
    rob_configs = [
        ("Lớp_1", "Lớp 1", 1, 4, False),
        ("Lớp_2", "Lớp 2", 2, 5, False),
        ("Lớp_3", "Lớp 3", 3, 6, False),
        ("Lớp_4", "Lớp 4", 4, 7, False),
        ("Lớp_5", "Lớp 5", 5, 8, True),
        ("Lớp_6", "Lớp 6", 6, 9, True),
        ("Lớp_7", "Lớp 7", 7, 10, True),
        ("Lớp_8", "Lớp 8", 8, 11, True),
    ]
    rob_count = process_mon("Robotics", PPCT_ROB, OUT_ROB, ROB_SCHEDULE, rob_configs)

    print(f"\n{'='*60}")
    print(f" HOÀN THÀNH! Tin học: {tin_count} files | Robotics: {rob_count} files")
    print(f" Tổng: {tin_count + rob_count} files KHBD đã tạo (chuẩn luật V2)")
    print(f"{'='*60}")


if __name__ == '__main__':
    main()
