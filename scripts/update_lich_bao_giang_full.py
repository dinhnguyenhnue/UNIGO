import os, sys, re, shutil
from datetime import datetime, timedelta
import openpyxl, docx
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

sys.stdout.reconfigure(encoding='utf-8')

def set_font_run(run, font_name="Times New Roman", size_pt=11, bold=None, italic=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
        
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
    rPr.append(rFonts)

def format_paragraph(p, font_name="Times New Roman", size_pt=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=None, italic=None, space_before=1, space_after=1):
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.05
    for run in p.runs:
        set_font_run(run, font_name=font_name, size_pt=size_pt, bold=bold, italic=italic)

def format_cell(cell, font_name="Times New Roman", size_pt=10.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=None, italic=None):
    tcPr = cell._element.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)
    for p in cell.paragraphs:
        format_paragraph(p, font_name=font_name, size_pt=size_pt, align=align, bold=bold, italic=italic, space_before=1, space_after=1)

def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)

def set_table_borders(table):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        if not tblPr[0].xpath('w:tblBorders'):
            borders_xml = parse_xml(
                f'<w:tblBorders {nsdecls("w")}>\n'
                f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
                f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
                f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
                f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
                f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
                f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
                f'</w:tblBorders>'
            )
            tblPr[0].append(borders_xml)

def load_khdh_lessons(subject, grade):
    if subject == 'Tin học':
        path = f'.\\Hệ thống mẫu văn bản\\Nguyên đã làm\\Kế hoạch dạy học Tin học từng lớp\\Kế hoạch dạy học môn Tin học - Lớp {grade} - 2026 - 2027.docx'
    else:
        path = f'.\\Hệ thống mẫu văn bản\\Nguyên đã làm\\Kế hoạch dạy học Robotics từng lớp\\Kế hoạch dạy học môn Robotics - Lớp {grade} - 2026 - 2027.docx'
        
    doc = docx.Document(path)
    lessons = []
    for t in doc.tables:
        if len(t.rows) > 5 and len(t.columns) == 5:
            for r in range(1, len(t.rows)):
                cells = [c.text.strip().replace('\n', ' ') for c in t.rows[r].cells]
                stt, name, so_tiet, ppct, yccd = cells[0], cells[1], cells[2], cells[3], cells[4]
                if stt == '' and ('chủ đề' in name.lower()):
                    continue
                lessons.append((ppct, name))
    return lessons

def parse_schedule_from_excel(xlsx_path):
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb['TKB_LOP_SC']

    classes = []
    for c in range(3, ws.max_column + 1, 2):
        cls_name = ws.cell(row=4, column=c).value
        if cls_name:
            classes.append((str(cls_name).strip(), c, c+1))

    schedule_sang = []
    schedule_chieu = []

    thu_map = {'2': 'Hai', '3': 'Ba', '4': 'Tư', '5': 'Năm', '6': 'Sáu'}

    current_thu = ''
    for r in range(6, ws.max_row + 1):
        thu_val = ws.cell(row=r, column=1).value
        tiet_val = ws.cell(row=r, column=2).value
        if thu_val:
            current_thu = str(thu_val).strip()
        if not tiet_val:
            continue
        tiet_str = str(tiet_val).strip()
        
        thu_name = thu_map.get(current_thu, current_thu)
        
        for cls_name, sang_col, chieu_col in classes:
            sang_val = str(ws.cell(row=r, column=sang_col).value or '').strip()
            chieu_val = str(ws.cell(row=r, column=chieu_col).value or '').strip()
            
            if 'nguyên' in sang_val.lower() or ('tin' in sang_val.lower() and 'nguyên' in sang_val.lower()) or ('rob' in sang_val.lower() and 'nguyên' in sang_val.lower()):
                mon_name = "Robotics" if 'rob' in sang_val.lower() else "Tin học"
                schedule_sang.append((thu_name, int(tiet_str) if tiet_str.isdigit() else tiet_str, mon_name, cls_name))
                
            if 'nguyên' in chieu_val.lower() or ('tin' in chieu_val.lower() and 'nguyên' in chieu_val.lower()) or ('rob' in chieu_val.lower() and 'nguyên' in chieu_val.lower()):
                mon_name = "Robotics" if 'rob' in chieu_val.lower() else "Tin học"
                schedule_chieu.append((thu_name, int(tiet_str) if tiet_str.isdigit() else tiet_str, mon_name, cls_name))

    day_order = {'Hai': 2, 'Ba': 3, 'Tư': 4, 'Năm': 5, 'Sáu': 6}
    schedule_sang.sort(key=lambda x: (day_order.get(x[0], 9), x[1]))
    schedule_chieu.sort(key=lambda x: (day_order.get(x[0], 9), x[1]))

    return schedule_sang, schedule_chieu

def extract_grade(lop_name):
    nums = re.findall(r'\d+', lop_name)
    if nums:
        return nums[0]
    return '1' # default for TTH or special

def main():
    xlsx_path = r"D:\UNIGO\TKB toàn trường CHECK - lần 2.1.xlsx"
    target_docx = r"D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Lịch báo giảng\Lịch báo giảng cả năm (35 tuần).docx"
    sync_docx = r"D:\UNIGO\Hệ thống mẫu văn bản\Lịch báo giảng.docx"
    
    schedule_sang, schedule_chieu = parse_schedule_from_excel(xlsx_path)
    
    # Load all KHDH lesson databases
    khdh_db = {}
    for sub in ['Tin học', 'Robotics']:
        for g in range(1, 9):
            khdh_db[(sub, str(g))] = load_khdh_lessons(sub, g)
            
    # Track current lesson index for each (cls_name, subject, session_slot)
    lesson_counters = {}
    
    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = Cm(1.5)
        s.bottom_margin = Cm(1.5)
        s.left_margin = Cm(1.5)
        s.right_margin = Cm(1.5)
        
    col_widths_cm = [1.8, 1.2, 1.5, 2.2, 1.3, 5.8, 2.7]
    hdr_titles = ['Thứ/ ngày', 'Tiết TKB', 'Tiết PPCT', 'Môn (phân môn)', 'Lớp', 'Tên bài', 'Đồ dùng']
    
    # Start date of Week 1: 03/08/2026 (Monday of week 1)
    base_monday = datetime(2026, 8, 3)
    
    for week_num in range(1, 36):
        w_monday = base_monday + timedelta(weeks=week_num - 1)
        w_friday = w_monday + timedelta(days=4)
        
        m_str = w_monday.strftime('%d/%m/%Y')
        f_str = w_friday.strftime('%d/%m/%Y')
        week_label = f"TUẦN {week_num:02d} (từ ngày {m_str} đến ngày {f_str})"
        
        # ---------------------------------------------------------
        # MORNING SESSION (PAGE 1 of Week)
        # ---------------------------------------------------------
        p_title = doc.add_paragraph(week_label)
        format_paragraph(p_title, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=0, space_after=2)
        
        p_header_info = doc.add_paragraph("Họ và tên giáo viên: Nguyễn                  Môn giảng dạy: Tin học & Robotics")
        format_paragraph(p_header_info, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4)
        
        p_sang_head = doc.add_paragraph(f"Buổi Sáng - Tuần {week_num:02d} (Từ ngày {m_str} đến ngày {f_str})")
        format_paragraph(p_sang_head, bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=2, space_after=4)
        
        t_sang = doc.add_table(rows=0, cols=7)
        t_sang.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t_sang)
        
        hdr_row = t_sang.add_row()
        for idx, title in enumerate(hdr_titles):
            cell = hdr_row.cells[idx]
            cell.text = title
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in [0, 1, 2, 4] else WD_ALIGN_PARAGRAPH.LEFT
            format_cell(cell, font_name="Times New Roman", size_pt=11, align=align, bold=True)
            set_cell_width(cell, col_widths_cm[idx])
            
        prev_day = ""
        for thu, tiet, mon, lop in schedule_sang:
            row = t_sang.add_row()
            day_str = thu if thu != prev_day else ""
            prev_day = thu
            
            # Check if this slot gets a lesson
            is_m1_blank = (week_num == 1 and thu == 'Hai')
            
            if is_m1_blank:
                ppct_val = ""
                ten_bai_val = ""
                dodung_val = ""
            else:
                grade = extract_grade(lop)
                key = (mon, grade)
                c_key = (mon, lop, 'Sáng', thu, tiet)
                curr_idx = lesson_counters.get(c_key, 0)
                
                lessons = khdh_db.get(key, [])
                if lessons and curr_idx < len(lessons):
                    ppct_val, ten_bai_val = lessons[curr_idx]
                    lesson_counters[c_key] = curr_idx + 1
                else:
                    # Fallback / wrap around / review
                    ppct_val = str(curr_idx + 1)
                    ten_bai_val = f"Ôn tập & Thực hành mở rộng {mon} {lop}"
                    lesson_counters[c_key] = curr_idx + 1
                    
                dodung_val = "Máy tính, Tivi" if mon == "Tin học" else "Bộ kit Robotics, Máy tính"
                
            row.cells[0].text = day_str
            format_cell(row.cells[0], font_name="Times New Roman", size_pt=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True if day_str else False)
            set_cell_width(row.cells[0], col_widths_cm[0])
            
            row.cells[1].text = str(tiet)
            format_cell(row.cells[1], font_name="Times New Roman", size_pt=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_width(row.cells[1], col_widths_cm[1])
            
            row.cells[2].text = ppct_val
            format_cell(row.cells[2], font_name="Times New Roman", size_pt=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_width(row.cells[2], col_widths_cm[2])
            
            row.cells[3].text = mon
            format_cell(row.cells[3], font_name="Times New Roman", size_pt=10.5, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_width(row.cells[3], col_widths_cm[3])
            
            row.cells[4].text = lop
            format_cell(row.cells[4], font_name="Times New Roman", size_pt=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_width(row.cells[4], col_widths_cm[4])
            
            row.cells[5].text = ten_bai_val
            format_cell(row.cells[5], font_name="Times New Roman", size_pt=10.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            set_cell_width(row.cells[5], col_widths_cm[5])
            
            row.cells[6].text = dodung_val
            format_cell(row.cells[6], font_name="Times New Roman", size_pt=10.5, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_width(row.cells[6], col_widths_cm[6])

        # Sign-off box for Morning
        p_space1 = doc.add_paragraph()
        format_paragraph(p_space1, space_before=1, space_after=1)
        
        t_sign_sang = doc.add_table(rows=1, cols=2)
        t_sign_sang.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t_sign_sang)
        c0 = t_sign_sang.rows[0].cells[0]
        c1 = t_sign_sang.rows[0].cells[1]
        c0.text = "Kiểm tra, nhận xét:\n……………………………………………………………………………………\n……………………………………………………………………………………"
        c1.text = f"Hà Nội, ngày {f_str[:2]} tháng {f_str[3:5]} năm 2026\nBan giám hiệu (Ký và đóng dấu)"
        format_cell(c0, font_name="Times New Roman", size_pt=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        format_cell(c1, font_name="Times New Roman", size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        set_cell_width(c0, Cm(9.5))
        set_cell_width(c1, Cm(7.0))

        # ---------------------------------------------------------
        # PAGE BREAK TO AFTERNOON (PAGE 2 of Week)
        # ---------------------------------------------------------
        doc.add_page_break()

        p_chieu_head = doc.add_paragraph(f"Buổi Chiều - Tuần {week_num:02d} (Từ ngày {m_str} đến ngày {f_str})")
        format_paragraph(p_chieu_head, bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=2, space_after=4)
        
        t_chieu = doc.add_table(rows=0, cols=7)
        t_chieu.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t_chieu)
        
        hdr_row_c = t_chieu.add_row()
        for idx, title in enumerate(hdr_titles):
            cell = hdr_row_c.cells[idx]
            cell.text = title
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in [0, 1, 2, 4] else WD_ALIGN_PARAGRAPH.LEFT
            format_cell(cell, font_name="Times New Roman", size_pt=11, align=align, bold=True)
            set_cell_width(cell, col_widths_cm[idx])
            
        prev_day = ""
        for thu, tiet, mon, lop in schedule_chieu:
            row = t_chieu.add_row()
            day_str = thu if thu != prev_day else ""
            prev_day = thu
            
            is_m1_blank = (week_num == 1 and thu == 'Hai')
            
            if is_m1_blank:
                ppct_val = ""
                ten_bai_val = ""
                dodung_val = ""
            else:
                grade = extract_grade(lop)
                key = (mon, grade)
                c_key = (mon, lop, 'Chiều', thu, tiet)
                curr_idx = lesson_counters.get(c_key, 0)
                
                lessons = khdh_db.get(key, [])
                if lessons and curr_idx < len(lessons):
                    ppct_val, ten_bai_val = lessons[curr_idx]
                    lesson_counters[c_key] = curr_idx + 1
                else:
                    ppct_val = str(curr_idx + 1)
                    ten_bai_val = f"Ôn tập & Thực hành mở rộng {mon} {lop}"
                    lesson_counters[c_key] = curr_idx + 1
                    
                dodung_val = "Máy tính, Tivi" if mon == "Tin học" else "Bộ kit Robotics, Máy tính"
                
            row.cells[0].text = day_str
            format_cell(row.cells[0], font_name="Times New Roman", size_pt=10.5, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True if day_str else False)
            set_cell_width(row.cells[0], col_widths_cm[0])
            
            row.cells[1].text = str(tiet)
            format_cell(row.cells[1], font_name="Times New Roman", size_pt=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_width(row.cells[1], col_widths_cm[1])
            
            row.cells[2].text = ppct_val
            format_cell(row.cells[2], font_name="Times New Roman", size_pt=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_width(row.cells[2], col_widths_cm[2])
            
            row.cells[3].text = mon
            format_cell(row.cells[3], font_name="Times New Roman", size_pt=10.5, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_width(row.cells[3], col_widths_cm[3])
            
            row.cells[4].text = lop
            format_cell(row.cells[4], font_name="Times New Roman", size_pt=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_width(row.cells[4], col_widths_cm[4])
            
            row.cells[5].text = ten_bai_val
            format_cell(row.cells[5], font_name="Times New Roman", size_pt=10.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            set_cell_width(row.cells[5], col_widths_cm[5])
            
            row.cells[6].text = dodung_val
            format_cell(row.cells[6], font_name="Times New Roman", size_pt=10.5, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_width(row.cells[6], col_widths_cm[6])

        # Sign-off box for Afternoon
        p_space2 = doc.add_paragraph()
        format_paragraph(p_space2, space_before=1, space_after=1)
        
        t_sign_chieu = doc.add_table(rows=1, cols=2)
        t_sign_chieu.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(t_sign_chieu)
        c0_c = t_sign_chieu.rows[0].cells[0]
        c1_c = t_sign_chieu.rows[0].cells[1]
        c0_c.text = "Kiểm tra, nhận xét:\n……………………………………………………………………………………\n……………………………………………………………………………………"
        c1_c.text = f"Hà Nội, ngày {f_str[:2]} tháng {f_str[3:5]} năm 2026\nBan giám hiệu (Ký và đóng dấu)"
        format_cell(c0_c, font_name="Times New Roman", size_pt=10, align=WD_ALIGN_PARAGRAPH.LEFT)
        format_cell(c1_c, font_name="Times New Roman", size_pt=10, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        set_cell_width(c0_c, Cm(9.5))
        set_cell_width(c1_c, Cm(7.0))
        
        # Add page break after each week except week 35
        if week_num < 35:
            doc.add_page_break()

    print("Generated all 35 weeks! Saving document...")
    
    os.makedirs(os.path.dirname(target_docx), exist_ok=True)
    doc.save(target_docx)
    print(f"Saved: {target_docx}")
    
    try:
        shutil.copy2(target_docx, sync_docx)
        print(f"Synced to: {sync_docx}")
    except Exception as e:
        print(f"Sync skipped: {e}")

if __name__ == "__main__":
    main()
