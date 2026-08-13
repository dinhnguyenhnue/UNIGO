# -*- coding: utf-8 -*-
"""
Script tạo Bộ Kế hoạch bài dạy (KHBD) Tuần 03 cho TOÀN BỘ 9 KHỐI LỚP
(Tiền Tiểu Học, Lớp 1, 2, 3, 4, 5, 6, 7, 8)
Tuân thủ Phụ lục IV (CV 5512), Thông tư 02/2025 (CV 3456 về Năng lực số), và quy định UNIGO.
"""
import os
import re
import sys
import glob
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')

OUT_BASE_DIR = r'd:\UNIGO\KHBD_Tin_học'

def set_table_borders(table, color="000000", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def apply_font(run, font_name="Times New Roman", size_pt=13, bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color

def format_paragraph(p, font_name="Times New Roman", size_pt=13, line_spacing=1.15, space_after=3, bold=False, italic=False):
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)
    for run in p.runs:
        apply_font(run, font_name=font_name, size_pt=size_pt, bold=bold, italic=italic)

def sanitize_filename(name):
    import unicodedata
    name = unicodedata.normalize('NFD', name).encode('ascii', 'ignore').decode('utf-8')
    name = re.sub(r'[^a-zA-Z0-9_\-]', '_', name)
    name = re.sub(r'_+', '_', name).strip('_')
    return name

def populate_primary_doc_tuan03(template_path, grade_str, lesson_title, tiet_ppct, yccd, khbd_detail):
    doc = docx.Document(template_path)
    
    # Safely update header text runs without affecting drawing Run 0
    for s in doc.sections:
        for hp in s.header.paragraphs:
            if len(hp.runs) >= 6:
                hp.runs[2].text = "Đậu Đình Nguyên"
                hp.runs[5].text = f"{grade_str} "

    # Paragraph 0: Week & Dates
    if len(doc.paragraphs) > 0:
        p0 = doc.paragraphs[0]
        p0.text = "TUẦN: 03\t\t\t\tNgày soạn: 10/09/2026\n\t\t\t\t\tNgày dạy: 17/09/2026"
        format_paragraph(p0, size_pt=12, italic=True)

    if len(doc.paragraphs) > 2:
        p2 = doc.paragraphs[2]
        p2.text = f"KẾ HOẠCH DẠY HỌC MÔN TIN HỌC - {grade_str.upper()}"
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_paragraph(p2, size_pt=14, bold=True)

    if len(doc.paragraphs) > 3:
        p3 = doc.paragraphs[3]
        p3.text = f"CHỦ ĐỀ: THẾ GIỚI CÔNG NGHỆ & TƯ DUY SỐ"
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_paragraph(p3, size_pt=13, bold=True)

    if len(doc.paragraphs) > 4:
        p4 = doc.paragraphs[4]
        p4.text = f"BÀI: {lesson_title.upper()} (Thời lượng: 1 tiết | Tiết PPCT: {tiet_ppct})"
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        format_paragraph(p4, size_pt=13, bold=True)

    # Search & update YÊU CẦU CẦN ĐẠT
    for i, p in enumerate(doc.paragraphs):
        if "- Sau bài học này" in p.text or "- Sau tiết học" in p.text:
            p.text = f"- Sau bài học này, học sinh sẽ: {yccd}"
            format_paragraph(p, italic=True)
            break

    # Search & update Năng lực môn học
    for i, p in enumerate(doc.paragraphs):
        if "2.1. Năng lực môn học" in p.text or "NLa" in p.text:
            p.text = (
                f"- 2.1. Năng lực môn học (Tin học):\n"
                f"  + NLa (Nhận biết & Khám phá): Nhận diện các thiết bị công nghệ, kỹ năng và quy tắc theo bài {lesson_title} (Đạt được thông qua HĐ 1, HĐ 2).\n"
                f"  + NLb (Sử dụng & Quản lý): Thao tác thành thạo bàn phím, chuột hoặc phần mềm theo đúng quy trình hướng dẫn (Đạt được thông qua HĐ 3).\n"
                f"- 2.2. Năng lực chung:\n"
                f"  + Tự chủ và tự học: Tự giác theo dõi thao tác mẫu của giáo viên, tự mình hoàn thành bài tập thực hành trên máy tính (Đạt được thông qua HĐ 2, HĐ 3).\n"
                f"  + Giao tiếp và hợp tác: Lắng nghe ý kiến của bạn trong nhóm đôi/nhóm lớn, phối hợp cùng hoàn thành nhiệm vụ chung (Đạt được thông qua HĐ 2, HĐ 4).\n"
                f"- 2.3. Năng lực số:\n"
                f"  + Khai thác thông tin số: Nhận biết và thao tác với thông tin hiển thị trên màn hình máy tính (Đạt được thông qua HĐ 2).\n"
                f"  + An toàn & Văn hóa số: Ý thức tuân thủ quy tắc an toàn về điện, bảo vệ mắt và giữ vệ sinh thiết bị công nghệ (Đạt được thông qua HĐ 1, HĐ 4)."
            )
            format_paragraph(p)
            break

    # Table 0: Primary 2-column activities
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        set_table_borders(t0)
        t0.alignment = WD_TABLE_ALIGNMENT.CENTER

        while len(t0.rows) < 5:
            t0.add_row()

        hdr = t0.rows[0].cells
        hdr[0].text = "HOẠT ĐỘNG CỦA GIÁO VIÊN"
        hdr[1].text = "HOẠT ĐỘNG CỦA HỌC SINH"
        for cell in hdr:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    apply_font(r, size_pt=13, bold=True)

        for row_i in range(1, 5):
            act_k = f"hd{row_i}"
            gv_txt = khbd_detail[act_k]["gv"]
            hs_txt = khbd_detail[act_k]["hs"]
            
            c = t0.rows[row_i].cells
            c[0].text = gv_txt
            c[1].text = hs_txt

        for row in t0.rows[1:]:
            for cell in row.cells:
                set_cell_margins(cell)
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_after = Pt(2)
                    for run in p.runs:
                        apply_font(run, size_pt=12)

    return doc

def populate_secondary_doc_tuan03(template_path, grade_str, lesson_title, tiet_ppct, yccd, khbd_detail):
    doc = docx.Document(template_path)
    
    # Table 0: Header info table
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        set_table_borders(t0)
        t0.rows[0].cells[0].text = "Trường: TH & THCS UNIGO\nTổ: Khoa học Tự nhiên & Công nghệ"
        t0.rows[0].cells[1].text = "CỘNG HOÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc"
        t0.rows[1].cells[0].text = "Họ và tên GV: Đậu Đình Nguyên"
        t0.rows[1].cells[1].text = f"Ngày soạn: 10/09/2026 | Ngày dạy: 17/09/2026\nLớp: {grade_str}"
        for row in t0.rows:
            for cell in row.cells:
                set_cell_margins(cell)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for r in p.runs:
                        apply_font(r, size_pt=12, bold=True)

    for p in doc.paragraphs:
        if "TÊN BÀI DẠY" in p.text:
            p.text = f"TÊN BÀI DẠY: {lesson_title.upper()}\n"
            r1 = p.add_run(f"MÔN HỌC: TIN HỌC - {grade_str.upper()}\n")
            apply_font(r1, size_pt=13, bold=True)
            r2 = p.add_run(f"Thời lượng: 1 tiết | Tiết theo PPCT: {tiet_ppct}")
            apply_font(r2, size_pt=12, italic=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            break

    for p in doc.paragraphs:
        if "1. Kiến thức:" in p.text:
            p.text = f"1. Kiến thức: {yccd}. Nắm vững các khái niệm, quy trình và kỹ năng sử dụng công nghệ trong bài {lesson_title} (Đạt được thông qua HĐ 1, HĐ 2)."
            format_paragraph(p)
        elif "2. Năng lực:" in p.text or "- Năng lực chung:" in p.text:
            p.text = (
                f"2. Năng lực:\n"
                f"  - Năng lực đặc thù (Tin học): Nhận biết, phân tích và thao tác thành thạo các phần mềm/công cụ trong bài {lesson_title} (Đạt được thông qua HĐ 2, HĐ 3).\n"
                f"  - Năng lực số: Khai thác thông tin số an toàn, bảo mật dữ liệu cá nhân, tuân thủ pháp luật và văn hóa ứng xử trong môi trường số (Đạt được thông qua HĐ 2, HĐ 4).\n"
                f"  - Năng lực chung: Tự chủ tự học nghiên cứu tài liệu; giao tiếp hợp tác làm việc nhóm; giải quyết vấn đề kỹ thuật số sáng tạo (Đạt được thông qua HĐ 1, HĐ 3)."
            )
            format_paragraph(p)
        elif "3. Phẩm chất:" in p.text:
            p.text = "3. Phẩm chất: Tác phong công nghiệp, trung thực trong học tập, có ý thức trách nhiệm bảo vệ thiết bị số và tài nguyên mạng (Đạt được thông qua HĐ 1, HĐ 4)."
            format_paragraph(p)

    if len(doc.tables) > 1:
        t1 = doc.tables[1]
        set_table_borders(t1)
        while len(t1.rows) < 5:
            t1.add_row()

        hdr = t1.rows[0].cells
        hdr[0].text = "HOẠT ĐỘNG CỦA GV VÀ HỌC SINH"
        hdr[1].text = "KẾT QUẢ CẦN ĐẠT / SẢN PHẨM"
        for cell in hdr:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    apply_font(r, size_pt=13, bold=True)

        for row_i in range(1, 5):
            act_k = f"hd{row_i}"
            c = t1.rows[row_i].cells
            c[0].text = khbd_detail[act_k]["gv"]
            c[1].text = khbd_detail[act_k]["sp"]

        for row in t1.rows[1:]:
            for cell in row.cells:
                set_cell_margins(cell)
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_after = Pt(2)
                    for run in p.runs:
                        apply_font(run, size_pt=12)

    return doc

# ─── Dữ liệu chi tiết 9 bài học Tuần 03 ───
LESSONS_KHBD_TUAN03 = [
    {
        "folder": "Tiền_tiểu_học", "grade_str": "Tiền tiểu học", "is_primary": True,
        "title": "Bài 2. Em ngồi máy tính an toàn", "tiet_ppct": 2,
        "yccd": "Học sinh nhận biết tư thế ngồi đúng khi sử dụng máy tính (lưng thẳng, mắt cách màn hình 50-70cm), ý thức giữ an toàn điện và bảo vệ mắt khi học bài.",
        "detail": {
            "hd1": {"gv": "1. Hoạt động MỞ ĐẦU (Khởi động) (5 phút)\na) Chuyển giao: GV cho HS xem 2 bức tranh (1 bạn ngồi còng lưng, 1 bạn ngồi thẳng lưng).\nb) Thực hiện: Hỏi 'Bạn nào ngồi đẹp và đúng cách hả các con?'\nc) Báo cáo: Mời 2-3 HS giơ tay chỉ vào bức tranh đúng.\nd) Kết luận: GV ngợi khen và giới thiệu bài học tư thế ngồi máy tính an toàn.",
                    "hs": "1. Hoạt động MỞ ĐẦU (Khởi động) (5 phút)\na) Tiếp nhận: Quan sát hai bức tranh trên màn chiếu.\nb) Thực hiện: Suy nghĩ và trao đổi nhỏ với bạn bên cạnh.\nc) Báo cáo: Hăng hái giơ tay trả lời.\nd) Kết luận: Lắng nghe GV dặn dò."},
            "hd2": {"gv": "2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI (12 phút)\na) Chuyển giao: GV hướng dẫn 3 quy tắc: Lưng thẳng chạm tựa ghế, hai chân chạm sàn, mắt cách màn hình một cánh tay (50-70cm).\nb) Thực hiện: GV đi từng bàn làm mẫu cho từng HS.\nc) Báo cáo: Mời 1 HS làm mẫu tư thế chuẩn trước lớp.\nd) Kết luận: Chốt quy tắc ngồi thẳng - mắt xa - chân chạm sàn.",
                    "hs": "2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI (12 phút)\na) Tiếp nhận: Quan sát cô làm mẫu tư thế ngồi chuẩn.\nb) Thực hiện: Tự điều chỉnh tư thế ngồi của bản thân.\nc) Báo cáo: HS đại diện lên làm mẫu tư thế chuẩn.\nd) Kết luận: Nhớ 3 quy tắc ngồi an toàn."},
            "hd3": {"gv": "3. HOẠT ĐỘNG LUYỆN TẬP - THỰC HÀNH (15 phút)\na) Chuyển giao: Cho HS chơi trò chơi 'Tạo dáng robot ngồi máy tính'.\nb) Thực hiện: Khi hô 'Robot sẵn sàng!', HS lập tức ngồi đúng tư thế chuẩn.\nc) Báo cáo: GV đi kiểm tra và khen thưởng các bạn làm đúng.\nd) Kết luận: Nhận xét tuyên dương cả lớp.",
                    "hs": "3. HOẠT ĐỘNG LUYỆN TẬP - THỰC HÀNH (15 phút)\na) Tiếp nhận: Lắng nghe luật chơi từ cô giáo.\nb) Thực hiện: Nhanh chóng điều chỉnh tư thế ngồi thẳng ngay khi nghe hiệu lệnh.\nc) Báo cáo: Sửa tư thế cho bạn bên cạnh.\nd) Kết luận: Hào hứng tham gia trò chơi."},
            "hd4": {"gv": "4. HOẠT ĐỘNG VẬN DỤNG (3 phút)\na) Chuyển giao: Dặn dò HS không chạm tay vào ổ điện và giữ khoảng cách mắt khi xem máy tính ở nhà.\nb) Thực hiện: Hướng dẫn HS xếp ghế gọn gàng.\nc) Báo cáo: HS hứa thực hiện tư thế đúng ở nhà.\nd) Kết luận: Kết thúc tiết học vui vẻ.",
                    "hs": "4. HOẠT ĐỘNG VẬN DỤNG (3 phút)\na) Tiếp nhận: Ghi nhớ lời cô dặn.\nb) Thực hiện: Đẩy ghế gọn gàng vào dưới bàn.\nc) Báo cáo: Chào cô giáo trước khi rời phòng máy.\nd) Kết luận: Ra về an toàn."}
        }
    },
    {
        "folder": "Lớp_1", "grade_str": "Lớp 1", "is_primary": True,
        "title": "Bài 2. Ôn và nâng cấp kỹ năng chuột", "tiet_ppct": 2,
        "yccd": "Học sinh thực hiện thành thạo 5 thao tác chuột cơ bản (di chuyển, nháy trái, nháy phải, nháy đôi, kéo thả), áp dụng thao tác chuột trong các trò chơi luyện tập đơn giản.",
        "detail": {
            "hd1": {"gv": "1. Hoạt động MỞ ĐẦU (Khởi động) (5 phút)\na) Chuyển giao: GV tổ chức trò chơi 'Ngón tay nháy chuột'. Hô lệnh nháy trái/nháy phải/nháy đôi.\nb) Thực hiện: HS thực hiện hành động bằng ngón tay.\nc) Báo cáo: Kiểm tra phản xạ của HS.\nd) Kết luận: Dẫn dắt vào bài luyện tập kỹ năng chuột nâng cao.",
                    "hs": "1. Hoạt động MỞ ĐẦU (Khởi động) (5 phút)\na) Tiếp nhận: Lắng nghe luật chơi từ thầy cô.\nb) Thực hiện: Co ngón trỏ/ngón giữa theo hiệu lệnh.\nc) Báo cáo: Cùng tham gia sôi nổi.\nd) Kết luận: Sẵn sàng thực hành chuột trên máy."},
            "hd2": {"gv": "2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI (12 phút)\na) Chuyển giao: GV hướng dẫn kỹ thuật 'Kéo thả chuột' (Drag and Drop): Đè giữ nút trái + Di chuyển + Thả ngón tay.\nb) Thực hiện: GV làm mẫu trên máy chiếu với việc di chuyển biểu tượng.\nc) Báo cáo: Gọi 1 HS lên bảng kéo thả một thư mục.\nd) Kết luận: Chốt 3 bước kéo thả chuột chính xác.",
                    "hs": "2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI (12 phút)\na) Tiếp nhận: Quan sát các bước kéo thả trên màn chiếu.\nb) Thực hiện: Đặt tay đúng vị trí trên chuột.\nc) Báo cáo: 1 HS làm mẫu thao tác kéo thả.\nd) Kết luận: Nắm vững quy trình 3 bước kéo thả."},
            "hd3": {"gv": "3. HOẠT ĐỘNG LUYỆN TẬP - THỰC HÀNH (15 phút)\na) Chuyển giao: Mở phần mềm luyện chuột Sebran / Tux Paint. Giao bài tập kéo hình ghép tranh.\nb) Thực hiện: GV di chuyển hỗ trợ các HS cầm chuột chưa vững.\nc) Báo cáo: Cho HS kiểm tra chéo kết quả hoàn thành bức tranh.\nd) Kết luận: Tuyên dương các em ghép hình nhanh và chính xác.",
                    "hs": "3. HOẠT ĐỘNG LUYỆN TẬP - THỰC HÀNH (15 phút)\na) Tiếp nhận: Mở phần mềm luyện tập trên máy tính.\nb) Thực hiện: Tự giác sử dụng thao tác kéo thả hoàn thành trò chơi.\nc) Báo cáo: Khoe bức tranh đã ghép hoàn chỉnh với bạn.\nd) Kết luận: Thành thạo thao tác chuột."},
            "hd4": {"gv": "4. HOẠT ĐỘNG VẬN DỤNG (3 phút)\na) Chuyển giao: Hỏi HS: Kéo thả chuột giúp em làm được những việc gì trên máy tính?\nb) Thực hiện: Hướng dẫn thoát phần mềm và đặt chuột gọn gàng.\nc) Báo cáo: HS trả lời câu hỏi.\nd) Kết luận: Nhận xét chung tiết học.",
                    "hs": "4. HOẠT ĐỘNG VẬN DỤNG (3 phút)\na) Tiếp nhận: Suy nghĩ ứng dụng của kéo thả chuột.\nb) Thực hiện: Đóng ứng dụng, cất chuột ngay ngắn.\nc) Báo cáo: Nêu ứng dụng kéo thả di chuyển tệp.\nd) Kết luận: Kết thúc buổi học."}
        }
    },
    {
        "folder": "Lớp_2", "grade_str": "Lớp 2", "is_primary": True,
        "title": "Bài 2. Ôn và nâng cấp kỹ năng chuột", "tiet_ppct": 2,
        "yccd": "Học sinh nâng cao kỹ năng nháy đôi chuột mở ứng dụng, thực hành cuộn con lăn và kéo chuột chọn nhiều đối tượng cùng lúc.",
        "detail": {
            "hd1": {"gv": "1. Hoạt động MỞ ĐẦU (Khởi động) (5 phút)\na) Chuyển giao: GV đố HS: Muốn mở ngay một thư mục trên máy tính, ta nháy chuột như thế nào?\nb) Thực hiện: Gợi ý nháy 1 lần hay 2 lần?\nc) Báo cáo: HS trả lời 'Nháy đôi chuột!'.\nd) Kết luận: Giới thiệu bài học nâng cấp thao tác nháy đôi và cuộn chuột.",
                    "hs": "1. Hoạt động MỞ ĐẦU (Khởi động) (5 phút)\na) Tiếp nhận: Lắng nghe câu hỏi gợi mở.\nb) Thực hiện: Suy nghĩ từ trải nghiệm tiết trước.\nc) Báo cáo: Giơ tay trả lời.\nd) Kết luận: Sẵn sàng bài học mới."},
            "hd2": {"gv": "2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI (12 phút)\na) Chuyển giao: GV hướng dẫn kỹ năng cuộn chuột (Scroll) để xem trang dài và kỹ thuật kéo khung chọn nhiều đối tượng.\nb) Thực hiện: Làm mẫu cuộn trang web và quét chuột chọn 5 biểu tượng.\nc) Báo cáo: Mời HS lên thao tác thử cuộn trang sách điện tử.\nd) Kết luận: Chốt công dụng của con lăn cuộn chuột.",
                    "hs": "2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI (12 phút)\na) Tiếp nhận: Quan sát thầy cô cuộn chuột và quét chọn nhiều tệp.\nb) Thực hiện: Lăn thử con lăn chuột bằng ngón trỏ.\nc) Báo cáo: HS thực hiện thao tác trước lớp.\nd) Kết luận: Ghi nhớ thao tác cuộn trang."},
            "hd3": {"gv": "3. HOẠT ĐỘNG LUYỆN TẬP - THỰC HÀNH (15 phút)\na) Chuyển giao: Giao bài tập: Mở thư mục bài tập bằng nháy đôi, cuộn chuột xem danh sách và kéo chọn 3 tệp ảnh.\nb) Thực hiện: GV theo dõi, chỉnh tư thế đặt tay cho HS.\nc) Báo cáo: HS kiểm tra kết quả bài làm của bạn bên cạnh.\nd) Kết luận: Đánh giá tiết thực hành đạt hiệu quả cao.",
                    "hs": "3. HOẠT ĐỘNG LUYỆN TẬP - THỰC HÀNH (15 phút)\na) Tiếp nhận: Nhận nhiệm vụ trên máy tính.\nb) Thực hiện: Thực hành nháy đôi, cuộn chuột và quét chọn tệp.\nc) Báo cáo: Báo cáo với GV khi hoàn thành.\nd) Kết luận: Thành thạo thao tác nâng cao."},
            "hd4": {"gv": "4. HOẠT ĐỘNG VẬN DỤNG (3 phút)\na) Chuyển giao: Dặn HS ứng dụng cuộn chuột khi đọc truyện tranh điện tử ở nhà.\nb) Thực hiện: Hướng dẫn tắt máy tính chuẩn.\nc) Báo cáo: HS nêu cảm nghĩ tiết học.\nd) Kết luận: Tổng kết tiết học.",
                    "hs": "4. HOẠT ĐỘNG VẬN DỤNG (3 phút)\na) Tiếp nhận: Ghi nhớ ứng dụng cuộn chuột.\nb) Thực hiện: Thực hiện Shutdown tắt máy đúng quy trình.\nc) Báo cáo: Rời phòng học trật tự.\nd) Kết luận: Hoàn thành bài học."}
        }
    },
    {
        "folder": "Lớp_3", "grade_str": "Lớp 3", "is_primary": True,
        "title": "Bài 2. Xử lí thông tin", "tiet_ppct": 2,
        "yccd": "Học sinh nhận biết 3 giai đoạn xử lý thông tin (Tiếp nhận -> Xử lý -> Quyết định), so sánh vai trò xử lý thông tin của bộ não con người và CPU máy tính.",
        "detail": {
            "hd1": {"gv": "1. Hoạt động MỞ ĐẦU (Khởi động) (5 phút)\na) Chuyển giao: GV đưa tình huống: 'Khi nghe tiếng chuông báo thức reo vào buổi sáng, em sẽ làm gì?'\nb) Thực hiện: Cho HS suy nghĩ 1 phút.\nc) Báo cáo: Mời 3 HS trả lời (Nghe -> Biết giờ -> Tắt chuông thức dậy).\nd) Kết luận: Dẫn dắt đó chính là chu trình xử lý thông tin!",
                    "hs": "1. Hoạt động MỞ ĐẦU (Khởi động) (5 phút)\na) Tiếp nhận: Lắng nghe tình huống chuông báo thức.\nb) Thực hiện: Suy nghĩ phản xạ của bản thân.\nc) Báo cáo: Trả lời tự nhiên các bước hành động.\nd) Kết luận: Hào hứng vào bài học."},
            "hd2": {"gv": "2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI (12 phút)\na) Chuyển giao: GV phân tích 3 giai đoạn: 1. Tiếp nhận thông tin vào -> 2. Bộ não suy nghĩ xử lý -> 3. Xuất kết quả / Quyết định. Liên hệ máy tính có CPU xử lý.\nb) Thực hiện: Chiếu sơ đồ khối so sánh Não người ↔ CPU Máy tính.\nc) Báo cáo: HS chỉ ra điểm tương đồng giữa Não và CPU.\nd) Kết luận: Chốt sơ đồ 3 giai đoạn xử lý thông tin.",
                    "hs": "2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI (12 phút)\na) Tiếp nhận: Quan sát sơ đồ 3 giai đoạn trên máy chiếu.\nb) Thực hiện: Ghi chép khái niệm chốt vào vở.\nc) Báo cáo: Trình bày vai trò của CPU máy tính.\nd) Kết luận: Nắm chắc mô hình xử lý thông tin."},
            "hd3": {"gv": "3. HOẠT ĐỘNG LUYỆN TẬP - THỰC HÀNH (15 phút)\na) Chuyển giao: Giao phiếu bài tập: Phân tích 3 tình huống (Đèn đỏ giao thông, Tiếng trống trường, Phép tính toán).\nb) Thực hiện: Cho HS thảo luận nhóm đôi điền vào bảng 3 cột.\nc) Báo cáo: Đại diện 2 nhóm lên bảng trình bày.\nd) Kết luận: Nhận xét, chốt đáp án đúng trên phiếu.",
                    "hs": "3. HOẠT ĐỘNG LUYỆN TẬP - THỰC HÀNH (15 phút)\na) Tiếp nhận: Nhận phiếu bài tập từ GV.\nb) Thực hiện: Thảo luận nhóm đôi phân tích các tình huống.\nc) Báo cáo: Đại diện lên bảng chia sẻ đáp án.\nd) Kết luận: Sửa bài và hoàn thiện phiếu."},
            "hd4": {"gv": "4. HOẠT ĐỘNG VẬN DỤNG (3 phút)\na) Chuyển giao: Đặt câu hỏi: Máy tính tính toán rất nhanh nhưng có thay thế hoàn toàn bộ não con người được không?\nb) Thực hiện: GV gợi mở về tính sáng tạo và cảm xúc của con người.\nc) Báo cáo: HS trả lời suy nghĩ cá nhân.\nd) Kết luận: Tổng kết bài học ý nghĩa.",
                    "hs": "4. HOẠT ĐỘNG VẬN DỤNG (3 phút)\na) Tiếp nhận: Suy nghĩ về ưu điểm của bộ não con người.\nb) Thực hiện: Thảo luận nhanh.\nc) Báo cáo: Khẳng định con người làm chủ máy tính.\nd) Kết luận: Kết thúc tiết học."}
        }
    },
    {
        "folder": "Lớp_4", "grade_str": "Lớp 4", "is_primary": True,
        "title": "Bài 2. Gõ bàn phím đúng cách", "tiet_ppct": 2,
        "yccd": "Học sinh nhận biết hàng phím cơ sở và 2 phím có gờ F, J; nắm vững quy tắc phân công 10 ngón tay và thực hành đặt tay đúng vị trí xuất phát.",
        "detail": {
            "hd1": {"gv": "1. Hoạt động MỞ ĐẦU (Khởi động) (5 phút)\na) Chuyển giao: GV cho HS sờ lên bàn phím và tìm 2 phím đặc biệt có gai/gờ nổi nhỏ.\nb) Thực hiện: HS tìm trên bàn phím cá nhân.\nc) Báo cáo: HS reo lên 'Đó là phím F và phím J!'.\nd) Kết luận: Giới thiệu bí mật hàng phím cơ sở và gõ 10 ngón.",
                    "hs": "1. Hoạt động MỞ ĐẦU (Khởi động) (5 phút)\na) Tiếp nhận: Nhận thử thách tìm phím có gờ.\nb) Thực hiện: Dùng ngón trỏ sờ khắp hàng phím chữ.\nc) Báo cáo: Phát hiện ra phím F và J có gờ nổi.\nd) Kết luận: Háo hức khám phá quy tắc gõ phím."},
            "hd2": {"gv": "2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI (12 phút)\na) Chuyển giao: GV giới thiệu Hàng phím cơ sở (A S D F G H J K L ;). Hướng dẫn đặt ngón trỏ trái lên F, ngón trỏ phải lên J.\nb) Thực hiện: Làm mẫu phân công ngón tay gõ phím trên mô hình bàn phím to.\nc) Báo cáo: Mời 2 HS phát biểu vị trí 8 ngón tay xuất phát.\nd) Kết luận: Chốt quy tắc xuất phát và trở về hàng phím cơ sở.",
                    "hs": "2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI (12 phút)\na) Tiếp nhận: Quan sát sơ đồ phân công 10 ngón tay.\nb) Thực hiện: Đặt thử 8 ngón tay lên hàng phím cơ sở.\nc) Báo cáo: Trả lời vị trí đặt ngón trỏ và ngón cái.\nd) Kết luận: Nắm vững quy tắc gõ 10 ngón."},
            "hd3": {"gv": "3. HOẠT ĐỘNG LUYỆN TẬP - THỰC HÀNH (15 phút)\na) Chuyển giao: Mở phần mềm RapidTyping / Mario Teaches Typing. Yêu cầu luyện bài gõ hàng phím cơ sở (asdf jkl;).\nb) Thực hiện: GV đi kiểm tra từng bàn, sửa tư thế ngón tay cho HS.\nc) Báo cáo: Xem bảng điểm độ chính xác trên phần mềm.\nd) Kết luận: Tuyên dương HS đạt độ chính xác > 95%.",
                    "hs": "3. HOẠT ĐỘNG LUYỆN TẬP - THỰC HÀNH (15 phút)\na) Tiếp nhận: Bật bài luyện gõ trên phần mềm.\nb) Thực hiện: Đặt tay đúng vị trí và tập gõ nhẹ nhàng không nhìn bàn phím.\nc) Báo cáo: Quan sát tốc độ và tỷ lệ gõ đúng của mình.\nd) Kết luận: Rèn luyện thói quen gõ đúng."},
            "hd4": {"gv": "4. HOẠT ĐỘNG VẬN DỤNG (3 phút)\na) Chuyển giao: Hỏi: Vì sao gõ bằng 10 ngón lại giúp em gõ nhanh hơn gõ 1 ngón?\nb) Thực hiện: Hướng dẫn HS lưu bài luyện gõ và tắt máy.\nc) Báo cáo: HS giải thích 'Vì các ngón chia nhau làm việc!'.\nd) Kết luận: Dặn dò luyện gõ 15 phút mỗi ngày.",
                    "hs": "4. HOẠT ĐỘNG VẬN DỤNG (3 phút)\na) Tiếp nhận: Suy nghĩ lý do gõ 10 ngón nhanh hơn.\nb) Thực hiện: Đóng phần mềm và tắt máy an toàn.\nc) Báo cáo: Trả lời nguyên lý phân công ngón tay.\nd) Kết luận: Hoàn thành bài học."}
        }
    },
    {
        "folder": "Lớp_5", "grade_str": "Lớp 5", "is_primary": True,
        "title": "Bài 2. Tìm kiếm thông tin trên website", "tiet_ppct": 2,
        "yccd": "Học sinh sử dụng được trình duyệt web và máy tìm kiếm, biết lựa chọn từ khóa ngắn gọn chính xác và đánh giá độ tin cậy của thông tin tìm kiếm được.",
        "detail": {
            "hd1": {"gv": "1. Hoạt động MỞ ĐẦU (Khởi động) (5 phút)\na) Chuyển giao: Đặt thử thách: 'Trong 1 phút, ai tìm được địa danh núi cao nhất Việt Nam?'\nb) Thực hiện: HS mở máy tính và sử dụng công cụ tìm kiếm.\nc) Báo cáo: HS reo lên 'Đỉnh Fansipan cao 3.143m!'.\nd) Kết luận: Giới thiệu bài học kỹ năng tìm kiếm thông tin hiệu quả.",
                    "hs": "1. Hoạt động MỞ ĐẦU (Khởi động) (5 phút)\na) Tiếp nhận: Thách thức tìm kiếm thông tin.\nb) Thực hiện: Mở Google và gõ nội dung tìm kiếm.\nc) Báo cáo: Trả lời nhanh kết quả tìm được.\nd) Kết luận: Hào hứng vào bài học."},
            "hd2": {"gv": "2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI (12 phút)\na) Chuyển giao: GV phân biệt Trình duyệt Web (Chrome, Edge) ↔ Máy tìm kiếm (Google, Bing). Hướng dẫn chọn 'Từ khóa' (Keywords) ngắn gọn.\nb) Thực hiện: Làm mẫu gõ từ khóa 'ảnh con hổ' thay vì cả câu văn dài.\nc) Báo cáo: HS nhận xét sự khác nhau về kết quả khi đổi từ khóa.\nd) Kết luận: Chốt quy tắc chọn từ khóa đắt giá.",
                    "hs": "2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI (12 phút)\na) Tiếp nhận: Quan sát GV minh họa chọn từ khóa trên máy chiếu.\nb) Thực hiện: Rút ra bí quyết chọn từ khóa ngắn gọn.\nc) Báo cáo: Đề xuất từ khóa cho các câu hỏi khác nhau.\nd) Kết luận: Ghi nhớ mẹo chọn từ khóa."},
            "hd3": {"gv": "3. HOẠT ĐỘNG LUYỆN TẬP - THỰC HÀNH (15 phút)\na) Chuyển giao: Giao bài tập thực hành 3 nhiệm vụ: 1. Tìm thông tin hoa sen; 2. Tìm hình ảnh Hệ Mặt Trời; 3. Lưu ảnh về máy tính.\nb) Thực hiện: GV di chuyển hướng dẫn HS cách lưu hình ảnh.\nc) Báo cáo: Kiểm tra sản phẩm tệp ảnh trên máy tính của HS.\nd) Kết luận: Đánh giá xuất sắc các nhóm hoàn thành đúng hạn.",
                    "hs": "3. HOẠT ĐỘNG LUYỆN TẬP - THỰC HÀNH (15 phút)\na) Tiếp nhận: Nhận phiếu nhiệm vụ tìm kiếm.\nb) Thực hiện: Gõ từ khóa, tìm kiếm và thao tác lưu hình ảnh.\nc) Báo cáo: Mở sản phẩm ảnh đã lưu cho thầy cô kiểm tra.\nd) Kết luận: Nắm vững kỹ năng tra cứu trực tuyến."},
            "hd4": {"gv": "4. HOẠT ĐỘNG VẬN DỤNG (3 phút)\na) Chuyển giao: Lưu ý HS về việc kiểm tra thông tin trên trang web uy tín (.edu, .gov, báo chính thống).\nb) Thực hiện: Hướng dẫn dọn dẹp lịch sử trình duyệt và tắt máy.\nc) Báo cáo: HS nêu 2 trang web uy tín em biết.\nd) Kết luận: Tổng kết tiết học.",
                    "hs": "4. HOẠT ĐỘNG VẬN DỤNG (3 phút)\na) Tiếp nhận: Ghi nhớ lưu ý chọn trang web uy tín.\nb) Thực hiện: Đóng trình duyệt và tắt máy tính an toàn.\nc) Báo cáo: Nêu ví dụ trang web tin cậy.\nd) Kết luận: Rời phòng máy trật tự."}
        }
    },
    {
        "folder": "Lớp_6", "grade_str": "Lớp 6", "is_primary": False,
        "title": "Bài 2. Xử lí thông tin", "tiet_ppct": 2,
        "yccd": "Học sinh mô tả được mô hình xử lý thông tin 4 bước (Thu nhận -> Lưu trữ -> Xử lý -> Xuất thông tin), giải thích được vai trò xử lý thông tin của máy tính trong đời sống.",
        "detail": {
            "hd1": {"gv": "Hoạt động 1: Khởi động (5 phút)\nBước 1. GV giao nhiệm vụ: Quan sát sơ đồ một chiếc máy tính bỏ túi khi thực hiện phép tính 8 x 9 = 72 và đặt câu hỏi về các công đoạn bên trong.\nBước 2. HS tiếp nhận nhiệm vụ và suy nghĩ.\nBước 3. Báo cáo: Mời 2 HS nêu ý kiến.\nBước 4. GV chốt và giới thiệu bài Mô hình xử lý thông tin.",
                    "sp": "Sản phẩm: Câu trả lời của HS nhận biết được các khâu: Nhập số -> Máy tính toán -> Hiện kết quả."},
            "hd2": {"gv": "Hoạt động 2: Hình thành kiến thức mới (12 phút)\nBước 1. GV giao nhiệm vụ: Nghiên cứu SGK và phân tích 4 bước xử lý thông tin: Thu nhận (Input) -> Lưu trữ (Storage) -> Xử lý (Processing) -> Xuất (Output).\nBước 2. HS thảo luận nhóm 4 người hoàn thành sơ đồ khối.\nBước 3. Báo cáo: Đại diện 1 nhóm trình bày trên bảng.\nBước 4. GV chuẩn hóa mô hình 4 bước.",
                    "sp": "Sản phẩm: Sơ đồ khối 4 bước xử lý thông tin hoàn chỉnh trong vở ghi."},
            "hd3": {"gv": "Hoạt động 3: Luyện tập - Thực hành (23 phút)\nBước 1. GV giao nhiệm vụ: Cho danh sách các thiết bị (Bàn phím, RAM, CPU, Màn hình, Micro, Loa, Ổ cứng). Yêu cầu phân loại vào 4 bước xử lý.\nBước 2. HS làm bài tập cá nhân trên phiếu học tập số.\nBước 3. Báo cáo: Cho HS chấm điểm chéo bài làm.\nBước 4. GV nhận xét chốt bảng phân loại chuẩn.",
                    "sp": "Sản phẩm: Bảng phân loại chính xác các thiết bị máy tính tương ứng với 4 chức năng xử lý."},
            "hd4": {"gv": "Hoạt động 4: Vận dụng & Nhiệm vụ về nhà (5 phút)\nBước 1. GV giao nhiệm vụ: Phân tích chu trình xử lý thông tin của hệ thống xe tự lái hoặc nhận diện khuôn mặt.\nBước 2. HS tiếp nhận nhiệm vụ về nhà.\nBước 3. GV hướng dẫn HS lưu tệp bài làm và tắt máy an toàn.\nBước 4. Đánh giá tiết học.",
                    "sp": "Sản phẩm: Báo cáo phân tích hệ thống tự động hóa trong cuộc sống thực tế."}
        }
    },
    {
        "folder": "Lớp_7", "grade_str": "Lớp 7", "is_primary": False,
        "title": "Bài 2. Phần mềm máy tính", "tiet_ppct": 2,
        "yccd": "Học sinh phân biệt được Phần mềm hệ thống (Hệ điều hành) và Phần mềm ứng dụng, hiểu mối quan hệ giữa Phần cứng, Hệ điều hành và Phần mềm ứng dụng.",
        "detail": {
            "hd1": {"gv": "Hoạt động 1: Khởi động (5 phút)\nBước 1. GV giao nhiệm vụ: Đặt tình huống: Mới mua một chiếc máy tính xách tay cấu hình rất mạnh nhưng chưa cài bất kỳ phần mềm nào. Liệu có dùng được không?\nBước 2. HS thảo luận nhanh.\nBước 3. Báo cáo: HS trả lời 'Không dùng được vì thiếu hệ điều hành!'.\nBước 4. GV kết nối vào bài Phần mềm máy tính.",
                    "sp": "Sản phẩm: HS nhận thức được tầm quan trọng bắt buộc của phần mềm đối với máy tính."},
            "hd2": {"gv": "Hoạt động 2: Hình thành kiến thức mới (12 phút)\nBước 1. GV giao nhiệm vụ: Phân biệt Phần mềm hệ thống (Windows, Android, iOS) ↔ Phần mềm ứng dụng (Word, Chrome, Canva). Phân tích sơ đồ 3 lớp: Phần cứng -> Hệ điều hành -> Ứng dụng.\nBước 2. HS nghiên cứu SGK và vẽ sơ đồ 3 lớp vào vở.\nBước 3. Báo cáo: 2 HS phát biểu mối quan hệ giữa 3 thành phần.\nBước 4. GV chốt khái niệm chuẩn.",
                    "sp": "Sản phẩm: Sơ đồ mối quan hệ 3 lớp Phần cứng - Hệ điều hành - Phần mềm ứng dụng."},
            "hd3": {"gv": "Hoạt động 3: Luyện tập - Thực hành (23 phút)\nBước 1. GV giao nhiệm vụ: Mở máy tính, kiểm tra tên Hệ điều hành đang sử dụng và liệt kê 5 phần mềm ứng dụng có trên máy.\nBước 2. HS thao tác trực tiếp trên máy tính phòng máy.\nBước 3. Báo cáo: HS ghi kết quả vào phiếu thu hoạch thực hành.\nBước 4. GV nhận xét, chấm điểm bài thực hành.",
                    "sp": "Sản phẩm: Phiếu kiểm tra thông tin hệ điều hành và danh mục phần mềm ứng dụng trên máy tính."},
            "hd4": {"gv": "Hoạt động 4: Vận dụng & Nhiệm vụ về nhà (5 phút)\nBước 1. GV giao nhiệm vụ: Tìm hiểu về bản quyền phần mềm và lý do không nên dùng phần mềm bẻ khóa (crack).\nBước 2. HS nhận nhiệm vụ thảo luận nhóm về nhà.\nBước 3. GV dặn dò tắt máy đúng quy trình.\nBước 4. Tổng kết bài học.",
                    "sp": "Sản phẩm: Bài viết ngắn về ý thức sử dụng phần mềm có bản quyền và an toàn mạng."}
        }
    },
    {
        "folder": "Lớp_8", "grade_str": "Lớp 8", "is_primary": False,
        "title": "Bài 2. Thông tin trong môi trường số", "tiet_ppct": 2,
        "yccd": "Học sinh trình bày được đặc điểm của thông tin trong môi trường số (khối lượng lớn, lan truyền nhanh, đa dạng), nắm vững 4 tiêu chí đánh giá độ tin cậy của thông tin và cách phòng chống tin giả.",
        "detail": {
            "hd1": {"gv": "Hoạt động 1: Khởi động (5 phút)\nBước 1. GV giao nhiệm vụ: Chiếu 2 bài đăng trên mạng xã hội về một sự kiện (1 bài tin chuẩn từ báo Chính phủ, 1 bài tin đồn giật gân). Hỏi HS bài nào đáng tin hơn?\nBước 2. HS quan sát và so sánh.\nBước 3. Báo cáo: HS chọn bài đăng chính thống.\nBước 4. GV dẫn dắt vào bài Thông tin trong môi trường số.",
                    "sp": "Sản phẩm: HS bước đầu nhận diện được sự khác biệt giữa thông tin chuẩn xác và tin đồn."},
            "hd2": {"gv": "Hoạt động 2: Hình thành kiến thức mới (12 phút)\nBước 1. GV giao nhiệm vụ: Phân tích 3 đặc điểm thông tin số (Khối lượng lớn, lan truyền nhanh, đa dạng). Giới thiệu 4 tiêu chí đánh giá độ tin cậy: Nguồn tin, Tính cập nhật, Mục đích, Kiểm chứng chéo.\nBước 2. HS thảo luận nhóm 4 người hoàn thành sơ đồ 4 tiêu chí.\nBước 3. Báo cáo: Đại diện nhóm trình bày trước lớp.\nBước 4. GV kết luận bộ quy tắc kiểm chứng thông tin.",
                    "sp": "Sản phẩm: Bộ 4 tiêu chí đánh giá độ tin cậy thông tin được hệ thống hóa chi tiết."},
            "hd3": {"gv": "Hoạt động 3: Luyện tập - Thực hành (23 phút)\nBước 1. GV giao nhiệm vụ: Đánh giá độ tin cậy của 4 nguồn tin thực tế (Website Chinhphu.vn, Bài đăng TikTok cá nhân, Web Bộ GD&ĐT, Blog không rõ tác giả).\nBước 2. HS làm bài tập đánh giá trên máy tính.\nBước 3. Báo cáo: Cho HS tranh luận ý kiến giữa các nhóm.\nBước 4. GV chốt đáp án phân loại chuẩn xác.",
                    "sp": "Sản phẩm: Báo cáo đánh giá độ tin cậy 4 nguồn tin kèm lý do lập luận thuyết phục."},
            "hd4": {"gv": "Hoạt động 4: Vận dụng & Nhiệm vụ về nhà (5 phút)\nBước 1. GV giao nhiệm vụ: Tìm 1 bài viết tin giả (Fake news) gần đây trên mạng và phân tích nguyên nhân sai lệch theo 4 tiêu chí đã học.\nBước 2. HS tiếp nhận bài tập dự án nhỏ.\nBước 3. GV dặn dò quy tắc 5 giây suy nghĩ trước khi nhấn nút Share.\nBước 4. Tổng kết tiết học.",
                    "sp": "Sản phẩm: Bài phân tích tin giả rèn luyện tư duy phản biện trong môi trường số."}
        }
    },
]

def main():
    print("=" * 70)
    print("  TẠO BỘ KẾ HOẠCH BÀI DẠY (KHBD) CHUẨN — TOÀN BỘ 9 KHỐI LỚP (TUẦN 03)")
    print("  Quy chuẩn: Phụ lục IV (CV 5512), Thông tư 02/2025 (CV 3456 về NLS)")
    print("=" * 70)

    total_created = 0

    for item in LESSONS_KHBD_TUAN03:
        folder_prefix = item["folder"]
        grade_str = item["grade_str"]
        is_primary = item["is_primary"]
        title = item["title"]
        tiet_ppct = item["tiet_ppct"]
        yccd = item["yccd"]
        detail = item["detail"]

        safe_title = sanitize_filename(title)
        out_dir = os.path.join(OUT_BASE_DIR, folder_prefix, "Tuần_03")
        os.makedirs(out_dir, exist_ok=True)

        filename = f"KHBD_Tin_hoc_{folder_prefix}_Bai02_{safe_title}.docx"
        out_file = os.path.join(out_dir, filename)

        # Find template from Tuần_02
        t2_dir = os.path.join(OUT_BASE_DIR, folder_prefix, "Tuần_02")
        t2_files = glob.glob(os.path.join(t2_dir, "*.docx"))
        if not t2_files:
            print(f"  [!] Không tìm thấy template Tuần 02 cho {folder_prefix}")
            continue
        
        template_path = t2_files[0]

        if is_primary:
            doc = populate_primary_doc_tuan03(template_path, grade_str, title, tiet_ppct, yccd, detail)
        else:
            doc = populate_secondary_doc_tuan03(template_path, grade_str, title, tiet_ppct, yccd, detail)

        try:
            doc.save(out_file)
            total_created += 1
            print(f"  [+] Đã tạo KHBD: {folder_prefix} -> Tuần_03 -> {filename}")
        except Exception as e:
            print(f"  [!] Lỗi khi lưu {out_file}: {e}")

    print(f"\n==========================================")
    print(f" HOÀN THÀNH TẠO {total_created} FILE KHBD TIN HỌC TUẦN 03")
    print(f"==========================================")

if __name__ == '__main__':
    main()
