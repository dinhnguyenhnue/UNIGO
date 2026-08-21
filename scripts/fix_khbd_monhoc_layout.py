"""
Script sửa layout phần sau TÊN BÀI DẠY trong KHBD:
- Tìm paragraph chứa "Môn học: ... Lớp: ... Thời lượng: ..."
- Bỏ phần "Lớp: ..." (vì đã có trong Table 0 phía trên)
- Tách thành các dòng riêng (hàng dọc):
    Môn học: Tin học
    Thời lượng: 1 tiết (45 phút)
    Tiết theo PPCT: X
"""
import sys
import os
import re
import copy

sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from lxml import etree


def fix_khbd_monhoc_layout(filepath: str, dry_run: bool = False) -> bool:
    """
    Fix the layout of 'Môn học' line in KHBD files.
    Returns True if file was modified, False otherwise.
    """
    doc = Document(filepath)

    # Find the paragraph with "Môn học:" and "Lớp:" on the same line
    target_idx = None
    target_para = None
    for i, p in enumerate(doc.paragraphs[:10]):
        text = p.text.strip()
        if 'Môn học:' in text and ('Lớp:' in text or 'Thời lượng:' in text):
            target_idx = i
            target_para = p
            break

    if target_para is None:
        return False

    text = target_para.text.strip()

    # Extract components using regex
    # Pattern: "Môn học: Tin học     Lớp: 6A1     Thời lượng: 1 tiết (45 phút)"
    mon_hoc_match = re.search(r'Môn học:\s*([^L\s][\w\s]*?)(?:\s{2,}|$)', text)
    if not mon_hoc_match:
        # Try alternative pattern
        mon_hoc_match = re.search(r'Môn học:\s*(.+?)(?:\s{2,}Lớp:|\s{2,}Thời lượng:|$)', text)

    thoi_luong_match = re.search(r'Thời lượng:\s*(.+?)$', text)

    # Extract "Môn học" value
    if mon_hoc_match:
        mon_hoc_val = mon_hoc_match.group(1).strip()
    else:
        # Fallback: extract between "Môn học:" and next field
        parts = text.split('Môn học:')
        if len(parts) > 1:
            rest = parts[1].strip()
            # Find next field separator
            for sep in ['Lớp:', 'Thời lượng:']:
                if sep in rest:
                    mon_hoc_val = rest.split(sep)[0].strip()
                    break
            else:
                mon_hoc_val = rest.strip()
        else:
            return False

    # Extract "Thời lượng" value
    if thoi_luong_match:
        thoi_luong_val = thoi_luong_match.group(1).strip()
    else:
        thoi_luong_val = "1 tiết (45 phút)"

    if dry_run:
        print(f"  Would fix: [{text}]")
        print(f"    -> Môn học: {mon_hoc_val}")
        print(f"    -> Thời lượng: {thoi_luong_val}")
        return True

    # Get formatting from existing paragraph
    existing_run = target_para.runs[0] if target_para.runs else None
    is_italic = existing_run.italic if existing_run else True
    is_bold = existing_run.bold if existing_run else True
    font_size = existing_run.font.size if existing_run and existing_run.font.size else Pt(13)
    font_name = existing_run.font.name if existing_run and existing_run.font.name else 'Times New Roman'

    def set_run_format(run, bold=True, italic=True):
        run.bold = bold
        run.italic = italic
        run.font.size = font_size
        run.font.name = font_name

    # Clear the existing paragraph and set new text
    # First, clear all runs
    for run in target_para.runs:
        run.text = ""
    # Set first run or add new
    if target_para.runs:
        target_para.runs[0].text = f"Môn học: {mon_hoc_val}"
        set_run_format(target_para.runs[0], bold=is_bold, italic=is_italic)
        # Remove extra runs
        p_elem = target_para._element
        runs_to_remove = []
        ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        for child in list(p_elem):
            if child.tag == f'{ns}r':
                runs_to_remove.append(child)
        # Keep first run, remove rest
        for r_elem in runs_to_remove[1:]:
            p_elem.remove(r_elem)
    else:
        run = target_para.add_run(f"Môn học: {mon_hoc_val}")
        set_run_format(run, bold=is_bold, italic=is_italic)

    # Now insert new paragraphs AFTER target paragraph for Thời lượng
    # We need to insert a new paragraph after the current one
    # The "Tiết theo PPCT" line should already exist as the next paragraph

    # Create new paragraph for "Thời lượng"
    body = doc.element.body
    ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'

    # Find the target paragraph element in the body
    target_elem = target_para._element

    # Create new paragraph element for "Thời lượng"
    new_p = copy.deepcopy(target_elem)
    # Clear runs in the new paragraph
    for child in list(new_p):
        if child.tag == f'{ns}r':
            new_p.remove(child)

    # Create a new run
    new_r = etree.SubElement(new_p, f'{ns}r')

    # Copy run properties from original
    if target_para.runs:
        orig_rpr = target_para.runs[0]._element.find(f'{ns}rPr')
        if orig_rpr is not None:
            new_r.insert(0, copy.deepcopy(orig_rpr))

    # Set text
    new_t = etree.SubElement(new_r, f'{ns}t')
    new_t.text = f"Thời lượng: {thoi_luong_val}"
    new_t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')

    # Insert after target paragraph
    target_elem.addnext(new_p)

    # Save
    doc.save(filepath)
    return True


def main():
    khbd_dirs = [r'D:\UNIGO\KHBD_Tin_học', r'D:\UNIGO\KHBD_Robotics']
    dry_run = '--dry-run' in sys.argv

    if dry_run:
        print("=== DRY RUN MODE ===\n")

    total = 0
    fixed = 0
    errors = []

    for d in khbd_dirs:
        if not os.path.exists(d):
            continue
        for root, dirs, files in os.walk(d):
            for f in sorted(files):
                if f.startswith('KHBD_') and f.endswith('.docx') and not f.startswith('~'):
                    fp = os.path.join(root, f)
                    total += 1
                    try:
                        if fix_khbd_monhoc_layout(fp, dry_run=dry_run):
                            fixed += 1
                            if not dry_run:
                                print(f"  ✓ Fixed: {f}")
                    except Exception as e:
                        errors.append((f, str(e)))
                        print(f"  ✗ Error: {f} -> {e}")

    print(f"\n{'=' * 60}")
    print(f"Total files scanned: {total}")
    print(f"Files {'to fix' if dry_run else 'fixed'}: {fixed}")
    print(f"Errors: {len(errors)}")
    if errors:
        for fn, err in errors:
            print(f"  - {fn}: {err}")


if __name__ == '__main__':
    main()
