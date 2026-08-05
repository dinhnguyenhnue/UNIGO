import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_border(cell, **kwargs):
    """
    Set cell borders
    kwargs: top, bottom, left, right, insideH, insideV
    values: dict(val='single', sz='4', color='000000', space='0')
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for border_name, border_props in kwargs.items():
        node = OxmlElement(f'w:{border_name}')
        for key, val in border_props.items():
            node.set(qn(f'w:{key}'), str(val))
        tcBorders.append(node)
    tcPr.append(tcBorders)

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_table_borders(table, color="1D2A64", sz="4"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>
            <w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

def format_run(run, font_name="Times New Roman", font_size=11, bold=False, italic=False, color_rgb=(0,0,0)):
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color_rgb)

def create_docx():
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    # Title
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_logo = p_logo.add_run("UNIGO COLLEGE")
    format_run(r_logo, font_size=18, bold=True, color_rgb=(29, 42, 100))
    
    p_sublogo = doc.add_paragraph()
    p_sublogo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sublogo.paragraph_format.space_after = Pt(12)
    r_sublogo = p_sublogo.add_run("JUNIOR & MIDDLE SCHOOL")
    format_run(r_sublogo, font_size=10, bold=True, color_rgb=(100, 110, 140))

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("THỜI GIAN BIỂU TRONG NGÀY")
    format_run(r_title, font_size=18, bold=True, color_rgb=(29, 42, 100))

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(14)
    r_sub = p_sub.add_run("THỰC HIỆN TỪ NGÀY 04/8/2025")
    format_run(r_sub, font_size=13, bold=True, color_rgb=(29, 42, 100))

    # Table creation: 5 columns
    data = [
        ("Ăn sáng", "7:30 - 7:55", "7:30 - 7:55", "7:30 - 7:55", "7:30 - 7:50", "meal"),
        ("Sinh hoạt đầu giờ", "", "", "7:55 - 8:15", "7:50 - 8:05", "act"),
        ("Chào cờ", "7:50 - 8:05", "7:50 - 8:05", "", "", "act"),
        ("Tiết 1", "8:15 - 8:50", "8:05 - 8:50", "8:15 - 8:50", "8:05 - 8:50", "lesson"),
        ("Tiết 2", "8:55 - 9:30", "8:55 - 9:40", "8:55 - 9:30", "8:55 - 9:40", "lesson"),
        ("Ăn phụ sáng, Ra chơi", "9:30 - 9:45", "9:40 - 9:45", "9:30 - 9:45", "9:40 - 9:45", "break"),
        ("Tiết 3", "9:45 - 10:20", "9:45 - 10:25", "9:45 - 10:20", "9:45 - 10:25", "lesson"),
        ("Tiết 4", "10:25 - 11:00", "10:30 - 11:10", "10:25 - 11:00", "10:30 - 11:10", "lesson"),
        ("Tiết 5 (Sáng)", "", "11:15 - 11:50", "", "11:15 - 11:50", "lesson"),
        ("Dọn dẹp lớp học, vệ sinh cá nhân\nĂn trưa, Nghỉ trưa", "11:00 - 12:50", "11:50 - 12:50", "11:00 - 12:50", "11:50 - 12:50", "meal"),
        ("Tiết 5 (Chiều)", "13:15 - 13:50", "13:05 - 13:50", "13:15 - 13:50", "13:05 - 13:50", "lesson"),
        ("Tiết 6", "13:55 - 14:30", "13:55 - 14:40", "13:55 - 14:30", "13:55 - 14:40", "lesson"),
        ("Ra chơi, ăn xế", "14:30 - 14:55", "14:40 - 14:50", "14:30 - 14:55", "14:40 - 14:50", "break"),
        ("Tiết 7", "14:55 - 15:30", "14:50 - 15:30", "14:55 - 15:30", "14:50 - 15:30", "lesson"),
        ("Tiết 8", "15:35 - 16:10", "15:35 - 16:20", "15:35 - 16:10", "15:35 - 16:20", "lesson"),
        ("Tổng kết cuối ngày & dọn dẹp", "16:10 - 16:15", "16:20 - 16:25", "16:10 - 16:15", "16:20 - 16:25", "act"),
        ("Trả học sinh / học sinh tự ra về", "16:15 - 17:00", "16:25 - 17:00", "16:10 - 17:00", "16:25 - 17:00", "act"),
    ]

    table = doc.add_table(rows=len(data) + 3, cols=5)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, color="1D2A64", sz="6")

    col_widths = [Inches(2.2), Inches(1.25), Inches(1.25), Inches(1.25), Inches(1.25)]

    # Row 0: Top Header
    r0 = table.rows[0]
    cell_act = r0.cells[0]
    cell_act.text = "Hoạt động"
    
    # Merge for Thứ 2 (col 1 & 2)
    r0.cells[1].merge(r0.cells[2])
    r0.cells[1].text = "Thứ 2"

    # Merge for Thứ 3,4,5,6 (col 3 & 4)
    r0.cells[3].merge(r0.cells[4])
    r0.cells[3].text = "Thứ 3, 4, 5, 6"

    # Row 1: Sub Header
    r1 = table.rows[1]
    r1.cells[0].text = "Hoạt động"
    r1.cells[1].text = "Tiểu học"
    r1.cells[2].text = "THCS"
    r1.cells[3].text = "Tiểu học"
    r1.cells[4].text = "THCS"

    # Merge cell_act vertically across row 0 & row 1
    cell_act.merge(r1.cells[0])
    cell_act.text = "Hoạt động"

    # Format Header Row 0 & 1
    for row_idx in [0, 1]:
        for cell in table.rows[row_idx].cells:
            set_cell_shading(cell, "1D2A64")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    format_run(run, font_size=11, bold=True, color_rgb=(255, 255, 255))

    # Fill data rows
    for idx, row_data in enumerate(data):
        r = table.rows[idx + 2]
        act_name, t2_th, t2_thcs, t36_th, t36_thcs, row_type = row_data

        cells = r.cells
        cells[0].text = act_name
        cells[1].text = t2_th
        cells[2].text = t2_thcs
        cells[3].text = t36_th
        cells[4].text = t36_thcs

        # Styling per cell
        for c_idx, cell in enumerate(cells):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = col_widths[c_idx]
            
            p = cell.paragraphs[0]
            if c_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                is_bold = row_type in ["lesson", "meal"]
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                is_bold = True if cell.text and cell.text != "-" else False

            # Shading for empty/na cells
            if not cell.text.strip():
                set_cell_shading(cell, "EAEAEA")

            for run in p.runs:
                color = (29, 42, 100) if is_bold else (50, 50, 50)
                format_run(run, font_size=10.5, bold=is_bold, color_rgb=color)

    # Bottom summary row (Row len(data) + 2)
    r_bot = table.rows[-1]
    r_bot.cells[0].text = ""
    r_bot.cells[1].merge(r_bot.cells[4])
    r_bot.cells[1].text = "Ra về"
    r_bot.cells[1].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p_bot = r_bot.cells[1].paragraphs[0]
    p_bot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in p_bot.runs:
        format_run(run, font_size=11, bold=True, color_rgb=(29, 42, 100))

    # Save
    out_path = "Thời gian biểu trong ngày UNIGO.docx"
    doc.save(out_path)
    print("Saved docx successfully.")

if __name__ == "__main__":
    create_docx()
