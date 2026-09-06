import sys
import openpyxl
import docx

sys.stdout.reconfigure(encoding='utf-8')

wb = openpyxl.load_workbook('d:/UNIGO/Thời khóa biểu giáo viên/Thời khóa biểu - Đậu Đình Nguyên.xlsx')
ws = wb.active

print('=== XLSX FORMATTING ===')
print('Dimensions:', ws.dimensions)
for r in range(1, ws.max_row + 1):
    row_data = []
    for c in range(1, ws.max_column + 1):
        cell = ws.cell(r, c)
        fill_color = cell.fill.start_color.rgb if cell.fill and cell.fill.start_color else None
        val = str(cell.value) if cell.value is not None else ''
        if val:
            row_data.append((c, val.replace('\n', ' '), fill_color, cell.font.name, cell.font.size, cell.font.bold))
    print(f'Row {r:2d}: {row_data}')

print('\nMerged ranges in xlsx:', [str(m) for m in ws.merged_cells.ranges])
print('Column widths:', {col: ws.column_dimensions[col].width for col in ['A', 'B', 'C', 'D', 'E', 'F', 'G', 'H']})

print('\n=== DOCX FORMATTING ===')
doc = docx.Document('d:/UNIGO/Thời khóa biểu giáo viên/Thời khóa biểu - Đậu Đình Nguyên.docx')
for p_idx, p in enumerate(doc.paragraphs):
    print(f'P{p_idx} (align={p.alignment}): {p.text}')

for t_idx, table in enumerate(doc.tables):
    print(f'Table {t_idx}: {len(table.rows)} rows x {len(table.columns)} cols')
    for r_idx, row in enumerate(table.rows):
        cells_text = [cell.text.replace('\n', ' ') for cell in row.cells]
        print(f'  Row {r_idx:2d}: {cells_text}')
