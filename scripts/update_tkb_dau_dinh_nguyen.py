import os
import sys
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.oxml import OxmlElement

sys.stdout.reconfigure(encoding='utf-8')

DIR_WORD = r'D:\UNIGO\Thời khóa biểu giáo viên'
DIR_MAU = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Thời khóa biểu'
TKB_EXCEL_SOURCE = r'D:\UNIGO\TKB toàn trường CHECK - chuẩn.xlsx'

os.makedirs(DIR_WORD, exist_ok=True)
os.makedirs(DIR_MAU, exist_ok=True)

# ─────────────── THỜI GIAN ───────────────────────────────
# TH = Tiểu học | THCS = Trung học cơ sở
# Thứ 2: Chào cờ 7:50-8:05, tiết bắt đầu từ 8:15
# Thứ 3,4,5,6: Sinh hoạt đầu giờ, tiết bắt đầu từ 8:05

TIMES_SANG_TH_T2 = {1: "8:15–8:50", 2: "8:55–9:30", 3: "9:45–10:20", 4: "10:25–11:00", 5: ""}
TIMES_SANG_TH    = {1: "8:15–8:50", 2: "8:55–9:30", 3: "9:45–10:20", 4: "10:25–11:00", 5: ""}
TIMES_SANG_THCS  = {1: "8:05–8:50", 2: "8:55–9:40", 3: "9:45–10:25", 4: "10:30–11:10", 5: "11:15–11:50"}

TIMES_CHIEU_TH   = {1: "13:15–13:50", 2: "13:55–14:30", 3: "14:55–15:30", 4: "15:35–16:10", 5: ""}
TIMES_CHIEU_THCS = {1: "13:05–13:50", 2: "13:55–14:40", 3: "14:50–15:30", 4: "15:35–16:20", 5: ""}

# ─────────────── MÀU SẮC ─────────────────────────────────
# Excel (hex strings)
CLR_TIN        = "D6E4FF"   # xanh nhạt — Tin học
CLR_ROBOTICS   = "FFE5CC"   # cam nhạt  — Robotics
CLR_HEADER_BG  = "1D2A64"   # đậm — tiêu đề
CLR_HEADER_FG  = "FFFFFF"   # trắng
CLR_SESSION_BG = "E8ECF5"   # header buổi
CLR_PERIOD_BG  = "F5F7FF"   # cột tiết
CLR_TIME_BG    = "EFF3FF"   # cột giờ

# DOCX (hex strings - no hash)
DOCX_TIN       = "D6E4FF"
DOCX_ROBOTICS  = "FFE5CC"
DOCX_HEADER_BG = "1D2A64"
DOCX_SESSION   = "E8ECF5"
DOCX_PERIOD    = "F5F7FF"
DOCX_TIME      = "EFF3FF"


def get_nguyen_data():
    wb = openpyxl.load_workbook(TKB_EXCEL_SOURCE, data_only=True)
    ws = wb['TKB_LOP_SC']

    classes = {}
    curr_lop = None
    for c in range(3, ws.max_column + 1):
        val = ws.cell(4, c).value
        if val:
            curr_lop = str(val).strip()
        buoi = str(ws.cell(5, c).value or '').strip()
        classes[c] = (curr_lop, buoi)

    slots = []
    curr_thu = None
    for r in range(6, ws.max_row + 1):
        t_val = ws.cell(r, 1).value
        if t_val:
            curr_thu = str(int(t_val))
        tiet_val = ws.cell(r, 2).value
        if tiet_val is None:
            continue
        tiet = int(tiet_val)
        for c in range(3, ws.max_column + 1):
            cell_val = ws.cell(r, c).value
            if cell_val and 'Nguyên' in str(cell_val):
                lop, buoi = classes[c]
                raw = str(cell_val).strip()
                mon = raw.split('-')[0].strip()
                is_thcs = any(x in lop for x in ['6', '7', '8'])
                slots.append((curr_thu, buoi, tiet, mon, lop, is_thcs))

    return {
        'full_name': 'Đậu Đình Nguyên',
        'subject': 'Tin học & Robotics',
        'slots': slots
    }

# ────────────────── DOCX HELPERS ──────────────────────────
def set_tbl_borders(table):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        for child in list(tblPr[0]):
            if child.tag.endswith('tblBorders'):
                tblPr[0].remove(child)
        borders_xml = parse_xml(
            r'<w:tblBorders %s>'
            r'<w:top w:val="single" w:sz="6" w:space="0" w:color="1D2A64"/>'
            r'<w:left w:val="single" w:sz="6" w:space="0" w:color="1D2A64"/>'
            r'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="1D2A64"/>'
            r'<w:right w:val="single" w:sz="6" w:space="0" w:color="1D2A64"/>'
            r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="AAAACC"/>'
            r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="AAAACC"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        tblPr[0].append(borders_xml)

def set_cell_bg(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_vert_align(cell, align="center"):
    tcPr = cell._element.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)

def add_run(para, text, font_name='Times New Roman', size=12, bold=False,
            italic=False, color=None):
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def style_header_cell(cell, text, bg_hex, fg=(255,255,255), size=12, bold=True):
    set_cell_bg(cell, bg_hex)
    set_cell_vert_align(cell)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    if fg != (0, 0, 0):
        run.font.color.rgb = RGBColor(*fg)

def create_docx_tkb(data, save_paths):
    doc = docx.Document()
    sec = doc.sections[0]
    # Landscape A4 for wide timetable
    sec.page_width  = int(29.7 * 360000)  # 29.7cm (landscape)
    sec.page_height = int(21 * 360000)    # 21cm
    sec.left_margin   = int(1.2 * 360000)
    sec.right_margin  = int(1.2 * 360000)
    sec.top_margin    = int(1.2 * 360000)
    sec.bottom_margin = int(1.2 * 360000)
    # Set landscape orientation in sectPr
    from docx.oxml.ns import qn as _qn
    sectPr = sec._sectPr
    pgSz = sectPr.find(_qn('w:pgSz'))
    if pgSz is None:
        pgSz = OxmlElement('w:pgSz')
        sectPr.append(pgSz)
    pgSz.set(_qn('w:orient'), 'landscape')

    # ── Title block ──────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "TRƯỜNG TIỂU HỌC VÀ THCS UNIGO\n", size=13, bold=True, color=(29,42,100))
    add_run(p, "THỜI KHOÁ BIỂU CÁ NHÂN GIÁO VIÊN\n", size=16, bold=True, color=(29,42,100))
    add_run(p, "Năm học 2026 – 2027", size=12, italic=True, color=(80,80,80))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after  = Pt(8)
    add_run(p2, f"Giáo viên: {data['full_name']}   |   Môn: {data['subject']}   |   Tổng: {len(data['slots'])} tiết/tuần",
            size=13, bold=True, color=(29,42,100))

    # ── Bảng: 13 hàng × 8 cột ────────────────────────────
    # Cols: [Tiết | TH – Sáng/Chiều | THCS – Sáng/Chiều | T2 | T3 | T4 | T5 | T6]
    table = doc.add_table(rows=13, cols=8)
    table.style = 'Table Grid'
    set_tbl_borders(table)

    COL_HEADERS = ["Tiết", "Thời gian\n(Tiểu học)", "Thời gian\n(THCS)", "Thứ Hai", "Thứ Ba", "Thứ Tư", "Thứ Năm", "Thứ Sáu"]

    # Row 0 – column headers
    for ci, hdr in enumerate(COL_HEADERS):
        c = table.rows[0].cells[ci]
        style_header_cell(c, hdr, DOCX_HEADER_BG, fg=(255,255,255), size=11, bold=True)

    def add_session_row(row_idx, label):
        cell0 = table.rows[row_idx].cells[0]
        for ci in range(1, 8):
            cell0.merge(table.rows[row_idx].cells[ci])
        style_header_cell(cell0, label, DOCX_SESSION, fg=(29,42,100), size=12, bold=True)

    # Row 1 – BUỔI SÁNG
    add_session_row(1, "▸  BUỔI SÁNG")

    # Row 7 – BUỔI CHIỀU
    add_session_row(7, "▸  BUỔI CHIỀU")

    # Điền cột Tiết + Thời gian
    for p_num in range(1, 6):
        # Sáng – rows 2..6
        row_s = table.rows[p_num + 1]
        set_cell_bg(row_s.cells[0], DOCX_PERIOD)
        set_cell_bg(row_s.cells[1], DOCX_TIME)
        set_cell_bg(row_s.cells[2], DOCX_TIME)
        for ci in (0, 1, 2):
            set_cell_vert_align(row_s.cells[ci])
            row_s.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(row_s.cells[0].paragraphs[0], f"Tiết {p_num}", size=12, bold=True, color=(29,42,100))
        add_run(row_s.cells[1].paragraphs[0], TIMES_SANG_TH.get(p_num, ""), size=11, italic=True, color=(60,60,80))
        add_run(row_s.cells[2].paragraphs[0], TIMES_SANG_THCS.get(p_num, ""), size=11, italic=True, color=(60,60,80))

        # Chiều – rows 8..12
        row_c = table.rows[p_num + 7]
        set_cell_bg(row_c.cells[0], DOCX_PERIOD)
        set_cell_bg(row_c.cells[1], DOCX_TIME)
        set_cell_bg(row_c.cells[2], DOCX_TIME)
        for ci in (0, 1, 2):
            set_cell_vert_align(row_c.cells[ci])
            row_c.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(row_c.cells[0].paragraphs[0], f"Tiết {p_num}", size=12, bold=True, color=(29,42,100))
        add_run(row_c.cells[1].paragraphs[0], TIMES_CHIEU_TH.get(p_num, ""), size=11, italic=True, color=(60,60,80))
        add_run(row_c.cells[2].paragraphs[0], TIMES_CHIEU_THCS.get(p_num, ""), size=11, italic=True, color=(60,60,80))

    # Điền tiết dạy – tô màu Tin vs Robotics
    day_to_col = {'2': 3, '3': 4, '4': 5, '5': 6, '6': 7}
    for slot in data['slots']:
        day, sess, p_num, subject, cls_name, is_thcs = slot
        ci = day_to_col[day]
        ri = (p_num + 1) if sess == 'Sáng' else (p_num + 7)

        cell = table.rows[ri].cells[ci]
        is_rob = "Robotics" in subject
        bg = DOCX_ROBOTICS if is_rob else DOCX_TIN
        fg = (140, 60, 0) if is_rob else (0, 60, 140)
        set_cell_bg(cell, bg)
        set_cell_vert_align(cell)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        subject_label = "🤖 Robotics" if is_rob else "💻 Tin học"
        add_run(cell.paragraphs[0], subject_label + "\n", size=11, bold=True, color=fg)
        add_run(cell.paragraphs[0], f"({cls_name})", size=10, bold=False, color=(80,80,80))

    for path in save_paths:
        try:
            doc.save(path)
            print(f"✓ Saved DOCX: {path}")
        except PermissionError:
            print(f"✗ Permission error: {path} (đóng file Word trước)")


# ────────────────── EXCEL ──────────────────────────────────
def xl_border(thin=True):
    st = 'thin' if thin else 'medium'
    s = Side(style=st, color='1D2A64')
    t = Side(style='thin', color='AAAACC')
    return Border(left=t, right=t, top=t, bottom=t)

def xl_thick_border():
    s = Side(style='medium', color='1D2A64')
    return Border(left=s, right=s, top=s, bottom=s)

def xl_fill(hex_color):
    return PatternFill(start_color=hex_color, end_color=hex_color, fill_type="solid")

def xl_font(size=12, bold=False, italic=False, color="000000", name='Times New Roman'):
    return Font(name=name, size=size, bold=bold, italic=italic, color=color)

def xl_align(h='center', v='center', wrap=True):
    return Alignment(horizontal=h, vertical=v, wrap_text=wrap)

THIN = Side(style='thin', color='AAAACC')
THICK = Side(style='medium', color='1D2A64')
INNER_BORDER = Border(left=THIN, right=THIN, top=THIN, bottom=THIN)
OUTER_BORDER = Border(left=THICK, right=THICK, top=THICK, bottom=THICK)

def create_excel_tkb(data, save_path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "TKB Đậu Đình Nguyên"

    # ── Rows 1-4: header block ────────────────────────────
    ws.merge_cells('A1:H1')
    ws['A1'] = "TRƯỜNG TIỂU HỌC VÀ THCS UNIGO"
    ws['A1'].font = xl_font(14, bold=True, color="FFFFFF")
    ws['A1'].fill = xl_fill("1D2A64")
    ws['A1'].alignment = xl_align()

    ws.merge_cells('A2:H2')
    ws['A2'] = "THỜI KHOÁ BIỂU CÁ NHÂN GIÁO VIÊN  —  Năm học 2026–2027"
    ws['A2'].font = xl_font(15, bold=True, color="FFFFFF")
    ws['A2'].fill = xl_fill("253580")
    ws['A2'].alignment = xl_align()

    ws.merge_cells('A3:H3')
    ws['A3'] = f"Giáo viên: {data['full_name']}   |   Bộ môn: {data['subject']}   |   Tổng: {len(data['slots'])} tiết/tuần"
    ws['A3'].font = xl_font(12, italic=True, color="1D2A64")
    ws['A3'].fill = xl_fill("EBF0FF")
    ws['A3'].alignment = xl_align()

    ws.row_dimensions[1].height = 28
    ws.row_dimensions[2].height = 30
    ws.row_dimensions[3].height = 24

    # ── Row 4 – legend ────────────────────────────────────
    ws.merge_cells('A4:D4')
    ws['A4'] = "🔵 Xanh nhạt: Tin học"
    ws['A4'].font = xl_font(10, color="003C8C")
    ws['A4'].fill = xl_fill(CLR_TIN)
    ws['A4'].alignment = Alignment(horizontal='center')

    ws.merge_cells('E4:H4')
    ws['E4'] = "🟠 Cam nhạt: Robotics"
    ws['E4'].font = xl_font(10, color="8C3C00")
    ws['E4'].fill = xl_fill(CLR_ROBOTICS)
    ws['E4'].alignment = Alignment(horizontal='center')

    ws.row_dimensions[4].height = 18

    # ── Row 5 – column headers ───────────────────────────
    col_headers = ["Tiết", "Giờ TH", "Giờ THCS", "Thứ Hai", "Thứ Ba", "Thứ Tư", "Thứ Năm", "Thứ Sáu"]
    for ci, h in enumerate(col_headers, 1):
        c = ws.cell(5, ci, h)
        c.font = xl_font(12, bold=True, color="FFFFFF")
        c.fill = xl_fill("1D2A64")
        c.alignment = xl_align()
        c.border = INNER_BORDER
    ws.row_dimensions[5].height = 30

    # ── Helper: draw a session header ────────────────────
    def draw_session(row, label):
        ws.merge_cells(f'A{row}:H{row}')
        c = ws[f'A{row}']
        c.value = label
        c.font = xl_font(12, bold=True, color="1D2A64")
        c.fill = xl_fill(CLR_SESSION_BG)
        c.alignment = xl_align()
        ws.row_dimensions[row].height = 20
        for ci in range(1, 9):
            ws.cell(row, ci).border = INNER_BORDER

    # ── Buổi sáng: rows 6(header) + 7..11(tiết 1..5) ────
    draw_session(6, "▸   BUỔI SÁNG")
    for p_num in range(1, 6):
        ri = 6 + p_num
        ws.row_dimensions[ri].height = 40
        # Tiết
        c_tiet = ws.cell(ri, 1, f"Tiết {p_num}")
        c_tiet.font = xl_font(12, bold=True, color="1D2A64")
        c_tiet.fill = xl_fill(CLR_PERIOD_BG)
        c_tiet.alignment = xl_align()
        c_tiet.border = INNER_BORDER
        # Giờ TH
        c_th = ws.cell(ri, 2, TIMES_SANG_TH.get(p_num, ""))
        c_th.font = xl_font(10, italic=True, color="404060")
        c_th.fill = xl_fill(CLR_TIME_BG)
        c_th.alignment = xl_align()
        c_th.border = INNER_BORDER
        # Giờ THCS
        c_thcs = ws.cell(ri, 3, TIMES_SANG_THCS.get(p_num, ""))
        c_thcs.font = xl_font(10, italic=True, color="404060")
        c_thcs.fill = xl_fill(CLR_TIME_BG)
        c_thcs.alignment = xl_align()
        c_thcs.border = INNER_BORDER
        # Empty day cells
        for ci in range(4, 9):
            cc = ws.cell(ri, ci)
            cc.border = INNER_BORDER
            cc.alignment = xl_align()

    # ── Buổi chiều: rows 12(header) + 13..17(tiết 1..5) ─
    draw_session(12, "▸   BUỔI CHIỀU")
    for p_num in range(1, 6):
        ri = 12 + p_num
        ws.row_dimensions[ri].height = 40
        c_tiet = ws.cell(ri, 1, f"Tiết {p_num}")
        c_tiet.font = xl_font(12, bold=True, color="1D2A64")
        c_tiet.fill = xl_fill(CLR_PERIOD_BG)
        c_tiet.alignment = xl_align()
        c_tiet.border = INNER_BORDER
        c_th = ws.cell(ri, 2, TIMES_CHIEU_TH.get(p_num, ""))
        c_th.font = xl_font(10, italic=True, color="404060")
        c_th.fill = xl_fill(CLR_TIME_BG)
        c_th.alignment = xl_align()
        c_th.border = INNER_BORDER
        c_thcs = ws.cell(ri, 3, TIMES_CHIEU_THCS.get(p_num, ""))
        c_thcs.font = xl_font(10, italic=True, color="404060")
        c_thcs.fill = xl_fill(CLR_TIME_BG)
        c_thcs.alignment = xl_align()
        c_thcs.border = INNER_BORDER
        for ci in range(4, 9):
            cc = ws.cell(ri, ci)
            cc.border = INNER_BORDER
            cc.alignment = xl_align()

    # ── Fill tiết dạy ─────────────────────────────────────
    day_to_col = {'2': 4, '3': 5, '4': 6, '5': 7, '6': 8}
    for slot in data['slots']:
        day, sess, p_num, subject, cls_name, is_thcs = slot
        ci = day_to_col[day]
        ri = (6 + p_num) if sess == 'Sáng' else (12 + p_num)

        is_rob = "Robotics" in subject
        bg = CLR_ROBOTICS if is_rob else CLR_TIN
        fg = "8C3C00" if is_rob else "003C8C"
        label = f"🤖 Robotics\n({cls_name})" if is_rob else f"💻 Tin học\n({cls_name})"

        cc = ws.cell(ri, ci, label)
        cc.font = xl_font(11, bold=True, color=fg)
        cc.fill = xl_fill(bg)
        cc.alignment = xl_align()
        cc.border = INNER_BORDER

    # ── Column widths ─────────────────────────────────────
    ws.column_dimensions['A'].width = 9
    ws.column_dimensions['B'].width = 13
    ws.column_dimensions['C'].width = 13
    for col in ['D', 'E', 'F', 'G', 'H']:
        ws.column_dimensions[col].width = 16

    # ── Page setup: Landscape A4, fit to 1 page wide ──────
    from openpyxl.worksheet.page import PageMargins
    ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
    ws.page_setup.paperSize   = ws.PAPERSIZE_A4
    ws.page_setup.fitToPage   = True
    ws.page_setup.fitToWidth  = 1
    ws.page_setup.fitToHeight = 0
    ws.page_margins = PageMargins(left=0.5, right=0.5, top=0.6, bottom=0.6,
                                  header=0.3, footer=0.3)
    ws.print_area = 'A1:H17'
    ws.sheet_properties.pageSetUpPr.fitToPage = True

    # Outer border around whole table (rows 5-17, cols A-H)
    from openpyxl.utils import get_column_letter
    for ri in range(5, 18):
        for ci in range(1, 9):
            cell = ws.cell(ri, ci)
            left = THICK.border_style if ci == 1 else THIN.border_style
            right = THICK.border_style if ci == 8 else THIN.border_style
            top = THICK.border_style if ri == 5 else THIN.border_style
            bottom = THICK.border_style if ri == 17 else THIN.border_style
            left_c = "1D2A64" if ci == 1 else "AAAACC"
            right_c = "1D2A64" if ci == 8 else "AAAACC"
            top_c = "1D2A64" if ri == 5 else "AAAACC"
            bottom_c = "1D2A64" if ri == 17 else "AAAACC"
            cell.border = Border(
                left=Side(style=left, color=left_c),
                right=Side(style=right, color=right_c),
                top=Side(style=top, color=top_c),
                bottom=Side(style=bottom, color=bottom_c)
            )

    try:
        wb.save(save_path)
        print(f"✓ Saved EXCEL: {save_path}")
    except PermissionError:
        print(f"✗ Permission error: {save_path} (đóng file Excel trước)")


def main():
    print("Đọc dữ liệu Đậu Đình Nguyên từ TKB toàn trường...")
    data = get_nguyen_data()
    print(f"  → {len(data['slots'])} tiết dạy / tuần")

    word1 = os.path.join(DIR_WORD, "Thời khóa biểu - Đậu Đình Nguyên.docx")
    word2 = os.path.join(DIR_MAU,  "Thời khóa biểu - Đậu Đình Nguyên.docx")
    excel1 = os.path.join(DIR_WORD, "Thời khóa biểu - Đậu Đình Nguyên.xlsx")

    create_docx_tkb(data, [word1, word2])
    create_excel_tkb(data, excel1)
    print("XONG! Thời khóa biểu đã được tô màu và thêm đầy đủ khung giờ TH & THCS.")

if __name__ == '__main__':
    main()
