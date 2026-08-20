"""
fix_nls_all_khbd.py — Cập nhật Năng lực số (NLS) cho toàn bộ KHBD
theo CV 3456/BGDĐT-GDPT đúng Miền, Bậc và Descriptor.

Usage:
    python fix_nls_all_khbd.py                  # Fix all files
    python fix_nls_all_khbd.py --dry-run        # Preview changes only
    python fix_nls_all_khbd.py --test 5         # Fix first 5 files only
"""
import os, sys, json, re, copy, argparse
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Emu
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ═══════════════════════════════════════════════════════════════════════
# 1. CONSTANTS & MAPPINGS
# ═══════════════════════════════════════════════════════════════════════

KHBD_DIRS = [
    r'D:\UNIGO\KHBD_Tin_học',
    r'D:\UNIGO\KHBD_Robotics',
]

CV3456_JSON = r'D:\UNIGO\.agents\skills\tao-khbd\references\cv3456_full_data.json'

# Grade folder name → (JSON key for bậc column, bậc number)
GRADE_TO_BAC = {
    'Tiền_tiểu_học': ('L1-3', 1),
    'Lớp_1': ('L1-3', 1),
    'Lớp_2': ('L1-3', 1),
    'Lớp_3': ('L1-3', 1),
    'Lớp_4': ('L4-5', 2),
    'Lớp_5': ('L4-5', 2),
    'Lớp_6': ('L6-7', 3),
    'Lớp_7': ('L6-7', 3),
    'Lớp_8': ('L8-9', 4),
    'Lớp_9': ('L8-9', 4),
}

# Miền Roman → Full name
MIEN_ROMAN = {
    1: 'I', 2: 'II', 3: 'III', 4: 'IV', 5: 'V', 6: 'VI',
}
MIEN_NAMES = {
    'I':   'Khai thác dữ liệu và thông tin',
    'II':  'Giao tiếp và Hợp tác',
    'III': 'Sáng tạo nội dung số',
    'IV':  'An toàn',
    'V':   'Giải quyết vấn đề',
    'VI':  'Ứng dụng trí tuệ nhân tạo',
}

# Thành tố key (matching JSON keys) → (Miền Roman, short thành tố code)
THANH_TO_MAP = {
    '1.1. Duyệt, tìm kiếm và lọc dữ liệu, thông tin và nội dung số': 'I',
    '1.2. Đánh giá dữ liệu, thông tin và nội dung số': 'I',
    '1.3. Quản lý dữ liệu, thông tin và nội dung số': 'I',
    '2.1. Tương tác thông qua công nghệ số': 'II',
    '2.2. Chia sẻ thông tin và nội dung thông qua công nghệ số': 'II',
    '2.3. Sử dụng công nghệ số để thực hiện trách nhiệm công dân': 'II',
    '2.4. Hợp tác thông qua công nghệ số': 'II',
    '2.5. Quy tắc ứng xử trên mạng': 'II',
    '2.6. Quản lý danh tính số': 'II',
    '3.1. Phát triển nội dung số': 'III',
    '3.2. Tích hợp và tạo lập lại nội dung số': 'III',
    '3.3. Thực thi bản quyền và giấy phép': 'III',
    '3.4. Lập trình': 'III',
    '4.1. Bảo vệ thiết bị': 'IV',
    '4.2. Bảo vệ dữ liệu cá nhân và quyền riêng tư': 'IV',
    '4.3. Bảo vệ sức khỏe và an sinh số': 'IV',
    '4.4. Bảo vệ môi trường': 'IV',
    '5.1. Giải quyết các vấn đề kỹ thuật': 'V',
    '5.2. Xác định nhu cầu và giải pháp công nghệ': 'V',
    '5.3. Sử dụng sáng tạo công nghệ số': 'V',
    '5.4. Xác định các vấn đề cần cải thiện về NLS': 'V',
    '6.1. Hiểu biết về trí tuệ nhân tạo': 'VI',
    '6.2. Sử dụng trí tuệ nhân tạo': 'VI',
    '6.3. Đánh giá trí tuệ nhân tạo': 'VI',
}

# ═══════════════════════════════════════════════════════════════════════
# 2. LESSON CATEGORY MAPPING (keyword → NLS domains)
# ═══════════════════════════════════════════════════════════════════════

# Each category: keywords (case-insensitive match on filename),
#   primary NLS (thành tố key), secondary NLS (thành tố key),
#   activity refs for primary/secondary
LESSON_CATEGORIES = [
    # === LẬP TRÌNH / THUẬT TOÁN ===
    {
        'name': 'programming',
        'keywords': ['lap_trinh', 'thuat_toan', 'chuong_trinh', 'cau_lenh',
                     're_nhanh', 'vong_lap', 'ham_trong', 'bien_trong',
                     'Thuattoan', 'tu_thuat_toan', 'scratch', 'python',
                     'du_an_hoc_tap', 'lenh_dieu_kien', 'lenh_lap'],
        'primary': '3.4. Lập trình',
        'secondary': '5.3. Sử dụng sáng tạo công nghệ số',
        'hd_primary': 'Hoạt động 2, Hoạt động 3',
        'hd_secondary': 'Hoạt động 3, Hoạt động 4',
    },
    # === INTERNET / MẠNG ===
    {
        'name': 'internet',
        'keywords': ['internet', 'mang_may_tinh', 'mang_thong_tin',
                     'trang_web', 'lien_ket', 'World_Wide_Web'],
        'primary': '1.1. Duyệt, tìm kiếm và lọc dữ liệu, thông tin và nội dung số',
        'secondary': '2.1. Tương tác thông qua công nghệ số',
        'hd_primary': 'Hoạt động 2, Hoạt động 3',
        'hd_secondary': 'Hoạt động 3, Hoạt động 4',
    },
    # === AN TOÀN / BẢO MẬT / ĐẠO ĐỨC SỐ ===
    {
        'name': 'safety',
        'keywords': ['an_toan', 'bao_mat', 'bao_ve', 'mat_khau',
                     'dao_duc', 'van_hoa', 'ung_xu', 'quyen_rieng_tu'],
        'primary': '4.2. Bảo vệ dữ liệu cá nhân và quyền riêng tư',
        'secondary': '2.5. Quy tắc ứng xử trên mạng',
        'hd_primary': 'Hoạt động 2, Hoạt động 3',
        'hd_secondary': 'Hoạt động 2, Hoạt động 4',
    },
    # === SOẠN THẢO VĂN BẢN / TRÌNH CHIẾU ===
    {
        'name': 'document',
        'keywords': ['van_ban', 'soan_thao', 'trinh_chieu', 'trang_chieu',
                     'dinh_dang', 'dau_trang', 'chan_trang', 'liet_ke',
                     'hinh_anh_trong_van', 'ban_mau', 'word', 'powerpoint'],
        'primary': '3.1. Phát triển nội dung số',
        'secondary': '1.3. Quản lý dữ liệu, thông tin và nội dung số',
        'hd_primary': 'Hoạt động 2, Hoạt động 3',
        'hd_secondary': 'Hoạt động 3, Hoạt động 4',
    },
    # === BẢNG TÍNH / DỮ LIỆU ===
    {
        'name': 'spreadsheet',
        'keywords': ['bang_tinh', 'sap_xep', 'loc_du_lieu', 'bieu_do',
                     'du_lieu_bang', 'excel', 'cong_thuc'],
        'primary': '1.3. Quản lý dữ liệu, thông tin và nội dung số',
        'secondary': '3.1. Phát triển nội dung số',
        'hd_primary': 'Hoạt động 2, Hoạt động 3',
        'hd_secondary': 'Hoạt động 3, Hoạt động 4',
    },
    # === ROBOTICS (động cơ, sensor, cơ khí) — PHẢI ĐẶT TRƯỚC AI ===
    {
        'name': 'robotics',
        'keywords': ['Robot', 'robot', 'ong_co', 'Dynamixel', 'cam_bien',
                     'sensor', 'ieu_khien', 'hut_bui', 'co_inh_vat',
                     'nhan_biet_am_thanh', 'nhan_biet_vat_the',
                     'me_cung', 'giao_thong', 'tang_giam_chieu_dai',
                     'Domino', 'hap_dan', 'ngau_nhien', 'phuc_vu_oi_song',
                     'Kit', 'kit', 'ieu_khien_tu_xa', 'Trien_lam_Robotics',
                     'Robotics'],
        'primary': '5.2. Xác định nhu cầu và giải pháp công nghệ',
        'secondary': '3.4. Lập trình',
        'hd_primary': 'Hoạt động 2, Hoạt động 3',
        'hd_secondary': 'Hoạt động 3, Hoạt động 4',
    },
    # === AI / TRÍ TUỆ NHÂN TẠO (keywords cụ thể, tránh false positive) ===
    {
        'name': 'ai',
        'keywords': ['AI_quanh', 'tri_tue_nhan_tao', 'may_thong_minh',
                     'Con_nguoi_va_may', '_AI_', 'Tiet33_AI', 'Tiet34_Con'],
        'primary': '6.1. Hiểu biết về trí tuệ nhân tạo',
        'secondary': '6.2. Sử dụng trí tuệ nhân tạo',
        'hd_primary': 'Hoạt động 2, Hoạt động 3',
        'hd_secondary': 'Hoạt động 3, Hoạt động 4',
    },
    # === SÁNG TẠO / VẼ / MULTIMEDIA ===
    {
        'name': 'creative',
        'keywords': ['tranh', 've_tranh', 'am_thanh', 'hinh_anh',
                     'mau_sac', 'hinh_dang', 'chuyen_ong', 'cau_chuyen',
                     'nhan_vat', 'sang_tao', 'ngoi_nha', 'khu_vuon',
                     'ke_chuyen', 'turtleart'],
        'primary': '3.1. Phát triển nội dung số',
        'secondary': '3.2. Tích hợp và tạo lập lại nội dung số',
        'hd_primary': 'Hoạt động 2, Hoạt động 3',
        'hd_secondary': 'Hoạt động 3, Hoạt động 4',
    },
    # === DỮ LIỆU / THÔNG TIN / XỬ LÝ ===
    {
        'name': 'data',
        'keywords': ['thong_tin', 'du_lieu', 'xu_li', 'luu_tru',
                     'file', 'thu_muc', 'khai_thac', 'tinh_toan',
                     'Luoc_su_cong_cu', 'the_gioi_ki_thuat_so'],
        'primary': '1.1. Duyệt, tìm kiếm và lọc dữ liệu, thông tin và nội dung số',
        'secondary': '1.2. Đánh giá dữ liệu, thông tin và nội dung số',
        'hd_primary': 'Hoạt động 2, Hoạt động 3',
        'hd_secondary': 'Hoạt động 2, Hoạt động 4',
    },
    # === LÀM QUEN MÁY TÍNH / THIẾT BỊ ===
    {
        'name': 'basic_computer',
        'keywords': ['may_tinh', 'lam_quen', 'bo_phan', 'chuot', 'ban_phim',
                     'khoi_dong', 'bieu_tuong', 'man_hinh', 'phim',
                     'go_ten', 'tu_on_gian', 'cong_nghe', 'keo_tha',
                     'Kham_pha'],
        'primary': '5.1. Giải quyết các vấn đề kỹ thuật',
        'secondary': '4.1. Bảo vệ thiết bị',
        'hd_primary': 'Hoạt động 2, Hoạt động 3',
        'hd_secondary': 'Hoạt động 3, Hoạt động 4',
    },
    # === TƯ DUY LOGIC / QUY LUẬT ===
    {
        'name': 'logic',
        'keywords': ['tung_buoc', 'trai_phai', 'truoc_sau', 'quy_luat',
                     'nghe_lenh', 'uong_i', 'kho_bau', 'lac_roi',
                     'lam_viec_cung_ban', 'lam_viec_theo'],
        'primary': '3.4. Lập trình',
        'secondary': '5.3. Sử dụng sáng tạo công nghệ số',
        'hd_primary': 'Hoạt động 2, Hoạt động 3',
        'hd_secondary': 'Hoạt động 3, Hoạt động 4',
    },
    # === NGHỀ NGHIỆP / ĐỊNH HƯỚNG ===
    {
        'name': 'career',
        'keywords': ['nghe_nghiep', 'inh_huong', 'tong_ket',
                     'san_pham_ki_thuat_so', 'Trien_lam'],
        'primary': '5.4. Xác định các vấn đề cần cải thiện về NLS',
        'secondary': '5.2. Xác định nhu cầu và giải pháp công nghệ',
        'hd_primary': 'Hoạt động 2, Hoạt động 3',
        'hd_secondary': 'Hoạt động 3, Hoạt động 4',
    },
    # === ÔN TẬP / ĐÁNH GIÁ (NLS tổng hợp) ===
    {
        'name': 'review',
        'keywords': ['on_tap', 'anh_gia', 'ANH_GIA', 'On_tap',
                     'Luyen_tap_Thuc_hanh'],
        'primary': '5.4. Xác định các vấn đề cần cải thiện về NLS',
        'secondary': '1.2. Đánh giá dữ liệu, thông tin và nội dung số',
        'hd_primary': 'Hoạt động 2, Hoạt động 3',
        'hd_secondary': 'Hoạt động 2, Hoạt động 3',
    },
]

# Fallback category if no keyword matches
FALLBACK_CATEGORY = {
    'name': 'general',
    'primary': '1.1. Duyệt, tìm kiếm và lọc dữ liệu, thông tin và nội dung số',
    'secondary': '4.1. Bảo vệ thiết bị',
    'hd_primary': 'Hoạt động 2, Hoạt động 3',
    'hd_secondary': 'Hoạt động 3, Hoạt động 4',
}


# ═══════════════════════════════════════════════════════════════════════
# 3. HELPER FUNCTIONS
# ═══════════════════════════════════════════════════════════════════════

def load_cv3456_data():
    """Load CV 3456 competency descriptors from JSON."""
    with open(CV3456_JSON, 'r', encoding='utf-8') as f:
        return json.load(f)


def detect_grade_folder(filepath):
    """Extract grade folder name from filepath."""
    parts = filepath.replace('\\', '/').split('/')
    for part in parts:
        if part in GRADE_TO_BAC:
            return part
    return None


def detect_category(filename, filepath=''):
    """Detect lesson category from filename using keyword matching.
    If file is in KHBD_Robotics folder, use robotics as fallback instead of general.
    """
    # Check each category's keywords
    for cat in LESSON_CATEGORIES:
        for kw in cat['keywords']:
            if kw.lower() in filename.lower():
                return cat

    # If in KHBD_Robotics directory, fallback to robotics
    if 'KHBD_Robotics' in filepath:
        for cat in LESSON_CATEGORIES:
            if cat['name'] == 'robotics':
                return cat

    return FALLBACK_CATEGORY


def get_first_bullets(descriptor_text, n=2):
    """Extract first N bullet points from CV 3456 descriptor text."""
    lines = [l.strip() for l in descriptor_text.strip().split('\n') if l.strip()]
    bullets = []
    for line in lines:
        clean = line.lstrip('- ').strip()
        if clean:
            # Remove trailing period and add it back
            clean = clean.rstrip('.')
            bullets.append(clean)
            if len(bullets) >= n:
                break
    return bullets


def compose_nls_line(mien_roman, thanh_to_key, bac_num, descriptor_bullets, hd_ref):
    """
    Compose a complete NLS line:
    - Miền [Roman]. [Tên Miền] (thành tố [Mã]. [Tên] – Bậc [X]): [Descriptor] (Đạt được thông qua HĐ X, HĐ Y).
    """
    mien_name = MIEN_NAMES[mien_roman]
    # Join bullets with "; "
    descriptor_text = '; '.join(descriptor_bullets)
    # Ensure first letter is lowercase after the colon
    # Actually keep original casing from CV 3456
    line = (f'- Miền {mien_roman}. {mien_name} '
            f'(thành tố {thanh_to_key} – Bậc {bac_num}): '
            f'{descriptor_text}. '
            f'(Đạt được thông qua {hd_ref}).')
    # Fix double period
    line = line.replace('..', '.')
    return line


def build_nls_lines(category, grade_folder, cv3456_data):
    """Build 2 NLS lines for a given category + grade."""
    bac_col, bac_num = GRADE_TO_BAC[grade_folder]

    lines = []
    for role in ['primary', 'secondary']:
        thanh_to_key = category[role]
        hd_ref = category[f'hd_{role}']
        mien_roman = THANH_TO_MAP[thanh_to_key]

        # Get descriptor from CV 3456 data
        if thanh_to_key in cv3456_data:
            raw_desc = cv3456_data[thanh_to_key]['descriptors'].get(bac_col, '')
            bullets = get_first_bullets(raw_desc, n=2)
        else:
            bullets = ['Phát triển năng lực số phù hợp với nội dung bài học']

        if not bullets:
            bullets = ['Phát triển năng lực số phù hợp với nội dung bài học']

        line = compose_nls_line(mien_roman, thanh_to_key, bac_num, bullets, hd_ref)
        lines.append(line)

    return lines


# ═══════════════════════════════════════════════════════════════════════
# 4. DOCUMENT MODIFICATION
# ═══════════════════════════════════════════════════════════════════════

def afont(run, bold=None, italic=None, size_pt=13):
    """Apply standard font formatting."""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), 'Times New Roman')


def replace_nls_in_doc(doc, new_nls_lines):
    """
    Find and replace NLS paragraphs in document.
    Looks for paragraphs containing "Miền" + ("thành tố" or "Bậc").
    Returns number of paragraphs replaced.
    """
    # Find indices of NLS paragraphs
    nls_indices = []
    for i, p in enumerate(doc.paragraphs):
        txt = p.text.strip()
        if 'Miền' in txt and ('thành tố' in txt or 'Bậc' in txt):
            nls_indices.append(i)

    if not nls_indices:
        return 0

    # Case 1: Same number of NLS lines → replace in place
    if len(nls_indices) == len(new_nls_lines):
        for idx, new_text in zip(nls_indices, new_nls_lines):
            p = doc.paragraphs[idx]
            # Clear all runs
            for run in p.runs:
                run.text = ''
            # Set text in first run, preserving its formatting
            if p.runs:
                p.runs[0].text = new_text
                afont(p.runs[0])
            else:
                r = p.add_run(new_text)
                afont(r)
        return len(nls_indices)

    # Case 2: Different count → replace first N, add/remove as needed
    # Replace existing ones
    for i, idx in enumerate(nls_indices):
        if i < len(new_nls_lines):
            p = doc.paragraphs[idx]
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = new_nls_lines[i]
                afont(p.runs[0])
            else:
                r = p.add_run(new_nls_lines[i])
                afont(r)
        else:
            # Extra NLS paragraph - clear it
            p = doc.paragraphs[idx]
            for run in p.runs:
                run.text = ''

    # If we need more NLS lines than existing paragraphs
    if len(new_nls_lines) > len(nls_indices):
        # Insert after the last NLS paragraph
        last_idx = nls_indices[-1]
        last_p = doc.paragraphs[last_idx]
        # Copy formatting from last NLS paragraph
        for extra_text in new_nls_lines[len(nls_indices):]:
            new_p = OxmlElement('w:p')
            # Copy paragraph properties
            src_pPr = last_p._p.find(qn('w:pPr'))
            if src_pPr is not None:
                new_pPr = copy.deepcopy(src_pPr)
                new_p.append(new_pPr)
            # Add run with text
            new_r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            rFonts = OxmlElement('w:rFonts')
            for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
                rFonts.set(qn(attr), 'Times New Roman')
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '26')
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '26')
            rPr.append(szCs)
            new_r.append(rPr)
            t = OxmlElement('w:t')
            t.text = extra_text
            t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            new_r.append(t)
            new_p.append(new_r)
            # Insert after last NLS paragraph
            last_p._p.addnext(new_p)

    return max(len(nls_indices), len(new_nls_lines))


def fix_file(filepath, cv3456_data, dry_run=False):
    """
    Fix NLS in a single KHBD file.
    Returns: (success: bool, category_name: str, old_nls: list, new_nls: list)
    """
    filename = os.path.basename(filepath)
    grade_folder = detect_grade_folder(filepath)

    if grade_folder is None:
        return False, 'unknown_grade', [], []

    category = detect_category(filename, filepath)
    new_nls_lines = build_nls_lines(category, grade_folder, cv3456_data)

    # Read current NLS
    try:
        doc = Document(filepath)
    except Exception as e:
        print(f'  ❌ Error opening: {e}')
        return False, category['name'], [], []

    old_nls = []
    for p in doc.paragraphs:
        txt = p.text.strip()
        if 'Miền' in txt and ('thành tố' in txt or 'Bậc' in txt):
            old_nls.append(txt)

    if not old_nls:
        return False, category['name'], [], new_nls_lines

    if dry_run:
        return True, category['name'], old_nls, new_nls_lines

    # Replace NLS
    replaced = replace_nls_in_doc(doc, new_nls_lines)

    if replaced == 0:
        return False, category['name'], old_nls, new_nls_lines

    # Save
    try:
        doc.save(filepath)
    except PermissionError:
        alt = filepath.replace('.docx', '_fixed.docx')
        doc.save(alt)
        print(f'  ⚠️ File locked, saved to: {alt}')

    # Verify header/footer
    verify = Document(filepath)
    hdr_ok = True
    try:
        hdr = verify.sections[0].header
        drawings = sum(1 for p in hdr.paragraphs
                       for r in p.runs if r._r.findall(qn('w:drawing')))
        if drawings < 1:
            hdr_ok = False
    except:
        pass

    return True, category['name'], old_nls, new_nls_lines


# ═══════════════════════════════════════════════════════════════════════
# 5. MAIN
# ═══════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description='Fix NLS in all KHBD files')
    parser.add_argument('--dry-run', action='store_true', help='Preview changes only')
    parser.add_argument('--test', type=int, default=0, help='Fix only first N files')
    args = parser.parse_args()

    print('═' * 70)
    print('  FIX NLS ALL KHBD — CV 3456/BGDĐT-GDPT')
    print('═' * 70)

    # Load CV 3456 data
    cv3456_data = load_cv3456_data()
    print(f'✅ Loaded CV 3456 data: {len(cv3456_data)} competencies')

    # Collect all KHBD files
    all_files = []
    for base_dir in KHBD_DIRS:
        for root, dirs, fnames in os.walk(base_dir):
            for f in sorted(fnames):
                if f.endswith('.docx') and not f.startswith('~'):
                    all_files.append(os.path.join(root, f))

    print(f'📁 Found {len(all_files)} KHBD files')

    if args.test > 0:
        all_files = all_files[:args.test]
        print(f'🔬 Test mode: processing first {args.test} files only')

    if args.dry_run:
        print('🔍 DRY RUN mode: no files will be modified\n')

    # Process files
    stats = {
        'fixed': 0, 'skipped_no_nls': 0, 'skipped_unknown': 0, 'errors': 0,
        'categories': {},
        'bac_fixed': {1: 0, 2: 0, 3: 0, 4: 0},
    }

    for i, fp in enumerate(all_files):
        filename = os.path.basename(fp)
        grade = detect_grade_folder(fp)
        if grade is None:
            stats['skipped_unknown'] += 1
            continue

        bac_col, bac_num = GRADE_TO_BAC[grade]

        success, cat_name, old_nls, new_nls = fix_file(fp, cv3456_data, dry_run=args.dry_run)

        if not old_nls and not success:
            stats['skipped_no_nls'] += 1
            continue

        if success:
            stats['fixed'] += 1
            stats['categories'][cat_name] = stats['categories'].get(cat_name, 0) + 1
            stats['bac_fixed'][bac_num] = stats['bac_fixed'].get(bac_num, 0) + 1

            # Print progress every 50 files or in test/dry-run mode
            if args.dry_run or args.test or (i + 1) % 50 == 0:
                print(f'  [{i+1}/{len(all_files)}] {grade}/{filename}')
                print(f'    Category: {cat_name} | Bậc: {bac_num}')
                if old_nls:
                    print(f'    OLD: {old_nls[0][:100]}...')
                if new_nls:
                    print(f'    NEW: {new_nls[0][:100]}...')
                print()
        else:
            stats['errors'] += 1

    # Summary
    print('\n' + '═' * 70)
    print('  SUMMARY')
    print('═' * 70)
    action = 'Would fix' if args.dry_run else 'Fixed'
    print(f'  {action}: {stats["fixed"]} files')
    print(f'  Skipped (no NLS): {stats["skipped_no_nls"]}')
    print(f'  Skipped (unknown grade): {stats["skipped_unknown"]}')
    print(f'  Errors: {stats["errors"]}')
    print()
    print('  By Bậc:')
    for bac, count in sorted(stats['bac_fixed'].items()):
        bac_names = {1: 'Cơ bản 1 (L1-3)', 2: 'Cơ bản 2 (L4-5)',
                     3: 'Trung cấp 1 (L6-7)', 4: 'Trung cấp 2 (L8-9)'}
        print(f'    Bậc {bac} ({bac_names.get(bac, "?")}): {count} files')
    print()
    print('  By Category:')
    for cat, count in sorted(stats['categories'].items(), key=lambda x: -x[1]):
        print(f'    {cat}: {count} files')
    print()


if __name__ == '__main__':
    main()
