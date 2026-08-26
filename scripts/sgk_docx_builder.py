"""
Stage 4: SGK DOCX Builder
Tạo file .docx cho từng bài, từng khối lớp từ kết quả phân tích.

Output: 1 file DOCX / bài / lớp, bao gồm:
- Tiêu đề bài
- Nội dung text có formatting
- Hình ảnh minh hoạ nhúng đúng vị trí
- Bảng biểu (nếu có)
- Phần bài tập, ghi nhớ

Usage:
    python scripts/sgk_docx_builder.py --grade 3
    python scripts/sgk_docx_builder.py --grade 3 --lessons 1-5
    python scripts/sgk_docx_builder.py --all
"""
import sys, os, io, json, glob, argparse, re
if hasattr(sys.stdout, 'buffer') and not isinstance(sys.stdout, io.TextIOWrapper):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

SGK_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))), 'SGK')


# ─── Formatting Helpers ──────────────────────────────────────────

def set_font(run, name='Times New Roman', size=13, bold=False, italic=False, color=None):
    """Thiết lập font cho một Run."""
    run.font.name = name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


def add_paragraph_with_font(doc, text, size=13, bold=False, italic=False, 
                             alignment=None, space_after=Pt(6), color=None,
                             style=None):
    """Thêm paragraph với font formatting chuẩn."""
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
    
    run = p.add_run(text)
    set_font(run, size=size, bold=bold, italic=italic, color=color)
    
    if alignment:
        p.alignment = alignment
    
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    
    return p


def add_heading_with_style(doc, text, level=1):
    """Thêm heading với style TNR."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return heading


def set_table_borders(table, border_color='000000', border_size='4'):
    """Thiết lập viền bảng đầy đủ."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    
    borders = OxmlElement('w:tblBorders')
    for edge in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), border_size)
        element.set(qn('w:color'), border_color)
        element.set(qn('w:space'), '0')
        borders.append(element)
    
    tblPr.append(borders)
    if tbl.tblPr is None:
        tbl.append(tblPr)


def add_image_from_rendered_page(doc, grade, page_number, bbox=None, max_width_inches=5.5):
    """
    Thêm hình ảnh từ trang đã render vào DOCX.
    
    Nếu có bbox, crop vùng ảnh trước khi chèn.
    Nếu không có bbox, chèn toàn bộ trang (hoặc skip).
    """
    pages_dir = os.path.join(SGK_DIR, f'Lớp_{grade}', 'ocr_output', 'pages')
    img_path = os.path.join(pages_dir, f'page_{page_number:03d}.png')
    
    if not os.path.exists(img_path):
        return None
    
    if bbox:
        # Crop image region
        try:
            from PIL import Image
            img = Image.open(img_path)
            w, h = img.size
            
            x = int(bbox.get('x', 0) * w)
            y = int(bbox.get('y', 0) * h)
            bw = int(bbox.get('w', 1) * w)
            bh = int(bbox.get('h', 1) * h)
            
            cropped = img.crop((x, y, x + bw, y + bh))
            
            # Save cropped image temporarily
            crop_dir = os.path.join(SGK_DIR, f'Lớp_{grade}', 'ocr_output', 'cropped')
            os.makedirs(crop_dir, exist_ok=True)
            crop_path = os.path.join(crop_dir, f'p{page_number:03d}_crop_{x}_{y}.png')
            cropped.save(crop_path)
            
            img_path = crop_path
        except ImportError:
            pass  # PIL not available, use full page image
        except Exception as e:
            print(f'    [!] Crop failed: {e}')
    
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(max_width_inches))
        p.paragraph_format.space_after = Pt(6)
        return p
    except Exception as e:
        print(f'    [!] Cannot add image: {e}')
        return None


def add_embedded_image(doc, grade, page_number, image_index, max_width_inches=5.0):
    """Thêm hình ảnh nhúng (extracted từ PDF) vào DOCX."""
    images_dir = os.path.join(SGK_DIR, f'Lớp_{grade}', 'ocr_output', 'images')
    
    # Tìm file ảnh matching page + index
    pattern = f'p{page_number:03d}_img{image_index:02d}_*'
    matches = glob.glob(os.path.join(images_dir, pattern))
    
    if not matches:
        return None
    
    img_path = matches[0]
    
    try:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_path, width=Inches(max_width_inches))
        p.paragraph_format.space_after = Pt(6)
        return p
    except Exception as e:
        print(f'    [!] Cannot add embedded image: {e}')
        return None


# ─── DOCX Builder ────────────────────────────────────────────────

def build_lesson_docx(grade, lesson_data, analysis_pages):
    """
    Tạo file DOCX cho 1 bài học.
    
    Args:
        grade: Khối lớp
        lesson_data: dict từ lesson_manifest
        analysis_pages: dict mapping page_number → analysis JSON
    
    Returns:
        str — đường dẫn file DOCX đã tạo
    """
    lesson_num = lesson_data['lesson_number']
    lesson_title = lesson_data['lesson_title']
    lesson_pages = lesson_data['pages']
    
    # Sanitize title cho filename
    safe_title = re.sub(r'[<>:"/\\|?*]', '', lesson_title)
    safe_title = safe_title.strip()[:50]
    
    # Create DOCX
    doc = Document()
    
    # Set default font
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    style.font.size = Pt(13)
    style.paragraph_format.space_after = Pt(6)
    
    # ── Header info ──
    add_paragraph_with_font(
        doc, f'SGK TIN HỌC {grade} — KẾT NỐI TRI THỨC',
        size=11, italic=True, color=(128, 128, 128),
        alignment=WD_ALIGN_PARAGRAPH.RIGHT
    )
    
    # ── Title ──
    add_heading_with_style(doc, f'Bài {lesson_num}. {lesson_title}', level=1)
    
    add_paragraph_with_font(
        doc, f'(Trang {lesson_pages[0]}–{lesson_pages[-1]} SGK)',
        size=11, italic=True, color=(100, 100, 100),
        alignment=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=Pt(12)
    )
    
    # ── Process elements from each page ──
    img_counter = 0
    
    for page_num in lesson_pages:
        page_data = analysis_pages.get(page_num)
        if not page_data:
            continue
        
        elements = page_data.get('elements', [])
        
        for elem in elements:
            elem_type = elem.get('type', '')
            text = elem.get('text', '').strip()
            
            if not text and elem_type != 'image_region':
                continue
            
            if elem_type == 'heading':
                level = min(elem.get('level', 2), 4)
                # Skip nếu trùng tiêu đề bài (đã add ở trên)
                if level == 1 and f'Bài {lesson_num}' in text:
                    continue
                add_heading_with_style(doc, text, level=level)
                
            elif elem_type == 'paragraph':
                add_paragraph_with_font(doc, text, size=13)
                
            elif elem_type == 'note':
                # Ghi nhớ / Highlight box
                p = add_paragraph_with_font(doc, '', size=13, bold=True)
                
                # Add border styling
                pPr = p._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                for edge in ['top', 'left', 'bottom', 'right']:
                    border = OxmlElement(f'w:{edge}')
                    border.set(qn('w:val'), 'single')
                    border.set(qn('w:sz'), '6')
                    border.set(qn('w:color'), '2E7D32')  # Green
                    border.set(qn('w:space'), '4')
                    pBdr.append(border)
                pPr.append(pBdr)
                
                # Add shading
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'E8F5E9')  # Light green
                pPr.append(shd)
                
                run = p.runs[0]
                run.text = text
                set_font(run, size=13, bold=True, color=(46, 125, 50))
                
            elif elem_type == 'exercise':
                # Bài tập
                ex_num = elem.get('exercise_number', '')
                prefix = f'Câu {ex_num}: ' if ex_num else ''
                p = add_paragraph_with_font(doc, f'{prefix}{text}', size=13)
                
                # Blue left border
                pPr = p._element.get_or_add_pPr()
                pBdr = OxmlElement('w:pBdr')
                left_border = OxmlElement('w:left')
                left_border.set(qn('w:val'), 'single')
                left_border.set(qn('w:sz'), '12')
                left_border.set(qn('w:color'), '1565C0')
                left_border.set(qn('w:space'), '8')
                pBdr.append(left_border)
                pPr.append(pBdr)
                
            elif elem_type == 'activity':
                # Hoạt động thực hành
                activity_type = elem.get('activity_type', 'practice')
                label_map = {
                    'practice': '🔧 THỰC HÀNH',
                    'explore': '🔍 KHÁM PHÁ', 
                    'apply': '💡 VẬN DỤNG',
                    'discuss': '💬 THẢO LUẬN'
                }
                label = label_map.get(activity_type, '📝 HOẠT ĐỘNG')
                
                add_paragraph_with_font(
                    doc, label, size=13, bold=True,
                    color=(21, 101, 192), space_after=Pt(3)
                )
                add_paragraph_with_font(doc, text, size=13)
                
            elif elem_type == 'image_region':
                # Hình ảnh minh hoạ
                img_counter += 1
                description = elem.get('description', f'Hình {img_counter}')
                bbox = elem.get('bbox')
                
                # Thử crop từ rendered page
                added = add_image_from_rendered_page(doc, grade, page_num, bbox, max_width_inches=5.0)
                
                # Thêm caption
                if added:
                    add_paragraph_with_font(
                        doc, description, size=10, italic=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        color=(100, 100, 100), space_after=Pt(10)
                    )
                else:
                    # Placeholder
                    add_paragraph_with_font(
                        doc, f'[Hình: {description}]', size=11, italic=True,
                        alignment=WD_ALIGN_PARAGRAPH.CENTER,
                        color=(150, 150, 150)
                    )
                
            elif elem_type == 'table':
                # Bảng
                caption = elem.get('caption', '')
                headers = elem.get('headers', [])
                rows = elem.get('rows', [])
                
                if caption:
                    add_paragraph_with_font(doc, caption, size=12, bold=True, italic=True)
                
                if headers and rows:
                    ncols = len(headers)
                    table = doc.add_table(rows=1 + len(rows), cols=ncols)
                    set_table_borders(table)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    
                    # Headers
                    for j, h in enumerate(headers):
                        cell = table.rows[0].cells[j]
                        cell.text = str(h)
                        for run in cell.paragraphs[0].runs:
                            set_font(run, size=12, bold=True)
                    
                    # Data rows
                    for i, row_data in enumerate(rows):
                        for j, val in enumerate(row_data):
                            if j < ncols:
                                cell = table.rows[i + 1].cells[j]
                                cell.text = str(val)
                                for run in cell.paragraphs[0].runs:
                                    set_font(run, size=12)
                
            elif elem_type == 'code_example':
                # Code / Scratch blocks
                language = elem.get('language', 'code')
                add_paragraph_with_font(
                    doc, f'[{language.upper()}]', size=10, bold=True,
                    color=(156, 39, 176)
                )
                
                # Monospace code block
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.font.name = 'Consolas'
                run.font.size = Pt(11)
                
                # Gray background
                pPr = p._element.get_or_add_pPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'F5F5F5')
                pPr.append(shd)
    
    # ── Page break + full-page images at the end ──
    # Add a section with full-page rendered images for reference
    doc.add_page_break()
    add_heading_with_style(doc, 'PHỤ LỤC: Trang SGK gốc', level=2)
    
    for page_num in lesson_pages:
        pages_dir = os.path.join(SGK_DIR, f'Lớp_{grade}', 'ocr_output', 'pages')
        img_path = os.path.join(pages_dir, f'page_{page_num:03d}.png')
        
        if os.path.exists(img_path):
            add_paragraph_with_font(
                doc, f'Trang {page_num}', size=10, bold=True,
                alignment=WD_ALIGN_PARAGRAPH.CENTER
            )
            try:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(6.0))
            except Exception:
                pass
    
    # ── Save ──
    docx_dir = os.path.join(SGK_DIR, f'Lớp_{grade}', 'ocr_output', 'docx')
    os.makedirs(docx_dir, exist_ok=True)
    
    filename = f'Bai_{lesson_num:02d}_{safe_title}.docx' if safe_title else f'Bai_{lesson_num:02d}.docx'
    out_path = os.path.join(docx_dir, filename)
    
    try:
        doc.save(out_path)
        return out_path
    except PermissionError:
        # File đang mở trong Word
        alt_path = out_path.replace('.docx', '_new.docx')
        doc.save(alt_path)
        print(f'    [!] File đang mở, lưu vào: {alt_path}')
        return alt_path


def process_grade(grade, lesson_range=None):
    """Xử lý Stage 4 cho 1 khối lớp."""
    grade_dir = os.path.join(SGK_DIR, f'Lớp_{grade}', 'ocr_output')
    
    # Load lesson manifest
    manifest_path = os.path.join(grade_dir, 'lesson_manifest.json')
    if not os.path.exists(manifest_path):
        print(f'  [!] Chưa có lesson_manifest.json — Chạy Stage 3 trước')
        return None
    
    with open(manifest_path, 'r', encoding='utf-8') as f:
        manifest = json.load(f)
    
    # Load analysis pages
    analysis_dir = os.path.join(grade_dir, 'analysis')
    analysis_pages = {}
    for fpath in glob.glob(os.path.join(analysis_dir, 'page_*.json')):
        with open(fpath, 'r', encoding='utf-8') as f:
            try:
                data = json.load(f)
                pnum = data.get('_meta', {}).get('page_number', 0)
                if pnum:
                    analysis_pages[pnum] = data
            except json.JSONDecodeError:
                pass
    
    print(f'\n{"=" * 60}')
    print(f'  STAGE 4 — LỚP {grade}: DOCX Builder')
    print(f'{"=" * 60}')
    print(f'  Total lessons: {manifest["total_lessons_found"]}')
    print(f'  Analysis pages loaded: {len(analysis_pages)}')
    
    # Filter lessons
    lessons = manifest.get('lessons', [])
    if lesson_range:
        lessons = [l for l in lessons if lesson_range[0] <= l['lesson_number'] <= lesson_range[1]]
    
    results = []
    for lesson in lessons:
        lesson_num = lesson['lesson_number']
        lesson_title = lesson['lesson_title']
        
        print(f'\n  Bài {lesson_num}: {lesson_title}', end='', flush=True)
        
        out_path = build_lesson_docx(grade, lesson, analysis_pages)
        
        if out_path:
            fsize = os.path.getsize(out_path) // 1024
            print(f' → ✓ {os.path.basename(out_path)} ({fsize}KB)')
            results.append({
                'lesson_number': lesson_num,
                'lesson_title': lesson_title,
                'docx_path': out_path,
                'file_size_kb': fsize
            })
        else:
            print(f' → ✗ Failed')
    
    # Save build manifest
    build_manifest = {
        'grade': grade,
        'total_docx_created': len(results),
        'files': results
    }
    
    build_path = os.path.join(grade_dir, 'docx', 'build_manifest.json')
    os.makedirs(os.path.dirname(build_path), exist_ok=True)
    with open(build_path, 'w', encoding='utf-8') as f:
        json.dump(build_manifest, f, ensure_ascii=False, indent=2)
    
    print(f'\n  ✓ Created {len(results)} DOCX files')
    return build_manifest


def parse_lesson_range(s):
    if '-' in s:
        parts = s.split('-')
        return (int(parts[0]), int(parts[1]))
    else:
        n = int(s)
        return (n, n)


def main():
    parser = argparse.ArgumentParser(description='SGK DOCX Builder — Stage 4')
    parser.add_argument('--grade', type=int, choices=[3, 4, 5, 6, 7, 8],
                        help='Khối lớp')
    parser.add_argument('--all', action='store_true',
                        help='Tạo DOCX cho tất cả các lớp')
    parser.add_argument('--lessons', type=str, default=None,
                        help='Phạm vi bài (VD: 1-5 hoặc 3)')
    
    args = parser.parse_args()
    
    lesson_range = parse_lesson_range(args.lessons) if args.lessons else None
    
    if args.all:
        grades = [3, 4, 5, 6, 7, 8]
    elif args.grade:
        grades = [args.grade]
    else:
        parser.print_help()
        return
    
    for grade in grades:
        process_grade(grade, lesson_range)
    
    print(f'\n{"=" * 60}')
    print('  STAGE 4 COMPLETE!')
    print(f'{"=" * 60}')


if __name__ == '__main__':
    main()
