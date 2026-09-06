import os
import sys
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

sys.stdout.reconfigure(encoding='utf-8')

# Paths
SOURCE_EXCEL = r'D:\UNIGO\TKB toàn trường CHECK - 26.8.xlsx'
OUTPUT_DIR_MAIN = r'D:\UNIGO\Thời khóa biểu giáo viên'
OUTPUT_DIR_MAU = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Thời khóa biểu'

os.makedirs(OUTPUT_DIR_MAIN, exist_ok=True)
os.makedirs(OUTPUT_DIR_MAU, exist_ok=True)

# ─────────────── THỜI GIAN TIẾT HỌC ───────────────────────────────
TIMES_SANG_TH    = {1: "8:15–8:50", 2: "8:55–9:30", 3: "9:45–10:20", 4: "10:25–11:00", 5: ""}
TIMES_SANG_THCS  = {1: "8:05–8:50", 2: "8:55–9:40", 3: "9:45–10:25", 4: "10:30–11:10", 5: "11:15–11:50"}

TIMES_CHIEU_TH   = {1: "13:15–13:50", 2: "13:55–14:30", 3: "14:55–15:30", 4: "15:35–16:10", 5: ""}
TIMES_CHIEU_THCS = {1: "13:05–13:50", 2: "13:55–14:40", 3: "14:50–15:30", 4: "15:35–16:20", 5: ""}

# ─────────────── MÀU SẮC CHUNG ───────────────────────────────────
CLR_HEADER_BG  = "1D2A64"   # Navy đậm UNIGO
CLR_SESSION_BG = "E8ECF5"   # Xám xanh nhạt thanh tiêu đề Buổi
CLR_PERIOD_BG  = "F5F7FF"   # Cột Tiết
CLR_TIME_BG    = "EFF3FF"   # Cột Giờ

DOCX_HEADER_BG = "1D2A64"
DOCX_SESSION   = "E8ECF5"
DOCX_PERIOD    = "F5F7FF"
DOCX_TIME      = "EFF3FF"

THIN_SIDE = Side(style='thin', color='AAAACC')
THICK_SIDE = Side(style='medium', color='1D2A64')
INNER_BORDER = Border(left=THIN_SIDE, right=THIN_SIDE, top=THIN_SIDE, bottom=THIN_SIDE)

def xl_fill(hex_color):
    return PatternFill(start_color=hex_color, end_color=hex_color, fill_type="solid")

def xl_font(size=12, bold=False, italic=False, color="000000", name='Times New Roman'):
    return Font(name=name, size=size, bold=bold, italic=italic, color=color)

def xl_align(h='center', v='center', wrap=True):
    return Alignment(horizontal=h, vertical=v, wrap_text=wrap)

def set_tbl_borders(table):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        for child in list(tblPr[0]):
            if child.tag.endswith('tblBorders'):
                tblPr[0].remove(child)
        borders_xml = parse_xml(
            r'<w:tblBorders %s>'
            r'<w:top w:val="single" w:sz="6" w:space="0" w:color="1D2A64"/>'
            r'<w:left w:val="single" w:sz="6" w:space="0" w:color="1D2A64"/>'
            r'<w:bottom w:val="single" w:sz="6" w:space="0" w:color="1D2A64"/>'
            r'<w:right w:val="single" w:sz="6" w:space="0" w:color="1D2A64"/>'
            r'<w:insideH w:val="single" w:sz="4" w:space="0" w:color="AAAACC"/>'
            r'<w:insideV w:val="single" w:sz="4" w:space="0" w:color="AAAACC"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        tblPr[0].append(borders_xml)

def set_cell_bg(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_vert_align(cell, align="center"):
    tcPr = cell._element.get_or_add_tcPr()
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), align)
    tcPr.append(vAlign)

def add_run(para, text, font_name='Times New Roman', size=12, bold=False,
            italic=False, color=None):
    run = para.add_run(text)
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)
    return run

def style_header_cell(cell, text, bg_hex, fg=(255,255,255), size=12, bold=True):
    set_cell_bg(cell, bg_hex)
    set_cell_vert_align(cell)
    para = cell.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.bold = bold
    if fg != (0, 0, 0):
        run.font.color.rgb = RGBColor(*fg)

# ─────────────── PARSING TKB TỪ EXCEL GỐC ───────────────────────────────
def parse_teacher_schedule(teacher_keyword, teacher_name, subject_title, config_subjects):
    """
    config_subjects: dict of {
        'Tin học': {'label': '💻 Tin học', 'bg_hex': 'D6E4FF', 'fg_hex': '003C8C', 'fg_rgb': (0, 60, 140), 'desc': '🔵 Xanh nhạt: Tin học'},
        'Robotics': {'label': '🤖 Robotics', 'bg_hex': 'FFE5CC', 'fg_hex': '8C3C00', 'fg_rgb': (140, 60, 0), 'desc': '🟠 Cam nhạt: Robotics'}
    }
    """
    wb = openpyxl.load_workbook(SOURCE_EXCEL, data_only=True)
    ws = wb['TKB_LOP_SC']

    col_to_info = {}
    for c in range(1, ws.max_column + 1):
        classname = ''
        for back_c in range(c, 0, -1):
            v = ws.cell(4, back_c).value
            if v and str(v).strip() not in ['THỨ', 'TIẾT', 'THU', 'TIET', '']:
                classname = str(v).strip()
                break
        sess = str(ws.cell(5, c).value or '').strip()
        if sess in ['Sáng', 'Chiều']:
            col_to_info[c] = (classname, sess)

    days_key = ['2', '3', '4', '5', '6']
    days_name = ['Thứ Hai', 'Thứ Ba', 'Thứ Tư', 'Thứ Năm', 'Thứ Sáu']

    slots = []
    for d_idx, start_r in enumerate([6, 11, 16, 21, 26]):
        day_k = days_key[d_idx]
        for tiet in range(1, 6):
            r = start_r + tiet - 1
            for c, (cname, sess) in col_to_info.items():
                val = str(ws.cell(r, c).value or '').strip()
                if teacher_keyword in val:
                    # Determine subject
                    matched_subj = None
                    for subj_k in config_subjects.keys():
                        if subj_k in val:
                            matched_subj = subj_k
                            break
                    if not matched_subj:
                        # default to first configured subject
                        matched_subj = list(config_subjects.keys())[0]
                    
                    is_thcs = any(x in cname for x in ['6', '7', '8'])
                    slots.append({
                        'day_key': day_k,
                        'day_name': days_name[d_idx],
                        'day_idx': d_idx,
                        'session': sess,
                        'tiet': tiet,
                        'subject_key': matched_subj,
                        'class_name': cname,
                        'is_thcs': is_thcs,
                        'raw': val
                    })

    return {
        'keyword': teacher_keyword,
        'full_name': teacher_name,
        'subject_title': subject_title,
        'config_subjects': config_subjects,
        'slots': slots,
        'total_periods': len(slots)
    }

# ─────────────── TẠO FILE DOCX ──────────────────────────────────────────
def export_docx(data, save_paths):
    doc = docx.Document()
    sec = doc.sections[0]
    
    # Landscape A4
    sec.page_width  = int(29.7 * 360000)
    sec.page_height = int(21.0 * 360000)
    sec.left_margin   = int(1.2 * 360000)
    sec.right_margin  = int(1.2 * 360000)
    sec.top_margin    = int(1.2 * 360000)
    sec.bottom_margin = int(1.2 * 360000)
    
    sectPr = sec._sectPr
    pgSz = sectPr.find(qn('w:pgSz'))
    if pgSz is None:
        pgSz = OxmlElement('w:pgSz')
        sectPr.append(pgSz)
    pgSz.set(qn('w:orient'), 'landscape')

    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    add_run(p, "TRƯỜNG TIỂU HỌC VÀ THCS UNIGO\n", size=13, bold=True, color=(29,42,100))
    add_run(p, "THỜI KHOÁ BIỂU CÁ NHÂN GIÁO VIÊN\n", size=16, bold=True, color=(29,42,100))
    add_run(p, "Năm học 2026 – 2027", size=12, italic=True, color=(80,80,80))

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after  = Pt(8)
    add_run(p2, f"Giáo viên: {data['full_name']}   |   Môn: {data['subject_title']}   |   Tổng: {data['total_periods']} tiết/tuần",
            size=13, bold=True, color=(29,42,100))

    # Bảng: 13 hàng × 8 cột
    table = doc.add_table(rows=13, cols=8)
    table.style = 'Table Grid'
    set_tbl_borders(table)

    COL_HEADERS = ["Tiết", "Thời gian\n(Tiểu học)", "Thời gian\n(THCS)", "Thứ Hai", "Thứ Ba", "Thứ Tư", "Thứ Năm", "Thứ Sáu"]

    for ci, hdr in enumerate(COL_HEADERS):
        c = table.rows[0].cells[ci]
        style_header_cell(c, hdr, DOCX_HEADER_BG, fg=(255,255,255), size=11, bold=True)

    def add_session_row(row_idx, label):
        cell0 = table.rows[row_idx].cells[0]
        for ci in range(1, 8):
            cell0.merge(table.rows[row_idx].cells[ci])
        style_header_cell(cell0, label, DOCX_SESSION, fg=(29,42,100), size=12, bold=True)

    add_session_row(1, "▸  BUỔI SÁNG")
    add_session_row(7, "▸  BUỔI CHIỀU")

    # Fill Tiết & Time
    for p_num in range(1, 6):
        # Sáng – rows 2..6
        row_s = table.rows[p_num + 1]
        set_cell_bg(row_s.cells[0], DOCX_PERIOD)
        set_cell_bg(row_s.cells[1], DOCX_TIME)
        set_cell_bg(row_s.cells[2], DOCX_TIME)
        for ci in (0, 1, 2):
            set_cell_vert_align(row_s.cells[ci])
            row_s.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(row_s.cells[0].paragraphs[0], f"Tiết {p_num}", size=12, bold=True, color=(29,42,100))
        add_run(row_s.cells[1].paragraphs[0], TIMES_SANG_TH.get(p_num, ""), size=11, italic=True, color=(60,60,80))
        add_run(row_s.cells[2].paragraphs[0], TIMES_SANG_THCS.get(p_num, ""), size=11, italic=True, color=(60,60,80))

        # Chiều – rows 8..12
        row_c = table.rows[p_num + 7]
        set_cell_bg(row_c.cells[0], DOCX_PERIOD)
        set_cell_bg(row_c.cells[1], DOCX_TIME)
        set_cell_bg(row_c.cells[2], DOCX_TIME)
        for ci in (0, 1, 2):
            set_cell_vert_align(row_c.cells[ci])
            row_c.cells[ci].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(row_c.cells[0].paragraphs[0], f"Tiết {p_num}", size=12, bold=True, color=(29,42,100))
        add_run(row_c.cells[1].paragraphs[0], TIMES_CHIEU_TH.get(p_num, ""), size=11, italic=True, color=(60,60,80))
        add_run(row_c.cells[2].paragraphs[0], TIMES_CHIEU_THCS.get(p_num, ""), size=11, italic=True, color=(60,60,80))

    # Fill Teaching Slots
    day_to_col = {'2': 3, '3': 4, '4': 5, '5': 6, '6': 7}
    for slot in data['slots']:
        ci = day_to_col[slot['day_key']]
        ri = (slot['tiet'] + 1) if slot['session'] == 'Sáng' else (slot['tiet'] + 7)

        cell = table.rows[ri].cells[ci]
        cfg = data['config_subjects'].get(slot['subject_key'], list(data['config_subjects'].values())[0])
        
        set_cell_bg(cell, cfg['bg_hex'])
        set_cell_vert_align(cell)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(cell.paragraphs[0], cfg['label'] + "\n", size=11, bold=True, color=cfg['fg_rgb'])
        add_run(cell.paragraphs[0], f"({slot['class_name']})", size=10, bold=False, color=(80,80,80))

    for path in save_paths:
        try:
            doc.save(path)
            print(f"✓ Đã lưu DOCX: {path}")
        except PermissionError:
            print(f"✗ Lỗi quyền ghi DOCX: {path} (vui lòng đóng file nếu đang mở trong Word)")

# ─────────────── TẠO FILE EXCEL ──────────────────────────────────────────
def export_excel(data, save_path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"TKB {data['full_name']}"

    # Header block
    ws.merge_cells('A1:H1')
    ws['A1'] = "TRƯỜNG TIỂU HỌC VÀ THCS UNIGO"
    ws['A1'].font = xl_font(14, bold=True, color="FFFFFF")
    ws['A1'].fill = xl_fill("1D2A64")
    ws['A1'].alignment = xl_align()

    ws.merge_cells('A2:H2')
    ws['A2'] = "THỜI KHOÁ BIỂU CÁ NHÂN GIÁO VIÊN  —  Năm học 2026–2027"
    ws['A2'].font = xl_font(15, bold=True, color="FFFFFF")
    ws['A2'].fill = xl_fill("253580")
    ws['A2'].alignment = xl_align()

    ws.merge_cells('A3:H3')
    ws['A3'] = f"Giáo viên: {data['full_name']}   |   Bộ môn: {data['subject_title']}   |   Tổng: {data['total_periods']} tiết/tuần"
    ws['A3'].font = xl_font(12, italic=True, color="1D2A64")
    ws['A3'].fill = xl_fill("EBF0FF")
    ws['A3'].alignment = xl_align()

    ws.row_dimensions[1].height = 28
    ws.row_dimensions[2].height = 30
    ws.row_dimensions[3].height = 24

    # Legend block
    subj_keys = list(data['config_subjects'].keys())
    if len(subj_keys) >= 2:
        k1, k2 = subj_keys[0], subj_keys[1]
        cfg1 = data['config_subjects'][k1]
        cfg2 = data['config_subjects'][k2]

        ws.merge_cells('A4:D4')
        ws['A4'] = cfg1['desc']
        ws['A4'].font = xl_font(10, color=cfg1['fg_hex'])
        ws['A4'].fill = xl_fill(cfg1['bg_hex'])
        ws['A4'].alignment = Alignment(horizontal='center')

        ws.merge_cells('E4:H4')
        ws['E4'] = cfg2['desc']
        ws['E4'].font = xl_font(10, color=cfg2['fg_hex'])
        ws['E4'].fill = xl_fill(cfg2['bg_hex'])
        ws['E4'].alignment = Alignment(horizontal='center')
    elif len(subj_keys) == 1:
        k1 = subj_keys[0]
        cfg1 = data['config_subjects'][k1]
        ws.merge_cells('A4:H4')
        ws['A4'] = cfg1['desc']
        ws['A4'].font = xl_font(10, color=cfg1['fg_hex'])
        ws['A4'].fill = xl_fill(cfg1['bg_hex'])
        ws['A4'].alignment = Alignment(horizontal='center')

    ws.row_dimensions[4].height = 18

    # Column headers
    col_headers = ["Tiết", "Giờ TH", "Giờ THCS", "Thứ Hai", "Thứ Ba", "Thứ Tư", "Thứ Năm", "Thứ Sáu"]
    for ci, h in enumerate(col_headers, 1):
        c = ws.cell(5, ci, h)
        c.font = xl_font(12, bold=True, color="FFFFFF")
        c.fill = xl_fill("1D2A64")
        c.alignment = xl_align()
        c.border = INNER_BORDER
    ws.row_dimensions[5].height = 30

    def draw_session(row, label):
        ws.merge_cells(f'A{row}:H{row}')
        c = ws[f'A{row}']
        c.value = label
        c.font = xl_font(12, bold=True, color="1D2A64")
        c.fill = xl_fill(CLR_SESSION_BG)
        c.alignment = xl_align()
        ws.row_dimensions[row].height = 20
        for ci in range(1, 9):
            ws.cell(row, ci).border = INNER_BORDER

    # Buổi sáng
    draw_session(6, "▸   BUỔI SÁNG")
    for p_num in range(1, 6):
        ri = 6 + p_num
        ws.row_dimensions[ri].height = 40
        c_tiet = ws.cell(ri, 1, f"Tiết {p_num}")
        c_tiet.font = xl_font(12, bold=True, color="1D2A64")
        c_tiet.fill = xl_fill(CLR_PERIOD_BG)
        c_tiet.alignment = xl_align()
        c_tiet.border = INNER_BORDER

        c_th = ws.cell(ri, 2, TIMES_SANG_TH.get(p_num, ""))
        c_th.font = xl_font(10, italic=True, color="404060")
        c_th.fill = xl_fill(CLR_TIME_BG)
        c_th.alignment = xl_align()
        c_th.border = INNER_BORDER

        c_thcs = ws.cell(ri, 3, TIMES_SANG_THCS.get(p_num, ""))
        c_thcs.font = xl_font(10, italic=True, color="404060")
        c_thcs.fill = xl_fill(CLR_TIME_BG)
        c_thcs.alignment = xl_align()
        c_thcs.border = INNER_BORDER

        for ci in range(4, 9):
            cc = ws.cell(ri, ci)
            cc.border = INNER_BORDER
            cc.alignment = xl_align()

    # Buổi chiều
    draw_session(12, "▸   BUỔI CHIỀU")
    for p_num in range(1, 6):
        ri = 12 + p_num
        ws.row_dimensions[ri].height = 40
        c_tiet = ws.cell(ri, 1, f"Tiết {p_num}")
        c_tiet.font = xl_font(12, bold=True, color="1D2A64")
        c_tiet.fill = xl_fill(CLR_PERIOD_BG)
        c_tiet.alignment = xl_align()
        c_tiet.border = INNER_BORDER

        c_th = ws.cell(ri, 2, TIMES_CHIEU_TH.get(p_num, ""))
        c_th.font = xl_font(10, italic=True, color="404060")
        c_th.fill = xl_fill(CLR_TIME_BG)
        c_th.alignment = xl_align()
        c_th.border = INNER_BORDER

        c_thcs = ws.cell(ri, 3, TIMES_CHIEU_THCS.get(p_num, ""))
        c_thcs.font = xl_font(10, italic=True, color="404060")
        c_thcs.fill = xl_fill(CLR_TIME_BG)
        c_thcs.alignment = xl_align()
        c_thcs.border = INNER_BORDER

        for ci in range(4, 9):
            cc = ws.cell(ri, ci)
            cc.border = INNER_BORDER
            cc.alignment = xl_align()

    # Fill Teaching Slots
    day_to_col = {'2': 4, '3': 5, '4': 6, '5': 7, '6': 8}
    for slot in data['slots']:
        ci = day_to_col[slot['day_key']]
        ri = (slot['tiet'] + 6) if slot['session'] == 'Sáng' else (slot['tiet'] + 12)

        cell = ws.cell(ri, ci)
        cfg = data['config_subjects'].get(slot['subject_key'], list(data['config_subjects'].values())[0])

        cell.value = f"{cfg['label']}\n({slot['class_name']})"
        cell.font = xl_font(11, bold=True, color=cfg['fg_hex'])
        cell.fill = xl_fill(cfg['bg_hex'])
        cell.alignment = xl_align()
        cell.border = INNER_BORDER

    # Column widths
    col_widths = {'A': 9, 'B': 13, 'C': 13, 'D': 16, 'E': 16, 'F': 16, 'G': 16, 'H': 16}
    for col, width in col_widths.items():
        ws.column_dimensions[col].width = width

    try:
        wb.save(save_path)
        print(f"✓ Đã lưu XLSX: {save_path}")
    except PermissionError:
        print(f"✗ Lỗi quyền ghi XLSX: {save_path} (vui lòng đóng file nếu đang mở trong Excel)")

# ─────────────── CHẠY TOÀN BỘ ───────────────────────────────────────────
if __name__ == '__main__':
    print("==========================================================")
    print("TRÍCH XUẤT VÀ TẠO THỜI KHÓA BIỂU CÁ NHÂN GIÁO VIÊN UNIGO")
    print(f"Nguồn: {SOURCE_EXCEL}")
    print("==========================================================")

    # 1. Thầy Đậu Đình Nguyên
    cfg_nguyen = {
        'Tin': {
            'label': '💻 Tin học',
            'bg_hex': 'D6E4FF',
            'fg_hex': '003C8C',
            'fg_rgb': (0, 60, 140),
            'desc': '🔵 Xanh nhạt: Tin học'
        },
        'Robotics': {
            'label': '🤖 Robotics',
            'bg_hex': 'FFE5CC',
            'fg_hex': '8C3C00',
            'fg_rgb': (140, 60, 0),
            'desc': '🟠 Cam nhạt: Robotics'
        }
    }
    data_nguyen = parse_teacher_schedule(
        teacher_keyword='Nguyên',
        teacher_name='Đậu Đình Nguyên',
        subject_title='Tin học & Robotics',
        config_subjects=cfg_nguyen
    )
    print(f"\n[1] ĐẬU ĐÌNH NGUYÊN: Tổng {data_nguyen['total_periods']} tiết")
    export_docx(data_nguyen, [
        os.path.join(OUTPUT_DIR_MAIN, 'Thời khóa biểu - Đậu Đình Nguyên.docx'),
        os.path.join(OUTPUT_DIR_MAU, 'Thời khóa biểu - Đậu Đình Nguyên.docx')
    ])
    export_excel(data_nguyen, os.path.join(OUTPUT_DIR_MAIN, 'Thời khóa biểu - Đậu Đình Nguyên.xlsx'))

    # 2. Cô Nguyệt
    cfg_nguyet = {
        'Toán': {
            'label': '📐 Toán',
            'bg_hex': 'D6E4FF',
            'fg_hex': '003C8C',
            'fg_rgb': (0, 60, 140),
            'desc': '🔵 Xanh nhạt: Toán'
        },
        'HĐTN': {
            'label': '🌟 HĐTN',
            'bg_hex': 'FFE5CC',
            'fg_hex': '8C3C00',
            'fg_rgb': (140, 60, 0),
            'desc': '🟠 Cam nhạt: Hoạt động trải nghiệm'
        }
    }
    data_nguyet = parse_teacher_schedule(
        teacher_keyword='Nguyệt',
        teacher_name='Cô Nguyệt',
        subject_title='Toán & HĐTN',
        config_subjects=cfg_nguyet
    )
    print(f"\n[2] CÔ NGUYỆT: Tổng {data_nguyet['total_periods']} tiết")
    export_docx(data_nguyet, [
        os.path.join(OUTPUT_DIR_MAIN, 'Thời khóa biểu - Cô Nguyệt.docx'),
        os.path.join(OUTPUT_DIR_MAU, 'Thời khóa biểu - Cô Nguyệt.docx')
    ])
    export_excel(data_nguyet, os.path.join(OUTPUT_DIR_MAIN, 'Thời khóa biểu - Cô Nguyệt.xlsx'))
    print("\n================== HOÀN THÀNH ==================")
