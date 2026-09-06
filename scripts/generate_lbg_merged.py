"""
generate_lbg_merged.py — Gop tat ca 35 tuan Lich bao giang vao 3 file duy nhat
================================================================================

OUTPUT:
  Lich bao giang - Ca nam.docx            <- ban goc tong hop 35 tuan
  Lich bao giang - Ca nam (TTH+TH).docx   <- Tien TH + TH, 35 tuan
  Lich bao giang - Ca nam (THCS).docx     <- THCS, 35 tuan

QUY TRINH:
  1. Tao tung tuan (1..35) bang ham generate_week_doc() (dua tren logic generate_lbg.py)
  2. Gop tat ca cac tuan vao 1 document, chen section break (new page) giua cac tuan
  3. Tach ban TTH+TH va THCS

CACH DUNG:
  python generate_lbg_merged.py
  python generate_lbg_merged.py 10       # Chi gop 10 tuan dau (test nhanh)
"""

import sys
import os
import copy
import re
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import date, timedelta

sys.stdout.reconfigure(encoding='utf-8')

# Import all utilities and data from generate_lbg
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from generate_lbg import (
    LBG_DIR, TEMPLATE, TUAN_01_START,
    week_dates, fmt_date, day_label,
    classify_lop, is_rotation_lop, rotation_mon, rotation_do_dung,
    get_grade_key, get_ppct_lesson, compute_ppct, compute_rotation_ppct,
    set_cell_text, save_safe,
    update_headers, update_table_data, update_ky_ten,
    remove_second_copy, fix_all_fonts,
    filter_for_cap, fix_page_breaks,
    TO_TRUONG_TH, TO_TRUONG_THCS,
    update_sign_names, compact_sign_tables,
    remove_cover_page, remove_end_evaluation_table,
)

TOTAL_WEEKS = 35


# ─────────────────────────────────────────────────────────────────────────────
# GENERATE 1 WEEK AS DOCUMENT (in memory)
# ─────────────────────────────────────────────────────────────────────────────

def generate_week_doc(tuan_so):
    """Tao Document cho 1 tuan (khong luu file)."""
    start_date, end_date = week_dates(tuan_so)
    loai = 'CHẴN → Tin học' if tuan_so % 2 == 0 else 'LẺ → Robotics'

    print(f'  📅 Tuần {tuan_so:02d}: {fmt_date(start_date)} -> {fmt_date(end_date)}  [{loai}]')

    doc = Document(TEMPLATE)
    remove_second_copy(doc)

    # File gộp cả năm: TẤT CẢ các tuần (kể cả tuần 1) đều bỏ trang bìa đầu và bảng nhận xét cuối
    remove_cover_page(doc)
    remove_end_evaluation_table(doc)

    update_headers(doc, tuan_so, start_date, end_date)
    update_table_data(doc, tuan_so, start_date)
    fix_all_fonts(doc)
    update_ky_ten(doc, start_date)
    update_sign_names(doc, '', start_date)
    compact_sign_tables(doc)
    fix_page_breaks(doc)

    return doc


# ─────────────────────────────────────────────────────────────────────────────
# MERGE UTILITIES
# ─────────────────────────────────────────────────────────────────────────────

def add_section_break_page(body, template_sectpr=None):
    """Chen mot section break (new page) vao cuoi body.
    Copy page size, margins va header/footer refs tu template_sectpr.
    """
    p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    sectPr = OxmlElement('w:sectPr')

    if template_sectpr is not None:
        # Copy header/footer references
        for child in template_sectpr:
            tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
            if tag in ('headerReference', 'footerReference'):
                sectPr.append(copy.deepcopy(child))

        # Copy pgSz
        pgSz_src = template_sectpr.find(qn('w:pgSz'))
        if pgSz_src is not None:
            sectPr.append(copy.deepcopy(pgSz_src))
        else:
            pgSz = OxmlElement('w:pgSz')
            pgSz.set(qn('w:w'), '11910')
            pgSz.set(qn('w:h'), '16840')
            sectPr.append(pgSz)

        # Copy pgMar
        pgMar_src = template_sectpr.find(qn('w:pgMar'))
        if pgMar_src is not None:
            sectPr.append(copy.deepcopy(pgMar_src))
    else:
        # Fallback: Portrait A4 with template margins
        pgSz = OxmlElement('w:pgSz')
        pgSz.set(qn('w:w'), '11910')
        pgSz.set(qn('w:h'), '16840')
        sectPr.append(pgSz)

        pgMar = OxmlElement('w:pgMar')
        pgMar.set(qn('w:top'), '184')
        pgMar.set(qn('w:right'), '720')
        pgMar.set(qn('w:bottom'), '280')
        pgMar.set(qn('w:left'), '1180')
        pgMar.set(qn('w:header'), '283')
        pgMar.set(qn('w:footer'), '0')
        pgMar.set(qn('w:gutter'), '0')
        sectPr.append(pgMar)

    pPr.append(sectPr)
    p.append(pPr)
    body.append(p)


def get_body_elements_without_sectpr(doc):
    """Lay tat ca body elements tru sectPr cuoi cung."""
    body = doc.element.body
    children = list(body)
    result = []
    for child in children:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'sectPr':
            continue
        result.append(child)
    return result


def copy_sectpr(source_doc):
    """Copy sectPr tu source doc."""
    body = source_doc.element.body
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'sectPr':
            return copy.deepcopy(child)
    return None


def merge_weeks(week_docs):
    """
    Gộp 35 tuần vào 1 Document duy nhất:
    - Trang 1: Trang bìa đầu (Table 0: Thông tin giáo viên)
    - Section break (new page)
    - Các tuần từ 1 đến 35 (mỗi tuần gồm Sáng + Chiều, phân cách nhau bằng section break new page)
    - Sau tuần 35: Section break (new page)
    - Trang cuối: Bảng nhận xét tiến độ thực hiện chương trình của BGH (Table 5)
    """
    if not week_docs:
        return Document()

    # Lấy cover, eval elements và sectPr từ TEMPLATE
    doc_tmpl = Document(TEMPLATE)
    tmpl_body = doc_tmpl.element.body

    # 1. Trích xuất cover elements (Table 0 + đoạn p trước tiêu đề TUẦN)
    tuan_el = None
    for c in list(tmpl_body):
        if c.tag.endswith('p'):
            txt = ''.join(c.itertext()).strip()
            if 'TUẦN' in txt.upper() and 'từ ngày' in txt:
                tuan_el = c
                break
    cover_elements = []
    if tuan_el is not None:
        idx = list(tmpl_body).index(tuan_el)
        for el in list(tmpl_body)[:idx]:
            if not el.tag.endswith('sectPr'):
                cover_elements.append(copy.deepcopy(el))

    # 2. Trích xuất eval elements (tiêu đề Nhận xét BGH + Table 5 + đoạn sau)
    nhanxet_el = None
    for c in list(tmpl_body):
        if c.tag.endswith('p'):
            txt = ''.join(c.itertext()).strip()
            if 'nhận xét' in txt.lower() and 'bgh' in txt.lower():
                nhanxet_el = c
                break
    eval_elements = []
    if nhanxet_el is not None:
        idx = list(tmpl_body).index(nhanxet_el)
        for el in list(tmpl_body)[idx:]:
            if not el.tag.endswith('sectPr'):
                eval_elements.append(copy.deepcopy(el))

    # 3. Trích xuất base_sectpr
    base_sectpr = copy_sectpr(doc_tmpl)

    # 4. Tạo document merged rỗng (giữ lại sectPr cuối)
    merged = Document(TEMPLATE)
    merged_body = merged.element.body
    for child in list(merged_body):
        if not child.tag.endswith('sectPr'):
            merged_body.remove(child)

    final_sectpr = list(merged_body)[0]
    merged_body.remove(final_sectpr)

    # A. Thêm Trang bìa đầu
    for el in cover_elements:
        merged_body.append(copy.deepcopy(el))

    # B. Thêm từng tuần (mỗi tuần có section break new page phía trước)
    for i, wdoc in enumerate(week_docs, start=1):
        print(f'    Merging tuần {i:02d}...')
        add_section_break_page(merged_body, base_sectpr)
        elements = get_body_elements_without_sectpr(wdoc)
        for elem in elements:
            merged_body.append(copy.deepcopy(elem))

    # C. Thêm Section break (new page) và Bảng nhận xét cuối của BGH
    if eval_elements:
        add_section_break_page(merged_body, base_sectpr)
        for el in eval_elements:
            merged_body.append(copy.deepcopy(el))

    # D. Đảm bảo final_sectpr ở cuối body
    merged_body.append(final_sectpr)

    return merged


# ─────────────────────────────────────────────────────────────────────────────
# FILTER + FIX PAGE BREAKS FOR MERGED DOC
# ─────────────────────────────────────────────────────────────────────────────

def filter_merged_for_cap(doc, keep_cap):
    """Filter toan bo tables trong merged doc (nhieu tuan)."""
    for tbl in doc.tables:
        # Chi xu ly bang 7 cot (bang LBG chinh)
        if len(tbl.columns) != 7:
            continue
        # Kiem tra header
        header_texts = [c.text.strip() for c in tbl.rows[0].cells]
        if 'Lớp' not in header_texts:
            continue
        # Tim cot Lop (thuong la cot 4)
        lop_col = header_texts.index('Lớp')
        for ri in range(1, len(tbl.rows)):
            row = tbl.rows[ri]
            lop = row.cells[lop_col].text.strip()
            # Bỏ qua nếu là ô nghỉ lễ hoặc tổng duyệt / sự kiện
            if any(kw in lop.upper() for kw in ['NGHỈ', 'LỄ', 'TỔNG DUYỆT', 'KHAI GIẢNG']):
                continue
            cap = classify_lop(lop)
            if cap is not None and cap != keep_cap:
                # Clear data columns (2..6)
                for ci in range(2, min(7, len(row.cells))):
                    for p in row.cells[ci].paragraphs:
                        for r in p.runs:
                            r.text = ''
                        if not p.runs and p.text.strip():
                            p.text = ''


# ─────────────────────────────────────────────────────────────────────────────
# HAM CHINH
# ─────────────────────────────────────────────────────────────────────────────

def main():
    max_weeks = TOTAL_WEEKS
    if len(sys.argv) >= 2:
        try:
            max_weeks = int(sys.argv[1])
        except ValueError:
            pass
    max_weeks = min(max_weeks, TOTAL_WEEKS)

    print(f'\n{"="*60}')
    print(f'  GENERATE LBG MERGED — {max_weeks} tuần')
    print(f'{"="*60}')

    # ── BUOC 1: Tao document cho tung tuan ──
    print('\n[1/3] Tạo document cho từng tuần...')
    week_docs = []
    for tuan in range(1, max_weeks + 1):
        wdoc = generate_week_doc(tuan)
        week_docs.append(wdoc)

    # ── BUOC 2: Merge vao 1 file ──
    print('\n[2/3] Gộp tất cả tuần vào 1 file...')
    merged = merge_weeks(week_docs)

    # Luu ban tong hop
    main_name = f'Lịch báo giảng - Cả năm.docx'
    main_path = os.path.join(LBG_DIR, main_name)
    saved_main = save_safe(merged, main_path)
    print(f'  ✅ Tổng hợp: {os.path.basename(saved_main)}')

    # ── BUOC 3: Tach TTH+TH va THCS ──
    print('\n[3/3] Tách bản TTH+TH và THCS...')
    for keep_cap, suffix, label, to_truong in [
        ('TTH_TH', 'TTH+TH', 'Tiền TH + TH', TO_TRUONG_TH),
        ('THCS', 'THCS', 'THCS', TO_TRUONG_THCS),
    ]:
        # Load lai file da luu de tach
        d2 = Document(saved_main)
        filter_merged_for_cap(d2, keep_cap)
        update_sign_names(d2, to_truong)
        compact_sign_tables(d2)

        split_name = f'Lịch báo giảng - Cả năm ({suffix}).docx'
        split_path = os.path.join(LBG_DIR, split_name)
        saved = save_safe(d2, split_path)
        print(f'  ✅ {label}: {os.path.basename(saved)}')

    print(f'\n{"="*60}')
    print(f'  HOÀN TẤT! 3 file trong: {LBG_DIR}')
    print(f'{"="*60}\n')


if __name__ == '__main__':
    main()
