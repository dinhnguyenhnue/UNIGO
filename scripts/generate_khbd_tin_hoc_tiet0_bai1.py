import os
import re
import sys
import shutil
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')

TPL_PRIMARY = r'd:\UNIGO\Hệ thống mẫu văn bản\Khung  giáo án Unigo 2026-2027 Thang 7.2026.docx'
TPL_SECONDARY = r'd:\UNIGO\Hệ thống mẫu văn bản\PL4-Khung kế hoạch bài dạy (THCS).docx'
OUT_BASE_DIR = r'd:\UNIGO\KHBD_Tin_học'

def set_table_borders(table, color="000000", sz="4", val="single"):
    tblPr = table._tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def apply_font(run, font_name="Times New Roman", size_pt=13, bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color

def add_clean_paragraph(doc, text="", font_name="Times New Roman", size_pt=13, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=3, line_spacing=1.15):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        apply_font(r, font_name=font_name, size_pt=size_pt, bold=bold, italic=italic)
    return p

def sanitize_filename(name):
    import unicodedata
    name = unicodedata.normalize('NFD', name).encode('ascii', 'ignore').decode('utf-8')
    name = re.sub(r'[^a-zA-Z0-9_\-]', '_', name)
    name = re.sub(r'_+', '_', name).strip('_')
    return name

# ==============================================================================
# KHBD BUILDER ACCORDING TO USER'S EXACT PROMPT SPECIFICATION
# ==============================================================================
def build_khbd_exact_prompt(grade_str, lesson_title, tiet_ppct, yccd, is_primary):
    template_path = TPL_PRIMARY if is_primary else TPL_SECONDARY
    doc = docx.Document(template_path)

    # Safely update header text without overwriting header logo drawing (Run 0)
    for s in doc.sections:
        for hp in s.header.paragraphs:
            if len(hp.runs) >= 6:
                hp.runs[2].text = "Đậu Đình Nguyên"
                hp.runs[5].text = f"{grade_str} "

    # Clean body elements preserving section header/footer logo properties (w:sectPr)
    for child in list(doc.element.body):
        if not child.tag.endswith('sectPr'):
            doc.element.body.remove(child)

    # Header Title
    add_clean_paragraph(doc, f"TÊN BÀI DẠY: {lesson_title.upper()}", size_pt=14, bold=True)
    add_clean_paragraph(doc, f"Môn học: Tin học | Lớp: {grade_str} | Thời lượng: 1 tiết (45 phút) | Tiết theo PPCT: {tiet_ppct}", size_pt=12, bold=True)
    add_clean_paragraph(doc, "Giáo viên thực hiện: Đậu Đình Nguyên", size_pt=12, italic=True)
    add_clean_paragraph(doc, "")

    # I. MỤC TIÊU
    add_clean_paragraph(doc, "I. MỤC TIÊU", size_pt=13, bold=True)

    # 1. Về kiến thức (Dùng Danh từ / Cụm danh từ)
    add_clean_paragraph(doc, "1. Về kiến thức:", size_pt=13, bold=True)
    add_clean_paragraph(doc, f"* Sự hiểu biết về nội dung bài học: {yccd}.")
    add_clean_paragraph(doc, f"* Khả năng nhận diện và phân tích nguyên lý vận hành các khái niệm, thiết bị trong bài {lesson_title}.")
    add_clean_paragraph(doc, f"* Sự phân biệt được các thao tác đúng và quy trình an toàn trên môi trường máy tính.")

    # 2. Về năng lực (Chia 3 nhóm, chỉ rõ mốc #Hoạt động)
    add_clean_paragraph(doc, "2. Về năng lực:", size_pt=13, bold=True)
    
    # Năng lực đặc thù (Tin học)
    add_clean_paragraph(doc, "* Năng lực đặc thù (Tin học):", bold=True)
    add_clean_paragraph(doc, f"  * NLa/NLb/NLc (Sử dụng & Giải quyết vấn đề với Tin học): Khả năng nhận biết, thao tác sử dụng phần mềm và giải quyết nhiệm vụ bài học {lesson_title} (Đạt được thông qua Hoạt động 1, Hoạt động 2, Hoạt động 3, Hoạt động 4).")
    
    # Năng lực số (Không ghi chữ BAT BUOC)
    add_clean_paragraph(doc, "* Năng lực số:", bold=True)
    add_clean_paragraph(doc, f"  * Năng lực [3.4 - Lập trình & Tư duy số]: Liệt kê và thực hiện các chỉ dẫn rõ ràng để giải quyết nhiệm vụ cụ thể trên môi trường số (Đạt được thông qua Hoạt động 2, Hoạt động 3 và Hoạt động 4).")
    add_clean_paragraph(doc, f"  * Năng lực [4.1 - An toàn thiết bị & Dữ liệu]: Ý thức tuân thủ quy tắc an toàn khi sử dụng thiết bị máy tính (Đạt được thông qua Hoạt động 1 và Hoạt động 4).")

    # Năng lực chung
    add_clean_paragraph(doc, "* Năng lực chung:", bold=True)
    add_clean_paragraph(doc, f"  * Năng lực Tự chủ và tự học: Khả năng tự nghiên cứu nhiệm vụ trong SGK và tự giác thực hiện bài tập thực hành (Đạt được thông qua Hoạt động 2).")
    add_clean_paragraph(doc, f"  * Năng lực Giao tiếp và hợp tác: Khả năng thảo luận nhóm, phân công nhiệm vụ và tiếp nhận phản hồi từ bạn học (Đạt được thông qua Hoạt động 2 và Hoạt động 3).")
    add_clean_paragraph(doc, f"  * Năng lực Giải quyết vấn đề và sáng tạo: Khả năng phân tích tình huống thực tế, đề xuất quy trình các bước tối ưu để giải quyết nhiệm vụ (Đạt được thông qua Hoạt động 1 và Hoạt động 4).")

    # 3. Về phẩm chất (Chỉ rõ mốc #Hoạt động)
    add_clean_paragraph(doc, "3. Về phẩm chất:", size_pt=13, bold=True)
    add_clean_paragraph(doc, f"* Chăm chỉ: Tích cực tham gia các hoạt động phát biểu và kiên trì thực hiện nhiệm vụ thực hành (Thông qua Hoạt động 2, Hoạt động 3).")
    add_clean_paragraph(doc, f"* Trách nhiệm: Giữ gìn vệ sinh phòng máy, bảo vệ thiết bị số và hoàn thành sản phẩm được giao (Thông qua Hoạt động 3, Hoạt động 4).")
    add_clean_paragraph(doc, f"* Trung thực: Tự giác làm bài tập, tôn trọng kết quả học tập của bản thân và bạn bè (Thông qua Hoạt động 3).")
    add_clean_paragraph(doc, "")

    # II. THIẾT BỊ DẠY HỌC VÀ HỌC LIỆU
    add_clean_paragraph(doc, "II. THIẾT BỊ DẠY HỌC VÀ HỌC LIỆU", size_pt=13, bold=True)
    add_clean_paragraph(doc, f"* Giáo viên: Máy tính kết nối mạng, máy chiếu (màn hình tivi lớn), bài giảng điện tử sinh động cho bài {lesson_title}, Phiếu học tập số 1, Phiếu học tập số 2, Rubric đánh giá (chuyển toàn bộ xuống phần V. Phụ lục).")
    add_clean_paragraph(doc, f"* Học sinh: SGK Tin học {grade_str}, vở ghi bài, máy tính thực hành.")
    add_clean_paragraph(doc, "")

    # III. TIẾN TRÌNH DẠY HỌC (4 Hoạt động)
    add_clean_paragraph(doc, "III. TIẾN TRÌNH DẠY HỌC", size_pt=13, bold=True)

    activities = [
        ("1. Hoạt động 1: Khởi động/Xác định vấn đề (Thời gian: 7 phút)",
         "Sự khơi gợi hứng thú, tạo mâu thuẫn nhận thức và dẫn dắt học sinh nhận diện bài học.",
         f"Học sinh tham gia hoạt động khởi động/trải nghiệm liên quan đến {lesson_title}.",
         "Câu trả lời của học sinh về việc nhận diện vấn đề bài học.",
         [("Chuyển giao", f"- GV chiếu hình ảnh/tình huống thực tế liên quan đến {lesson_title} và đặt câu hỏi gợi mở.", "- HS quan sát hình ảnh trên màn chiếu, lắng nghe câu hỏi và suy nghĩ."),
          ("Thực hiện", "- GV quan sát, gợi mở suy nghĩ của học sinh.", "- HS trao đổi nhanh với bạn bên cạnh."),
          ("Báo cáo", "- GV gọi 2-3 đại diện HS phát biểu ý kiến.", "- HS giơ tay trả lời phát biểu ý kiến cá nhân."),
          ("Kết luận", f"- GV chốt kiến thức chuẩn và dẫn dắt vào bài {lesson_title}.", "- HS đối chiếu ghi nhận nội dung bài học vào vở.")]),

        ("2. Hoạt động 2: Hình thành kiến thức mới (Thời gian: 18 phút)",
         f"Sự hiểu biết về các khái niệm, quy trình và nội dung kiến thức bài {lesson_title}.",
         "Học sinh nghiên cứu SGK, quan sát giáo viên thao tác mẫu và hoàn thành Phiếu học tập số 1.",
         "Phiếu học tập số 1 hoàn thiện; ghi nhận kiến thức chuẩn.",
         [("Chuyển giao", f"- GV hướng dẫn HS nghiên cứu SGK bài {lesson_title} và phát Phiếu học tập số 1.", "- HS nhận Phiếu học tập số 1 từ GV."),
          ("Thực hiện", "- GV đi xung quanh quan sát, hỗ trợ giải đáp cho HS.", "- HS đọc SGK, thảo luận nhóm đôi hoàn thành Phiếu học tập số 1."),
          ("Báo cáo", "- GV gọi đại diện nhóm báo cáo kết quả Phiếu học tập số 1.", "- Đại diện nhóm đứng dậy trình bày kết quả thảo luận."),
          ("Kết luận", f"- GV chuẩn hóa kiến thức bài {lesson_title} trên màn hình.", "- HS tự kiểm tra đối chiếu và ghi lại nội dung vào vở.")]),

        ("3. Hoạt động 3: Luyện tập (Thời gian: 12 phút)",
         "Vận dụng kiến thức đã học để giải quyết bài tập luyện tập thực hành trên máy tính.",
         "Các nhóm thực hiện bài tập thực hành theo yêu cầu trong Phiếu học tập số 2.",
         f"Sản phẩm thực hành hoàn chỉnh bài {lesson_title} trên máy tính.",
         [("Chuyển giao", f"- GV chia nhóm HS và giao nhiệm vụ thực hành bài tập {lesson_title} (Phiếu học tập số 2).", "- Các nhóm tiếp nhận nhiệm vụ, phân công thao tác."),
          ("Thực hiện", "- GV đi quan sát hỗ trợ các nhóm gặp khó khăn kĩ thuật.", "- Các thành viên thực hiện bài tập thực hành trên máy tính."),
          ("Báo cáo", "- GV gọi 1 nhóm treo/trình diễn sản phẩm thực hành.", "- Đại diện nhóm thuyết trình sản phẩm. Các nhóm khác nhận xét chéo."),
          ("Kết luận", "- GV nhận xét bài làm của các nhóm, tuyên dương tinh thần hợp tác.", "- HS lắng nghe đánh giá của GV, tự sửa lỗi nếu có.")]),

        ("4. Hoạt động 4: Vận dụng (Thời gian: 8 phút - Giao việc tại lớp, hoàn thiện ở nhà)",
         "Phát triển năng lực giải quyết vấn đề thực tiễn thông qua tình huống mở rộng.",
         f"GV đưa ra bài toán thực tế yêu cầu HS vận dụng kiến thức bài {lesson_title} để giải quyết.",
         "Bài làm sản phẩm vận dụng thực tế của học sinh.",
         [("Chuyển giao", f"- GV giao nhiệm vụ vận dụng thực tế bài {lesson_title} và phổ biến Rubric 2.", "- HS quan sát nhiệm vụ trên màn hình, ghi lại yêu cầu vào vở."),
          ("Thực hiện", "- GV hướng dẫn định hướng xử lý tình huống thực tế.", "- HS thảo luận nhanh ý tưởng xử lý tại lớp."),
          ("Báo cáo", "- GV yêu cầu HS nộp sản phẩm hoàn thiện vào tiết học sau.", "- HS ghi nhận mốc thời gian nộp bài."),
          ("Kết luận", f"- GV chốt lại thông điệp bài học {lesson_title}.", "- HS lắng nghe, ghi nhận và kết thúc tiết học.")])
    ]

    for title_act, mục_tiêu, nội_dung, sản_phẩm, steps in activities:
        add_clean_paragraph(doc, title_act, size_pt=13, bold=True)
        add_clean_paragraph(doc, f"a) Mục tiêu: {mục_tiêu}")
        add_clean_paragraph(doc, f"b) Nội dung: {nội_dung}")
        add_clean_paragraph(doc, f"c) Sản phẩm: {sản_phẩm}")
        add_clean_paragraph(doc, "d) Tổ chức thực hiện:")

        # Exact 3-Column Table ("Bước", "Hoạt động của GV", "Hoạt động của HS")
        t = doc.add_table(rows=5, cols=3)
        set_table_borders(t)
        t.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = t.rows[0].cells
        hdr[0].text = "Bước"
        hdr[1].text = "Hoạt động của GV"
        hdr[2].text = "Hoạt động của HS"
        for cell in hdr:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    apply_font(r, size_pt=13, bold=True)

        for step_idx, (bước, gv_act, hs_act) in enumerate(steps, start=1):
            row_cells = t.rows[step_idx].cells
            row_cells[0].text = bước
            row_cells[1].text = gv_act
            row_cells[2].text = hs_act

            for cell in row_cells:
                set_cell_margins(cell)
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.15
                    p.paragraph_format.space_after = Pt(2)
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    for r in p.runs:
                        apply_font(r, size_pt=12)

        add_clean_paragraph(doc, "")

    # IV. ĐÁNH GIÁ KẾT QUẢ
    add_clean_paragraph(doc, "IV. ĐÁNH GIÁ KẾT QUẢ", size_pt=13, bold=True)
    add_clean_paragraph(doc, f"* Đánh giá thường xuyên: Thông qua mức độ tích cực tham gia phát biểu và kết quả hoàn thành Phiếu học tập số 1 ở Hoạt động 2.")
    add_clean_paragraph(doc, f"* Đánh giá định biên (Nhóm): Thông qua sản phẩm bài tập nhóm ở Hoạt động 3 (Đánh giá bằng Rubric 1 ở Phụ lục 3).")
    add_clean_paragraph(doc, f"* Đánh giá sản phẩm vận dụng: Đánh giá bài tập về nhà ở Hoạt động 4 dựa trên tiêu chí chính xác logic và sáng tạo (Đánh giá bằng Rubric 2 ở Phụ lục 4).")
    add_clean_paragraph(doc, "")

    # V. PHỤ LỤC (Chứa toàn bộ Phiếu học tập & Rubric)
    add_clean_paragraph(doc, "V. PHỤ LỤC", size_pt=13, bold=True)

    add_clean_paragraph(doc, f"Phụ lục 1: Phiếu học tập số 1 (Cá nhân - Hoạt động 2)")
    add_clean_paragraph(doc, f"Nhiệm vụ: Điền từ thích hợp vào chỗ trống để hoàn thiện định nghĩa và quy tắc bài {lesson_title}.")
    add_clean_paragraph(doc, "")

    add_clean_paragraph(doc, f"Phụ lục 2: Phiếu học tập số 2 (Hoạt động nhóm - Luyện tập Hoạt động 3)")
    add_clean_paragraph(doc, f"Lớp: {grade_str} | Thời gian thực hiện: 5 phút")
    add_clean_paragraph(doc, f"Nhiệm vụ: Thực hiện bài tập phân tích và hoàn thiện sản phẩm thực hành bài {lesson_title}.")
    add_clean_paragraph(doc, "")

    add_clean_paragraph(doc, "Phụ lục 3: Bảng tiêu chí đánh giá hoạt động nhóm (Rubric 1 - Dành cho Hoạt động 3)", size_pt=13, bold=True)

    # Rubric 1 Table (4 cols)
    t_rub1 = doc.add_table(rows=4, cols=4)
    set_table_borders(t_rub1)
    t_rub1.alignment = WD_TABLE_ALIGNMENT.CENTER

    r1_hdr = t_rub1.rows[0].cells
    r1_hdr[0].text = "Tiêu chí"
    r1_hdr[1].text = "Mức Tốt (9-10đ)"
    r1_hdr[2].text = "Mức Khá (7-8đ)"
    r1_hdr[3].text = "Mức Đạt (5-6đ)"
    for cell in r1_hdr:
        set_cell_margins(cell)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                apply_font(r, size_pt=12, bold=True)

    rub1_data = [
        ("Tính chính xác logic", "Sắp xếp đúng hoàn toàn quy trình các bước thực hiện.", "Sai 1 vị trí bước xử lý nhưng vẫn đảm bảo cấu trúc.", "Sai từ 2 vị trí bước trở lên, logic bị đảo lộn."),
        ("Sử dụng phần mềm/công cụ", "Thao tác thành thạo, không mắc lỗi kĩ thuật.", "Có 1-2 thao tác chưa tối ưu nhưng vẫn ra kết quả.", "Mắc nhiều lỗi kĩ thuật, cần GV hỗ trợ trực tiếp."),
        ("Tinh thần hợp tác", "Nhóm thảo luận sôi nổi, tất cả thành viên tham gia tích cực.", "Nhóm thảo luận nhưng còn 1 thành viên chưa tập trung.", "Làm việc rời rạc, chỉ có 1-2 thành viên làm việc.")
    ]

    for idx, (tc, tot, kha, dat) in enumerate(rub1_data, start=1):
        row_c = t_rub1.rows[idx].cells
        row_c[0].text = tc
        row_c[1].text = tot
        row_c[2].text = kha
        row_c[3].text = dat
        for cell in row_c:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(2)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    apply_font(r, size_pt=11)

    add_clean_paragraph(doc, "")

    add_clean_paragraph(doc, "Phụ lục 4: Bảng tiêu chí đánh giá sản phẩm vận dụng (Rubric 2 - Bài về nhà Hoạt động 4)", size_pt=13, bold=True)
    add_clean_paragraph(doc, f"* Xác định mục tiêu bài toán (3 điểm): Nhận diện chính xác yêu cầu bài toán thực tế.")
    add_clean_paragraph(doc, f"* Cấu trúc xử lý (5 điểm): Mô tả đầy đủ, rõ ràng và đúng trình tự các bước theo đúng bài học {lesson_title}.")
    add_clean_paragraph(doc, f"* Hình thức & Sáng tạo (2 điểm): Sản phẩm trình bày đẹp mắt, sạch sẽ, có màu sắc minh họa sinh động.")

    return doc

# ==============================================================================
# MAIN EXECUTOR
# ==============================================================================
def main():
    print("--- 1. Xóa toàn bộ các file KHBD Tin học hiện tại ---")
    if os.path.exists(OUT_BASE_DIR):
        for root, dirs, files in os.walk(OUT_BASE_DIR, topdown=False):
            for f in files:
                try:
                    os.remove(os.path.join(root, f))
                except Exception:
                    pass
            for d in dirs:
                try:
                    os.rmdir(os.path.join(root, d))
                except Exception:
                    pass
        print("  [+] Đã dọn dẹp thư mục KHBD_Tin_học.")

    os.makedirs(OUT_BASE_DIR, exist_ok=True)

    targets = [
        ("Tiền_tiểu_học", "Tiền tiểu học", True,
         "Tiết 0: Định hướng môn học - Em làm quen với thế giới công nghệ",
         "Nhận biết lớp học Tin học, nội quy phòng máy và làm quen với thế giới công nghệ số xung quanh em.",
         "Bài 1. Máy tính xung quanh em",
         "Nhận biết các thiết bị công nghệ quen thuộc trong đời sống và vai trò của máy tính."),

        ("Lớp_1", "Lớp 1", True,
         "Tiết 0: Định hướng môn học - Nội quy và an toàn phòng máy tính",
         "Nắm được nội quy phòng máy tính, tư thế ngồi học đúng và các quy tắc an toàn về điện khi sử dụng máy tính.",
         "Bài 1. Em làm quen với máy tính",
         "Nhận biết các bộ phận chính của máy tính bàn (Thân máy, Màn hình, Bàn phím, Chuột) và thao tác bật/tắt máy an toàn."),

        ("Lớp_2", "Lớp 2", True,
         "Tiết 0: Định hướng môn học - Em trở thành nhà sáng tạo số",
         "Khảo sát kĩ năng số ban đầu, nắm mục tiêu năm học và hình thành ý thức sử dụng công nghệ văn minh.",
         "Bài 1. Máy tính và các thiết bị số thông minh",
         "Phân biệt máy tính để bàn, máy tính xách tay, máy tính bảng và điện thoại thông minh; nêu lợi ích của thiết bị số."),

        ("Lớp_3", "Lớp 3", True,
         "Tiết 0: Định hướng môn học - Khám phá môn Tin học 3",
         "Giới thiệu cấu trúc môn Tin học Lớp 3, phương pháp học tập kết hợp lý thuyết và thực hành trên máy tính.",
         "Bài 1. Thông tin và quyết định",
         "Nêu ví dụ đơn giản minh họa vai trò của thông tin trong việc ra quyết định của con người; phân biệt 3 dạng thông tin cơ bản."),

        ("Lớp_4", "Lớp 4", True,
         "Tiết 0: Định hướng môn học - Khám phá môn Tin học 4",
         "Tổng quan chương trình Tin học 4, định hướng học tập chủ đề phần cứng, phần mềm và an toàn trên Internet.",
         "Bài 1. Phần cứng và phần mềm máy tính",
         "Kể tên một số thiết bị phần cứng và phần mềm; nêu vai trò và mối quan hệ phụ thuộc giữa phần cứng và phần mềm."),

        ("Lớp_5", "Lớp 5", True,
         "Tiết 0: Định hướng môn học - Khám phá môn Tin học 5",
         "Tổng quan chương trình Tin học 5, phát triển tư duy thuật toán và kĩ năng sáng tạo sản phẩm số.",
         "Bài 1. Em có thể làm gì với máy tính?",
         "Nêu ví dụ máy tính giúp học tập, giải trí, tìm kiếm, trao đổi thông tin, hợp tác và tạo sản phẩm số."),

        ("Lớp_6", "Lớp 6", False,
         "Tiết 0: Định hướng môn học Tin học 6 - Phương pháp học tập và an toàn số",
         "Giới thiệu tổng quan chương trình Tin học 6 THCS, phương pháp nghiên cứu tài liệu, thực hành dự án và văn hóa ứng xử trên mạng.",
         "Bài 1. Thông tin và dữ liệu",
         "Giải thích được sự khác nhau giữa thông tin và dữ liệu; nêu được tầm quan trọng của thông tin và dữ liệu trong đời sống."),

        ("Lớp_7", "Lớp 7", False,
         "Tiết 0: Định hướng môn học Tin học 7 - Tổng quan chương trình và kỹ năng số",
         "Tổng quan môn Tin học 7 THCS, định hướng khai thác thiết bị vào-ra, phần mềm bảng tính và làm việc nhóm trực tuyến.",
         "Bài 1. Thiết bị vào - ra",
         "Nhận biết và phân biệt được chức năng của các thiết bị vào và thiết bị ra trong hệ thống máy tính."),

        ("Lớp_8", "Lớp 8", False,
         "Tiết 0: Định hướng môn học Tin học 8 - Định hướng học tập và nghiên cứu công nghệ",
         "Tổng quan chương trình Tin học 8 THCS, phương pháp phân tích thuật toán, xử lý dữ liệu và định hướng nghề nghiệp CNTT.",
         "Bài 1. Lược sử công cụ tính toán",
         "Trình bày sơ lược lịch sử phát triển của công cụ tính toán và máy tính qua các thời kỳ; nêu tác động đến xã hội."),

        ("Lớp_9", "Lớp 9", False,
         "Tiết 0: Định hướng môn học Tin học 9 - Tổng quan chương trình và định hướng công nghệ tương lai",
         "Tổng quan chương trình Tin học 9 THCS, định hướng công nghệ số thông minh, trí tuệ nhân tạo và tạo lập trang web.",
         "Bài 1. Thế giới kĩ thuật số",
         "Nêu được ví dụ thiết bị số thông minh và tác động hai mặt của thế giới kỹ thuật số đối với đời sống con người.")
    ]

    total_files = 0
    print("\n--- 2. Tiến hành tạo Tiết 0 và Bài 1 theo đúng Prompt quy chuẩn mới ---")

    for folder_prefix, grade_str, is_primary, t0_title, t0_yccd, b1_title, b1_yccd in targets:
        print(f"\n================ Đang xử lý {grade_str.upper()} ================")
        
        # Tiết 0
        safe_t0 = sanitize_filename(t0_title)
        dir_t0 = os.path.join(OUT_BASE_DIR, folder_prefix, "Tiết_00")
        os.makedirs(dir_t0, exist_ok=True)
        file_t0 = os.path.join(dir_t0, f"KHBD_Tin_hoc_{folder_prefix}_Tiet00_{safe_t0}.docx")

        doc_t0 = build_khbd_exact_prompt(grade_str, t0_title, 1, t0_yccd, is_primary)
        try:
            doc_t0.save(file_t0)
            print(f"  [+] Đã tạo Tiết 0: {folder_prefix} -> Tiết_00 -> {os.path.basename(file_t0)}")
        except Exception as e:
            file_t0_alt = file_t0.replace(".docx", "_new.docx")
            doc_t0.save(file_t0_alt)
            print(f"  [!] File locked by Word, saved to: {os.path.basename(file_t0_alt)}")

        # Bài 1
        safe_b1 = sanitize_filename(b1_title)
        dir_b1 = os.path.join(OUT_BASE_DIR, folder_prefix, "Bài_01")
        os.makedirs(dir_b1, exist_ok=True)
        file_b1 = os.path.join(dir_b1, f"KHBD_Tin_hoc_{folder_prefix}_Bai01_{safe_b1}.docx")

        doc_b1 = build_khbd_exact_prompt(grade_str, b1_title, 2, b1_yccd, is_primary)
        try:
            doc_b1.save(file_b1)
            print(f"  [+] Đã tạo Bài 1 : {folder_prefix} -> Bài_01 -> {os.path.basename(file_b1)}")
        except Exception as e:
            file_b1_alt = file_b1.replace(".docx", "_new.docx")
            doc_b1.save(file_b1_alt)
            print(f"  [!] File locked by Word, saved to: {os.path.basename(file_b1_alt)}")

    print(f"\n==========================================")
    print(f" HOÀN THÀNH CHUẨN HÓA 20 FILE TIẾT 0 & BÀI 1 CHO 10 KHỐI LỚP")
    print(f"==========================================")

if __name__ == '__main__':
    main()
