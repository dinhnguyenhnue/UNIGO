"""
generate_lbg.py — Tu dong tao Lich bao giang theo tuan tu TKB Giao vien
========================================================================

QUY TAC ROTATION:
  - Lop 5, 6, 7, 8: Moi lop co 2 tiet / tuan
      Tuan CHAN (2, 4, 6...): ca 2 tiet -> Tin hoc (Phong Tin hoc)
      Tuan LE  (3, 5, 7...): ca 2 tiet -> Robotics (Bo Kit Robotics)
  - Cac lop con lai: giu nguyen mon hoc goc tu TKB

QUY TAC PPCT (Khong offset):
  - Tuan 1: Tat ca = Tiet 0 (Dinh huong)
  - Tuan 2+: Tat ca cac Thu (T2 -> T6) deu dong bo ppct = tuan_so - 1
  - Rotation classes (5,6,7,8): PPCT rieng cho Tin/Robotics
      Khi co 2 tiet lien tiep cung mon -> PPCT lien tiep (Tuan 2: PPCT 1,2; Tuan 4: PPCT 3,4...)
      Tuan 3 (Le - Robotics): PPCT 1, 2; Tuan 5: PPCT 3, 4...

MASTER SCHEDULE (Tu 'Thoi khoa bieu - Dau Dinh Nguyen.xlsx'):
  SANG (14 tiet):
    - Thu Hai: T3 Rob 4C1, T4 Tin 1A1
    - Thu Ba: T1 Tin 1C1, T3-T4 Tin-Rob 5C1
    - Thu Tu: T2 Rob 3A1, T4 Tin 3C1
    - Thu Nam: T1 Tin TT3, T2 Rob 1A1, T4 Rob 2C1
    - Thu Sau: T1 Tin 2A1, T2 Rob 3C1, T4-T5 Tin-Rob 6A1
  CHIEU (11 tiet):
    - Thu Hai: T3 Tin 2C1
    - Thu Ba: T3-T4 Tin-Rob 7A1
    - Thu Tu: T3 Tin 4C1, T4 Rob 2A1
    - Thu Nam: T1 Tin 3A1, T2 Tin TTH 1, T3 Tin TTH2, T4 Rob 1C1
    - Thu Sau: T1-T2 Tin-Rob 8A1
"""

import sys
import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import date, timedelta

sys.stdout.reconfigure(encoding='utf-8')

LBG_DIR = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Lịch báo giảng'
TEMPLATE = os.path.join(LBG_DIR, 'template_lbg_goc.docx')
if not os.path.exists(TEMPLATE):
    TEMPLATE = os.path.join(LBG_DIR, 'Lịch báo giảng - Tuần 01.docx')
TUAN_01_START = date(2026, 8, 3)

ROTATION_PREFIXES = ('5', '6', '7', '8')
THU_LABELS = {0: 'Hai', 1: 'Ba', 2: 'Tư', 3: 'Năm', 4: 'Sáu'}

# Tên tổ trưởng theo cấp học
TO_TRUONG_THCS = 'Nguyễn Thị Ngọc Ánh'
TO_TRUONG_TH = 'Nguyễn Thị Ngọc'

# Số dòng chấm nhận xét tối đa trong bảng chữ ký (giữ gọn 1 trang)
MAX_DONG_CHAM = 4

PPCT_TIN = {
    'TT3': {
        0: 'Tiết 0: Định hướng môn học',
        1: 'Em làm quen với thế giới công nghệ',
        2: 'Máy tính quanh em',
        3: 'Em ngồi máy tính an toàn',
        4: 'Làm quen với chuột',
        5: 'Kéo – thả thật vui',
        6: 'Khám phá bàn phím',
        7: 'Chữ, hình và âm thanh',
        8: 'Ôn tập Đánh giá định kỳ 1',
        9: 'Đánh giá định kỳ 1',
        10: 'Màu sắc và hình dạng',
        11: 'Em tạo bức tranh đầu tiên',
        12: 'Thông tin quanh em',
        13: 'Máy tính làm việc theo lệnh',
        14: 'Một việc – nhiều bước',
        15: 'Chỉ đường cho nhân vật',
        16: 'Mê cung của em',
        17: 'Ôn tập Đánh giá định kỳ 2',
        18: 'Đánh giá định kỳ 2',
        19: 'Sửa lỗi đường đi',
        20: 'Làm quen lập trình trực quan',
        21: 'Kể chuyện số',
        22: 'Tạo nhân vật',
        23: 'Tạo bối cảnh',
        24: 'Nhân vật chuyển động',
        25: 'Âm thanh và lời kể',
        26: 'Ôn tập Đánh giá định kỳ 3',
        27: 'Đánh giá định kỳ 3',
        28: 'Dự án: Câu chuyện số',
        29: 'Chia sẻ và làm việc nhóm',
        30: 'Robot quanh em',
        31: 'Điều khiển robot',
        32: 'Ôn tập Đánh giá định kỳ 4',
        33: 'Đánh giá định kỳ 4',
        34: 'Tổng kết năm học',
        35: 'Tổng kết năm học',
    },
    '1': {
        0: 'Tiết 0: Định hướng môn học',
        1: 'Em làm quen với thế giới công nghệ',
        2: 'Ôn và nâng cấp kỹ năng chuột',
        3: 'Bàn phím và gõ câu ngắn',
        4: 'Mở, đóng và chuyển đổi ứng dụng',
        5: 'Tệp là gì?',
        6: 'Thư mục là gì?',
        7: 'Đặt tên và lưu sản phẩm',
        8: 'Ôn tập Đánh giá định kỳ 1',
        9: 'Đánh giá định kỳ 1',
        10: 'Mở lại và sắp xếp sản phẩm',
        11: 'Tạo một sản phẩm có tổ chức',
        12: 'Thông tin và dữ liệu',
        13: 'Thông tin dạng chữ, hình, âm thanh, video',
        14: 'Thu nhận và xử lý thông tin',
        15: 'Một nhiệm vụ thành nhiều bước',
        16: 'Thuật toán bằng lời và hình',
        17: 'Ôn tập Đánh giá định kỳ 2',
        18: 'Đánh giá định kỳ 2',
        19: 'Lập trình trực quan: chuỗi lệnh',
        20: 'Lặp lại một hành động',
        21: 'Chạy thử và sửa lỗi',
        22: 'Thiết kế câu chuyện số',
        23: 'Thiết kế nhân vật',
        24: 'Thiết kế bối cảnh',
        25: 'Animation: chuyển động',
        26: 'Ôn tập Đánh giá định kỳ 3',
        27: 'Đánh giá định kỳ 3',
        28: 'Tương tác và âm thanh',
        29: 'Dự án: Hoạt hình ngắn',
        30: 'Thiết kế sản phẩm nhóm',
        31: 'Kiểm thử và hoàn thiện',
        32: 'Ôn tập Đánh giá định kỳ 4',
        33: 'Đánh giá định kỳ 4',
        34: 'Tổng kết năm học',
        35: 'Tổng kết năm học',
    },
    '2': {
        0: 'Tiết 0: Định hướng môn học',
        1: 'Em trở thành nhà sáng tạo số',
        2: 'Ôn và nâng cấp kỹ năng chuột',
        3: 'Bàn phím và gõ câu ngắn',
        4: 'Mở, đóng và chuyển đổi ứng dụng',
        5: 'Tệp là gì?',
        6: 'Thư mục là gì?',
        7: 'Đặt tên và lưu sản phẩm',
        8: 'Ôn tập Đánh giá định kỳ 1',
        9: 'Đánh giá định kỳ 1',
        10: 'Mở lại và sắp xếp sản phẩm',
        11: 'Tạo một sản phẩm có tổ chức',
        12: 'Thông tin và dữ liệu',
        13: 'Thông tin dạng chữ, hình, âm thanh, video',
        14: 'Thu nhận và xử lý thông tin',
        15: 'Một nhiệm vụ thành nhiều bước',
        16: 'Thuật toán bằng lời và hình',
        17: 'Ôn tập Đánh giá định kỳ 2',
        18: 'Đánh giá định kỳ 2',
        19: 'Lập trình trực quan: chuỗi lệnh',
        20: 'Lặp lại một hành động',
        21: 'Chạy thử và sửa lỗi',
        22: 'Thiết kế câu chuyện số',
        23: 'Thiết kế nhân vật',
        24: 'Thiết kế bối cảnh',
        25: 'Animation: chuyển động',
        26: 'Ôn tập Đánh giá định kỳ 3',
        27: 'Đánh giá định kỳ 3',
        28: 'Tương tác và âm thanh',
        29: 'Dự án: Hoạt hình ngắn',
        30: 'Thiết kế sản phẩm nhóm',
        31: 'Kiểm thử và hoàn thiện',
        32: 'Ôn tập Đánh giá định kỳ 4',
        33: 'Đánh giá định kỳ 4',
        34: 'Tổng kết năm học',
        35: 'Tổng kết năm học',
    },
    '3': {
        0: 'Tiết 0: Định hướng môn học',
        1: 'Bài 1: Thông tin và quyết định (Tiết 1)',
        2: 'Bài 1: Thông tin và quyết định (Tiết 2)',
        3: 'Bài 2: Xử lí thông tin (Tiết 1)',
        4: 'Bài 2: Xử lí thông tin (Tiết 2)',
        5: 'Bài 3: Máy tính và em (Tiết 1)',
        6: 'Bài 3: Máy tính và em (Tiết 2)',
        7: 'Bài 4: Làm việc với máy tính (Tiết 1)',
        8: 'Ôn tập Đánh giá định kỳ 1',
        9: 'Đánh giá định kỳ 1',
        10: 'Bài 4: Làm việc với máy tính (Tiết 2)',
        11: 'Bài 5: Sử dụng bàn phím (Tiết 1)',
        12: 'Bài 5: Sử dụng bàn phím (Tiết 2)',
        13: 'Bài 6: Khám phá thông tin trên Internet (Tiết 1)',
        14: 'Bài 6: Khám phá thông tin trên Internet (Tiết 2)',
        15: 'Bài 7: Sắp xếp để dễ tìm (Tiết 1)',
        16: 'Bài 7: Sắp xếp để dễ tìm (Tiết 2)',
        17: 'Ôn tập Đánh giá định kỳ 2',
        18: 'Đánh giá định kỳ 2',
        19: 'Bài 8: Sơ đồ hình cây. Tổ chức thông tin trong máy tính (Tiết 1)',
        20: 'Bài 8: Sơ đồ hình cây. Tổ chức thông tin trong máy tính (Tiết 2)',
        21: 'Bài 9: Thực hành với tệp và thư mục trong máy tính (Tiết 1)',
        22: 'Bài 9: Thực hành với tệp và thư mục trong máy tính (Tiết 2)',
        23: 'Bài 10: Bảo vệ thông tin khi dùng máy tính',
        24: 'Bài 11: Bài trình chiếu của em (Tiết 1)',
        25: 'Bài 11: Bài trình chiếu của em (Tiết 2)',
        26: 'Ôn tập Đánh giá định kỳ 3',
        27: 'Đánh giá định kỳ 3',
        28: 'Bài 12: Tìm hiểu về thế giới tự nhiên',
        29: 'Bài 13: Luyện tập sử dụng chuột',
        30: 'Bài 14: Em thực hiện công việc như thế nào? (Tiết 1)',
        31: 'Bài 14: Em thực hiện công việc như thế nào? (Tiết 2)',
        32: 'Ôn tập Đánh giá định kỳ 4',
        33: 'Đánh giá định kỳ 4',
        34: 'Tổng kết năm học',
        35: 'Tổng kết năm học',
    },
    '4': {
        0: 'Tiết 0: Định hướng môn học',
        1: 'Bài 1: Phần cứng và phần mềm máy tính (Tiết 1)',
        2: 'Bài 1: Phần cứng và phần mềm máy tính (Tiết 2)',
        3: 'Bài 2: Gõ bàn phím đúng cách (Tiết 1)',
        4: 'Bài 2: Gõ bàn phím đúng cách (Tiết 2)',
        5: 'Bài 3: Thông tin trên trang Web (Tiết 1)',
        6: 'Bài 3: Thông tin trên trang Web (Tiết 2)',
        7: 'Bài 4: Tìm kiếm thông tin trên Internet (Tiết 1)',
        8: 'Ôn tập Đánh giá định kỳ 1',
        9: 'Đánh giá định kỳ 1',
        10: 'Bài 4: Tìm kiếm thông tin trên Internet (Tiết 2)',
        11: 'Bài 5: Thao tác với tệp và thư mục (Tiết 1)',
        12: 'Bài 5: Thao tác với tệp và thư mục (Tiết 2)',
        13: 'Bài 6: Sử dụng phần mềm khi được phép',
        14: 'Bài 7: Tạo bài trình chiếu (Tiết 1)',
        15: 'Bài 7: Tạo bài trình chiếu (Tiết 2)',
        16: 'Bài 8: Định dạng văn bản trên trang chiếu (Tiết 1)',
        17: 'Ôn tập Đánh giá định kỳ 2',
        18: 'Đánh giá định kỳ 2',
        19: 'Bài 8: Định dạng văn bản trên trang chiếu (Tiết 2)',
        20: 'Bài 9: Hiệu ứng chuyển trang (Tiết 1)',
        21: 'Bài 9: Hiệu ứng chuyển trang (Tiết 2)',
        22: 'Bài 10: Phần mềm soạn thảo văn bản (Tiết 1)',
        23: 'Bài 10: Phần mềm soạn thảo văn bản (Tiết 2)',
        24: 'Bài 11: Chỉnh sửa văn bản (Tiết 1)',
        25: 'Bài 11: Chỉnh sửa văn bản (Tiết 2)',
        26: 'Ôn tập Đánh giá định kỳ 3',
        27: 'Đánh giá định kỳ 3',
        28: 'Bài 13: Chơi với máy tính (Tiết 1)',
        29: 'Bài 13: Chơi với máy tính (Tiết 2)',
        30: 'Bài 14: Khám phá môi trường lập trình trực quan (Tiết 1)',
        31: 'Bài 14: Khám phá môi trường lập trình trực quan (Tiết 2)',
        32: 'Ôn tập Đánh giá định kỳ 4',
        33: 'Đánh giá định kỳ 4',
        34: 'Tổng kết năm học',
        35: 'Tổng kết năm học',
    },
    '5': {
        0: 'Tiết 0: Định hướng môn học',
        1: 'Bài 1: Em có thể làm gì với máy tính? (Tiết 1)',
        2: 'Bài 1: Em có thể làm gì với máy tính? (Tiết 2)',
        3: 'Bài 2: Tìm kiếm thông tin trên website (Tiết 1)',
        4: 'Bài 2: Tìm kiếm thông tin trên website (Tiết 2)',
        5: 'Bài 3: Tìm kiếm thông tin trong giải quyết vấn đề (Tiết 1)',
        6: 'Bài 3: Tìm kiếm thông tin trong giải quyết vấn đề (Tiết 2)',
        7: 'Bài 4: Cây thư mục (Tiết 1)',
        8: 'Bài 4: Cây thư mục (Tiết 2)',
        9: 'Ôn tập Đánh giá định kỳ 1',
        10: 'Đánh giá định kỳ 1',
        11: 'Bài 5: Bản quyền nội dung thông tin (Tiết 1)',
        12: 'Bài 5: Bản quyền nội dung thông tin (Tiết 2)',
        13: 'Bài 6: Định dạng kí tự và bố trí hình ảnh trong văn bản (Tiết 1)',
        14: 'Bài 6: Định dạng kí tự và bố trí hình ảnh trong văn bản (Tiết 2)',
        15: 'Bài 7: Thực hành soạn thảo văn bản (Tiết 1)',
        16: 'Bài 7: Thực hành soạn thảo văn bản (Tiết 2)',
        17: 'Bài 8A: Làm quen với phần mềm đồ họa',
        18: 'Ôn tập Đánh giá định kỳ 2',
        19: 'Đánh giá định kỳ 2',
        20: 'Bài 9A: Sử dụng phần mềm đồ hoạ tạo sản phẩm số',
        21: 'Bài 10: Cấu trúc tuần tự (Tiết 1)',
        22: 'Bài 10: Cấu trúc tuần tự (Tiết 2)',
        23: 'Bài 11: Cấu trúc lặp (Tiết 1)',
        24: 'Bài 11: Cấu trúc lặp (Tiết 2)',
        25: 'Bài 12: Thực hành sử dụng lệnh lặp (Tiết 1)',
        26: 'Bài 12: Thực hành sử dụng lệnh lặp (Tiết 2)',
        27: 'Ôn tập Đánh giá định kỳ 3',
        28: 'Đánh giá định kỳ 3',
        29: 'Bài 13: Cấu trúc rẽ nhánh (Tiết 1)',
        30: 'Bài 13: Cấu trúc rẽ nhánh (Tiết 2)',
        31: 'Bài 14: Sử dụng biến trong chương trình (Tiết 1)',
        32: 'Bài 14: Sử dụng biến trong chương trình (Tiết 2)',
        33: 'Ôn tập Đánh giá định kỳ 4',
        34: 'Đánh giá định kỳ 4',
        35: 'Tổng kết năm học',
    },
    '6': {
        0: 'Tiết 0: Định hướng môn học',
        1: 'Bài 1. Thông tin và dữ liệu (Tiết 1)',
        2: 'Bài 1. Thông tin và dữ liệu (Tiết 2)',
        3: 'Bài 2. Xử lí thông tin (Tiết 1)',
        4: 'Bài 2. Xử lí thông tin (Tiết 2)',
        5: 'Bài 3. Thông tin trong máy tính (Tiết 1)',
        6: 'Bài 3. Thông tin trong máy tính (Tiết 2)',
        7: 'Bài 4. Mạng máy tính (Tiết 1)',
        8: 'Bài 4. Mạng máy tính (Tiết 2)',
        9: 'Ôn tập Đánh giá định kỳ 1',
        10: 'Đánh giá định kỳ 1',
        11: 'Bài 5. Internet (Tiết 1)',
        12: 'Bài 5. Internet (Tiết 2)',
        13: 'Bài 6. Mạng thông tin toàn cầu (Tiết 1)',
        14: 'Bài 6. Mạng thông tin toàn cầu (Tiết 2)',
        15: 'Bài 7. Tìm kiếm thông tin trên Internet (Tiết 1)',
        16: 'Bài 7. Tìm kiếm thông tin trên Internet (Tiết 2)',
        17: 'Bài 8. Thư điện tử (Tiết 1)',
        18: 'Ôn tập Đánh giá định kỳ 2',
        19: 'Đánh giá định kỳ 2',
        20: 'Bài 8. Thư điện tử (Tiết 2)',
        21: 'Bài 9. An toàn thông tin trên Internet (Tiết 1)',
        22: 'Bài 9. An toàn thông tin trên Internet (Tiết 2)',
        23: 'Bài 10. Sơ đồ tư duy (Tiết 1)',
        24: 'Bài 10. Sơ đồ tư duy (Tiết 2)',
        25: 'Bài 11. Định dạng văn bản (Tiết 1)',
        26: 'Bài 11. Định dạng văn bản (Tiết 2)',
        27: 'Ôn tập Đánh giá định kỳ 3',
        28: 'Đánh giá định kỳ 3',
        29: 'Bài 12. Trình bày thông tin ở dạng bảng (Tiết 1)',
        30: 'Bài 12. Trình bày thông tin ở dạng bảng (Tiết 2)',
        31: 'Bài 13. Thực hành: Tìm kiếm và thay thế',
        32: 'Bài 14. Thực hành tổng hợp',
        33: 'Ôn tập Đánh giá định kỳ 4',
        34: 'Đánh giá định kỳ 4',
        35: 'Tổng kết năm học',
    },
    '7': {
        0: 'Tiết 0: Định hướng môn học',
        1: 'Bài 1. Thiết bị vào - ra (Tiết 1)',
        2: 'Bài 1. Thiết bị vào - ra (Tiết 2)',
        3: 'Bài 2. Phần mềm máy tính (Tiết 1)',
        4: 'Bài 2. Phần mềm máy tính (Tiết 2)',
        5: 'Bài 3. Quản lý dữ liệu trong máy tính (Tiết 1)',
        6: 'Bài 3. Quản lý dữ liệu trong máy tính (Tiết 2)',
        7: 'Bài 4. Mạng xã hội và một số kênh trao đổi thông tin trên Internet (Tiết 1)',
        8: 'Bài 4. Mạng xã hội và một số kênh trao đổi thông tin trên Internet (Tiết 2)',
        9: 'Ôn tập Đánh giá định kỳ 1',
        10: 'Đánh giá định kỳ 1',
        11: 'Bài 5. Ứng xử trên mạng (Tiết 1)',
        12: 'Bài 5. Ứng xử trên mạng (Tiết 2)',
        13: 'Bài 6. Làm quen với phần mềm bảng tính (Tiết 1)',
        14: 'Bài 6. Làm quen với phần mềm bảng tính (Tiết 2)',
        15: 'Bài 7. Tính toán tự động trên bảng tính (Tiết 1)',
        16: 'Bài 7. Tính toán tự động trên bảng tính (Tiết 2)',
        17: 'Bài 8. Công cụ hỗ trợ tính toán (Tiết 1)',
        18: 'Ôn tập Đánh giá định kỳ 2',
        19: 'Đánh giá định kỳ 2',
        20: 'Bài 8. Công cụ hỗ trợ tính toán (Tiết 2)',
        21: 'Bài 9. Trình bày bảng tính (Tiết 1)',
        22: 'Bài 9. Trình bày bảng tính (Tiết 2)',
        23: 'Bài 10. Hoàn thiện bảng tính (Tiết 1)',
        24: 'Bài 10. Hoàn thiện bảng tính (Tiết 2)',
        25: 'Bài 11. Tạo bài trình chiếu (Tiết 1)',
        26: 'Bài 11. Tạo bài trình chiếu (Tiết 2)',
        27: 'Ôn tập Đánh giá định kỳ 3',
        28: 'Đánh giá định kỳ 3',
        29: 'Bài 12. Định dạng đối tượng trên trang chiếu (Tiết 1)',
        30: 'Bài 12. Định dạng đối tượng trên trang chiếu (Tiết 2)',
        31: 'Bài 13. Thực hành tổng hợp: Hoàn thiện bài trình chiếu',
        32: 'Bài 14. Thuật toán tìm kiếm tuần tự',
        33: 'Ôn tập Đánh giá định kỳ 4',
        34: 'Đánh giá định kỳ 4',
        35: 'Tổng kết năm học',
    },
    '8': {
        0: 'Tiết 0: Định hướng môn học',
        1: 'Bài 1. Lược sử công cụ tính toán (Tiết 1)',
        2: 'Bài 1. Lược sử công cụ tính toán (Tiết 2)',
        3: 'Bài 2. Thông tin trong môi trường số (Tiết 1)',
        4: 'Bài 2. Thông tin trong môi trường số (Tiết 2)',
        5: 'Bài 3. Thực hành khai thác thông tin số (Tiết 1)',
        6: 'Bài 3. Thực hành khai thác thông tin số (Tiết 2)',
        7: 'Bài 4. Đạo đức và văn hóa trong sử dụng công nghệ số',
        8: 'Bài 5. Sử dụng bảng tính giải quyết bài toán thực tế (Tiết 1)',
        9: 'Ôn tập Đánh giá định kỳ 1',
        10: 'Đánh giá định kỳ 1',
        11: 'Bài 5. Sử dụng bảng tính giải quyết bài toán thực tế (Tiết 2)',
        12: 'Bài 6. Sắp xếp và lọc dữ liệu (Tiết 1)',
        13: 'Bài 6. Sắp xếp và lọc dữ liệu (Tiết 2)',
        14: 'Bài 7. Trình bày dữ liệu bằng biểu đồ (Tiết 1)',
        15: 'Bài 7. Trình bày dữ liệu bằng biểu đồ (Tiết 2)',
        16: 'Bài 8a. Làm việc với danh sách dạng liệt kê và hình ảnh trong văn bản (Tiết 1)',
        17: 'Bài 8a. Làm việc với danh sách dạng liệt kê và hình ảnh trong văn bản (Tiết 2)',
        18: 'Ôn tập Đánh giá định kỳ 2',
        19: 'Đánh giá định kỳ 2',
        20: 'Bài 9a. Tạo đầu trang, chân trang cho văn bản (Tiết 1)',
        21: 'Bài 9a. Tạo đầu trang, chân trang cho văn bản (Tiết 2)',
        22: 'Bài 10a. Định dạng nâng cao cho trang chiếu (Tiết 1)',
        23: 'Bài 10a. Định dạng nâng cao cho trang chiếu (Tiết 2)',
        24: 'Bài 11a. Sử dụng bản mẫu cho bài trình chiếu (Tiết 1)',
        25: 'Bài 11a. Sử dụng bản mẫu cho bài trình chiếu (Tiết 2)',
        26: 'Bài 12. Từ thuật toán đến chương trình (Tiết 1)',
        27: 'Ôn tập Đánh giá định kỳ 3',
        28: 'Đánh giá định kỳ 3',
        29: 'Bài 12. Từ thuật toán đến chương trình (Tiết 2)',
        30: 'Bài 13. Biểu diễn dữ liệu (Tiết 1)',
        31: 'Bài 13. Biểu diễn dữ liệu (Tiết 2)',
        32: 'Bài 16. Tin học với nghề nghiệp',
        33: 'Ôn tập Đánh giá định kỳ 4',
        34: 'Đánh giá định kỳ 4',
        35: 'Tổng kết năm học',
    },
}

PPCT_ROB = {
    '1': {
        0: 'Tiết 0: Định hướng môn học',
        1: 'Bài 1. Tập thể dục nào!',
        2: 'Bài 2. Chú cún dễ thương',
        3: 'Bài 3. Tăng cường sức khỏe',
        4: 'Bài 4. Chú ốc sên chậm chạp',
        5: 'Bài 5. Xe cảnh sát tuần tra',
        6: 'Bài 6. Khám phá xe cứu hoả',
        7: 'Bài 7. Người giao hàng đã đến',
        8: 'Ôn tập Đánh giá định kỳ 1',
        9: 'Đánh giá định kỳ 1',
        10: 'Bài 8. Thế giới khủng long',
        11: 'Bài 9. Hồ bơi mùa hè',
        12: 'Bài 10. Khám phá đại dương',
        13: 'Bài 11. Chú cua cứng cáp',
        14: 'Bài 12. Tôi có thể di chuyển đến bất cứ đâu',
        15: 'Bài 13. Hoạt động mùa hè',
        16: 'Bài 14. Bắt đầu chuyến hành trình cùng tàu hoả',
        17: 'Ôn tập Đánh giá định kỳ 2',
        18: 'Đánh giá định kỳ 2',
        19: 'Bài 15. Phía trên bầu trời',
        20: 'Bài 16. Khám phá vũ trụ rộng lớn',
        21: 'Bài 17. Máy bắn đá khổng lồ',
        22: 'Bài 18. Trò chơi dân gian',
        23: 'Bài 19. Khám phá trò chơi truyền thống các nước',
        24: 'Bài 20. Đấu vật thú vị',
        25: 'Bài 21. Sóc nhỏ dễ thương',
        26: 'Ôn tập Đánh giá định kỳ 3',
        27: 'Đánh giá định kỳ 3',
        28: 'Bài 22. Chú hươu tuyệt đẹp',
        29: 'Bài 23. Chú rùa thông minh',
        30: 'Bài 24. Đôi chân mạnh mẽ của chuột túi',
        31: 'Bài 25. Chú sư tử dũng mãnh',
        32: 'Ôn tập Đánh giá định kỳ 4',
        33: 'Đánh giá định kỳ 4',
        34: 'Tổng kết môn học & Triển lãm Robotics',
        35: 'Tổng kết môn học & Triển lãm Robotics',
    },
    '2': {
        0: 'Tiết 0: Định hướng môn học',
        1: 'Bài 1. Hãy nấu những món ăn ngon',
        2: 'Bài 2. Hàm răng trắng sáng',
        3: 'Bài 3. Sự phản xạ ánh sáng',
        4: 'Bài 4. Đàn gà con',
        5: 'Bài 5. Những bạn nhỏ lễ phép',
        6: 'Bài 6. Động vật thân mềm',
        7: 'Bài 7. Khu phố của chúng ta',
        8: 'Ôn tập Đánh giá định kỳ 1',
        9: 'Đánh giá định kỳ 1',
        10: 'Bài 8. Rèn luyện sức khỏe',
        11: 'Bài 9. Môi trường biển',
        12: 'Bài 10. Cùng câu cá nào!',
        13: 'Bài 11. Thế giới khủng long',
        14: 'Bài 12. Người hiệp sĩ dũng cảm',
        15: 'Bài 13. Vận chuyển đồ vật',
        16: 'Bài 14. Sức mạnh của máy ủi',
        17: 'Ôn tập Đánh giá định kỳ 2',
        18: 'Đánh giá định kỳ 2',
        19: 'Bài 15. Máy xúc',
        20: 'Bài 16. Cùng nhau đi khắp thế giới',
        21: 'Bài 17. Nhắm và bắn!',
        22: 'Bài 18. Độ đàn hồi, lực đẩy của cung tên',
        23: 'Bài 19. Ba! Hai! Một! Bắn!!!',
        24: 'Bài 20. Cùng nhau tham quan thành phố',
        25: 'Bài 21. Khám phá trang phục người da đỏ',
        26: 'Ôn tập Đánh giá định kỳ 3',
        27: 'Đánh giá định kỳ 3',
        28: 'Bài 22. Cá sấu thật ngầu!',
        29: 'Bài 23. Những nốt nhạc vui',
        30: 'Bài 24. Choo Choo! Tàu hỏa',
        31: 'Bài 25. Có sáu chân thật tuyệt!!!',
        32: 'Ôn tập Đánh giá định kỳ 4',
        33: 'Đánh giá định kỳ 4',
        34: 'Tổng kết môn học & Triển lãm Robotics',
        35: 'Tổng kết môn học & Triển lãm Robotics',
    },
}

PPCT_ROB['3'] = PPCT_ROB['2'].copy()
PPCT_ROB['4'] = PPCT_ROB['2'].copy()
PPCT_ROB['TT'] = PPCT_ROB['1'].copy()

# Rotation Robotics (5, 6, 7, 8) - Taught on ODD weeks (Tuần 3, 5, 7, 9... 2 tiết/tuần)
PPCT_ROB['5'] = {
    0: 'Tiết 0: Định hướng môn học',
    1: 'Bài 1. Động cơ là gì?',
    2: 'Bài 2. Robot cố định vật',
    3: 'Bài 3. Robot nhận biết âm thanh',
    4: 'Bài 4. Bộ điều khiển từ xa',
    5: 'Bài 5. Động cơ Dynamixel hoạt động thế nào?',
    6: 'Bài 6. Phương tiện giao thông qua các thời đại',
    7: 'Ôn tập Đánh giá định kỳ 1',
    8: 'Đánh giá định kỳ 1',
    9: 'Bài 7. Cơ chế tăng giảm chiều dài tự động',
    10: 'Bài 8. Robot nhận biết vật thể bằng cách nào?',
    11: 'Bài 9. Số ngẫu nhiên',
    12: 'Bài 10. Robot hút bụi',
    13: 'Bài 11. Lực hấp dẫn',
    14: 'Bài 12. Domino trong Robotics',
    15: 'Bài 13. Robot phục vụ đời sống',
    16: 'Luyện tập & Thực hành sáng tạo Robotics',
    17: 'Ôn tập Đánh giá định kỳ 2',
    18: 'Đánh giá định kỳ 2',
    19: 'Bài 14. Robot thám hiểm',
    20: 'Bài 15. Cảm biến siêu âm',
    21: 'Bài 16. Tránh vật cản tự động',
    22: 'Bài 17. Robot theo dõi đường line',
    23: 'Bài 18. Lập trình vòng lặp trong điều khiển robot',
    24: 'Bài 19. Khám phá cánh tay robot',
    25: 'Ôn tập Đánh giá định kỳ 3',
    26: 'Đánh giá định kỳ 3',
    27: 'Bài 20. Robot phân loại sản phẩm',
    28: 'Bài 21. Cơ chế kẹp gắp vật thể',
    29: 'Bài 22. Dự án sáng tạo robot tự hành',
    30: 'Bài 23. Tối ưu hóa thuật toán điều khiển robot',
    31: 'Ôn tập Đánh giá định kỳ 4',
    32: 'Đánh giá định kỳ 4',
    33: 'Tổng kết môn học & Triển lãm Robotics',
    34: 'Tổng kết môn học & Triển lãm Robotics',
    35: 'Tổng kết môn học & Triển lãm Robotics',
}

PPCT_ROB['6'] = PPCT_ROB['5'].copy()
PPCT_ROB['7'] = PPCT_ROB['5'].copy()
PPCT_ROB['8'] = PPCT_ROB['5'].copy()


# ─────────────────────────────────────────────────────────────────────────────
# MASTER SCHEDULE (Tu 'Thoi khoa bieu - Dau Dinh Nguyen.xlsx')
# (day_idx 0..4, tiet_tkb 1..5): (lop, mon_display, is_rotation)
# ─────────────────────────────────────────────────────────────────────────────

MASTER_SCHEDULE_SANG = {
    (0, 3): ('4C1', 'Robotics', False),
    (0, 4): ('1A1', 'Tin học', False),
    (1, 1): ('1C1', 'Tin học', False),
    (1, 3): ('5C1', 'Tin - Robotics', True),
    (1, 4): ('5C1', 'Tin - Robotics', True),
    (2, 2): ('3A1', 'Robotics', False),
    (2, 4): ('3C1', 'Tin học', False),
    (3, 1): ('TT3', 'Tin học', False),
    (3, 2): ('1A1', 'Robotics', False),
    (3, 4): ('2C1', 'Robotics', False),
    (4, 1): ('2A1', 'Tin học', False),
    (4, 2): ('3C1', 'Robotics', False),
    (4, 4): ('6A1', 'Tin - Robotics', True),
    (4, 5): ('6A1', 'Tin - Robotics', True),
}

MASTER_SCHEDULE_CHIEU = {
    (0, 3): ('2C1', 'Tin học', False),
    (1, 3): ('7A1', 'Tin - Robotics', True),
    (1, 4): ('7A1', 'Tin - Robotics', True),
    (2, 3): ('4C1', 'Tin học', False),
    (2, 4): ('2A1', 'Robotics', False),
    (3, 1): ('3A1', 'Tin học', False),
    (3, 2): ('TTH 1', 'Tin học', False),
    (3, 3): ('TTH2', 'Tin học', False),
    (3, 4): ('1C1', 'Robotics', False),
    (4, 1): ('8A1', 'Tin - Robotics', True),
    (4, 2): ('8A1', 'Tin - Robotics', True),
}


def week_dates(tuan_so):
    delta = timedelta(weeks=tuan_so - 1)
    start = TUAN_01_START + delta
    end = start + timedelta(days=4)
    return start, end


def fmt_date(d):
    return f'{d.day:02d}/{d.month:02d}/{d.year}'


def day_label(d):
    thu = THU_LABELS[d.weekday()]
    return f'{thu} ({d.day:02d}/{d.month:02d})'


def classify_lop(lop):
    s = lop.strip().upper()
    if not s:
        return None
    if 'NGHỈ' in s or 'KHAI GIẢNG' in s or 'TỔNG DUYỆT' in s or 'LỄ' in s:
        return None
    if re.match(r'^[6789]', s):
        return 'THCS'
    return 'TTH_TH'


def is_rotation_lop(lop):
    s = lop.strip().upper()
    return any(s.startswith(p) for p in ROTATION_PREFIXES)


def rotation_mon(tuan_so):
    return 'Tin học' if tuan_so % 2 == 0 else 'Robotics'


def rotation_do_dung(tuan_so):
    return 'Phòng Tin học' if tuan_so % 2 == 0 else 'Bộ Kit Robotics'


def get_grade_key(lop):
    s = lop.strip().upper()
    if s.startswith('TT'):
        return 'TT3'
    m = re.match(r'^(\d)', s)
    if m:
        return m.group(1)
    return None


def get_ppct_lesson(lop, mon, ppct_num):
    grade = get_grade_key(lop)
    if grade is None:
        return ''
    if mon == 'Tin học':
        data = PPCT_TIN.get(grade, {})
    elif mon == 'Robotics':
        data = PPCT_ROB.get(grade, {})
    else:
        return ''
    if not data:
        return ''
    if ppct_num in data:
        return data[ppct_num]
    max_k = max(data.keys())
    if ppct_num > max_k:
        return data[max_k]
    return ''


def compute_ppct(tuan_so, day_idx=0):
    if tuan_so <= 1:
        return 0
    return tuan_so - 1


def compute_rotation_ppct(tuan_so, day_idx, is_even_week, period_index_in_day):
    if tuan_so <= 1:
        return 0
    if is_even_week:
        if tuan_so < 2:
            return 0
        k = (tuan_so - 2) // 2 + 1
        return (k - 1) * 2 + period_index_in_day + 1
    else:
        if tuan_so < 3:
            return 0
        k = (tuan_so - 3) // 2 + 1
        return (k - 1) * 2 + period_index_in_day + 1


def set_cell_text(cell, text):
    for p in cell.paragraphs:
        if p.runs:
            p.runs[0].text = text
            for r in p.runs[1:]:
                r.text = ''
            if p.runs[0].font.name != 'Times New Roman':
                p.runs[0].font.name = 'Times New Roman'
            if p.runs[0].font.size != Pt(13):
                p.runs[0].font.size = Pt(13)
            return
    p = cell.paragraphs[0]
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)


def save_safe(doc, path):
    try:
        doc.save(path)
        return path
    except PermissionError:
        base, ext = os.path.splitext(path)
        alt = base + '_v2' + ext
        doc.save(alt)
        print(f'   ⚠ Bi khoa, luu: {os.path.basename(alt)}')
        return alt


def update_headers(doc, tuan_so, start_date, end_date):
    tuan_str = f'TUẦN {tuan_so:02d}'
    fs = fmt_date(start_date)
    fe = fmt_date(end_date)
    ngay_soan = start_date - timedelta(days=3)

    for p in doc.paragraphs:
        txt = p.text.strip()
        if 'TUẦN' in txt.upper() and 'từ ngày' in txt:
            new = f'{tuan_str} (từ ngày {fs} đến ngày {fe})'
        elif re.match(r'^Ngày\s+\d+\s+tháng\s+\d+\s+năm\s+\d+', txt) and len(txt) < 45:
            new = f'Ngày {ngay_soan.day:02d} tháng {ngay_soan.month:02d} năm {ngay_soan.year}'
        elif 'sáng' in txt.lower() and 'tuần' in txt.lower():
            new = f'Buổi…sáng……..Tuần…{tuan_so:02d}…(Từ ngày…{fs} …đến ngày:… {fe}….)'
        elif 'chiều' in txt.lower() and 'tuần' in txt.lower():
            new = f'Buổi…chiều…..Tuần…{tuan_so:02d}…(Từ ngày…{fs}…đến ngày:… {fe}….)'
        else:
            continue

        if p.runs:
            p.runs[0].text = new
            for r in p.runs[1:]:
                r.text = ''
        else:
            p.text = new


def merge_and_format_block(tbl, start_row, end_row, start_col, end_col, text):
    """Merge a rectangular block of cells in tbl (from start_row..end_row and start_col..end_col)
    and format the resulting cell with centered text in Times New Roman 13pt.
    """
    for r in range(start_row, end_row + 1):
        if r >= len(tbl.rows):
            continue
        for c in range(start_col, min(end_col + 1, len(tbl.rows[r].cells))):
            cell = tbl.cell(r, c)
            for p in cell.paragraphs:
                p.clear()

    top_left = tbl.cell(start_row, start_col)
    bottom_right = tbl.cell(end_row, end_col)
    try:
        merged = top_left.merge(bottom_right)
    except Exception:
        merged = top_left

    while len(merged.paragraphs) > 1:
        p = merged.paragraphs[-1]
        p._element.getparent().remove(p._element)

    p = merged.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)

    tcPr = merged._tc.get_or_add_tcPr()
    for existing_va in tcPr.findall(qn('w:vAlign')):
        tcPr.remove(existing_va)
    vAlign = OxmlElement('w:vAlign')
    vAlign.set(qn('w:val'), 'center')
    tcPr.append(vAlign)
    return merged


def populate_lbg_table(tbl, schedule_map, tuan_so, start_date, is_sang=True):
    is_even = (tuan_so % 2 == 0)
    for day_idx in range(5):
        d = start_date + timedelta(days=day_idx)
        d_lbl = day_label(d)
        for tiet in range(1, 6):
            ri = day_idx * 5 + tiet
            if ri >= len(tbl.rows):
                continue
            row = tbl.rows[ri]
            set_cell_text(row.cells[0], d_lbl)
            set_cell_text(row.cells[1], str(tiet))
            key = (day_idx, tiet)
            if key in schedule_map:
                lop, mon_display, is_rot = schedule_map[key]
                if is_rot:
                    mon_int = rotation_mon(tuan_so)
                    dd = rotation_do_dung(tuan_so)
                    mon_lbl = 'Tin - Robotics'
                    p_idx = 0 if (day_idx, tiet - 1) not in schedule_map else 1
                    ppct = compute_rotation_ppct(tuan_so, day_idx, is_even, p_idx)
                    lesson = get_ppct_lesson(lop, mon_int, ppct) if tuan_so > 1 else 'Tiết 0: Định hướng môn học'
                else:
                    mon_lbl = mon_display
                    dd = 'Phòng Tin học' if 'Tin' in mon_lbl else 'Bộ Kit Robotics'
                    ppct = compute_ppct(tuan_so, day_idx) if tuan_so > 1 else 0
                    lesson = get_ppct_lesson(lop, mon_lbl, ppct) if tuan_so > 1 else 'Tiết 0: Định hướng môn học'
                ppct_str = str(ppct) if tuan_so > 1 else '1'
                if not lesson and ppct == 0:
                    lesson = 'Tiết 0: Định hướng môn học'
                set_cell_text(row.cells[2], ppct_str)
                set_cell_text(row.cells[3], mon_lbl)
                set_cell_text(row.cells[4], lop)
                set_cell_text(row.cells[5], lesson)
                set_cell_text(row.cells[6], dd)
            else:
                set_cell_text(row.cells[2], '')
                set_cell_text(row.cells[3], '')
                set_cell_text(row.cells[4], '')
                set_cell_text(row.cells[5], '')
                set_cell_text(row.cells[6], '')

    # ── Tuần 5: Block merge Nghỉ lễ 02/09 & Tổng duyệt khai giảng ──
    if tuan_so == 5:
        # Thứ 2, 3, 4 (Dòng 1..15, Cột 2..6): Nghỉ lễ 02/09 (cả Sáng và Chiều)
        merge_and_format_block(tbl, start_row=1, end_row=15, start_col=2, end_col=6, text='Nghỉ lễ 02/09')

        # Sáng Thứ 5, 6 (Dòng 16..25, Cột 2..6): Tổng duyệt khai giảng
        if is_sang:
            merge_and_format_block(tbl, start_row=16, end_row=25, start_col=2, end_col=6, text='Tổng duyệt khai giảng')
        # Chiều Thứ 5, 6: giữ nguyên lịch dạy bình thường theo TKB


def update_table_data(doc, tuan_so, start_date=None):
    if start_date is None:
        start_date, _ = week_dates(tuan_so)
    lbg_tbls = [t for t in doc.tables
                if len(t.columns) == 7
                and 'Lớp' in [c.text.strip() for c in t.rows[0].cells]]
    if len(lbg_tbls) >= 1:
        populate_lbg_table(lbg_tbls[0], MASTER_SCHEDULE_SANG, tuan_so, start_date, is_sang=True)
    if len(lbg_tbls) >= 2:
        populate_lbg_table(lbg_tbls[1], MASTER_SCHEDULE_CHIEU, tuan_so, start_date, is_sang=False)


def update_ky_ten(doc, start_date):
    ngay_soan = start_date - timedelta(days=3)
    ngay_str = f'Ngày {ngay_soan.day} tháng {ngay_soan.month} năm {ngay_soan.year}'
    # Update all sign tables (1r x 2c)
    for tbl in doc.tables:
        if len(tbl.rows) != 1 or len(tbl.columns) != 2:
            continue
        for row in tbl.rows:
            for cell in row.cells:
                txt = cell.text
                if 'Ngày' in txt and 'tháng' in txt and 'năm' in txt:
                    for p in cell.paragraphs:
                        for r in p.runs:
                            if 'Ngày' in r.text and 'tháng' in r.text:
                                m = re.search(r'Ngày\s*\d*\s*tháng\s*\d*\s*năm\s*\d+', r.text)
                                if m:
                                    r.text = r.text[:m.start()] + ngay_str + r.text[m.end():]


def update_sign_names(doc, to_truong_name, start_date=None):
    """Cập nhật phần ký tên của Tổ trưởng trong tất cả bảng chữ ký (1 hàng x 2 cột):
    - Dòng 1: Ngày tháng soạn (nếu có start_date thì tính Ngày d tháng m năm y, hoặc giữ ngày đã có)
    - Dòng 2: Tổ trưởng
    - Khoảng cách dòng trống cho chữ ký (3 dòng trống)
    - Dòng cuối: Tên tổ trưởng (Nguyễn Thị Ngọc cho TH / Nguyễn Thị Ngọc Ánh cho THCS)
    - TUYỆT ĐỐI BỎ chữ '(Ký tên, đóng dấu)'
    - Căn giữa toàn bộ, font Times New Roman 13pt
    """
    default_ngay = None
    if start_date is not None:
        ngay_soan = start_date - timedelta(days=3)
        default_ngay = f'Ngày {ngay_soan.day} tháng {ngay_soan.month} năm {ngay_soan.year}'

    for tbl in doc.tables:
        if len(tbl.rows) != 1 or len(tbl.columns) != 2:
            continue
        cell = tbl.rows[0].cells[1]  # Ô bên phải (Tổ trưởng)
        txt = cell.text
        if 'Tổ trưởng' not in txt:
            continue

        # Tìm chuỗi ngày tháng soạn
        m = re.search(r'Ngày\s*\d*\s*tháng\s*\d*\s*năm\s*\d+', txt)
        if default_ngay:
            ngay_str = default_ngay
        elif m and re.search(r'\d+', m.group(0).split('tháng')[0]):
            ngay_str = m.group(0).strip()
        else:
            ngay_str = 'Ngày   tháng  năm 2026'

        # Xóa sạch các paragraph cũ, chỉ giữ 1 paragraph
        while len(cell.paragraphs) > 1:
            p_extra = cell.paragraphs[-1]
            p_extra._element.getparent().remove(p_extra._element)

        p = cell.paragraphs[0]
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Run 1: Ngày tháng năm
        r1 = p.add_run(ngay_str + '\n')
        r1.font.name = 'Times New Roman'
        r1.font.size = Pt(13)

        # Run 2: 'Tổ trưởng' + 3 dòng trống (tổng 4 dấu xuống dòng) để chỗ cho chữ ký
        # BỎ HOÀN TOÀN '(Ký tên, đóng dấu)'
        spacing = '\n\n\n\n' if to_truong_name else '\n\n\n'
        r2 = p.add_run('Tổ trưởng' + spacing)
        r2.font.name = 'Times New Roman'
        r2.font.size = Pt(13)

        # Run 3: Tên tổ trưởng (nếu có)
        if to_truong_name:
            r3 = p.add_run(to_truong_name)
            r3.font.name = 'Times New Roman'
            r3.font.size = Pt(13)


def compact_sign_tables(doc):
    """Reduce dòng chấm nhận xét in sign tables to MAX_DONG_CHAM lines
    to ensure the LBG table + sign section fit within 1 page.

    Sign table cell [0,0] run structure:
      Run 0: 'Kiểm tra, nhận xét'
      Run 1..N: '\n…………………………………………………………………'
    We keep at most MAX_DONG_CHAM dòng chấm runs.
    """
    for tbl in doc.tables:
        if len(tbl.rows) != 1 or len(tbl.columns) != 2:
            continue
        cell = tbl.rows[0].cells[0]  # Left cell (Kiểm tra, nhận xét)
        for p in cell.paragraphs:
            runs = p.runs
            if not runs:
                continue
            # Find runs with dòng chấm
            dot_runs = []
            for ri, r in enumerate(runs):
                if '……' in r.text:
                    dot_runs.append((ri, r))
            # Keep only MAX_DONG_CHAM dòng chấm runs
            if len(dot_runs) > MAX_DONG_CHAM:
                for ri, r in dot_runs[MAX_DONG_CHAM:]:
                    r.text = ''

    # Also remove excessive empty paragraphs between tables
    body = doc.element.body
    children = list(body)
    for i in range(len(children) - 1):
        c = children[i]
        tag = c.tag.split('}')[-1] if '}' in c.tag else c.tag
        if tag == 'tbl':
            # Check if this is a sign table
            is_sign = False
            for tbl in doc.tables:
                if tbl._element is c and len(tbl.rows) == 1 and len(tbl.columns) == 2:
                    is_sign = True
                    break
            if is_sign:
                # Remove excessive empty paragraphs after sign table
                # Keep at most 0 empty paragraphs
                j = i + 1
                empty_count = 0
                while j < len(children):
                    nc = children[j]
                    nc_tag = nc.tag.split('}')[-1] if '}' in nc.tag else nc.tag
                    if nc_tag != 'p':
                        break
                    txt = ''.join(t.text or '' for t in nc.findall('.//' + qn('w:t')))
                    if txt.strip():
                        break
                    empty_count += 1
                    j += 1
                # Remove all empty paragraphs after sign table
                if empty_count > 0:
                    for nc in children[i+1:i+1+empty_count]:
                        nc_tag = nc.tag.split('}')[-1] if '}' in nc.tag else nc.tag
                        if nc_tag == 'p' and nc_tag != 'sectPr':
                            txt = ''.join(t.text or '' for t in nc.findall('.//' + qn('w:t')))
                            if not txt.strip():
                                body.remove(nc)


def remove_second_copy(doc):
    body = doc.element.body
    children = list(body)
    tuan_paras = []
    for c in children:
        tag = c.tag.split('}')[-1] if '}' in c.tag else c.tag
        if tag == 'p':
            txt = ''.join(t.text or '' for t in c.findall('.//' + qn('w:t')))
            if 'TUẦN' in txt.upper() and 'từ ngày' in txt:
                tuan_paras.append(c)
    if len(tuan_paras) < 2:
        return 0
    nhanxet_el = None
    for c in children:
        tag = c.tag.split('}')[-1] if '}' in c.tag else c.tag
        if tag == 'p':
            txt = ''.join(t.text or '' for t in c.findall('.//' + qn('w:t')))
            if 'nhận xét' in txt.lower() and 'bgh' in txt.lower():
                nhanxet_el = c
                break
    if nhanxet_el is None:
        return 0
    children = list(body)
    start_idx = children.index(tuan_paras[1])
    end_idx = children.index(nhanxet_el)
    removed = 0
    for el in children[start_idx:end_idx]:
        tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
        if tag != 'sectPr':
            body.remove(el)
            removed += 1
    return removed


def remove_cover_page(doc):
    """Xóa trang bìa đầu (bảng thông tin giáo viên và các đoạn văn trước tiêu đề TUẦN)."""
    body = doc.element.body
    children = list(body)
    tuan_el = None
    for c in children:
        tag = c.tag.split('}')[-1] if '}' in c.tag else c.tag
        if tag == 'p':
            txt = ''.join(t.text or '' for t in c.findall('.//' + qn('w:t')))
            if 'TUẦN' in txt.upper() and 'từ ngày' in txt:
                tuan_el = c
                break
    if tuan_el is not None:
        idx = children.index(tuan_el)
        for el in children[:idx]:
            tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
            if tag != 'sectPr':
                body.remove(el)


def remove_end_evaluation_table(doc):
    """Xóa bảng nhận xét tiến độ cuối cùng (tiêu đề + bảng 42 dòng x 3 cột + các đoạn sau đó)."""
    body = doc.element.body
    children = list(body)
    nhanxet_el = None
    for c in children:
        tag = c.tag.split('}')[-1] if '}' in c.tag else c.tag
        if tag == 'p':
            txt = ''.join(t.text or '' for t in c.findall('.//' + qn('w:t')))
            if 'nhận xét' in txt.lower() and 'bgh' in txt.lower():
                nhanxet_el = c
                break
    if nhanxet_el is not None:
        idx = children.index(nhanxet_el)
        for el in children[idx:]:
            tag = el.tag.split('}')[-1] if '}' in el.tag else el.tag
            if tag != 'sectPr':
                body.remove(el)


def fix_all_fonts(doc):
    fixed = 0
    for tbl in doc.tables:
        if len(tbl.columns) != 7:
            continue
        header_texts = [c.text.strip() for c in tbl.rows[0].cells]
        if 'Lớp' not in header_texts:
            continue
        for ri in range(0, len(tbl.rows)):
            for ci in range(len(tbl.rows[ri].cells)):
                cell = tbl.rows[ri].cells[ci]
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.font.name != 'Times New Roman':
                            run.font.name = 'Times New Roman'
                            fixed += 1
                        if run.font.size is None or run.font.size != Pt(13):
                            run.font.size = Pt(13)
    return fixed


def add_page_break_after_table(tbl_element):
    nxt = tbl_element.getnext()
    if nxt is not None and (nxt.tag.split('}')[-1] if '}' in nxt.tag else nxt.tag) == 'p':
        brs = nxt.findall('.//' + qn('w:br'))
        for b in brs:
            if b.get(qn('w:type')) == 'page':
                return
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    tbl_element.addnext(p)


def fix_page_breaks(doc):
    body = doc.element.body
    children = list(body)
    sign_tables = []
    for c in children:
        tag = c.tag.split('}')[-1] if '}' in c.tag else c.tag
        if tag == 'tbl':
            for tbl in doc.tables:
                if tbl._element is c and len(tbl.rows) == 1 and len(tbl.columns) == 2:
                    sign_tables.append(c)
                    break
    if len(sign_tables) < 2:
        return
    sign_sang = sign_tables[0]
    sign_chieu = sign_tables[1]

    def find_para(keyword):
        for c in list(body):
            if (c.tag.split('}')[-1] if '}' in c.tag else c.tag) == 'p':
                txt = ''.join(t.text or '' for t in c.findall('.//' + qn('w:t')))
                if keyword.lower() in txt.lower():
                    return c
        return None

    chieu_el = find_para('chiều')
    nhanxet_el = find_para('nhận xét của bgh')

    def remove_empty_between(el_a, el_b):
        ch = list(body)
        if el_a not in ch or el_b not in ch:
            return 0
        ia, ib = ch.index(el_a), ch.index(el_b)
        removed = 0
        for c in ch[ia + 1:ib]:
            if (c.tag.split('}')[-1] if '}' in c.tag else c.tag) == 'p':
                if not ''.join(t.text or '' for t in c.findall('.//' + qn('w:t'))).strip():
                    body.remove(c)
                    removed += 1
        return removed

    if chieu_el is not None:
        remove_empty_between(sign_sang, chieu_el)
    if nhanxet_el is not None:
        remove_empty_between(sign_chieu, nhanxet_el)

    add_page_break_after_table(sign_sang)
    if nhanxet_el is not None:
        add_page_break_after_table(sign_chieu)


def clear_row_data(row):
    for ci in range(2, min(7, len(row.cells))):
        for p in row.cells[ci].paragraphs:
            for r in p.runs:
                r.text = ''
            if not p.runs and p.text.strip():
                p.text = ''


def filter_for_cap(doc, keep_cap):
    for tbl in doc.tables:
        if len(tbl.columns) != 7:
            continue
        header_texts = [c.text.strip() for c in tbl.rows[0].cells]
        if 'Lớp' not in header_texts:
            continue
        for ri in range(1, len(tbl.rows)):
            row = tbl.rows[ri]
            lop = row.cells[4].text.strip()
            # Bỏ qua nếu là ô nghỉ lễ hoặc tổng duyệt / sự kiện
            if any(kw in lop.upper() for kw in ['NGHỈ', 'LỄ', 'TỔNG DUYỆT', 'KHAI GIẢNG']):
                continue
            cap = classify_lop(lop)
            if cap is not None and cap != keep_cap:
                clear_row_data(row)


def generate_lbg(tuan_so):
    start_date, end_date = week_dates(tuan_so)
    loai = 'CHẴN → Tin học' if tuan_so % 2 == 0 else 'LẺ → Robotics'
    print(f'📅 Tuần {tuan_so:02d}: {fmt_date(start_date)} -> {fmt_date(end_date)} [{loai}]')

    doc = Document(TEMPLATE)
    remove_second_copy(doc)

    # User yêu cầu: các tuần lẻ không cần để bìa và bảng nhận xét cuối (chỉ file gộp cả năm mới cần)
    remove_cover_page(doc)
    remove_end_evaluation_table(doc)

    update_headers(doc, tuan_so, start_date, end_date)
    update_table_data(doc, tuan_so, start_date)
    fix_all_fonts(doc)
    update_ky_ten(doc, start_date)
    update_sign_names(doc, '', start_date)
    compact_sign_tables(doc)
    fix_page_breaks(doc)

    main_name = f'Lịch báo giảng - Tuần {tuan_so:02d}.docx'
    main_path = os.path.join(LBG_DIR, main_name)
    saved_main = save_safe(doc, main_path)

    for keep_cap, suffix, label, to_truong in [
        ('TTH_TH', 'TTH+TH', 'TTH + TH', TO_TRUONG_TH),
        ('THCS', 'THCS', 'THCS', TO_TRUONG_THCS),
    ]:
        d2 = Document(saved_main)
        filter_for_cap(d2, keep_cap)
        update_sign_names(d2, to_truong, start_date)
        compact_sign_tables(d2)
        fix_page_breaks(d2)
        split_name = f'Lịch báo giảng - Tuần {tuan_so:02d} ({suffix}).docx'
        split_path = os.path.join(LBG_DIR, split_name)
        save_safe(d2, split_path)


if __name__ == '__main__':
    if len(sys.argv) < 2:
        print('Cách dùng: python generate_lbg.py <so_tuan> [den_tuan]')
        print('Ví dụ:     python generate_lbg.py 3')
        print('           python generate_lbg.py 3 35')
        sys.exit(1)
    try:
        t_start = int(sys.argv[1])
        t_end = int(sys.argv[2]) if len(sys.argv) > 2 else t_start
    except ValueError:
        print('Lỗi: Số tuần phải là số nguyên.')
        sys.exit(1)

    for w in range(t_start, t_end + 1):
        generate_lbg(w)
    print(f'\\n✅ Đã tạo xong từ tuần {t_start:02d} đến {t_end:02d}!')
