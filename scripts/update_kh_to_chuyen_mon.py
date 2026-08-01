import docx
import os
import sys
import copy
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding='utf-8')

SRC_PATH = r'C:\Users\bmngu\Downloads\09.07.26. Kế hoạch tổ chuyên môn (THCS).docx'
TGT_DIR = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch tổ chuyên môn'
TGT_FILE_30 = os.path.join(TGT_DIR, '30.07.26. Kế hoạch tổ chuyên môn (THCS).docx')
TGT_FILE_09 = os.path.join(TGT_DIR, '09.07.26. Kế hoạch tổ chuyên môn (THCS).docx')
TGT_FILE_BASE = os.path.join(TGT_DIR, 'Kế hoạch tổ chuyên môn (THCS).docx')

def set_table_borders(table):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        for child in list(tblPr[0]):
            if child.tag.endswith('tblBorders'):
                tblPr[0].remove(child)
        borders_xml = parse_xml(
            r'<w:tblBorders %s>'
            r'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        tblPr[0].append(borders_xml)

def set_font_all(doc):
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.name = 'Times New Roman'
    for tbl in doc.tables:
        set_table_borders(tbl)
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = 'Times New Roman'

def process_merge():
    src_doc = docx.Document(SRC_PATH)
    tgt_doc = docx.Document(TGT_FILE_30)

    src_body = src_doc.element.body
    tgt_body = tgt_doc.element.body

    src_children = list(src_body)
    tgt_children = list(tgt_body)

    # 1. Find SRC boundaries
    src_vi_idx = None
    src_viii_idx = None
    src_sectpr_idx = None

    for idx, child in enumerate(src_children):
        tag = child.tag.split('}')[-1]
        if tag == 'sectPr':
            src_sectpr_idx = idx
            continue
        text = ""
        if tag == 'p':
            p = docx.text.paragraph.Paragraph(child, src_doc)
            text = p.text.strip()
        elif tag == 'tbl':
            tbl = docx.table.Table(child, src_doc)
            text = tbl.cell(0, 0).text.strip() if len(tbl.rows) > 0 and len(tbl.columns) > 0 else ""

        if text.startswith('VI. KẾ HOẠCH GIẢNG DẠY') and src_vi_idx is None:
            src_vi_idx = idx
        elif text.startswith('VIII. Nhiệm vụ:') and src_viii_idx is None:
            src_viii_idx = idx

    # 2. Find TGT boundaries
    tgt_vi_idx = None
    tgt_viii_idx = None
    tgt_sectpr_idx = None

    for idx, child in enumerate(tgt_children):
        tag = child.tag.split('}')[-1]
        if tag == 'sectPr':
            tgt_sectpr_idx = idx
            continue
        text = ""
        if tag == 'p':
            p = docx.text.paragraph.Paragraph(child, tgt_doc)
            text = p.text.strip()
        elif tag == 'tbl':
            tbl = docx.table.Table(child, tgt_doc)
            text = tbl.cell(0, 0).text.strip() if len(tbl.rows) > 0 and len(tbl.columns) > 0 else ""

        if text.startswith('VI. KẾ HOẠCH GIẢNG DẠY') and tgt_vi_idx is None:
            tgt_vi_idx = idx
        elif text.startswith('VIII. Nhiệm vụ:') and tgt_viii_idx is None:
            tgt_viii_idx = idx

    # Clear target body completely except sectPr
    for child in list(tgt_body):
        tgt_body.remove(child)

    # Re-assemble body elements
    # Part A: SRC Intro + I + II + III + IV + V (from 0 to src_vi_idx)
    for elem in src_children[0:src_vi_idx]:
        tgt_body.append(copy.deepcopy(elem))

    # Part B: TGT VI + VII (from tgt_vi_idx to tgt_viii_idx)
    for elem in tgt_children[tgt_vi_idx:tgt_viii_idx]:
        tgt_body.append(copy.deepcopy(elem))

    # Part C: SRC VIII (from src_viii_idx to src_sectpr_idx)
    for elem in src_children[src_viii_idx:src_sectpr_idx]:
        tgt_body.append(copy.deepcopy(elem))

    # Part D: TGT sectPr
    tgt_body.append(copy.deepcopy(tgt_children[tgt_sectpr_idx]))

    # Set font & table borders
    set_font_all(tgt_doc)

    # Save updated target file
    tgt_doc.save(TGT_FILE_30)
    print(f"Successfully updated and saved merged document to {TGT_FILE_30}")

    # Synchronize to other identical copies in folder
    for sync_target in [TGT_FILE_09, TGT_FILE_BASE]:
        if os.path.exists(sync_target):
            tgt_doc.save(sync_target)
            print(f"Synchronized to {sync_target}")

def verify():
    doc = docx.Document(TGT_FILE_30)
    print("=== VERIFYING SECTION HEADINGS & INDEXES ===")
    for idx, child in enumerate(doc.element.body):
        tag = child.tag.split('}')[-1]
        if tag == 'p':
            p = docx.text.paragraph.Paragraph(child, doc)
            txt = p.text.strip()
            for prefix in ['I.', 'II.', 'III.', 'IV.', 'V.', 'VI.', 'VII.', 'VIII.']:
                if txt.startswith(prefix) and len(txt) < 80:
                    print(f"Element [{idx:3d}] P: {txt}")

    print(f"\nTotal body elements: {len(list(doc.element.body))}")
    print(f"Total sections in doc: {len(doc.sections)}")

    sec = doc.sections[0]
    h_drawings = sec.header._element.xpath('.//*[local-name()="drawing"]')
    print(f"Header drawings count: {len(h_drawings)}")
    print(f"Footer paragraphs count: {len(sec.footer.paragraphs)}")

if __name__ == '__main__':
    process_merge()
    print()
    verify()
