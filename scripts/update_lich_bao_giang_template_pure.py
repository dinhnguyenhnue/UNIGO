import os, sys, re, shutil
import openpyxl, docx
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

sys.stdout.reconfigure(encoding='utf-8')

def parse_schedule_from_excel(xlsx_path):
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb['TKB_LOP_SC']

    classes = []
    for c in range(3, ws.max_column + 1, 2):
        cls_name = ws.cell(row=4, column=c).value
        if cls_name:
            classes.append((str(cls_name).strip(), c, c+1))

    # Map (thu_str, tiet_int) -> (mon_name, lop_name)
    sang_dict = {}
    chieu_dict = {}

    thu_map = {'2': 'Hai', '3': 'Ba', '4': 'Tư', '5': 'Năm', '6': 'Sáu'}

    current_thu = ''
    for r in range(6, ws.max_row + 1):
        thu_val = ws.cell(row=r, column=1).value
        tiet_val = ws.cell(row=r, column=2).value
        if thu_val:
            current_thu = str(thu_val).strip()
        if not tiet_val:
            continue
        tiet_str = str(tiet_val).strip()
        if not tiet_str.isdigit():
            continue
        tiet_num = int(tiet_str)
        thu_name = thu_map.get(current_thu, current_thu)
        
        for cls_name, sang_col, chieu_col in classes:
            sang_val = str(ws.cell(row=r, column=sang_col).value or '').strip()
            chieu_val = str(ws.cell(row=r, column=chieu_col).value or '').strip()
            
            if 'nguyên' in sang_val.lower() or ('tin' in sang_val.lower() and 'nguyên' in sang_val.lower()) or ('rob' in sang_val.lower() and 'nguyên' in sang_val.lower()):
                mon_name = "Robotics" if 'rob' in sang_val.lower() else "Tin học"
                sang_dict[(thu_name, tiet_num)] = (mon_name, cls_name)
                
            if 'nguyên' in chieu_val.lower() or ('tin' in chieu_val.lower() and 'nguyên' in chieu_val.lower()) or ('rob' in chieu_val.lower() and 'nguyên' in chieu_val.lower()):
                mon_name = "Robotics" if 'rob' in chieu_val.lower() else "Tin học"
                chieu_dict[(thu_name, tiet_num)] = (mon_name, cls_name)

    return sang_dict, chieu_dict

def format_cell_text(cell, text, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.text = text
    cell.paragraphs[0].alignment = align
    cell.paragraphs[0].paragraph_format.space_before = Pt(1)
    cell.paragraphs[0].paragraph_format.space_after = Pt(1)
    cell.paragraphs[0].paragraph_format.line_spacing = 1.05
    for run in cell.paragraphs[0].runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.bold = bold

def update_template_pure(pristine_template_path, target_docx_path, sang_dict, chieu_dict):
    # Restore pristine template to ensure header/footer/logo/structure are 100% intact
    doc = docx.Document(pristine_template_path)
    
    # Update Paragraph titles without touching header/footer/sections
    for p in doc.paragraphs:
        txt = p.text.strip()
        if 'TUẦN 01' in txt or 'TUẦN' in txt:
            p.text = '                                                   TUẦN 01 (từ ngày 04/08/2026 đến ngày 08/08/2026)'
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
                r.bold = True
        elif 'Ngày ........ tháng........ năm 202' in txt:
            p.text = 'Ngày 04 tháng 08 năm 2026'
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)

    # Update Table 0 (Teacher info) if needed
    if len(doc.tables) > 0:
        cell0 = doc.tables[0].rows[0].cells[0]
        if 'Đậu Đình Nguyên' not in cell0.text:
            cell0.text = cell0.text.replace('Họ và tên giáo viên  :              ', 'Họ và tên giáo viên  :  Đậu Đình Nguyên            ')
            cell0.text = cell0.text.replace('Chức vụ hiện nay                  : Giáo viên', 'Chức vụ hiện nay                  : Giáo viên Tin học & Robotics')
            for r in cell0.paragraphs[0].runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)

    # Populate Table 1 (Buổi Sáng) and Table 3 (Buổi Chiều)
    day_list = ['Hai', 'Ba', 'Tư', 'Năm', 'Sáu']
    
    def fill_table_dict(table, sched_dict):
        # Table has 26 rows (row 0 is header, rows 1..25 are 5 days x 5 periods)
        row_idx = 1
        for day in day_list:
            for tiet in range(1, 6):
                if row_idx < len(table.rows):
                    row = table.rows[row_idx]
                    
                    # Col 1: Tiết TKB
                    format_cell_text(row.cells[1], str(tiet), align=WD_ALIGN_PARAGRAPH.CENTER)
                    
                    if (day, tiet) in sched_dict:
                        mon, lop = sched_dict[(day, tiet)]
                        is_monday = (day == 'Hai')
                        ppct_val = "" if is_monday else "1"
                        ten_bai_val = "" if is_monday else "Tiết 0: Định hướng môn học"
                        dodung_val = "" if is_monday else ("Máy tính, Tivi" if mon == "Tin học" else "Bộ kit Robotics, Máy tính")
                        
                        format_cell_text(row.cells[2], ppct_val, align=WD_ALIGN_PARAGRAPH.CENTER)
                        format_cell_text(row.cells[3], mon, align=WD_ALIGN_PARAGRAPH.LEFT)
                        format_cell_text(row.cells[4], lop, align=WD_ALIGN_PARAGRAPH.CENTER)
                        format_cell_text(row.cells[5], ten_bai_val, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
                        format_cell_text(row.cells[6], dodung_val, align=WD_ALIGN_PARAGRAPH.LEFT)
                    else:
                        # Clear empty slot
                        format_cell_text(row.cells[2], "", align=WD_ALIGN_PARAGRAPH.CENTER)
                        format_cell_text(row.cells[3], "", align=WD_ALIGN_PARAGRAPH.LEFT)
                        format_cell_text(row.cells[4], "", align=WD_ALIGN_PARAGRAPH.CENTER)
                        format_cell_text(row.cells[5], "", align=WD_ALIGN_PARAGRAPH.JUSTIFY)
                        format_cell_text(row.cells[6], "", align=WD_ALIGN_PARAGRAPH.LEFT)
                row_idx += 1

    if len(doc.tables) > 1:
        fill_table_dict(doc.tables[1], sang_dict)
        
    if len(doc.tables) > 3:
        fill_table_dict(doc.tables[3], chieu_dict)

    doc.save(target_docx_path)
    print(f"Successfully updated template in-place preserving 100% header/footer/logo at: {target_docx_path}")

def main():
    pristine_template_path = r"D:\UNIGO\Hệ thống mẫu văn bản\Lịch báo giảng.docx"
    xlsx_path = r"D:\UNIGO\TKB toàn trường CHECK - lần 2.1.xlsx"
    target_docx_path = r"D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Lịch báo giảng\Lịch báo giảng.docx"
    
    sang_dict, chieu_dict = parse_schedule_from_excel(xlsx_path)
    update_template_pure(pristine_template_path, target_docx_path, sang_dict, chieu_dict)

if __name__ == "__main__":
    main()
