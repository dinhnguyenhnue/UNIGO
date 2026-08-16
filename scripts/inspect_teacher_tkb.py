import sys, os
sys.stdout.reconfigure(encoding='utf-8')
import openpyxl
from docx import Document

tkb_dir = r'D:\UNIGO\Thời khóa biểu giáo viên'

print("=== INSPECTING: Thời khóa biểu - Đậu Đình Nguyên.xlsx ===")
path_xlsx = os.path.join(tkb_dir, 'Thời khóa biểu - Đậu Đình Nguyên.xlsx')
if os.path.exists(path_xlsx):
    wb = openpyxl.load_workbook(path_xlsx, data_only=True)
    for sname in wb.sheetnames:
        ws = wb[sname]
        print(f"\nSheet: {sname} ({ws.max_row}x{ws.max_column})")
        for r in range(1, min(ws.max_row + 1, 25)):
            vals = [ws.cell(r, c).value for c in range(1, min(ws.max_column + 1, 15))]
            if any(vals):
                print(f"Row {r:2d}: {vals}")

print("\n=== INSPECTING: Thời khóa biểu - Đậu Đình Nguyên (NEW).xlsx ===")
path_xlsx_new = os.path.join(tkb_dir, 'Thời khóa biểu - Đậu Đình Nguyên (NEW).xlsx')
if os.path.exists(path_xlsx_new):
    wb = openpyxl.load_workbook(path_xlsx_new, data_only=True)
    for sname in wb.sheetnames:
        ws = wb[sname]
        print(f"\nSheet: {sname} ({ws.max_row}x{ws.max_column})")
        for r in range(1, min(ws.max_row + 1, 25)):
            vals = [ws.cell(r, c).value for c in range(1, min(ws.max_column + 1, 15))]
            if any(vals):
                print(f"Row {r:2d}: {vals}")

print("\n=== INSPECTING: Thời khóa biểu - Đậu Đình Nguyên.docx ===")
path_docx = os.path.join(tkb_dir, 'Thời khóa biểu - Đậu Đình Nguyên.docx')
if os.path.exists(path_docx):
    doc = Document(path_docx)
    for i, p in enumerate(doc.paragraphs[:10]):
        if p.text.strip():
            print(f"P[{i}]: {p.text.strip()}")
    print(f"Tables: {len(doc.tables)}")
    if doc.tables:
        for ri, r in enumerate(doc.tables[0].rows[:15]):
            print(f"Table row {ri}: {[c.text.strip().replace('\n', ' ') for c in r.cells]}")
