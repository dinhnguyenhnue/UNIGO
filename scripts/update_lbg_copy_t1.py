import os
import sys
import docx
from docx.shared import Pt
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding='utf-8')

SRC_FILE = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Lịch báo giảng\Lịch báo giảng - Tuần 01.docx'
DST_FILE = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Lịch báo giảng\Lịch báo giảng.docx'

def set_table_borders(table):
    """Sets full black grid borders."""
    table.style = 'Table Grid'
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        for child in list(tblPr[0]):
            if child.tag.endswith('tblBorders'):
                tblPr[0].remove(child)
        borders_xml = parse_xml(
            r'<w:tblBorders %s>'
            r'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        tblPr[0].append(borders_xml)

def set_font_all(doc):
    """Sets Times New Roman 13pt across all paragraphs and tables."""
    for p in doc.paragraphs:
        p.style.font.name = 'Times New Roman'
        p.style.font.size = Pt(13)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)
            
    for t in doc.tables:
        set_table_borders(t)
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.style.font.name = 'Times New Roman'
                    p.style.font.size = Pt(13)
                    for r in p.runs:
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(13)

def get_slot_map(table_src):
    """Parses teaching slots from source table."""
    day_map = {'Hai': 0, 'Ba': 1, 'Tư': 2, 'Năm': 3, 'Sáu': 4}
    slot_map = {}
    for r in table_src.rows[1:]:
        day_text = r.cells[0].text.strip()
        period_text = r.cells[1].text.strip()
        
        day_idx = None
        for k, v in day_map.items():
            if k in day_text:
                day_idx = v
                break
                
        if day_idx is not None and period_text.isdigit():
            period = int(period_text)
            slot_map[(day_idx, period)] = (
                r.cells[2].text.strip(),  # ppct
                r.cells[3].text.strip(),  # mon
                r.cells[4].text.strip(),  # lop
                r.cells[5].text.strip(),  # ten_bai
                r.cells[6].text.strip()   # do_dung
            )
    return slot_map

DAYS_LABEL = [
    'Hai (03/08)',
    'Ba (04/08)',
    'Tư (05/08)',
    'Năm (06/08)',
    'Sáu (07/08)'
]

def fill_25_row_table(table_dst, slot_map):
    """Populates 25 rows (5 days x 5 periods), preserving empty periods."""
    for day_idx in range(5):
        day_label = DAYS_LABEL[day_idx]
        for period in range(1, 6):
            r_idx = day_idx * 5 + period
            if r_idx < len(table_dst.rows):
                row = table_dst.rows[r_idx]
                row.cells[0].text = day_label
                row.cells[1].text = str(period)
                
                if (day_idx, period) in slot_map:
                    ppct, mon, lop, ten_bai, do_dung = slot_map[(day_idx, period)]
                    row.cells[2].text = ppct
                    row.cells[3].text = mon
                    row.cells[4].text = lop
                    row.cells[5].text = ten_bai
                    row.cells[6].text = do_dung
                else:
                    row.cells[2].text = ''
                    row.cells[3].text = ''
                    row.cells[4].text = ''
                    row.cells[5].text = ''
                    row.cells[6].text = ''

def main():
    print(f"Reading source {SRC_FILE}...")
    doc_src = docx.Document(SRC_FILE)
    
    print(f"Reading destination template {DST_FILE}...")
    doc_dst = docx.Document(DST_FILE)
    
    # 1. Update Table 0 (Teacher info)
    doc_dst.tables[0].rows[0].cells[0].text = doc_src.tables[0].rows[0].cells[0].text
    
    # 2. Update Paragraphs
    for p in doc_dst.paragraphs:
        txt = p.text.strip()
        if 'TUẦN 01' in txt or 'TUẦN 1' in txt or 'từ ngày ……đến' in txt:
            p.text = '                                                   TUẦN 01 (từ ngày 03/08/2026 đến ngày 07/08/2026)'
        elif 'Ngày ........ tháng' in txt:
            p.text = 'Ngày 31 tháng 07 năm 2026'
        elif 'Buổi' in txt and 'Tuần' in txt:
            if 'Sáng' in txt or 'sáng' in txt or p == doc_dst.paragraphs[3]:
                p.text = 'Buổi……sáng…..Tuần…01…(Từ ngày…03/08/2026 …đến ngày:…07/08/2026….)                                                                                           '
            elif 'Chiều' in txt or 'chiều' in txt or p == doc_dst.paragraphs[9]:
                p.text = 'Buổi……chiều…..Tuần…01…(Từ ngày…03/08/2026 …đến ngày:… 07/08/2026….)                                                                                           '

    # 3. Parse maps & Populate 25-row tables
    map_sang = get_slot_map(doc_src.tables[1])
    map_chieu = get_slot_map(doc_src.tables[3])
    
    print("Populating Table 1 (Sáng) - 25 rows...")
    fill_25_row_table(doc_dst.tables[1], map_sang)
    
    print("Populating Table 3 (Chiều) - 25 rows...")
    fill_25_row_table(doc_dst.tables[3], map_chieu)
    
    # 4. Update signature / review blocks
    if len(doc_src.tables) > 2 and len(doc_dst.tables) > 2:
        doc_dst.tables[2].rows[0].cells[0].text = doc_src.tables[2].rows[0].cells[0].text
        doc_dst.tables[2].rows[0].cells[1].text = doc_src.tables[2].rows[0].cells[1].text
        
    if len(doc_src.tables) > 4 and len(doc_dst.tables) > 4:
        doc_dst.tables[4].rows[0].cells[0].text = doc_src.tables[4].rows[0].cells[0].text
        doc_dst.tables[4].rows[0].cells[1].text = doc_src.tables[4].rows[0].cells[1].text

    # 5. Apply font & borders
    set_font_all(doc_dst)
    
    try:
        doc_dst.save(DST_FILE)
        print(f"Successfully saved {DST_FILE}")
    except PermissionError:
        print("ERROR: Permission denied when saving file. Please close Word if open.")

if __name__ == '__main__':
    main()
