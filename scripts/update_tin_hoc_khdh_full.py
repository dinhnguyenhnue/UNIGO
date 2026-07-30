import os
import sys
import shutil
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

sys.stdout.reconfigure(encoding='utf-8')

SRC_4COT = r'D:\UNIGO\Phân phối chương trình\Ke-hoach\Ke_hoach_Tin_hoc_Tien_tieu_hoc_Lop_1_Lop_2_4_cot.docx'

DIR_MAU = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm'
DIR_TUNG_LOP = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch dạy học Tin học từng lớp'
DIR_DIST = r'D:\UNIGO\Phân phối chương trình\Tin học'

MAIN_ALL = os.path.join(DIR_MAU, 'Kế hoạch dạy học môn Tin học 2026-2027.docx')
MAIN_TH = os.path.join(DIR_MAU, 'Kế hoạch dạy học môn Tin học (TH) - 2026 - 2027.docx')

FILE_L1 = os.path.join(DIR_TUNG_LOP, 'Kế hoạch dạy học môn Tin học - Lớp 1 - 2026 - 2027.docx')
FILE_L2 = os.path.join(DIR_TUNG_LOP, 'Kế hoạch dạy học môn Tin học - Lớp 2 - 2026 - 2027.docx')
FILE_TTH = os.path.join(DIR_TUNG_LOP, 'Kế hoạch dạy học môn Tin học - Tiền tiểu học - 2026 - 2027.docx')

COL_WIDTHS = [0.9, 3.8, 1.1, 1.4, 9.3] # total ~ 16.5 cm

def set_table_borders(table):
    """Sets explicit XML full black grid borders for a table."""
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        for child in list(tblPr[0]):
            if child.tag.endswith('tblBorders'):
                tblPr[0].remove(child)
                
        borders_xml = parse_xml(
            r'<w:tblBorders %s>'
            r'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        tblPr[0].append(borders_xml)

def set_font_all(doc):
    """Sets Times New Roman 13pt for all paragraphs, runs, tables, cells in the document."""
    for p in doc.paragraphs:
        p.style.font.name = 'Times New Roman'
        p.style.font.size = Pt(13)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)
            
    for t in doc.tables:
        set_table_borders(t)
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.style.font.name = 'Times New Roman'
                    p.style.font.size = Pt(13)
                    for r in p.runs:
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(13)

def set_cell_width_and_align(cell, width_cm, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Sets exact cell width and text alignment."""
    cell.width = Cm(width_cm)
    tcPr = cell._element.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567))) # 1 cm ~ 567 dxa
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)
    
    for p in cell.paragraphs:
        p.alignment = align
        p.style.font.name = 'Times New Roman'
        p.style.font.size = Pt(13)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)

def replace_table_with_5col(doc, target_table, data_rows):
    """Replaces target_table in doc with a newly created 5-column table holding data_rows."""
    tbl_elem = target_table._element
    parent = tbl_elem.getparent()
    tbl_idx = parent.index(tbl_elem)
    
    new_tbl = doc.add_table(rows=0, cols=5)
    new_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(new_tbl)
    
    hdr_row = new_tbl.add_row()
    headers = ['STT', 'Bài học (1)', 'Số tiết (2)', 'Tiết theo PPCT', 'Yêu cầu cần đạt (3)']
    for i, h_text in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h_text
        set_cell_width_and_align(cell, COL_WIDTHS[i], WD_ALIGN_PARAGRAPH.CENTER)
        p = cell.paragraphs[0]
        for r in p.runs:
            r.bold = True
            
    for r_data in data_rows:
        row = new_tbl.add_row()
        if r_data['is_topic']:
            a_cell = row.cells[0]
            a_cell.text = r_data['title']
            for c_idx in range(1, 5):
                a_cell.merge(row.cells[c_idx])
            p = a_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(13)
        else:
            values = [
                str(r_data['stt']),
                r_data['bai'],
                str(r_data['so_tiet']),
                str(r_data['ppct']),
                r_data['yccd']
            ]
            aligns = [
                WD_ALIGN_PARAGRAPH.CENTER,  # STT
                WD_ALIGN_PARAGRAPH.JUSTIFY, # Bài học (1)
                WD_ALIGN_PARAGRAPH.CENTER,  # Số tiết (2)
                WD_ALIGN_PARAGRAPH.CENTER,  # Tiết theo PPCT
                WD_ALIGN_PARAGRAPH.JUSTIFY  # Yêu cầu cần đạt (3)
            ]
            for c_idx in range(5):
                cell = row.cells[c_idx]
                cell.text = values[c_idx]
                set_cell_width_and_align(cell, COL_WIDTHS[c_idx], aligns[c_idx])
                
    new_elem = new_tbl._element
    parent.remove(new_elem)
    parent.insert(tbl_idx, new_elem)
    parent.remove(tbl_elem)

def standardize_existing_5col_table(table):
    """Standardizes width, font, alignment, and borders of an existing 5-column lesson table."""
    set_table_borders(table)
    aligns = [
        WD_ALIGN_PARAGRAPH.CENTER,  # STT
        WD_ALIGN_PARAGRAPH.JUSTIFY, # Bài học (1)
        WD_ALIGN_PARAGRAPH.CENTER,  # Số tiết (2)
        WD_ALIGN_PARAGRAPH.CENTER,  # Tiết theo PPCT
        WD_ALIGN_PARAGRAPH.JUSTIFY  # Yêu cầu cần đạt (3)
    ]
    for row_idx, row in enumerate(table.rows):
        cells_txt = [c.text.strip() for c in row.cells]
        if len(set(cells_txt)) == 1 and "Chủ đề" in cells_txt[0]:
            p = row.cells[0].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(13)
        elif len(row.cells) == 5:
            is_header = (row_idx == 0 or cells_txt[0] == 'STT')
            for c_idx in range(5):
                c_align = WD_ALIGN_PARAGRAPH.CENTER if is_header else aligns[c_idx]
                set_cell_width_and_align(row.cells[c_idx], COL_WIDTHS[c_idx], c_align)
                if is_header:
                    for r in row.cells[c_idx].paragraphs[0].runs:
                        r.bold = True

def parse_source_grade_table(doc_src, table_idx):
    """Parses source 4-column table into structured data rows with PPCT."""
    t = doc_src.tables[table_idx]
    data_rows = []
    curr_ppct = 0
    stt_counter = 1
    
    for r_idx in range(1, len(t.rows)):
        cells = [c.text.strip().replace('\n', ' ') for c in t.rows[r_idx].cells]
        raw_stt, bai, so_tiet_str, yccd = cells[0], cells[1], cells[2], cells[3]
        
        unique_c = list(dict.fromkeys(cells))
        if len(unique_c) == 1 or "Chủ đề" in bai:
            data_rows.append({
                'is_topic': True,
                'title': unique_c[0]
            })
            continue
            
        try:
            st = int(so_tiet_str)
        except:
            st = 1
            
        if st == 1:
            curr_ppct += 1
            ppct_str = str(curr_ppct)
        else:
            ppct_str = f"{curr_ppct + 1} - {curr_ppct + st}"
            curr_ppct += st
            
        data_rows.append({
            'is_topic': False,
            'stt': stt_counter,
            'bai': bai,
            'so_tiet': st,
            'ppct': ppct_str,
            'yccd': yccd
        })
        stt_counter += 1
        
    return data_rows

def main():
    print("Reading source document...")
    doc_src = docx.Document(SRC_4COT)
    
    data_tth = parse_source_grade_table(doc_src, 0)
    data_g1 = parse_source_grade_table(doc_src, 1)
    data_g2 = parse_source_grade_table(doc_src, 2)
    
    print(f"Parsed Tiền tiểu học: {len(data_tth)} rows")
    print(f"Parsed Lớp 1: {len(data_g1)} rows")
    print(f"Parsed Lớp 2: {len(data_g2)} rows")

    # 1. Update MAIN_ALL (Kế hoạch dạy học môn Tin học 2026-2027.docx)
    print(f"\nProcessing {MAIN_ALL}...")
    doc_main = docx.Document(MAIN_ALL)
    
    replace_table_with_5col(doc_main, doc_main.tables[5], data_g2)
    replace_table_with_5col(doc_main, doc_main.tables[4], data_g1)
    replace_table_with_5col(doc_main, doc_main.tables[3], data_tth)
    
    for t_idx in range(6, 12):
        standardize_existing_5col_table(doc_main.tables[t_idx])
        
    set_font_all(doc_main)
    doc_main.save(MAIN_ALL)
    print(f"Saved: {MAIN_ALL}")
    
    # 2. Update MAIN_TH (Kế hoạch dạy học môn Tin học (TH) - 2026 - 2027.docx)
    print(f"\nProcessing {MAIN_TH}...")
    doc_th = docx.Document(MAIN_TH)
    
    replace_table_with_5col(doc_th, doc_th.tables[5], data_g2)
    replace_table_with_5col(doc_th, doc_th.tables[4], data_tth)
    replace_table_with_5col(doc_th, doc_th.tables[3], data_g1)
    
    for t_idx in range(6, 9):
        standardize_existing_5col_table(doc_th.tables[t_idx])
        
    set_font_all(doc_th)
    doc_th.save(MAIN_TH)
    print(f"Saved: {MAIN_TH}")
    
    # 3. Update Individual Class Files
    print(f"\nProcessing individual class files...")
    
    # Lớp 1
    doc_l1 = docx.Document(FILE_L1)
    replace_table_with_5col(doc_l1, doc_l1.tables[3], data_g1)
    set_font_all(doc_l1)
    doc_l1.save(FILE_L1)
    print(f"Saved: {FILE_L1}")
    
    # Lớp 2
    doc_l2 = docx.Document(FILE_L2)
    replace_table_with_5col(doc_l2, doc_l2.tables[3], data_g2)
    set_font_all(doc_l2)
    doc_l2.save(FILE_L2)
    print(f"Saved: {FILE_L2}")
    
    # Tiền tiểu học
    doc_tth = docx.Document(FILE_L1)
    replace_table_with_5col(doc_tth, doc_tth.tables[3], data_tth)
    for p in doc_tth.paragraphs:
        if "LỚP 1" in p.text.upper() or "LỚP 1" in p.text:
            p.text = p.text.replace("Lớp 1", "Tiền tiểu học").replace("LỚP 1", "TIỀN TIỀU HỌC")
        if "Môn Tin lớp 1" in p.text:
            p.text = p.text.replace("Môn Tin lớp 1", "Môn Tư duy lập trình Tiền Tiểu học")
    set_font_all(doc_tth)
    doc_tth.save(FILE_TTH)
    print(f"Created & Saved: {FILE_TTH}")
    
    # 4. Synchronize to Phân phối chương trình\Tin học
    print(f"\nSynchronizing updated files to {DIR_DIST}...")
    shutil.copyfile(MAIN_ALL, os.path.join(DIR_DIST, 'Kế hoạch dạy học môn Tin học 2026-2027.docx'))
    shutil.copyfile(MAIN_TH, os.path.join(DIR_DIST, 'Kế hoạch dạy học môn Tin học (TH) - 2026 - 2027.docx'))
    print("Synchronization completed successfully!")

if __name__ == '__main__':
    try:
        main()
    except PermissionError as pe:
        print(f"PERMISSION ERROR: Please close open Word files. Details: {pe}")
        sys.exit(1)
