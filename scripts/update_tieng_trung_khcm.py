"""
update_tieng_trung_khcm.py — Điền dữ liệu môn Tiếng Trung vào Kế hoạch tổ chuyên môn THCS
======================================================================================
Theo đúng quy định AGENTS.md:
- TUYỆT ĐỐI GIỮ NGUYÊN FORM MẪU FILE GỐC.
- Điền dữ liệu vào đúng Bảng 7 CỘT và Bảng 5 CỘT sẵn có trong file gốc cho 4.1 (6A1), 4.2 (7A1), 4.3 (8A1).
- Xóa bỏ đoạn ghi chú placeholder.
- Cập nhật số tiết: Cả năm 72 tiết, HK1 36 tiết, HK2 36 tiết.
- Font Times New Roman 13pt, đường viền bảng XML tblBorders.
"""

import sys
import os
import re
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')

SRC_PATH = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch tổ chuyên môn\Kế hoạch dạy môn tiếng Trung - cô Châu.docx'
TARGET_PATH = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch tổ chuyên môn\30.07.26. Kế hoạch tổ chuyên môn (THCS).docx'


def set_table_borders(table):
    tblPr = table._tbl.tblPr
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tblPr.append(tblBorders)


def set_run_font(run, name='Times New Roman', size=13, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def format_cell(cell, text='', font_name='Times New Roman', size=13, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = text
    p = cell.paragraphs[0]
    p.alignment = align
    for r in p.runs:
        set_run_font(r, font_name, size, bold, italic)


def parse_source_data():
    sdoc = Document(SRC_PATH)
    t3 = sdoc.tables[3]
    t4 = sdoc.tables[4]

    hk1_items = []
    hk2_items = []

    current_hk = 1
    last_tuan = ''
    
    for ri in range(3, len(t3.rows)):
        row = t3.rows[ri]
        vals = [c.text.strip().replace('\n', ' ') for c in row.cells]
        tuan, chu_de, ten_bai, tiet_ppct, nd_dc, ghi_chu = vals[:6]

        if 'SEMESTER 2' in tuan.upper() or 'SEMESTER 2' in chu_de.upper():
            current_hk = 2
            continue

        if not tiet_ppct and not ten_bai and not chu_de:
            continue

        if tuan:
            last_tuan = tuan

        if not ten_bai:
            if chu_de:
                display_title = chu_de
            else:
                display_title = 'Luyện tập thực hành phát âm và hội thoại'
        else:
            display_title = ten_bai

        if 'giữa kỳ 1' in chu_de.lower() or 'giữa kỳ 1' in ten_bai.lower():
            goal = 'Kiểm tra kiến thức phát âm (nguyên âm, phụ âm, thanh điệu, vần) và khả năng nghe - viết phiên âm.'
        elif 'cuối kì 1' in chu_de.lower() or 'cuối kì 1' in ten_bai.lower():
            goal = 'Kiểm tra toàn bộ kiến thức Học kỳ 1: phát âm, từ vựng, cấu trúc câu cơ bản, viết chữ Hán cơ bản.'
        elif 'giữa kỳ 2' in chu_de.lower() or 'giữa kỳ 2' in ten_bai.lower():
            goal = 'Kiểm tra các chủ điểm từ Bài 9 - Bài 14: gia đình, tuổi tác, nghề nghiệp, sở thích, tính từ.'
        elif 'cuối kì 2' in chu_de.lower() or 'cuối kì 2' in ten_bai.lower():
            goal = 'Kiểm tra kiến thức toàn năm học; trọng tâm các mẫu câu và từ vựng HK2.'
        else:
            goal = 'Rèn luyện kỹ năng nghe, nói, đọc, viết Tiếng Trung sơ cấp; nắm vững phát âm, từ vựng và cấu trúc câu.'

        item = {
            'tuan': last_tuan,
            'chu_de': chu_de,
            'ten_bai': display_title,
            'tiet': tiet_ppct,
            'goal': goal
        }

        if current_hk == 1:
            hk1_items.append(item)
        else:
            hk2_items.append(item)

    ktdg_items = []
    for ri in range(1, len(t4.rows)):
        row = t4.rows[ri]
        vals = [c.text.strip().replace('\n', ' ') for c in row.cells]
        ten_bai_kt, thoi_gian, thoi_diem, yc_can_dat, hinh_thuc = vals[:5]
        
        if 'Giữa HK1' in ten_bai_kt:
            moc = 'Đánh giá định kỳ 1 (Giữa HK1)'
        elif 'Cuối HK1' in ten_bai_kt:
            moc = 'Đánh giá định kỳ 2 (Cuối HK1)'
        elif 'Giữa HK2' in ten_bai_kt:
            moc = 'Đánh giá định kỳ 3 (Giữa HK2)'
        elif 'Cuối HK2' in ten_bai_kt:
            moc = 'Đánh giá định kỳ 4 (Cuối HK2)'
        else:
            moc = f'Đánh giá định kỳ {ri}'

        ktdg_items.append({
            'tt': str(ri),
            'moc': moc,
            'noi_dung': f'{yc_can_dat} ({thoi_diem})',
            'hinh_thuc': hinh_thuc
        })

    return hk1_items, hk2_items, ktdg_items


def populate_ppct_table(tbl, hk1_items, hk2_items):
    set_table_borders(tbl)

    # Format Header Row (Row 0)
    headers = ['TT', 'Bài/chủ đề', 'Tổng số tiết', 'Tuần', 'Tiết theo PPCT', 'Nội dung', 'Mục tiêu bài học']
    for i, h in enumerate(headers):
        format_cell(tbl.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Add HK1 Header Row
    r1 = tbl.add_row()
    mcell1 = r1.cells[0].merge(r1.cells[6])
    format_cell(mcell1, 'HỌC KỲ I: 36 TIẾT (18 TUẦN)', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Add HK1 rows
    for idx, item in enumerate(hk1_items, start=1):
        r = tbl.add_row()
        format_cell(r.cells[0], str(idx), align=WD_ALIGN_PARAGRAPH.CENTER)
        format_cell(r.cells[1], item['ten_bai'], align=WD_ALIGN_PARAGRAPH.LEFT)
        format_cell(r.cells[2], '1', align=WD_ALIGN_PARAGRAPH.CENTER)
        format_cell(r.cells[3], item['tuan'], align=WD_ALIGN_PARAGRAPH.CENTER)
        format_cell(r.cells[4], item['tiet'], align=WD_ALIGN_PARAGRAPH.CENTER)
        format_cell(r.cells[5], item['ten_bai'], align=WD_ALIGN_PARAGRAPH.LEFT)
        format_cell(r.cells[6], item['goal'], align=WD_ALIGN_PARAGRAPH.LEFT)

    # Add HK2 Header Row
    r2 = tbl.add_row()
    mcell2 = r2.cells[0].merge(r2.cells[6])
    format_cell(mcell2, 'HỌC KỲ II: 36 TIẾT (18 TUẦN)', bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Add HK2 rows
    for idx, item in enumerate(hk2_items, start=37):
        r = tbl.add_row()
        format_cell(r.cells[0], str(idx), align=WD_ALIGN_PARAGRAPH.CENTER)
        format_cell(r.cells[1], item['ten_bai'], align=WD_ALIGN_PARAGRAPH.LEFT)
        format_cell(r.cells[2], '1', align=WD_ALIGN_PARAGRAPH.CENTER)
        format_cell(r.cells[3], item['tuan'], align=WD_ALIGN_PARAGRAPH.CENTER)
        format_cell(r.cells[4], item['tiet'], align=WD_ALIGN_PARAGRAPH.CENTER)
        format_cell(r.cells[5], item['ten_bai'], align=WD_ALIGN_PARAGRAPH.LEFT)
        format_cell(r.cells[6], item['goal'], align=WD_ALIGN_PARAGRAPH.LEFT)


def fill_ktdg_table(tbl, ktdg_items, lop_name):
    set_table_borders(tbl)

    headers = ['TT', 'Lớp', 'Bài kiểm tra', 'Nội dung', 'Hình thức']
    for i, h in enumerate(headers):
        format_cell(tbl.rows[0].cells[i], h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Fill existing 4 data rows (rows 1..4)
    for idx, item in enumerate(ktdg_items, start=1):
        if idx < len(tbl.rows):
            r = tbl.rows[idx]
        else:
            r = tbl.add_row()

        format_cell(r.cells[0], str(idx), align=WD_ALIGN_PARAGRAPH.CENTER)
        format_cell(r.cells[1], lop_name, align=WD_ALIGN_PARAGRAPH.CENTER)
        format_cell(r.cells[2], item['moc'], align=WD_ALIGN_PARAGRAPH.LEFT)
        format_cell(r.cells[3], item['noi_dung'], align=WD_ALIGN_PARAGRAPH.LEFT)
        format_cell(r.cells[4], item['hinh_thuc'], align=WD_ALIGN_PARAGRAPH.LEFT)


def main():
    print('Parsing source data...')
    hk1_items, hk2_items, ktdg_items = parse_source_data()
    print(f'Parsed: {len(hk1_items)} HK1 lessons, {len(hk2_items)} HK2 lessons, {len(ktdg_items)} KTĐG milestones.')

    print('Opening target document...')
    doc = Document(TARGET_PATH)

    # Identify tables following section 4 headings
    body = doc.element.body
    
    sections = [
        ('4.1.', 'Lớp 6A1'),
        ('4.2.', 'Lớp 7A1'),
        ('4.3.', 'Lớp 8A1'),
    ]

    for prefix, lop_name in sections:
        print(f'\nProcessing {prefix} ({lop_name})...')
        
        header_p_idx = None
        for p_idx, p in enumerate(doc.paragraphs):
            txt = p.text.strip()
            if txt.startswith(f'{prefix} Kế hoạch dạy học môn Tiếng Trung'):
                header_p_idx = p_idx
                break

        if header_p_idx is None:
            print(f'   ⚠ Could not find section {prefix}')
            continue

        # Update summary lines following header_p_idx
        for j in range(header_p_idx, header_p_idx + 10):
            if j < len(doc.paragraphs):
                ptxt = doc.paragraphs[j].text.strip()
                if ptxt.startswith('Cả năm:'):
                    doc.paragraphs[j].text = 'Cả năm: 72 Tiết'
                    for r in doc.paragraphs[j].runs: set_run_font(r)
                elif ptxt.startswith('Học kì 1:'):
                    doc.paragraphs[j].text = 'Học kì 1: 36 Tiết'
                    for r in doc.paragraphs[j].runs: set_run_font(r)
                elif ptxt.startswith('Học kì 2:'):
                    doc.paragraphs[j].text = 'Học kì 2: 36 Tiết'
                    for r in doc.paragraphs[j].runs: set_run_font(r)

        # Remove placeholder note paragraph if present
        for j in range(header_p_idx, header_p_idx + 12):
            if j < len(doc.paragraphs):
                ptxt = doc.paragraphs[j].text.strip()
                if 'Mục môn Tiếng Trung - Đã ghi nhận form mẫu' in ptxt:
                    p_to_remove = doc.paragraphs[j]
                    parent = p_to_remove._element.getparent()
                    parent.remove(p_to_remove._element)
                    print(f'   Removed placeholder note paragraph')
                    break

        # Find 7-column table and 5-column table after header_p_idx in body order
        header_p = doc.paragraphs[header_p_idx]
        body_children = list(body)
        h_elem_idx = body_children.index(header_p._element)

        ppct_tbl_elem = None
        ktdg_tbl_elem = None

        for elem in body_children[h_elem_idx:]:
            tag = elem.tag.split('}')[-1]
            if tag == 'p':
                txt = ''.join(t.text or '' for t in elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t')).strip()
                # Stop if next section reached
                if txt.startswith('4.2.') and prefix == '4.1.': break
                if txt.startswith('4.3.') and prefix == '4.2.': break
                if txt.startswith('5. Môn:') and prefix == '4.3.': break
            elif tag == 'tbl':
                rows = elem.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tr')
                if rows:
                    cols = rows[0].findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tc')
                    if len(cols) == 7 and ppct_tbl_elem is None:
                        ppct_tbl_elem = elem
                    elif len(cols) == 5 and ktdg_tbl_elem is None:
                        ktdg_tbl_elem = elem

        if ppct_tbl_elem is not None:
            # Find matching Table object in doc.tables
            for tbl in doc.tables:
                if tbl._element == ppct_tbl_elem:
                    populate_ppct_table(tbl, hk1_items, hk2_items)
                    print(f'   Populated 7-col PPCT table for {prefix}')
                    break

        if ktdg_tbl_elem is not None:
            for tbl in doc.tables:
                if tbl._element == ktdg_tbl_elem:
                    fill_ktdg_table(tbl, ktdg_items, lop_name)
                    print(f'   Filled 5-col KTĐG table for {prefix}')
                    break

    print('\nApplying Times New Roman 13pt to all paragraphs and tables...')
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.name = 'Times New Roman'
            if r.font.size is None:
                r.font.size = Pt(13)

    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = 'Times New Roman'
                        if r.font.size is None:
                            r.font.size = Pt(13)

    print('Saving updated document...')
    doc.save(TARGET_PATH)
    print(f'✅ Successfully updated: {TARGET_PATH}')


if __name__ == '__main__':
    main()
