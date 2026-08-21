"""
verify_khbd_all.py
==================
Kiểm tra tất cả file KHBD đã tạo theo đúng chuẩn luật:
- Bảng 2 cột (KHÔNG PHẢI 3 cột)
- Tên trường đúng
- Kiến thức không chứa cụm cấm
- Viền bảng đúng chuẩn
- Header drawing còn nguyên
"""

import os
import sys
from docx import Document
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')

DIRS = [
    r"D:\UNIGO\KHBD_Tin_học",
    r"D:\UNIGO\KHBD_Robotics",
]

BANNED_PHRASES = [
    'Sự hiểu biết về',
    'Khả năng nhận diện',
    'Khả năng phân tích',
    'Khả năng vận dụng',
    'Sự nhận biết',
]


def check_file(filepath):
    """Check a single KHBD file. Returns list of issues."""
    issues = []
    try:
        doc = Document(filepath)
    except Exception as e:
        return [f"Không thể mở file: {e}"]

    # 1. Check header drawing
    has_drawing = False
    for sec in doc.sections:
        header = sec.header
        for p in header.paragraphs:
            for run in p.runs:
                if run._element.findall(f'.//{qn("w:drawing")}'):
                    has_drawing = True
                    break
            if not has_drawing:
                for child in p._p:
                    if child.tag.endswith('drawing'):
                        has_drawing = True
                        break
    if not has_drawing:
        issues.append("⚠ Header drawing (logo) không tìm thấy")

    # 2. Check tables
    tables = doc.tables
    for idx, table in enumerate(tables):
        num_cols = len(table.columns)
        if num_cols == 3 and idx > 0 and idx < len(tables) - 1:
            # Check if this is a 3-column activity table (BAD)
            first_cell_text = table.rows[0].cells[0].text.strip()
            if 'Bước' in first_cell_text or 'Hoạt động của GV' in first_cell_text:
                issues.append(f"❌ Bảng {idx}: Vẫn dùng 3 cột (Bước/GV/HS) — phải là 2 cột")

    # 3. Check school name
    full_text = '\n'.join(p.text for p in doc.paragraphs)
    for t in tables:
        for row in t.rows:
            for cell in row.cells:
                full_text += '\n' + cell.text

    if 'TH&THCS UNIGO' in full_text or 'TH & THCS UNIGO' in full_text:
        issues.append("❌ Tên trường sai: phải là 'Trường Tiểu học và THCS UNIGO'")
    if 'Trường: Tiểu học' in full_text or 'Trường: TH' in full_text:
        issues.append("❌ Tên trường sai format: không dùng 'Trường:' mà dùng 'Trường Tiểu học và THCS UNIGO'")

    # 4. Check banned knowledge phrases
    for phrase in BANNED_PHRASES:
        if phrase in full_text:
            issues.append(f"❌ Chứa cụm cấm: '{phrase}'")

    # 5. Check duplicate "Tên tiết:"
    if 'Tên tiết:' in full_text:
        issues.append("❌ Có 'Tên tiết:' trùng lặp — chỉ nên có 'TÊN BÀI DẠY:'")

    # 6. Check table borders
    for idx, table in enumerate(tables):
        tbl_borders = table._tbl.tblPr.find(qn('w:tblBorders')) if table._tbl.tblPr else None
        if tbl_borders is not None:
            top_border = tbl_borders.find(qn('w:top'))
            if top_border is not None:
                border_val = top_border.get(qn('w:val'))
                # First and last tables should be nil (no border)
                if idx == 0 or idx == len(tables) - 1:
                    if border_val not in ('nil', 'none', None):
                        issues.append(f"⚠ Bảng {idx} (info/ký tên) phải NO BORDER, hiện: {border_val}")

    return issues


def main():
    total_files = 0
    total_issues = 0
    files_with_issues = 0

    for base_dir in DIRS:
        if not os.path.exists(base_dir):
            print(f"[!] Thư mục không tồn tại: {base_dir}")
            continue

        print(f"\n{'='*60}")
        print(f" KIỂM TRA: {base_dir}")
        print(f"{'='*60}")

        for root, dirs, files in os.walk(base_dir):
            for f in files:
                if not f.endswith('.docx') or f.startswith('~'):
                    continue
                filepath = os.path.join(root, f)
                total_files += 1
                issues = check_file(filepath)
                if issues:
                    files_with_issues += 1
                    total_issues += len(issues)
                    rel = os.path.relpath(filepath, base_dir)
                    print(f"\n  📄 {rel}")
                    for issue in issues:
                        print(f"     {issue}")

    print(f"\n{'='*60}")
    print(f" KẾT QUẢ KIỂM TRA:")
    print(f"   Tổng files:      {total_files}")
    print(f"   Files có lỗi:    {files_with_issues}")
    print(f"   Tổng vấn đề:     {total_issues}")
    if files_with_issues == 0:
        print(f"   ✅ TẤT CẢ ĐẠT CHUẨN!")
    else:
        print(f"   ⚠ CẦN SỬA {files_with_issues} FILES!")
    print(f"{'='*60}")


if __name__ == '__main__':
    main()
