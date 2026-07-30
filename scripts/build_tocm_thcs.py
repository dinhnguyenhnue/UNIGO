import docx
import os
import sys
import copy
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding='utf-8')

DIR_TOCM = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch tổ chuyên môn'
DIR_MAU = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm'
MASTER_FILE = os.path.join(DIR_TOCM, 'Kế hoạch tổ chuyên môn (THCS).docx')

def set_table_borders(table):
    """Sets explicit full black grid borders XML without relying on style name."""
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        for child in list(tblPr[0]):
            if child.tag.endswith('tblBorders'):
                tblPr[0].remove(child)
        borders_xml = parse_xml(
            r'<w:tblBorders %s>'
            r'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        tblPr[0].append(borders_xml)

def append_p(doc, text, bold=False, italic=False, level=0, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=13):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(font_size)
    r.bold = bold
    r.italic = italic
    return p

def append_table_copy(doc, src_table):
    set_table_borders(src_table)
    tbl_elem = copy.deepcopy(src_table._element)
    doc._element.body.append(tbl_elem)

def main():
    print("Reading master document structure...")
    doc_master = docx.Document(MASTER_FILE)
    
    doc_out = docx.Document()
    
    copy_mode = True
    section_vi_reached = False
    section_vii_paragraphs = []
    section_vii_reached = False
    
    for p in doc_master.paragraphs:
        txt = p.text.strip()
        if txt.startswith("VI. KẾ HOẠCH GIẢNG DẠY"):
            section_vi_reached = True
            copy_mode = False
            continue
        elif txt.startswith("VII. KẾ HOẠCH CỤ THỂ"):
            section_vii_reached = True
            copy_mode = False
            
        if section_vii_reached:
            section_vii_paragraphs.append(p)
        elif copy_mode:
            p_elem = copy.deepcopy(p._element)
            doc_out._element.body.append(p_elem)
            
    for t_i in range(min(9, len(doc_master.tables))):
        append_table_copy(doc_out, doc_master.tables[t_i])
        append_p(doc_out, "")
        
    print("Building Section VI. KẾ HOẠCH GIẢNG DẠY CÁC MÔN HỌC (14 subjects in order)...")
    
    append_p(doc_out, "VI. KẾ HOẠCH GIẢNG DẠY CÁC MÔN HỌC", bold=True, font_size=14)
    append_p(doc_out, "Khung kế hoạch dạy học chi tiết cho các môn học tổ THCS (Lớp 6, Lớp 7, Lớp 8) năm học 2026 - 2027:", italic=True)
    append_p(doc_out, "")
    
    # Define source paths
    f_van_gddp = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch dạy học môn Văn_GDDP (THCS)-PPCT.docx')
    f_toan = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch môn Toán THCS (Kế hoạch cá nhân).docx')
    f_anh = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch dạy học môn Tiếng Anh (THCS)-PPCT.docx')
    f_khtn = os.path.join(DIR_TOCM, 'PL1. Tổng hợp khung kế hoạch dạy học môn KHTN6,7,8.26-27.docx')
    f_tin = os.path.join(DIR_MAU, 'Kế hoạch dạy học môn Tin học (THCS) - 2026 - 2027.docx')
    f_robotics = os.path.join(DIR_MAU, 'Kế hoạch dạy học môn Robotics (THCS) - 2026 - 2027.docx')
    f_hdtn = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch môn Toán + HĐTN TH, THCS. Bùi Thị Phương Anh(Kế hoạch cá nhân).docx')
    f_mythuat = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch dạy học môn Mĩ Thuật (THCS)-PPCT.docx')
    f_gdtc = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch dạy học môn GDTC 26-27.docx')

    # 1. Văn (Ngữ văn)
    print("1/14: Processing Văn...")
    append_p(doc_out, "1. Môn: Ngữ văn", bold=True, font_size=13)
    if os.path.exists(f_van_gddp):
        doc_v = docx.Document(f_van_gddp)
        for i, t in enumerate(doc_v.tables[:3]):
            append_p(doc_out, f"1.{i+1}. Kế hoạch dạy học môn Ngữ văn - Lớp {i+6}", bold=True)
            append_table_copy(doc_out, t)
            append_p(doc_out, "")

    # 2. Toán
    print("2/14: Processing Toán...")
    append_p(doc_out, "2. Môn: Toán", bold=True, font_size=13)
    if os.path.exists(f_toan):
        doc_t = docx.Document(f_toan)
        for i, t in enumerate(doc_t.tables):
            hdr_txt = " ".join([c.text.strip() for c in t.rows[0].cells])
            if "Bài" in hdr_txt or "chủ đề" in hdr_txt.lower() or "STT" in hdr_txt:
                append_table_copy(doc_out, t)
                append_p(doc_out, "")

    # 3. Tiếng Anh
    print("3/14: Processing Tiếng Anh...")
    append_p(doc_out, "3. Môn: Tiếng Anh", bold=True, font_size=13)
    if os.path.exists(f_anh):
        doc_a = docx.Document(f_anh)
        for i, t in enumerate(doc_a.tables):
            if len(t.rows) > 5:
                append_table_copy(doc_out, t)
                append_p(doc_out, "")

    # 4. Tiếng Trung (Placeholder)
    print("4/14: Processing Tiếng Trung (Placeholder)...")
    append_p(doc_out, "4. Môn: Tiếng Trung", bold=True, font_size=13)
    append_p(doc_out, "(Mục môn Tiếng Trung - Đã ghi nhận tên mục và sẽ bổ sung chi tiết Khung kế hoạch dạy học sau)", italic=True)
    append_p(doc_out, "")

    # 5. KHTN (Khoa học tự nhiên)
    print("5/14: Processing Khoa học tự nhiên...")
    append_p(doc_out, "5. Môn: Khoa học tự nhiên (KHTN)", bold=True, font_size=13)
    if os.path.exists(f_khtn):
        doc_k = docx.Document(f_khtn)
        for i, t in enumerate(doc_k.tables):
            if len(t.rows) > 5:
                append_table_copy(doc_out, t)
                append_p(doc_out, "")

    # 6. Lịch sử và Địa lý (Placeholder)
    print("6/14: Processing Lịch sử và Địa lý (Placeholder)...")
    append_p(doc_out, "6. Môn: Lịch sử và Địa lý", bold=True, font_size=13)
    append_p(doc_out, "(Mục môn Lịch sử và Địa lý - Đã ghi nhận tên mục và sẽ bổ sung chi tiết Khung kế hoạch dạy học sau)", italic=True)
    append_p(doc_out, "")

    # 7. Tin học
    print("7/14: Processing Tin học...")
    append_p(doc_out, "7. Môn: Tin học", bold=True, font_size=13)
    if os.path.exists(f_tin):
        doc_tin = docx.Document(f_tin)
        for i, t in enumerate(doc_tin.tables):
            if len(t.rows) > 5:
                append_table_copy(doc_out, t)
                append_p(doc_out, "")

    # 8. Robotics
    print("8/14: Processing Robotics...")
    append_p(doc_out, "8. Môn: Robotics", bold=True, font_size=13)
    if os.path.exists(f_robotics):
        doc_rob = docx.Document(f_robotics)
        for i, t in enumerate(doc_rob.tables):
            if len(t.rows) > 5:
                append_table_copy(doc_out, t)
                append_p(doc_out, "")

    # 9. GDCD (Placeholder)
    print("9/14: Processing Giáo dục công dân (Placeholder)...")
    append_p(doc_out, "9. Môn: Giáo dục công dân (GDCD)", bold=True, font_size=13)
    append_p(doc_out, "(Mục môn Giáo dục công dân - Đã ghi nhận tên mục và sẽ bổ sung chi tiết Khung kế hoạch dạy học sau)", italic=True)
    append_p(doc_out, "")

    # 10. GD địa phương
    print("10/14: Processing Giáo dục địa phương...")
    append_p(doc_out, "10. Môn: Giáo dục địa phương", bold=True, font_size=13)
    if os.path.exists(f_van_gddp):
        doc_v = docx.Document(f_van_gddp)
        if len(doc_v.tables) >= 6:
            for i, t in enumerate(doc_v.tables[3:6]):
                append_p(doc_out, f"10.{i+1}. Kế hoạch dạy học môn Giáo dục địa phương - Lớp {i+6}", bold=True)
                append_table_copy(doc_out, t)
                append_p(doc_out, "")

    # 11. Hoạt động trải nghiệm
    print("11/14: Processing Hoạt động trải nghiệm...")
    append_p(doc_out, "11. Môn: Hoạt động trải nghiệm, hướng nghiệp", bold=True, font_size=13)
    if os.path.exists(f_hdtn):
        doc_h = docx.Document(f_hdtn)
        for i, t in enumerate(doc_h.tables):
            if len(t.rows) > 10:
                append_table_copy(doc_out, t)
                append_p(doc_out, "")

    # 12. Âm nhạc (Placeholder)
    print("12/14: Processing Âm nhạc (Placeholder)...")
    append_p(doc_out, "12. Môn: Âm nhạc", bold=True, font_size=13)
    append_p(doc_out, "(Mục môn Âm nhạc - Đã ghi nhận tên mục và sẽ bổ sung chi tiết Khung kế hoạch dạy học sau)", italic=True)
    append_p(doc_out, "")

    # 13. Mỹ thuật
    print("13/14: Processing Mỹ thuật...")
    append_p(doc_out, "13. Môn: Mỹ thuật", bold=True, font_size=13)
    if os.path.exists(f_mythuat):
        doc_m = docx.Document(f_mythuat)
        for i, t in enumerate(doc_m.tables):
            if len(t.rows) > 5:
                append_table_copy(doc_out, t)
                append_p(doc_out, "")

    # 14. Thể chất (GDTC)
    print("14/14: Processing Thể chất (GDTC)...")
    append_p(doc_out, "14. Môn: Giáo dục thể chất (GDTC)", bold=True, font_size=13)
    if os.path.exists(f_gdtc):
        doc_g = docx.Document(f_gdtc)
        for i, t in enumerate(doc_g.tables):
            hdr_txt = " ".join([c.text.strip() for c in t.rows[0].cells])
            if ("Chạy cự li" in hdr_txt or "Bài học" in hdr_txt or "Tuần" in hdr_txt) and len(t.rows) > 10:
                append_table_copy(doc_out, t)
                append_p(doc_out, "")

    # Append Section VII and signatures
    print("Appending Section VII. KẾ HOẠCH CỤ THỂ and signature blocks...")
    if section_vii_paragraphs:
        for p in section_vii_paragraphs:
            p_elem = copy.deepcopy(p._element)
            doc_out._element.body.append(p_elem)
            
    # Set Times New Roman 13pt
    for p in doc_out.paragraphs:
        p.style.font.name = 'Times New Roman'
        for r in p.runs:
            r.font.name = 'Times New Roman'
            
    for t in doc_out.tables:
        set_table_borders(t)
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.style.font.name = 'Times New Roman'
                    for r in p.runs:
                        r.font.name = 'Times New Roman'

    try:
        doc_out.save(MASTER_FILE)
        print(f"\nSuccessfully built and saved master document: {MASTER_FILE}")
    except PermissionError:
        fallback = os.path.join(DIR_TOCM, 'Kế hoạch tổ chuyên môn (THCS) - Mới.docx')
        doc_out.save(fallback)
        print(f"\nPermission error saving {MASTER_FILE}, saved fallback: {fallback}")

if __name__ == '__main__':
    main()
