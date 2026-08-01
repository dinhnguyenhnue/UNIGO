import os
import sys
import docx
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

sys.stdout.reconfigure(encoding='utf-8')

# Source & target paths
SRC_4COT = r'D:\UNIGO\Phân phối chương trình\Ke-hoach\Ke_hoach_Tin_hoc_Tien_tieu_hoc_Lop_1_Lop_2_4_cot.docx'

MAIN_ALL = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch dạy học môn Tin học 2026-2027.docx'
MAIN_TH = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch dạy học môn Tin học (TH) - 2026 - 2027.docx'

INDIV_DIR = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch dạy học Tin học từng lớp'
DIST_DIR = r'D:\UNIGO\Phân phối chương trình\Tin học'

def set_font_all(doc):
    """Sets Times New Roman 13pt for all paragraphs, runs, tables, cells in the document."""
    for p in doc.paragraphs:
        p.style.font.name = 'Times New Roman'
        p.style.font.size = Pt(13)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)
            
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.style.font.name = 'Times New Roman'
                    p.style.font.size = Pt(13)
                    for r in p.runs:
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(13)

def set_cell_width_and_align(cell, width_cm, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Sets width and alignment for a single cell."""
    cell.width = Cm(width_cm)
    tcPr = cell._element.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567))) # 1 cm ~ 567 dxa
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)
    
    for p in cell.paragraphs:
        p.alignment = align
        p.style.font.name = 'Times New Roman'
        p.style.font.size = Pt(13)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)

def format_5col_table(table, data_rows):
    """
    Formats an existing table or populates a table with 5 columns:
    STT | Bài học (1) | Số tiết (2) | Tiết theo PPCT | Yêu cầu cần đạt (3)
    """
    col_widths = [0.9, 3.8, 1.1, 1.4, 9.3] # total ~ 16.5 cm
    
    # Ensure table has 5 columns by clearing and building rows
    # Clear existing rows except keep table structure
    # Re-build rows
    tblPr = table._element.xpath('w:tblPr')
    
    # We will reset rows
    # Delete all rows
    for row in list(table.rows):
        table._element.remove(row._element)
        
    # Add Header Row
    hdr_row = table.add_row()
    headers = ['STT', 'Bài học (1)', 'Số tiết (2)', 'Tiết theo PPCT', 'Yêu cầu cần đạt (3)']
    for i, h_text in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h_text
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)
        set_cell_width_and_align(cell, col_widths[i], WD_ALIGN_PARAGRAPH.CENTER)
        
    # Populate data rows
    for r_data in data_rows:
        row = table.add_row()
        if r_data['is_topic']:
            # Topic row merged across 5 columns
            a_cell = row.cells[0]
            a_cell.text = r_data['title']
            # Merge cells 0..4
            for c_idx in range(1, 5):
                a_cell.merge(row.cells[c_idx])
            p = a_cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(13)
        else:
            # 5 data cells
            values = [
                str(r_data['stt']),
                r_data['bai'],
                str(r_data['so_tiet']),
                str(r_data['ppct']),
                r_data['yccd']
            ]
            aligns = [
                WD_ALIGN_PARAGRAPH.CENTER,  # STT
                WD_ALIGN_PARAGRAPH.JUSTIFY, # Bài học
                WD_ALIGN_PARAGRAPH.CENTER,  # Số tiết
                WD_ALIGN_PARAGRAPH.CENTER,  # Tiết theo PPCT
                WD_ALIGN_PARAGRAPH.JUSTIFY  # YCCĐ
            ]
            for c_idx in range(5):
                cell = row.cells[c_idx]
                cell.text = values[c_idx]
                set_cell_width_and_align(cell, col_widths[c_idx], aligns[c_idx])

def parse_source_grade_table(doc_src, table_idx):
    """Parses source 4-column table into structured data rows with PPCT."""
    t = doc_src.tables[table_idx]
    data_rows = []
    curr_ppct = 0
    stt_counter = 1
    
    for r_idx in range(1, len(t.rows)):
        cells = [c.text.strip().replace('\n', ' ') for c in t.rows[r_idx].cells]
        raw_stt, bai, so_tiet_str, yccd = cells[0], cells[1], cells[2], cells[3]
        
        unique_c = list(dict.fromkeys(cells))
        if len(unique_c) == 1 or "Chủ đề" in bai:
            data_rows.append({
                'is_topic': True,
                'title': unique_c[0]
            })
            continue
            
        try:
            st = int(so_tiet_str)
        except:
            st = 1
            
        if st == 1:
            curr_ppct += 1
            ppct_str = str(curr_ppct)
        else:
            ppct_str = f"{curr_ppct + 1} - {curr_ppct + st}"
            curr_ppct += st
            
        data_rows.append({
            'is_topic': False,
            'stt': stt_counter,
            'bai': bai,
            'so_tiet': st,
            'ppct': ppct_str,
            'yccd': yccd
        })
        stt_counter += 1
        
    return data_rows

def main():
    print("Reading source document...")
    doc_src = docx.Document(SRC_4COT)
    
    data_tth = parse_source_grade_table(doc_src, 0)
    data_g1 = parse_source_grade_table(doc_src, 1)
    data_g2 = parse_source_grade_table(doc_src, 2)
    
    print(f"Parsed Tiền tiểu học: {len(data_tth)} rows")
    print(f"Parsed Lớp 1: {len(data_g1)} rows")
    print(f"Parsed Lớp 2: {len(data_g2)} rows")
    
    # 1. Update MAIN_ALL (Kế hoạch dạy học môn Tin học 2026-2027.docx)
    print(f"\nProcessing {MAIN_ALL}...")
    doc_main = docx.Document(MAIN_ALL)
    
    # Table 3 = Tiền tiểu học, Table 4 = Lớp 1, Table 5 = Lớp 2
    format_5col_table(doc_main.tables[3], data_tth)
    format_5col_table(doc_main.tables[4], data_g1)
    format_5col_table(doc_main.tables[5], data_g2)
    
    # Also standardize width & alignment of all remaining lesson tables (Lớp 3..8: Table 6..11)
    for t_idx in range(6, 12):
        t = doc_main.tables[t_idx]
        col_widths = [0.9, 3.8, 1.1, 1.4, 9.3]
        aligns = [
            WD_ALIGN_PARAGRAPH.CENTER,  # STT
            WD_ALIGN_PARAGRAPH.JUSTIFY, # Bài học
            WD_ALIGN_PARAGRAPH.CENTER,  # Số tiết
            WD_ALIGN_PARAGRAPH.CENTER,  # Tiết theo PPCT
            WD_ALIGN_PARAGRAPH.JUSTIFY  # YCCĐ
        ]
        for row in t.rows:
            # check if topic row
            cells_txt = [c.text.strip() for c in row.cells]
            if len(set(cells_txt)) == 1 and "Chủ đề" in cells_txt[0]:
                p = row.cells[0].paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.bold = True
            elif len(row.cells) == 5:
                is_header = (row.cells[0].text.strip() == 'STT')
                for c_idx in range(5):
                    c_align = WD_ALIGN_PARAGRAPH.CENTER if is_header else aligns[c_idx]
                    set_cell_width_and_align(row.cells[c_idx], col_widths[c_idx], c_align)
                    
    set_font_all(doc_main)
    doc_main.save(MAIN_ALL)
    print(f"Successfully saved {MAIN_ALL}")
    
    # 2. Update MAIN_TH (Kế hoạch dạy học môn Tin học (TH) - 2026 - 2027.docx)
    print(f"\nProcessing {MAIN_TH}...")
    doc_th = docx.Document(MAIN_TH)
    # Check table structure in MAIN_TH
    # Table 3 = Lớp 1 (or Tiền tiểu học), Table 4 = Tiền tiểu học (or Lớp 2)
    # Let's inspect MAIN_TH tables first or apply to matching tables
    for t_idx, t in enumerate(doc_th.tables):
        if len(t.rows) > 0 and len(t.columns) in [4, 5]:
            hdr = [c.text.strip() for c in t.rows[0].cells]
            if 'Bài học (1)' in hdr or 'Tên bài' in hdr:
                # Format 5 cols
                col_widths = [0.9, 3.8, 1.1, 1.4, 9.3]
                aligns = [WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.JUSTIFY, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.CENTER, WD_ALIGN_PARAGRAPH.JUSTIFY]
                for row in t.rows:
                    cells_txt = [c.text.strip() for c in row.cells]
                    if len(set(cells_txt)) == 1 and "Chủ đề" in cells_txt[0]:
                        p = row.cells[0].paragraphs[0]
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                        for r in p.runs:
                            r.bold = True
                    elif len(row.cells) == 5:
                        is_header = (row.cells[0].text.strip() == 'STT')
                        for c_idx in range(5):
                            c_align = WD_ALIGN_PARAGRAPH.CENTER if is_header else aligns[c_idx]
                            set_cell_width_and_align(row.cells[c_idx], col_widths[c_idx], c_align)

    set_font_all(doc_th)
    doc_th.save(MAIN_TH)
    print(f"Successfully saved {MAIN_TH}")

if __name__ == '__main__':
    main()
