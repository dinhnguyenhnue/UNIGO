import docx
import os
import sys
import copy
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding='utf-8')

DIR_TOCM = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch tổ chuyên môn'
DIR_MAU = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm'
FILE_30_07_26 = os.path.join(DIR_TOCM, '30.07.26. Kế hoạch tổ chuyên môn (THCS).docx')
TARGET_FILE = os.path.join(DIR_TOCM, 'Kế hoạch tổ chuyên môn (THCS).docx')

def clean_xml_element(elem):
    for bm in elem.xpath('.//w:bookmarkStart | .//w:bookmarkEnd'):
        bm.getparent().remove(bm)
    for node in elem.xpath('.//*[@r:id]'):
        for attr in list(node.attrib):
            if 'id' in attr.lower():
                del node.attrib[attr]
    return elem

def set_table_borders(table):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        for child in list(tblPr[0]):
            if child.tag.endswith('tblBorders'):
                tblPr[0].remove(child)
        borders_xml = parse_xml(
            r'<w:tblBorders %s>'
            r'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        tblPr[0].append(borders_xml)

def append_element_before_sectpr(doc, elem):
    body = doc._element.body
    clean_elem = clean_xml_element(elem)
    sectPr = body.xpath('w:sectPr')
    if sectPr:
        sectPr[-1].addprevious(clean_elem)
    else:
        body.append(clean_elem)

def append_p(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, font_size=13):
    p = doc.add_paragraph()
    p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(font_size)
    r.bold = bold
    r.italic = italic
    
    p_elem = p._element
    p_elem.getparent().remove(p_elem)
    append_element_before_sectpr(doc, p_elem)
    return p

def create_form_table_7col(doc, items_hk1, items_hk2):
    """Creates a 7-column table in the EXACT template form of 30.07.26."""
    table = doc.add_table(rows=0, cols=7)
    set_table_borders(table)
    
    # Header row
    hdr_row = table.add_row()
    headers = ['TT', 'Bài/chủ đề', 'Tổng số tiết', 'Tuần', 'Tiết theo PPCT', 'Nội dung', 'Mục tiêu bài học']
    for idx, text in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.text = text
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)
            
    # Row HK1 header
    hk1_row = table.add_row()
    cell_hk1 = hk1_row.cells[0]
    for c_i in range(1, 7):
        cell_hk1.merge(hk1_row.cells[c_i])
    cell_hk1.text = "Học kì 1"
    p = cell_hk1.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.runs[0].font.bold = True
    p.runs[0].font.name = 'Times New Roman'
    
    # HK1 items
    for item in items_hk1:
        row = table.add_row()
        for idx in range(min(7, len(item))):
            cell = row.cells[idx]
            cell.text = str(item[idx])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in [0, 2, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(13)
                
    # Row HK2 header
    hk2_row = table.add_row()
    cell_hk2 = hk2_row.cells[0]
    for c_i in range(1, 7):
        cell_hk2.merge(hk2_row.cells[c_i])
    cell_hk2.text = "Học kì 2"
    p = cell_hk2.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.runs[0].font.bold = True
    p.runs[0].font.name = 'Times New Roman'
    
    # HK2 items
    for item in items_hk2:
        row = table.add_row()
        for idx in range(min(7, len(item))):
            cell = row.cells[idx]
            cell.text = str(item[idx])
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in [0, 2, 3, 4] else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(13)

    t_elem = table._element
    t_elem.getparent().remove(t_elem)
    append_element_before_sectpr(doc, t_elem)

def create_form_table_5col(doc, grade_name):
    """Creates a 5-column assessment table in the EXACT template form of 30.07.26."""
    table = doc.add_table(rows=0, cols=5)
    set_table_borders(table)
    
    hdr_row = table.add_row()
    headers = ['TT', 'Lớp', 'Bài kiểm tra', 'Nội dung', 'Hình thức']
    for idx, text in enumerate(headers):
        cell = hdr_row.cells[idx]
        cell.text = text
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)
            
    assessments = [
        ('1', grade_name, 'Đánh giá định kỳ 1 (Giữa HK1)', 'Nội dung Học kỳ 1 (Tuần 1 - 9)', 'Viết trên giấy / Thực hành'),
        ('2', grade_name, 'Đánh giá định kỳ 2 (Cuối HK1)', 'Nội dung Học kỳ 1 (Tuần 1 - 18)', 'Viết trên giấy / Thực hành'),
        ('3', grade_name, 'Đánh giá định kỳ 3 (Giữa HK2)', 'Nội dung Học kỳ 2 (Tuần 19 - 27)', 'Viết trên giấy / Thực hành'),
        ('4', grade_name, 'Đánh giá định kỳ 4 (Cuối HK2)', 'Nội dung Học kỳ 2 & Tổng hợp cả năm', 'Viết trên giấy / Thực hành'),
    ]
    
    for item in assessments:
        row = table.add_row()
        for idx in range(5):
            cell = row.cells[idx]
            cell.text = item[idx]
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if idx in [0, 1, 4] else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(13)
                
    t_elem = table._element
    t_elem.getparent().remove(t_elem)
    append_element_before_sectpr(doc, t_elem)

def parse_lessons_from_source(doc_src, table_idx):
    """Extracts lessons from a source table and converts to 7-column format."""
    hk1_items = []
    hk2_items = []

    if table_idx >= len(doc_src.tables):
        return hk1_items, hk2_items

    t = doc_src.tables[table_idx]
    current_hk = 1

    for row in t.rows[1:]:
        row_txt = " ".join([c.text.strip() for c in row.cells])
        if "Học kì 2" in row_txt or "HỌC KÌ II" in row_txt or "Học kỳ 2" in row_txt:
            current_hk = 2
            continue
        elif "Học kì 1" in row_txt or "HỌC KÌ I" in row_txt or "Học kỳ 1" in row_txt:
            current_hk = 1
            continue

        cell_vals = [c.text.strip() for c in row.cells]
        if len(cell_vals) >= 4 and cell_vals[0]:
            stt = cell_vals[0]
            bai = cell_vals[1] if len(cell_vals) > 1 else ''
            so_tiet = cell_vals[2] if len(cell_vals) > 2 else '1'
            tuan = ''
            tiet_ppct = cell_vals[2] if len(cell_vals) > 2 else ''
            noi_dung = bai
            yccd = cell_vals[3] if len(cell_vals) > 3 else ''

            if len(cell_vals) >= 5:
                tiet_ppct = cell_vals[2]
                yccd = cell_vals[4] if len(cell_vals) > 4 else cell_vals[3]

            item = [stt, bai, so_tiet, tuan, tiet_ppct, noi_dung, yccd]
            if current_hk == 1:
                hk1_items.append(item)
            else:
                hk2_items.append(item)

    return hk1_items, hk2_items

def add_subject_section(doc_out, s_idx, s_name, has_data, source_doc, table_indices):
    """Adds a subject section in EXACT 30.07.26 TEMPLATE FORM."""
    append_p(doc_out, f"{s_idx}. Môn: {s_name}", bold=True, font_size=14)

    for g_num in range(6, 9):
        g_name = f"{s_name} {g_num}"
        append_p(doc_out, f"{s_idx}.{g_num-5}. Kế hoạch dạy học môn {g_name}", bold=True, font_size=13)
        append_p(doc_out, f"{s_idx}.{g_num-5}.1. Kế hoạch dạy học chính khoá", bold=True)
        append_p(doc_out, "Cả năm: 35 Tiết")
        append_p(doc_out, "Học kì 1: 18 Tiết")
        append_p(doc_out, "Học kì 2: 17 Tiết")
        append_p(doc_out, "Số điểm kiểm tra thường xuyên: 4/kỳ")
        append_p(doc_out, "Số điểm kiểm tra định kỳ:  2/kỳ")
        append_p(doc_out, "")

        if has_data and source_doc and (g_num - 6) < len(table_indices):
            t_idx = table_indices[g_num - 6]
            hk1, hk2 = parse_lessons_from_source(source_doc, t_idx)
            create_form_table_7col(doc_out, hk1, hk2)
        else:
            # Empty template form
            create_form_table_7col(doc_out, [], [])
            append_p(doc_out, f"(Ghi chú: Mục môn {s_name} - Đã ghi nhận form mẫu và sẽ bổ sung chi tiết Khung PPCT sau)", italic=True)

        append_p(doc_out, "")
        append_p(doc_out, f"{s_idx}.{g_num-5}.2. Kế hoạch dạy học tăng cường", bold=True)
        append_p(doc_out, ".....")
        append_p(doc_out, "")
        append_p(doc_out, f"{s_idx}.{g_num-5}.3. Kế hoạch kiểm tra đánh giá", bold=True)
        create_form_table_5col(doc_out, f"Lớp {g_num}A1")
        append_p(doc_out, "")

def main():
    print(f"Reading 30.07.26 master template from {FILE_30_07_26}...")
    doc_base = docx.Document(FILE_30_07_26)

    doc_out = docx.Document()

    copy_mode = True
    section_vii_paragraphs = []
    section_vii_reached = False

    for p in doc_base.paragraphs:
        txt = p.text.strip()
        if txt.startswith("VI. KẾ HOẠCH GIẢNG DẠY"):
            copy_mode = False
            continue
        elif txt.startswith("VII. KẾ HOẠCH CỤ THỂ") or txt.startswith("VIII. Nhiệm vụ"):
            section_vii_reached = True
            copy_mode = False

        if section_vii_reached:
            section_vii_paragraphs.append(p)
        elif copy_mode:
            p_elem = copy.deepcopy(p._element)
            append_element_before_sectpr(doc_out, p_elem)

    # Copy general tables (Tables 0..8 before Section VI)
    for t_i in range(min(9, len(doc_base.tables))):
        tbl_elem = copy.deepcopy(doc_base.tables[t_i]._element)
        append_element_before_sectpr(doc_out, tbl_elem)
        append_p(doc_out, "")

    print("Building Section VI. KẾ HOẠCH GIẢNG DẠY strictly using 30.07.26 form template for 14 subjects...")
    append_p(doc_out, "VI. KẾ HOẠCH GIẢNG DẠY", bold=True, font_size=14)
    append_p(doc_out, "")

    # Source documents
    f_van_gddp = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch dạy học môn Văn_GDDP (THCS)-PPCT.docx')
    f_toan = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch môn Toán THCS (Kế hoạch cá nhân).docx')
    f_anh = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch dạy học môn Tiếng Anh (THCS)-PPCT.docx')
    f_khtn = os.path.join(DIR_TOCM, 'PL1. Tổng hợp khung kế hoạch dạy học môn KHTN6,7,8.26-27.docx')
    f_tin = os.path.join(DIR_MAU, 'Kế hoạch dạy học môn Tin học (THCS) - 2026 - 2027.docx')
    f_robotics = os.path.join(DIR_MAU, 'Kế hoạch dạy học môn Robotics (THCS) - 2026 - 2027.docx')
    f_hdtn = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch môn Toán + HĐTN TH, THCS. Bùi Thị Phương Anh(Kế hoạch cá nhân).docx')
    f_mythuat = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch dạy học môn Mĩ Thuật (THCS)-PPCT.docx')
    f_gdtc = os.path.join(DIR_TOCM, 'PL1. Khung kế hoạch dạy học môn GDTC 26-27.docx')

    doc_van = docx.Document(f_van_gddp) if os.path.exists(f_van_gddp) else None
    doc_toan = docx.Document(f_toan) if os.path.exists(f_toan) else None
    doc_anh = docx.Document(f_anh) if os.path.exists(f_anh) else None
    doc_khtn = docx.Document(f_khtn) if os.path.exists(f_khtn) else None
    doc_tin = docx.Document(f_tin) if os.path.exists(f_tin) else None
    doc_rob = docx.Document(f_robotics) if os.path.exists(f_robotics) else None
    doc_hdtn = docx.Document(f_hdtn) if os.path.exists(f_hdtn) else None
    doc_myt = docx.Document(f_mythuat) if os.path.exists(f_mythuat) else None
    doc_gdtc = docx.Document(f_gdtc) if os.path.exists(f_gdtc) else None

    # 1. Ngữ văn
    add_subject_section(doc_out, 1, "Ngữ văn", True, doc_van, [0, 1, 2])
    # 2. Toán
    add_subject_section(doc_out, 2, "Toán", True, doc_toan, [0, 1, 2])
    # 3. Tiếng Anh
    add_subject_section(doc_out, 3, "Tiếng Anh", True, doc_anh, [2, 3, 4])
    # 4. Tiếng Trung
    add_subject_section(doc_out, 4, "Tiếng Trung", False, None, [])
    # 5. Khoa học tự nhiên
    add_subject_section(doc_out, 5, "Khoa học tự nhiên", True, doc_khtn, [0, 1, 2])
    # 6. Lịch sử và Địa lý
    add_subject_section(doc_out, 6, "Lịch sử và Địa lý", False, None, [])
    # 7. Tin học
    add_subject_section(doc_out, 7, "Tin học", True, doc_tin, [0, 1, 2])
    # 8. Robotics
    add_subject_section(doc_out, 8, "Robotics", True, doc_rob, [0, 1, 2])
    # 9. Giáo dục công dân
    add_subject_section(doc_out, 9, "Giáo dục công dân", False, None, [])
    # 10. Giáo dục địa phương
    add_subject_section(doc_out, 10, "Giáo dục địa phương", True, doc_van, [3, 4, 5])
    # 11. Hoạt động trải nghiệm
    add_subject_section(doc_out, 11, "Hoạt động trải nghiệm, hướng nghiệp", True, doc_hdtn, [5, 6, 7])
    # 12. Âm nhạc
    add_subject_section(doc_out, 12, "Âm nhạc", False, None, [])
    # 13. Mỹ thuật
    add_subject_section(doc_out, 13, "Mỹ thuật", True, doc_myt, [3, 4, 5])
    # 14. Giáo dục thể chất
    add_subject_section(doc_out, 14, "Giáo dục thể chất", True, doc_gdtc, [15, 16, 17])

    # Append Section VII & VIII
    if section_vii_paragraphs:
        for p in section_vii_paragraphs:
            p_elem = copy.deepcopy(p._element)
            append_element_before_sectpr(doc_out, p_elem)

    # Set Times New Roman 13pt
    for p in doc_out.paragraphs:
        p.style.font.name = 'Times New Roman'
        for r in p.runs:
            r.font.name = 'Times New Roman'

    for t in doc_out.tables:
        set_table_borders(t)
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.style.font.name = 'Times New Roman'
                    for r in p.runs:
                        r.font.name = 'Times New Roman'

    # Save to BOTH 30.07.26 file and target master file
    for path in [FILE_30_07_26, TARGET_FILE]:
        try:
            doc_out.save(path)
            print(f"Successfully saved cleanly to {path}")
        except PermissionError:
            print(f"Permission error saving {path}")

if __name__ == '__main__':
    main()
