import os
import sys
import shutil
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

sys.stdout.reconfigure(encoding='utf-8')

TMPL_FILE = r'D:\UNIGO\Hệ thống mẫu văn bản\Lịch báo giảng.docx'
OUT_DIR = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Lịch báo giảng'

OUT_LBG = os.path.join(OUT_DIR, 'Lịch báo giảng.docx')
OUT_LBG_FULL = os.path.join(OUT_DIR, 'Lịch báo giảng cả năm (35 tuần).docx')

os.makedirs(OUT_DIR, exist_ok=True)

# Timetable data for Teacher Nguyên (Tin học & Robotics)
# Format: day_idx (0=Hai, 1=Ba, 2=Tư, 3=Năm, 4=Sáu), period (1..5), session ('sang'/'chieu'), mon, lop, ppct, ten_bai, do_dung
SCHEDULE = [
    # Thứ Hai (03/08)
    (0, 4, 'sang', 'Tin học', '1A1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (0, 1, 'chieu', 'Tin học', '7A1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    
    # Thứ Ba (04/08)
    (1, 3, 'sang', 'Tin học', '5C1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (1, 4, 'sang', 'Robotics', '5C1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
    (1, 1, 'chieu', 'Tin học', '2A1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (1, 2, 'chieu', 'Tin học', '1C1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (1, 3, 'chieu', 'Tin học', '7A1', '2', 'Bài 1. Thiết bị vào - ra', 'Phòng Tin học'),
    (1, 4, 'chieu', 'Robotics', '7A1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
    
    # Thứ Tư (05/08)
    (2, 2, 'sang', 'Robotics', '3A1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
    (2, 4, 'sang', 'Tin học', '3C1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (2, 3, 'chieu', 'Tin học', '4C1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (2, 4, 'chieu', 'Robotics', '2A1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
    
    # Thứ Năm (06/08)
    (3, 2, 'sang', 'Robotics', '1A1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
    (3, 4, 'sang', 'Robotics', '2C1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
    (3, 1, 'chieu', 'Tin học', '3A1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (3, 2, 'chieu', 'Tin học', 'TTH1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (3, 3, 'chieu', 'Tin học', 'TTH2', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (3, 4, 'chieu', 'Robotics', '1C1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
    
    # Thứ Sáu (07/08)
    (4, 1, 'sang', 'Tin học', '2C1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (4, 2, 'sang', 'Robotics', '3C1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
    (4, 4, 'sang', 'Tin học', '6A1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (4, 5, 'sang', 'Robotics', '6A1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
    (4, 1, 'chieu', 'Tin học', '8A1', '1', 'Tiết 0: Định hướng môn học', 'Phòng Tin học'),
    (4, 2, 'chieu', 'Robotics', '8A1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
    (4, 3, 'chieu', 'Robotics', '4C1', '1', 'Tiết 0: Định hướng môn học', 'Bộ Kit Robotics'),
]

DAYS_DATES = [
    ("Hai (03/08)", 0),
    ("Ba (04/08)", 1),
    ("Tư (05/08)", 2),
    ("Năm (06/08)", 3),
    ("Sáu (07/08)", 4)
]

def set_table_borders(table):
    """Sets full black grid borders."""
    table.style = 'Table Grid'
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        for child in list(tblPr[0]):
            if child.tag.endswith('tblBorders'):
                tblPr[0].remove(child)
        borders_xml = parse_xml(
            r'<w:tblBorders %s>'
            r'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        tblPr[0].append(borders_xml)

def set_font_all(doc):
    """Sets Times New Roman 13pt across all paragraphs and tables."""
    for p in doc.paragraphs:
        p.style.font.name = 'Times New Roman'
        p.style.font.size = Pt(13)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)
            
    for t in doc.tables:
        set_table_borders(t)
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.style.font.name = 'Times New Roman'
                    p.style.font.size = Pt(13)
                    for r in p.runs:
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(13)

def populate_lbg_table(table, session_type):
    """Fills in a 25-row timetable (5 days x 5 periods) for Sáng or Chiều."""
    # Organize schedule lookup dictionary: (day_idx, period) -> item
    slot_dict = {}
    for item in SCHEDULE:
        d_idx, period, sess, mon, lop, ppct, ten_bai, do_dung = item
        if sess == session_type:
            slot_dict[(d_idx, period)] = {
                'mon': mon, 'lop': lop, 'ppct': ppct, 'ten_bai': ten_bai, 'do_dung': do_dung
            }
            
    # Iterate through rows 1..25
    # Row structure: 5 days x 5 periods = 25 rows
    for day_idx in range(5):
        day_name, _ = DAYS_DATES[day_idx]
        for period in range(1, 6):
            r_idx = day_idx * 5 + period
            if r_idx < len(table.rows):
                row = table.rows[r_idx]
                row.cells[0].text = day_name
                row.cells[1].text = str(period)
                
                if (day_idx, period) in slot_dict:
                    data = slot_dict[(day_idx, period)]
                    row.cells[2].text = data['ppct']
                    row.cells[3].text = data['mon']
                    row.cells[4].text = data['lop']
                    row.cells[5].text = data['ten_bai']
                    row.cells[6].text = data['do_dung']
                else:
                    row.cells[2].text = ""
                    row.cells[3].text = ""
                    row.cells[4].text = ""
                    row.cells[5].text = ""
                    row.cells[6].text = ""

def process_lbg():
    print(f"Reading template {TMPL_FILE}...")
    doc = docx.Document(TMPL_FILE)
    
    # Update teacher info table 0
    t0 = doc.tables[0]
    info_text = (
        "Họ và tên giáo viên: Nguyễn Đình Nguyên\n"
        "Tổ chuyên môn: Tin học - Robotics\n"
        "Năm học: 2026 - 2027"
    )
    t0.rows[0].cells[0].text = info_text
    
    # Update paragraphs (Week header)
    for p in doc.paragraphs:
        txt = p.text.strip()
        if "TUẦN 01" in txt or "TUẦN 1" in txt or "từ ngày" in txt:
            p.text = "TUẦN 01 (từ ngày 03/08/2026 đến ngày 07/08/2026)"
        if "Buổi" in txt and "Tuần" in txt:
            if "Sáng" in txt or "SÁNG" in txt:
                p.text = "Buổi: Sáng    Tuần: 01 (Từ ngày 03/08/2026 đến ngày 07/08/2026)"
            elif "Chiều" in txt or "CHIỀU" in txt:
                p.text = "Buổi: Chiều    Tuần: 01 (Từ ngày 03/08/2026 đến ngày 07/08/2026)"
                
    # Populate Table 1 (Sáng) and Table 3 (Chiều)
    print("Populating Table 1 (Sáng)...")
    populate_lbg_table(doc.tables[1], 'sang')
    
    print("Populating Table 3 (Chiều)...")
    populate_lbg_table(doc.tables[3], 'chieu')
    
    set_font_all(doc)
    
    doc.save(OUT_LBG)
    print(f"Successfully saved {OUT_LBG}")
    
    doc.save(OUT_LBG_FULL)
    print(f"Successfully saved {OUT_LBG_FULL}")

if __name__ == '__main__':
    process_lbg()
