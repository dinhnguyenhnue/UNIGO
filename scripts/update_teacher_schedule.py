import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter

import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def update_excel_schedule():
    xlsx_path = r'Thời khóa biểu giáo viên/Thời khóa biểu - Đậu Đình Nguyên.xlsx'
    wb = openpyxl.load_workbook(xlsx_path)
    
    # Use or create sheet
    if 'TKB Cá Nhân' in wb.sheetnames:
        ws = wb['TKB Cá Nhân']
    else:
        ws = wb.active

    # Unmerge all existing cells to start fresh layout
    merged_ranges = list(ws.merged_cells.ranges)
    for rng in merged_ranges:
        ws.unmerge_cells(str(rng))

    ws.delete_rows(1, ws.max_row)

    # Styles
    font_title = Font(name='Segoe UI', size=14, bold=True, color='FFFFFF')
    font_subtitle = Font(name='Segoe UI', size=11, bold=True, color='1D2A64')
    font_header = Font(name='Segoe UI', size=11, bold=True, color='FFFFFF')
    font_buoi = Font(name='Segoe UI', size=11, bold=True, color='1D2A64')
    font_tiet = Font(name='Segoe UI', size=10, bold=True, color='0F172A')
    font_time = Font(name='Segoe UI', size=9.5, italic=False, bold=True, color='334155')
    
    font_tin = Font(name='Segoe UI', size=10.5, bold=True, color='0369A1')
    font_rob = Font(name='Segoe UI', size=10.5, bold=True, color='B45309')
    font_empty = Font(name='Segoe UI', size=10, color='94A3B8')

    fill_title = PatternFill(start_color='1D2A64', end_color='1D2A64', fill_type='solid')
    fill_subtitle = PatternFill(start_color='E0E7FF', end_color='E0E7FF', fill_type='solid')
    fill_header = PatternFill(start_color='1D2A64', end_color='1D2A64', fill_type='solid')
    fill_buoi_sang = PatternFill(start_color='F1F5F9', end_color='F1F5F9', fill_type='solid')
    fill_buoi_chieu = PatternFill(start_color='E2E8F0', end_color='E2E8F0', fill_type='solid')
    fill_tin = PatternFill(start_color='E0F2FE', end_color='E0F2FE', fill_type='solid')
    fill_rob = PatternFill(start_color='FEF3C7', end_color='FEF3C7', fill_type='solid')
    fill_empty = PatternFill(start_color='FFFFFF', end_color='FFFFFF', fill_type='solid')

    thin_border = Border(
        left=Side(style='thin', color='CBD5E1'),
        right=Side(style='thin', color='CBD5E1'),
        top=Side(style='thin', color='CBD5E1'),
        bottom=Side(style='thin', color='CBD5E1')
    )

    align_center = Alignment(horizontal='center', vertical='center', wrap_text=True)
    align_left = Alignment(horizontal='left', vertical='center', wrap_text=True)

    # Row 1: Title
    ws.merge_cells('A1:I1')
    c1 = ws['A1']
    c1.value = "UNIGO COLLEGE — THỜI KHOÁ BIỂU CÁ NHÂN GIÁO VIÊN"
    c1.font = font_title
    c1.fill = fill_title
    c1.alignment = align_center
    ws.row_dimensions[1].height = 35

    # Row 2: Subtitle
    ws.merge_cells('A2:I2')
    c2 = ws['A2']
    c2.value = "Giáo viên: ĐẬU ĐÌNH NGUYÊN  |  Bộ môn: Tin học & Robotics  |  Tổng số tiết/tuần: 26 tiết"
    c2.font = font_subtitle
    c2.fill = fill_subtitle
    c2.alignment = align_center
    ws.row_dimensions[2].height = 26

    # Row 3: Headers
    headers = ['BUỔI', 'TIẾT', 'THỜI GIAN', 'THỨ 2', 'THỨ 3', 'THỨ 4', 'THỨ 5', 'THỨ 6', 'THỨ 7']
    ws.append(headers)
    ws.row_dimensions[3].height = 28
    for col_num in range(1, 10):
        cell = ws.cell(3, col_num)
        cell.font = font_header
        cell.fill = fill_header
        cell.alignment = align_center
        cell.border = thin_border

    # Data Rows definition
    # format: (buoi, tiet, time_str, [thu2, thu3, thu4, thu5, thu6, thu7])
    data = [
        # SÁNG
        ('SÁNG', 'Tiết 1', '08:15 - 08:50\n(08:05-08:50)', ['', '', '', 'Tin - TT3', '', '']),
        ('SÁNG', 'Tiết 2', '08:55 - 9:30\n(08:55-09:40)', ['', '', 'Robotics - 3A1', 'Robotics - 1A1', 'Robotics - 3C1', '']),
        ('SÁNG', 'Tiết 3', '09:45 - 10:20\n(09:45-10:25)', ['Tin - 7A1', 'Tin - 5C1', '', '', 'Tin - 2C1', '']),
        ('SÁNG', 'Tiết 4', '10:25 - 11:00\n(10:30-11:10)', ['Tin - 1A1', 'Robotics - 5C1', 'Tin - 3C1', 'Robotics - 2C1', 'Tin - 6A1', '']),
        ('SÁNG', 'Tiết 5', '11:15 - 11:50\n(THCS)', ['', '', '', '', 'Robotics - 6A1', '']),
        
        # CHIỀU
        ('CHIỀU', 'Tiết 1', '13:15 - 13:50\n(13:05-13:50)', ['', 'Tin - 2A1', '', 'Tin - 3A1', 'Tin - 8A1', '']),
        ('CHIỀU', 'Tiết 2', '13:55 - 14:30\n(13:55-14:40)', ['', 'Tin - 1C1', '', 'Tin - TTH 1', 'Robotics - 8A1', '']),
        ('CHIỀU', 'Tiết 3', '14:55 - 15:30\n(14:50-15:30)', ['', 'Tin - 7A1', 'Tin - 4C1', 'Tin - TTH2', 'Robotics - 4C1', '']),
        ('CHIỀU', 'Tiết 4', '15:35 - 16:10\n(15:35-16:20)', ['', 'Robotics - 7A1', 'Robotics - 2A1', 'Robotics - 1C1', '', '']),
        ('CHIỀU', 'Tiết 5', '16:10 - 16:25\n(Dọn dẹp)', ['', '', '', '', '', '']),
    ]

    for idx, row_item in enumerate(data):
        buoi, tiet, time_val, days = row_item
        row_num = idx + 4
        
        ws.cell(row_num, 1, buoi)
        ws.cell(row_num, 2, tiet)
        ws.cell(row_num, 3, time_val)

        for col_idx, subject in enumerate(days):
            ws.cell(row_num, col_idx + 4, subject)

        ws.row_dimensions[row_num].height = 36

    # Formatting data rows
    for r in range(4, 14):
        for c in range(1, 10):
            cell = ws.cell(r, c)
            cell.border = thin_border
            cell.alignment = align_center

            val = str(cell.value or '').strip()

            if c == 1:
                cell.font = font_buoi
                cell.fill = fill_buoi_sang if r <= 8 else fill_buoi_chieu
            elif c == 2:
                cell.font = font_tiet
                cell.fill = PatternFill(start_color='FAFAFA', fill_type='solid')
            elif c == 3:
                cell.font = font_time
                cell.fill = PatternFill(start_color='F8FAFC', fill_type='solid')
            else:
                if 'Tin' in val:
                    cell.font = font_tin
                    cell.fill = fill_tin
                elif 'Robotics' in val:
                    cell.font = font_rob
                    cell.fill = fill_rob
                else:
                    cell.font = font_empty
                    cell.fill = fill_empty

    # Merge BUỔI cells
    ws.merge_cells('A4:A8')
    ws.merge_cells('A9:A13')
    ws['A4'].alignment = align_center
    ws['A9'].alignment = align_center

    # Auto column widths
    col_widths = {
        'A': 11,
        'B': 10,
        'C': 16,
        'D': 18,
        'E': 18,
        'F': 18,
        'G': 18,
        'H': 18,
        'I': 12
    }
    for col_letter, width in col_widths.items():
        ws.column_dimensions[col_letter].width = width

    # Page print settings
    ws.page_setup.orientation = ws.ORIENTATION_LANDSCAPE
    ws.page_setup.paperSize = ws.PAPERSIZE_A4
    ws.sheet_properties.pageSetUpPr.fitToPage = True

    wb.save(xlsx_path)
    print("Excel timetable updated successfully.")

def update_docx_schedule():
    docx_path = r'Thời khóa biểu giáo viên/Thời khóa biểu - Đậu Đình Nguyên.docx'
    doc = Document()

    section = doc.sections[0]
    section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11.69)
    section.page_height = Inches(8.27)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)

    def set_cell_border(cell, color="CBD5E1", sz="4"):
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for b in ['top', 'left', 'bottom', 'right']:
            node = OxmlElement(f'w:{b}')
            node.set(qn('w:val'), 'single')
            node.set(qn('w:sz'), sz)
            node.set(qn('w:space'), '0')
            node.set(qn('w:color'), color)
            tcBorders.append(node)
        tcPr.append(tcBorders)

    def set_cell_shading(cell, color_hex):
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
        cell._tc.get_or_add_tcPr().append(shd)

    def format_run(run, font_name="Times New Roman", font_size=11, bold=False, italic=False, color_rgb=(0,0,0)):
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic
        run.font.color.rgb = RGBColor(*color_rgb)

    # Header title
    p_sublogo = doc.add_paragraph()
    p_sublogo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sublogo.paragraph_format.space_after = Pt(2)
    r_sublogo = p_sublogo.add_run("UNIGO COLLEGE — JUNIOR & MIDDLE SCHOOL")
    format_run(r_sublogo, font_size=10, bold=True, color_rgb=(100, 110, 140))

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run("THỜI KHOÁ BIỂU CÁ NHÂN GIÁO VIÊN: ĐẬU ĐÌNH NGUYÊN")
    format_run(r_title, font_size=16, bold=True, color_rgb=(29, 42, 100))

    p_info = doc.add_paragraph()
    p_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_info.paragraph_format.space_after = Pt(12)
    r_info = p_info.add_run("Bộ môn: Tin học & Robotics  |  Tổng số tiết/tuần: 26 tiết  |  Áp dụng từ 04/08/2025")
    format_run(r_info, font_size=11, bold=True, color_rgb=(15, 23, 42))

    # Table 9 columns
    headers = ['BUỔI', 'TIẾT', 'THỜI GIAN', 'THỨ 2', 'THỨ 3', 'THỨ 4', 'THỨ 5', 'THỨ 6', 'THỨ 7']
    data = [
        ('SÁNG', 'Tiết 1', '08:15 - 08:50\n(08:05-08:50)', '', '', '', 'Tin - TT3', '', ''),
        ('SÁNG', 'Tiết 2', '08:55 - 09:30\n(08:55-09:40)', '', '', 'Robotics - 3A1', 'Robotics - 1A1', 'Robotics - 3C1', ''),
        ('SÁNG', 'Tiết 3', '09:45 - 10:20\n(09:45-10:25)', 'Tin - 7A1', 'Tin - 5C1', '', '', 'Tin - 2C1', ''),
        ('SÁNG', 'Tiết 4', '10:25 - 11:00\n(10:30-11:10)', 'Tin - 1A1', 'Robotics - 5C1', 'Tin - 3C1', 'Robotics - 2C1', 'Tin - 6A1', ''),
        ('SÁNG', 'Tiết 5', '11:15 - 11:50\n(THCS)', '', '', '', '', 'Robotics - 6A1', ''),
        
        ('CHIỀU', 'Tiết 1', '13:15 - 13:50\n(13:05-13:50)', '', 'Tin - 2A1', '', 'Tin - 3A1', 'Tin - 8A1', ''),
        ('CHIỀU', 'Tiết 2', '13:55 - 14:30\n(13:55-14:40)', '', 'Tin - 1C1', '', 'Tin - TTH 1', 'Robotics - 8A1', ''),
        ('CHIỀU', 'Tiết 3', '14:55 - 15:30\n(14:50-15:30)', '', 'Tin - 7A1', 'Tin - 4C1', 'Tin - TTH2', 'Robotics - 4C1', ''),
        ('CHIỀU', 'Tiết 4', '15:35 - 16:10\n(15:35-16:20)', '', 'Robotics - 7A1', 'Robotics - 2A1', 'Robotics - 1C1', '', ''),
        ('CHIỀU', 'Tiết 5', '16:10 - 16:25\n(Dọn dẹp)', '', '', '', '', '', ''),
    ]

    table = doc.add_table(rows=len(data) + 1, cols=9)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_widths = [Inches(0.9), Inches(0.8), Inches(1.4), Inches(1.3), Inches(1.3), Inches(1.3), Inches(1.3), Inches(1.3), Inches(0.9)]

    # Format Header
    hdr_cells = table.rows[0].cells
    for i, title in enumerate(headers):
        hdr_cells[i].text = title
        set_cell_shading(hdr_cells[i], "1D2A64")
        set_cell_border(hdr_cells[i], color="1D2A64")
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        hdr_cells[i].width = col_widths[i]
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            format_run(r, font_size=10.5, bold=True, color_rgb=(255, 255, 255))

    # Format Data Rows
    for r_idx, row_item in enumerate(data):
        row_cells = table.rows[r_idx + 1].cells
        buoi, tiet, time_val, *days = row_item

        row_vals = [buoi, tiet, time_val] + days
        for c_idx, val in enumerate(row_vals):
            cell = row_cells[c_idx]
            cell.text = val
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell.width = col_widths[c_idx]
            set_cell_border(cell, color="CBD5E1")

            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Cell shading & fonts
            if c_idx == 0:
                set_cell_shading(cell, "F1F5F9" if r_idx < 5 else "E2E8F0")
                for r in p.runs:
                    format_run(r, font_size=10, bold=True, color_rgb=(29, 42, 100))
            elif c_idx == 1:
                set_cell_shading(cell, "FAFAFA")
                for r in p.runs:
                    format_run(r, font_size=10, bold=True, color_rgb=(15, 23, 42))
            elif c_idx == 2:
                set_cell_shading(cell, "F8FAFC")
                for r in p.runs:
                    format_run(r, font_size=9, bold=True, color_rgb=(51, 65, 85))
            else:
                if 'Tin' in val:
                    set_cell_shading(cell, "E0F2FE")
                    for r in p.runs:
                        format_run(r, font_size=10.5, bold=True, color_rgb=(3, 105, 161))
                elif 'Robotics' in val:
                    set_cell_shading(cell, "FEF3C7")
                    for r in p.runs:
                        format_run(r, font_size=10.5, bold=True, color_rgb=(180, 83, 9))
                else:
                    set_cell_shading(cell, "FFFFFF")
                    for r in p.runs:
                        format_run(r, font_size=9.5, color_rgb=(148, 163, 184))

    # Merge BUỔI cells for SÁNG and CHIỀU
    # SÁNG: rows 1 to 5 (cells table.rows[1].cells[0] to table.rows[5].cells[0])
    table.rows[1].cells[0].merge(table.rows[5].cells[0])
    table.rows[1].cells[0].text = "SÁNG"
    table.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_run(table.rows[1].cells[0].paragraphs[0].runs[0], font_size=11, bold=True, color_rgb=(29, 42, 100))

    # CHIỀU: rows 6 to 10
    table.rows[6].cells[0].merge(table.rows[10].cells[0])
    table.rows[6].cells[0].text = "CHIỀU"
    table.rows[6].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    format_run(table.rows[6].cells[0].paragraphs[0].runs[0], font_size=11, bold=True, color_rgb=(29, 42, 100))

    doc.save(docx_path)
    print("DOCX timetable updated successfully.")

if __name__ == '__main__':
    update_excel_schedule()
    update_docx_schedule()
