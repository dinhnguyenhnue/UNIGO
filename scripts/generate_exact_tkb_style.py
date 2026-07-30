import os
import sys
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

sys.stdout.reconfigure(encoding='utf-8')

OUT_DIR_WORD = r'D:\UNIGO\Thời khóa biểu giáo viên'
OUT_DIR_MAU = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Thời khóa biểu'

os.makedirs(OUT_DIR_WORD, exist_ok=True)
os.makedirs(OUT_DIR_MAU, exist_ok=True)

# 4 Teachers Schedule Data from TKB CHECK lần 2.1
TEACHERS = {
    'Nguyên': {
        'full_name': 'Đậu Đình Nguyên',
        'subject': 'Tin học & Robotics',
        'slots': [
            # (day '2'..'7', session 'SÁNG'/'CHIỀU', period 1..5, text_val, subject_type)
            ('2', 'SÁNG', 3, 'Tin - 7A1', 'Tin'),
            ('2', 'SÁNG', 4, 'Tin - 1A1', 'Tin'),
            ('3', 'SÁNG', 3, 'Tin - 5C1', 'Tin'),
            ('3', 'SÁNG', 4, 'Robotics - 5C1', 'Robotics'),
            ('3', 'CHIỀU', 1, 'Tin - 2A1', 'Tin'),
            ('3', 'CHIỀU', 2, 'Tin - 1C1', 'Tin'),
            ('3', 'CHIỀU', 3, 'Tin - 7A1', 'Tin'),
            ('3', 'CHIỀU', 4, 'Robotics - 7A1', 'Robotics'),
            ('4', 'SÁNG', 2, 'Robotics - 3A1', 'Robotics'),
            ('4', 'SÁNG', 4, 'Tin - 3C1', 'Tin'),
            ('4', 'CHIỀU', 3, 'Tin - 4C1', 'Tin'),
            ('4', 'CHIỀU', 4, 'Robotics - 2A1', 'Robotics'),
            ('5', 'SÁNG', 1, 'Tin - TT3', 'Tin'),
            ('5', 'SÁNG', 2, 'Robotics - 1A1', 'Robotics'),
            ('5', 'SÁNG', 4, 'Robotics - 2C1', 'Robotics'),
            ('5', 'CHIỀU', 1, 'Tin - 3A1', 'Tin'),
            ('5', 'CHIỀU', 2, 'Tin - TTH 1', 'Tin'),
            ('5', 'CHIỀU', 3, 'Tin - TTH2', 'Tin'),
            ('5', 'CHIỀU', 4, 'Robotics - 1C1', 'Robotics'),
            ('6', 'SÁNG', 2, 'Robotics - 3C1', 'Robotics'),
            ('6', 'SÁNG', 3, 'Tin - 2C1', 'Tin'),
            ('6', 'SÁNG', 4, 'Tin - 6A1', 'Tin'),
            ('6', 'SÁNG', 5, 'Robotics - 6A1', 'Robotics'),
            ('6', 'CHIỀU', 1, 'Tin - 8A1', 'Tin'),
            ('6', 'CHIỀU', 2, 'Robotics - 8A1', 'Robotics'),
            ('6', 'CHIỀU', 3, 'Robotics - 4C1', 'Robotics'),
        ]
    },
    'Hà': {
        'full_name': 'Nguyễn Thị Hà',
        'subject': 'Ngữ văn & LS-ĐL-GDCD',
        'slots': [
            ('2', 'SÁNG', 1, 'LS&ĐL - 6A1', 'Van'),
            ('2', 'SÁNG', 2, 'LS&ĐL - 8A1', 'Van'),
            ('2', 'CHIỀU', 2, 'Văn - 6A1', 'Van'),
            ('2', 'CHIỀU', 3, 'Văn - 6A1', 'Van'),
            ('2', 'CHIỀU', 4, 'LS&ĐL - 7A1', 'Van'),
            ('3', 'SÁNG', 2, 'GDCD - 8A1', 'Other'),
            ('3', 'SÁNG', 3, 'Văn - 6A1', 'Van'),
            ('3', 'CHIỀU', 2, 'Văn - 6A1', 'Van'),
            ('3', 'CHIỀU', 3, 'Đọc sách - 6A1', 'Other'),
            ('4', 'SÁNG', 3, 'Văn - 6A1', 'Van'),
            ('4', 'SÁNG', 4, 'Văn - 6A1', 'Van'),
            ('4', 'CHIỀU', 2, 'LS&ĐL - 7A1', 'Van'),
            ('4', 'CHIỀU', 3, 'GDCD - 6A1', 'Other'),
            ('4', 'CHIỀU', 4, 'LS&ĐL - 8A1', 'Van'),
            ('5', 'SÁNG', 2, 'LS&ĐL - 8A1', 'Van'),
            ('5', 'SÁNG', 4, 'LS&ĐL - 6A1', 'Van'),
            ('6', 'SÁNG', 1, 'LS&ĐL - 6A1', 'Van'),
            ('6', 'SÁNG', 2, 'GDCD - 7A1', 'Other'),
            ('6', 'SÁNG', 3, 'LS&ĐL - 7A1', 'Van'),
            ('6', 'CHIỀU', 3, 'Văn - 6A1', 'Van'),
            ('6', 'CHIỀU', 4, 'Văn - 6A1', 'Van'),
        ]
    },
    'Ánh NV': {
        'full_name': 'Nguyễn Thị Ánh (Ánh NV)',
        'subject': 'Ngữ văn & GDĐP-Đọc sách',
        'slots': [
            ('2', 'SÁNG', 1, 'GDĐP - 7A1', 'Other'),
            ('2', 'SÁNG', 4, 'Văn - 8A1', 'Van'),
            ('2', 'SÁNG', 5, 'Văn - 8A1', 'Van'),
            ('2', 'CHIỀU', 1, 'GDĐP - 6A1', 'Other'),
            ('2', 'CHIỀU', 2, 'Văn - 7A1', 'Van'),
            ('3', 'SÁNG', 5, 'Văn - 7A1', 'Van'),
            ('3', 'CHIỀU', 3, 'Văn - 8A1', 'Van'),
            ('3', 'CHIỀU', 4, 'Văn - 8A1', 'Van'),
            ('4', 'SÁNG', 1, 'Văn - 7A1', 'Van'),
            ('4', 'SÁNG', 4, 'Văn - 7A1', 'Van'),
            ('4', 'SÁNG', 5, 'Đọc sách - 8A1', 'Other'),
            ('4', 'CHIỀU', 4, 'Đọc sách - 7A1', 'Other'),
            ('5', 'SÁNG', 3, 'Văn - 8A1', 'Van'),
            ('5', 'SÁNG', 4, 'Văn - 8A1', 'Van'),
            ('5', 'SÁNG', 5, 'GDĐP - 8A1', 'Other'),
            ('5', 'CHIỀU', 3, 'Văn - 7A1', 'Van'),
            ('5', 'CHIỀU', 4, 'Văn - 7A1', 'Van'),
            ('6', 'SÁNG', 4, 'Văn - 8A1', 'Van'),
            ('6', 'SÁNG', 5, 'Văn - 8A1', 'Van'),
            ('6', 'CHIỀU', 3, 'Văn - 7A1', 'Van'),
            ('6', 'CHIỀU', 4, 'Văn - 7A1', 'Van'),
        ]
    },
    'Tân': {
        'full_name': 'Cô Tân',
        'subject': 'Toán',
        'slots': [
            ('2', 'SÁNG', 1, 'Toán - 8A1', 'Toan'),
            ('2', 'SÁNG', 5, 'Toán - 6A1', 'Toan'),
            ('2', 'CHIỀU', 1, 'Toán - 7A1', 'Toan'),
            ('2', 'CHIỀU', 4, 'Toán - 8A1', 'Toan'),
            ('3', 'SÁNG', 2, 'Toán - 7A1', 'Toan'),
            ('3', 'SÁNG', 3, 'Toán - 8A1', 'Toan'),
            ('3', 'SÁNG', 5, 'Toán - 8A1', 'Toan'),
            ('3', 'CHIỀU', 1, 'Toán - 6A1', 'Toan'),
            ('3', 'CHIỀU', 2, 'Toán - 7A1', 'Toan'),
            ('3', 'CHIỀU', 4, 'Toán - 6A1', 'Toan'),
            ('4', 'SÁNG', 1, 'Toán - 8A1', 'Toan'),
            ('4', 'CHIỀU', 1, 'Toán - 7A1', 'Toan'),
            ('4', 'CHIỀU', 4, 'Toán - 6A1', 'Toan'),
            ('5', 'SÁNG', 2, 'Toán - 7A1', 'Toan'),
            ('5', 'SÁNG', 5, 'Toán - 6A1', 'Toan'),
            ('5', 'CHIỀU', 1, 'Toán - 6A1', 'Toan'),
            ('5', 'CHIỀU', 2, 'Toán - 7A1', 'Toan'),
            ('5', 'CHIỀU', 4, 'Toán - 8A1', 'Toan'),
            ('6', 'SÁNG', 2, 'Toán - 6A1', 'Toan'),
            ('6', 'SÁNG', 3, 'Toán - 6A1', 'Toan'),
            ('6', 'CHIỀU', 1, 'Toán - 7A1', 'Toan'),
            ('6', 'CHIỀU', 2, 'Toán - 7A1', 'Toan'),
            ('6', 'CHIỀU', 3, 'Toán - 8A1', 'Toan'),
            ('6', 'CHIỀU', 4, 'Toán - 8A1', 'Toan'),
        ]
    }
}

# Colors matching user's image exactly:
# Navy blue header: 1F497D
# Soft lavender fill for Tin / Van / Toan: D9E1F2
# Soft green fill for Robotics / Other: E2EFDA
FILL_HEADER = PatternFill(start_color="1F497D", end_color="1F497D", fill_type="solid")
FILL_TIN = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
FILL_ROBOTICS = PatternFill(start_color="E2EFDA", end_color="E2EFDA", fill_type="solid")
FILL_OTHER = PatternFill(start_color="FFF2CC", end_color="FFF2CC", fill_type="solid")

FONT_HEADER = Font(name='Times New Roman', size=13, bold=True, color="FFFFFF")
FONT_BUOI = Font(name='Times New Roman', size=13, bold=True, color="1F497D")
FONT_TIET = Font(name='Times New Roman', size=12)
FONT_DATA = Font(name='Times New Roman', size=12, bold=True)

BORDER_THIN = Border(
    left=Side(style='thin', color='BFBFBF'),
    right=Side(style='thin', color='BFBFBF'),
    top=Side(style='thin', color='BFBFBF'),
    bottom=Side(style='thin', color='BFBFBF')
)

def create_excel_exact_style(t_key, data, save_path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "TKB Cá Nhân"

    # Show gridlines explicitly
    ws.views.sheetView[0].showGridLines = True

    # Row 1: Header (BUỔI | TIẾT | THỨ 2 | THỨ 3 | THỨ 4 | THỨ 5 | THỨ 6 | THỨ 7)
    headers = ["BUỔI", "TIẾT", "THỨ 2", "THỨ 3", "THỨ 4", "THỨ 5", "THỨ 6", "THỨ 7"]
    for c_i, h in enumerate(headers, 1):
        cell = ws.cell(1, c_i, h)
        cell.font = FONT_HEADER
        cell.fill = FILL_HEADER
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.border = BORDER_THIN

    # SÁNG: Rows 2..6 (Tiết 1 to Tiết 5)
    ws.merge_cells('A2:A6')
    cell_sang = ws['A2']
    cell_sang.value = "SÁNG"
    cell_sang.font = FONT_BUOI
    cell_sang.alignment = Alignment(horizontal='center', vertical='center')

    for p in range(1, 6):
        r = p + 1
        ws.cell(r, 2, f"Tiết {p}").font = FONT_TIET
        ws.cell(r, 2).alignment = Alignment(horizontal='center', vertical='center')
        for c in range(1, 9):
            ws.cell(r, c).border = BORDER_THIN

    # CHIỀU: Rows 7..11 (Tiết 1 to Tiết 5)
    ws.merge_cells('A7:A11')
    cell_chieu = ws['A7']
    cell_chieu.value = "CHIỀU"
    cell_chieu.font = FONT_BUOI
    cell_chieu.alignment = Alignment(horizontal='center', vertical='center')

    for p in range(1, 6):
        r = p + 6
        ws.cell(r, 2, f"Tiết {p}").font = FONT_TIET
        ws.cell(r, 2).alignment = Alignment(horizontal='center', vertical='center')
        for c in range(1, 9):
            ws.cell(r, c).border = BORDER_THIN

    # Map slot items
    day_col_map = {'2': 3, '3': 4, '4': 5, '5': 6, '6': 7, '7': 8}

    for slot in data['slots']:
        day, sess, p_num, text_val, s_type = slot
        c_idx = day_col_map[day]
        r_idx = (p_num + 1) if sess == 'SÁNG' else (p_num + 6)

        cell = ws.cell(r_idx, c_idx)
        cell.value = text_val
        cell.font = FONT_DATA
        cell.alignment = Alignment(horizontal='center', vertical='center')

        if s_type == 'Robotics':
            cell.fill = FILL_ROBOTICS
        elif s_type in ['Tin', 'Van', 'Toan']:
            cell.fill = FILL_TIN
        else:
            cell.fill = FILL_OTHER

    # Column widths
    ws.column_dimensions['A'].width = 12
    ws.column_dimensions['B'].width = 12
    for c_let in ['C', 'D', 'E', 'F', 'G', 'H']:
        ws.column_dimensions[c_let].width = 20

    wb.save(save_path)
    print(f"Saved Excel exact style: {save_path}")

def set_table_borders(table):
    """Sets explicit full black grid borders."""
    table.style = 'Table Grid'
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        for child in list(tblPr[0]):
            if child.tag.endswith('tblBorders'):
                tblPr[0].remove(child)
        borders_xml = parse_xml(
            r'<w:tblBorders %s>'
            r'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        tblPr[0].append(borders_xml)

def set_cell_background(cell, fill_hex):
    """Sets background color for a docx cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_docx_exact_style(t_key, data, save_paths):
    doc = docx.Document()
    
    # Title Header
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_header.add_run(f"THỜI KHOÁ BIỂU CÁ NHÂN GIÁO VIÊN: {data['full_name'].upper()}\n")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(15)
    r1.bold = True
    
    r2 = p_header.add_run(f"Bộ môn: {data['subject']} | Tổng số tiết/tuần: {len(data['slots'])} tiết\n")
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(13)
    r2.italic = True
    
    # Table 11 rows x 8 columns
    table = doc.add_table(rows=11, cols=8)
    set_table_borders(table)
    
    headers = ["BUỔI", "TIẾT", "THỨ 2", "THỨ 3", "THỨ 4", "THỨ 5", "THỨ 6", "THỨ 7"]
    
    # Row 0 Header
    for c_idx in range(8):
        cell = table.rows[0].cells[c_idx]
        cell.text = headers[c_idx]
        set_cell_background(cell, "1F497D")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(255, 255, 255)
            
    # Merge BUỔI cells for SÁNG (rows 1..5)
    cell_sang = table.rows[1].cells[0]
    for r_i in range(2, 6):
        cell_sang.merge(table.rows[r_i].cells[0])
    cell_sang.text = "SÁNG"
    p = cell_sang.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(31, 73, 125)

    # Merge BUỔI cells for CHIỀU (rows 6..10)
    cell_chieu = table.rows[6].cells[0]
    for r_i in range(7, 11):
        cell_chieu.merge(table.rows[r_i].cells[0])
    cell_chieu.text = "CHIỀU"
    p = cell_chieu.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
        r.font.color.rgb = RGBColor(31, 73, 125)

    # Populate Tiết 1..5
    for p_num in range(1, 6):
        r_sang = table.rows[p_num]
        r_sang.cells[1].text = f"Tiết {p_num}"
        p = r_sang.cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.name = 'Times New Roman'

        r_chieu = table.rows[p_num + 5]
        r_chieu.cells[1].text = f"Tiết {p_num}"
        p = r_chieu.cells[1].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.name = 'Times New Roman'

    day_col_map = {'2': 2, '3': 3, '4': 4, '5': 5, '6': 6, '7': 7}

    for slot in data['slots']:
        day, sess, p_num, text_val, s_type = slot
        c_idx = day_col_map[day]
        r_idx = p_num if sess == 'SÁNG' else (p_num + 5)

        cell = table.rows[r_idx].cells[c_idx]
        cell.text = text_val
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
            r.bold = True

        if s_type == 'Robotics':
            set_cell_background(cell, "E2EFDA")
        elif s_type in ['Tin', 'Van', 'Toan']:
            set_cell_background(cell, "D9E1F2")
        else:
            set_cell_background(cell, "FFF2CC")

    for path in save_paths:
        try:
            doc.save(path)
            print(f"Saved DOCX exact style: {path}")
        except PermissionError:
            print(f"Permission error saving {path}")

def main():
    print("Generating Timetable documents in EXACT USER VISUAL STYLE...")
    for t_key, data in TEACHERS.items():
        fname_word1 = os.path.join(OUT_DIR_WORD, f"Thời khóa biểu - {data['full_name']}.docx")
        fname_word2 = os.path.join(OUT_DIR_MAU, f"Thời khóa biểu - {data['full_name']}.docx")
        fname_excel1 = os.path.join(OUT_DIR_WORD, f"Thời khóa biểu - {data['full_name']}.xlsx")
        fname_excel2 = os.path.join(OUT_DIR_MAU, f"Thời khóa biểu - {data['full_name']}.xlsx")
        
        create_docx_exact_style(t_key, data, [fname_word1, fname_word2])
        create_excel_exact_style(t_key, data, fname_excel1)
        create_excel_exact_style(t_key, data, fname_excel2)

if __name__ == '__main__':
    main()
