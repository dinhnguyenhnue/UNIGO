"""
fix_khbd_robotics_dates.py — Sửa toàn bộ file KHBD Robotics (Lớp 1-8):
  1. Xóa file thừa/trùng lặp ở Tuần_01 (*_Tiet_0_inh_huong_mon_hoc.docx)
  2. Xóa dòng "Lớp" thừa & ngày rác trong Table[0] Row[1] Cell[1] (THCS Lớp 6, 7, 8)
  3. Chuẩn hóa Ngày soạn (Thứ 7 tuần trước) / Ngày dạy (theo LBG) cho toàn bộ các tuần
  4. Chuẩn hóa tên lớp THCS: 6A1, 7A1, 8A1 và TH: 5C1, 4C1, 3A1, 2A1, 1A1
"""
import os, sys, re
sys.stdout.reconfigure(encoding='utf-8')
from datetime import date, timedelta
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ─── SCHEDULE CONFIG ───────────────────────────────────────────────────────
TUAN_01_START = date(2026, 8, 3)

ROBOTICS_SCHEDULE = {
    '1': (3, '1A1'),   # Thứ Năm
    '2': (2, '2A1'),   # Thứ Tư
    '3': (2, '3A1'),   # Thứ Tư
    '4': (4, '4C1'),   # Thứ Sáu
    '5': (1, '5C1'),   # Thứ Ba
    '6': (4, '6A1'),   # Thứ Sáu
    '7': (1, '7A1'),   # Thứ Ba
    '8': (4, '8A1'),   # Thứ Sáu
}

KHBD_BASE = r'd:\UNIGO\KHBD_Robotics'

def compute_dates(tuan_so, day_of_week):
    week_start = TUAN_01_START + timedelta(weeks=tuan_so - 1)
    ngay_day = week_start + timedelta(days=day_of_week)
    ngay_soan = week_start - timedelta(days=2)  # Thứ 7 tuần trước
    return ngay_soan.strftime('%d/%m/%Y'), ngay_day.strftime('%d/%m/%Y')

def make_paragraph(text, bold=True):
    """Create a w:p element with a single run, TNR 13pt."""
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '26')  # 13pt
    rPr.append(sz)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    return p

def set_run_font(run):
    """Set font to TNR 13pt on a run."""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), 'Times New Roman')

# ─── FIX THCS (Lớp 6, 7, 8) ──────────────────────────────────────────────
def fix_robotics_thcs(filepath, lop_key, tuan_so):
    day_of_week, ten_lop = ROBOTICS_SCHEDULE[lop_key]
    ngay_soan, ngay_day = compute_dates(tuan_so, day_of_week)
    
    doc = Document(filepath)
    modified = False
    
    # Fix Table[0] Row[1] Cell[1]
    if doc.tables:
        tbl = doc.tables[0]
        rows = tbl._tbl.findall(qn('w:tr'))
        if len(rows) >= 2:
            cells_r1 = rows[1].findall(qn('w:tc'))
            if len(cells_r1) >= 2:
                tc = cells_r1[1]
                # Remove all old paragraphs
                for old_p in tc.findall(qn('w:p')):
                    tc.remove(old_p)
                # Create 2 clean paragraphs only
                tc.append(make_paragraph(f'Ngày soạn: {ngay_soan}   Ngày dạy: {ngay_day}'))
                tc.append(make_paragraph(f'Lớp: {ten_lop}'))
                modified = True
    
    # Fix paragraph class names
    for p in doc.paragraphs:
        if p.text.strip().startswith('MÔN HỌC:') and f'ROBOTICS {lop_key}' in p.text:
            p.text = p.text.replace(f'ROBOTICS {lop_key}', f'ROBOTICS {ten_lop}')
            modified = True
        elif 'Lớp:' in p.text and f'Lớp: {lop_key}' in p.text:
            for run in p.runs:
                if f'Lớp: {lop_key}' in run.text:
                    run.text = run.text.replace(f'Lớp: {lop_key}', f'Lớp: {ten_lop}')
                    modified = True
    
    if modified:
        try:
            doc.save(filepath)
            print(f'  ✅ Fixed: {os.path.basename(filepath)}')
        except PermissionError:
            alt = filepath.replace('.docx', '_fixed.docx')
            doc.save(alt)
            print(f'  ⚠️ Saved as: {os.path.basename(alt)} (original locked)')
    else:
        print(f'  ⏩ No changes: {os.path.basename(filepath)}')

# ─── FIX TH (Lớp 1, 2, 3, 4, 5) ──────────────────────────────────────────
def fix_robotics_th(filepath, lop_key, tuan_so):
    day_of_week, ten_lop = ROBOTICS_SCHEDULE[lop_key]
    ngay_soan, ngay_day = compute_dates(tuan_so, day_of_week)
    
    doc = Document(filepath)
    modified = False
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if ('Thứ' in text and 'ngày' in text) or text.startswith('Ngày soạn:') or text.startswith('TUẦN:'):
            p.clear()
            r = p.add_run(f'Ngày soạn: {ngay_soan}   Ngày dạy: {ngay_day}')
            r.italic = True
            set_run_font(r)
            modified = True
            break
            
    # Sửa tên lớp ở tiêu đề nếu có
    for p in doc.paragraphs[:5]:
        if f'LỚP {lop_key}' in p.text and lop_key in ('5', '4', '3', '2', '1'):
            if f'LỚP {ten_lop}' not in p.text:
                p.text = p.text.replace(f'LỚP {lop_key}', f'LỚP {ten_lop}')
                modified = True
                
    if modified:
        try:
            doc.save(filepath)
            print(f'  ✅ Fixed: {os.path.basename(filepath)}')
        except PermissionError:
            alt = filepath.replace('.docx', '_fixed.docx')
            doc.save(alt)
            print(f'  ⚠️ Saved as: {os.path.basename(alt)} (original locked)')
    else:
        print(f'  ⏩ No changes: {os.path.basename(filepath)}')

# ─── MAIN ──────────────────────────────────────────────────────────────────
def main():
    # 1. Dọn dẹp các file trùng lặp ở Tuần_01
    print('🧹 Xóa file thừa/trùng lặp ở Tuần_01...')
    for lop_key in ['1', '2', '3', '4', '5', '6', '7', '8']:
        dup_path = os.path.join(KHBD_BASE, f'Lớp_{lop_key}', 'Tuần_01', f'KHBD_Robotics_Lớp_{lop_key}_Tiet00_Tiet_0_inh_huong_mon_hoc.docx')
        if os.path.exists(dup_path):
            os.remove(dup_path)
            print(f'  🗑️ Đã xóa file trùng: {os.path.basename(dup_path)}')

    # 2. Xử lý THCS (Lớp 6, 7, 8)
    for lop_key in ['6', '7', '8']:
        lop_folder = f'Lớp_{lop_key}'
        print(f'\n🟠 Fixing THCS Robotics {lop_folder}...')
        dir_lop = os.path.join(KHBD_BASE, lop_folder)
        if not os.path.isdir(dir_lop):
            continue
        for tuan_folder in sorted(os.listdir(dir_lop)):
            dir_tuan = os.path.join(dir_lop, tuan_folder)
            if not os.path.isdir(dir_tuan) or not tuan_folder.startswith('Tuần_'):
                continue
            m = re.search(r'\d+', tuan_folder)
            if not m:
                continue
            tuan_so = int(m.group(0))
            for f in sorted(os.listdir(dir_tuan)):
                if f.endswith('.docx') and 'KHBD' in f:
                    fix_robotics_thcs(os.path.join(dir_tuan, f), lop_key, tuan_so)

    # 3. Xử lý TH (Lớp 1, 2, 3, 4, 5)
    for lop_key in ['1', '2', '3', '4', '5']:
        lop_folder = f'Lớp_{lop_key}'
        print(f'\n🔵 Fixing TH Robotics {lop_folder}...')
        dir_lop = os.path.join(KHBD_BASE, lop_folder)
        if not os.path.isdir(dir_lop):
            continue
        for tuan_folder in sorted(os.listdir(dir_lop)):
            dir_tuan = os.path.join(dir_lop, tuan_folder)
            if not os.path.isdir(dir_tuan) or not tuan_folder.startswith('Tuần_'):
                continue
            m = re.search(r'\d+', tuan_folder)
            if not m:
                continue
            tuan_so = int(m.group(0))
            for f in sorted(os.listdir(dir_tuan)):
                if f.endswith('.docx') and 'KHBD' in f:
                    fix_robotics_th(os.path.join(dir_tuan, f), lop_key, tuan_so)

    print('\n🎉 Hoàn thành sửa toàn bộ KHBD Robotics!')

if __name__ == '__main__':
    main()
