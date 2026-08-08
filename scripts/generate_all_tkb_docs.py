import os
import sys
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

sys.stdout.reconfigure(encoding='utf-8')

OUT_DIR_WORD = r'D:\UNIGO\Thời khóa biểu giáo viên'
OUT_DIR_MAU = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Thời khóa biểu'
TKB_EXCEL_SOURCE = r'D:\UNIGO\TKB toàn trường CHECK - chuẩn.xlsx'

os.makedirs(OUT_DIR_WORD, exist_ok=True)
os.makedirs(OUT_DIR_MAU, exist_ok=True)

KNOWN_TEACHERS = [
    'Nguyên', 'Hà', 'Ánh NV', 'Tân', 'Nhung', 'Chánh', 'Châu', 'Thanh', 
    'Hồng', 'Ngọc Ánh', 'Nhi', 'Linh TH', 'Xuân', 'Ngọc', 'Ly', 'Phương', 
    'Trang ÂN', 'Phương Anh', 'Trang', 'Bích', 'Chi', 'Quỳnh', 'Trà', 
    'Huyền', 'Aziz', 'Tay', 'Hiền', 'MN1', 'MN2', 'Linh', 'Ngân', 'Phượng', 
    'Dung', 'Kel'
]

TEACHER_DETAILS = {
    'Nguyên': {'full_name': 'Đậu Đình Nguyên', 'subject': 'Tin học & Robotics'},
    'Hà': {'full_name': 'Nguyễn Thị Hà', 'subject': 'Ngữ văn & LS-ĐL-GDCD'},
    'Ánh NV': {'full_name': 'Nguyễn Thị Ánh (Ánh NV)', 'subject': 'Ngữ văn & GDĐP-Đọc sách'},
    'Tân': {'full_name': 'Cô Tân', 'subject': 'Toán'},
    'Nhung': {'full_name': 'Nguyễn Thị Nhung', 'subject': 'Khoa học tự nhiên'},
    'Chánh': {'full_name': 'Lê Thị Chánh', 'subject': 'Ngữ văn & GVVN'},
    'Thanh': {'full_name': 'Trần Thị Thanh', 'subject': 'Tiếng Trung'},
    'Châu': {'full_name': 'Bùi Thị Châu', 'subject': 'Tiếng Trung'},
    'Hồng': {'full_name': 'Trần Thị Hồng', 'subject': 'Giáo dục thể chất'},
    'Ngọc Ánh': {'full_name': 'Nguyễn Thị Ngọc Ánh', 'subject': 'Giáo viên Chủ nhiệm / Tiểu học'},
    'Nhi': {'full_name': 'Nguyễn Thị Nhi', 'subject': 'Giáo viên Chủ nhiệm / Tiểu học'},
    'Linh TH': {'full_name': 'Cô Linh (TH)', 'subject': 'Giáo viên Chủ nhiệm / Tiểu học'},
    'Xuân': {'full_name': 'Nguyễn Thị Xuân', 'subject': 'Giáo viên Chủ nhiệm / Tiểu học'},
    'Ngọc': {'full_name': 'Nguyễn Thị Ngọc', 'subject': 'Giáo viên Chủ nhiệm / Tiểu học'},
    'Ly': {'full_name': 'Nguyễn Thị Ly', 'subject': 'Lịch sử - Địa lý & Tiểu học'},
    'Phương': {'full_name': 'Trương Thị Phương', 'subject': 'Mĩ thuật'},
    'Trang ÂN': {'full_name': 'Cô Trang (Âm nhạc)', 'subject': 'Âm nhạc'},
    'Phương Anh': {'full_name': 'Đào Phương Anh', 'subject': 'Toán & Hoạt động trải nghiệm'},
    'Trang': {'full_name': 'Trịnh Kiều Trang', 'subject': 'Giáo viên Việt Nam / Tiền Tiểu học'},
    'Bích': {'full_name': 'Cô Bích', 'subject': 'Science & GVVN'},
    'Chi': {'full_name': 'Cô Chi', 'subject': 'Toán & Hoạt động trải nghiệm'},
    'Quỳnh': {'full_name': 'Cô Quỳnh', 'subject': 'Science & GVVN'},
    'Trà': {'full_name': 'Cô Trà', 'subject': 'Science & GVVN'},
    'Huyền': {'full_name': 'Cô Huyền', 'subject': 'GVVN & GVNN'},
    'Aziz': {'full_name': 'Thầy Aziz', 'subject': 'Giáo viên Nước ngoài (GVNN)'},
    'Tay': {'full_name': 'Thầy Tay', 'subject': 'Giáo viên Nước ngoài (GVNN)'},
    'Hiền': {'full_name': 'Cô Hiền', 'subject': 'Tiền Tiểu học TTH 1'},
    'MN1': {'full_name': 'Cô MN1', 'subject': 'Tiền Tiểu học TTH 2'},
    'MN2': {'full_name': 'Cô MN2', 'subject': 'Tiền Tiểu học TT3'},
    'Linh': {'full_name': 'Cô Linh', 'subject': 'Mầm Mẫu'},
    'Ngân': {'full_name': 'Cô Ngân', 'subject': 'Giáo viên Nước ngoài (GVNN)'},
    'Phượng': {'full_name': 'Cô Phượng', 'subject': 'Giáo viên Nước ngoài (GVNN)'},
    'Dung': {'full_name': 'Cô Dung', 'subject': 'Giáo viên Nước ngoài (GVNN)'},
    'Kel': {'full_name': 'Thầy Kel', 'subject': 'Giáo viên Nước ngoài (GVNN)'},
}

def extract_teachers_from_str(teachers_str):
    found = []
    t_str = teachers_str.replace('LinhTH', 'Linh TH')
    for k in sorted(KNOWN_TEACHERS, key=len, reverse=True):
        if k in t_str:
            found.append(k)
            t_str = t_str.replace(k, '')
    if not found:
        found = [t.strip() for t in teachers_str.split(',') if t.strip()]
    return found

def load_teacher_data(xlsx_path):
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb['TKB_LOP_SC']

    class_map = {}
    for c in range(1, ws.max_column + 1):
        v4 = ws.cell(4, c).value
        v5 = ws.cell(5, c).value
        if v4 and str(v4).strip() not in ['THỨ', 'TIẾT']:
            curr_class = str(v4).strip()
            buoi = str(v5).strip() if v5 else ''
            class_map[c] = (curr_class, buoi)
        elif v5 and (c-1) in class_map:
            curr_class = class_map[c-1][0]
            buoi = str(v5).strip()
            class_map[c] = (curr_class, buoi)

    thu_map = {6: '2', 11: '3', 16: '4', 21: '5', 26: '6'}

    teacher_data = {}

    for start_r, thu_val in thu_map.items():
        for offset in range(5):
            r = start_r + offset
            tiet_val = int(offset + 1)
            for c in class_map:
                cls_name, buoi = class_map[c]
                cell_val = ws.cell(r, c).value
                if cell_val and str(cell_val).strip() not in ['0', 'None', '']:
                    txt = str(cell_val).strip()
                    if '-' in txt:
                        parts = txt.split('-', 1)
                        subj = parts[0].strip()
                        teachers_part = parts[1].strip()
                        gvs = extract_teachers_from_str(teachers_part)
                        for gv in gvs:
                            details = TEACHER_DETAILS.get(gv, {'full_name': f'Giáo viên {gv}', 'subject': subj})
                            teacher_data.setdefault(gv, {
                                'full_name': details['full_name'],
                                'subject': details['subject'],
                                'slots': []
                            })
                            teacher_data[gv]['slots'].append((thu_val, buoi, tiet_val, subj, cls_name))
    return teacher_data

def set_table_borders(table):
    table.style = 'Table Grid'
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        for child in list(tblPr[0]):
            if child.tag.endswith('tblBorders'):
                tblPr[0].remove(child)
        borders_xml = parse_xml(
            r'<w:tblBorders %s>'
            r'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        tblPr[0].append(borders_xml)

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_docx_tkb(gv_key, data, save_paths):
    doc = docx.Document()
    
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_header.add_run("TRƯỜNG TIỂU HỌC VÀ THCS UNIGO\n")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(13)
    r1.bold = True
    
    r2 = p_header.add_run("THỜI KHOÁ BIỂU CÁ NHÂN GIÁO VIÊN\n")
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(15)
    r2.bold = True
    
    r3 = p_header.add_run("Năm học 2026 - 2027 (Theo TKB toàn trường CHECK - chuẩn)\n")
    r3.font.name = 'Times New Roman'
    r3.font.size = Pt(12)
    r3.italic = True
    
    p_info = doc.add_paragraph()
    r_info = p_info.add_run(
        f"Họ và tên giáo viên: {data['full_name']}\n"
        f"Môn học đảm nhận: {data['subject']}\n"
        f"Tổng số tiết/tuần: {len(data['slots'])} tiết"
    )
    r_info.font.name = 'Times New Roman'
    r_info.font.size = Pt(13)
    r_info.bold = True
    
    table = doc.add_table(rows=13, cols=6)
    set_table_borders(table)
    
    headers = ["Buổi / Tiết", "Thứ Hai", "Thứ Ba", "Thứ Tư", "Thứ Năm", "Thứ Sáu"]
    
    for c_idx in range(6):
        cell = table.rows[0].cells[c_idx]
        cell.text = headers[c_idx]
        set_cell_background(cell, "D9E1F2")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)
            
    cell_sang = table.rows[1].cells[0]
    for c_idx in range(1, 6):
        cell_sang.merge(table.rows[1].cells[c_idx])
    cell_sang.text = "BUỔI SÁNG"
    set_cell_background(cell_sang, "F2F2F2")
    p = cell_sang.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)

    cell_chieu = table.rows[7].cells[0]
    for c_idx in range(1, 6):
        cell_chieu.merge(table.rows[7].cells[c_idx])
    cell_chieu.text = "BUỔI CHIỀU"
    set_cell_background(cell_chieu, "F2F2F2")
    p = cell_chieu.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)

    for p_num in range(1, 6):
        r_sang = table.rows[p_num + 1]
        r_sang.cells[0].text = f"Tiết {p_num}"
        p = r_sang.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.bold = True
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(13)

        r_chieu = table.rows[p_num + 7]
        r_chieu.cells[0].text = f"Tiết {p_num}"
        p = r_chieu.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.bold = True
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(13)

    day_to_col = {'2': 1, '3': 2, '4': 3, '5': 4, '6': 5}

    for slot in data['slots']:
        day, sess, p_num, subject, cls_name = slot
        col_i = day_to_col[day]
        row_i = (p_num + 1) if sess == 'Sáng' else (p_num + 7)
        
        cell = table.rows[row_i].cells[col_i]
        curr_text = cell.text.strip()
        new_text = f"{subject} ({cls_name})"
        cell.text = f"{curr_text}, {new_text}" if curr_text else new_text
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)

    for path in save_paths:
        try:
            doc.save(path)
            print(f"Saved DOCX: {path}")
        except PermissionError:
            print(f"Permission error saving {path}")

def create_excel_tkb(gv_key, data, save_path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "TKB Cá Nhân"

    thin_border = Border(
        left=Side(style='thin', color='000000'),
        right=Side(style='thin', color='000000'),
        top=Side(style='thin', color='000000'),
        bottom=Side(style='thin', color='000000')
    )

    ws.merge_cells('A1:F1')
    ws['A1'] = "TRƯỜNG TIỂU HỌC VÀ THCS UNIGO"
    ws['A1'].font = Font(name='Times New Roman', size=13, bold=True)
    ws['A1'].alignment = Alignment(horizontal='center')

    ws.merge_cells('A2:F2')
    ws['A2'] = "THỜI KHOÁ BIỂU CÁ NHÂN GIÁO VIÊN"
    ws['A2'].font = Font(name='Times New Roman', size=15, bold=True)
    ws['A2'].alignment = Alignment(horizontal='center')

    ws.merge_cells('A3:F3')
    ws['A3'] = f"Giáo viên: {data['full_name']} | Bộ môn: {data['subject']} | Tổng số: {len(data['slots'])} tiết/tuần (Theo TKB toàn trường CHECK - chuẩn)"
    ws['A3'].font = Font(name='Times New Roman', size=12, italic=True)
    ws['A3'].alignment = Alignment(horizontal='center')

    headers = ["Buổi / Tiết", "Thứ Hai", "Thứ Ba", "Thứ Tư", "Thứ Năm", "Thứ Sáu"]
    for c_i, h in enumerate(headers, 1):
        cell = ws.cell(5, c_i, h)
        cell.font = Font(name='Times New Roman', size=13, bold=True)
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
        cell.border = thin_border

    ws.merge_cells('A6:F6')
    ws['A6'] = "BUỔI SÁNG"
    ws['A6'].font = Font(name='Times New Roman', size=13, bold=True)
    ws['A6'].alignment = Alignment(horizontal='center')
    ws['A6'].fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")

    for c_i in range(1, 7):
        ws.cell(6, c_i).border = thin_border

    for p_num in range(1, 6):
        r_idx = 6 + p_num
        ws.cell(r_idx, 1, f"Tiết {p_num}").font = Font(name='Times New Roman', size=13, bold=True)
        ws.cell(r_idx, 1).alignment = Alignment(horizontal='center', vertical='center')
        ws.cell(r_idx, 1).border = thin_border
        for col_i in range(2, 7):
            ws.cell(r_idx, col_i).border = thin_border
            ws.cell(r_idx, col_i).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    ws.merge_cells('A12:F12')
    ws['A12'] = "BUỔI CHIỀU"
    ws['A12'].font = Font(name='Times New Roman', size=13, bold=True)
    ws['A12'].alignment = Alignment(horizontal='center')
    ws['A12'].fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")

    for c_i in range(1, 7):
        ws.cell(12, c_i).border = thin_border

    for p_num in range(1, 6):
        r_idx = 12 + p_num
        ws.cell(r_idx, 1, f"Tiết {p_num}").font = Font(name='Times New Roman', size=13, bold=True)
        ws.cell(r_idx, 1).alignment = Alignment(horizontal='center', vertical='center')
        ws.cell(r_idx, 1).border = thin_border
        for col_i in range(2, 7):
            ws.cell(r_idx, col_i).border = thin_border
            ws.cell(r_idx, col_i).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    day_to_col = {'2': 2, '3': 3, '4': 4, '5': 5, '6': 6}

    for slot in data['slots']:
        day, sess, p_num, subject, cls_name = slot
        col_i = day_to_col[day]
        row_i = (6 + p_num) if sess == 'Sáng' else (12 + p_num)
        
        curr_val = ws.cell(row_i, col_i).value or ''
        new_val = f"{subject} ({cls_name})"
        ws.cell(row_i, col_i, f"{curr_val}, {new_val}" if curr_val else new_val)

    ws.column_dimensions['A'].width = 16
    for c_let in ['B', 'C', 'D', 'E', 'F']:
        ws.column_dimensions[c_let].width = 22

    wb.save(save_path)

def main():
    print("Loading Timetable data for ALL teachers from 'TKB toàn trường CHECK - chuẩn.xlsx'...")
    teacher_data = load_teacher_data(TKB_EXCEL_SOURCE)
    
    for gv_key, data in teacher_data.items():
        full_name = data['full_name']
        fname_word1 = os.path.join(OUT_DIR_WORD, f"Thời khóa biểu - {full_name}.docx")
        fname_word2 = os.path.join(OUT_DIR_MAU, f"Thời khóa biểu - {full_name}.docx")
        fname_excel = os.path.join(OUT_DIR_WORD, f"Thời khóa biểu - {full_name}.xlsx")
        
        create_docx_tkb(gv_key, data, [fname_word1, fname_word2])
        create_excel_tkb(gv_key, data, fname_excel)

    print("All timetable files updated successfully.")

if __name__ == '__main__':
    main()
