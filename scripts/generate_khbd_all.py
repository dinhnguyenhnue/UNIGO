"""
generate_khbd_all.py — Tạo lại toàn bộ KHBD Bài 1 từ Tiền tiểu học đến Lớp 8
Dựa trên luật KHBD_TIEU_HOC.md và KHBD_THCS.md đã được phân tích từ template thực tế.
"""
import os, sys, shutil, copy
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Pt, Emu, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ─── TEMPLATE PATHS ────────────────────────────────────────────────────────
TPL_TH   = r'd:\UNIGO\Hệ thống mẫu văn bản\Khung  giáo án Unigo 2026-2027 Thang 7.2026.docx'
TPL_THCS = r'd:\UNIGO\Hệ thống mẫu văn bản\PL4-Khung kế hoạch bài dạy (THCS).docx'
OUT_BASE = r'd:\UNIGO\KHBD_Tin_học'

# ─── FONT HELPER ───────────────────────────────────────────────────────────
def afont(run, bold=None, italic=None, size_pt=13, color=None):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    if bold is not None: run.bold = bold
    if italic is not None: run.italic = italic
    if color: run.font.color.rgb = RGBColor(*color)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), 'Times New Roman')

def set_line_spacing(para, ratio):
    """Set line spacing by ratio (1.5 → 360, 1.15 → 276)"""
    pPr = para._p.get_or_add_pPr()
    sp = pPr.find(qn('w:spacing'))
    if sp is None:
        sp = OxmlElement('w:spacing')
        pPr.append(sp)
    sp.set(qn('w:line'), str(int(240 * ratio)))
    sp.set(qn('w:lineRule'), 'auto')

def add_para(doc, text='', bold=False, italic=False, align=None,
             indent_first=None, indent_left=None, line_ratio=None,
             sp_before=None, sp_after=None, size_pt=13):
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text)
        afont(r, bold=bold, italic=italic, size_pt=size_pt)
    if align is not None: p.alignment = align
    pf = p.paragraph_format
    if indent_first is not None: pf.first_line_indent = Emu(indent_first)
    if indent_left is not None: pf.left_indent = Emu(indent_left)
    if sp_before is not None: pf.space_before = Pt(sp_before)
    if sp_after is not None: pf.space_after = Pt(sp_after)
    if line_ratio is not None: set_line_spacing(p, line_ratio)
    return p

# ─── TABLE BORDERS ─────────────────────────────────────────────────────────
def set_table_borders(table):
    tblPr = table._tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:color'), '000000')
        b.set(qn('w:space'), '0')
        tblBorders.append(b)
    tblPr.append(tblBorders)

def fill_cell(cell, text, bold=False, italic=False, line_ratio=1.5, align=None):
    for p in cell.paragraphs:
        for r in p.runs: r.text = ''
    if not cell.paragraphs:
        cell.add_paragraph()
    p = cell.paragraphs[0]
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i > 0: p = cell.add_paragraph()
        if line:
            r = p.add_run(line)
            afont(r, bold=bold, italic=italic)
    if line_ratio: set_line_spacing(cell.paragraphs[0], line_ratio)
    if align: cell.paragraphs[0].alignment = align

# ─── SAVE HELPER ───────────────────────────────────────────────────────────
def save_doc(doc, path):
    os.makedirs(os.path.dirname(path), exist_ok=True)
    try:
        doc.save(path)
        return path
    except PermissionError:
        alt = path.replace('.docx', '_new.docx')
        doc.save(alt)
        return alt

# ─── CLEAN BODY ────────────────────────────────────────────────────────────
def clean_body(doc):
    """Xóa toàn bộ body elements, giữ sectPr"""
    body = doc.element.body
    for child in list(body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'sectPr':
            body.remove(child)

# ════════════════════════════════════════════════════════════════════════════
# PHẦN A: TẠO KHBD TIỂU HỌC
# Theo luật KHBD_TIEU_HOC.md:
# - Font: TNR 13pt, Line spacing: 1.5, Indent: 457200 EMU
# - Mục tiêu: Phẩm chất TRƯỚC → Năng lực SAU
# - Bảng 2 cột GV/HS với hàng gộp gridSpan=2 cho tiêu đề HĐ
# ════════════════════════════════════════════════════════════════════════════

def add_merged_row_th(table, text_lines, bold=True):
    """Thêm hàng tiêu đề gộp 2 cột cho TH"""
    row = table.add_row()
    tc0 = row.cells[0]._tc
    tc1 = row.cells[1]._tc
    # Set gridSpan=2
    tcPr = tc0.get_or_add_tcPr()
    gs = OxmlElement('w:gridSpan')
    gs.set(qn('w:val'), '2')
    tcPr.append(gs)
    # Remove 2nd cell
    row._tr.remove(tc1)
    # Clear existing paragraph in tc0 and fill text
    existing_p = tc0.find(qn('w:p'))
    if existing_p is not None:
        tc0.remove(existing_p)
    for i, line in enumerate(text_lines):
        new_p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        # line spacing 1.5
        sp_el = OxmlElement('w:spacing')
        sp_el.set(qn('w:line'), '360')
        sp_el.set(qn('w:lineRule'), 'auto')
        pPr.append(sp_el)
        new_p.append(pPr)
        if line:
            new_r = OxmlElement('w:r')
            rPr = OxmlElement('w:rPr')
            if bold:
                b_el = OxmlElement('w:b')
                rPr.append(b_el)
            # Font
            rFonts = OxmlElement('w:rFonts')
            for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
                rFonts.set(qn(attr), 'Times New Roman')
            rPr.append(rFonts)
            sz = OxmlElement('w:sz')
            sz.set(qn('w:val'), '26')  # 13pt = 26 half-points
            rPr.append(sz)
            szCs = OxmlElement('w:szCs')
            szCs.set(qn('w:val'), '26')
            rPr.append(szCs)
            new_r.append(rPr)
            t_el = OxmlElement('w:t')
            t_el.text = line
            t_el.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            new_r.append(t_el)
            new_p.append(new_r)
        tc0.append(new_p)
    return row

def add_content_row_th(table, gv_lines, hs_lines):
    """Thêm hàng nội dung GV | HS"""
    row = table.add_row()
    fill_cell(row.cells[0], '\n'.join(gv_lines), line_ratio=1.5)
    fill_cell(row.cells[1], '\n'.join(hs_lines), line_ratio=1.5)
    return row

def create_khbd_th(grade_name, grade_folder, bai_so, ten_bai, chu_diem,
                   tiet_ppct, ngay_day,
                   pham_chat_items, nang_luc_mon_items, nang_luc_chung_items,
                   do_dung_gv, do_dung_hs,
                   phuong_phap, ki_thuat,
                   activities, nang_luc_so_items=None):
    """
    Tạo KHBD Tiểu học hoàn chỉnh.
    activities: list of dicts:
      {
        'ten': 'Tên hoạt động (...phút)',
        'muc_tieu': 'Mục tiêu',
        'sub': [  # Sub-hoạt động trong cùng nhóm (cho HĐ2, HĐ3)
          {'ten': '2.1. Tên HĐ', 'gv': [...], 'hs': [...]},
          ...
        ],
        'gv': ['GV làm...'],  # Nếu không có sub
        'hs': ['HS làm...'],
      }
    """
    doc = Document(TPL_TH)
    clean_body(doc)

    # ── Phần đầu ─────────────────────────────────────────────────────────
    # P[0]: Ngày
    p = add_para(doc, f'Thứ …… ngày {ngay_day}',
                 align=WD_ALIGN_PARAGRAPH.CENTER, line_ratio=1.5)
    # Italic cho phần thứ
    p.clear()
    r0 = p.add_run('Thứ …… ngày ')
    afont(r0, italic=True)
    r1 = p.add_run(ngay_day)
    afont(r1, italic=True)
    set_line_spacing(p, 1.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # P[1]: GV
    p1 = add_para(doc, 'Họ và tên Giáo viên: Đậu Đình Nguyên',
                  align=WD_ALIGN_PARAGRAPH.CENTER, line_ratio=1.5)

    # P[2]: Tên môn
    add_para(doc, 'KẾ HOẠCH DẠY HỌC MÔN TIN HỌC',
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, line_ratio=1.5)

    # P[3]: Chủ điểm
    p3 = add_para(doc, line_ratio=1.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    r_cd_label = p3.add_run('CHỦ ĐIỂM: ')
    afont(r_cd_label, bold=True)
    r_cd_val = p3.add_run(chu_diem)
    afont(r_cd_val, bold=True)

    # P[4]: Bài + Tiết PPCT
    p4 = add_para(doc, line_ratio=1.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    r_bai = p4.add_run(f'BÀI: {ten_bai}  ')
    afont(r_bai, bold=True)
    r_tiet_l = p4.add_run('(Tiết: ')
    afont(r_tiet_l)
    r_tiet_v = p4.add_run(str(tiet_ppct))
    afont(r_tiet_v, italic=True)
    r_tiet_r = p4.add_run(')')
    afont(r_tiet_r)

    # P[5]: Trống
    add_para(doc, '')

    # ── I. YÊU CẦU CẦN ĐẠT ──────────────────────────────────────────────
    add_para(doc, 'I. YÊU CẦU CẦN ĐẠT:', bold=True, line_ratio=1.5,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_para(doc, '- Sau tiết học, học sinh sẽ:', bold=True,
             indent_first=457200, line_ratio=1.5,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # 1. Phẩm chất (TRƯỚC)
    add_para(doc, '1. Phát triển phẩm chất', bold=True,
             indent_first=457200, line_ratio=1.5,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    for pc in pham_chat_items:
        add_para(doc, pc, indent_first=450215, line_ratio=1.33,
                 align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # 2. Năng lực (SAU)
    add_para(doc, '2. Phát triển năng lực', bold=True,
             indent_first=457200, line_ratio=1.5,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_para(doc, '2.1. Năng lực môn học:', indent_first=457200,
             line_ratio=1.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    for nl in nang_luc_mon_items:
        add_para(doc, nl, indent_first=450215, line_ratio=1.33,
                 align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    # 2.1b. Năng lực số (tách riêng, hanging indent)
    if nang_luc_so_items:
        add_para(doc, '- Năng lực số:', bold=True,
                 indent_first=457200, sp_after=0, line_ratio=1.33,
                 align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        for nl in nang_luc_so_items:
            p = add_para(doc, f'  {nl}', sp_after=0, line_ratio=1.33,
                         indent_left=720000)
            p.paragraph_format.first_line_indent = Emu(-262000)
    add_para(doc, '2.2. Năng lực chung và đặc thù:', indent_first=457200,
             line_ratio=1.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    for nl in nang_luc_chung_items:
        add_para(doc, nl, indent_first=450215, line_ratio=1.33,
                 align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ── II. ĐỒ DÙNG ──────────────────────────────────────────────────────
    add_para(doc, 'II. ĐỒ DÙNG DẠY HỌC :', bold=True, line_ratio=None)
    add_para(doc, f'\t1. Giáo viên: {do_dung_gv}', line_ratio=1.5,
             indent_left=228600)
    add_para(doc, f'2. Học sinh: {do_dung_hs}', line_ratio=1.5,
             indent_first=457200, align=WD_ALIGN_PARAGRAPH.JUSTIFY)

    # ── III. PHƯƠNG PHÁP ─────────────────────────────────────────────────
    add_para(doc, 'III. PHƯƠNG PHÁP, KĨ THUẬT DẠY HỌC', bold=True,
             line_ratio=1.5, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_para(doc, f'- Phương pháp: {phuong_phap}',
             indent_first=457200, line_ratio=None,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    add_para(doc, f'\t- Kĩ thuật: {ki_thuat}',
             line_ratio=None)

    # ── IV. CÁC HOẠT ĐỘNG DẠY-HỌC ────────────────────────────────────────
    add_para(doc, 'IV. CÁC HOẠT ĐỘNG DẠY - HỌC CHỦ YẾU: ', bold=True,
             line_ratio=1.5)

    # Tạo bảng 2 cột
    table = doc.add_table(rows=1, cols=2)
    set_table_borders(table)

    # Header row
    hdr_row = table.rows[0]
    fill_cell(hdr_row.cells[0], 'Hoạt động của GV', bold=True,
              line_ratio=1.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell(hdr_row.cells[1], 'Hoạt động của HS', bold=True,
              line_ratio=1.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Điền các hoạt động
    for act in activities:
        if act.get('sub'):
            # Tiêu đề nhóm HĐ (gộp 2 cột)
            header_lines = [act['ten']]
            if act.get('muc_tieu'):
                header_lines.append(f'*Mục tiêu: {act["muc_tieu"]}')
            add_merged_row_th(table, header_lines)
            for sub in act['sub']:
                sub_header = [sub['ten']]
                if sub.get('muc_tieu'):
                    sub_header.append(f'*Mục tiêu: {sub["muc_tieu"]}')
                add_merged_row_th(table, sub_header)
                add_content_row_th(table, sub.get('gv', ['']), sub.get('hs', ['']))
        else:
            # Hàng đơn gộp tiêu đề
            header_lines = [act['ten']]
            if act.get('muc_tieu'):
                header_lines.append(f'*Mục tiêu: {act["muc_tieu"]}')
            if len(activities) == 1 or not act.get('sub'):
                # Nếu là HĐ cuối (Vận dụng), không gộp mà để 2 cột
                add_content_row_th(table,
                    [act['ten']] + act.get('gv', ['']),
                    act.get('hs', ['...............']))
            else:
                add_merged_row_th(table, header_lines)
                add_content_row_th(table, act.get('gv', ['']), act.get('hs', ['']))

    # ── V. ĐIỀU CHỈNH BỔ SUNG — copy từ template gốc (chỉ P[33]-P[35], BỎ P[36]+) ──
    doc_tpl = Document(TPL_TH)
    tpl_paras = doc_tpl.paragraphs
    # P[33] đến P[35]: "V. ĐIỀU CHỈNH...", ghi chú GV, hàng dấu chấm. Bỏ P[36]+ (* Lưu ý...)
    tail_start = 33
    for i in range(tail_start, 36):
        new_p = copy.deepcopy(tpl_paras[i]._p)
        # Chèn trước sectPr
        sect_pr = doc.element.body.find(qn('w:sectPr'))
        if sect_pr is not None:
            doc.element.body.insert(list(doc.element.body).index(sect_pr), new_p)
        else:
            doc.element.body.append(new_p)

    # ── Save ──────────────────────────────────────────────────────────────
    if bai_so == 0:
        out_dir = os.path.join(OUT_BASE, grade_folder, 'Tiết_00')
        clean_name = ten_bai.replace(":", "").replace("/", "_").replace("\\", "_").replace(" ", "_")[:40]
        filename = f'KHBD_Tin_hoc_{grade_folder}_Tiet00_{clean_name}.docx'
    else:
        out_dir = os.path.join(OUT_BASE, grade_folder, f'Bài_{bai_so:02d}')
        clean_name = ten_bai.replace(":", "").replace("/", "_").replace("\\", "_").replace(" ", "_")[:40]
        filename = f'KHBD_Tin_hoc_{grade_folder}_Bai{bai_so:02d}_{clean_name}.docx'
    out_path = os.path.join(out_dir, filename)
    saved = save_doc(doc, out_path)

    # ── Verify ────────────────────────────────────────────────────────────
    verify = Document(saved)
    hdr_drawings = sum(1 for r in verify.sections[0].header.paragraphs[0].runs
                       if r._r.findall(qn('w:drawing')))
    ftr_paras = len(verify.sections[0].footer.paragraphs)
    empty_start = sum(1 for p in verify.paragraphs[:10] if not p.text.strip())
    print(f'  ✅ Saved: {saved}')
    print(f'     Header drawings: {hdr_drawings} | Footer paras: {ftr_paras} | Empty start: {empty_start}')
    return saved


# ════════════════════════════════════════════════════════════════════════════
# PHẦN B: TẠO KHBD THCS
# Theo luật KHBD_THCS.md (Phương án B):
# - Font: TNR 13pt, Line spacing: 1.15
# - Mục tiêu: Kiến thức → Năng lực (chung/đặc thù/số) → Phẩm chất
# - Tiến trình: mỗi HĐ có a)b)c)d) + BẢNG 3 CỘT (Bước | HĐ GV | HĐ HS)
# - GIỮ NGUYÊN Table[0] thông tin và Table[2] ký tên từ template
# ════════════════════════════════════════════════════════════════════════════

def add_activity_table_b(doc, buoc_data, is_last=False):
    """
    Tạo bảng 3 cột cho 1 hoạt động (Phương án B).
    buoc_data: list of 4 tuples [(gv_text, hs_text), ...] cho Bước 1-4
    """
    buoc_labels_default = [
        'Bước 1:\nChuyển giao\nnhiệm vụ học tập',
        'Bước 2:\nHọc sinh tiếp nhận\nnhiệm vụ học tập',
        'Bước 3:\nBáo cáo kết quả\nhoạt động',
        'Bước 4:\nĐánh giá kết quả\nthực hiện nhiệm vụ',
    ]
    buoc_labels_last = [
        'Bước 1:\nChuyển giao\nnhiệm vụ học tập',
        'Bước 2:\nHọc sinh tiếp nhận\nnhiệm vụ học tập',
        'Bước 3:\nBáo cáo kết quả\nhoạt động',
        'Bước 4:\nGiáo viên nhắc nhở\nnhiệm vụ về nhà',
    ]
    buoc_labels = buoc_labels_last if is_last else buoc_labels_default

    table = doc.add_table(rows=1, cols=3)
    set_table_borders(table)

    # Thiết lập chiều rộng cột: 20% | 40% | 40% (tổng ~9026 units = 16cm)
    tbl = table._tbl
    tblGrid = OxmlElement('w:tblGrid')
    for w in [1805, 3610, 3611]:  # twips (units)
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(w))
        tblGrid.append(gridCol)
    # Chèn tblGrid vào sau tblPr
    tblPr_el = tbl.find(qn('w:tblPr'))
    if tblPr_el is not None:
        tblPr_el.addnext(tblGrid)
    else:
        tbl.insert(0, tblGrid)

    # Header row
    hdr = table.rows[0].cells
    fill_cell(hdr[0], 'Bước', bold=True, line_ratio=1.15,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell(hdr[1], 'Hoạt động của GV', bold=True, line_ratio=1.15,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell(hdr[2], 'Hoạt động của HS', bold=True, line_ratio=1.15,
              align=WD_ALIGN_PARAGRAPH.CENTER)

    # 4 data rows
    for i, (gv_text, hs_text) in enumerate(buoc_data):
        row = table.add_row()
        fill_cell(row.cells[0], buoc_labels[i], bold=True, italic=True,
                  line_ratio=1.15, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(row.cells[1], gv_text, line_ratio=1.15)
        fill_cell(row.cells[2], hs_text, line_ratio=1.15)

    # Thêm paragraph trống sau bảng để phân cách
    add_para(doc, '', line_ratio=1.15)
    return table


def create_khbd_thcs(grade_name, grade_folder, bai_so, ten_bai,
                     mon_hoc, lop, thoi_luong, tiet_ppct,
                     ngay_soan, ngay_day,
                     kien_thuc_items, nang_luc_chung_items,
                     nang_luc_dac_thu_items, nang_luc_so_items,
                     pham_chat_items,
                     thiet_bi, hoc_lieu,
                     hoat_dong_list):
    """
    Tạo KHBD THCS hoàn chỉnh theo Phương án B (bảng 3 cột).
    hoat_dong_list: list of dicts:
      {
        'stt': 1,
        'ten': 'Khởi động (Xác định vấn đề/nhiệm vụ học tập/Mở đầu)',
        'muc_tieu': '...',
        'noi_dung': '...',
        'san_pham': '...',
        'to_chuc': '...',          # tùy chọn, mô tả hình thức tổ chức
        'buoc1_gv': '...GV làm...',
        'buoc1_hs': '...HS làm...',
        'buoc2_gv': '...',
        'buoc2_hs': '...',
        'buoc3_gv': '...',
        'buoc3_hs': '...',
        'buoc4_gv': '...',
        'buoc4_hs': '...',
      }
    """
    doc = Document(TPL_THCS)
    clean_body(doc)

    # ── Table[0]: Thông tin trường/GV — copy từ template ──────────────────
    doc_tpl = Document(TPL_THCS)
    tbl_info = copy.deepcopy(doc_tpl.tables[0]._tbl)

    # Sửa ngày soạn/dạy trong Row[1], Col[1]
    rows = tbl_info.findall(qn('w:tr'))
    if len(rows) >= 2:
        cells_r1 = rows[1].findall(qn('w:tc'))
        if len(cells_r1) >= 2:
            tc_ngay = cells_r1[1]
            for p_el in tc_ngay.findall(qn('w:p')):
                for r_el in p_el.findall(qn('w:r')):
                    t_el = r_el.find(qn('w:t'))
                    if t_el is not None and t_el.text:
                        if 'Ngày soạn' in t_el.text:
                            t_el.text = f'Ngày soạn: {ngay_soan}   Ngày dạy: {ngay_day}'

    # Insert table vào body
    sect_pr = doc.element.body.find(qn('w:sectPr'))
    if sect_pr is not None:
        idx = list(doc.element.body).index(sect_pr)
        doc.element.body.insert(idx, tbl_info)
    else:
        doc.element.body.append(tbl_info)

    # ── Tên bài + Tiết PPCT ───────────────────────────────────────────────
    add_para(doc, f'TÊN BÀI DẠY: {ten_bai.upper()}',
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, line_ratio=1.15)
    add_para(doc, f'Môn học: {mon_hoc}   Lớp: {lop}   Thời lượng: {thoi_luong}',
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, line_ratio=1.15)
    add_para(doc, f'Tiết theo PPCT: {tiet_ppct}',
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, line_ratio=1.15)
    add_para(doc, f'Tên tiết: {ten_bai}',
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, line_ratio=1.15)

    # ── I. MỤC TIÊU ──────────────────────────────────────────────────────
    add_para(doc, 'I. Mục tiêu', bold=True, line_ratio=1.15)

    add_para(doc, '1. Kiến thức: ', bold=True,
             indent_first=180340, line_ratio=1.15)
    for kt in kien_thuc_items:
        add_para(doc, kt, indent_first=180340, sp_after=0, line_ratio=1.15)

    p_nl = add_para(doc, '', indent_first=180340, line_ratio=1.15)
    r_nl_label = p_nl.add_run('2. Năng lực: ')
    afont(r_nl_label, bold=True)
    r_nl_desc = p_nl.add_run('Nêu cụ thể yêu cầu học sinh làm được gì')
    afont(r_nl_desc)

    add_para(doc, '- Năng lực chung:', bold=True,
             indent_first=180340, sp_after=0, line_ratio=1.15)
    for nl in nang_luc_chung_items:
        add_para(doc, f'  {nl}', indent_first=180340, sp_after=0, line_ratio=1.15)

    add_para(doc, '- Năng lực đặc thù:', bold=True,
             indent_first=180340, sp_after=0, line_ratio=1.15)
    for nl in nang_luc_dac_thu_items:
        add_para(doc, f'  {nl}', indent_first=180340, sp_after=0, line_ratio=1.15)

    add_para(doc, '- Năng lực số:', bold=True,
             indent_first=180340, sp_after=0, line_ratio=1.15)
    for nl in nang_luc_so_items:
        p = add_para(doc, f'  {nl}', sp_after=0, line_ratio=1.15,
                     indent_left=540000)
        p.paragraph_format.first_line_indent = Emu(-180000)

    add_para(doc, '3. Phẩm chất: ', bold=True,
             indent_first=180340, line_ratio=1.15)
    for pc in pham_chat_items:
        add_para(doc, pc, indent_first=180340, sp_after=0, line_ratio=1.15)

    # ── II. THIẾT BỊ ─────────────────────────────────────────────────────
    p_tb = add_para(doc, '', line_ratio=1.15)
    r_tb = p_tb.add_run('II. Thiết bị dạy học và học liệu: ')
    afont(r_tb, bold=True)

    add_para(doc, '1. Thiết bị', bold=True,
             indent_first=180340, sp_after=0, line_ratio=1.15)
    add_para(doc, thiet_bi, indent_first=180340, sp_after=0, line_ratio=1.15)

    add_para(doc, '2. Học liệu', bold=True,
             indent_first=180340, sp_after=0, line_ratio=1.15)
    add_para(doc, hoc_lieu, indent_first=180340, sp_after=0, line_ratio=1.15)

    # ── III. TIẾN TRÌNH DẠY HỌC (Phương án B: bảng 3 cột) ────────────────
    add_para(doc, 'III. Tiến trình dạy học', bold=True, line_ratio=1.15)

    for hd in hoat_dong_list:
        stt = hd['stt']
        ten_hd = hd['ten']
        is_last = (stt == len(hoat_dong_list))

        # Tiêu đề hoạt động
        add_para(doc, f'{stt}. Hoạt động {stt}. {ten_hd}', bold=True,
                 indent_first=180340, line_ratio=1.15)

        # a) Mục tiêu
        p_a = add_para(doc, '', indent_first=180340, sp_after=0, line_ratio=1.15)
        r_a_label = p_a.add_run('a) Mục tiêu: ')
        afont(r_a_label, italic=True)
        r_a_val = p_a.add_run(hd.get('muc_tieu', ''))
        afont(r_a_val)

        # b) Nội dung
        p_b = add_para(doc, '', indent_first=180340, sp_after=0, line_ratio=1.15)
        r_b_label = p_b.add_run('b) Nội dung: ')
        afont(r_b_label, italic=True)
        r_b_val = p_b.add_run(hd.get('noi_dung', ''))
        afont(r_b_val)

        # c) Sản phẩm
        p_c = add_para(doc, '', indent_first=180340, sp_after=0, line_ratio=1.15)
        r_c_label = p_c.add_run('c) Sản phẩm: ')
        afont(r_c_label, italic=True)
        r_c_val = p_c.add_run(hd.get('san_pham', ''))
        afont(r_c_val)

        # d) Tổ chức thực hiện
        p_d = add_para(doc, '', indent_first=180340, line_ratio=1.15)
        r_d_label = p_d.add_run('d) Tổ chức thực hiện: ')
        afont(r_d_label, italic=True)
        to_chuc = hd.get('to_chuc', '')
        if to_chuc:
            r_d_val = p_d.add_run(to_chuc)
            afont(r_d_val)


        # Bảng 3 cột Bước 1-4
        # Hỗ trợ cả key mới (buoc1_gv/buoc1_hs) và key cũ (buoc1)
        def _get_buoc(hd, n):
            gv = hd.get(f'buoc{n}_gv') or hd.get(f'buoc{n}', '')
            hs = hd.get(f'buoc{n}_hs', '')
            return gv, hs
        buoc_data = [_get_buoc(hd, 1), _get_buoc(hd, 2),
                     _get_buoc(hd, 3), _get_buoc(hd, 4)]
        add_activity_table_b(doc, buoc_data, is_last=is_last)


    # ── RÚT KINH NGHIỆM ──────────────────────────────────────────────────
    add_para(doc, 'RÚT KINH NGHIỆM SAU BÀI DẠY', bold=True, line_ratio=1.15)
    dot_line = '........................................................................................................................'
    for _ in range(4):
        add_para(doc, dot_line, line_ratio=1.15)
    add_para(doc, '')


    # ── Table ký tên — copy từ template ───────────────────────────────────
    tbl_sign = copy.deepcopy(doc_tpl.tables[2]._tbl)
    sect_pr2 = doc.element.body.find(qn('w:sectPr'))
    if sect_pr2 is not None:
        idx2 = list(doc.element.body).index(sect_pr2)
        doc.element.body.insert(idx2, tbl_sign)
    else:
        doc.element.body.append(tbl_sign)

    # ── Save ──────────────────────────────────────────────────────────────
    if bai_so == 0:
        out_dir = os.path.join(OUT_BASE, grade_folder, 'Tiết_00')
        clean_name = ten_bai.replace(":", "").replace("/", "_").replace("\\", "_").replace(" ", "_")[:40]
        filename = f'KHBD_Tin_hoc_{grade_folder}_Tiet00_{clean_name}.docx'
    else:
        out_dir = os.path.join(OUT_BASE, grade_folder, f'Bài_{bai_so:02d}')
        clean_name = ten_bai.replace(":", "").replace("/", "_").replace("\\", "_").replace(" ", "_")[:40]
        filename = f'KHBD_Tin_hoc_{grade_folder}_Bai{bai_so:02d}_{clean_name}.docx'
    out_path = os.path.join(out_dir, filename)
    saved = save_doc(doc, out_path)

    # ── Verify ────────────────────────────────────────────────────────────
    verify = Document(saved)
    hdr = verify.sections[0].header
    hdr_drawings = sum(1 for p in hdr.paragraphs
                       for r in p.runs if r._r.findall(qn('w:drawing')))
    tbl_count = len(verify.tables)
    print(f'  ✅ Saved: {saved}')
    print(f'     Header drawings: {hdr_drawings} | Tables: {tbl_count} (expected ≥6)')
    return saved


# ════════════════════════════════════════════════════════════════════════════
# DỮ LIỆU NỘI DUNG CHO TỪNG LỚP - Bài 1
# ════════════════════════════════════════════════════════════════════════════


def build_all():
    results = []

    # ════════════════════════════════════════════════════════════════════
    # TIẾT 0 — ĐỊNH HƯỚNG MÔN HỌC (TẤT CẢ CÁC KHỐI LỚP)
    # ════════════════════════════════════════════════════════════════════

    # 1. Tiền tiểu học - Tiết 0
    print('\n🔵 Tạo KHBD Tiền tiểu học - Tiết 0...')
    saved = create_khbd_th(
        grade_name='Tiền tiểu học', grade_folder='Tiền_tiểu_học',
        bai_so=0, ten_bai='TIẾT 0. ĐỊNH HƯỚNG MÔN HỌC - EM LÀM QUEN VỚI THẾ GIỚI CÔNG NGHỆ',
        chu_diem='Định hướng môn học', tiet_ppct='0', ngay_day='   /   /2026',
        pham_chat_items=[
            '- Chăm chỉ: Tích cực lắng nghe thầy cô giới thiệu phòng học Tin học và thiết bị.',
            '- Trách nhiệm: Giữ gìn an toàn thiết bị máy tính, tuân thủ hướng dẫn của thầy cô.',
            '- Nhân ái: Yêu thương, nhường nhịn và giúp đỡ bạn bè trong phòng máy.',
        ],
        nang_luc_mon_items=[
            '- NLa: Nhận diện phòng thực hành Tin học là không gian học tập đặc biệt.',
            '- NLb: Biết vị trí chỗ ngồi và cách làm quen nhẹ nhàng với máy tính.',
        ],
        nang_luc_chung_items=[
            '- Tự chủ và tự học: Bước đầu có thói quen tự giác vào lớp và ngồi đúng vị trí.',
            '- Giao tiếp và hợp tác: Làm quen với thầy cô và các bạn trong lớp.',
        ],
        do_dung_gv='Phòng máy tính ngăn nắp, máy chiếu, nhạc vui tươi, bảng nội quy có hình ảnh minh họa.',
        do_dung_hs='Tâm thế vui tươi, sẵn sàng học tập.',
        phuong_phap='trực quan, làm quen, trò chơi học tập, trò chuyện thân thiện',
        ki_thuat='đặt câu hỏi nhẹ nhàng, trình bày 1 phút, động viên khen thưởng',
        activities=[
            {
                'ten': '1. Hoạt động MỞ ĐẦU (5 phút)',
                'muc_tieu': 'Chào đón HS, tạo khí thế vui tươi khi đến với phòng Tin học.',
                'gv': [
                    'GV bật bản nhạc vui tươi chào đón HS bước vào phòng thực hành Tin học.',
                    'GV tự giới thiệu bản thân và chào mừng các em học sinh.',
                ],
                'hs': [
                    '- HS đi vào phòng máy theo hàng, mỉm cười chào thầy cô.',
                    '- HS vỗ tay theo giai điệu nhạc.',
                ],
            },
            {
                'ten': '2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI',
                'sub': [
                    {
                        'ten': '2.1. Khám phá phòng máy tính của em (10 phút)',
                        'muc_tieu': 'HS nhận biết được phòng máy tính và chỗ ngồi cố định của mình.',
                        'gv': [
                            'GV giới thiệu xung quanh phòng máy: máy tính, màn chiếu, bàn ghế.',
                            'GV hướng dẫn từng cặp HS về đúng vị trí chỗ ngồi của mình.',
                        ],
                        'hs': [
                            '- HS quan sát xung quanh phòng máy.',
                            '- HS về đúng chỗ ngồi được phân công và ngồi ngay ngắn.',
                        ],
                    },
                    {
                        'ten': '2.2. Quy tắc 3 NÊN và 3 KHÔNG (8 phút)',
                        'muc_tieu': 'HS ghi nhớ những quy tắc an toàn cơ bản nhất.',
                        'gv': [
                            'GV chiếu tranh minh họa 3 NÊN (Lắng nghe, Giữ vệ sinh, Ngồi đúng tư thế) và 3 KHÔNG (Không ăn uống, Không sờ vào dây điện, Không đùa giỡn).',
                            'GV cùng HS đọc thuộc 3 NÊN và 3 KHÔNG qua bài vè vui.',
                        ],
                        'hs': [
                            '- HS xem tranh minh họa.',
                            '- HS đọc vè theo GV.',
                        ],
                    },
                ],
            },
            {
                'ten': '3. HĐ LUYỆN TẬP-THỰC HÀNH',
                'sub': [
                    {
                        'ten': '3.1. Trò chơi "Bé ngoan phòng máy" (7 phút)',
                        'muc_tieu': 'Rèn luyện phản xạ ngồi đúng tư thế và tuân thủ quy tắc.',
                        'gv': [
                            'GV hô: "Tư thế đẹp!" → HS ngồi thẳng lưng, tay đặt lên bàn.',
                            'GV khen thưởng các bé thực hiện đúng nhanh nhất.',
                        ],
                        'hs': [
                            '- HS tham gia trò chơi hào hứng.',
                            '- HS thi đua ngồi đẹp.',
                        ],
                    },
                ],
            },
            {
                'ten': '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                'gv': [
                    '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                    'GV hỏi: "Hôm nay bước vào phòng máy em thấy thế nào?"',
                    'GV dặn: Về nhà kể cho bố mẹ nghe về thầy cô và phòng Tin học nhé!',
                ],
                'hs': [
                    '- HS phát biểu cảm xúc vui vẻ.',
                    '- HS vẫy tay chào thầy cô khi ra về.',
                ],
            },
        ]
    )
    results.append(saved)

    # 2. Lớp 1 - Tiết 0
    print('\n🔵 Tạo KHBD Lớp 1 - Tiết 0...')
    saved = create_khbd_th(
        grade_name='Lớp 1', grade_folder='Lớp_1',
        bai_so=0, ten_bai='TIẾT 0. ĐỊNH HƯỚNG MÔN HỌC - NỘI QUY VÀ AN TOÀN PHÒNG MÁY TÍNH',
        chu_diem='Định hướng môn học', tiet_ppct='0', ngay_day='   /   /2026',
        pham_chat_items=[
            '- Chăm chỉ: Hăng hái tham gia thảo luận quy tắc an toàn phòng máy.',
            '- Trách nhiệm: Giữ gìn an toàn thiết bị, bảo vệ bản thân và bạn bè khỏi nguy cơ về điện.',
            '- Trung thực: Nhận lỗi ngay khi lỡ làm rớt hay chạm nhầm thiết bị.',
        ],
        nang_luc_mon_items=[
            '- NLa: Nắm được quy trình học tập môn Tin học lớp 1 và các đồ dùng cần chuẩn bị.',
            '- NLb: Biết các quy tắc an toàn điện và vệ sinh phòng thực hành.',
        ],
        nang_luc_chung_items=[
            '- Tự chủ và tự học: Chủ động chuẩn bị sách vở và đồ dùng học tập môn Tin học.',
            '- Giao tiếp và hợp tác: Thảo luận nhóm đôi về nội quy phòng máy.',
        ],
        do_dung_gv='Máy chiếu, bảng nội quy phòng máy tính, video hoạt hình an toàn điện.',
        do_dung_hs='SGK Tin học 1, vở ghi, bút.',
        phuong_phap='xem video, vấn đáp, thảo luận nhóm, thực hành tư thế',
        ki_thuat='đặt câu hỏi, đóng vai tình huống, chia sẻ nhóm đôi',
        activities=[
            {
                'ten': '1. Hoạt động MỞ ĐẦU (5 phút)',
                'muc_tieu': 'Tạo không khí học tập hào hứng và kết nối bài học.',
                'gv': [
                    'GV cho HS xem 1 video ngắn hoạt hình vui nhộn về máy tính.',
                    'GV hỏi: "Học môn Tin học chúng ta được làm những gì?"',
                ],
                'hs': [
                    '- HS xem video vui vẻ.',
                    '- HS phát biểu: được chơi game học tập, được vẽ tranh, gõ phím...',
                ],
            },
            {
                'ten': '2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI',
                'sub': [
                    {
                        'ten': '2.1. Nội quy phòng thực hành Tin học (10 phút)',
                        'muc_tieu': 'HS nắm rõ các điều khoản trong bảng nội quy.',
                        'gv': [
                            'GV chiếu Bảng nội quy phòng Tin học gồm 5 điều.',
                            'GV giải thích ý nghĩa từng điều nội quy qua hình ảnh thực tế.',
                        ],
                        'hs': [
                            '- HS quan sát và lắng nghe GV giải thích.',
                            '- HS đọc đồng thanh 5 điều nội quy.',
                        ],
                    },
                    {
                        'ten': '2.2. An toàn điện và tư thế ngồi đúng (5 phút)',
                        'muc_tieu': 'HS biết cách phòng tránh nguy cơ giật điện và bảo vệ mắt, cột sống.',
                        'gv': [
                            'GV chỉ rõ các vị trí dây điện, ổ cắm — nhắc tuyệt đối KHÔNG chạm vào.',
                            'GV thao tác mẫu tư thế ngồi: lưng thẳng, mắt cách màn hình 50cm.',
                        ],
                        'hs': [
                            '- HS quan sát các vị trí an toàn.',
                            '- HS điều chỉnh dáng ngồi theo mẫu của GV.',
                        ],
                    },
                ],
            },
            {
                'ten': '3. HĐ LUYỆN TẬP-THỰC HÀNH',
                'sub': [
                    {
                        'ten': '3.1. Thực hành dáng ngồi đẹp (7 phút)',
                        'muc_tieu': 'Củng cố tư thế ngồi chuẩn khi làm việc với máy tính.',
                        'gv': ['GV đi từng bàn kiểm tra và chỉnh sửa tư thế cho HS.'],
                        'hs': ['- HS ngồi đúng tư thế, tay đặt lên bàn máy tính.'],
                    },
                ],
            },
            {
                'ten': '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                'gv': [
                    '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                    'GV: "Em hãy nhủ bạn bên cạnh 1 việc KHÔNG được làm trong phòng máy."',
                    'GV dặn dò chuẩn bị SGK và vở cho tiết sau.',
                ],
                'hs': [
                    '- HS nhắc nhở bạn ngồi cùng.',
                    '- HS ghi nhớ dặn dò.',
                ],
            },
        ]
    )
    results.append(saved)

    # 3. Lớp 2 - Tiết 0
    print('\n🔵 Tạo KHBD Lớp 2 - Tiết 0...')
    saved = create_khbd_th(
        grade_name='Lớp 2', grade_folder='Lớp_2',
        bai_so=0, ten_bai='TIẾT 0. ĐỊNH HƯỚNG MÔN HỌC - EM TRỞ THÀNH NHÀ SÁNG TẠO SỐ',
        chu_diem='Định hướng môn học', tiet_ppct='0', ngay_day='   /   /2026',
        pham_chat_items=[
            '- Chăm chỉ: Tích cực tìm hiểu mục tiêu môn Tin học 2.',
            '- Trách nhiệm: Giữ gìn vệ sinh chung, bảo quản thiết bị.',
            '- Nhân ái: Sẵn sàng hợp tác và hỗ trợ bạn bè cùng tiến bộ.',
        ],
        nang_luc_mon_items=[
            '- NLa: Hiểu được định hướng môn Tin học 2 giúp em sáng tạo ra các sản phẩm số đơn giản.',
            '- NLb: Nắm rõ cách thức đánh giá và ôn tập trong năm học.',
        ],
        nang_luc_chung_items=[
            '- Tự chủ và tự học: Biết lập mục tiêu cá nhân cho môn học.',
            '- Giao tiếp và hợp tác: Thảo luận nhóm về các sản phẩm số yêu thích.',
        ],
        do_dung_gv='Slide giới thiệu chương trình Tin 2, video các sản phẩm HS lớp trước đã làm (tranh vẽ, bài gõ).',
        do_dung_hs='SGK Tin học 2, vở, bút.',
        phuong_phap='trình bày trực quan, thảo luận nhóm, chia sẻ cảm xúc',
        ki_thuat='động não, trình bày 1 phút, hỏi đáp',
        activities=[
            {
                'ten': '1. Hoạt động MỞ ĐẦU (5 phút)',
                'muc_tieu': 'Khơi gợi niềm yêu thích môn Tin học lớp 2.',
                'gv': [
                    'GV chiếu một số bức tranh vẽ đẹp và bài gõ văn bản của HS lớp trước.',
                    'GV hỏi: "Em có muốn tự tay tạo nên những sản phẩm đẹp như vậy không?"',
                ],
                'hs': [
                    '- HS quan sát các sản phẩm.',
                    '- HS hào hứng reo hò phát biểu.',
                ],
            },
            {
                'ten': '2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI',
                'sub': [
                    {
                        'ten': '2.1. Hành trình khám phá Tin học 2 (10 phút)',
                        'muc_tieu': 'HS nắm được tổng quan các chủ đề học trong năm.',
                        'gv': [
                            'GV giới thiệu 4 chủ đề chính: Máy tính & em, Tập gõ bàn phím, Vẽ tranh sáng tạo, An toàn số.',
                        ],
                        'hs': [
                            '- HS mở mục lục SGK Tin 2, quan sát và đọc tên các chủ đề.',
                        ],
                    },
                    {
                        'ten': '2.2. Tiêu chí đánh giá và khen thưởng (5 phút)',
                        'muc_tieu': 'HS hiểu cách thức đánh giá và tích điểm thi đua.',
                        'gv': [
                            'GV giới thiệu bảng tích sao thi đua "Nhà sáng tạo số nhỏ tuổi".',
                        ],
                        'hs': [
                            '- HS lắng nghe và đặt mục tiêu tích sao.',
                        ],
                    },
                ],
            },
            {
                'ten': '3. HĐ LUYỆN TẬP-THỰC HÀNH',
                'sub': [
                    {
                        'ten': '3.1. Lập mục tiêu học tập (8 phút)',
                        'muc_tieu': 'HS viết/vẽ ước mơ mục tiêu đạt được trong môn Tin học 2.',
                        'gv': ['GV phát phiếu mục tiêu nhỏ, hướng dẫn HS điền.'],
                        'hs': ['- HS viết mục tiêu (VD: Đạt 10 điểm, Vẽ được bức tranh ngôi nhà...).'],
                    },
                ],
            },
            {
                'ten': '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                'gv': [
                    '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                    'GV: "Hãy dán tấm phiếu mục tiêu này vào trang đầu vở Tin học của em!"',
                ],
                'hs': ['- HS dán phiếu mục tiêu vào vở.'],
            },
        ]
    )
    results.append(saved)

    # 4. Lớp 3 - Tiết 0
    print('\n🔵 Tạo KHBD Lớp 3 - Tiết 0...')
    saved = create_khbd_th(
        grade_name='Lớp 3', grade_folder='Lớp_3',
        bai_so=0, ten_bai='TIẾT 0. ĐỊNH HƯỚNG MÔN HỌC - KHÁM PHÁ MÔN TIN HỌC 3',
        chu_diem='Định hướng môn học', tiet_ppct='0', ngay_day='   /   /2026',
        pham_chat_items=[
            '- Chăm chỉ: Thích học hỏi, chủ động khám phá tri thức số mới.',
            '- Trách nhiệm: Tuân thủ quy định phòng máy, bảo vệ thiết bị.',
            '- Trung thực: Tự giác trong kiểm tra đánh giá.',
        ],
        nang_luc_mon_items=[
            '- NLa: Hiểu được vai trò của Tin học 3 theo Chương trình GDPT 2018.',
            '- NLc: Biết phương pháp học tập trải nghiệm kết hợp thực hành trên máy.',
        ],
        nang_luc_chung_items=[
            '- Tự chủ và tự học: Biết tự quản lý thời gian học tập phòng máy.',
            '- Giao tiếp và hợp tác: Hợp tác tốt với bạn cùng bàn thực hành.',
        ],
        do_dung_gv='Máy chiếu, sơ đồ tư duy chương trình Tin học 3, phiếu khảo sát ban đầu.',
        do_dung_hs='SGK Tin học 3, vở ghi.',
        phuong_phap='sơ đồ tư duy, hỏi đáp gợi mở, làm việc nhóm đôi',
        ki_thuat='động não, chia sẻ nhóm đôi, KWL',
        activities=[
            {
                'ten': '1. Hoạt động MỞ ĐẦU (5 phút)',
                'muc_tieu': 'Tạo tâm thế chủ động, khám phá môn học mới bắt buộc.',
                'gv': [
                    'GV chúc mừng HS bước vào Lớp 3 — năm đầu tiên học Tin học chính thức!',
                    'GV đặt câu hỏi: "Em mong chờ điều gì nhất ở môn Tin học 3?"',
                ],
                'hs': [
                    '- HS vỗ tay chúc mừng.',
                    '- HS chia sẻ mong muốn cá nhân.',
                ],
            },
            {
                'ten': '2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI',
                'sub': [
                    {
                        'ten': '2.1. Cấu trúc chương trình Tin học 3 (10 phút)',
                        'muc_tieu': 'HS nắm bắt 6 chủ đề lớn của môn Tin học 3.',
                        'gv': ['GV giới thiệu sơ đồ tư duy 6 chủ đề (A, B, C, D, E, F).'],
                        'hs': ['- HS theo dõi sơ đồ và đối chiếu với SGK.'],
                    },
                    {
                        'ten': '2.2. Phương pháp học tập hiệu quả (5 phút)',
                        'muc_tieu': 'HS biết cách kết hợp giữa lý thuyết và thực hành.',
                        'gv': ['GV nêu 3 bí quyết: Quan sát kỹ -> Thao tác đúng -> Sáng tạo thêm.'],
                        'hs': ['- HS ghi nhớ 3 bí quyết vào vở.'],
                    },
                ],
            },
            {
                'ten': '3. HĐ LUYỆN TẬP-THỰC HÀNH',
                'sub': [
                    {
                        'ten': '3.1. Điền phiếu KWL (7 phút)',
                        'muc_tieu': 'Khảo sát nhu cầu và kiến thức sẵn có của HS.',
                        'gv': ['GV phát phiếu KWL: Điền cột K (Đã biết) và W (Muốn biết).'],
                        'hs': ['- HS hoàn thành phiếu cá nhân.'],
                    },
                ],
            },
            {
                'ten': '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                'gv': [
                    '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                    'GV dặn dò HS bọc sách vở Tin học 3 cẩn thận và mang đủ cho buổi học sau.',
                ],
                'hs': ['- HS ghi lại dặn dò.'],
            },
        ]
    )
    results.append(saved)

    # 5. Lớp 4 - Tiết 0
    print('\n🔵 Tạo KHBD Lớp 4 - Tiết 0...')
    saved = create_khbd_th(
        grade_name='Lớp 4', grade_folder='Lớp_4',
        bai_so=0, ten_bai='TIẾT 0. ĐỊNH HƯỚNG MÔN HỌC - KHÁM PHÁ MÔN TIN HỌC 4',
        chu_diem='Định hướng môn học', tiet_ppct='0', ngay_day='   /   /2026',
        pham_chat_items=[
            '- Chăm chỉ: Tích cực khám phá nội dung kiến thức nâng cao lớp 4.',
            '- Trách nhiệm: Giữ gìn an toàn thông tin cá nhân và an toàn mạng.',
            '- Trung thực: Không sao chép bài làm của bạn.',
        ],
        nang_luc_mon_items=[
            '- NLa: Nhận biết được các mạch kiến thức trọng tâm của Tin học 4.',
            '- NLc: Biết định hướng làm dự án học tập nhỏ và lập trình Scratch.',
        ],
        nang_luc_chung_items=[
            '- Tự chủ và tự học: Tự đặt ra kế hoạch rèn luyện gõ phím và tìm kiếm mạng.',
            '- Giao tiếp và hợp tác: Hợp tác tích cực trong các bài tập nhóm.',
        ],
        do_dung_gv='Máy chiếu, demo dự án Scratch mẫu, bài trình chiếu slide ấn tượng.',
        do_dung_hs='SGK Tin học 4, vở ghi.',
        phuong_phap='demo dự án, thảo luận, định hướng nhiệm vụ',
        ki_thuat='đặt câu hỏi gợi mở, chia sẻ nhóm',
        activities=[
            {
                'ten': '1. Hoạt động MỞ ĐẦU (5 phút)',
                'muc_tieu': 'Tạo sự hứng thú với nội dung lập trình và thiết kế slide ở Lớp 4.',
                'gv': [
                    'GV chiếu 1 minigame Scratch đơn giản do học sinh làm.',
                    'GV hỏi: "Các em có muốn tự tay lập trình trò chơi này không?"',
                ],
                'hs': [
                    '- HS xem demo trò chơi và vô cùng thích thú.',
                ],
            },
            {
                'ten': '2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI',
                'sub': [
                    {
                        'ten': '2.1. Tổng quan mạch kiến thức Tin 4 (10 phút)',
                        'muc_tieu': 'HS nắm bắt các điểm mới: Phần cứng-mềm, Tìm kiếm Internet, Lập trình Scratch.',
                        'gv': ['GV giới thiệu chi tiết 6 chủ đề chính của lớp 4.'],
                        'hs': ['- HS theo dõi và đánh dấu bài học yêu thích trong SGK.'],
                    },
                    {
                        'ten': '2.2. Yêu cầu sản phẩm học tập (5 phút)',
                        'muc_tieu': 'HS hiểu rõ các yêu cầu về bài trình chiếu và trò chơi Scratch.',
                        'gv': ['GV nêu quy định về sản phẩm cuối chủ đề.'],
                        'hs': ['- HS ghi nhớ quy định.'],
                    },
                ],
            },
            {
                'ten': '3. HĐ LUYỆN TẬP-THỰC HÀNH',
                'sub': [
                    {
                        'ten': '3.1. Thảo luận nhóm chọn chủ đề yêu thích (8 phút)',
                        'muc_tieu': 'Kích thích định hướng học tập tự chủ.',
                        'gv': ['GV chia nhóm 4, yêu cầu chọn 1 chủ đề nhóm mong muốn khám phá nhất.'],
                        'hs': ['- Các nhóm thảo luận và đại diện báo cáo.'],
                    },
                ],
            },
            {
                'ten': '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                'gv': [
                    '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                    'GV dặn chuẩn bị máy tính ở nhà (nếu có) và SGK cho Bài 1.',
                ],
                'hs': ['- HS ghi chép nhiệm vụ.'],
            },
        ]
    )
    results.append(saved)

    # 6. Lớp 5 - Tiết 0
    print('\n🔵 Tạo KHBD Lớp 5 - Tiết 0...')
    saved = create_khbd_th(
        grade_name='Lớp 5', grade_folder='Lớp_5',
        bai_so=0, ten_bai='TIẾT 0. ĐỊNH HƯỚNG MÔN HỌC - KHÁM PHÁ MÔN TIN HỌC 5',
        chu_diem='Định hướng môn học', tiet_ppct='0', ngay_day='   /   /2026',
        pham_chat_items=[
            '- Chăm chỉ: Tích cực chuẩn bị cho năm học cuối cấp Tiểu học.',
            '- Trách nhiệm: Gương mẫu thực hiện quy định phòng máy, làm nòng cốt giúp đỡ các em lớp dưới.',
            '- Trung thực: Tôn trọng bản quyền tác giả khi thu thập thông tin.',
        ],
        nang_luc_mon_items=[
            '- NLa: Nắm vững mục tiêu Tin học 5 — hoàn thiện năng lực số cấp Tiểu học.',
            '- NLe: Ý thức rõ về bản quyền, văn hóa ứng xử trên môi trường số.',
        ],
        nang_luc_chung_items=[
            '- Tự chủ và tự học: Tự định hướng phát triển kỹ năng số cá nhân.',
            '- Giao tiếp và hợp tác: Đóng vai trò trưởng nhóm trong các dự án nhỏ.',
        ],
        do_dung_gv='Slide tổng quan Tin 5, video giới thiệu kỹ năng số thế kỷ 21.',
        do_dung_hs='SGK Tin học 5, sổ tay học tập.',
        phuong_phap='thuyết trình, định hướng dự án, thảo luận nhóm',
        ki_thuat='đặt vấn đề, sơ đồ tư duy',
        activities=[
            {
                'ten': '1. Hoạt động MỞ ĐẦU (5 phút)',
                'muc_tieu': 'Tạo động lực học tập cho năm cuối cấp Tiểu học.',
                'gv': [
                    'GV chúc mừng HS lớp 5 và nhấn mạnh tầm quan trọng của kỹ năng số khi bước sang THCS.',
                ],
                'hs': ['- HS lắng nghe và thể hiện quyết tâm.'],
            },
            {
                'ten': '2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI',
                'sub': [
                    {
                        'ten': '2.1. Những điểm mới trong Tin học 5 (10 phút)',
                        'muc_tieu': 'HS biết được các nội dung nâng cao: Thuật toán, Xử lý ảnh, Bản quyền số.',
                        'gv': ['GV giới thiệu các chủ đề cốt lõi của Lớp 5.'],
                        'hs': ['- HS theo dõi SGK và ghi lại các từ khóa chính.'],
                    },
                    {
                        'ten': '2.2. Văn hóa số và Trách nhiệm người dùng (5 phút)',
                        'muc_tieu': 'Nhấn mạnh quy tắc ứng xử văn minh trên không gian mạng.',
                        'gv': ['GV phân tích quy tắc ứng xử văn minh và tôn trọng bản quyền.'],
                        'hs': ['- HS lắng nghe và cam kết thực hiện.'],
                    },
                ],
            },
            {
                'ten': '3. HĐ LUYỆN TẬP-THỰC HÀNH',
                'sub': [
                    {
                        'ten': '3.1. Cam kết văn hóa phòng máy (8 phút)',
                        'muc_tieu': 'Xây dựng môi trường học tập văn minh.',
                        'gv': ['GV hướng dẫn cả lớp xây dựng Bảng cam kết văn hóa phòng máy Lớp 5.'],
                        'hs': ['- Đại diện HS lên ký cam kết.'],
                    },
                ],
            },
            {
                'ten': '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                'gv': [
                    '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                    'GV dặn dò chuẩn bị bài học tuần sau.',
                ],
                'hs': ['- HS ghi nhớ dặn dò.'],
            },
        ]
    )
    results.append(saved)

    # 7. Lớp 6 - Tiết 0 (THCS)
    print('\n🟠 Tạo KHBD Lớp 6 - Tiết 0 (THCS)...')
    saved = create_khbd_thcs(
        grade_name='Lớp 6', grade_folder='Lớp_6',
        bai_so=0, ten_bai='Định hướng môn học Tin học 6 - Phương pháp học tập và An toàn số',
        mon_hoc='Tin học', lop='6', thoi_luong='1 tiết (45 phút)', tiet_ppct='0',
        ngay_soan='   /   /2026', ngay_day='   /   /2026',
        kien_thuc_items=[
            '- Sự hiểu biết về cấu trúc môn Tin học 6 theo Chương trình GDPT 2018.',
            '- Khả năng nhận biết các phương pháp học tập chủ động và nghiên cứu môn Tin học ở cấp THCS.',
            '- Sự hiểu biết về các nguyên tắc an toàn thông tin và văn hóa ứng xử trên mạng.',
        ],
        nang_luc_chung_items=[
            '- Tự chủ và tự học: Xác định được mục tiêu và kế hoạch học tập môn Tin học 6. (Đạt được thông qua Hoạt động 2, 3)',
            '- Giao tiếp và hợp tác: Thảo luận nhóm về các phương pháp học tập hiệu quả. (Đạt được thông qua Hoạt động 2)',
        ],
        nang_luc_dac_thu_items=[
            '- NLa (Hiểu biết máy tính): Nhận biết hệ thống phòng thực hành và các quy định an toàn THCS. (Đạt được thông qua Hoạt động 1, 2)',
            '- NLe (Ứng xử phù hợp): Hiểu được các quy tắc an toàn và văn hóa số khi tham gia không gian mạng. (Đạt được thông qua Hoạt động 3)',
        ],
        nang_luc_so_items=[
            '- Năng lực 1.1 (Hiểu biết số): Nắm bắt tổng quan về năng lực số cần đạt trong cấp THCS. (Đạt được thông qua Hoạt động 2)',
        ],
        pham_chat_items=[
            '- Chăm chỉ: Chủ động học tập, thích ứng nhanh với môi trường học tập cấp THCS. (Thông qua Hoạt động 1, 2)',
            '- Trách nhiệm: Giữ gìn an toàn thiết bị phòng máy và có trách nhiệm trên không gian số. (Thông qua Hoạt động 3, 4)',
        ],
        thiet_bi='Máy tính GV, máy chiếu, sơ đồ chương trình Tin 6, video an toàn mạng.',
        hoc_lieu='SGK Tin học 6, sổ tay học tập THCS.',
        hoat_dong_list=[
            {
                'stt': 1,
                'ten': 'Khởi động (Chào mừng HS đến với môn Tin học THCS)',
                'muc_tieu': 'Giúp HS cảm thấy tự tin và hào hứng khi bước vào cấp học mới.',
                'noi_dung': 'GV chiếu video chào mừng và đặt câu hỏi giao lưu: "Tin học THCS có gì khác so với Tiểu học?"',
                'san_pham': 'HS chia sẻ suy nghĩ và nêu được cảm nhận khi học Tin học ở cấp THCS.',
                'to_chuc': 'GV trình chiếu và dẫn dắt thảo luận.',
                'buoc1_gv': 'GV phát video chào mừng khối 6 THCS (1-2 phút), sau đó đặt câu hỏi giao lưu.',
                'buoc1_hs': 'HS xem video, theo dõi nội dung và sẵn sàng trả lời câu hỏi.',
                'buoc2_gv': 'GV quan sát HS, gợi ý thêm nếu HS chưa phát biểu.',
                'buoc2_hs': 'HS suy nghĩ câu hỏi: "Tin học THCS có gì khác Tiểu học?"',
                'buoc3_gv': 'GV ghi nhanh ý kiến HS lên bảng, khích lệ thêm người phát biểu.',
                'buoc3_hs': '2-3 HS phát biểu suy nghĩ cá nhân trước lớp.',
                'buoc4_gv': 'GV tổng kết, nhấn mạnh bước chuyển mình quan trọng sang cấp THCS và dẫn vào bài.',
                'buoc4_hs': 'HS lắng nghe và ghi chép những điểm mới của môn Tin học THCS.',
            },
            {
                'stt': 2,
                'ten': 'Hình thành kiến thức mới (Tổng quan chương trình & Phương pháp học)',
                'muc_tieu': 'HS hiểu rõ cấu trúc 6 chủ đề Tin 6 và các tiêu chí đánh giá.',
                'noi_dung': 'GV giới thiệu 6 chủ đề chính (Máy tính, Mạng máy tính, An toàn thông tin, Biểu diễn thông tin, Thuật toán, Lập trình).',
                'san_pham': 'Sơ đồ tư duy học tập cá nhân trong vở ghi của HS.',
                'to_chuc': 'GV giảng kết hợp làm việc nhóm.',
                'buoc1_gv': 'GV chiếu sơ đồ 6 chủ đề môn Tin 6, giới thiệu từng chủ đề ngắn gọn.',
                'buoc1_hs': 'HS theo dõi sơ đồ và mở SGK Tin 6 trang đầu để đối chiếu.',
                'buoc2_gv': 'GV đặt câu hỏi: "Em thấy chủ đề nào thú vị nhất? Vì sao?"',
                'buoc2_hs': 'HS quan sát SGK Tin 6, gạch chân các chủ đề và ghi vở sơ đồ tư duy cá nhân.',
                'buoc3_gv': 'GV trình bày phương pháp học chủ động: Học qua dự án, Học qua giải quyết vấn đề.',
                'buoc3_hs': 'HS ghi chép phương pháp học hiệu quả và đặt câu hỏi nếu cần làm rõ.',
                'buoc4_gv': 'GV chốt các hình thức đánh giá (Thường xuyên, Định kỳ, Sản phẩm thực hành).',
                'buoc4_hs': 'HS ghi chép tiêu chí đánh giá và kế hoạch phấn đấu cá nhân.',
            },
            {
                'stt': 3,
                'ten': 'Luyện tập (Nội quy phòng máy & An toàn số)',
                'muc_tieu': 'Củng cố các quy định phòng thực hành THCS và nguyên tắc an toàn số.',
                'noi_dung': 'HS thảo luận nhóm phân tích các tình huống vi phạm an toàn điện và an toàn mạng.',
                'san_pham': 'Kết quả xử lý tình huống của các nhóm trên bảng phụ.',
                'to_chuc': 'Thảo luận nhóm 4 người.',
                'buoc1_gv': 'GV giao 3 phiếu tình huống thực tế cho các nhóm: tình huống an toàn điện, mạng, bản quyền.',
                'buoc1_hs': 'HS nhận phiếu, đọc tình huống và chuẩn bị thảo luận trong nhóm.',
                'buoc2_gv': 'GV quan sát, gợi ý nhóm gặp khó khăn, đặt câu hỏi phụ nếu cần.',
                'buoc2_hs': 'HS thảo luận phân tích đúng/sai từng tình huống và đề xuất giải pháp xử lý.',
                'buoc3_gv': 'GV mời đại diện nhóm trình bày, ghi kết quả lên bảng.',
                'buoc3_hs': 'Đại diện nhóm trình bày kết quả, các nhóm khác lắng nghe và nhận xét.',
                'buoc4_gv': 'GV nhận xét, chuẩn hóa quy tắc an toàn và chốt 5 nguyên tắc vàng.',
                'buoc4_hs': 'HS ghi chép 5 nguyên tắc an toàn số vào vở.',
            },
            {
                'stt': 4,
                'ten': 'Mở rộng (Nhiệm vụ về nhà)',
                'muc_tieu': 'Chuẩn bị kế hoạch cá nhân cho môn học.',
                'noi_dung': 'HS lập bảng kế hoạch cá nhân "Chinh phục môn Tin học 6".',
                'san_pham': 'Bảng kế hoạch cá nhân hoàn chỉnh nộp vào tiết sau.',
                'buoc1_gv': 'GV hướng dẫn cấu trúc bảng kế hoạch: Mục tiêu — Kế hoạch tuần — Điểm tự đánh giá.',
                'buoc1_hs': 'HS lắng nghe, ghi cấu trúc bảng kế hoạch vào vở.',
                'buoc2_gv': 'GV dặn dò: Hoàn thành bảng kế hoạch tại nhà và mang nộp đầu tiết sau.',
                'buoc2_hs': 'HS ghi nhiệm vụ vào sổ tay và hỏi nếu chưa rõ.',
                'buoc3_gv': 'GV nhắc nhở: Chuẩn bị SGK Bài 1 và đọc trước trang 6-8.',
                'buoc3_hs': 'HS ghi nhớ nhiệm vụ, chuẩn bị SGK cho tiết sau.',
                'buoc4_gv': 'GV tổng kết tiết học, chúc HS có năm học Tin học thú vị.',
                'buoc4_hs': 'HS lắng nghe, thu dọn sách vở và rời phòng máy đúng quy định.',
            },
        ]
    )
    results.append(saved)

    # 8. Lớp 7 - Tiết 0 (THCS)
    print('\n🟠 Tạo KHBD Lớp 7 - Tiết 0 (THCS)...')
    saved = create_khbd_thcs(
        grade_name='Lớp 7', grade_folder='Lớp_7',
        bai_so=0, ten_bai='Định hướng môn học Tin học 7 - Tổng quan chương trình và Kỹ năng số',
        mon_hoc='Tin học', lop='7', thoi_luong='1 tiết (45 phút)', tiet_ppct='0',
        ngay_soan='   /   /2026', ngay_day='   /   /2026',
        kien_thuc_items=[
            '- Sự hiểu biết về mục tiêu và nội dung cốt lõi môn Tin học 7.',
            '- Khả năng định hướng ứng dụng phần mềm bảng tính và biên tập đa phương tiện.',
            '- Sự hiểu biết về khung năng lực số cấp THCS theo CV 3456.',
        ],
        nang_luc_chung_items=[
            '- Tự chủ và tự học: Biết xác định kỹ năng số cần rèn luyện trong năm học. (Đạt được thông qua Hoạt động 2)',
            '- Giao tiếp và hợp tác: Thảo luận nhóm xây dựng tiêu chí làm việc nhóm hiệu quả. (Đạt được thông qua Hoạt động 3)',
        ],
        nang_luc_dac_thu_items=[
            '- NLa (Hiểu biết máy tính): Nắm vững cấu trúc bài học và tiêu chí đánh giá sản phẩm thực hành. (Đạt được thông qua Hoạt động 2)',
            '- NLc (Giải quyết vấn đề): Biết sử dụng phần mềm bảng tính để xử lý dữ liệu thực tế. (Đạt được thông qua Hoạt động 3)',
        ],
        nang_luc_so_items=[
            '- Năng lực số (CV 3456): Khai thác và quản lý dữ liệu số an toàn, hiệu quả. (Đạt được thông qua Hoạt động 2, 3)',
        ],
        pham_chat_items=[
            '- Chăm chỉ: Tích cực thực hành rèn luyện kỹ năng xử lý dữ liệu.',
            '- Trách nhiệm: Tuân thủ quy định sử dụng thiết bị và bản quyền phần mềm.',
        ],
        thiet_bi='Máy tính GV, máy chiếu, demo các sản phẩm Bảng tính & Slide Lớp 7.',
        hoc_lieu='SGK Tin học 7, tài liệu hướng dẫn kỹ năng số.',
        hoat_dong_list=[
            {
                'stt': 1,
                'ten': 'Khởi động (Kết nối tri thức Lớp 6 & Lớp 7)',
                'muc_tieu': 'Ôn lại nền tảng Lớp 6 và định hướng yêu cầu cao hơn ở Lớp 7.',
                'noi_dung': 'GV đặt câu hỏi nhắm vào kỹ năng đã có và giới thiệu điểm mới ở Lớp 7.',
                'san_pham': 'HS nêu được các công cụ số mình đã thành thạo.',
                'to_chuc': 'GV nêu vấn đề và điều phối.',
                'buoc1_gv': 'GV chiếu bảng khảo sát nhanh: "Em đã thành thạo công cụ nào từ Lớp 6?"',
                'buoc1_hs': 'HS đọc câu hỏi khảo sát và chuẩn bị câu trả lời.',
                'buoc2_gv': 'GV gọi HS trả lời, ghi nhanh lên bảng các kỹ năng HS đã có.',
                'buoc2_hs': 'HS trả lời cá nhân, nêu các công cụ số mình đã sử dụng được từ Lớp 6.',
                'buoc3_gv': 'GV hỏi thêm: "Lớp 7 có gì mới? Em muốn học gì thêm?"',
                'buoc3_hs': '3 HS chia sẻ kinh nghiệm và kỳ vọng cho Lớp 7.',
                'buoc4_gv': 'GV chốt: Lớp 7 tiếp nối và nâng cao từ Lớp 6 — dẫn vào tổng quan chương trình.',
                'buoc4_hs': 'HS lắng nghe và ghi chú điểm khác biệt giữa Tin 6 và Tin 7.',
            },
            {
                'stt': 2,
                'ten': 'Hình thành kiến thức mới (Tổng quan Tin 7 & Khung Năng lực số)',
                'muc_tieu': 'HS nắm được trọng tâm: Bảng tính Excel, Biên tập ảnh/video, Đạo đức số.',
                'noi_dung': 'GV trình bày các chủ đề lớn của Lớp 7 và Khung Năng lực số CV 3456.',
                'san_pham': 'Ghi chép trọng tâm chương trình vào vở.',
                'to_chuc': 'GV trình chiếu minh họa.',
                'buoc1_gv': 'GV giới thiệu các chủ đề cốt lõi Tin 7: Bảng tính, Biên tập đa phương tiện, An toàn số nâng cao.',
                'buoc1_hs': 'HS theo dõi slide và mở SGK Tin 7 đối chiếu mục lục.',
                'buoc2_gv': 'GV trình bày Khung Năng lực số (CV 3456): 5 lĩnh vực, 3 mức độ.',
                'buoc2_hs': 'HS đối chiếu với SGK Tin 7, ghi chú chủ đề quan tâm.',
                'buoc3_gv': 'GV giải thích các mức độ năng lực số cần đạt trong năm học.',
                'buoc3_hs': 'HS đặt câu hỏi về những nội dung chưa rõ.',
                'buoc4_gv': 'GV chốt các cột mốc kiểm tra đánh giá và thời điểm.',
                'buoc4_hs': 'HS ghi chép lịch kiểm tra và tiêu chí đánh giá sản phẩm.',
            },
            {
                'stt': 3,
                'ten': 'Luyện tập (Xây dựng Quy tắc làm việc nhóm phòng thực hành)',
                'muc_tieu': 'Thống nhất nguyên tắc làm việc nhóm khi thực hành bài tập lớn.',
                'noi_dung': 'Các nhóm thảo luận xây dựng 5 tiêu chí nhóm hiệu quả.',
                'san_pham': 'Bản tiêu chí hoạt động nhóm trên giấy A4.',
                'to_chuc': 'Thảo luận nhóm 4.',
                'buoc1_gv': 'GV giao nhiệm vụ: Mỗi nhóm 4 người xây dựng 5 tiêu chí làm việc nhóm hiệu quả.',
                'buoc1_hs': 'HS nhận nhiệm vụ, phân công nhóm trưởng và thư ký.',
                'buoc2_gv': 'GV quan sát, hỗ trợ nhóm nào chưa có ý tưởng bằng gợi ý mẫu.',
                'buoc2_hs': 'HS thảo luận và viết 5 tiêu chí ra giấy A4.',
                'buoc3_gv': 'GV mời đại diện từng nhóm treo bản tiêu chí lên bảng.',
                'buoc3_hs': 'Đại diện nhóm treo bản tiêu chí và trình bày lý do chọn.',
                'buoc4_gv': 'GV tổng kết, chọn lọc và công bố quy tắc chung cho cả lớp trong cả năm học.',
                'buoc4_hs': 'HS ghi quy tắc nhóm chung vào vở và cam kết thực hiện.',
            },
            {
                'stt': 4,
                'ten': 'Mở rộng (Nhiệm vụ về nhà)',
                'muc_tieu': 'Chuẩn bị bài học tiếp theo.',
                'noi_dung': 'Đọc trước Bài 1: Thiết bị vào - ra.',
                'san_pham': 'Câu hỏi thắc mắc ban đầu về Bài 1.',
                'buoc1_gv': 'GV dặn HS đọc trước Bài 1 SGK Tin 7 trang 6-10 và liệt kê thiết bị ở nhà.',
                'buoc1_hs': 'HS ghi nhiệm vụ vào sổ tay.',
                'buoc2_gv': 'GV hỏi: "Ở nhà em có những thiết bị máy tính nào?"',
                'buoc2_hs': 'HS suy nghĩ và chuẩn bị danh sách thiết bị ở nhà.',
                'buoc3_gv': 'GV nhắc: Chụp ảnh hoặc ghi tên các thiết bị để chia sẻ đầu tiết sau.',
                'buoc3_hs': 'HS ghi chú cách chuẩn bị bài tập về nhà.',
                'buoc4_gv': 'GV kết thúc tiết học, nhắc HS thu dọn và thoát khỏi máy đúng quy trình.',
                'buoc4_hs': 'HS thoát phần mềm, tắt màn hình và rời phòng máy đúng quy định.',
            },
        ]
    )
    results.append(saved)

    # 9. Lớp 8 - Tiết 0 (THCS)
    print('\n🟠 Tạo KHBD Lớp 8 - Tiết 0 (THCS)...')
    saved = create_khbd_thcs(
        grade_name='Lớp 8', grade_folder='Lớp_8',
        bai_so=0, ten_bai='Định hướng môn học Tin học 8 - Định hướng học tập và Nghiên cứu công nghệ',
        mon_hoc='Tin học', lop='8', thoi_luong='1 tiết (45 phút)', tiet_ppct='0',
        ngay_soan='   /   /2026', ngay_day='   /   /2026',
        kien_thuc_items=[
            '- Sự hiểu biết về cấu trúc chương trình Tin học 8 (Lược sử công nghệ, Đồ họa vector, Xử lý dữ liệu nâng cao, Lập trình).',
            '- Khả năng định hướng nghiên cứu và ứng dụng công nghệ giải quyết bài toán thực tế.',
            '- Sự hiểu biết về tác động của công nghệ số và trí tuệ nhân tạo (AI) trong đời sống.',
        ],
        nang_luc_chung_items=[
            '- Tự chủ và tự học: Chủ động lập kế hoạch tự nghiên cứu công nghệ mới. (Đạt được thông qua Hoạt động 2, 4)',
            '- Giải quyết vấn đề và sáng tạo: Đề xuất ý tưởng ứng dụng Tin học vào các môn học khác. (Đạt được thông qua Hoạt động 3)',
        ],
        nang_luc_dac_thu_items=[
            '- NLa (Hiểu biết máy tính): Trình bày được định hướng môn học và phương pháp tư duy máy tính (Computational Thinking). (Đạt được thông qua Hoạt động 2)',
            '- NLc (Giải quyết vấn đề): Sử dụng công nghệ số như công cụ sáng tạo và nghiên cứu học tập. (Đạt được thông qua Hoạt động 3)',
        ],
        nang_luc_so_items=[
            '- Năng lực số (CV 3456): Khai thác an toàn các công cụ trí tuệ nhân tạo và dịch vụ đám mây phục vụ học tập. (Đạt được thông qua Hoạt động 2, 3)',
        ],
        pham_chat_items=[
            '- Chăm chỉ: Tích cực tìm hiểu công nghệ mới, có tinh thần ham học hỏi.',
            '- Trách nhiệm: Sử dụng công nghệ có trách nhiệm, tuân thủ đạo đức trí tuệ nhân tạo.',
        ],
        thiet_bi='Máy tính GV, máy chiếu, demo sản phẩm đồ họa vector Inkscape và lập trình Python/Scratch nâng cao.',
        hoc_lieu='SGK Tin học 8, tài liệu tham khảo công nghệ số.',
        hoat_dong_list=[
            {
                'stt': 1,
                'ten': 'Khởi động (Công nghệ thay đổi thế giới như thế nào?)',
                'muc_tieu': 'Truyền cảm hứng nghiên cứu công nghệ cho học sinh Lớp 8.',
                'noi_dung': 'GV chiếu clip 2 phút về sự phát triển của AI và công nghệ hiện đại.',
                'san_pham': 'HS phát biểu suy nghĩ về vai trò của công nghệ với tương lai bản thân.',
                'to_chuc': 'GV trình chiếu và dẫn dắt.',
                'buoc1_gv': 'GV chiếu video ngắn (2 phút) về AI, robot và công nghệ tương lai.',
                'buoc1_hs': 'HS xem video, chú ý những hình ảnh công nghệ ấn tượng nhất.',
                'buoc2_gv': 'GV đặt câu hỏi: "Công nghệ nào em thấy thú vị nhất? Tại sao?"',
                'buoc2_hs': 'HS suy nghĩ và ghi nhanh ý tưởng cá nhân vào giấy nháp.',
                'buoc3_gv': 'GV gọi 3-4 HS chia sẻ ước mơ nghề nghiệp liên quan đến công nghệ.',
                'buoc3_hs': 'HS chia sẻ ước mơ nghề nghiệp công nghệ: lập trình, thiết kế AI, kỹ sư robot...',
                'buoc4_gv': 'GV tổng kết: Tin học 8 sẽ trang bị nền tảng để HS hiểu và tạo ra công nghệ — dẫn vào bài.',
                'buoc4_hs': 'HS lắng nghe và ghi mục tiêu cá nhân cho năm học Tin 8.',
            },
            {
                'stt': 2,
                'ten': 'Hình thành kiến thức mới (Tổng quan Tin 8 & Tư duy máy tính)',
                'muc_tieu': 'HS nắm được nội dung trọng tâm Lớp 8 và khái niệm Tư duy máy tính (Computational Thinking).',
                'noi_dung': 'GV giới thiệu 5 mạch kiến thức Tin 8 và 4 trụ cột của tư duy máy tính (Tách nhỏ, Tìm quy luật, Trừu tượng hóa, Thuật toán).',
                'san_pham': 'HS ghi chép 4 trụ cột tư duy máy tính vào vở.',
                'to_chuc': 'GV trình bày kết hợp ví dụ minh họa.',
                'buoc1_gv': 'GV chiếu sơ đồ 5 mạch kiến thức Tin 8 và giới thiệu từng mảng.',
                'buoc1_hs': 'HS mở SGK Tin 8, đối chiếu mục lục với sơ đồ GV chiếu.',
                'buoc2_gv': 'GV trình bày 4 trụ cột Tư duy máy tính kèm ví dụ minh họa đời thực.',
                'buoc2_hs': 'HS ghi chép 4 trụ cột: Tách nhỏ — Tìm quy luật — Trừu tượng hóa — Thuật toán.',
                'buoc3_gv': 'GV đặt câu hỏi: "Các em đã từng dùng tư duy này chưa? Khi nào?"',
                'buoc3_hs': 'HS thảo luận cặp đôi, tìm ví dụ thực tế về tư duy máy tính trong cuộc sống.',
                'buoc4_gv': 'GV chốt phương pháp học Tin 8 theo dự án và nghiên cứu.',
                'buoc4_hs': 'HS ghi chép phương pháp học và thời gian biểu thực hành dự án.',
            },
            {
                'stt': 3,
                'ten': 'Luyện tập (Ứng dụng Tư duy máy tính giải quyết tình huống)',
                'muc_tieu': 'HS vận dụng 4 trụ cột tư duy máy tính để phân tích 1 bài toán thực tế.',
                'noi_dung': 'Bài toán: "Tổ chức một buổi triển lãm sản phẩm công nghệ của trường".',
                'san_pham': 'Bản phân tích công việc theo 4 trụ cột tư duy trên giấy nháp.',
                'to_chuc': 'Thảo luận nhóm 4.',
                'buoc1_gv': 'GV giao bài toán: "Hãy tổ chức một buổi triển lãm công nghệ — áp dụng 4 trụ cột tư duy máy tính".',
                'buoc1_hs': 'HS nhận nhiệm vụ, phân công vai trò trong nhóm.',
                'buoc2_gv': 'GV quan sát và hỗ trợ nhóm chưa biết cách tách nhỏ công việc.',
                'buoc2_hs': 'Nhóm thảo luận, tách nhỏ công việc và đề xuất giải pháp từng bước.',
                'buoc3_gv': 'GV mời đại diện 2 nhóm trình bày, ghi kết quả lên bảng.',
                'buoc3_hs': 'Đại diện nhóm trình bày dàn ý, nhóm khác nhận xét và bổ sung.',
                'buoc4_gv': 'GV nhận xét, chuẩn hóa và liên hệ với lập trình: "Tư duy này chính là cách máy tính giải bài toán".',
                'buoc4_hs': 'HS ghi chép kết luận về Tư duy máy tính và ứng dụng thực tiễn.',
            },
            {
                'stt': 4,
                'ten': 'Mở rộng (Nhiệm vụ về nhà)',
                'muc_tieu': 'Chuẩn bị cho Bài 1.',
                'noi_dung': 'Tìm hiểu lịch sử ra đời chiếc máy tính đầu tiên.',
                'san_pham': 'Ghi chú ngắn 3 câu về chiếc máy tính ENIAC.',
                'buoc1_gv': 'GV dặn HS đọc trước Bài 1 SGK Tin 8 (trang 6-14) và tìm hiểu về ENIAC.',
                'buoc1_hs': 'HS ghi nhiệm vụ: Đọc trước Bài 1 và tìm hiểu về chiếc máy tính ENIAC.',
                'buoc2_gv': 'GV gợi ý từ khóa tìm kiếm: "ENIAC máy tính đầu tiên", "lược sử máy tính".',
                'buoc2_hs': 'HS ghi chú nguồn tài liệu tham khảo.',
                'buoc3_gv': 'GV nhắc: Đầu tiết sau sẽ chia sẻ nhanh 1-2 câu về ENIAC.',
                'buoc3_hs': 'HS chuẩn bị 3 câu ghi chú ngắn về ENIAC để chia sẻ.',
                'buoc4_gv': 'GV tổng kết tiết học và nhắc nhở quy trình tắt máy.',
                'buoc4_hs': 'HS tắt máy đúng quy trình và rời phòng máy.',
            },
        ]
    )
    results.append(saved)

    # ════════════════════════════════════════════════════════════════════
    # BÀI 1 — NỘI DUNG BÀI HỌC ĐẦU TIÊN (TẤT CẢ CÁC KHỐI LỚP)
    # ════════════════════════════════════════════════════════════════════

    saved = create_khbd_th(
        grade_name='Tiền tiểu học', grade_folder='Tiền_tiểu_học',
        bai_so=1, ten_bai='BÀI 1. MÁY TÍNH XUNG QUANH EM', chu_diem='Làm quen với công nghệ',
        tiet_ppct='1', ngay_day='   /   /2026',
        pham_chat_items=[
            '- Chăm chỉ: Tích cực tham gia các hoạt động học tập, hoàn thành nhiệm vụ được giao.',
            '- Nhân ái: Biết chia sẻ, hợp tác với bạn bè trong các hoạt động nhóm.',
            '- Trách nhiệm: Giữ gìn thiết bị máy tính, sử dụng đúng cách.',
        ],
        nang_luc_mon_items=[
            '- NLa (Sử dụng và quản lí thiết bị công nghệ số): Nhận diện được hình dạng và chức năng cơ bản của máy tính.',
            '- NLb (Kết nối và cộng tác): Biết làm việc cùng bạn trong các hoạt động khám phá thiết bị.',
        ],
        nang_luc_chung_items=[
            '- Tự chủ và tự học: Quan sát, nhận biết các thiết bị xung quanh.',
            '- Giao tiếp và hợp tác: Chia sẻ những gì em biết về máy tính với bạn.',
        ],
        do_dung_gv='Máy chiếu, máy tính, tranh ảnh về các loại máy tính, thẻ từ hình ảnh thiết bị.',
        do_dung_hs='Bút, vở, tranh ảnh chuẩn bị sẵn về máy tính.',
        phuong_phap='vấn đáp, quan sát, trò chơi học tập, hoạt động nhóm đôi',
        ki_thuat='đặt câu hỏi, trình bày 1 phút, chia sẻ nhóm đôi',
        activities=[
            {
                'ten': '1. Hoạt động MỞ ĐẦU (5 phút)',
                'muc_tieu': 'Tạo hứng thú, kết nối bài học với cuộc sống.',
                'gv': [
                    'GV chiếu hình ảnh nhiều loại máy tính (laptop, máy tính bảng, điện thoại).',
                    'GV hỏi: "Em đã thấy những thiết bị này ở đâu chưa?"',
                    '* Nhận xét, dẫn vào bài: Hôm nay chúng ta cùng khám phá máy tính xung quanh em!',
                ],
                'hs': [
                    '- HS quan sát hình ảnh trên màn chiếu.',
                    '- HS xung phong trả lời: ở nhà, ở trường, trong siêu thị...',
                    '- HS lắng nghe, vỗ tay.',
                ],
            },
            {
                'ten': '2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI',
                'muc_tieu': '',
                'sub': [
                    {
                        'ten': '2.1. Nhận biết các loại máy tính (10 phút)',
                        'muc_tieu': 'HS phân biệt được laptop, máy tính bảng và điện thoại thông minh.',
                        'gv': [
                            'GV giới thiệu từng loại máy tính qua hình ảnh lớn trên màn chiếu.',
                            'GV đặt câu hỏi: "Cái này gọi là gì? Để làm gì?"',
                            '* GV kết luận tên gọi và chức năng từng loại.',
                        ],
                        'hs': [
                            '- HS quan sát, lắng nghe.',
                            '- HS đoán tên và trả lời theo hiểu biết.',
                            '- HS lắp lại tên gọi theo GV.',
                        ],
                    },
                    {
                        'ten': '2.2. Các bộ phận của máy tính (5 phút)',
                        'muc_tieu': 'HS chỉ đúng màn hình, bàn phím và chuột máy tính.',
                        'gv': [
                            'GV chỉ vào máy tính thật, giới thiệu: màn hình, bàn phím, chuột.',
                            'GV cho HS lần lượt lên chỉ và gọi tên bộ phận.',
                        ],
                        'hs': [
                            '- HS quan sát, lắng nghe.',
                            '- Từng HS lên chỉ và nói tên bộ phận.',
                        ],
                    },
                ],
            },
            {
                'ten': '3. HĐ LUYỆN TẬP-THỰC HÀNH',
                'muc_tieu': '',
                'sub': [
                    {
                        'ten': '3.1. Trò chơi: Tìm đúng tên (7 phút)',
                        'muc_tieu': 'Củng cố nhận biết tên các thiết bị máy tính.',
                        'gv': [
                            'GV phát thẻ hình ảnh thiết bị, HS ghép với thẻ tên tương ứng.',
                            '* Nhận xét, chữa bài, tuyên dương nhóm đúng nhanh nhất.',
                        ],
                        'hs': [
                            '- HS làm việc nhóm đôi, ghép thẻ hình-tên.',
                            '- HS đổi bài, kiểm tra chéo.',
                        ],
                    },
                    {
                        'ten': '3.2. Em tô màu máy tính (3 phút)',
                        'muc_tieu': 'HS ghi nhớ hình dạng thiết bị qua hoạt động sáng tạo.',
                        'gv': [
                            'GV phát phiếu tô màu hình máy tính.',
                            '* Nhận xét sản phẩm của HS.',
                        ],
                        'hs': [
                            '- HS tô màu theo ý thích.',
                            '- HS giơ bài, cùng chia sẻ với bạn ngồi cạnh.',
                        ],
                    },
                ],
            },
            {
                'ten': '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                'gv': [
                    '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                    'GV hỏi: "Về nhà, em thấy máy tính/điện thoại ở đâu? Ai dùng?"',
                    'GV dặn dò: Về nhà kể cho bố mẹ nghe tên 3 bộ phận của máy tính.',
                ],
                'hs': [
                    '- HS chia sẻ trước lớp.',
                    '- HS lắng nghe dặn dò.',
                ],
            },
        ]
    )
    results.append(saved)

    # ────────────────────────────────────────────────────────────────────
    # LỚP 1 — Bài 1
    # ────────────────────────────────────────────────────────────────────
    print('\n🔵 Tạo KHBD Lớp 1 - Bài 1...')
    saved = create_khbd_th(
        grade_name='Lớp 1', grade_folder='Lớp_1',
        bai_so=1, ten_bai='BÀI 1. CHIẾC MÁY TÍNH CỦA EM', chu_diem='Em và máy tính',
        tiet_ppct='1', ngay_day='   /   /2026',
        pham_chat_items=[
            '- Chăm chỉ: Tích cực học tập, hoàn thành bài tập đúng thời gian.',
            '- Trách nhiệm: Giữ gìn và sử dụng thiết bị đúng cách.',
            '- Nhân ái: Biết nhường bạn, hợp tác trong học nhóm.',
        ],
        nang_luc_mon_items=[
            '- NLa: Nhận biết và gọi tên được các bộ phận chính của máy tính (màn hình, bàn phím, chuột, thân máy).',
            '- NLb: Biết cách bật/tắt máy tính đúng quy trình.',
        ],
        nang_luc_chung_items=[
            '- Tự chủ và tự học: Chủ động thực hiện các thao tác được hướng dẫn.',
            '- Giao tiếp và hợp tác: Trao đổi, chia sẻ với bạn về các bộ phận máy tính.',
        ],
        do_dung_gv='Máy chiếu, máy tính demo, tranh poster các bộ phận máy tính.',
        do_dung_hs='SGK Tin học 1, bút chì.',
        phuong_phap='vấn đáp, quan sát trực quan, thực hành, hoạt động nhóm đôi',
        ki_thuat='đặt câu hỏi, tia chớp, chia sẻ nhóm đôi',
        activities=[
            {
                'ten': '1. Hoạt động MỞ ĐẦU (5 phút)',
                'muc_tieu': 'Kết nối kiến thức cũ, tạo hứng thú học bài mới.',
                'gv': [
                    'GV chiếu hình ảnh câu đố: "Đây là vật gì? Dùng để làm gì?"',
                    'GV mời HS trả lời, dẫn vào bài học.',
                ],
                'hs': [
                    '- HS quan sát, đoán và xung phong trả lời.',
                    '- HS cùng đọc tên bài học.',
                ],
            },
            {
                'ten': '2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI',
                'sub': [
                    {
                        'ten': '2.1. Khám phá các bộ phận của máy tính (10 phút)',
                        'muc_tieu': 'HS nhận biết và gọi tên đúng các bộ phận chính của máy tính.',
                        'gv': [
                            'GV chỉ từng bộ phận trên máy tính thật và màn chiếu.',
                            'GV gọi HS lên chỉ và đọc tên.',
                            '* Kết luận: Máy tính gồm màn hình, bàn phím, chuột và thân máy.',
                        ],
                        'hs': [
                            '- HS quan sát, lắng nghe.',
                            '- HS lần lượt lên chỉ và đọc tên bộ phận.',
                            '- HS đọc đồng thanh tên các bộ phận.',
                        ],
                    },
                    {
                        'ten': '2.2. Quy trình bật và tắt máy tính (5 phút)',
                        'muc_tieu': 'HS biết thứ tự các bước bật và tắt máy tính đúng cách.',
                        'gv': [
                            'GV thao tác mẫu bật máy, giải thích từng bước.',
                            'GV thao tác mẫu tắt máy đúng quy trình.',
                        ],
                        'hs': [
                            '- HS quan sát theo dõi từng bước.',
                            '- HS nhắc lại thứ tự bước bật/tắt máy.',
                        ],
                    },
                ],
            },
            {
                'ten': '3. HĐ LUYỆN TẬP-THỰC HÀNH',
                'sub': [
                    {
                        'ten': '3.1. Nối hình với tên (5 phút)',
                        'muc_tieu': 'Củng cố nhận biết tên các bộ phận máy tính.',
                        'gv': [
                            'GV phát phiếu bài tập: HS nối hình bộ phận với tên gọi đúng.',
                            '* GV chữa bài, nhận xét.',
                        ],
                        'hs': [
                            '- HS làm bài tập nối hình.',
                            '- HS đổi bài kiểm tra chéo, báo cáo kết quả.',
                        ],
                    },
                    {
                        'ten': '3.2. Thực hành bật/tắt máy (5 phút)',
                        'muc_tieu': 'HS thực hiện được thao tác bật/tắt máy đúng quy trình.',
                        'gv': [
                            'GV hướng dẫn HS thực hành theo từng bước.',
                            '* GV quan sát, hỗ trợ HS gặp khó khăn.',
                        ],
                        'hs': [
                            '- HS thực hành bật máy theo hướng dẫn.',
                            '- HS thực hành tắt máy đúng quy trình.',
                        ],
                    },
                ],
            },
            {
                'ten': '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                'gv': [
                    '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                    'GV hỏi: "Em có thể kể tên 3 bộ phận của máy tính cho bố mẹ nghe không?"',
                    'GV dặn: Về nhà quan sát máy tính ở gia đình.',
                ],
                'hs': [
                    '- HS thi đua kể tên bộ phận nhanh nhất.',
                    '- HS lắng nghe và ghi nhớ nhiệm vụ về nhà.',
                ],
            },
        ]
    )
    results.append(saved)

    # ────────────────────────────────────────────────────────────────────
    # LỚP 2 — Bài 1
    # ────────────────────────────────────────────────────────────────────
    print('\n🔵 Tạo KHBD Lớp 2 - Bài 1...')
    saved = create_khbd_th(
        grade_name='Lớp 2', grade_folder='Lớp_2',
        bai_so=1, ten_bai='BÀI 1. MÁY TÍNH LÀ NGƯỜI BẠN CỦA EM', chu_diem='Máy tính và cuộc sống',
        tiet_ppct='1', ngay_day='   /   /2026',
        pham_chat_items=[
            '- Chăm chỉ: Chủ động tìm hiểu và ghi nhớ các ứng dụng của máy tính.',
            '- Trách nhiệm: Sử dụng máy tính đúng mục đích, không lạm dụng.',
            '- Nhân ái: Biết chia sẻ kiến thức về máy tính với bạn bè.',
        ],
        nang_luc_mon_items=[
            '- NLa: Biết được máy tính có thể giúp em học tập, giải trí và kết nối với mọi người.',
            '- NLb: Nhận biết được một số ứng dụng phổ biến trên máy tính.',
        ],
        nang_luc_chung_items=[
            '- Tự chủ và tự học: Tự tìm hiểu và chia sẻ những điều em biết về máy tính.',
            '- Giao tiếp và hợp tác: Thảo luận nhóm về cách máy tính giúp ích cuộc sống.',
        ],
        do_dung_gv='Máy chiếu, video ngắn về ứng dụng máy tính trong cuộc sống, phiếu học tập.',
        do_dung_hs='SGK Tin học 2, bút màu.',
        phuong_phap='vấn đáp, xem video, thảo luận nhóm, trình bày',
        ki_thuat='động não, chia sẻ nhóm đôi, trình bày 1 phút',
        activities=[
            {
                'ten': '1. Hoạt động MỞ ĐẦU (5 phút)',
                'muc_tieu': 'Khơi gợi hiểu biết của HS về lợi ích của máy tính.',
                'gv': [
                    'GV hỏi: "Các em đã dùng máy tính để làm gì rồi?"',
                    'GV ghi nhanh các ý kiến lên bảng, dẫn vào bài.',
                ],
                'hs': [
                    '- HS tự do phát biểu.',
                    '- HS lắng nghe, bổ sung ý kiến của nhau.',
                ],
            },
            {
                'ten': '2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI',
                'sub': [
                    {
                        'ten': '2.1. Máy tính giúp em học tập (8 phút)',
                        'muc_tieu': 'HS biết được máy tính hỗ trợ học tập như thế nào.',
                        'gv': [
                            'GV chiếu video ngắn: HS dùng máy tính tra từ điển, đọc sách điện tử, làm bài tập online.',
                            'GV đặt câu hỏi: "Máy tính giúp em học tập như thế nào?"',
                            '* Kết luận: Máy tính là công cụ học tập hiệu quả.',
                        ],
                        'hs': [
                            '- HS xem video, quan sát.',
                            '- HS thảo luận nhóm đôi trả lời câu hỏi.',
                            '- Đại diện nhóm chia sẻ.',
                        ],
                    },
                    {
                        'ten': '2.2. Máy tính giúp em vui chơi và kết nối (7 phút)',
                        'muc_tieu': 'HS biết máy tính có thể dùng để giải trí và kết nối an toàn.',
                        'gv': [
                            'GV giới thiệu: máy tính dùng nghe nhạc, xem phim, chơi game học tập, video call.',
                            'GV nhấn mạnh: cần dùng đúng cách, đúng thời gian.',
                        ],
                        'hs': [
                            '- HS lắng nghe, ghi nhớ.',
                            '- HS chia sẻ: "Em đã dùng máy tính để làm gì ở nhà?"',
                        ],
                    },
                ],
            },
            {
                'ten': '3. HĐ LUYỆN TẬP-THỰC HÀNH',
                'sub': [
                    {
                        'ten': '3.1. Phân loại việc dùng máy tính (8 phút)',
                        'muc_tieu': 'HS phân biệt được dùng máy tính đúng cách và chưa đúng cách.',
                        'gv': [
                            'GV phát phiếu: HS phân loại các tình huống vào 2 cột Đúng/Chưa đúng.',
                            '* GV nhận xét, chốt đáp án.',
                        ],
                        'hs': [
                            '- HS đọc và phân loại tình huống.',
                            '- HS giải thích lý do lựa chọn.',
                        ],
                    },
                    {
                        'ten': '3.2. Vẽ máy tính của em (2 phút)',
                        'muc_tieu': 'HS thể hiện sự sáng tạo và ghi nhớ hình dạng thiết bị.',
                        'gv': ['GV hướng dẫn HS vẽ nhanh máy tính yêu thích vào vở.'],
                        'hs': ['- HS vẽ và tô màu.'],
                    },
                ],
            },
            {
                'ten': '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                'gv': [
                    '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                    'GV: "Em hãy kể cho bạn nghe 3 điều máy tính giúp ích cho em."',
                    'Dặn dò: Về nhà nói chuyện với bố mẹ về việc dùng máy tính đúng cách.',
                ],
                'hs': [
                    '- HS chia sẻ cặp đôi.',
                    '- Đại diện một số HS chia sẻ trước lớp.',
                ],
            },
        ]
    )
    results.append(saved)

    # ────────────────────────────────────────────────────────────────────
    # LỚP 3 — Bài 1: Thông tin và quyết định
    # ────────────────────────────────────────────────────────────────────
    print('\n🔵 Tạo KHBD Lớp 3 - Bài 1...')
    saved = create_khbd_th(
        grade_name='Lớp 3', grade_folder='Lớp_3',
        bai_so=1, ten_bai='BÀI 1. THÔNG TIN VÀ QUYẾT ĐỊNH', chu_diem='Thông tin quanh ta',
        tiet_ppct='1', ngay_day='   /   /2026',
        pham_chat_items=[
            '- Chăm chỉ: Chủ động tìm hiểu thông tin trong cuộc sống hằng ngày.',
            '- Trung thực: Nhận biết và không lan truyền thông tin sai lệch.',
            '- Trách nhiệm: Biết kiểm tra thông tin trước khi ra quyết định.',
        ],
        nang_luc_mon_items=[
            '- NLa (Sử dụng và quản lí phương tiện kĩ thuật số): Hiểu được vai trò của thông tin trong cuộc sống.',
            '- NLc (Giải quyết vấn đề với sự hỗ trợ của công nghệ thông tin): Biết thu thập và sử dụng thông tin để đưa ra quyết định đơn giản.',
            '- Năng lực số (CV 3456): Biết thông tin có nhiều dạng khác nhau (hình ảnh, âm thanh, văn bản).',
        ],
        nang_luc_chung_items=[
            '- Tự chủ và tự học: Thu thập thông tin từ nhiều nguồn khác nhau.',
            '- Giao tiếp và hợp tác: Chia sẻ thông tin và thảo luận quyết định cùng nhóm.',
            '- Giải quyết vấn đề và sáng tạo: Vận dụng thông tin để giải quyết tình huống thực tế.',
        ],
        do_dung_gv='Máy chiếu, tranh ảnh minh họa thông tin, phiếu bài tập tình huống.',
        do_dung_hs='SGK Tin học 3, bút, vở bài tập.',
        phuong_phap='vấn đáp, thảo luận nhóm, giải quyết vấn đề, trò chơi học tập',
        ki_thuat='đặt câu hỏi, động não, chia sẻ nhóm đôi',
        activities=[
            {
                'ten': '1. Hoạt động MỞ ĐẦU (5 phút)',
                'muc_tieu': 'Kích hoạt suy nghĩ về thông tin trong cuộc sống hằng ngày.',
                'gv': [
                    'GV kể tình huống: "Sáng nay, trước khi ra khỏi nhà, em xem dự báo thời tiết thấy trời sẽ mưa. Em quyết định mang ô theo. Tại sao em làm vậy?"',
                    'GV hỏi: "Dự báo thời tiết là loại thông tin gì? Nó giúp em quyết định điều gì?"',
                ],
                'hs': [
                    '- HS lắng nghe tình huống.',
                    '- HS phát biểu suy nghĩ.',
                    '- HS thảo luận cặp đôi, chia sẻ đáp án.',
                ],
            },
            {
                'ten': '2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI',
                'sub': [
                    {
                        'ten': '2.1. Thông tin là gì? (8 phút)',
                        'muc_tieu': 'HS nêu được khái niệm thông tin và các dạng thông tin.',
                        'gv': [
                            'GV chiếu hình ảnh: biển báo giao thông, tiếng chuông, sách vở, video.',
                            'GV hỏi: "Mỗi hình ảnh truyền đến em điều gì?"',
                            '* Kết luận: Thông tin là những gì cho chúng ta biết về sự vật, hiện tượng. Thông tin có 3 dạng: văn bản, hình ảnh, âm thanh.',
                        ],
                        'hs': [
                            '- HS quan sát từng hình ảnh.',
                            '- HS trả lời câu hỏi theo suy nghĩ.',
                            '- HS ghi nhớ 3 dạng thông tin.',
                        ],
                    },
                    {
                        'ten': '2.2. Thông tin giúp ta quyết định (7 phút)',
                        'muc_tieu': 'HS hiểu mối liên hệ giữa thông tin và quyết định.',
                        'gv': [
                            'GV nêu ví dụ: "Nhìn đèn giao thông đỏ → quyết định dừng lại."',
                            'GV phát phiếu tình huống: HS viết thông tin → quyết định.',
                            '* GV chốt: Thông tin đúng giúp ta quyết định đúng.',
                        ],
                        'hs': [
                            '- HS lắng nghe ví dụ.',
                            '- HS làm bài tập tình huống.',
                            '- HS chia sẻ đáp án.',
                        ],
                    },
                ],
            },
            {
                'ten': '3. HĐ LUYỆN TẬP-THỰC HÀNH',
                'sub': [
                    {
                        'ten': '3.1. Trò chơi "Thông tin – Quyết định" (8 phút)',
                        'muc_tieu': 'Củng cố hiểu biết về mối quan hệ thông tin-quyết định.',
                        'gv': [
                            'GV phát thẻ: mỗi nhóm nhận thẻ thông tin và thẻ quyết định, ghép đôi cho đúng.',
                            '* GV chữa bài, tuyên dương nhóm làm đúng và nhanh.',
                        ],
                        'hs': [
                            '- HS làm việc nhóm 4, ghép thẻ.',
                            '- Đại diện nhóm giải thích lý do ghép.',
                        ],
                    },
                    {
                        'ten': '3.2. Viết tình huống của em (2 phút)',
                        'muc_tieu': 'HS vận dụng kiến thức vào tình huống thực tế.',
                        'gv': ['GV yêu cầu HS tự viết 1 ví dụ thông tin → quyết định trong cuộc sống em.'],
                        'hs': ['- HS viết vào vở.', '- 2-3 HS đọc ví dụ của mình.'],
                    },
                ],
            },
            {
                'ten': '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                'gv': [
                    '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                    'GV: "Hôm nay em đã dùng thông tin nào để quyết định điều gì?"',
                    'Dặn dò: Quan sát và ghi lại 3 tình huống dùng thông tin để quyết định trong ngày mai.',
                ],
                'hs': [
                    '- HS suy nghĩ và chia sẻ.',
                    '- HS lắng nghe dặn dò.',
                ],
            },
        ]
    )
    results.append(saved)

    # ────────────────────────────────────────────────────────────────────
    # LỚP 4 — Bài 1: Phần cứng và phần mềm máy tính
    # ────────────────────────────────────────────────────────────────────
    print('\n🔵 Tạo KHBD Lớp 4 - Bài 1...')
    saved = create_khbd_th(
        grade_name='Lớp 4', grade_folder='Lớp_4',
        bai_so=1, ten_bai='BÀI 1. PHẦN CỨNG VÀ PHẦN MỀM MÁY TÍNH', chu_diem='Cấu tạo máy tính',
        tiet_ppct='1', ngay_day='   /   /2026',
        pham_chat_items=[
            '- Chăm chỉ: Tích cực tìm hiểu, ghi nhớ kiến thức về phần cứng và phần mềm.',
            '- Trách nhiệm: Sử dụng đúng cách, bảo quản thiết bị phần cứng.',
            '- Trung thực: Sử dụng phần mềm hợp pháp, không sao chép trái phép.',
        ],
        nang_luc_mon_items=[
            '- NLa: Phân biệt được phần cứng và phần mềm máy tính; liệt kê được ví dụ cụ thể.',
            '- NLc: Biết phần cứng và phần mềm phải phối hợp mới tạo nên hệ thống hoàn chỉnh.',
            '- Năng lực số (CV 3456): Hiểu được cơ bản về hệ thống máy tính (hardware/software).',
        ],
        nang_luc_chung_items=[
            '- Tự chủ và tự học: Phân loại thiết bị máy tính vào đúng nhóm phần cứng/mềm.',
            '- Giao tiếp và hợp tác: Thảo luận nhóm về các ví dụ phần cứng, phần mềm.',
        ],
        do_dung_gv='Máy chiếu, ảnh/video các thiết bị phần cứng, màn hình giới thiệu phần mềm.',
        do_dung_hs='SGK Tin học 4, phiếu bài tập.',
        phuong_phap='vấn đáp, quan sát, thảo luận nhóm, phân loại',
        ki_thuat='đặt câu hỏi, sơ đồ tư duy, chia sẻ nhóm',
        activities=[
            {
                'ten': '1. Hoạt động MỞ ĐẦU (5 phút)',
                'muc_tieu': 'Kích thích sự tò mò về cấu tạo của máy tính.',
                'gv': [
                    'GV đặt câu hỏi: "Máy tính gồm những gì? Em có thể cầm lấy phần mềm không?"',
                    'GV dẫn vào bài: Hôm nay chúng ta tìm hiểu sự khác nhau giữa phần cứng và phần mềm.',
                ],
                'hs': [
                    '- HS suy nghĩ và trả lời.',
                    '- HS tò mò và nêu câu hỏi.',
                ],
            },
            {
                'ten': '2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI',
                'sub': [
                    {
                        'ten': '2.1. Phần cứng là gì? (8 phút)',
                        'muc_tieu': 'HS định nghĩa và liệt kê được ví dụ phần cứng máy tính.',
                        'gv': [
                            'GV chiếu ảnh các thiết bị: CPU, RAM, ổ cứng, màn hình, bàn phím, chuột.',
                            'GV hỏi: "Điểm chung của những thứ này là gì?"',
                            '* Kết luận: Phần cứng là những bộ phận vật lí của máy tính — có thể nhìn thấy và chạm vào được.',
                        ],
                        'hs': [
                            '- HS quan sát hình ảnh.',
                            '- HS phát biểu nhận xét.',
                            '- HS ghi nhớ định nghĩa phần cứng.',
                        ],
                    },
                    {
                        'ten': '2.2. Phần mềm là gì? (7 phút)',
                        'muc_tieu': 'HS định nghĩa và liệt kê được ví dụ phần mềm.',
                        'gv': [
                            'GV mở máy tính, giới thiệu: Word, Chrome, Paint, hệ điều hành Windows.',
                            'GV hỏi: "Những thứ này khác phần cứng như thế nào?"',
                            '* Kết luận: Phần mềm là các chương trình chạy trên máy tính — không thể chạm vào.',
                        ],
                        'hs': [
                            '- HS quan sát màn hình máy tính.',
                            '- HS phát biểu sự khác biệt.',
                            '- HS cho thêm ví dụ phần mềm em biết.',
                        ],
                    },
                ],
            },
            {
                'ten': '3. HĐ LUYỆN TẬP-THỰC HÀNH',
                'sub': [
                    {
                        'ten': '3.1. Phân loại phần cứng/phần mềm (8 phút)',
                        'muc_tieu': 'Củng cố khả năng phân biệt phần cứng và phần mềm.',
                        'gv': [
                            'GV chiếu danh sách 10 ví dụ, HS xếp vào bảng phân loại.',
                            '* GV chữa bài, giải thích các trường hợp khó.',
                        ],
                        'hs': [
                            '- HS làm việc cá nhân, điền vào phiếu bài tập.',
                            '- HS so sánh kết quả với bạn cạnh.',
                        ],
                    },
                    {
                        'ten': '3.2. Sơ đồ tư duy (2 phút)',
                        'muc_tieu': 'HS tổng hợp kiến thức bài học bằng sơ đồ.',
                        'gv': ['GV hướng dẫn HS vẽ sơ đồ Máy tính → Phần cứng / Phần mềm.'],
                        'hs': ['- HS vẽ sơ đồ vào vở.'],
                    },
                ],
            },
            {
                'ten': '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                'gv': [
                    '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                    'GV: "Về nhà, em hãy quan sát và liệt kê 5 phần cứng và 5 phần mềm."',
                    'Dặn dò: Chia sẻ danh sách với bố mẹ, giải thích sự khác nhau.',
                ],
                'hs': [
                    '- HS lắng nghe nhiệm vụ về nhà.',
                    '- HS ghi vào vở.',
                ],
            },
        ]
    )
    results.append(saved)

    # ────────────────────────────────────────────────────────────────────
    # LỚP 5 — Bài 1: Em có thể làm gì với máy tính
    # ────────────────────────────────────────────────────────────────────
    print('\n🔵 Tạo KHBD Lớp 5 - Bài 1...')
    saved = create_khbd_th(
        grade_name='Lớp 5', grade_folder='Lớp_5',
        bai_so=1, ten_bai='BÀI 1. EM CÓ THỂ LÀM GÌ VỚI MÁY TÍNH', chu_diem='Máy tính và cuộc sống',
        tiet_ppct='1', ngay_day='   /   /2026',
        pham_chat_items=[
            '- Chăm chỉ: Tích cực khám phá các ứng dụng hữu ích của máy tính trong học tập.',
            '- Trách nhiệm: Sử dụng máy tính đúng mục đích, không lướt web vô ích.',
            '- Trung thực: Thành thật về thời gian và cách sử dụng máy tính.',
        ],
        nang_luc_mon_items=[
            '- NLa: Biết và sử dụng được máy tính cho các hoạt động học tập, sáng tạo.',
            '- NLb: Biết cách tìm kiếm thông tin học tập an toàn trên Internet.',
            '- NLe: Nhận biết được những nguy cơ khi sử dụng máy tính không đúng cách.',
            '- Năng lực số (CV 3456): Sử dụng thiết bị và phần mềm cơ bản phục vụ học tập và sáng tạo.',
        ],
        nang_luc_chung_items=[
            '- Tự chủ và tự học: Lên kế hoạch sử dụng máy tính hợp lí.',
            '- Giao tiếp và hợp tác: Chia sẻ kinh nghiệm sử dụng máy tính học tập với bạn.',
            '- Giải quyết vấn đề và sáng tạo: Đề xuất các cách dùng máy tính sáng tạo.',
        ],
        do_dung_gv='Máy chiếu, máy tính kết nối mạng, slide ví dụ ứng dụng máy tính.',
        do_dung_hs='SGK Tin học 5, vở ghi.',
        phuong_phap='thảo luận nhóm, trình bày, thực hành, giải quyết vấn đề',
        ki_thuat='động não, chia sẻ nhóm đôi, đặt câu hỏi',
        activities=[
            {
                'ten': '1. Hoạt động MỞ ĐẦU (5 phút)',
                'muc_tieu': 'Khởi động tư duy về các ứng dụng đa dạng của máy tính.',
                'gv': [
                    'GV đặt câu hỏi: "Trong 1 ngày, em đã dùng máy tính/điện thoại bao nhiêu lần và để làm gì?"',
                    'GV ghi nhanh ý kiến lên bảng, dẫn vào bài.',
                ],
                'hs': [
                    '- HS phát biểu tự do.',
                    '- HS so sánh câu trả lời của nhau.',
                ],
            },
            {
                'ten': '2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI',
                'sub': [
                    {
                        'ten': '2.1. Máy tính trong học tập và sáng tạo (10 phút)',
                        'muc_tieu': 'HS biết được các ứng dụng máy tính hỗ trợ học tập và sáng tạo.',
                        'gv': [
                            'GV trình chiếu: Soạn văn bản, vẽ tranh, làm slide, học tiếng Anh online, lập trình Scratch.',
                            'GV thao tác nhanh demo 1-2 ứng dụng.',
                            '* Kết luận: Máy tính là công cụ mạnh mẽ cho học tập và sáng tạo.',
                        ],
                        'hs': [
                            '- HS quan sát và đặt câu hỏi.',
                            '- HS chia sẻ ứng dụng em đã dùng.',
                        ],
                    },
                    {
                        'ten': '2.2. Sử dụng máy tính an toàn (5 phút)',
                        'muc_tieu': 'HS nhận biết và tuân thủ các quy tắc dùng máy tính an toàn.',
                        'gv': [
                            'GV nêu 5 quy tắc: Giới hạn thời gian, bảo vệ mật khẩu, không chia sẻ thông tin cá nhân, không xem nội dung không phù hợp, giữ tư thế đúng.',
                        ],
                        'hs': [
                            '- HS lắng nghe và ghi nhớ 5 quy tắc.',
                            '- HS cho ví dụ về vi phạm quy tắc.',
                        ],
                    },
                ],
            },
            {
                'ten': '3. HĐ LUYỆN TẬP-THỰC HÀNH',
                'sub': [
                    {
                        'ten': '3.1. Lập kế hoạch sử dụng máy tính (8 phút)',
                        'muc_tieu': 'HS tự lập được lịch sử dụng máy tính hợp lí.',
                        'gv': [
                            'GV phát mẫu lịch tuần, HS điền thời gian và mục đích dùng máy tính.',
                            '* GV nhận xét, tư vấn điều chỉnh cho hợp lí.',
                        ],
                        'hs': [
                            '- HS điền vào mẫu lịch.',
                            '- HS chia sẻ kế hoạch với bạn cạnh.',
                        ],
                    },
                    {
                        'ten': '3.2. Thực hành mở ứng dụng (2 phút)',
                        'muc_tieu': 'HS thực hành mở và đóng một ứng dụng trên máy tính.',
                        'gv': ['GV hướng dẫn HS mở Paint, vẽ một hình tự do, lưu file.'],
                        'hs': ['- HS thực hành trên máy tính.'],
                    },
                ],
            },
            {
                'ten': '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                'gv': [
                    '4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)',
                    'GV: "Em hãy chọn 1 ứng dụng trên máy tính để giúp em hoàn thành một bài tập ở nhà. Chia sẻ kết quả tuần sau."',
                    'Dặn dò: Áp dụng kế hoạch sử dụng máy tính đã lập, chia sẻ với bố mẹ.',
                ],
                'hs': [
                    '- HS lên kế hoạch cụ thể.',
                    '- HS lắng nghe và ghi lại nhiệm vụ.',
                ],
            },
        ]
    )
    results.append(saved)

    # ────────────────────────────────────────────────────────────────────
    # LỚP 6 — THCS — Bài 1: Thông tin và dữ liệu
    # ────────────────────────────────────────────────────────────────────
    print('\n🟠 Tạo KHBD Lớp 6 - Bài 1 (THCS)...')
    saved = create_khbd_thcs(
        grade_name='Lớp 6', grade_folder='Lớp_6',
        bai_so=1, ten_bai='Thông tin và dữ liệu',
        mon_hoc='Tin học', lop='6', thoi_luong='1 tiết (45 phút)', tiet_ppct='1',
        ngay_soan='   /   /2026', ngay_day='   /   /2026',
        kien_thuc_items=[
            '- Sự hiểu biết về khái niệm thông tin và dữ liệu trong máy tính.',
            '- Khả năng phân biệt thông tin (ý nghĩa con người hiểu) với dữ liệu (biểu diễn trong máy tính).',
            '- Sự nhận biết các dạng dữ liệu cơ bản: số, văn bản, hình ảnh, âm thanh, video.',
        ],
        nang_luc_chung_items=[
            '- Tự chủ và tự học: Chủ động tìm kiếm ví dụ về thông tin và dữ liệu trong cuộc sống. (Đạt được thông qua Hoạt động 1, 2)',
            '- Giao tiếp và hợp tác: Thảo luận nhóm, chia sẻ ví dụ và đưa ra kết luận chung. (Đạt được thông qua Hoạt động 2, 3)',
        ],
        nang_luc_dac_thu_items=[
            '- NLa (Hiểu biết máy tính): Mô tả được khái niệm thông tin và dữ liệu; phân biệt được các dạng dữ liệu. (Đạt được thông qua Hoạt động 2)',
            '- NLc (Giải quyết vấn đề): Áp dụng hiểu biết về dữ liệu để phân tích tình huống thực tế. (Đạt được thông qua Hoạt động 3)',
        ],
        nang_luc_so_items=[
            '- Năng lực 1.1 (Sử dụng thiết bị kĩ thuật số): Hiểu cách máy tính biểu diễn thông tin dưới dạng dữ liệu. (Đạt được thông qua Hoạt động 2)',
        ],
        pham_chat_items=[
            '- Chăm chỉ: Tích cực tìm hiểu và ghi nhớ khái niệm về thông tin và dữ liệu. (Thông qua Hoạt động 1, 2)',
            '- Trung thực: Đánh giá đúng giá trị của thông tin, không bịa đặt dữ liệu. (Thông qua Hoạt động 3)',
            '- Trách nhiệm: Sử dụng thông tin có trách nhiệm, không chia sẻ thông tin sai. (Thông qua Hoạt động 4)',
        ],
        thiet_bi='Máy tính GV, máy chiếu, bảng tương tác, phiếu học tập nhóm.',
        hoc_lieu='SGK Tin học 6 (trang 6-12), phiếu bài tập phân loại dữ liệu.',
        hoat_dong_list=[
            {
                'stt': 1,
                'ten': 'Khởi động (Xác định vấn đề/nhiệm vụ học tập/Mở đầu)',
                'muc_tieu': 'Giúp học sinh xác định được mối quan hệ giữa thông tin và dữ liệu trong thực tế.',
                'noi_dung': 'GV chiếu tình huống: "Em nhận được dãy số 36.5 — đây là dữ liệu gì? Có ý nghĩa gì?" HS thảo luận cặp đôi.',
                'san_pham': 'HS phát biểu được rằng dãy số 36.5 có thể là nhiệt độ cơ thể (thông tin: không sốt) hoặc điểm số.',
                'to_chuc': 'GV chiếu câu hỏi tình huống, yêu cầu HS thảo luận cặp đôi trong 2 phút.',
                'buoc1_gv': 'GV chiếu hình ảnh: dãy số "36.5" và đặt câu hỏi: "Đây là gì? Có nghĩa là gì trong từng ngữ cảnh?"',
                'buoc1_hs': 'HS quan sát câu hỏi và chuẩn bị thảo luận cặp đôi.',
                'buoc2_gv': 'GV quan sát các cặp thảo luận, hỗ trợ gợi ý nếu HS chưa nghĩ ra ví dụ.',
                'buoc2_hs': 'HS thảo luận cặp đôi (2 phút), mỗi cặp nêu ít nhất 1 cách hiểu khác nhau về dãy số 36.5.',
                'buoc3_gv': 'GV mời 2-3 cặp trình bày, ghi các cách hiểu lên bảng.',
                'buoc3_hs': '2-3 cặp HS trình bày, HS khác lắng nghe và nhận xét.',
                'buoc4_gv': 'GV nhận xét, nhấn mạnh: cùng 1 dữ liệu → nhiều thông tin khác nhau tùy ngữ cảnh → Dẫn vào bài.',
                'buoc4_hs': 'HS lắng nghe, ghi nhận câu kết luận ban đầu vào vở.',
            },
            {
                'stt': 2,
                'ten': 'Hình thành kiến thức mới/giải quyết vấn đề',
                'muc_tieu': 'HS phân biệt được thông tin và dữ liệu; nhận biết 5 dạng dữ liệu cơ bản.',
                'noi_dung': 'GV giảng khái niệm thông tin và dữ liệu qua SGK trang 6-8; HS quan sát ví dụ và phân loại.',
                'san_pham': 'HS ghi được vào vở: định nghĩa thông tin, dữ liệu và 5 dạng dữ liệu với ví dụ minh họa.',
                'to_chuc': 'GV trình bày kiến thức kết hợp đặt câu hỏi tương tác.',
                'buoc1_gv': 'GV chiếu slide khái niệm: Thông tin = ý nghĩa con người hiểu. Dữ liệu = biểu diễn thông tin trong máy tính.',
                'buoc1_hs': 'HS theo dõi slide, mở SGK tr.6 và gạch chân định nghĩa.',
                'buoc2_gv': 'GV hỏi tương tác: "Cho ví dụ 1 thông tin và dữ liệu tương ứng trong cuộc sống?"',
                'buoc2_hs': 'HS đọc SGK tr.6-8, gạch chân định nghĩa, ví dụ và chuẩn bị trả lời câu hỏi.',
                'buoc3_gv': 'GV gọi 3-4 HS trả lời và giới thiệu 5 dạng dữ liệu: số, văn bản, hình ảnh, âm thanh, video.',
                'buoc3_hs': 'HS trả lời cá nhân, quan sát 5 dạng dữ liệu và ghi ví dụ minh họa vào vở.',
                'buoc4_gv': 'GV tổng kết, chiếu bảng 5 dạng dữ liệu với ví dụ cụ thể.',
                'buoc4_hs': 'HS ghi vào vở định nghĩa thông tin, dữ liệu và bảng 5 dạng dữ liệu.',
            },
            {
                'stt': 3,
                'ten': 'Luyện tập',
                'muc_tieu': 'Củng cố khả năng phân biệt thông tin/dữ liệu và phân loại dạng dữ liệu.',
                'noi_dung': 'HS thực hành phân loại 10 ví dụ vào đúng dạng dữ liệu; giải thích lý do.',
                'san_pham': 'Phiếu học tập hoàn chỉnh với 10/10 ví dụ được phân loại đúng và có giải thích.',
                'to_chuc': 'Hoạt động cá nhân, sau đó đổi phiếu kiểm tra chéo.',
                'buoc1_gv': 'GV phát phiếu học tập gồm 10 ví dụ (ảnh chụp, file MP3, bảng Excel...) và yêu cầu HS phân loại.',
                'buoc1_hs': 'HS nhận phiếu, đọc qua các ví dụ và bắt đầu phân loại từng ví dụ.',
                'buoc2_gv': 'GV quan sát, nhắc nhở HS làm việc cá nhân và hỗ trợ HS gặp khó khăn.',
                'buoc2_hs': 'HS làm việc cá nhân trong 5 phút, phân loại và viết giải thích ngắn cho mỗi ví dụ.',
                'buoc3_gv': 'GV hướng dẫn đổi phiếu kiểm tra chéo và gọi 2 cặp báo cáo kết quả.',
                'buoc3_hs': 'HS đổi phiếu theo cặp, kiểm tra chéo và báo cáo số câu đúng.',
                'buoc4_gv': 'GV chiếu đáp án, giải thích các trường hợp khó và tổng kết điểm cần nhớ.',
                'buoc4_hs': 'HS đối chiếu đáp án, sửa phiếu và ghi lại những điểm còn chưa hiểu rõ.',
            },
            {
                'stt': 4,
                'ten': 'Mở rộng (Nhiệm vụ về nhà)',
                'muc_tieu': 'HS vận dụng kiến thức vào thực tế và chuẩn bị bài học tiếp theo.',
                'noi_dung': 'HS quan sát dữ liệu trong cuộc sống (điện thoại, máy tính ở nhà) và ghi lại 5 ví dụ về các dạng dữ liệu khác nhau.',
                'san_pham': 'Danh sách 5 ví dụ dữ liệu từ cuộc sống thực với chú thích dạng dữ liệu.',
                'buoc1_gv': 'GV giao nhiệm vụ: Quan sát điện thoại/máy tính ở nhà, tìm 5 ví dụ về 5 dạng dữ liệu khác nhau.',
                'buoc1_hs': 'HS ghi nhiệm vụ vào sổ tay và đặt câu hỏi nếu chưa rõ yêu cầu.',
                'buoc2_gv': 'GV nhắc: Ghi rõ tên dữ liệu, dạng và ý nghĩa (thông tin) của nó.',
                'buoc2_hs': 'HS ghi mẫu bài tập về nhà theo cấu trúc GV hướng dẫn.',
                'buoc3_gv': 'GV dặn: Đầu tiết sau sẽ chia sẻ kết quả và nhận xét chéo.',
                'buoc3_hs': 'HS lắng nghe và chuẩn bị hoàn thành bài trước tiết sau.',
                'buoc4_gv': 'GV tổng kết tiết học và nhắc HS tắt máy đúng quy trình.',
                'buoc4_hs': 'HS tắt máy đúng quy trình, thu dọn sách vở và rời phòng máy.',
            },
        ]
    )
    results.append(saved)

    # ────────────────────────────────────────────────────────────────────
    # LỚP 7 — THCS — Bài 1: Thiết bị vào - ra
    # ────────────────────────────────────────────────────────────────────
    print('\n🟠 Tạo KHBD Lớp 7 - Bài 1 (THCS)...')
    saved = create_khbd_thcs(
        grade_name='Lớp 7', grade_folder='Lớp_7',
        bai_so=1, ten_bai='Thiết bị vào - ra',
        mon_hoc='Tin học', lop='7', thoi_luong='1 tiết (45 phút)', tiet_ppct='1',
        ngay_soan='   /   /2026', ngay_day='   /   /2026',
        kien_thuc_items=[
            '- Sự hiểu biết về khái niệm thiết bị vào (Input devices) và thiết bị ra (Output devices).',
            '- Khả năng phân loại và liệt kê các thiết bị vào/ra phổ biến của máy tính.',
            '- Sự hiểu biết về chức năng cơ bản của từng loại thiết bị vào-ra.',
        ],
        nang_luc_chung_items=[
            '- Tự chủ và tự học: Tự tìm hiểu thêm về các thiết bị vào-ra mới trên thị trường. (Đạt được thông qua HĐ 2, 4)',
            '- Giao tiếp và hợp tác: Thảo luận nhóm phân loại thiết bị và trình bày kết quả. (Đạt được thông qua HĐ 2, 3)',
        ],
        nang_luc_dac_thu_items=[
            '- NLa (Hiểu biết máy tính): Phân loại chính xác các thiết bị vào-ra của máy tính. (Đạt được thông qua HĐ 2)',
            '- NLb (Kết nối và cộng tác): Mô tả chức năng của từng thiết bị trong hệ thống máy tính. (Đạt được thông qua HĐ 3)',
        ],
        nang_luc_so_items=[
            '- Năng lực 1.2 (Hiểu biết về thiết bị): Nhận biết và sử dụng đúng cách các thiết bị vào-ra phổ biến. (Đạt được thông qua HĐ 3)',
        ],
        pham_chat_items=[
            '- Chăm chỉ: Tích cực tham gia thảo luận nhóm và hoàn thành phiếu học tập. (Thông qua HĐ 2, 3)',
            '- Trách nhiệm: Sử dụng và bảo quản thiết bị vào-ra đúng cách. (Thông qua HĐ 4)',
        ],
        thiet_bi='Máy chiếu, máy tính, bàn phím, chuột, loa ngoài, tai nghe, webcam (nếu có).',
        hoc_lieu='SGK Tin học 7 (trang 6-10), phiếu học tập phân loại thiết bị.',
        hoat_dong_list=[
            {
                'stt': 1,
                'ten': 'Khởi động (Xác định vấn đề/nhiệm vụ học tập/Mở đầu)',
                'muc_tieu': 'Kích hoạt hiểu biết về các thiết bị máy tính và tạo nhu cầu học.',
                'noi_dung': 'GV đặt câu hỏi: "Kể tên tất cả thiết bị em thấy xung quanh máy tính trong phòng học này." HS quan sát và kể.',
                'san_pham': 'Danh sách 5-10 thiết bị xung quanh máy tính được HS liệt kê.',
                'to_chuc': 'Hoạt động cá nhân nhanh (2 phút) sau đó chia sẻ với lớp.',
                'buoc1_gv': 'GV yêu cầu HS quan sát phòng máy và liệt kê nhanh các thiết bị nhìn thấy trong 2 phút.',
                'buoc1_hs': 'HS quan sát xung quanh phòng máy và ghi nhanh tên thiết bị vào giấy nháp.',
                'buoc2_gv': 'GV theo dõi HS quan sát, nhắc nhở liệt kê tất cả thiết bị kể cả nhỏ như dây cáp.',
                'buoc2_hs': 'HS tiếp tục ghi chép, cố gắng liệt kê nhiều thiết bị nhất có thể.',
                'buoc3_gv': 'GV mời HS xung phong đọc danh sách, tổng hợp và ghi lên bảng.',
                'buoc3_hs': 'HS xung phong đọc danh sách, bổ sung nếu bạn còn thiếu.',
                'buoc4_gv': 'GV đặt câu hỏi: "Chúng ta sẽ phân loại những thiết bị này như thế nào?" → Dẫn vào bài.',
                'buoc4_hs': 'HS suy nghĩ cách phân loại và sẵn sàng vào nội dung bài học.',
            },
            {
                'stt': 2,
                'ten': 'Hình thành kiến thức mới/giải quyết vấn đề',
                'muc_tieu': 'HS phân biệt và định nghĩa được thiết bị vào và thiết bị ra.',
                'noi_dung': 'GV giảng khái niệm: thiết bị vào (đưa thông tin vào máy), thiết bị ra (đưa thông tin ra); HS làm việc nhóm phân loại.',
                'san_pham': 'HS ghi được định nghĩa và ví dụ 5 thiết bị vào, 5 thiết bị ra vào vở.',
                'to_chuc': 'GV giảng kết hợp hỏi đáp; sau đó nhóm 4 làm bài tập phân loại.',
                'buoc1_gv': 'GV chiếu slide: Thiết bị vào = đưa thông tin/lệnh vào máy (bàn phím, chuột, micro, camera). Thiết bị ra = đưa kết quả ra ngoài (màn hình, máy in, loa).',
                'buoc1_hs': 'HS theo dõi slide và đối chiếu với danh sách thiết bị vừa liệt kê.',
                'buoc2_gv': 'GV hỏi từng thiết bị trong danh sách: "Thiết bị này là vào hay ra? Tại sao?"',
                'buoc2_hs': 'HS đọc SGK tr.6-8, gạch chân ví dụ, trả lời câu hỏi GV đặt ra.',
                'buoc3_gv': 'GV giao phiếu học tập 10 thiết bị cho nhóm 4, yêu cầu phân loại và cử đại diện trình bày.',
                'buoc3_hs': 'Nhóm 4 thảo luận phân loại 10 thiết bị, đại diện trình bày kết quả.',
                'buoc4_gv': 'GV nhận xét, giải thích trường hợp đặc biệt (màn hình cảm ứng = vừa vào vừa ra).',
                'buoc4_hs': 'HS ghi vở định nghĩa, ví dụ 5 thiết bị vào và 5 thiết bị ra.',
            },
            {
                'stt': 3,
                'ten': 'Luyện tập',
                'muc_tieu': 'Củng cố kĩ năng phân loại thiết bị vào-ra; liên hệ thực tế.',
                'noi_dung': 'HS thực hành phân loại thiết bị và kết hợp thực hành sử dụng bàn phím, chuột.',
                'san_pham': 'Phiếu bài tập hoàn chỉnh; HS gõ đúng đoạn văn mẫu bằng bàn phím.',
                'to_chuc': 'Kết hợp bài tập lý thuyết và thực hành máy tính.',
                'buoc1_gv': 'GV phát phiếu bài tập 10 câu phân loại thiết bị + 2 câu tình huống; hướng dẫn HS mở máy thực hành gõ văn bản.',
                'buoc1_hs': 'HS nhận phiếu, bắt đầu làm bài tập lý thuyết và mở máy thực hành.',
                'buoc2_gv': 'GV quan sát HS làm bài và thực hành, hỗ trợ HS gặp khó khăn với bàn phím.',
                'buoc2_hs': 'HS làm bài tập lý thuyết (5 phút) và thực hành gõ bàn phím (5 phút).',
                'buoc3_gv': 'GV hướng dẫn HS đổi phiếu kiểm tra chéo và chiếu đáp án.',
                'buoc3_hs': 'HS đổi phiếu kiểm tra chéo, báo cáo số câu đúng và kết quả gõ bàn phím.',
                'buoc4_gv': 'GV nhận xét kết quả thực hành, nhấn mạnh tư thế ngồi và cách đặt tay đúng.',
                'buoc4_hs': 'HS lắng nghe, điều chỉnh tư thế ngồi và cách đặt tay đúng.',
            },
            {
                'stt': 4,
                'ten': 'Mở rộng (Nhiệm vụ về nhà)',
                'muc_tieu': 'HS tìm hiểu thêm về các thiết bị vào-ra hiện đại.',
                'noi_dung': 'HS tìm hiểu 2 thiết bị vào-ra hiện đại (VR headset, bút stylus, máy in 3D...) và mô tả chức năng.',
                'san_pham': 'Đoạn văn ngắn (5-7 câu) mô tả 2 thiết bị vào-ra hiện đại em tìm được.',
                'buoc1_gv': 'GV giao nhiệm vụ: Tìm hiểu 2 thiết bị vào-ra hiện đại chưa học trong bài, viết mô tả ngắn 5-7 câu.',
                'buoc1_hs': 'HS ghi nhiệm vụ vào sổ tay và hỏi nếu chưa rõ yêu cầu.',
                'buoc2_gv': 'GV gợi ý: VR headset, bút stylus, máy in 3D, găng tay haptic...',
                'buoc2_hs': 'HS ghi gợi ý và bắt đầu nghĩ về thiết bị sẽ tìm hiểu.',
                'buoc3_gv': 'GV nhắc: Nộp bài đầu tiết học sau, bài hay sẽ được chia sẻ với lớp.',
                'buoc3_hs': 'HS chuẩn bị hoàn thành bài tập trước tiết sau.',
                'buoc4_gv': 'GV gợi ý từ khóa tìm kiếm: "thiết bị vào ra hiện đại 2024" và kết thúc tiết.',
                'buoc4_hs': 'HS ghi từ khóa tìm kiếm, tắt máy đúng quy trình và rời phòng.',
            },
        ]
    )
    results.append(saved)

    # ────────────────────────────────────────────────────────────────────
    # LỚP 8 — THCS — Bài 1: Lược sử công cụ tính toán
    # ────────────────────────────────────────────────────────────────────
    print('\n🟠 Tạo KHBD Lớp 8 - Bài 1 (THCS)...')
    saved = create_khbd_thcs(
        grade_name='Lớp 8', grade_folder='Lớp_8',
        bai_so=1, ten_bai='Lược sử công cụ tính toán',
        mon_hoc='Tin học', lop='8', thoi_luong='1 tiết (45 phút)', tiet_ppct='1',
        ngay_soan='   /   /2026', ngay_day='   /   /2026',
        kien_thuc_items=[
            '- Sự hiểu biết về các mốc lịch sử phát triển công cụ tính toán từ thủ công đến hiện đại.',
            '- Khả năng mô tả đặc điểm và vai trò của từng thế hệ máy tính (Thế hệ 1-5).',
            '- Sự nhận biết xu hướng phát triển công nghệ tính toán hiện nay và tương lai.',
        ],
        nang_luc_chung_items=[
            '- Tự chủ và tự học: Chủ động tìm kiếm và hệ thống hóa kiến thức về lịch sử máy tính. (Đạt được thông qua HĐ 2, 4)',
            '- Giao tiếp và hợp tác: Thảo luận và thuyết trình về lịch sử phát triển công nghệ. (Đạt được thông qua HĐ 2, 3)',
        ],
        nang_luc_dac_thu_items=[
            '- NLa (Hiểu biết máy tính): Trình bày được lịch sử phát triển máy tính theo đúng thứ tự thời gian. (Đạt được thông qua HĐ 2)',
            '- NLe (Ứng xử phù hợp): Nhận thức được tác động của công nghệ máy tính đến cuộc sống xã hội. (Đạt được thông qua HĐ 4)',
        ],
        nang_luc_so_items=[
            '- Năng lực 5.1 (Hiểu về lịch sử công nghệ): Mô tả được sự phát triển của công nghệ số qua các thế hệ. (Đạt được thông qua HĐ 2, 3)',
        ],
        pham_chat_items=[
            '- Chăm chỉ: Tích cực nghiên cứu và ghi chép thông tin lịch sử máy tính. (Thông qua HĐ 2)',
            '- Trách nhiệm: Trân trọng thành tựu khoa học và có ý thức ứng dụng công nghệ đúng đắn. (Thông qua HĐ 4)',
        ],
        thiet_bi='Máy chiếu, máy tính, video lược sử máy tính (2-3 phút), tranh ảnh minh họa các thế hệ máy tính.',
        hoc_lieu='SGK Tin học 8 (trang 6-14), phiếu học tập dòng thời gian.',
        hoat_dong_list=[
            {
                'stt': 1,
                'ten': 'Khởi động (Xác định vấn đề/nhiệm vụ học tập/Mở đầu)',
                'muc_tieu': 'Tạo hứng thú và kích hoạt tư duy về lịch sử phát triển công nghệ tính toán.',
                'noi_dung': 'GV chiếu 2 hình ảnh: chiếc bàn tính abacus cổ đại và siêu máy tính hiện đại. Hỏi: "Có điểm gì giống nhau giữa 2 thiết bị này?"',
                'san_pham': 'HS nêu được ít nhất 1 điểm chung (đều dùng để tính toán) và nêu được sự khác biệt về công nghệ.',
                'to_chuc': 'Hoạt động quan sát và thảo luận cặp đôi (2 phút).',
                'buoc1_gv': 'GV chiếu 2 hình ảnh (abacus và siêu máy tính) và đặt câu hỏi tư duy.',
                'buoc1_hs': 'HS thực hiện theo yêu cầu.',
                'buoc2_gv': 'HS quan sát, thảo luận cặp đôi tìm điểm giống và khác nhau.',
                'buoc2_hs': 'HS thực hiện theo yêu cầu.',
                'buoc3_gv': '3 cặp HS trình bày. GV ghi ý kiến lên bảng.',
                'buoc3_hs': 'HS thực hiện theo yêu cầu.',
                'buoc4_gv': 'GV dẫn vào bài: "Hành trình từ abacus đến siêu máy tính là hành trình hàng ngàn năm. Hôm nay chúng ta cùng khám phá lược sử đó."',
                'buoc4_hs': 'HS thực hiện theo yêu cầu.',
            },
            {
                'stt': 2,
                'ten': 'Hình thành kiến thức mới/giải quyết vấn đề',
                'muc_tieu': 'HS nắm được các mốc chính trong lịch sử phát triển công cụ tính toán từ cổ đại đến hiện đại.',
                'noi_dung': 'GV trình bày lịch sử qua timeline: Abacus → Pascaline → ENIAC → Microprocessor → AI/Cloud. HS điền vào phiếu dòng thời gian.',
                'san_pham': 'Phiếu dòng thời gian hoàn chỉnh với 5 mốc lịch sử, đặc điểm và năm ra đời.',
                'to_chuc': 'GV trình bày timeline kết hợp video ngắn; HS điền phiếu đồng thời.',
                'buoc1_gv': 'GV chiếu timeline và video lược sử máy tính (2 phút). Phát phiếu dòng thời gian cho HS.',
                'buoc1_hs': 'HS thực hiện theo yêu cầu.',
                'buoc2_gv': 'HS theo dõi video, đọc SGK tr.6-10, điền vào phiếu dòng thời gian: tên, năm, đặc điểm.',
                'buoc2_hs': 'HS thực hiện theo yêu cầu.',
                'buoc3_gv': 'HS so sánh phiếu với bạn cạnh, bổ sung thiếu sót. GV gọi 1 HS đọc timeline của mình.',
                'buoc3_hs': 'HS thực hiện theo yêu cầu.',
                'buoc4_gv': 'GV chốt đáp án, giải thích thêm về ENIAC và chiếc máy tính đầu tiên dùng transistor. HS sửa phiếu.',
                'buoc4_hs': 'HS thực hiện theo yêu cầu.',
            },
            {
                'stt': 3,
                'ten': 'Luyện tập',
                'muc_tieu': 'Củng cố kiến thức về lịch sử máy tính và liên hệ thực tế.',
                'noi_dung': 'HS thực hành bài tập sắp xếp thứ tự thời gian và câu hỏi liên hệ về công nghệ hiện đại.',
                'san_pham': 'Phiếu bài tập hoàn chỉnh và đoạn văn ngắn liên hệ thực tế.',
                'to_chuc': 'Bài tập cá nhân sau đó thảo luận nhóm.',
                'buoc1_gv': 'GV phát phiếu bài tập: Sắp xếp 10 sự kiện theo đúng thứ tự lịch sử; 2 câu liên hệ.',
                'buoc1_hs': 'HS thực hiện theo yêu cầu.',
                'buoc2_gv': 'HS làm bài cá nhân (7 phút). GV quan sát, gợi ý HS gặp khó khăn.',
                'buoc2_hs': 'HS thực hiện theo yêu cầu.',
                'buoc3_gv': 'HS đổi bài kiểm tra chéo. Nhóm 4 thảo luận câu hỏi liên hệ: "Theo em, 20 năm nữa máy tính sẽ như thế nào?"',
                'buoc3_hs': 'HS thực hiện theo yêu cầu.',
                'buoc4_gv': 'Đại diện 2 nhóm trình bày dự đoán. GV nhận xét, giới thiệu xu hướng AI và Quantum Computing.',
                'buoc4_hs': 'HS thực hiện theo yêu cầu.',
            },
            {
                'stt': 4,
                'ten': 'Mở rộng (Nhiệm vụ về nhà)',
                'muc_tieu': 'HS tìm hiểu sâu hơn về một nhân vật lịch sử hoặc mốc phát triển máy tính mà em quan tâm.',
                'noi_dung': 'HS chọn 1 trong số: Alan Turing, Grace Hopper, Steve Jobs, hay sự ra đời của Internet; viết đoạn thuyết trình 7-10 câu.',
                'san_pham': 'Đoạn văn 7-10 câu về nhân vật/mốc lịch sử được chọn, nộp đầu tiết học sau.',
                'buoc1_gv': 'GV giới thiệu 4 chủ đề lựa chọn và giao nhiệm vụ.',
                'buoc1_hs': 'HS thực hiện theo yêu cầu.',
                'buoc2_gv': 'HS ghi nhiệm vụ vào sổ tay, chọn chủ đề yêu thích.',
                'buoc2_hs': 'HS thực hiện theo yêu cầu.',
                'buoc3_gv': 'Tuần sau HS nộp bài, GV chọn 3-4 bài hay đọc trước lớp.',
                'buoc3_hs': 'HS thực hiện theo yêu cầu.',
                'buoc4_gv': 'GV gợi ý nguồn tham khảo: Britannica, Wikipedia tiếng Việt, kênh YouTube khoa học.',
                'buoc4_hs': 'HS thực hiện theo yêu cầu.',
            },
        ]
    )
    results.append(saved)

    return results


# ════════════════════════════════════════════════════════════════════════════
# MAIN
# ════════════════════════════════════════════════════════════════════════════
if __name__ == '__main__':
    print('🚀 Bắt đầu tạo toàn bộ KHBD Bài 1 (Tiền TH → Lớp 8)...\n')
    results = build_all()
    print(f'\n✅ HOÀN TẤT! Đã tạo {len(results)} file KHBD.')
    for r in results:
        print(f'  📄 {r}')
