import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn

files_to_check = [
    r'd:\UNIGO\KHBD_Tin_học\Lớp_7\Tuần_03\KHBD_Tin_hoc_Lớp_7_Bai02_Bai_2_Phan_mem_may_tinh.docx',
    r'd:\UNIGO\KHBD_Tin_học\Lớp_7\Tuần_02\KHBD_Tin_hoc_Lớp_7_Bai01_Thiết_bị_vào_-_ra.docx',
    r'd:\UNIGO\KHBD_Tin_học\Lớp_7\Tuần_01\KHBD_Tin_hoc_Lớp_7_Tiet00_Định_hướng_môn_học_Tin_học_7_-_Tổng_quan.docx',
    r'd:\UNIGO\KHBD_Tin_học\Lớp_6\Tuần_03\KHBD_Tin_hoc_Lớp_6_Bai02_Bai_2_Xu_li_thong_tin.docx',
    r'd:\UNIGO\KHBD_Tin_học\Lớp_8\Tuần_03\KHBD_Tin_hoc_Lớp_8_Bai02_Bai_2_Thong_tin_trong_moi_truong_so.docx',
]

for p in files_to_check:
    print(f"\n=== {os.path.basename(p)} ===")
    if not os.path.exists(p):
        print("Not found")
        continue
    doc = Document(p)
    print("Table count:", len(doc.tables))
    if doc.tables:
        t0 = doc.tables[0]
        print("Table 0 style:", t0.style.name if t0.style else "None")
        tblPr = t0._tbl.find(qn('w:tblPr'))
        borders = tblPr.find(qn('w:tblBorders')) if tblPr is not None else None
        if borders is not None:
            for b in borders:
                print(" ", b.tag.split('}')[-1], b.attrib)
        else:
            print("  No tblBorders in tblPr")
    for i, para in enumerate(doc.paragraphs[:6]):
        print(f"  P[{i}]: {para.text[:80]}")
