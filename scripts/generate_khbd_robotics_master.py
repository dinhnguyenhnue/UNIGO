"""
generate_khbd_robotics_master.py
================================
Tự động tạo toàn bộ KHBD Robotics (Lớp 1 -> Lớp 8) chuẩn 100% theo quy chuẩn UNIGO & LBG:
1. ĐỒNG BỘ TUẦN DẠY VỚI LỊCH BÁO GIẢNG (LBG):
   - Tuần 1: Tiết 0: Định hướng môn học (Tuần_01)
   - Lớp 5, 6, 7, 8 (Rotation Tuần Lẻ: 2 tiết/tuần):
       * Tuần 3: Tiết 1 (Bài 1), Tiết 2 (Bài 2)
       * Tuần 5: Tiết 3 (Bài 3), Tiết 4 (Bài 4)
       * Tuần 7: Tiết 5 (Ôn tập ĐGĐK1), Tiết 6 (ĐGĐK1)
       * Tuần 9: Tiết 7 (Bài 5), Tiết 8 (Bài 6)
       * Tuần 11: Tiết 9 (Bài 7), Tiết 10 (Bài 8)
       * Tuần 13: Tiết 11 (Ôn tập ĐGĐK2), Tiết 12 (ĐGĐK2)
       * Tuần 15: Tiết 13 (Bài 9), Tiết 14 (Bài 10)
       * Tuần 17: Tiết 15 (Bài 11), Tiết 16 (Bài 12)
       * Tuần 19: Tiết 17 (Ôn tập ĐGĐK3), Tiết 18 (ĐGĐK3)
       * Tuần 21: Tiết 19 (Bài 13), Tiết 20 (Luyện tập/Thực hành sáng tạo)
       * Tuần 23: Tiết 21 (Ôn tập ĐGĐK4), Tiết 22 (ĐGĐK4)
       * Tuần 25: Tiết 23 (Tổng kết môn học & Triển lãm Robotics)
   - Lớp 1, 2, 3, 4 (Học 1 tiết/tuần, từ Tuần 2 đến Tuần 35):
       * Tuần 2: Tiết 1 (Bài 1)
       * Tuần 3: Tiết 2 (Bài 2)
       * Tuần 4: Tiết 3 (Bài 3)
       * ...
       * Tuần N: Tiết N-1 (Bài N-1)

2. QUY CHUẨN ĐỊNH DẠNG THEO AGENTS.MD:
   - Font: Times New Roman 13pt toàn bộ
   - Table 0 (Bảng thông tin 3x2): KHÔNG VIỀN (NO BORDER)
   - Bố cục tiêu đề bài: Chuẩn theo cấp TH và THCS
   - Kiến thức: Danh từ / Cụm danh từ trực tiếp (CẤM "Sự hiểu biết về...", "Khả năng nhận diện...")
   - Năng lực đặc thù Robotics: NL1 -> NL5 theo chủ đề bài học, kèm (#Hoạt động)
   - Năng lực số: CV 3456 chuẩn Bậc (L1-3=Bậc1, L4-5=Bậc2, L6-7=Bậc3, L8=Bậc4)
   - Bảng Hoạt động (Tiến trình dạy học 2 cột): CÓ VIỀN
   - Rút kinh nghiệm & Bảng chữ ký 3x3: KHÔNG VIỀN (NO BORDER)
"""

import os
import re
import sys
import shutil
from datetime import date, timedelta
import docx
from docx import Document
from docx.shared import Pt, Cm, Emu, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')

# ─── PATH CONFIG ───────────────────────────────────────────────────────────
PPCT_ROB = r"D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch dạy học môn Robotics 2026-2027.docx"
TPL_DOC = r"D:\UNIGO\Hệ thống mẫu văn bản\PL4-Khung kế hoạch bài dạy (THCS).docx"
OUT_BASE = r"D:\UNIGO\KHBD_Robotics"

TUAN_01_START = date(2026, 8, 3)  # Thứ Hai Tuần 1

# Schedule mapping: (day_of_week_idx, ten_lop)
# day_of_week_idx: 0=Thứ Hai, 1=Thứ Ba, 2=Thứ Tư, 3=Thứ Năm, 4=Thứ Sáu
ROB_SCHEDULE = {
    '1': (3, '1A1'),  # Thứ Năm sáng (T2)
    '2': (2, '2A1'),  # Thứ Tư chiều (T4)
    '3': (2, '3A1'),  # Thứ Tư sáng (T2)
    '4': (0, '4C1'),  # Thứ Hai sáng (T3)
    '5': (1, '5C1'),  # Thứ Ba sáng (T3-T4)
    '6': (4, '6A1'),  # Thứ Sáu sáng (T4-T5)
    '7': (1, '7A1'),  # Thứ Ba chiều (T3-T4)
    '8': (4, '8A1'),  # Thứ Sáu chiều (T1-T2)
}

FONT_NAME = "Times New Roman"
FONT_SIZE_PT = 13

# EMU Indents
INDENT_0 = 0
INDENT_1 = 180340       # ~0.5cm (Mục con cấp 1)
INDENT_2 = 360045       # ~1.0cm (Mục con cấp 2)
INDENT_BULLET = 540000   # ~1.5cm (Nội dung gạch đầu dòng)
INDENT_TH_1 = 457200    # ~1.27cm
INDENT_TH_2 = 450215    # ~1.25cm


# ─── FORMATTING HELPERS ────────────────────────────────────────────────────
def afont(run, bold=False, italic=False, size_pt=FONT_SIZE_PT):
    run.font.name = FONT_NAME
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), FONT_NAME)
    rFonts.set(qn('w:hAnsi'), FONT_NAME)
    rFonts.set(qn('w:cs'), FONT_NAME)
    rFonts.set(qn('w:eastAsia'), FONT_NAME)


def clean_body(doc):
    for child in list(doc.element.body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'sectPr':
            doc.element.body.remove(child)


def add_p(doc, text="", bold=False, italic=False, first_indent=None,
          left_indent=None, alignment=WD_ALIGN_PARAGRAPH.JUSTIFY,
          space_after_pt=3, space_before_pt=0, line_spacing=1.15, size_pt=FONT_SIZE_PT):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = Pt(space_after_pt)
    pf.space_before = Pt(space_before_pt)
    if first_indent is not None:
        pf.first_line_indent = Emu(first_indent)
    if left_indent is not None:
        pf.left_indent = Emu(left_indent)
    if text:
        run = p.add_run(text)
        afont(run, bold=bold, italic=italic, size_pt=size_pt)
    return p


def add_multi_run(doc, runs_data, first_indent=None, left_indent=None,
                  alignment=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.15):
    p = doc.add_paragraph()
    p.alignment = alignment
    pf = p.paragraph_format
    pf.line_spacing = line_spacing
    pf.space_after = Pt(3)
    pf.space_before = Pt(0)
    if first_indent is not None:
        pf.first_line_indent = Emu(first_indent)
    if left_indent is not None:
        pf.left_indent = Emu(left_indent)
    for text, bold, italic in runs_data:
        run = p.add_run(text)
        afont(run, bold=bold, italic=italic)
    return p


def set_borders(table, val="single", sz="4", color="000000"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    borders_el = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), val)
        b.set(qn('w:sz'), sz)
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        borders_el.append(b)
    tblPr.append(borders_el)


def set_no_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    borders_el = OxmlElement('w:tblBorders')
    for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'nil')
        borders_el.append(b)
    tblPr.append(borders_el)


def set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, v in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)


def fill_cell(cell, text, bold=False, italic=False,
              align=WD_ALIGN_PARAGRAPH.LEFT, size_pt=FONT_SIZE_PT,
              space_after=2, line_spacing=1.15):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)
    set_cell_margins(cell)
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            p = cell.add_paragraph()
            p.alignment = align
            p.paragraph_format.line_spacing = line_spacing
            p.paragraph_format.space_after = Pt(space_after)
        run = p.add_run(line)
        afont(run, bold=bold, italic=italic, size_pt=size_pt)


def sanitize(name):
    import unicodedata
    name = unicodedata.normalize('NFD', name).encode('ascii', 'ignore').decode('utf-8')
    name = re.sub(r'[^a-zA-Z0-9_\-]', '_', name)
    name = re.sub(r'_+', '_', name).strip('_')
    return name


def compute_dates(tuan_so, day_of_week):
    week_start = TUAN_01_START + timedelta(weeks=tuan_so - 1)
    ngay_day = week_start + timedelta(days=day_of_week)
    ngay_soan = week_start - timedelta(days=2)  # Thứ 7 tuần trước
    return ngay_soan.strftime('%d/%m/%Y'), ngay_day.strftime('%d/%m/%Y')


def get_bac_nls(grade):
    if grade <= 3:
        return 1
    elif grade <= 5:
        return 2
    elif grade <= 7:
        return 3
    else:
        return 4


def get_kit_name(grade):
    if grade in [1, 2]:
        return "OLLO Kinder"
    elif grade in [3, 4]:
        return "OLLO Initiate"
    else:
        return "OLLO Excel 1"


def get_rob_competencies(grade, title, yccd="", kit_name=""):
    """
    Map Robotics lesson to appropriate competencies (NL1 - NL5)
    with specific descriptors based on CT ROBOHUB.
    """
    t_lower = (title + " " + yccd).lower()
    
    if any(k in t_lower for k in ['định hướng', 'tiết 0']):
        c1 = f"- NL1 (Nhận thức công nghệ): Nhận biết tổng quan chương trình môn học, an toàn sử dụng bộ kit {kit_name} và nội quy phòng học Robotics. (Đạt được thông qua Hoạt động 1, Hoạt động 2)"
        c2 = f"- NL2 (Sử dụng công nghệ): Tuân thủ quy tắc làm việc nhóm, bảo quản linh kiện và sử dụng thiết bị đúng hướng dẫn. (Đạt được thông qua Hoạt động 3, Hoạt động 4)"
        return [c1, c2]

    if any(k in t_lower for k in ['triển lãm', 'tổng kết', 'trình bày', 'chia sẻ', 'giới thiệu', 'giao lưu']):
        c1 = f"- NL5 (Giao tiếp công nghệ): Tự tin thuyết minh nguyên lý hoạt động và giới thiệu sản phẩm robot {title} trước tập thể. (Đạt được thông qua Hoạt động 3, Hoạt động 4)"
        c2 = f"- NL4 (Đánh giá công nghệ): Lắng nghe, nhận xét và đánh giá sản phẩm robot của nhóm bạn theo tiêu chí kỹ thuật. (Đạt được thông qua Hoạt động 4)"
        return [c1, c2]
        
    elif any(k in t_lower for k in ['thử nghiệm', 'đo', 'so sánh', 'lực', 'quán tính', 'tốc độ', 'góc', 'vận tốc', 'gia tốc', 'lực hấp dẫn', 'đàn hồi', 'ném', 'chuyển động']):
        c1 = f"- NL4 (Đánh giá công nghệ): Quan sát, đo đạc thông số và đánh giá hiệu quả vận hành của mô hình robot {title} trong các điều kiện thử nghiệm khác nhau. (Đạt được thông qua Hoạt động 3, Hoạt động 4)"
        c2 = f"- NL3 (Thiết kế kĩ thuật): Điều chỉnh kết cấu cơ khí và thông số điều khiển để tối ưu hóa khả năng hoạt động của robot. (Đạt được thông qua Hoạt động 3)"
        return [c1, c2]

    elif any(k in t_lower for k in ['sáng tạo', 'cải tiến', 'thiết kế', 'cánh tay', 'kẹp', 'gắp', 'tự động', 'domino', 'nhiệm vụ', 'hút bụi']):
        c1 = f"- NL3 (Thiết kế kĩ thuật): Lắp ráp, phối hợp các cơ cấu truyền động và sáng tạo cải tiến mô hình robot {title} thực hiện nhiệm vụ đặt ra. (Đạt được thông qua Hoạt động 3, Hoạt động 4)"
        c2 = f"- NL2 (Sử dụng công nghệ): Sử dụng thành thạo các module, khớp nối và linh kiện bộ kit {kit_name} theo đúng tiêu chuẩn an toàn. (Đạt được thông qua Hoạt động 2, Hoạt động 3)"
        return [c1, c2]

    elif any(k in t_lower for k in ['động cơ', 'cảm biến', 'âm thanh', 'hồng ngoại', 'dynamixel', 'bộ điều khiển', 'linh kiện', 'chốt', 'khung']):
        c1 = f"- NL1 (Nhận thức công nghệ): Nhận biết chính xác tên gọi, cấu tạo và vai trò của động cơ/cảm biến/bộ điều khiển trong mô hình {title}. (Đạt được thông qua Hoạt động 2)"
        c2 = f"- NL2 (Sử dụng công nghệ): Thao tác đấu nối đúng kỹ thuật, kiểm tra tín hiệu đầu vào/ra và vận hành robot an toàn. (Đạt được thông qua Hoạt động 3)"
        return [c1, c2]

    elif any(k in t_lower for k in ['ôn tập', 'đánh giá định kỳ']):
        c1 = f"- NL1 (Nhận thức công nghệ): Hệ thống hóa kiến thức về cấu tạo, nguyên lý hoạt động của các mô hình robot đã học. (Đạt được thông qua Hoạt động 2)"
        c2 = f"- NL2 (Sử dụng công nghệ): Thực hành thành thạo các thao tác lắp ráp và vận hành robot theo yêu cầu kiểm tra. (Đạt được thông qua Hoạt động 3)"
        return [c1, c2]

    else:
        c1 = f"- NL1 (Nhận thức công nghệ): Nhận biết mô hình {title} mô phỏng sự vật/hiện tượng thực tế và hiểu nguyên lý hoạt động cơ bản. (Đạt được thông qua Hoạt động 2)"
        c2 = f"- NL2 (Sử dụng công nghệ): Thực hiện lắp ráp đúng quy trình từng bước mô hình {title} từ bộ kit {kit_name}, vận hành chạy thử và kiểm tra hoạt động. (Đạt được thông qua Hoạt động 3)"
        return [c1, c2]


def yccd_to_noun(item):
    item = item.strip().strip('-').strip()
    if not item:
        return ''
    verb_patterns = [
        r'^Nhận biết được\s+', r'^Phân biệt được\s+', r'^Nêu được\s+',
        r'^Giải thích được\s+', r'^Biết\s+', r'^Hiểu được\s+',
        r'^Trình bày được\s+', r'^Mô tả được\s+', r'^Vận dụng được\s+',
        r'^Thực hiện được\s+', r'^Sử dụng được\s+', r'^Xác định được\s+',
        r'^Liệt kê được\s+', r'^So sánh được\s+', r'^Phân tích được\s+',
        r'^Đánh giá được\s+', r'^Tạo được\s+', r'^Viết được\s+',
        r'^Lập được\s+', r'^Thiết kế được\s+', r'^Lắp ráp được\s+',
        r'^Kể tên được\s+', r'^Nhận diện được\s+', r'^Tìm hiểu\s+',
        r'^Chế tạo\s+', r'^Lắp\s+',
    ]
    for pat in verb_patterns:
        item = re.sub(pat, '', item, count=1)
    banned = ['Sự hiểu biết về ', 'Khả năng nhận diện ', 'Khả năng phân tích ',
              'Khả năng vận dụng ', 'Sự nhận biết ']
    for b in banned:
        item = item.replace(b, '')
    if item:
        item = item[0].upper() + item[1:]
    item = item.rstrip('.')
    return item


def parse_yccd_bullets(yccd_raw):
    if not yccd_raw or len(yccd_raw.strip()) < 3:
        return []
    items = re.split(r'\s*[;–\-]\s+|\.\s+', yccd_raw.strip())
    result = []
    for item in items:
        item = item.strip()
        if not item:
            continue
        noun = yccd_to_noun(item)
        if noun and len(noun) > 3:
            result.append(noun)
    return result


def parse_yccd_for_th(yccd_raw):
    if not yccd_raw or len(yccd_raw.strip()) < 3:
        return []
    text = yccd_raw.strip()
    if ';' in text:
        parts = text.split(';')
    elif ' - ' in text:
        parts = re.split(r'\s*-\s+', text)
    else:
        parts = [text]

    cleaned = []
    for it in parts:
        it = it.lstrip('-').strip()
        if it:
            it = it[0].upper() + it[1:]
            if not it.endswith('.'):
                it += '.'
            cleaned.append(it)
    return cleaned


def format_th_title(raw_title, tiet_ppct):
    t = raw_title.strip()
    m = re.match(r'^(?:Bài|BÀI)\s*(\d+)[\.:\s]*(.*)$', t)
    if m:
        bai_num, bai_name = m.group(1), m.group(2).strip()
        return f"BÀI {bai_num}. {bai_name.upper()} (Tiết: {tiet_ppct} theo PPCT)"
    if t.lower().startswith('bài'):
        clean = re.sub(r'^(?:Bài|BÀI)[\s:]*', '', t).strip()
        return f"BÀI: {clean.upper()} (Tiết: {tiet_ppct} theo PPCT)"
    if any(t.lower().startswith(x) for x in ['tiết 0', 'ôn tập', 'đánh giá', 'tổng kết']):
        return f"{t.upper()} (Tiết: {tiet_ppct} theo PPCT)"
    return f"BÀI: {t.upper()} (Tiết: {tiet_ppct} theo PPCT)"


# ─── TIỂU HỌC BUILDER (Lớp 1-5) ───────────────────────────────────────────
def build_khbd_th(grade, ten_lop, title, tiet_ppct, yccd, tuan_so, day_of_week, kit_name):
    doc = Document(TPL_DOC)
    clean_body(doc)

    for sec in doc.sections:
        sec.left_margin = Cm(3.0)
        sec.right_margin = Cm(2.0)
        sec.top_margin = Cm(2.0)
        sec.bottom_margin = Cm(2.0)

    ngay_soan, ngay_day = compute_dates(tuan_so, day_of_week)
    ls = 1.33

    # Table 0: 3x2 — NO BORDER
    tbl_info = doc.add_table(rows=3, cols=2)
    set_no_borders(tbl_info)
    fill_cell(tbl_info.rows[0].cells[0], 'Trường: Tiểu học và THCS UNIGO', bold=True, line_spacing=ls)
    fill_cell(tbl_info.rows[0].cells[1], f'Ngày soạn: {ngay_soan}', bold=True, line_spacing=ls)
    fill_cell(tbl_info.rows[1].cells[0], 'GV: Đậu Đình Nguyên', line_spacing=ls)
    fill_cell(tbl_info.rows[1].cells[1], f'Ngày dạy: {ngay_day}', line_spacing=ls)
    fill_cell(tbl_info.rows[2].cells[0], 'Tổ: Tổ chuyên môn Tiểu học', line_spacing=ls)
    fill_cell(tbl_info.rows[2].cells[1], f'Lớp: {ten_lop}', line_spacing=ls)

    # Tiêu đề bài dạy
    add_p(doc, '', line_spacing=ls)
    add_p(doc, 'KẾ HOẠCH DẠY HỌC MÔN ROBOTICS', bold=True,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls, size_pt=14)
    add_p(doc, f'CHỦ ĐIỂM: BỘ THIẾT BỊ: {kit_name.upper()}', bold=True,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)
    title_formatted = format_th_title(title, tiet_ppct)
    add_p(doc, title_formatted, bold=True,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)
    add_p(doc, '', line_spacing=ls)

    # I. YÊU CẦU CẦN ĐẠT
    add_p(doc, 'I. YÊU CẦU CẦN ĐẠT:', bold=True, first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, '- Sau bài học này em sẽ:', italic=True, first_indent=INDENT_TH_1, line_spacing=ls)

    yccd_items = parse_yccd_for_th(yccd)
    if yccd_items:
        for y_item in yccd_items:
            add_p(doc, f'+ {y_item}', left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=1.2)
    else:
        add_p(doc, f'+ Nắm vững các kiến thức, kỹ năng lắp ráp và vận hành mô hình trong bài {title}.',
              left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=1.2)

    # 1. Phẩm chất
    add_p(doc, '1. Phát triển phẩm chất', bold=True, first_indent=INDENT_TH_1, line_spacing=ls)
    pham_chat_items = [
        '- Chăm chỉ: Hăng hái tham gia hoạt động tìm hiểu linh kiện, kiên trì thực hiện thao tác lắp ráp robot. (Đạt được thông qua Hoạt động 2, Hoạt động 3)',
        '- Trách nhiệm: Giữ gìn cẩn thận các chi tiết trong bộ Kit Robotics, thu dọn phân loại gọn gàng thiết bị. (Đạt được thông qua Hoạt động 4)',
        '- Trung thực: Tự giác làm việc nhóm, tôn trọng kết quả thử nghiệm mô hình robot của bản thân và bạn bè. (Đạt được thông qua Hoạt động 3)',
    ]
    for item in pham_chat_items:
        add_p(doc, item, left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=1.2)

    # 2. Năng lực
    add_p(doc, '2. Phát triển năng lực', bold=True, first_indent=INDENT_TH_1, line_spacing=ls)
    bac = get_bac_nls(grade)

    # 2.1. Năng lực đặc thù
    add_p(doc, '2.1. Năng lực đặc thù (Robotics):', bold=True,
          first_indent=INDENT_TH_2, line_spacing=1.2)
    comps = get_rob_competencies(grade, title, yccd, kit_name)
    for c_item in comps:
        add_p(doc, c_item, left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=1.2)

    # 2.2. Năng lực số
    add_p(doc, '2.2. Năng lực số (Thông tư 02/2025 – CV 3456):', bold=True,
          first_indent=INDENT_TH_2, line_spacing=1.2)
    add_p(doc, f'- Miền V. Giải quyết vấn đề (thành tố 5.3. Sử dụng sáng tạo công nghệ số – Bậc {bac}): Ứng dụng công nghệ và thiết bị thông minh để giải quyết nhiệm vụ mô hình {title}. (Đạt được thông qua Hoạt động 2, Hoạt động 3)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=1.2)

    # 2.3. Năng lực chung
    add_p(doc, '2.3. Năng lực chung:', bold=True, first_indent=INDENT_TH_2, line_spacing=1.2)
    add_p(doc, '- Tự chủ và tự học: Quan sát sơ đồ 2D/3D hướng dẫn lắp ráp, chủ động chuẩn bị đúng linh kiện. (Đạt được thông qua Hoạt động 2)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=1.2)
    add_p(doc, '- Giao tiếp và hợp tác: Phân công nhiệm vụ nhóm ăn ý khi thực hành lắp ráp và thử nghiệm sản phẩm. (Đạt được thông qua Hoạt động 3)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=1.2)
    add_p(doc, '- Giải quyết vấn đề và sáng tạo: Phát hiện và xử lý lỗi sai lắp ráp, đề xuất ý tưởng cải tiến mô hình. (Đạt được thông qua Hoạt động 3, Hoạt động 4)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=1.2)

    # II. ĐỒ DÙNG DẠY HỌC
    add_p(doc, 'II. ĐỒ DÙNG DẠY HỌC :', bold=True, first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, f'1. Giáo viên: Bộ Kit Robotics {kit_name} mẫu, máy tính GV, máy chiếu, bài trình chiếu slide hướng dẫn lắp ráp từng bước mô hình {title}, phiếu hướng dẫn thực hành.',
          first_indent=INDENT_TH_1, line_spacing=ls)
    add_p(doc, f'2. Học sinh: Bộ Kit Robotics {kit_name} theo nhóm, dụng cụ tháo chốt nhựa, vở ghi bài.',
          first_indent=INDENT_TH_1, line_spacing=ls)

    # III. PHƯƠNG PHÁP, KĨ THUẬT DẠY HỌC
    add_p(doc, 'III. PHƯƠNG PHÁP, KĨ THUẬT DẠY HỌC', bold=True,
          first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, '- Phương pháp dạy học: Trực quan mô hình, hướng dẫn thực hành step-by-step, làm việc nhóm, giải quyết vấn đề (STEM/STEAM).',
          first_indent=INDENT_TH_1, line_spacing=ls)
    add_p(doc, '- Kĩ thuật dạy học: Think-Pair-Share, động não, giao nhiệm vụ phân tầng, trình bày 1 phút.',
          first_indent=INDENT_TH_1, line_spacing=ls)

    # IV. CÁC HOẠT ĐỘNG DẠY - HỌC CHỦ YẾU
    add_p(doc, 'IV. CÁC HOẠT ĐỘNG DẠY - HỌC CHỦ YẾU:', bold=True,
          first_indent=INDENT_0, line_spacing=ls)

    table = doc.add_table(rows=1, cols=2)
    set_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    fill_cell(hdr[0], 'HOẠT ĐỘNG CỦA GV – HS', bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.33)
    fill_cell(hdr[1], 'KẾT QUẢ CẦN ĐẠT', bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=1.33)

    activities = [
        (f'1. Hoạt động MỞ ĐẦU (5 phút)\n*Mục tiêu: Kích hoạt hứng thú, tạo tình huống kết nối vào bài học {title}.',
         f'GV: Chiếu video/hình ảnh liên quan đến {title}. Đặt câu hỏi gợi mở tình huống thực tế.\n- Mời 2-3 HS phát biểu.\n- Chốt kiến thức, dẫn dắt vào bài mới.\nHS: Quan sát màn chiếu và suy nghĩ.\n- Thảo luận nhanh với bạn bên cạnh.\n- Hăng hái giơ tay trả lời, mở bộ Kit chuẩn bị.',
         f'HS nhận diện được vấn đề kỹ thuật của bài học.\nHS hứng thú, sẵn sàng bước vào hoạt động chính.'),
         
        (f'2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI (15 phút)\n*Mục tiêu: Tìm hiểu linh kiện và nguyên lý cơ khí của {title}.',
         f'GV: Giới thiệu linh kiện cần dùng và phân tích sơ đồ lắp ráp {title}.\n- Hướng dẫn HS quan sát cách ghép chốt và khớp nối.\n- Kiểm tra khay linh kiện từng nhóm.\nHS: Theo dõi slide hướng dẫn.\n- Nhặt đúng và đủ số lượng linh kiện ra khay chứa.\n- Ghi nhớ thứ tự các bước ghép linh kiện.',
         f'HS nhận biết chính xác tên gọi, vị trí các linh kiện.\nHS nắm vững quy trình các bước lắp ráp.'),
         
        (f'3. HĐ LUYỆN TẬP-THỰC HÀNH (10 phút)\n*Mục tiêu: Tiến hành lắp ráp và vận hành chạy thử robot {title}.',
         f'GV: Giao nhiệm vụ cho các nhóm thực hành lắp ráp robot {title}.\n- Quan sát, hỗ trợ nhóm gặp khó khăn, nhắc an toàn.\n- Cho các nhóm bật nguồn vận hành chạy thử.\nHS: Phân công nhiệm vụ trong nhóm (lấy chốt, giữ khung, lắp chi tiết).\n- Tiến hành lắp ráp cẩn thận theo sơ đồ.\n- Đặt robot thử nghiệm, quan sát chuyển động và tinh chỉnh.',
         f'Mô hình robot {title} hoàn thiện đúng cấu trúc.\nRobot vận hành chính xác theo yêu cầu bài học.'),
         
        (f'4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)\n*Mục tiêu: Cải tiến sản phẩm và thu dọn phân loại linh kiện',
         f'GV: Đặt câu hỏi mở rộng cải tiến tính năng cho robot {title}.\n- Hướng dẫn tháo dỡ và kiểm kê linh kiện.\n- Nhận xét tiết học, dặn dò bài sau.\nHS: Suy nghĩ ý tưởng nâng cấp robot.\n- Tháo rời nhẹ nhàng, xếp linh kiện vào đúng ngăn trong hộp Kit.\n- Đóng hộp Kit, cất gọn gàng vào tủ thiết bị.',
         f'HS đề xuất ý tưởng cải tiến mô hình sáng tạo.\nBộ Kit Robotics được thu dọn và bảo quản ngăn nắp.')
    ]

    for title_hdr, gv_hs_txt, ket_qua_txt in activities:
        row_hdr = table.add_row()
        c0 = row_hdr.cells[0]
        c1 = row_hdr.cells[1]
        tcPr = c0._tc.get_or_add_tcPr()
        gs = OxmlElement('w:gridSpan')
        gs.set(qn('w:val'), '2')
        tcPr.append(gs)
        row_hdr._tr.remove(c1._tc)
        fill_cell(c0, title_hdr, bold=True, size_pt=13, line_spacing=1.33)

        row_cnt = table.add_row()
        fill_cell(row_cnt.cells[0], gv_hs_txt, line_spacing=1.33)
        fill_cell(row_cnt.cells[1], ket_qua_txt, line_spacing=1.33)

    # V. ĐIỀU CHỈNH - BỔ SUNG
    add_p(doc, '', line_spacing=ls)
    add_p(doc, 'V. ĐIỀU CHỈNH - BỔ SUNG SAU TIẾT DẠY :', bold=True,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)
    add_p(doc, '(GV ghi những nội dung mà mình đã bổ sung ngoài KHBD đã lên...)', bold=True,
          italic=True, alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)
    add_p(doc, '.........................................................................................................................',
          line_spacing=ls)
    add_p(doc, '', line_spacing=ls)

    # Bảng chữ ký 3x3 — NO BORDER
    tbl_sign = doc.add_table(rows=3, cols=3)
    set_no_borders(tbl_sign)
    for i, txt in enumerate(['DUYỆT CỦA BGH', 'DUYỆT CỦA TỔ CM', 'NGƯỜI SOẠN']):
        fill_cell(tbl_sign.rows[0].cells[i], txt, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)
    for i in range(3):
        fill_cell(tbl_sign.rows[1].cells[i], '(Ký, ghi rõ họ tên)',
                  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)
    for i in range(3):
        fill_cell(tbl_sign.rows[2].cells[i], '\n\n\n',
                  align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)
    fill_cell(tbl_sign.rows[2].cells[2], '\n\n\nĐậu Đình Nguyên', bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)

    return doc


# ─── THCS BUILDER (Lớp 6-8) ───────────────────────────────────────────────
def build_khbd_thcs(grade, ten_lop, title, tiet_ppct, yccd, tuan_so, day_of_week, kit_name):
    doc = Document(TPL_DOC)
    clean_body(doc)

    for sec in doc.sections:
        sec.left_margin = Cm(2.54)
        sec.right_margin = Cm(1.27)
        sec.top_margin = Cm(1.27)
        sec.bottom_margin = Cm(1.27)

    ngay_soan, ngay_day = compute_dates(tuan_so, day_of_week)
    ls = 1.15

    # Table 0: 3x2 — NO BORDER
    tbl_info = doc.add_table(rows=3, cols=2)
    set_no_borders(tbl_info)
    fill_cell(tbl_info.rows[0].cells[0], 'Trường: Tiểu học và THCS UNIGO', bold=True, line_spacing=ls)
    fill_cell(tbl_info.rows[0].cells[1], f'Ngày soạn: {ngay_soan}', bold=True, line_spacing=ls)
    fill_cell(tbl_info.rows[1].cells[0], 'GV: Đậu Đình Nguyên', line_spacing=ls)
    fill_cell(tbl_info.rows[1].cells[1], f'Ngày dạy: {ngay_day}', line_spacing=ls)
    fill_cell(tbl_info.rows[2].cells[0], 'Tổ: Tổ chuyên môn THCS', line_spacing=ls)
    fill_cell(tbl_info.rows[2].cells[1], f'Lớp: {ten_lop}', line_spacing=ls)

    # Tiêu đề bài dạy
    add_p(doc, '', line_spacing=ls)
    add_p(doc, f'TÊN BÀI DẠY: {title.upper()}', bold=True,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)
    add_p(doc, 'Môn học: Robotics', bold=True, italic=True,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)
    add_p(doc, 'Thời lượng: 1 tiết (45 phút)', bold=True, italic=True,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)
    add_p(doc, f'Tiết theo PPCT: {tiet_ppct}', bold=True,
          alignment=WD_ALIGN_PARAGRAPH.CENTER, line_spacing=ls)

    # I. Mục tiêu
    add_p(doc, 'I. Mục tiêu', bold=True, first_indent=INDENT_0, line_spacing=ls)

    # 1. Kiến thức — Danh từ trực tiếp
    add_p(doc, '1. Kiến thức:', bold=True, first_indent=INDENT_1, line_spacing=ls)
    yccd_bullets = parse_yccd_bullets(yccd)
    if yccd_bullets:
        for bullet in yccd_bullets:
            add_p(doc, f'- {bullet}.', left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
    else:
        add_p(doc, f'- Cấu tạo cơ khí, nguyên lý truyền động và thuật toán điều khiển trong mô hình {title}.',
              left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)

    # 2. Năng lực
    add_p(doc, '2. Năng lực:', bold=True, first_indent=INDENT_1, line_spacing=ls)
    bac = get_bac_nls(grade)

    # 2.1. Năng lực đặc thù
    add_p(doc, '2.1. Năng lực đặc thù (Robotics):', bold=True,
          first_indent=INDENT_2, line_spacing=ls)
    comps = get_rob_competencies(grade, title, yccd, kit_name)
    for c_item in comps:
        add_p(doc, c_item, left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)

    # 2.2. Năng lực số
    add_p(doc, '2.2. Năng lực số (Thông tư 02/2025 – CV 3456):', bold=True,
          first_indent=INDENT_2, line_spacing=ls)
    add_p(doc, f'- Miền V. Giải quyết vấn đề (thành tố 5.3. Sử dụng sáng tạo công nghệ số – Bậc {bac}): Vận hành thiết bị điều khiển thông minh, ứng dụng phần mềm nạp lập trình tự động cho robot {title}. (Đạt được thông qua Hoạt động 2, Hoạt động 3)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, f'- Miền IV. An toàn (thành tố 4.1. Bảo vệ thiết bị – Bậc {bac}): Tuân thủ quy tắc an toàn thiết bị số, nguồn điện và bảo quản bộ linh kiện. (Đạt được thông qua Hoạt động 3, Hoạt động 4)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)

    # 2.3. Năng lực chung
    add_p(doc, '2.3. Năng lực chung:', bold=True, first_indent=INDENT_2, line_spacing=ls)
    add_p(doc, '- Tự chủ và tự học: Chủ động tự nghiên cứu tài liệu sơ đồ kỹ thuật, phân tích nguyên lý truyền động. (Đạt được thông qua Hoạt động 1, Hoạt động 2)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, '- Giao tiếp và hợp tác: Giao tiếp và làm việc nhóm hiệu quả trong quá trình lắp ráp và thử nghiệm sản phẩm. (Đạt được thông qua Hoạt động 3)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, '- Giải quyết vấn đề và sáng tạo: Phát hiện lỗi sai cơ khí/lập trình và đưa ra giải pháp tối ưu cho robot. (Đạt được thông qua Hoạt động 3, Hoạt động 4)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)

    # 3. Phẩm chất
    add_p(doc, '3. Phẩm chất:', bold=True, first_indent=INDENT_1, line_spacing=ls)
    add_p(doc, '- Chăm chỉ: Tác phong công nghiệp, tư duy khoa học và kiên trì trong thực hành chế tạo robot. (Đạt được thông qua Hoạt động 2, Hoạt động 3)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, '- Trung thực: Trung thực trong báo cáo kết quả thử nghiệm và đánh giá sản phẩm của nhóm. (Đạt được thông qua Hoạt động 3)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, '- Trách nhiệm: Quản lý, bảo vệ thiết bị công nghệ và giữ gìn vệ sinh phòng thực hành bộ môn. (Đạt được thông qua Hoạt động 3, Hoạt động 4)',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)

    # II. Thiết bị dạy học và học liệu
    add_p(doc, 'II. Thiết bị dạy học và học liệu:', bold=True,
          first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, '1. Thiết bị:', bold=True, first_indent=INDENT_1, line_spacing=ls)
    add_p(doc, f'- Bộ Kit Robotics {kit_name}, máy tính giáo viên và máy tính nhóm học sinh.',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, '- Phần mềm nạp lập trình, máy chiếu bài giảng 3D.',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, '2. Học liệu:', bold=True, first_indent=INDENT_1, line_spacing=ls)
    add_p(doc, f'- Phiếu học tập thực hành, sơ đồ nguyên lý cơ khí và thuật toán điều khiển {title}.',
          left_indent=INDENT_BULLET, first_indent=INDENT_0, line_spacing=ls)

    # III. Tiến trình dạy học
    add_p(doc, 'III. Tiến trình dạy học', bold=True, first_indent=INDENT_0, line_spacing=ls)

    buoc1 = 'Bước 1: Chuyển giao nhiệm vụ học tập'
    buoc2 = 'Bước 2: Học sinh tiếp nhận nhiệm vụ học tập'
    buoc3 = 'Bước 3: Báo cáo kết quả hoạt động'
    buoc4 = 'Bước 4: Đánh giá kết quả thực hiện nhiệm vụ'
    buoc4_last = 'Bước 4: Giáo viên nhắc nhở nhiệm vụ về nhà'

    activities_thcs = [
        ('1. Hoạt động 1. Khởi động (Xác định vấn đề/nhiệm vụ học tập/Mở đầu) (7 phút)',
         f'Kích hoạt tư duy kỹ thuật và kết nối tình huống vào bài học {title}.',
         f'GV chiếu hình ảnh/video thực tế về {title} và đặt câu hỏi phân tích cơ chế hoạt động.',
         'Câu trả lời của HS xác định được vấn đề kỹ thuật cần giải quyết.',
         f'{buoc1}:\n- GV trình chiếu clip/hình ảnh thực tế liên quan đến {title}.\n- Yêu cầu HS quan sát và phân tích cơ chế chuyển động.\n{buoc2}:\n- HS tập trung quan sát màn chiếu, lắng nghe câu hỏi dẫn dắt.\n- HS thảo luận nhanh theo nhóm bàn, đối chiếu hiện tượng.\n{buoc3}:\n- 2 đại diện HS trả lời câu hỏi khởi động.\n- Nêu giả thuyết và nguyên lý vận hành ban đầu.\n{buoc4}:\n- GV nhận xét, chuẩn hóa kiến thức và dẫn dắt vào bài mới.',
         f'HS nhận diện được vấn đề kỹ thuật của bài học.\nHS sẵn sàng bước vào hoạt động chính.'),

        ('2. Hoạt động 2. Hình thành kiến thức mới/giải quyết vấn đề (18 phút)',
         f'Nghiên cứu sơ đồ thiết kế và tìm hiểu linh kiện cho mô hình {title}.',
         'GV yêu cầu HS tìm hiểu linh kiện khung, chốt, động cơ và cơ cấu truyền động.',
         'Sơ đồ khối nguyên lý cơ khí và danh mục linh kiện chính xác.',
         f'{buoc1}:\n- GV phát phiếu hướng dẫn và trình chiếu sơ đồ 2D/3D của {title}.\n- Hướng dẫn nhận biết chiều lắp chốt và cổng cắm động cơ.\n{buoc2}:\n- HS tiếp nhận phiếu, quan sát các góc ghép linh kiện.\n- HS đối chiếu danh mục, nhặt đúng số lượng chốt và khung nối.\n{buoc3}:\n- GV kiểm tra khay linh kiện từng nhóm tại bàn.\n- Đại diện nhóm giơ khay linh kiện để kiểm tra.\n{buoc4}:\n- GV chốt quy trình các bước lắp ráp mô hình robot.',
         f'HS nắm vững quy trình các bước lắp ráp.\nDanh mục linh kiện đầy đủ, chuẩn xác.'),

        ('3. Hoạt động 3. Luyện tập (12 phút)',
         f'Thực hành lắp ráp, lập trình và chạy thử mô hình robot {title}.',
         'GV yêu cầu các nhóm tiến hành lắp ráp phần cứng và thử nghiệm vận hành.',
         f'Mô hình robot {title} hoàn thiện, hoạt động chính xác theo yêu cầu.',
         f'{buoc1}:\n- GV giao nhiệm vụ cho từng nhóm lắp ráp mô hình {title}.\n{buoc2}:\n- HS phân công nhiệm vụ (xem sơ đồ, lắp khung, chọn chốt).\n- HS tiến hành lắp ráp cẩn thận từng chi tiết theo sơ đồ.\n{buoc3}:\n- GV yêu cầu các nhóm cấp nguồn/nạp code và cho robot vận hành thử.\n- HS đặt robot lên bàn thử nghiệm, quan sát hoạt động.\n{buoc4}:\n- GV đánh giá sản phẩm thực hành, chấm điểm tiêu chí kỹ thuật.\n- HS tinh chỉnh lại chốt nối nếu robot bị kẹt hoặc di chuyển lệch.',
         f'Mô hình robot {title} hoàn thiện.\nRobot vận hành chính xác, trơn tru.'),

        ('4. Hoạt động 4. Vận dụng (8 phút)',
         'Vận dụng sáng tạo, nâng cấp tính năng robot và dọn dẹp phòng học.',
         'GV đặt yêu cầu cải tiến tối ưu thuật toán hoặc thêm chi tiết trang trí.',
         'Báo cáo ý tưởng cải tiến mô hình và bộ Kit được sắp xếp ngăn nắp.',
         f'{buoc1}:\n- GV đặt câu hỏi: "Em có thể cải tiến cấu trúc nào để {title} hoạt động tốt hơn?"\n{buoc2}:\n- HS suy nghĩ ý tưởng sáng tạo mở rộng.\n- HS nhẹ nhàng tháo chốt, phân loại linh kiện về đúng ngăn trong hộp Kit.\n{buoc3}:\n- 1-2 HS trình bày ý tưởng nâng cấp robot trước lớp.\n{buoc4_last}:\n- GV dặn dò nhiệm vụ chuẩn bị cho bài học tiếp theo.\n- HS thu dọn bàn học, nộp lại hộp Kit Robotics ngăn nắp.',
         f'HS đề xuất được ý tưởng cải tiến mô hình robot.\nBộ Kit Robotics được phân loại và bảo quản tốt.')
    ]

    for idx, (act_title, muc_tieu, noi_dung, san_pham, gv_hs_col, ket_qua_col) in enumerate(activities_thcs, 1):
        add_p(doc, act_title, bold=True, first_indent=INDENT_1, line_spacing=ls)
        add_multi_run(doc, [
            ('a) Mục tiêu: ', True, False), (muc_tieu, False, True)
        ], first_indent=INDENT_1, line_spacing=ls)
        add_multi_run(doc, [
            ('b) Nội dung: ', True, False), (noi_dung, False, True)
        ], first_indent=INDENT_1, line_spacing=ls)
        add_multi_run(doc, [
            ('c) Sản phẩm: ', True, False), (san_pham, False, True)
        ], first_indent=INDENT_1, line_spacing=ls)
        add_multi_run(doc, [
            ('d) Tổ chức thực hiện:', True, False),
        ], first_indent=INDENT_1, line_spacing=ls)

        tbl = doc.add_table(rows=1, cols=2)
        set_borders(tbl)
        for cell in tbl.columns[0].cells:
            cell.width = Cm(9.0)
        for cell in tbl.columns[1].cells:
            cell.width = Cm(7.0)

        fill_cell(tbl.rows[0].cells[0], 'HOẠT ĐỘNG CỦA GV – HS', bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(tbl.rows[0].cells[1], 'KẾT QUẢ CẦN ĐẠT', bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)

        row = tbl.add_row()
        fill_cell(row.cells[0], gv_hs_col)
        fill_cell(row.cells[1], ket_qua_col)

        add_p(doc, '', line_spacing=ls)

    # Rút kinh nghiệm & Bảng chữ ký
    add_p(doc, 'RÚT KINH NGHIỆM SAU BÀI DẠY:', bold=True,
          first_indent=INDENT_0, line_spacing=ls)
    add_p(doc, '...........................................................................................................................', line_spacing=ls)
    add_p(doc, '...........................................................................................................................', line_spacing=ls)
    add_p(doc, '', line_spacing=ls)

    tbl_sign = doc.add_table(rows=3, cols=3)
    set_no_borders(tbl_sign)
    for i, txt in enumerate(['DUYỆT CỦA BGH', 'DUYỆT CỦA TỔ CM', 'NGƯỜI SOẠN']):
        fill_cell(tbl_sign.rows[0].cells[i], txt, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    for i in range(3):
        fill_cell(tbl_sign.rows[1].cells[i], '(Ký, ghi rõ họ tên)',
                  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    for i in range(3):
        fill_cell(tbl_sign.rows[2].cells[i], '\n\n\n',
                  align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell(tbl_sign.rows[2].cells[2], '\n\n\nĐậu Đình Nguyên', bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER)

    return doc


# ─── MAIN GENERATOR ────────────────────────────────────────────────────────
def main():
    print("=" * 60)
    print(" BẮT ĐẦU TẠO LẠI TOÀN BỘ KHBD ROBOTICS (CHUẨN LBG & AGENTS.MD)")
    print("=" * 60)

    if os.path.exists(OUT_BASE):
        shutil.rmtree(OUT_BASE)
        print("  [+] Đã xóa sạch thư mục KHBD_Robotics cũ.")
    os.makedirs(OUT_BASE, exist_ok=True)

    doc_ppct = Document(PPCT_ROB)
    tables = doc_ppct.tables
    created = 0

    for grade in range(1, 9):
        g_str = str(grade)
        day_of_week, ten_lop = ROB_SCHEDULE[g_str]
        kit_name = get_kit_name(grade)
        is_rotation = (grade >= 5)

        t_idx = 3 + grade  # Table 4 to 11
        t = tables[t_idx]
        print(f"\n---> Đang xử lý Lớp {grade} ({ten_lop}) - Table {t_idx} ({len(t.rows)-1} bài)...")

        for row_idx, row in enumerate(t.rows[1:], start=1):
            cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
            title = cells[1] if len(cells) > 1 else ''
            yccd = cells[4] if len(cells) > 4 else (cells[3] if len(cells) > 3 else title)

            if not title:
                continue

            # row_idx 1 is always Tiết 0
            if row_idx == 1:
                ppct_num = 0
                tiet_ppct = 0
                tuan_so = 1
                safe_title = "Tiet_0_Dinh_huong_mon_hoc"
                filename = f"KHBD_Robotics_Lớp_{grade}_Tiet00_{safe_title}.docx"
            else:
                ppct_num = row_idx - 1  # Row 2 -> PPCT 1, Row 3 -> PPCT 2...
                tiet_ppct = ppct_num
                safe_title = sanitize(title)
                filename = f"KHBD_Robotics_Lớp_{grade}_Tiet{tiet_ppct:02d}_{safe_title}.docx"

                if is_rotation:
                    # Lớp 5, 6, 7, 8: Tuần LẺ (3, 5, 7, 9...) - mỗi tuần 2 tiết
                    # PPCT 1, 2 -> Tuần 3
                    # PPCT 3, 4 -> Tuần 5
                    # PPCT 5, 6 -> Tuần 7...
                    tuan_so = 2 * ((ppct_num - 1) // 2) + 3
                else:
                    # Lớp 1, 2, 3, 4: Mỗi tuần 1 tiết (Tuần 2, 3, 4, 5...)
                    # PPCT 1 -> Tuần 2
                    # PPCT 2 -> Tuần 3
                    # PPCT 3 -> Tuần 4...
                    tuan_so = ppct_num + 1

            tuan_folder = f"Tuần_{tuan_so:02d}"
            out_dir = os.path.join(OUT_BASE, f"Lớp_{grade}", tuan_folder)
            os.makedirs(out_dir, exist_ok=True)
            out_file = os.path.join(out_dir, filename)

            # Build Document
            if grade >= 6:
                doc = build_khbd_thcs(grade, ten_lop, title, tiet_ppct, yccd, tuan_so, day_of_week, kit_name)
            else:
                doc = build_khbd_th(grade, ten_lop, title, tiet_ppct, yccd, tuan_so, day_of_week, kit_name)

            try:
                doc.save(out_file)
                created += 1
                print(f"  [+] Lớp {grade} -> {tuan_folder} -> {filename}")
            except Exception as e:
                print(f"  [!] Lỗi khi lưu {out_file}: {e}")

    print("\n" + "=" * 60)
    print(f" HOÀN THÀNH TẠO {created} FILE KHBD ROBOTICS ĐỒNG BỘ 100% VỚI LỊCH BÁO GIẢNG!")
    print("=" * 60)


if __name__ == '__main__':
    main()
