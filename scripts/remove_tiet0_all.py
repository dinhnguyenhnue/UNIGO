import os
import sys
import shutil
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

sys.stdout.reconfigure(encoding='utf-8')

DIR_MAU = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm'
DIR_TUNG_LOP = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch dạy học Tin học từng lớp'
DIR_DIST = r'D:\UNIGO\Phân phối chương trình\Tin học'

MAIN_ALL = os.path.join(DIR_MAU, 'Kế hoạch dạy học môn Tin học 2026-2027.docx')
MAIN_TH = os.path.join(DIR_MAU, 'Kế hoạch dạy học môn Tin học (TH) - 2026 - 2027.docx')
MAIN_THCS = os.path.join(DIR_MAU, 'Kế hoạch dạy học môn Tin học (THCS) - 2026 - 2027.docx')

COL_WIDTHS = [0.9, 3.8, 1.1, 1.4, 9.3]

def set_table_borders(table):
    """Sets explicit XML full black grid borders for a table."""
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        for child in list(tblPr[0]):
            if child.tag.endswith('tblBorders'):
                tblPr[0].remove(child)
                
        borders_xml = parse_xml(
            r'<w:tblBorders %s>'
            r'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        tblPr[0].append(borders_xml)

def set_cell_width_and_align(cell, width_cm, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    """Sets exact cell width and text alignment."""
    cell.width = Cm(width_cm)
    tcPr = cell._element.get_or_add_tcPr()
    tcW = OxmlElement('w:tcW')
    tcW.set(qn('w:w'), str(int(width_cm * 567)))
    tcW.set(qn('w:type'), 'dxa')
    tcPr.append(tcW)
    
    for p in cell.paragraphs:
        p.alignment = align
        p.style.font.name = 'Times New Roman'
        p.style.font.size = Pt(13)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)

def set_font_all(doc):
    """Sets Times New Roman 13pt for all paragraphs, runs, tables, cells in the document."""
    for p in doc.paragraphs:
        p.style.font.name = 'Times New Roman'
        p.style.font.size = Pt(13)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)
            
    for t in doc.tables:
        set_table_borders(t)
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.style.font.name = 'Times New Roman'
                    p.style.font.size = Pt(13)
                    for r in p.runs:
                        r.font.name = 'Times New Roman'
                        r.font.size = Pt(13)

def clean_and_recount_table(table):
    """Removes Tiết 0 row if present, recounts STT & PPCT sequentially from 1."""
    if len(table.rows) == 0:
        return False
        
    hdr = [c.text.strip() for c in table.rows[0].cells]
    if not (len(hdr) == 5 and hdr[0] == 'STT' and 'Bài học' in hdr[1]):
        return False
        
    # Find Tiết 0 row
    tiet0_idx = None
    for r_idx, row in enumerate(table.rows):
        txt_all = " ".join([c.text.strip() for c in row.cells])
        if "Tiết 0" in txt_all or "Định hướng môn học" in txt_all:
            tiet0_idx = r_idx
            break
            
    if tiet0_idx is not None:
        table._element.remove(table.rows[tiet0_idx]._element)
        
    # Recount STT and PPCT
    curr_ppct = 0
    stt_counter = 1
    
    aligns = [
        WD_ALIGN_PARAGRAPH.CENTER,  # STT
        WD_ALIGN_PARAGRAPH.JUSTIFY, # Bài học (1)
        WD_ALIGN_PARAGRAPH.CENTER,  # Số tiết (2)
        WD_ALIGN_PARAGRAPH.CENTER,  # Tiết theo PPCT
        WD_ALIGN_PARAGRAPH.JUSTIFY  # Yêu cầu cần đạt (3)
    ]
    
    for row_idx, row in enumerate(table.rows):
        if row_idx == 0:
            # Header
            for c_idx in range(5):
                set_cell_width_and_align(row.cells[c_idx], COL_WIDTHS[c_idx], WD_ALIGN_PARAGRAPH.CENTER)
                for r in row.cells[c_idx].paragraphs[0].runs:
                    r.bold = True
            continue
            
        cells_txt = [c.text.strip() for c in row.cells]
        if len(set(cells_txt)) == 1 and "Chủ đề" in cells_txt[0]:
            p = row.cells[0].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(13)
            continue
            
        if len(row.cells) == 5:
            so_tiet_str = row.cells[2].text.strip()
            try:
                st = int(so_tiet_str)
            except:
                st = 1
                
            if st == 1:
                curr_ppct += 1
                ppct_str = str(curr_ppct)
            else:
                ppct_str = f"{curr_ppct + 1} - {curr_ppct + st}"
                curr_ppct += st
                
            row.cells[0].text = str(stt_counter)
            row.cells[3].text = ppct_str
            stt_counter += 1
            
            for c_idx in range(5):
                set_cell_width_and_align(row.cells[c_idx], COL_WIDTHS[c_idx], aligns[c_idx])
                
    set_table_borders(table)
    return True

def process_file(filepath):
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        return
        
    print(f"Processing: {filepath}")
    doc = docx.Document(filepath)
    modified = False
    
    for i, t in enumerate(doc.tables):
        if clean_and_recount_table(t):
            modified = True
            
    set_font_all(doc)
    doc.save(filepath)
    print(f"Successfully updated & saved: {filepath}")

def main():
    # Process main documents
    process_file(MAIN_ALL)
    process_file(MAIN_TH)
    process_file(MAIN_THCS)
    
    # Process individual class files
    for grade_num in range(1, 9):
        f_path = os.path.join(DIR_TUNG_LOP, f'Kế hoạch dạy học môn Tin học - Lớp {grade_num} - 2026 - 2027.docx')
        process_file(f_path)
        
    f_tth = os.path.join(DIR_TUNG_LOP, 'Kế hoạch dạy học môn Tin học - Tiền tiểu học - 2026 - 2027.docx')
    process_file(f_tth)
    
    # Synchronize to Phân phối chương trình\Tin học
    print(f"\nSynchronizing to {DIR_DIST}...")
    shutil.copyfile(MAIN_ALL, os.path.join(DIR_DIST, 'Kế hoạch dạy học môn Tin học 2026-2027.docx'))
    shutil.copyfile(MAIN_TH, os.path.join(DIR_DIST, 'Kế hoạch dạy học môn Tin học (TH) - 2026 - 2027.docx'))
    if os.path.exists(MAIN_THCS):
        shutil.copyfile(MAIN_THCS, os.path.join(DIR_DIST, 'Kế hoạch dạy học môn Tin học (THCS) - 2026 - 2027.docx'))
    print("Synchronization completed successfully!")

if __name__ == '__main__':
    try:
        main()
    except PermissionError as pe:
        print(f"PERMISSION ERROR: Please close open Word files. Details: {pe}")
        sys.exit(1)
