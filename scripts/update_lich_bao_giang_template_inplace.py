import os, sys, re, shutil
import openpyxl, docx
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

sys.stdout.reconfigure(encoding='utf-8')

def set_font_run(run, font_name="Times New Roman", size_pt=12, bold=None, italic=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
        
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
    rPr.append(rFonts)

def format_paragraph(p, font_name="Times New Roman", size_pt=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=None, italic=None):
    p.alignment = align
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.05
    for run in p.runs:
        set_font_run(run, font_name=font_name, size_pt=size_pt, bold=bold, italic=italic)

def format_cell(cell, font_name="Times New Roman", size_pt=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=None, italic=None):
    tcPr = cell._element.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)
    for p in cell.paragraphs:
        format_paragraph(p, font_name=font_name, size_pt=size_pt, align=align, bold=bold, italic=italic)

def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)

def parse_schedule_from_excel(xlsx_path):
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb['TKB_LOP_SC']

    classes = []
    for c in range(3, ws.max_column + 1, 2):
        cls_name = ws.cell(row=4, column=c).value
        if cls_name:
            classes.append((str(cls_name).strip(), c, c+1))

    schedule_sang = []
    schedule_chieu = []

    thu_map = {'2': 'Hai', '3': 'Ba', '4': 'Tư', '5': 'Năm', '6': 'Sáu'}

    current_thu = ''
    for r in range(6, ws.max_row + 1):
        thu_val = ws.cell(row=r, column=1).value
        tiet_val = ws.cell(row=r, column=2).value
        if thu_val:
            current_thu = str(thu_val).strip()
        if not tiet_val:
            continue
        tiet_str = str(tiet_val).strip()
        
        thu_name = thu_map.get(current_thu, current_thu)
        
        for cls_name, sang_col, chieu_col in classes:
            sang_val = str(ws.cell(row=r, column=sang_col).value or '').strip()
            chieu_val = str(ws.cell(row=r, column=chieu_col).value or '').strip()
            
            if 'nguyên' in sang_val.lower() or ('tin' in sang_val.lower() and 'nguyên' in sang_val.lower()) or ('rob' in sang_val.lower() and 'nguyên' in sang_val.lower()):
                mon_name = "Robotics" if 'rob' in sang_val.lower() else "Tin học"
                schedule_sang.append((thu_name, int(tiet_str) if tiet_str.isdigit() else tiet_str, mon_name, cls_name))
                
            if 'nguyên' in chieu_val.lower() or ('tin' in chieu_val.lower() and 'nguyên' in chieu_val.lower()) or ('rob' in chieu_val.lower() and 'nguyên' in chieu_val.lower()):
                mon_name = "Robotics" if 'rob' in chieu_val.lower() else "Tin học"
                schedule_chieu.append((thu_name, int(tiet_str) if tiet_str.isdigit() else tiet_str, mon_name, cls_name))

    day_order = {'Hai': 2, 'Ba': 3, 'Tư': 4, 'Năm': 5, 'Sáu': 6}
    schedule_sang.sort(key=lambda x: (day_order.get(x[0], 9), x[1]))
    schedule_chieu.sort(key=lambda x: (day_order.get(x[0], 9), x[1]))

    return schedule_sang, schedule_chieu

def update_template_inplace(pristine_template_path, target_docx_path, schedule_sang, schedule_chieu):
    doc = docx.Document(pristine_template_path)
    
    # Update Paragraph titles without touching headers/footers/sections/images
    for p in doc.paragraphs:
        txt = p.text.strip()
        if 'TUẦN 01' in txt:
            p.text = '                                                   TUẦN 01 (từ ngày 04/08/2026 đến ngày 08/08/2026)'
            format_paragraph(p, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT)
        elif 'Ngày ........ tháng........ năm 202' in txt:
            p.text = 'Ngày 04 tháng 08 năm 2026'
            format_paragraph(p, align=WD_ALIGN_PARAGRAPH.LEFT)

    # Update Table 0 (Teacher Info)
    if len(doc.tables) > 0:
        t0_cell = doc.tables[0].rows[0].cells[0]
        if 'Đậu Đình Nguyên' not in t0_cell.text:
            t0_cell.text = t0_cell.text.replace('Họ và tên giáo viên  :              ', 'Họ và tên giáo viên  :  Đậu Đình Nguyên            ')
            t0_cell.text = t0_cell.text.replace('Chức vụ hiện nay                  : Giáo viên', 'Chức vụ hiện nay                  : Giáo viên Tin học & Robotics')
            format_cell(t0_cell, font_name="Times New Roman", size_pt=12, align=WD_ALIGN_PARAGRAPH.LEFT)

    col_widths_cm = [1.8, 1.2, 1.5, 2.2, 1.3, 5.8, 2.7]

    def fill_table_data(table, schedule_list):
        req_rows = len(schedule_list)
        curr_rows = len(table.rows) - 1
        
        if curr_rows < req_rows:
            for _ in range(req_rows - curr_rows):
                table.add_row()
        elif curr_rows > req_rows:
            for _ in range(curr_rows - req_rows):
                row_elem = table.rows[-1]._element
                row_elem.getparent().remove(row_elem)
                
        prev_day = ""
        for i, (thu, tiet, mon, lop) in enumerate(schedule_list):
            row = table.rows[i + 1]
            day_str = thu if thu != prev_day else ""
            prev_day = thu
            
            is_monday = (thu == 'Hai')
            ppct_val = "" if is_monday else "1"
            ten_bai_val = "" if is_monday else "Tiết 0: Định hướng môn học"
            dodung_val = "" if is_monday else ("Máy tính, Tivi" if mon == "Tin học" else "Bộ kit Robotics, Máy tính")
            
            row.cells[0].text = day_str
            format_cell(row.cells[0], font_name="Times New Roman", size_pt=11, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True if day_str else False)
            set_cell_width(row.cells[0], col_widths_cm[0])
            
            row.cells[1].text = str(tiet)
            format_cell(row.cells[1], font_name="Times New Roman", size_pt=11, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_width(row.cells[1], col_widths_cm[1])
            
            row.cells[2].text = ppct_val
            format_cell(row.cells[2], font_name="Times New Roman", size_pt=11, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_width(row.cells[2], col_widths_cm[2])
            
            row.cells[3].text = mon
            format_cell(row.cells[3], font_name="Times New Roman", size_pt=11, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_width(row.cells[3], col_widths_cm[3])
            
            row.cells[4].text = lop
            format_cell(row.cells[4], font_name="Times New Roman", size_pt=11, align=WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_width(row.cells[4], col_widths_cm[4])
            
            row.cells[5].text = ten_bai_val
            format_cell(row.cells[5], font_name="Times New Roman", size_pt=11, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
            set_cell_width(row.cells[5], col_widths_cm[5])
            
            row.cells[6].text = dodung_val
            format_cell(row.cells[6], font_name="Times New Roman", size_pt=11, align=WD_ALIGN_PARAGRAPH.LEFT)
            set_cell_width(row.cells[6], col_widths_cm[6])

    # Table 1 is Morning schedule
    if len(doc.tables) > 1:
        fill_table_data(doc.tables[1], schedule_sang)
        
    # Table 3 is Afternoon schedule
    if len(doc.tables) > 3:
        fill_table_data(doc.tables[3], schedule_chieu)

    try:
        doc.save(target_docx_path)
        print(f"Successfully updated template in-place at: {target_docx_path}")
    except PermissionError:
        alt_path = target_docx_path.replace('.docx', ' - InPlace.docx')
        doc.save(alt_path)
        print(f"File locked. Saved in-place template to: {alt_path}")

def main():
    pristine_template_path = r"D:\UNIGO\Hệ thống mẫu văn bản\Lịch báo giảng.docx"
    xlsx_path = r"D:\UNIGO\TKB toàn trường CHECK - lần 2.1.xlsx"
    target_docx_path = r"D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Lịch báo giảng\Lịch báo giảng.docx"
    
    schedule_sang, schedule_chieu = parse_schedule_from_excel(xlsx_path)
    update_template_inplace(pristine_template_path, target_docx_path, schedule_sang, schedule_chieu)

if __name__ == "__main__":
    main()
