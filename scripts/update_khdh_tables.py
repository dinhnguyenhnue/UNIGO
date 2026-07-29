import os, sys, re, shutil
import docx
from docx.shared import Pt, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

sys.stdout.reconfigure(encoding='utf-8')

def set_font_run(run, font_name="Times New Roman", size_pt=13, bold=None, italic=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
        
    rPr = run._r.get_or_add_rPr()
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:ascii="{font_name}" w:hAnsi="{font_name}" w:cs="{font_name}"/>')
    rPr.append(rFonts)

def format_paragraph(p, font_name="Times New Roman", size_pt=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=None, italic=None):
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = 1.15
    for run in p.runs:
        set_font_run(run, font_name=font_name, size_pt=size_pt, bold=bold, italic=italic)

def format_cell(cell, font_name="Times New Roman", size_pt=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=None, italic=None):
    tcPr = cell._element.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)
    for p in cell.paragraphs:
        format_paragraph(p, font_name=font_name, size_pt=size_pt, align=align, bold=bold, italic=italic)

def set_cell_width(cell, width_cm):
    cell.width = Cm(width_cm)

def set_table_borders(table):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders_xml = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>\n'
            f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
            f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
            f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
            f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
            f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
            f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>\n'
            f'</w:tblBorders>'
        )
        tblPr[0].append(borders_xml)

def parse_periods_and_continuous_ppct(rows_data):
    transformed = []
    stt_counter = 1
    current_ppct = 1
    
    for is_topic, old_stt, name, old_so_tiet, yccd in rows_data:
        if is_topic:
            topic_title = name if 'chủ đề' in name.lower() else old_stt
            topic_title = re.sub(r'^\d+[\.\:]\s*', '', topic_title).strip()
            transformed.append((True, '', topic_title, '', '', ''))
        else:
            is_tiet_0 = 'tiết 0' in name.lower() or old_so_tiet == '0'
            if is_tiet_0:
                so_tiet = 1
                ppct_str = str(current_ppct)
                current_ppct += 1
            else:
                nums = [int(x) for x in re.findall(r'\d+', old_so_tiet)]
                if not nums:
                    so_tiet = 1
                elif '-' in old_so_tiet and len(nums) == 2 and nums[1] > nums[0]:
                    so_tiet = nums[1] - nums[0] + 1
                else:
                    so_tiet = max(1, len(nums))
                
                if so_tiet == 1:
                    ppct_str = str(current_ppct)
                elif so_tiet == 2:
                    ppct_str = f"{current_ppct}, {current_ppct + 1}"
                else:
                    ppct_str = f"{current_ppct} - {current_ppct + so_tiet - 1}"
                
                current_ppct += so_tiet
                
            transformed.append((False, str(stt_counter), name, str(so_tiet), ppct_str, yccd))
            stt_counter += 1
            
    return transformed

def transform_docx_tables(file_path):
    print(f"Processing: {file_path}")
    doc = docx.Document(file_path)
    modified = False
    
    for t_idx, table in enumerate(doc.tables):
        if len(table.rows) <= 5:
            continue
            
        header_cells = [c.text.strip().replace('\n', ' ') for c in table.rows[0].cells]
        is_khdh_table = False
        if len(header_cells) in [4, 5] and any('bài học' in h.lower() for h in header_cells):
            is_khdh_table = True
            
        if not is_khdh_table:
            continue
            
        print(f"  Found target table {t_idx} ({len(table.rows)} rows, {len(table.columns)} cols)")
        modified = True
        
        # 1. Extract raw row data
        rows_data = []
        for r_idx in range(1, len(table.rows)):
            row = table.rows[r_idx]
            c_text = [c.text.strip().replace('\n', ' ') for c in row.cells]
            if len(c_text) == 4:
                old_stt, name, old_st, yccd = c_text[0], c_text[1], c_text[2], c_text[3]
            elif len(c_text) >= 5:
                # If table was already expanded to 5 cols in previous run
                old_stt, name, old_st, yccd = c_text[0], c_text[1], c_text[2], c_text[4]
            else:
                continue
                
            is_topic = ('chủ đề' in name.lower() or 'chủ đề' in old_stt.lower()) and ('bài ' not in name.lower())
            rows_data.append((is_topic, old_stt, name, old_st, yccd))
            
        # 2. Parse transformed row data
        transformed_rows = parse_periods_and_continuous_ppct(rows_data)
        
        # 3. Rebuild table element cleanly in docx XML with explicit borders
        table_elem = table._element
        parent = table_elem.getparent()
        pos = parent.index(table_elem)
        
        new_table = doc.add_table(rows=0, cols=5)
        new_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        set_table_borders(new_table)
        
        # Set table width & grid
        col_widths_cm = [1.2, 5.5, 1.5, 2.2, 6.0] # Total ~16.4 cm
        
        # Add Header Row
        hdr_row = new_table.add_row()
        hdr_titles = ['STT', 'Bài học (1)', 'Số tiết (2)', 'Tiết theo PPCT', 'Yêu cầu cần đạt (3)']
        for idx, title in enumerate(hdr_titles):
            cell = hdr_row.cells[idx]
            cell.text = title
            align = WD_ALIGN_PARAGRAPH.CENTER if idx in [0, 2, 3] else WD_ALIGN_PARAGRAPH.LEFT
            format_cell(cell, font_name="Times New Roman", size_pt=13, align=align, bold=True)
            set_cell_width(cell, col_widths_cm[idx])
            
        # Add Data Rows
        for is_topic, stt_str, name, so_tiet_str, ppct_str, yccd in transformed_rows:
            row = new_table.add_row()
            if is_topic:
                # Merge cell 0..4
                cell_0 = row.cells[0]
                cell_0.merge(row.cells[4])
                cell_0.text = name
                format_cell(cell_0, font_name="Times New Roman", size_pt=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY, bold=True)
            else:
                # 5 columns
                row.cells[0].text = stt_str
                format_cell(row.cells[0], font_name="Times New Roman", size_pt=13, align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_width(row.cells[0], col_widths_cm[0])
                
                row.cells[1].text = name
                format_cell(row.cells[1], font_name="Times New Roman", size_pt=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
                set_cell_width(row.cells[1], col_widths_cm[1])
                
                row.cells[2].text = so_tiet_str
                format_cell(row.cells[2], font_name="Times New Roman", size_pt=13, align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_width(row.cells[2], col_widths_cm[2])
                
                row.cells[3].text = ppct_str
                format_cell(row.cells[3], font_name="Times New Roman", size_pt=13, align=WD_ALIGN_PARAGRAPH.CENTER)
                set_cell_width(row.cells[3], col_widths_cm[3])
                
                row.cells[4].text = yccd
                format_cell(row.cells[4], font_name="Times New Roman", size_pt=13, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
                set_cell_width(row.cells[4], col_widths_cm[4])
                
        # Swap new table into old table position
        parent.insert(pos, new_table._element)
        parent.remove(table_elem)
        
    if modified:
        doc.save(file_path)
        print(f"  Successfully updated {file_path}")

def main():
    workspace = r"d:\UNIGO"
    
    target_files = [
        r".\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch dạy học môn Tin học 2026-2027.docx",
        r".\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch dạy học môn Tin học (TH) - 2026 - 2027.docx",
        r".\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch dạy học môn Tin học (THCS) - 2026 - 2027.docx",
        r".\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch dạy học môn Robotics 2026-2027.docx",
        r".\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch dạy học môn Robotics (TH) - 2026 - 2027.docx",
        r".\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch dạy học môn Robotics (THCS) - 2026 - 2027.docx",
    ]
    
    for grade in range(1, 9):
        target_files.append(f".\\Hệ thống mẫu văn bản\\Nguyên đã làm\\Kế hoạch dạy học Tin học từng lớp\\Kế hoạch dạy học môn Tin học - Lớp {grade} - 2026 - 2027.docx")
        target_files.append(f".\\Hệ thống mẫu văn bản\\Nguyên đã làm\\Kế hoạch dạy học Robotics từng lớp\\Kế hoạch dạy học môn Robotics - Lớp {grade} - 2026 - 2027.docx")
        
    for rel_path in target_files:
        full_path = os.path.join(workspace, rel_path)
        if os.path.exists(full_path):
            try:
                transform_docx_tables(full_path)
            except Exception as e:
                print(f"Error processing {full_path}: {e}")
        else:
            print(f"File not found: {full_path}")
            
    print("\n--- Synchronizing to Phân phối chương trình directory ---")
    sync_map = [
        (r".\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch dạy học môn Tin học 2026-2027.docx", r".\Phân phối chương trình\Tin học\Kế hoạch dạy học môn Tin học 2026-2027.docx"),
        (r".\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch dạy học môn Tin học (TH) - 2026 - 2027.docx", r".\Phân phối chương trình\Tin học\Kế hoạch dạy học môn Tin học (TH) - 2026 - 2027.docx"),
        (r".\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch dạy học môn Tin học (THCS) - 2026 - 2027.docx", r".\Phân phối chương trình\Tin học\Kế hoạch dạy học môn Tin học (THCS) - 2026 - 2027.docx"),
        (r".\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch dạy học môn Robotics 2026-2027.docx", r".\Phân phối chương trình\Robotics\Kế hoạch dạy học môn Robotics 2026-2027.docx"),
        (r".\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch dạy học môn Robotics (TH) - 2026 - 2027.docx", r".\Phân phối chương trình\Robotics\Kế hoạch dạy học môn Robotics (TH) - 2026 - 2027.docx"),
        (r".\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch dạy học môn Robotics (THCS) - 2026 - 2027.docx", r".\Phân phối chương trình\Robotics\Kế hoạch dạy học môn Robotics (THCS) - 2026 - 2027.docx"),
    ]
    for src, dst in sync_map:
        src_full = os.path.join(workspace, src)
        dst_full = os.path.join(workspace, dst)
        if os.path.exists(src_full):
            os.makedirs(os.path.dirname(dst_full), exist_ok=True)
            shutil.copy2(src_full, dst_full)
            print(f"  Copied {src} -> {dst}")
            
    print("\n--- ALL COMPLETED SUCCESSFULLY WITH BORDERS ---")

if __name__ == "__main__":
    main()
