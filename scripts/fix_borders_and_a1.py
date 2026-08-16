import sys, os, re
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.oxml import OxmlElement

def make_borderless_table(tbl):
    """Set Table 0 borders to val='none' (borderless)."""
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    
    # Remove existing tblBorders
    old_b = tblPr.find(qn('w:tblBorders'))
    if old_b is not None:
        tblPr.remove(old_b)
        
    borders_xml = parse_xml(
        r'<w:tblBorders %s>'
        r'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        r'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        r'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        r'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        r'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        r'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        r'</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders_xml)
    
    # Also clean tcBorders in each cell
    for r in tbl._tbl.findall(qn('w:tr')):
        for c in r.findall(qn('w:tc')):
            tcPr = c.find(qn('w:tcPr'))
            if tcPr is not None:
                tcB = tcPr.find(qn('w:tcBorders'))
                if tcB is not None:
                    tcPr.remove(tcB)

def clean_a1_duplication(text):
    """Clean repeated A1 patterns like 7A1A1A1 -> 7A1."""
    # Matches (6|7|8|1|2|3)(A1)+ or similar
    text = re.sub(r'(6|7|8|1|2|3)A1(A1)+', r'\1A1', text)
    text = re.sub(r'(5|4)C1(C1)+', r'\1C1', text)
    return text

def fix_file(filepath):
    doc = Document(filepath)
    modified = False
    
    # 1. Fix Table 0: make it borderless
    if doc.tables:
        t0 = doc.tables[0]
        # Only make Table 0 borderless if it's the header info table (has "Trường:" or "Họ và tên" or "Ngày soạn")
        t0_text = "".join(c.text for r in t0.rows for c in r.cells)
        if any(k in t0_text for k in ["Trường", "Họ tên", "Họ và tên", "Ngày soạn", "Bộ môn"]):
            make_borderless_table(t0)
            modified = True
            
        # Clean A1 dup in Table 0
        for r in t0.rows:
            for c in r.cells:
                for p in c.paragraphs:
                    for run in p.runs:
                        new_t = clean_a1_duplication(run.text)
                        if new_t != run.text:
                            run.text = new_t
                            modified = True
                            
    # 2. Fix A1 duplication in all paragraphs
    for p in doc.paragraphs:
        for run in p.runs:
            new_t = clean_a1_duplication(run.text)
            if new_t != run.text:
                run.text = new_t
                modified = True

    # 3. Check for duplicated header block in THCS files
    # In some files (like Lớp 7 Bài 2), there is:
    # P[0]: TÊN BÀI DẠY: BÀI 2. PHẦN MỀM MÁY TÍNH...
    # P[1]: Môn học: Tin học Lớp: 7A1 Thời lượng: 1 tiết (45 phút)
    # P[2]: Tiết theo PPCT: 1  <-- wrong PPCT and title of Bai 1
    # P[3]: Tên tiết: Thiết bị vào - ra
    if len(doc.paragraphs) >= 4:
        p0_text = doc.paragraphs[0].text.strip()
        p1_text = doc.paragraphs[1].text.strip()
        p2_text = doc.paragraphs[2].text.strip()
        p3_text = doc.paragraphs[3].text.strip()
        if p0_text.startswith("TÊN BÀI DẠY: BÀI 2") and p2_text == "Tiết theo PPCT: 1" and "Thiết bị vào" in p3_text:
            # Remove p1, p2, p3 or clear them
            print(f"  🧹 Fixing duplicated header block in: {os.path.basename(filepath)}")
            # Clear paragraphs 1, 2, 3
            p1_elem = doc.paragraphs[1]._p
            p2_elem = doc.paragraphs[2]._p
            p3_elem = doc.paragraphs[3]._p
            p1_elem.getparent().remove(p1_elem)
            p2_elem.getparent().remove(p2_elem)
            p3_elem.getparent().remove(p3_elem)
            modified = True
        elif p0_text.startswith("TÊN BÀI DẠY: BÀI 2") and p2_text == "Tiết theo PPCT: 1" and "Thông tin và dữ liệu" in p3_text:
            # Lớp 6 Bài 2
            print(f"  🧹 Fixing duplicated header block in: {os.path.basename(filepath)}")
            p1_elem = doc.paragraphs[1]._p
            p2_elem = doc.paragraphs[2]._p
            p3_elem = doc.paragraphs[3]._p
            p1_elem.getparent().remove(p1_elem)
            p2_elem.getparent().remove(p2_elem)
            p3_elem.getparent().remove(p3_elem)
            modified = True
        elif p0_text.startswith("TÊN BÀI DẠY: BÀI 2") and p2_text == "Tiết theo PPCT: 1" and "Lược sử" in p3_text:
            # Lớp 8 Bài 2
            print(f"  🧹 Fixing duplicated header block in: {os.path.basename(filepath)}")
            p1_elem = doc.paragraphs[1]._p
            p2_elem = doc.paragraphs[2]._p
            p3_elem = doc.paragraphs[3]._p
            p1_elem.getparent().remove(p1_elem)
            p2_elem.getparent().remove(p2_elem)
            p3_elem.getparent().remove(p3_elem)
            modified = True

    if modified:
        try:
            doc.save(filepath)
            print(f"  ✅ Fixed: {os.path.basename(filepath)}")
        except PermissionError:
            print(f"  ✗ Permission error: {filepath}")

def process_directory(base_dir):
    print(f"\n==========================================")
    print(f"Processing: {base_dir}")
    print(f"==========================================")
    for root, dirs, files in os.walk(base_dir):
        for f in files:
            if f.endswith('.docx') and not f.startswith('~$'):
                fix_file(os.path.join(root, f))

if __name__ == '__main__':
    process_directory(r'd:\UNIGO\KHBD_Tin_học')
    process_directory(r'd:\UNIGO\KHBD_Robotics')
    print("\n🎉 HOÀN THÀNH SỬA TOÀN BỘ VIỀN BẢNG VÀ LỖI A1A1!")
