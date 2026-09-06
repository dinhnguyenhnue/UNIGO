import os
import sys
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_borders(cell, top=True, bottom=True, left=True, right=True, color="000000", sz="4"):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side, active in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        if active:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'single')
            b.set(qn('w:sz'), sz)
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), color)
            tcBorders.append(b)
        else:
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'nil')
            tcBorders.append(b)
    tcPr.append(tcBorders)

def set_table_borders(table, color="000000", sz="4"):
    tblPr = table._tbl.tblPr
    tblBorders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>\n'
        f'  <w:top w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:left w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:bottom w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:right w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideH w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'  <w:insideV w:val="single" w:sz="{sz}" w:space="0" w:color="{color}"/>\n'
        f'</w:tblBorders>'
    )
    tblPr.append(tblBorders)

def set_cell_background(cell, color_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:val="clear" w:color="auto" w:fill="{color_hex}"/>')
    tcPr.append(shd)

def apply_font_all(doc, font_name="Times New Roman", font_size=13):
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.name = font_name
            if not r.font.size:
                r.font.size = Pt(font_size)
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = font_name
                        if not r.font.size:
                            r.font.size = Pt(font_size)

def build_vanmieu_proposal(output_path):
    doc = Document()
    
    # Page setup - Margins: Top 2cm, Bottom 2cm, Left 2.5cm, Right 2cm
    for s in doc.sections:
        s.top_margin = Inches(0.79)
        s.bottom_margin = Inches(0.79)
        s.left_margin = Inches(0.98)
        s.right_margin = Inches(0.79)

    # 1. Header Organization Table (2 columns, no border)
    head_tbl = doc.add_table(rows=1, cols=2)
    head_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    head_tbl.autofit = False
    
    # Left: School info
    c0 = head_tbl.rows[0].cells[0]
    c0.width = Inches(3.2)
    p0 = c0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_after = Pt(2)
    r = p0.add_run("TRƯỜNG TIỂU HỌC & THCS UNIGO\n")
    r.font.size = Pt(12)
    r.font.bold = True
    r = p0.add_run("TỔ CHUYÊN MÔN THCS\n")
    r.font.size = Pt(12)
    r.font.bold = True
    r = p0.add_run("Số: ... /BC-KH-UNIGO")
    r.font.size = Pt(11)
    r.font.italic = True

    # Right: National Header
    c1 = head_tbl.rows[0].cells[1]
    c1.width = Inches(3.8)
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_after = Pt(2)
    r = p1.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
    r.font.size = Pt(12)
    r.font.bold = True
    r = p1.add_run("Độc lập – Tự do – Hạnh phúc\n")
    r.font.size = Pt(12)
    r.font.bold = True
    r = p1.add_run("---------------------------\n")
    r.font.size = Pt(10)
    r = p1.add_run("Hà Nội, ngày 31 tháng 08 năm 2026")
    r.font.size = Pt(12)
    r.font.italic = True

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_after = Pt(6)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(4)
    r = p_title.add_run("KẾ HOẠCH & BÁO CÁO ĐỀ XUẤT")
    r.font.size = Pt(15)
    r.font.bold = True

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(12)
    r = p_sub.add_run("THAM QUAN TRẢI NGHIỆM GIÁO DỤC DI SẢN TẠI VĂN MIẾU – QUỐC TỬ GIÁM\n")
    r.font.size = Pt(13.5)
    r.font.bold = True
    r = p_sub.add_run("(Dành cho học sinh khối Trung học cơ sở: Lớp 6, Lớp 7, Lớp 8)")
    r.font.size = Pt(12.5)
    r.font.italic = True

    # Kính gửi
    p_to = doc.add_paragraph()
    p_to.paragraph_format.space_after = Pt(8)
    r = p_to.add_run("Kính gửi: ")
    r.font.bold = True
    r = p_to.add_run("Ban Giám hiệu Trường Tiểu học và THCS UNIGO\n")
    r.font.bold = True
    r = p_to.add_run("                  Đồng kính gửi: Quý Phụ huynh học sinh Khối THCS (Lớp 6, 7, 8)")
    r.font.italic = True

    # Section I
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("I. MỤC ĐÍCH & Ý NGHĨA HOẠT ĐỘNG")
    r.font.bold = True
    r.font.size = Pt(13)

    bullet_points_I = [
        ("Giáo dục truyền thống & đạo lý dân tộc: ", "Giúp học sinh trực tiếp tìm hiểu lịch sử trường đại học đầu tiên của Việt Nam, thấm nhuần truyền thống hiếu học, tinh thần tôn sư trọng đạo và đạo lý “Uống nước nhớ nguồn”."),
        ("Tích hợp giáo dục liên môn & thực tiễn: ", "Gắn kết kiến thức các môn học: Lịch sử & Địa lý (lịch sử Thăng Long, các triều đại Lý – Trần – Lê), Ngữ văn (văn bia tiến sĩ, thơ văn cổ), Giáo dục địa phương (di tích Hà Nội), Hoạt động trải nghiệm hướng nghiệp và Tin học (truyền thông số, nhiếp ảnh di sản)."),
        ("Phát triển năng lực & kỹ năng mềm: ", "Rèn luyện năng lực quan sát, kỹ năng làm việc nhóm, kỹ năng giải mã thông tin qua phiếu điều tra di sản, kỹ năng thuyết trình và năng lực cảm thụ nghệ thuật kiến trúc cổ truyền.")
    ]
    for b_title, b_desc in bullet_points_I:
        bp = doc.add_paragraph()
        bp.paragraph_format.left_indent = Inches(0.25)
        bp.paragraph_format.space_after = Pt(3)
        r1 = bp.add_run("• " + b_title)
        r1.font.bold = True
        r2 = bp.add_run(b_desc)

    # Section II
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("II. TỔNG HỢP CÁC GÓI CHỦ ĐỀ TRẢI NGHIỆM TẠI VĂN MIẾU – QUỐC TỬ GIÁM")
    r.font.bold = True
    r.font.size = Pt(13)

    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = Pt(6)
    p_intro.add_run("Căn cứ chương trình Giáo dục di sản chuẩn của Trung tâm Hoạt động Văn hóa Khoa học Văn Miếu – Quốc Tử Giám (cập nhật chính thức tại vanmieu.gov.vn), các chủ đề được phân bổ theo cấp học và lứa tuổi như sau:")

    # Table 1: Packages overview
    tbl1 = doc.add_table(rows=1, cols=4)
    set_table_borders(tbl1)
    tbl1.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl1.autofit = False

    col_widths = [Inches(0.6), Inches(2.2), Inches(1.3), Inches(2.9)]
    headers = ["STT", "Tên chủ đề / Gói trải nghiệm", "Đối tượng phù hợp", "Nội dung & Hoạt động thực hành chính"]
    
    for i, h in enumerate(headers):
        c = tbl1.rows[0].cells[i]
        c.width = col_widths[i]
        set_cell_background(c, "EAEAEA")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(12)

    topics_data = [
        # Nhóm Lớp 6 (thuộc dải 4-6)
        ("1", "Đi tìm linh vật trên kiến trúc cổ Văn Miếu", "Lớp 6, 7", "Quan sát, phân biệt và giải mã ý nghĩa các linh vật (Rồng, Phượng, Nghê, Rùa) trên Tứ trụ, bia đá, mái đình, cổng Đại Thành."),
        ("2", "Khám phá bia Tiến sĩ & In dập hoa văn", "Lớp 6, 7, 8", "Tìm hiểu giá trị 82 bia Tiến sĩ (Di sản tư liệu thế giới), thực hành in dập các mảng hoa văn trang trí văn bia bằng mực nho và giấy dó."),
        ("3", "Khám phá nghệ thuật Khuê Văn Các", "Lớp 6", "Tìm hiểu biểu tượng văn hiến Thăng Long, cấu trúc 'trời tròn đất vuông', tỷ lệ kiến trúc gỗ và ý nghĩa sao Khuê."),
        ("4", "Lớp học xưa & Tập viết Thư pháp", "Lớp 6, 7", "Tìm hiểu nếp học của học trò xưa, đạo thầy trò, công cụ bút nghiên mực lông và thực hành viết chữ thư pháp Hán - Nôm cơ bản."),
        ("5", "Chương trình 'Ơ kìa con nghê!'", "Lớp 6", "Khám phá hình tượng Nghê thuần Việt ở các vị trí (Tứ trụ, bia đá, cổng Đại Thành, áo Khổng Tử) và làm bài tập phân biệt linh vật."),
        # Nhóm Lớp 7 - 8 (thuộc dải 7-12)
        ("6", "Khám phá Sách học & Ván khắc in sách Quốc Tử Giám", "Lớp 7, 8\n(Đặc sắc)", "Tiếp cận di sản mộc bản, tự tay làm ván in, thực hành kỹ thuật in ấn dập mực cổ truyền và trải nghiệm đóng sách xưa."),
        ("7", "Chương trình 'Thi Hương, Thi Hội, Thi Đình'", "Lớp 7, 8\n(Trọng tâm)", "Tìm hiểu lịch sử thi cử Nho học, tự tay dựng lều chõng như sĩ tử xưa, mô phỏng quy chế thi cử, kỷ luật trường thi và lễ xướng danh."),
        ("8", "Tìm hiểu Quốc Tử Giám ở Thăng Long & Danh nhân", "Lớp 7, 8", "Tìm hiểu lịch sử hình thành trường Quốc học đầu tiên qua các thời Lý - Trần - Lê, tấm gương Danh sư Chu Văn An, Thân Nhân Trung."),
        ("9", "Chương trình 'Vinh quy bái tổ'", "Lớp 7, 8", "Thảo luận nhóm về truyền thống 'Uống nước nhớ nguồn', ý nghĩa của lễ rước Trạng nguyên và trách nhiệm của người tài với đất nước."),
        ("10", "Văn Miếu xưa và nay & Nhiếp ảnh trẻ", "Lớp 7, 8", "So sánh tư liệu ảnh Văn Miếu xưa và hiện trạng, thảo luận sự thay đổi/bảo tồn và thực hành làm 'nhiếp ảnh gia trẻ' ghi lại vẻ đẹp di sản."),
        ("11", "Đánh giá bảo vệ di sản & Hiến kế môi trường", "Lớp 7, 8", "Khảo sát thực tế di tích, nhận diện hư hại, thảo luận nhóm và lập đề xuất 'Hiến kế' nâng cao tiện ích và bảo tồn cảnh quan di sản.")
    ]

    for row_data in topics_data:
        row = tbl1.add_row()
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            cell.width = col_widths[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            r = p.add_run(text)
            r.font.size = Pt(11.5)
            if idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif idx == 2:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if "Đặc sắc" in text or "Trọng tâm" in text:
                    r.font.bold = True
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Section III: Proposed Packages
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("III. ĐỀ XUẤT CÁC GÓI TRẢI NGHIỆM TỐI ƯU CHO HỌC SINH UNIGO (LỚP 6-8)")
    r.font.bold = True
    r.font.size = Pt(13)

    packages_detail = [
        ("GÓI ĐỀ XUẤT 1: 'HÀNH TRÌNH SĨ TỬ VÀ NGHỀ IN SÁCH CỔ' (Khuyên chọn số 1 cho Khối 7, 8)", [
            "Hoạt động 1: Lễ dâng hương báo công tại Nhà Tiền Đường/Nhà Thái Học và nghe thuyết minh tổng quan về di tích.",
            "Hoạt động 2: Trải nghiệm 'Thi Hương - Thi Hội - Thi Đình': Học sinh chia nhóm tự tay dựng lều chõng, nghe phổ biến quy chế trường thi Nho học.",
            "Hoạt động 3: Trải nghiệm 'Mộc bản & Đóng sách': Tự tay quét mực, in dập văn bản từ mộc bản và đóng sách chỉ truyền thống mang về làm kỷ niệm.",
            "Hoạt động 4: Khám phá 82 bia Tiến sĩ, giải mã bài học 'Hiền tài là nguyên khí của quốc gia' qua phiếu học tập tương tác."
        ]),
        ("GÓI ĐỀ XUẤT 2: 'GIẢI MÃ BÁU VẬT DI SẢN & NHIẾP ẢNH TRẺ' (Khuyên chọn cho Khối 6, 7)", [
            "Hoạt động 1: Lễ dâng hương tri ân các bậc Danh sư tiền bối.",
            "Hoạt động 2: Trò chơi khám phá 'Đi tìm linh vật trên kiến trúc cổ': Học sinh dùng bản đồ truy tìm Rồng, Phượng, Nghê, Rùa tại các phân khu.",
            "Hoạt động 3: Thực hành in dập hoa văn trang trí bia Tiến sĩ bằng mực và giấy dó cổ truyền.",
            "Hoạt động 4: Hoạt động 'Văn Miếu xưa và nay': So sánh ảnh lịch sử với thực tế và thực hành chụp ảnh ghi lại các góc kiến trúc độc đáo."
        ]),
        ("GÓI ĐỀ XUẤT 3: 'THƯ PHÁP CỔ & HIẾN KẾ BẢO VỆ DI SẢN' (Gói Tích hợp Liên môn Ngữ văn - GDĐP)", [
            "Hoạt động 1: Thuyết minh chuyên đề 'Quốc Tử Giám - Đỉnh cao giáo dục Nho học Việt Nam'.",
            "Hoạt động 2: Trải nghiệm 'Lớp học xưa': Ngồi chiếu hoa, mài mực lông, học lễ nghĩa và viết chữ Thư pháp khai tâm (chữ 'Tâm', 'Trí', 'Đức', 'Thành').",
            "Hoạt động 3: Thảo luận nhóm 'Hiến kế bảo tồn di sản': Khảo sát hiện trạng, lập biên bản ghi nhận và thuyết trình giải pháp bảo vệ di tích."
        ])
    ]

    for g_title, g_acts in packages_detail:
        p_g = doc.add_paragraph()
        p_g.paragraph_format.left_indent = Inches(0.15)
        p_g.paragraph_format.space_before = Pt(4)
        p_g.paragraph_format.space_after = Pt(2)
        r = p_g.add_run("★ " + g_title)
        r.font.bold = True
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(140, 20, 20)
        for act in g_acts:
            p_a = doc.add_paragraph()
            p_a.paragraph_format.left_indent = Inches(0.4)
            p_a.paragraph_format.space_after = Pt(2)
            r_a = p_a.add_run("– " + act)
            r_a.font.size = Pt(11.5)

    # Section IV: Detailed Budget & Cost estimation
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("IV. DỰ TOÁN KINH PHÍ & BẢNG GIÁ CHI TIẾT")
    r.font.bold = True
    r.font.size = Pt(13)

    p_cost_note = doc.add_paragraph()
    p_cost_note.paragraph_format.space_after = Pt(4)
    p_cost_note.add_run("Dự toán kinh phí được xây dựng trên cơ sở chính sách giá vé di tích, chi phí dịch vụ giáo dục di sản chính thức và các chi phí tổ chức hậu cần an toàn cho học sinh:")

    # Table 2: Budget Breakdown
    tbl2 = doc.add_table(rows=1, cols=5)
    set_table_borders(tbl2)
    tbl2.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl2.autofit = False

    col_widths2 = [Inches(0.5), Inches(2.3), Inches(1.3), Inches(1.4), Inches(1.5)]
    headers2 = ["STT", "Hạng mục chi phí", "Đơn giá dự kiến (VNĐ/HS)", "Định mức / Ghi chú", "Thành tiền (Dự kiến)"]
    
    for i, h in enumerate(headers2):
        c = tbl2.rows[0].cells[i]
        c.width = col_widths2[i]
        set_cell_background(c, "EAEAEA")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(11.5)

    budget_data = [
        ("1", "Vé vào cổng tham quan di tích", "0 – 35.000 đ", "Học sinh dưới 16 tuổi miễn phí / Có thẻ HS giảm 50%", "Miễn phí hoặc 35.000 đ"),
        ("2", "Phí chương trình Giáo dục di sản & Thuyết minh", "40.000 – 60.000 đ", "Bao gồm HDV chuyên tuyến + Phiếu học tập di sản", "50.000 đ"),
        ("3", "Vật tư thực hành trải nghiệm (Giấy dó, mực in, mộc bản, khung lều)", "30.000 – 40.000 đ", "Học sinh được mang sản phẩm (sách, bản in, chữ) về", "35.000 đ"),
        ("4", "Xe ô tô đưa đón chất lượng cao (45 chỗ)", "40.000 – 50.000 đ", "Đưa đón 2 chiều UNIGO ↔ Văn Miếu", "45.000 đ"),
        ("5", "Nước uống, đồ ăn nhẹ & Y tế dự phòng", "15.000 – 20.000 đ", "Nước suối đóng chai + bánh ngọt dinh dưỡng", "15.000 đ"),
        ("6", "Bảo hiểm du lịch trải nghiệm", "5.000 đ", "Mức trách nhiệm bảo hiểm 50.000.000 đ/vụ", "5.000 đ"),
        ("7", "Hương hoa, lễ phẩm dâng hương báo công", "–", "Chi phí chung của toàn đoàn (Nhà trường hỗ trợ)", "Nhà trường tài trợ")
    ]

    for row_data in budget_data:
        row = tbl2.add_row()
        for idx, text in enumerate(row_data):
            cell = row.cells[idx]
            cell.width = col_widths2[idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.space_before = Pt(2)
            r = p.add_run(text)
            r.font.size = Pt(11)
            if idx in [0, 2, 4]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Total Summary Box
    p_sum = doc.add_paragraph()
    p_sum.paragraph_format.space_before = Pt(6)
    p_sum.paragraph_format.space_after = Pt(4)
    r = p_sum.add_run("► TỔNG KINH PHÍ ĐỀ XUẤT THEO PHƯƠNG ÁN TỔ CHỨC:\n")
    r.font.bold = True
    r.font.size = Pt(12)
    
    p_opt1 = doc.add_paragraph()
    p_opt1.paragraph_format.left_indent = Inches(0.25)
    p_opt1.paragraph_format.space_after = Pt(2)
    r1 = p_opt1.add_run("• Phương án 1 (Nửa ngày - Sáng hoặc Chiều, không ăn trưa): ")
    r1.font.bold = True
    r2 = p_opt1.add_run("Khoảng ")
    r3 = p_opt1.add_run("150.000 VNĐ / học sinh")
    r3.font.bold = True
    r4 = p_opt1.add_run(" (Đã bao gồm trọn gói: xe đưa đón, vé, hướng dẫn viên, học liệu thực hành, nước uống, bảo hiểm).")

    p_opt2 = doc.add_paragraph()
    p_opt2.paragraph_format.left_indent = Inches(0.25)
    p_opt2.paragraph_format.space_after = Pt(4)
    r1 = p_opt2.add_run("• Phương án 2 (Cả ngày - Kết hợp ăn trưa & dã ngoại Bảo tàng Lịch sử / Hoàng thành): ")
    r1.font.bold = True
    r2 = p_opt2.add_run("Khoảng ")
    r3 = p_opt2.add_run("260.000 – 290.000 VNĐ / học sinh")
    r3.font.bold = True
    r4 = p_opt2.add_run(" (Bao gồm thêm suất ăn trưa tiêu chuẩn 60.000 - 70.000 đ + vé điểm tham quan thứ 2).")

    # Section V: Suggested Schedule
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("V. LỊCH TRÌNH THỰC HIỆN DỰ KIẾN (LỰA CHỌN PHƯƠNG ÁN NỬA NGÀY BUỔI SÁNG)")
    r.font.bold = True
    r.font.size = Pt(13)

    schedule_data = [
        ("07h30 – 07h50", "Học sinh tập trung tại sân trường UNIGO, GVCN điểm danh, kiểm tra trang phục, phát thẻ đoàn và phổ biến an toàn."),
        ("07h50 – 08h30", "Học sinh lên xe ô tô, di chuyển đến Di tích Văn Miếu – Quốc Tử Giám."),
        ("08h30 – 08h45", "Tập trung tại cổng Văn Miếu (Cổng Đại Trung), chỉnh đốn trang phục, di chuyển vào sân Nhà Thái Học."),
        ("08h45 – 09h00", "Làm Lễ dâng hương báo công và tưởng niệm Danh sư Chu Văn An cùng các bậc Tiên thánh, Tiên hiền."),
        ("09h00 – 10h00", "Chia lớp thành 3 nhóm chuyên đề tham quan dưới sự dẫn dắt của Cán bộ Giáo dục di sản:\n+ Nhóm 1: Khám phá 82 Bia Tiến sĩ & Giếng Thiên Quang.\n+ Nhóm 2: Khám phá kiến trúc Khuê Văn Các & Điện Đại Thành.\n+ Nhóm 3: Khảo sát hiện trạng di sản & giải mã linh vật Nghê, Rồng."),
        ("10h00 – 11h00", "Tập trung tại Không gian Trải nghiệm Giáo dục di sản:\n+ Thực hành Dựng lều chõng sĩ tử (Khối 7, 8).\n+ Thực hành In dập ván khắc gỗ & đóng sách xưa / viết chữ thư pháp.\n+ Hoàn thiện phiếu thu hoạch cá nhân/nhóm."),
        ("11h00 – 11h20", "Tổng kết hoạt động, trao phần thưởng cho các nhóm xuất sắc, chụp ảnh lưu niệm toàn khối trước Khuê Văn Các."),
        ("11h20 – 12h00", "Tập hợp học sinh, kiểm diện sĩ số, lên xe trở về trường UNIGO. Kết thúc chuyến tham quan an toàn, ý nghĩa.")
    ]

    tbl3 = doc.add_table(rows=1, cols=2)
    set_table_borders(tbl3)
    tbl3.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl3.autofit = False

    col_widths3 = [Inches(1.5), Inches(5.5)]
    headers3 = ["Thời gian", "Nội dung hoạt động chi tiết"]
    for i, h in enumerate(headers3):
        c = tbl3.rows[0].cells[i]
        c.width = col_widths3[i]
        set_cell_background(c, "EAEAEA")
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(11.5)

    for time_str, desc in schedule_data:
        row = tbl3.add_row()
        c0 = row.cells[0]
        c0.width = col_widths3[0]
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(time_str)
        r0.font.bold = True
        r0.font.size = Pt(11)

        c1 = row.cells[1]
        c1.width = col_widths3[1]
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.LEFT
        r1 = p1.add_run(desc)
        r1.font.size = Pt(11)

    # Section VI: Safety & Rules
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run("VI. QUY ĐỊNH AN TOÀN, NỘI QUY & PHÂN CÔNG NHIỆM VỤ")
    r.font.bold = True
    r.font.size = Pt(13)

    rules = [
        ("Nội quy Di tích (Bắt buộc): ", "Học sinh tuyệt đối KHÔNG chạm tay, xoa đầu rùa, không viết/vẽ/đứng/ngồi lên bia Tiến sĩ; không bẻ cành hái hoa, không vứt rác bừa bãi; giữ thái độ trang nghiêm tại khu vực thờ tự Điện Đại Thành."),
        ("Quy định trang phục: ", "100% học sinh mặc đồng phục trường UNIGO chỉnh tề, đi giày thể thao hoặc quai hậu thuận tiện di chuyển ngoài trời."),
        ("Dụng cụ học sinh cần mang: ", "Mũ nón che nắng, bình nước uống cá nhân, bút viết và sổ tay ghi chép (nhà trường phát phiếu thu hoạch và tài liệu trải nghiệm)."),
        ("Phân công quản lý học sinh: ", "Tỷ lệ quản lý đảm bảo tối thiểu 1 Giáo viên / Cán bộ phụ trách 12 – 15 học sinh. Giáo viên chủ nhiệm các lớp 6, 7, 8 theo sát lớp trong suốt hành trình và phối hợp chặt chẽ với Cán bộ Giáo dục của Văn Miếu.")
    ]
    for r_title, r_desc in rules:
        bp = doc.add_paragraph()
        bp.paragraph_format.left_indent = Inches(0.25)
        bp.paragraph_format.space_after = Pt(3)
        r1 = bp.add_run("• " + r_title)
        r1.font.bold = True
        r2 = bp.add_run(r_desc)

    # Section VII: Contact info
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.space_before = Pt(8)
    r = p.add_run("VII. THÔNG TIN LIÊN HỆ ĐĂNG KÝ VỚI BAN QUẢN LÝ DI TÍCH")
    r.font.bold = True
    r.font.size = Pt(13)

    p_contact = doc.add_paragraph()
    p_contact.paragraph_format.left_indent = Inches(0.25)
    p_contact.paragraph_format.space_after = Pt(8)
    p_contact.add_run("• Đơn vị tổ chức: Trung tâm Hoạt động Văn hóa Khoa học Văn Miếu – Quốc Tử Giám\n"
                      "• Phòng Giáo dục Truyền thông: 024.38235601 (Máy lẻ: 12)\n"
                      "• Hotline hỗ trợ đặt đoàn & giáo dục di sản: 0983.050.103 / 0932.070.109 / 0973.211.901 / 0986.060.566\n"
                      "• Website: https://vanmieu.gov.vn | Email: tthdvhkhvmqtg_svhtt@hanoi.gov.vn")

    # Section VIII: Conclusion & Signatures (Table 3x3 no border)
    p_rec = doc.add_paragraph()
    p_rec.paragraph_format.space_before = Pt(6)
    p_rec.paragraph_format.space_after = Pt(12)
    r_rec = p_rec.add_run("Kính trình Ban Giám hiệu xem xét, phê duyệt chủ trương và phương án tổ chức để Tổ Chuyên môn phối hợp cùng Ban Phụ huynh triển khai công tác chuẩn bị chu đáo nhất./.")
    r_rec.font.italic = True

    sig_tbl = doc.add_table(rows=3, cols=3)
    sig_tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    sig_tbl.autofit = False

    sig_widths = [Inches(2.3), Inches(2.3), Inches(2.4)]
    
    # Row 0: Titles
    titles = ["DUYỆT CỦA BAN GIÁM HIỆU", "TỔ TRƯỞNG CHUYÊN MÔN THCS", "NGƯỜI LẬP KẾ HOẠCH"]
    for i, t in enumerate(titles):
        c = sig_tbl.rows[0].cells[i]
        c.width = sig_widths[i]
        set_cell_borders(c, top=False, bottom=False, left=False, right=False)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(t)
        r.font.bold = True
        r.font.size = Pt(12)

    # Row 1: Subtitles
    subtitles = ["(Ký và ghi rõ họ tên)", "(Ký và ghi rõ họ tên)", "(Ký và ghi rõ họ tên)"]
    for i, st in enumerate(subtitles):
        c = sig_tbl.rows[1].cells[i]
        c.width = sig_widths[i]
        set_cell_borders(c, top=False, bottom=False, left=False, right=False)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(st)
        r.font.italic = True
        r.font.size = Pt(11)

    # Row 2: Names
    names = ["\n\n\n\n........................................", "\n\n\n\n........................................", "\n\n\n\nĐậu Đình Nguyên"]
    for i, n in enumerate(names):
        c = sig_tbl.rows[2].cells[i]
        c.width = sig_widths[i]
        set_cell_borders(c, top=False, bottom=False, left=False, right=False)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(n)
        r.font.bold = True
        r.font.size = Pt(12)

    # Apply font formatting across all paragraphs and tables
    apply_font_all(doc, "Times New Roman", 13)

    # Ensure output directories exist
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Successfully generated: {output_path}")

if __name__ == "__main__":
    out1 = r"d:\UNIGO\Kế hoạch tham quan Văn Miếu - Quốc Tử Giám (Lớp 6-8).docx"
    out2 = r"d:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Kế hoạch tham quan Văn Miếu - Quốc Tử Giám (Lớp 6-8).docx"
    build_vanmieu_proposal(out1)
    build_vanmieu_proposal(out2)
