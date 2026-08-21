"""
fix_khbd_robotics_header_and_signature.py
Sửa toàn bộ KHBD Robotics (Lớp 1 - 8):
1. Phần thứ ngày & thông tin trường/lớp (Table 0):
   - Row 0: Trường: TH&THCS UNIGO | Ngày soạn: DD/MM/YYYY
   - Row 1: GV: Đậu Đình Nguyên\nTổ: Tổ chuyên môn... | Ngày dạy: DD/MM/YYYY\nLớp: XA1
   - Table 0 NO BORDER
2. Xóa các paragraph thừa ở đầu trang (TUẦN: ..., Ngày soạn/dạy trùng lặp)
3. Chuẩn hóa tên lớp trong tiêu đề (Lớp 1A1, 2A1, 3A1, 4C1, 5C1, 6A1, 7A1, 8A1)
4. Phần Rút kinh nghiệm sau bài dạy
5. Bảng chữ ký cuối bài (Table 3x3 NO BORDER):
   - Row 0: DUYỆT CỦA BGH | DUYỆT CỦA TỔ CM | NGƯỜI SOẠN (In đậm, Căn giữa)
   - Row 1: (Ký, ghi rõ họ tên) | (Ký, ghi rõ họ tên) | (Ký, ghi rõ họ tên) (In nghiêng, Căn giữa)
   - Row 2: \n\n\n | \n\n\n | \n\n\nĐậu Đình Nguyên (Căn giữa, tên in đậm)
"""

import os, sys, re
sys.stdout.reconfigure(encoding='utf-8')
from datetime import date, timedelta
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

TUAN_01_START = date(2026, 8, 3)

ROBOTICS_SCHEDULE = {
    '1': (3, '1A1'),   # Thứ Năm
    '2': (2, '2A1'),   # Thứ Tư
    '3': (2, '3A1'),   # Thứ Tư
    '4': (4, '4C1'),   # Thứ Sáu
    '5': (1, '5C1'),   # Thứ Ba
    '6': (4, '6A1'),   # Thứ Sáu
    '7': (1, '7A1'),   # Thứ Ba
    '8': (4, '8A1'),   # Thứ Sáu
}

KHBD_BASE = r'd:\UNIGO\KHBD_Robotics'

def compute_dates(tuan_so, day_of_week):
    week_start = TUAN_01_START + timedelta(weeks=tuan_so - 1)
    ngay_day = (week_start + timedelta(days=day_of_week)).strftime('%d/%m/%Y')
    ngay_soan = (week_start - timedelta(days=2)).strftime('%d/%m/%Y')
    return ngay_soan, ngay_day

def set_table_no_borders(tbl):
    """Xóa toàn bộ viền bảng (Table borders = none & tcBorders = removed)."""
    tblPr = tbl._tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl._tbl.insert(0, tblPr)
    old_b = tblPr.find(qn('w:tblBorders'))
    if old_b is not None:
        tblPr.remove(old_b)
    borders_xml = parse_xml(
        r'<w:tblBorders %s>'
        r'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        r'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        r'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        r'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        r'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        r'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        r'</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(borders_xml)
    for r in tbl._tbl.findall(qn('w:tr')):
        for c in r.findall(qn('w:tc')):
            tcPr = c.find(qn('w:tcPr'))
            if tcPr is not None:
                tcB = tcPr.find(qn('w:tcBorders'))
                if tcB is not None:
                    tcPr.remove(tcB)

def set_run_font(run, size_pt=13, bold=False, italic=False):
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), 'Times New Roman')

def set_cell_paragraph(cell, p_idx, text_runs, align=WD_ALIGN_PARAGRAPH.LEFT):
    """
    text_runs: list of (text, bold, italic, size_pt)
    """
    while len(cell.paragraphs) <= p_idx:
        cell.add_paragraph()
    p = cell.paragraphs[p_idx]
    p.text = ''
    p.alignment = align
    p.paragraph_format.line_spacing = 1.15
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    for txt, bold, italic, sz in text_runs:
        r = p.add_run(txt)
        set_run_font(r, size_pt=sz, bold=bold, italic=italic)

def clean_cell_extra_paragraphs(cell, keep_count):
    while len(cell.paragraphs) > keep_count:
        p_elem = cell.paragraphs[-1]._p
        p_elem.getparent().remove(p_elem)

def clean_a1_duplication(text):
    text = re.sub(r'(6|7|8|1|2|3)A1(A1)+', r'\1A1', text)
    text = re.sub(r'(5|4)C1(C1)+', r'\1C1', text)
    return text

def fix_robotics_file(filepath):
    # Determine grade and week from filepath
    m_grade = re.search(r'Lớp_(\d+)', filepath)
    m_week = re.search(r'Tuần_(\d+)', filepath)
    if not m_grade or not m_week:
        print(f'  ⚠️ Cannot parse grade/week: {filepath}')
        return False
    
    grade_key = m_grade.group(1)
    tuan_so = int(m_week.group(1))
    
    if grade_key not in ROBOTICS_SCHEDULE:
        print(f'  ⚠️ Grade not in schedule: {grade_key}')
        return False
    
    day_of_week, ten_lop = ROBOTICS_SCHEDULE[grade_key]
    ngay_soan, ngay_day = compute_dates(tuan_so, day_of_week)
    
    is_thcs = int(grade_key) >= 6
    to_chuyen_mon = "Tổ chuyên môn THCS" if is_thcs else "Tổ chuyên môn Tiểu học"
    
    doc = Document(filepath)
    modified = False

    # -------------------------------------------------------------
    # 1. Sửa Table 0 (Bảng thông tin trường/lớp/ngày soạn/ngày dạy)
    # -------------------------------------------------------------
    if doc.tables:
        t0 = doc.tables[0]
        set_table_no_borders(t0)
        
        # Đảm bảo Table 0 có 2 rows x 2 cols
        while len(t0.rows) < 2:
            t0.add_row()
            
        # Row 0, Cell 0: Trường: TH&THCS UNIGO
        set_cell_paragraph(t0.rows[0].cells[0], 0, [('Trường: TH&THCS UNIGO', True, False, 13)])
        clean_cell_extra_paragraphs(t0.rows[0].cells[0], 1)
        
        # Row 0, Cell 1: Ngày soạn: DD/MM/YYYY
        set_cell_paragraph(t0.rows[0].cells[1], 0, [(f'Ngày soạn: {ngay_soan}', True, False, 13)])
        clean_cell_extra_paragraphs(t0.rows[0].cells[1], 1)
        
        # Row 1, Cell 0: GV: Đậu Đình Nguyên \n Tổ: Tổ chuyên môn...
        set_cell_paragraph(t0.rows[1].cells[0], 0, [
            ('GV: Đậu Đình Nguyên\n', False, False, 13),
            (f'Tổ: {to_chuyen_mon}', False, False, 13)
        ])
        clean_cell_extra_paragraphs(t0.rows[1].cells[0], 1)
        
        # Row 1, Cell 1: Ngày dạy: DD/MM/YYYY \n Lớp: XA1
        set_cell_paragraph(t0.rows[1].cells[1], 0, [
            (f'Ngày dạy: {ngay_day}\n', False, False, 13),
            (f'Lớp: {ten_lop}', False, False, 13)
        ])
        clean_cell_extra_paragraphs(t0.rows[1].cells[1], 1)
        
        modified = True

    # -------------------------------------------------------------
    # 2. Dọn dẹp paragraphs thừa ở đầu trang và chuẩn hóa tên lớp
    # -------------------------------------------------------------
    for p in doc.paragraphs:
        txt = p.text.strip()
        # Xóa dòng TUẦN / Ngày soạn rác nếu xuất hiện trong paragraph
        if txt.startswith('TUẦN:') and ('Ngày soạn:' in txt or 'Ngày dạy:' in txt):
            p.text = ''
            modified = True
        elif txt.startswith('Thứ …… ngày') or txt.startswith('Thứ ..... ngày'):
            p.text = ''
            modified = True

        # Sửa tên lớp trong paragraph
        if 'MÔN HỌC: ROBOTICS' in p.text or 'Môn học: Robotics' in p.text or 'KẾ HOẠCH DẠY HỌC MÔN ROBOTICS' in p.text:
            new_t = re.sub(r'ROBOTICS\s*(\d+)', f'ROBOTICS {ten_lop}', p.text)
            new_t = re.sub(r'LỚP\s*(\d+)', f'LỚP {ten_lop}', new_t)
            new_t = clean_a1_duplication(new_t)
            if new_t != p.text:
                p.text = new_t
                for r in p.runs:
                    set_run_font(r, size_pt=13, bold=True)
                modified = True
                
        if 'Lớp:' in p.text:
            new_t = re.sub(r'Lớp:\s*[0-9A-Za-z]+', f'Lớp: {ten_lop}', p.text)
            new_t = clean_a1_duplication(new_t)
            if new_t != p.text:
                p.text = new_t
                for r in p.runs:
                    set_run_font(r, size_pt=13, bold=True)
                modified = True

    # -------------------------------------------------------------
    # 3. Sửa phần Rút kinh nghiệm & Bảng chữ ký cuối bài (Table ký)
    # -------------------------------------------------------------
    # Tìm bảng ký tên (bảng cuối hoặc bảng có chứa BAN GIÁM HIỆU / DUYỆT CỦA BGH / TỔ CHUYÊN MÔN / NGƯỜI SOẠN)
    sign_table = None
    for tbl in reversed(doc.tables):
        tbl_text = ' '.join(c.text for r in tbl.rows for c in r.cells)
        if any(k in tbl_text.upper() for k in ['BAN GIÁM HIỆU', 'DUYỆT CỦA BGH', 'TỔ CHUYÊN MÔN', 'DUYỆT CỦA TỔ CM', 'NGƯỜI SOẠN']):
            sign_table = tbl
            break

    if sign_table is not None:
        set_table_no_borders(sign_table)
        # Đảm bảo bảng ký có 3 rows x 3 cols
        while len(sign_table.rows) < 3:
            sign_table.add_row()
        while len(sign_table.rows) > 3:
            # Xóa bớt hàng thừa nếu có
            tr_elem = sign_table.rows[-1]._tr
            tr_elem.getparent().remove(tr_elem)

        # Row 0: Tiêu đề duyệt
        headers = ['DUYỆT CỦA BGH', 'DUYỆT CỦA TỔ CM', 'NGƯỜI SOẠN']
        for col_idx, h_text in enumerate(headers):
            cell = sign_table.rows[0].cells[col_idx]
            set_cell_paragraph(cell, 0, [(h_text, True, False, 13)], align=WD_ALIGN_PARAGRAPH.CENTER)
            clean_cell_extra_paragraphs(cell, 1)

        # Row 1: (Ký, ghi rõ họ tên)
        for col_idx in range(3):
            cell = sign_table.rows[1].cells[col_idx]
            set_cell_paragraph(cell, 0, [('(Ký, ghi rõ họ tên)', False, True, 13)], align=WD_ALIGN_PARAGRAPH.CENTER)
            clean_cell_extra_paragraphs(cell, 1)

        # Row 2: Khoảng cách và tên
        for col_idx in range(2):
            cell = sign_table.rows[2].cells[col_idx]
            set_cell_paragraph(cell, 0, [('\n\n\n', False, False, 13)], align=WD_ALIGN_PARAGRAPH.CENTER)
            clean_cell_extra_paragraphs(cell, 1)

        # Row 2 Col 2: Đậu Đình Nguyên
        cell_nguoi_soan = sign_table.rows[2].cells[2]
        set_cell_paragraph(cell_nguoi_soan, 0, [
            ('\n\n\n', False, False, 13),
            ('Đậu Đình Nguyên', True, False, 13)
        ], align=WD_ALIGN_PARAGRAPH.CENTER)
        clean_cell_extra_paragraphs(cell_nguoi_soan, 1)
        modified = True

    # Kiểm tra xem có mục RÚT KINH NGHIỆM SAU BÀI DẠY trước bảng ký chưa
    has_rut_kn = any('RÚT KINH NGHIỆM' in p.text.upper() for p in doc.paragraphs)
    if not has_rut_kn and sign_table is not None:
        # Chèn Rút kinh nghiệm trước bảng ký
        tbl_elem = sign_table._tbl
        p_kn = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        sp = OxmlElement('w:spacing')
        sp.set(qn('w:line'), '276') # 1.15
        sp.set(qn('w:lineRule'), 'auto')
        sp.set(qn('w:after'), '60')
        pPr.append(sp)
        p_kn.append(pPr)
        r_kn = OxmlElement('w:r')
        rPr_kn = OxmlElement('w:rPr')
        rFonts_kn = OxmlElement('w:rFonts')
        for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
            rFonts_kn.set(qn(attr), 'Times New Roman')
        rPr_kn.append(rFonts_kn)
        sz_kn = OxmlElement('w:sz'); sz_kn.set(qn('w:val'), '26')
        rPr_kn.append(sz_kn)
        b_kn = OxmlElement('w:b')
        rPr_kn.append(b_kn)
        r_kn.append(rPr_kn)
        t_kn = OxmlElement('w:t')
        t_kn.text = 'RÚT KINH NGHIỆM SAU BÀI DẠY:'
        r_kn.append(t_kn)
        p_kn.append(r_kn)
        tbl_elem.addprevious(p_kn)

        # Thêm 2 dòng chấm
        for _ in range(2):
            p_dot = OxmlElement('w:p')
            pPr_dot = OxmlElement('w:pPr')
            sp_d = OxmlElement('w:spacing')
            sp_d.set(qn('w:line'), '276')
            sp_d.set(qn('w:after'), '60')
            pPr_dot.append(sp_d)
            p_dot.append(pPr_dot)
            r_d = OxmlElement('w:r')
            rPr_d = OxmlElement('w:rPr')
            rFonts_d = OxmlElement('w:rFonts')
            for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
                rFonts_d.set(qn(attr), 'Times New Roman')
            rPr_d.append(rFonts_d)
            sz_d = OxmlElement('w:sz'); sz_d.set(qn('w:val'), '26')
            rPr_d.append(sz_d)
            r_d.append(rPr_d)
            t_d = OxmlElement('w:t')
            t_d.text = '...........................................................................................................................'
            r_d.append(t_d)
            p_dot.append(r_d)
            tbl_elem.addprevious(p_dot)
        
        modified = True

    if modified:
        try:
            doc.save(filepath)
            print(f'  ✅ Fixed: {os.path.basename(filepath)}')
            return True
        except PermissionError:
            alt = filepath.replace('.docx', '_fixed.docx')
            doc.save(alt)
            print(f'  ⚠️ Saved as: {os.path.basename(alt)} (original locked)')
            return True
    else:
        print(f'  ⏩ No change: {os.path.basename(filepath)}')
        return False

def main():
    print("==================================================")
    print(" BẮT ĐẦU SỬA THỨ NGÀY VÀ PHẦN KÝ TOÀN BỘ KHBD ROBOTICS")
    print("==================================================")
    
    total = 0
    fixed = 0
    for root, dirs, files in os.walk(KHBD_BASE):
        for f in sorted(files):
            if f.endswith('.docx') and not f.startswith('~$'):
                total += 1
                full_path = os.path.join(root, f)
                if fix_robotics_file(full_path):
                    fixed += 1

    print(f"\n🎉 HOÀN THÀNH: Đã xử lý {fixed}/{total} files KHBD Robotics!")

if __name__ == '__main__':
    main()
