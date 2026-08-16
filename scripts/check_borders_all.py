import sys, os
sys.stdout.reconfigure(encoding='utf-8')
from docx import Document
from docx.oxml.ns import qn

def check_all(base_dir):
    print(f"=== CHECKING ALL TABLES IN: {base_dir} ===")
    for root, dirs, files in os.walk(base_dir):
        for f in files:
            if f.endswith('.docx') and not f.startswith('~$'):
                full_p = os.path.join(root, f)
                try:
                    d = Document(full_p)
                    if d.tables:
                        t0 = d.tables[0]
                        tblPr = t0._tbl.find(qn('w:tblPr'))
                        tblBorders = tblPr.find(qn('w:tblBorders')) if tblPr is not None else None
                        st = t0.style.name if t0.style else ""
                        
                        has_single = False
                        if tblBorders is not None:
                            for b in tblBorders:
                                if b.attrib.get(qn('w:val')) == 'single':
                                    has_single = True
                        elif 'Grid' in st:
                            has_single = True
                            
                        # Also check if any cell has single borders
                        for r in t0._tbl.findall(qn('w:tr')):
                            for c in r.findall(qn('w:tc')):
                                tcPr = c.find(qn('w:tcPr'))
                                if tcPr is not None:
                                    tcB = tcPr.find(qn('w:tcBorders'))
                                    if tcB is not None:
                                        for b in tcB:
                                            if b.attrib.get(qn('w:val')) == 'single':
                                                has_single = True
                        if has_single or 'Grid' in st:
                            print(f"  [HAS BORDER] {f} (Style: {st})")
                except Exception as e:
                    pass

check_all(r'd:\UNIGO\KHBD_Tin_học')
check_all(r'd:\UNIGO\KHBD_Robotics')
