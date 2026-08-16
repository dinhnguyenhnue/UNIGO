"""
fix_khbd_dates.py — Sửa file KHBD đã tạo cho Lớp 5-8 (Tuần 1-3):
  1. Xóa dòng "Lớp" thừa trong Table[0] Row[1] Cell[1] (THCS)
  2. Điền ngày soạn / ngày dạy đúng theo LBG
  3. Đổi tên lớp THCS: 6→6A1, 7→7A1, 8→8A1
  4. Sửa paragraph "Môn học: ... Lớp: ..." (THCS)
  5. Sửa header "Ngày soạn/Ngày dạy" cho TH (Lớp 5)
"""
import os, sys, glob, re
sys.stdout.reconfigure(encoding='utf-8')
from datetime import date, timedelta
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

# ─── SCHEDULE CONFIG ───────────────────────────────────────────────────────
TUAN_01_START = date(2026, 8, 3)

LOP_SCHEDULE = {
    # Tiền Tiểu học
    'TTH':  (3, 'TT3'),    # Thứ Năm
    # Tiểu học
    '1':  (0, '1A1'),       # Thứ Hai
    '2':  (1, '2A1'),       # Thứ Ba
    '3':  (3, '3A1'),       # Thứ Năm
    '4':  (2, '4C1'),       # Thứ Tư
    '5':  (1, '5C1'),       # Thứ Ba
    # THCS
    '6':  (4, '6A1'),       # Thứ Sáu
    '7':  (1, '7A1'),       # Thứ Ba
    '8':  (4, '8A1'),       # Thứ Sáu
}

KHBD_BASE = r'd:\UNIGO\KHBD_Tin_học'

def compute_dates(tuan_so, day_of_week):
    week_start = TUAN_01_START + timedelta(weeks=tuan_so - 1)
    ngay_day = week_start + timedelta(days=day_of_week)
    ngay_soan = week_start - timedelta(days=2)
    return ngay_soan.strftime('%d/%m/%Y'), ngay_day.strftime('%d/%m/%Y')

def make_paragraph(text, bold=True):
    """Create a w:p element with a single run, TNR 13pt."""
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), '26')  # 13pt
    rPr.append(sz)
    if bold:
        b = OxmlElement('w:b')
        rPr.append(b)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    r.append(t)
    p.append(r)
    return p

def set_run_font(run):
    """Set font to TNR 13pt on a run."""
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
        rFonts.set(qn(attr), 'Times New Roman')

# ─── FIX THCS (Lớp 6, 7, 8) ──────────────────────────────────────────────
def fix_thcs(filepath, lop_key, tuan_so):
    day_of_week, ten_lop = LOP_SCHEDULE[lop_key]
    ngay_soan, ngay_day = compute_dates(tuan_so, day_of_week)
    
    doc = Document(filepath)
    modified = False
    
    # Fix Table[0] Row[1] Cell[1] & Make borderless
    if doc.tables:
        tbl = doc.tables[0]
        # Make borderless
        tblPr = tbl._tbl.find(qn('w:tblPr'))
        if tblPr is not None:
            old_b = tblPr.find(qn('w:tblBorders'))
            if old_b is not None:
                tblPr.remove(old_b)
            borders_xml = parse_xml(
                r'<w:tblBorders %s>'
                r'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                r'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                r'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                r'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                r'<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                r'<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                r'</w:tblBorders>' % nsdecls('w')
            )
            tblPr.append(borders_xml)

        rows = tbl._tbl.findall(qn('w:tr'))
        if len(rows) >= 2:
            cells_r1 = rows[1].findall(qn('w:tc'))
            if len(cells_r1) >= 2:
                tc = cells_r1[1]
                # Remove all old paragraphs
                for old_p in tc.findall(qn('w:p')):
                    tc.remove(old_p)
                # Create new paragraphs: 2 lines only
                tc.append(make_paragraph(f'Ngày soạn: {ngay_soan}   Ngày dạy: {ngay_day}'))
                tc.append(make_paragraph(f'Lớp: {ten_lop}'))
                modified = True
    
    # Fix paragraph "Môn học: ... Lớp: X ..."
    for p in doc.paragraphs:
        if p.text.strip().startswith('Môn học:') and 'Lớp:' in p.text:
            new_text = re.sub(r'Lớp:\s*[0-9A-Za-z]+', f'Lớp: {ten_lop}', p.text)
            if new_text != p.text:
                p.text = new_text
                # Format font
                for r in p.runs:
                    set_run_font(r)
                    r.bold = True
                modified = True
            break
    
    if modified:
        try:
            doc.save(filepath)
            print(f'  ✅ Fixed: {os.path.basename(filepath)}')
        except PermissionError:
            alt = filepath.replace('.docx', '_fixed.docx')
            doc.save(alt)
            print(f'  ⚠️ Saved as: {os.path.basename(alt)} (original locked)')
    else:
        print(f'  ⏩ No changes: {os.path.basename(filepath)}')

# ─── FIX TH (Tiền TH + Lớp 1-5) ──────────────────────────────────────────
def fix_th_lop5(filepath, tuan_so, lop_key='5'):
    day_of_week, ten_lop = LOP_SCHEDULE[lop_key]
    ngay_soan, ngay_day = compute_dates(tuan_so, day_of_week)
    
    doc = Document(filepath)
    modified = False
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if ('Thứ' in text and 'ngày' in text) or text.startswith('Ngày soạn:') or text.startswith('TUẦN:'):
            p.clear()
            r = p.add_run(f'Ngày soạn: {ngay_soan}   Ngày dạy: {ngay_day}')
            r.italic = True
            set_run_font(r)
            modified = True
            break
    
    if modified:
        try:
            doc.save(filepath)
            print(f'  ✅ Fixed: {os.path.basename(filepath)}')
        except PermissionError:
            alt = filepath.replace('.docx', '_fixed.docx')
            doc.save(alt)
            print(f'  ⚠️ Saved as: {os.path.basename(alt)} (original locked)')
    else:
        print(f'  ⏩ No changes: {os.path.basename(filepath)}')

# ─── MAIN ──────────────────────────────────────────────────────────────────
def main():
    tuan_map = {'Tuần_01': 1, 'Tuần_02': 2, 'Tuần_03': 3}
    
    # THCS (Lớp 6, 7, 8)
    for lop_key in ['6', '7', '8']:
        lop_folder = f'Lớp_{lop_key}'
        print(f'\n🟠 Fixing THCS {lop_folder}...')
        for tuan_folder, tuan_so in tuan_map.items():
            dir_path = os.path.join(KHBD_BASE, lop_folder, tuan_folder)
            if not os.path.isdir(dir_path):
                continue
            for f in os.listdir(dir_path):
                if f.endswith('.docx') and 'KHBD' in f:
                    fix_thcs(os.path.join(dir_path, f), lop_key, tuan_so)
    
    # Tiểu học (Tiền TH + Lớp 1-5)
    th_mapping = {
        'Tiền_tiểu_học': 'TTH',
        'Lớp_1': '1',
        'Lớp_2': '2',
        'Lớp_3': '3',
        'Lớp_4': '4',
        'Lớp_5': '5',
    }
    for lop_folder, lop_key in th_mapping.items():
        print(f'\n🔵 Fixing TH {lop_folder}...')
        for tuan_folder, tuan_so in tuan_map.items():
            dir_path = os.path.join(KHBD_BASE, lop_folder, tuan_folder)
            if not os.path.isdir(dir_path):
                continue
            for f in os.listdir(dir_path):
                if f.endswith('.docx') and 'KHBD' in f:
                    fix_th_lop5(os.path.join(dir_path, f), tuan_so, lop_key)
    
    print('\n✅ Done!')

if __name__ == '__main__':
    main()

