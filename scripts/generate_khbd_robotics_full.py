import os
import re
import sys
import copy
import shutil
from datetime import date, timedelta
import docx
from docx.shared import Inches, Pt, Emu, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.stdout.reconfigure(encoding='utf-8')

# === INDENT EMU CONSTANTS ===
INDENT_0 = 0
INDENT_1 = 180340
INDENT_2 = 360045
INDENT_BULLET = 540000

TPL_KHBD = r'd:\UNIGO\Hệ thống mẫu văn bản\PL4-Khung kế hoạch bài dạy (THCS).docx'
PPCT_PATH = r'd:\UNIGO\Phân phối chương trình\Robotics\Kế hoạch dạy học môn Robotics 2026-2027.docx'
BASE_OUT_DIR = r'd:\UNIGO\KHBD_Robotics'

# --- Helpers ---
def set_table_borders(table, color="000000", sz="4", val="single"):
    """Set table borders via XML, handling missing tblPr and removing old borders."""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    old = tblPr.find(qn('w:tblBorders'))
    if old is not None:
        tblPr.remove(old)
    tblBorders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), val)
        border.set(qn('w:sz'), sz)
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tblBorders.append(border)
    tblPr.append(tblBorders)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def afont(run, font_name="Times New Roman", size_pt=13, bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color

def format_paragraph(p, font_name="Times New Roman", size_pt=13, line_spacing=1.15, space_after=3, bold=False, italic=False):
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)
    for run in p.runs:
        afont(run, font_name=font_name, size_pt=size_pt, bold=bold, italic=italic)

def sanitize_filename(name):
    import unicodedata
    name = unicodedata.normalize('NFD', name).encode('ascii', 'ignore').decode('utf-8')
    name = re.sub(r'[^a-zA-Z0-9_\-]', '_', name)
    name = re.sub(r'_+', '_', name).strip('_')
    return name

def get_kit_name(grade):
    if grade in [1, 2]:
        return "OLLO Kinder"
    elif grade in [3, 4]:
        return "OLLO Initiate"
    else:
        return "OLLO Excel 1"

def clean_body_preserve_sectpr(doc):
    for child in list(doc.element.body):
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag != 'sectPr':
            doc.element.body.remove(child)

def fill_cell(cell, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, size_pt=13, space_after=2, line_spacing=1.15):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(space_after)
    set_cell_margins(cell)
    
    lines = text.split('\n')
    for i, line in enumerate(lines):
        if i > 0:
            p = cell.add_paragraph()
            p.alignment = align
            p.paragraph_format.line_spacing = line_spacing
            p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(line)
        afont(r, size_pt=size_pt, bold=bold, italic=italic)

# --- Generator for Primary (Lớp 1-5) ---
def build_khbd_primary(grade, title, tiet_ppct, yccd):
    doc = docx.Document(TPL_KHBD)
    clean_body_preserve_sectpr(doc)
    kit = get_kit_name(grade)
    
    # Margins: L3cm, R2cm, T2cm, B2cm
    for sec in doc.sections:
        sec.top_margin = Inches(0.787)
        sec.bottom_margin = Inches(0.787)
        sec.left_margin = Inches(1.181)
        sec.right_margin = Inches(0.787)
        
    tuan_so = (tiet_ppct - 1) // 2 + 1 if tiet_ppct > 1 else 1

    # Header paragraphs
    p0 = doc.add_paragraph()
    p0.paragraph_format.line_spacing = 1.15
    p0.paragraph_format.space_after = Pt(2)
    r0 = p0.add_run(f"TUẦN: {tuan_so:02d}\t\t\t\tNgày soạn: 01/09/2026\n\t\t\t\t\tNgày dạy: 05/09/2026")
    afont(r0, size_pt=12, italic=True)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.line_spacing = 1.2
    p_title.paragraph_format.space_after = Pt(4)
    r1 = p_title.add_run(f"KẾ HOẠCH DẠY HỌC MÔN ROBOTICS - LỚP {grade}\n")
    afont(r1, size_pt=14, bold=True)
    r2 = p_title.add_run(f"BỘ THIẾT BỊ: {kit.upper()}\n")
    afont(r2, size_pt=13, bold=True)
    r3 = p_title.add_run(f"BÀI: {title.upper()}\n")
    afont(r3, size_pt=14, bold=True)
    r4 = p_title.add_run(f"(Thời lượng: 1 tiết | Tiết PPCT: {tiet_ppct})")
    afont(r4, size_pt=12, italic=True)

    # I. YÊU CẦU CẦN ĐẠT
    p = doc.add_paragraph()
    r = p.add_run("I. YÊU CẦU CẦN ĐẠT:")
    afont(r, size_pt=13, bold=True)

    p = doc.add_paragraph()
    r = p.add_run("- Sau tiết học, học sinh sẽ:")
    afont(r, size_pt=13, italic=True)

    # 1. Phẩm chất FIRST for Primary
    p = doc.add_paragraph()
    r = p.add_run("1. Phát triển phẩm chất:")
    afont(r, size_pt=13, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.33
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Inches(0.25)
    r = p.add_run(
        "- Chăm chỉ: Hăng hái tham gia hoạt động tìm hiểu linh kiện, kiên trì thực hiện thao tác lắp ráp mô hình robot theo hướng dẫn (Đạt được thông qua Hoạt động 2, Hoạt động 3).\n"
        "- Trách nhiệm: Giữ gìn cẩn thận các linh kiện trong bộ Kit Robotics, thu dọn phân loại gọn gàng thiết bị sau giờ học (Đạt được thông qua Hoạt động 4).\n"
        "- Trung thực: Tự giác làm việc nhóm, tôn trọng kết quả thử nghiệm mô hình robot của bản thân và nhóm bạn (Đạt được thông qua Hoạt động 3)."
    )
    afont(r, size_pt=13)

    # 2. Năng lực SECOND for Primary
    p = doc.add_paragraph()
    r = p.add_run("2. Phát triển năng lực:")
    afont(r, size_pt=13, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.33
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Emu(INDENT_2)
    r = p.add_run(
        f"2.1. Năng lực đặc thù (Robotics):\n"
    )
    afont(r, size_pt=13, bold=True)
    r2 = p.add_run(
        f"- NLa (Sử dụng và quản lí các phương tiện ICT): Nhận biết tên gọi, hình dạng, chức năng các chi tiết khung, chốt nối, động cơ, cảm biến bộ kit {kit}. (Đạt được thông qua Hoạt động 2)\n"
        f"- NLd (Ứng dụng ICT trong học và tự học): Thực hiện lắp ráp đúng quy trình từng bước mô hình {title}, vận hành chạy thử và điều chỉnh. (Đạt được thông qua Hoạt động 3)"
    )
    afont(r2, size_pt=13)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.33
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Emu(INDENT_2)
    r = p.add_run(
        f"2.2. Năng lực số (Thông tư 02/2025 – CV 3456):\n"
    )
    afont(r, size_pt=13, bold=True)
    r2 = p.add_run(
        f"- Miền V. Giải quyết vấn đề (thành tố 5.3. Sử dụng sáng tạo công nghệ – Bậc 1): Nhận biết linh kiện điện tử thông minh, thao tác an toàn với mạch điều khiển, dây dẫn và pin. (Đạt được thông qua Hoạt động 2, Hoạt động 3)"
    )
    afont(r2, size_pt=13)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.33
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.first_line_indent = Emu(INDENT_2)
    r = p.add_run(
        f"2.3. Năng lực chung:\n"
    )
    afont(r, size_pt=13, bold=True)
    r2 = p.add_run(
        f"- Tự chủ và tự học: Quan sát sơ đồ 2D/3D hướng dẫn lắp ráp, chủ động chuẩn bị đúng linh kiện. (Đạt được thông qua Hoạt động 2)\n"
        f"- Giao tiếp và hợp tác: Phân công nhiệm vụ nhóm ăn ý khi thực hành lắp ráp và thử nghiệm sản phẩm. (Đạt được thông qua Hoạt động 3)\n"
        f"- Giải quyết vấn đề và sáng tạo: Phát hiện và xử lý lỗi sai lắp ráp, đề xuất ý tưởng cải tiến mô hình. (Đạt được thông qua Hoạt động 3, Hoạt động 4)"
    )
    afont(r2, size_pt=13)

    # II. ĐỒ DÙNG DẠY HỌC
    p = doc.add_paragraph()
    r = p.add_run("II. ĐỒ DÙNG DẠY HỌC :")
    afont(r, size_pt=13, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(
        f"1. Giáo viên: Bộ Kit Robotics {kit} mẫu, máy tính GV, máy chiếu, bài trình chiếu slide hướng dẫn lắp ráp từng bước mô hình {title}, phiếu hướng dẫn thực hành.\n"
        f"2. Học sinh: Bộ Kit Robotics {kit} theo nhóm/cá nhân, dụng cụ tháo chốt nhựa, vở ghi bài."
    )
    afont(r, size_pt=13)

    # III. PHƯƠNG PHÁP, KĨ THUẬT DẠY HỌC
    p = doc.add_paragraph()
    r = p.add_run("III. PHƯƠNG PHÁP, KĨ THUẬT DẠY HỌC")
    afont(r, size_pt=13, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(
        "- Phương pháp dạy học: Trực quan mô hình, hướng dẫn thực hành step-by-step, làm việc nhóm, giải quyết vấn đề (STEM/STEAM).\n"
        "- Kĩ thuật dạy học: Think-Pair-Share, động não, giao nhiệm vụ phân tầng, trình bày 1 phút."
    )
    afont(r, size_pt=13)

    # IV. CÁC HOẠT ĐỘNG DẠY - HỌC CHỦ YẾU (BẢNG 2 CỘT)
    p = doc.add_paragraph()
    r = p.add_run("IV. CÁC HOẠT ĐỘNG DẠY - HỌC CHỦ YẾU:")
    afont(r, size_pt=13, bold=True)

    table = doc.add_table(rows=1, cols=2)
    set_table_borders(table)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Hdr row
    hdr = table.rows[0].cells
    fill_cell(hdr[0], "Hoạt động của GV", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    fill_cell(hdr[1], "Hoạt động của HS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    activities = [
        ("1. Hoạt động MỞ ĐẦU (5 phút)\n*Mục tiêu: Kích hoạt hứng thú, tạo tình huống kết nối vào bài học " + title,
         f"a) Chuyển giao: GV chiếu video/hình ảnh liên quan đến {title}. Đặt câu hỏi gợi mở tình huống thực tế.\nb) Thực hiện: Quant sát và lắng nghe câu hỏi.\nc) Báo cáo: Mời 2-3 HS phát biểu.\nd) Kết luận: Chốt kiến thức, dẫn dắt vào bài mới.",
         f"a) Tiếp nhận: Quan sát màn chiếu và suy nghĩ.\nb) Thực hiện: Thảo luận nhanh với bạn bên cạnh.\nc) Báo cáo: Hăng hái giơ tay trả lời.\nd) Kết luận: Chuẩn bị bộ Kit Robotics."),
         
        ("2. HOẠT ĐỘNG HÌNH THÀNH KIẾN THỨC MỚI (15 phút)\n*Mục tiêu: Tìm hiểu linh kiện và nguyên lý cơ khí của " + title,
         f"a) Chuyển giao: Giới thiệu linh kiện cần dùng và phân tích sơ đồ lắp ráp {title}.\nb) Thực hiện: Hướng dẫn HS quan sát cách ghép chốt và khớp nối.\nc) Báo cáo: Kiểm tra khay linh kiện từng nhóm.\nd) Kết luận: Chuẩn hóa tên linh kiện và quy trình các bước.",
         f"a) Tiếp nhận: Theo dõi slide hướng dẫn.\nb) Thực hiện: Nhặt đúng và đủ số lượng linh kiện ra khay chứa.\nc) Báo cáo: Đại diện nhóm giơ khay linh kiện kiểm tra.\nd) Kết luận: Ghi nhớ thứ tự các bước ghép."),
         
        ("3. HĐ LUYỆN TẬP-THỰC HÀNH (10 phút)\n*Mục tiêu: Tiến hành lắp ráp và vận hành chạy thử robot " + title,
         f"a) Chuyển giao: Giao nhiệm vụ cho các nhóm thực hành lắp ráp robot {title}.\nb) Thực hiện: Quan sát, hỗ trợ nhóm gặp khó khăn, nhắc an toàn.\nc) Báo cáo: Cho các nhóm bật nguồn vận hành chạy thử.\nd) Kết luận: Đánh giá sản phẩm hoàn thiện, tuyên dương nhóm làm tốt.",
         f"a) Tiếp nhận: Phân công nhiệm vụ trong nhóm (lấy chốt, giữ khung, lắp chi tiết).\nb) Thực hiện: Tiến hành lắp ráp cẩn thận theo sơ đồ.\nc) Báo cáo: Đặt robot thử nghiệm, quan sát chuyển động.\nd) Kết luận: Điều chỉnh lại chốt/khớp nếu robot kẹt."),
         
        ("4. HOẠT ĐỘNG VẬN DỤNG, TRẢI NGHIỆM (5 phút)\n*Mục tiêu: Cải tiến sản phẩm và thu dọn phân loại linh kiện",
         f"a) Chuyển giao: Đặt câu hỏi mở rộng cải tiến tính năng cho robot {title}.\nb) Thực hiện: Hướng dẫn tháo dỡ và kiểm kê linh kiện.\nc) Báo cáo: Lắng nghe ý tưởng sáng tạo.\nd) Kết luận: Nhận xét tiết học, dặn dò bài sau.",
         f"a) Tiếp nhận: Suy nghĩ ý tưởng nâng cấp robot.\nb) Thực hiện: Tháo rời nhẹ nhàng, xếp linh kiện vào đúng ngăn.\nc) Báo cáo: Chia sẻ ý tưởng cải tiến.\nd) Kết luận: Đóng hộp Kit, cất gọn gàng.")
    ]

    for title_hdr, gv_txt, hs_txt in activities:
        # Add merged header row
        row_hdr = table.add_row()
        cell0 = row_hdr.cells[0]
        cell1 = row_hdr.cells[1]
        tcPr = cell0._tc.get_or_add_tcPr()
        gs = OxmlElement('w:gridSpan')
        gs.set(qn('w:val'), '2')
        tcPr.append(gs)
        row_hdr._tr.remove(cell1._tc)
        
        fill_cell(cell0, title_hdr, bold=True, size_pt=12, line_spacing=1.15)
        
        # Add content row
        row_cnt = table.add_row()
        fill_cell(row_cnt.cells[0], gv_txt, size_pt=12, line_spacing=1.15)
        fill_cell(row_cnt.cells[1], hs_txt, size_pt=12, line_spacing=1.15)

    # Copy Section V (ĐIỀU CHỈNH - BỔ SUNG) from template P[33:36]
    p_adj1 = doc.add_paragraph()
    p_adj1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_adj1.add_run("V. ĐIỀU CHỈNH - BỔ SUNG SAU TIẾT DẠY :")
    afont(r, size_pt=13, bold=True)

    p_adj2 = doc.add_paragraph()
    p_adj2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p_adj2.add_run("(GV ghi những nội dung mà mình đã bổ sung ngoài KHBD đã lên...)")
    afont(r, size_pt=13, bold=True, italic=True)

    p_dots = doc.add_paragraph()
    r = p_dots.add_run("...........................................................................................................................................................................")
    afont(r, size_pt=13)

    return doc

# --- Generator for THCS (Lớp 6-8) ---
def build_khbd_thcs(grade, title, tiet_ppct, yccd):
    doc = docx.Document(TPL_KHBD)
    kit = get_kit_name(grade)
    
    # Preserve Table 0 (school/teacher info)
    t0 = doc.tables[0]
    t0.rows[0].cells[0].paragraphs[0].text = "Trường: TH & THCS UNIGO\nTổ: Robotics"
    t0.rows[0].cells[1].paragraphs[0].text = "Họ tên giáo viên: Đậu Đình Nguyên"
    t0.rows[1].cells[0].paragraphs[0].text = "Bộ môn: Robotics"
    t0.rows[1].cells[1].paragraphs[0].text = f"Ngày soạn: 01/09/2026   Ngày dạy: 05/09/2026\nLớp: {grade}"

    for row in t0.rows:
        for cell in row.cells:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    afont(r, size_pt=12, bold=True)

    # Clean body content between Table 0 and Table 2
    # In THCS template, Table 0 is first child, Table 2 is signature table near end.
    # We remove middle paragraphs/tables and rebuild section I, II, III, then keep Table 2.
    body = doc.element.body
    children = list(body)
    
    # Find index of table 0 and table 2
    t0_elem = t0._element
    t2_elem = doc.tables[2]._element
    
    removing = False
    for child in children:
        if child == t0_elem:
            removing = True
            continue
        if child == t2_elem:
            removing = False
            continue
        if removing:
            body.remove(child)

    # Set THCS Margins: L2.54cm, R1.27cm, T1.27cm, B1.27cm
    for sec in doc.sections:
        sec.top_margin = Inches(0.5)
        sec.bottom_margin = Inches(0.5)
        sec.left_margin = Inches(1.0)
        sec.right_margin = Inches(0.5)

    # Insert content before Table 2
    # We will build elements and insert before t2_elem
    def add_p_before_t2(text="", bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=3, size_pt=13):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            r = p.add_run(text)
            afont(r, size_pt=size_pt, bold=bold, italic=italic)
        body.remove(p._p)
        t2_elem.addprevious(p._p)
        return p

    def add_table_before_t2(table):
        body.remove(table._element)
        t2_elem.addprevious(table._element)

    # Title Block
    add_p_before_t2(f"TÊN BÀI DẠY: {title.upper()}", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=14)
    add_p_before_t2(f"MÔN HỌC: ROBOTICS {grade} (BỘ KIT {kit.upper()})", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=13)
    add_p_before_t2(f"Thời lượng: 1 tiết | Tiết theo PPCT: {tiet_ppct}", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=12)

    # I. Mục tiêu
    add_p_before_t2("I. Mục tiêu", bold=True)
    
    # 1. Kiến thức FIRST for THCS (using Nouns / Noun phrases)
    p_kt = add_p_before_t2()
    p_kt.paragraph_format.left_indent = Inches(0.2)
    r = p_kt.add_run(
        f"1. Kiến thức:\n"
        f"   - Sự hiểu biết về trọng tâm bài học: {yccd}\n"
        f"   - Khả năng phân tích cấu tạo cơ khí, nguyên lý truyền động bánh răng/động cơ servo Dynamixel và thuật toán điều khiển thông minh trong mô hình {title}."
    )
    afont(r, size_pt=13)

    p_nl = add_p_before_t2()
    p_nl.paragraph_format.first_line_indent = Emu(INDENT_1)
    r = p_nl.add_run("2. Năng lực:")
    afont(r, size_pt=13, bold=True)

    p_nldt = add_p_before_t2()
    p_nldt.paragraph_format.first_line_indent = Emu(INDENT_2)
    r = p_nldt.add_run("2.1. Năng lực đặc thù (Robotics):")
    afont(r, size_pt=13, bold=True)

    p_nldt_c = add_p_before_t2()
    p_nldt_c.paragraph_format.left_indent = Emu(INDENT_BULLET)
    r = p_nldt_c.add_run(
        f"- NLa (Sử dụng và quản lí các phương tiện ICT): Năng lực thiết kế mô phỏng 3D, kĩ năng thao tác lắp ráp chuẩn xác các khớp nối. (Đạt được thông qua Hoạt động 2, Hoạt động 3)\n"
        f"- NLd (Ứng dụng ICT trong học và tự học): Nạp code lập trình và hiệu chỉnh robot. (Đạt được thông qua Hoạt động 3)"
    )
    afont(r, size_pt=13)

    p_nls = add_p_before_t2()
    p_nls.paragraph_format.first_line_indent = Emu(INDENT_2)
    r = p_nls.add_run("2.2. Năng lực số (Thông tư 02/2025 – CV 3456):")
    afont(r, size_pt=13, bold=True)

    p_nls_c = add_p_before_t2()
    p_nls_c.paragraph_format.left_indent = Emu(INDENT_BULLET)
    r = p_nls_c.add_run(
        f"- Miền V. Giải quyết vấn đề (thành tố 5.3. Sử dụng sáng tạo công nghệ – Bậc 2): Kĩ năng vận hành thiết bị điều khiển thông minh, ứng dụng phần mềm nạp lập trình tự động. (Đạt được thông qua Hoạt động 2, Hoạt động 3)\n"
        f"- Miền IV. An toàn (thành tố 4.1. Bảo vệ thiết bị – Bậc 2): Tuân thủ quy tắc an toàn thiết bị số và nguồn điện. (Đạt được thông qua Hoạt động 3)"
    )
    afont(r, size_pt=13)

    p_nlc = add_p_before_t2()
    p_nlc.paragraph_format.first_line_indent = Emu(INDENT_2)
    r = p_nlc.add_run("2.3. Năng lực chung:")
    afont(r, size_pt=13, bold=True)

    p_nlc_c = add_p_before_t2()
    p_nlc_c.paragraph_format.left_indent = Emu(INDENT_BULLET)
    r = p_nlc_c.add_run(
        f"- Tự chủ và tự học: Chủ động tự học, tự nghiên cứu tài liệu sơ đồ kỹ thuật. (Đạt được thông qua Hoạt động 2)\n"
        f"- Giao tiếp và hợp tác: Giao tiếp và làm việc nhóm hiệu quả. (Đạt được thông qua Hoạt động 3)\n"
        f"- Giải quyết vấn đề và sáng tạo: Giải quyết vấn đề sáng tạo kỹ thuật. (Đạt được thông qua Hoạt động 3, Hoạt động 4)"
    )
    afont(r, size_pt=13)

    # 3. Phẩm chất THIRD for THCS
    p_pc = add_p_before_t2()
    p_pc.paragraph_format.first_line_indent = Emu(INDENT_1)
    r = p_pc.add_run("3. Phẩm chất:")
    afont(r, size_pt=13, bold=True)

    p_pc_c = add_p_before_t2()
    p_pc_c.paragraph_format.left_indent = Emu(INDENT_BULLET)
    r = p_pc_c.add_run(
        f"- Trách nhiệm: Quản lý, bảo vệ thiết bị công nghệ. (Đạt được thông qua Hoạt động 3, Hoạt động 4)\n"
        f"- Trung thực: Trung thực trong báo cáo kết quả thử nghiệm mô hình. (Đạt được thông qua Hoạt động 3)\n"
        f"- Chăm chỉ: Tác phong công nghiệp và tư duy khoa học. (Đạt được thông qua Hoạt động 2, Hoạt động 3)"
    )
    afont(r, size_pt=13)

    # II. Thiết bị dạy học và học liệu
    add_p_before_t2("II. Thiết bị dạy học và học liệu:", bold=True)
    p_tb = add_p_before_t2()
    p_tb.paragraph_format.left_indent = Inches(0.2)
    r = p_tb.add_run(
        f"1. Thiết bị: Bộ Kit Robotics {kit}, máy tính giáo viên và máy tính nhóm học sinh, phần mềm nạp lập trình, máy chiếu bài giảng 3D.\n"
        f"2. Học liệu: Phiếu học tập thực hành, sơ đồ nguyên lý cơ khí và thuật toán điều khiển {title}."
    )
    afont(r, size_pt=13)

    # III. Tiến trình dạy học (Option B: 3-column table per activity)
    add_p_before_t2("III. Tiến trình dạy học", bold=True)

    activities_thcs = [
        ("1. Hoạt động 1. Khởi động (Xác định vấn đề/nhiệm vụ học tập/Mở đầu) (7 phút)",
         f"Mục tiêu kích hoạt tư duy kỹ thuật và kết nối tình huống vào bài học {title}.",
         f"GV chiếu hình ảnh/video thực tế về {title} và đặt câu hỏi phân tích cơ chế.",
         f"Câu trả lời của HS xác định được vấn đề kỹ thuật cần giải quyết.",
         f"Chuyển giao nhiệm vụ khởi động.",
         [
             (f"GV trình chiếu clip/hình ảnh thực tế liên quan đến {title}, yêu cầu HS quan sát.", "HS tập trung quan sát màn chiếu, lắng nghe câu hỏi dẫn dắt."),
             ("GV theo dõi, gợi mở các câu hỏi liên quan đến nguyên lý chuyển động.", "HS thảo luận nhanh theo nhóm bàn, phân tích các hiện tượng kỹ thuật."),
             ("GV gọi 2 đại diện HS trả lời câu hỏi khởi động.", "Đại diện HS nêu giả thuyết và nguyên lý vận hành ban đầu."),
             ("GV nhận xét, chuẩn hóa kiến thức và dẫn dắt vào bài mới.", "HS lắng nghe, ghi chép và mở bộ Kit Robotics chuẩn bị thực hành.")
         ]),
         
        ("2. Hoạt động 2. Hình thành kiến thức mới/giải quyết vấn đề (18 phút)",
         f"Nghiên cứu sơ đồ thiết kế và thuật toán cho mô hình {title}.",
         f"GV yêu cầu HS tìm hiểu các linh kiện khung, chốt, động cơ và cơ cấu truyền động.",
         f"Sơ đồ khối nguyên lý cơ khí và danh mục linh kiện được chọn chính xác.",
         f"Chuyển giao nhiệm vụ khám phá sơ đồ.",
         [
             (f"GV phát phiếu hướng dẫn và trình chiếu sơ đồ 2D/3D của {title}.", "HS tiếp nhận phiếu hướng dẫn, quan sát các góc ghép linh kiện."),
             ("GV hướng dẫn HS cách nhận biết chiều lắp chốt và cổng cắm động cơ.", "HS đối chiếu danh mục linh kiện, nhặt đúng số lượng chốt và khung nối."),
             ("GV kiểm tra khay linh kiện của từng nhóm tại bàn.", "Đại diện nhóm giơ khay linh kiện đã chuẩn bị để GV kiểm tra."),
             ("GV chốt quy trình 4 bước lắp ráp mô hình robot.", "HS ghi nhớ thứ tự các bước và chuẩn bị công cụ lắp ráp.")
         ]),

        ("3. Hoạt động 3. Luyện tập (12 phút)",
         f"Thực hành lắp ráp, lập trình và chạy thử mô hình robot {title}.",
         f"GV yêu cầu các nhóm tiến hành lắp ráp phần cứng và thử nghiệm vận hành.",
         f"Mô hình robot {title} hoàn thiện, hoạt động chính xác theo yêu cầu.",
         f"Chuyển giao nhiệm vụ thực hành lắp ráp.",
         [
             (f"GV giao nhiệm vụ cho từng nhóm tiến hành lắp ráp mô hình {title}.", "HS phân công nhiệm vụ (1 bạn xem sơ đồ, 1 bạn lắp khung, 1 bạn chọn chốt)."),
             ("GV quan sát trực tiếp tại các bàn, hỗ trợ nhóm gặp vướng mắc kỹ thuật.", "HS tiến hành lắp ráp cẩn thận từng chi tiết, kiểm tra độ mượt của khớp nối."),
             ("GV yêu cầu các nhóm cấp nguồn/nạp code và cho robot vận hành thử.", "HS đặt robot lên bàn thử nghiệm, bật công tắc quan sát hoạt động."),
             ("GV đánh giá sản phẩm thực hành, chấm điểm tiêu chí kỹ thuật.", "HS tinh chỉnh lại chốt nối nếu robot bị kẹt hoặc di chuyển lệch.")
         ]),

        ("4. Hoạt động 4. Vận dụng (8 phút)",
         f"Vận dụng sáng tạo, nâng cấp tính năng robot và dọn dẹp vệ sinh.",
         f"GV đặt yêu cầu cải tiến tối ưu thuật toán hoặc thêm chi tiết trang trí.",
         f"Báo cáo ý tưởng cải tiến mô hình và bộ Kit được sắp xếp ngăn nắp.",
         f"Chuyển giao nhiệm vụ vận dụng và dọn dẹp.",
         [
             (f"GV đặt câu hỏi: 'Em có thể cải tiến cấu trúc nào để {title} hoạt động tốt hơn?'", "HS tiếp nhận câu hỏi suy nghĩ ý tưởng sáng tạo mở rộng."),
             ("GV hướng dẫn quy trình tháo dỡ robot và kiểm kê linh kiện.", "HS nhẹ nhàng tháo chốt, phân loại linh kiện về đúng vị trí trong hộp Kit."),
             ("GV gọi 1-2 HS trình bày ý tưởng nâng cấp robot.", "HS phát biểu ý tưởng sáng tạo cá nhân trước lớp."),
             ("GV dặn dò nhiệm vụ chuẩn bị cho bài học tiếp theo.", "HS thu dọn bàn học, đẩy ghế ngăn nắp và nộp lại hộp Kit Robotics.")
         ])
    ]

    buoc_labels = [
        'Bước 1:\nChuyển giao\nnhiệm vụ',
        'Bước 2:\nHọc sinh tiếp nhận\nnhiệm vụ',
        'Bước 3:\nBáo cáo kết quả\nhoạt động',
        'Bước 4:\nĐánh giá kết quả\nthực hiện'
    ]

    for idx, (act_title, act_mt, act_nd, act_sp, act_tc, buocs) in enumerate(activities_thcs, start=1):
        add_p_before_t2(act_title, bold=True, size_pt=13)
        
        p_sub = add_p_before_t2()
        p_sub.paragraph_format.left_indent = Inches(0.2)
        r = p_sub.add_run(
            f"a) Mục tiêu: {act_mt}\n"
            f"b) Nội dung: {act_nd}\n"
            f"c) Sản phẩm: {act_sp}\n"
            f"d) Tổ chức thực hiện:"
        )
        afont(r, size_pt=13, italic=True)

        # 3-column table for activity
        table = doc.add_table(rows=1, cols=3)
        set_table_borders(table)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = table.rows[0].cells
        fill_cell(hdr[0], "Bước", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(hdr[1], "Hoạt động của GV", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
        fill_cell(hdr[2], "Hoạt động của HS", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

        b_labels = list(buoc_labels)
        if idx == 4:
            b_labels[3] = 'Bước 4:\nNhắc nhở nhiệm vụ\nvề nhà'

        for i, (gv_t, hs_t) in enumerate(buocs):
            row = table.add_row().cells
            fill_cell(row[0], b_labels[i], bold=True, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size_pt=11)
            fill_cell(row[1], gv_t, size_pt=12, line_spacing=1.15)
            fill_cell(row[2], hs_t, size_pt=12, line_spacing=1.15)

        add_table_before_t2(table)
        add_p_before_t2() # spacing

    # Format Table 2 (signatures)
    t2 = doc.tables[2]
    for row in t2.rows:
        for cell in row.cells:
            set_cell_margins(cell)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs:
                    afont(r, size_pt=12, bold=True)

    return doc

# --- Main execution ---
def main():
    print("==================================================")
    print(" BẮT ĐẦU TẠO TOÀN BỘ KHBD ROBOTICS (THEO TUẦN)")
    print("==================================================")

    # 1. Xóa sạch thư mục KHBD_Robotics cũ
    if os.path.exists(BASE_OUT_DIR):
        shutil.rmtree(BASE_OUT_DIR)
        print("  [+] Đã xóa sạch thư mục KHBD_Robotics cũ.")

    os.makedirs(BASE_OUT_DIR, exist_ok=True)

    doc_ppct = docx.Document(PPCT_PATH)
    tables = doc_ppct.tables
    created_count = 0

    for grade, t_idx in enumerate(range(4, 12), start=1):
        t = tables[t_idx]
        print(f"\n---> Đang xử lý LỚP {grade} (Table {t_idx}, {len(t.rows)-1} mục)...")

        lesson_counter = 0
        for row_idx, row in enumerate(t.rows[1:], start=1):
            cells = [c.text.strip().replace('\n', ' ') for c in row.cells]
            stt = cells[0]
            title = cells[1]
            so_tiet = cells[2]
            tiet_ppct_str = cells[3] if len(cells) > 3 else str(row_idx)
            yccd = cells[4] if len(cells) > 4 else title

            if not title:
                continue

            try:
                tiet_ppct = int(tiet_ppct_str)
            except ValueError:
                tiet_ppct = row_idx

            safe_title = sanitize_filename(title)

            # Tính Tuần theo đúng logic Lịch báo giảng:
            if grade >= 5:
                # Lớp 5-8: Tuần 1 = Tiết 0; Tuần lẻ 3, 5, 7... = Robotics (2 tiết/tuần)
                if tiet_ppct == 0:
                    tuan_so = 1
                else:
                    tuan_so = 2 * ((tiet_ppct - 1) // 2) + 3
            else:
                # Lớp 1-4: Tuần 1 = Tiết 0; Tuần 2+ = ppct + 1
                tuan_so = tiet_ppct + 1 if tiet_ppct >= 1 else 1

            tuan_folder = f"Tuần_{tuan_so:02d}"

            if tiet_ppct == 0 or 'Tiết 0' in title or 'Định hướng' in title:
                filename = f"KHBD_Robotics_Lớp_{grade}_Tiet00_{safe_title}.docx"
            else:
                lesson_counter += 1
                filename = f"KHBD_Robotics_Lớp_{grade}_Tiet{tiet_ppct:02d}_{safe_title}.docx"

            out_folder = os.path.join(BASE_OUT_DIR, f"Lớp_{grade}", tuan_folder)
            os.makedirs(out_folder, exist_ok=True)
            out_filepath = os.path.join(out_folder, filename)

            # Use THCS template for ALL grades
            doc = build_khbd_thcs(grade, title, tiet_ppct, yccd)

            try:
                doc.save(out_filepath)
                created_count += 1
                print(f"  [+] Đã tạo: Lớp {grade} -> {tuan_folder} -> {filename}")
            except Exception as e:
                print(f"  [!] Lỗi khi lưu {out_filepath}: {e}")

    print(f"\n==================================================")
    print(f" HOÀN THÀNH TẠO {created_count} FILE KHBD ROBOTICS CHUẨN (THEO TUẦN)!")
    print(f"==================================================")

if __name__ == '__main__':
    main()
