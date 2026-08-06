import os
import sys
import docx
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import openpyxl
from openpyxl.styles import Font, Alignment, Border, Side, PatternFill

sys.stdout.reconfigure(encoding='utf-8')

OUT_DIR_WORD = r'D:\UNIGO\Thời khóa biểu giáo viên'
OUT_DIR_MAU = r'D:\UNIGO\Hệ thống mẫu văn bản\Nguyên đã làm\Thời khóa biểu'
TKB_EXCEL_SOURCE = r'D:\UNIGO\TKB toàn trường CHECK - chuẩn.xlsx'

os.makedirs(OUT_DIR_WORD, exist_ok=True)
os.makedirs(OUT_DIR_MAU, exist_ok=True)

def load_teacher_data_from_excel(xlsx_path):
    wb = openpyxl.load_workbook(xlsx_path, data_only=True)
    ws = wb['TKB_LOP_SC']

    classes = {}
    curr_lop = None
    for c in range(3, ws.max_column + 1):
        val = ws.cell(4, c).value
        if val:
            curr_lop = str(val).strip()
        buoi = str(ws.cell(5, c).value or '').strip()
        classes[c] = (curr_lop, buoi)

    teacher_map = {
        'Nguyên': {'full_name': 'Đậu Đình Nguyên', 'subject': 'Tin học & Robotics'},
        'Hà': {'full_name': 'Nguyễn Thị Hà', 'subject': 'Ngữ văn & LS-ĐL-GDCD'},
        'Ánh NV': {'full_name': 'Nguyễn Thị Ánh (Ánh NV)', 'subject': 'Ngữ văn & GDĐP-Đọc sách'},
        'Tân': {'full_name': 'Cô Tân', 'subject': 'Toán'}
    }

    data = {k: {'full_name': teacher_map[k]['full_name'], 'subject': teacher_map[k]['subject'], 'slots': []} for k in teacher_map}

    curr_thu = None
    for r in range(6, ws.max_row + 1):
        t_val = ws.cell(r, 1).value
        if t_val:
            curr_thu = str(int(t_val))
        tiet_val = ws.cell(r, 2).value
        if tiet_val is None:
            continue
        tiet = int(tiet_val)
        for c in range(3, ws.max_column + 1):
            cell_val = ws.cell(r, c).value
            if not cell_val:
                continue
            cell_str = str(cell_val).strip()
            
            for t_key in data:
                matched = False
                if t_key == 'Nguyên' and 'Nguyên' in cell_str:
                    matched = True
                elif t_key == 'Hà' and 'Hà' in cell_str and 'Ngọc Ánh' not in cell_str:
                    matched = True
                elif t_key == 'Ánh NV' and ('Ánh NV' in cell_str or ('Ánh' in cell_str and 'Ngọc' not in cell_str)):
                    matched = True
                elif t_key == 'Tân' and 'Tân' in cell_str:
                    matched = True
                
                if matched:
                    lop, buoi = classes[c]
                    mon = cell_str.split('-')[0].strip() if '-' in cell_str else cell_str
                    data[t_key]['slots'].append((curr_thu, buoi, tiet, mon, lop))
    return data

def set_table_borders(table):
    """Sets explicit full black grid borders."""
    table.style = 'Table Grid'
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        for child in list(tblPr[0]):
            if child.tag.endswith('tblBorders'):
                tblPr[0].remove(child)
        borders_xml = parse_xml(
            r'<w:tblBorders %s>'
            r'  <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>'
            r'</w:tblBorders>' % nsdecls('w')
        )
        tblPr[0].append(borders_xml)

def set_cell_background(cell, fill_hex):
    """Sets background color for a docx cell."""
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def create_docx_tkb(t_key, data, save_paths):
    doc = docx.Document()
    
    # Title Header
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p_header.add_run("TRƯỜNG TIỂU HỌC VÀ THCS UNIGO\n")
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(13)
    r1.bold = True
    
    r2 = p_header.add_run("THỜI KHOÁ BIỂU CÁ NHÂN GIÁO VIÊN\n")
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(15)
    r2.bold = True
    
    r3 = p_header.add_run("Năm học 2026 - 2027 (Theo TKB toàn trường CHECK - chuẩn)\n")
    r3.font.name = 'Times New Roman'
    r3.font.size = Pt(12)
    r3.italic = True
    
    # Info paragraph
    p_info = doc.add_paragraph()
    r_info = p_info.add_run(
        f"Họ và tên giáo viên: {data['full_name']}\n"
        f"Môn học đảm nhận: {data['subject']}\n"
        f"Tổng số tiết/tuần: {len(data['slots'])} tiết"
    )
    r_info.font.name = 'Times New Roman'
    r_info.font.size = Pt(13)
    r_info.bold = True
    
    # Create Timetable Grid Table (13 rows x 6 cols)
    table = doc.add_table(rows=13, cols=6)
    set_table_borders(table)
    
    headers = ["Buổi / Tiết", "Thứ Hai", "Thứ Ba", "Thứ Tư", "Thứ Năm", "Thứ Sáu"]
    
    # Row 0: Column headers
    for c_idx in range(6):
        cell = table.rows[0].cells[c_idx]
        cell.text = headers[c_idx]
        set_cell_background(cell, "D9E1F2")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)
            
    # Session Headers
    # Row 1: BUỔI SÁNG
    cell_sang = table.rows[1].cells[0]
    for c_idx in range(1, 6):
        cell_sang.merge(table.rows[1].cells[c_idx])
    cell_sang.text = "BUỔI SÁNG"
    set_cell_background(cell_sang, "F2F2F2")
    p = cell_sang.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)

    # Row 7: BUỔI CHIỀU
    cell_chieu = table.rows[7].cells[0]
    for c_idx in range(1, 6):
        cell_chieu.merge(table.rows[7].cells[c_idx])
    cell_chieu.text = "BUỔI CHIỀU"
    set_cell_background(cell_chieu, "F2F2F2")
    p = cell_chieu.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs:
        r.bold = True
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)

    # Populate period names in Col 0
    for p_num in range(1, 6):
        # Sáng: row p_num + 1 (rows 2..6)
        r_sang = table.rows[p_num + 1]
        r_sang.cells[0].text = f"Tiết {p_num}"
        p = r_sang.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.bold = True
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(13)

        # Chiều: row p_num + 7 (rows 8..12)
        r_chieu = table.rows[p_num + 7]
        r_chieu.cells[0].text = f"Tiết {p_num}"
        p = r_chieu.cells[0].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].font.bold = True
        p.runs[0].font.name = 'Times New Roman'
        p.runs[0].font.size = Pt(13)

    # Map slot items into grid
    day_to_col = {'2': 1, '3': 2, '4': 3, '5': 4, '6': 5}

    for slot in data['slots']:
        day, sess, p_num, subject, cls_name = slot
        col_i = day_to_col[day]
        row_i = (p_num + 1) if sess == 'Sáng' else (p_num + 7)
        
        cell = table.rows[row_i].cells[col_i]
        cell.text = f"{subject}\n({cls_name})"
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)
            if "Robotics" in subject:
                r.font.bold = True

    for path in save_paths:
        try:
            doc.save(path)
            print(f"Saved DOCX: {path}")
        except PermissionError:
            print(f"Permission error saving {path}")

def create_excel_tkb(t_key, data, save_path):
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "TKB Cá Nhân"

    thin_border = Border(
        left=Side(style='thin', color='000000'),
        right=Side(style='thin', color='000000'),
        top=Side(style='thin', color='000000'),
        bottom=Side(style='thin', color='000000')
    )

    ws.merge_cells('A1:F1')
    ws['A1'] = "TRƯỜNG TIỂU HỌC VÀ THCS UNIGO"
    ws['A1'].font = Font(name='Times New Roman', size=13, bold=True)
    ws['A1'].alignment = Alignment(horizontal='center')

    ws.merge_cells('A2:F2')
    ws['A2'] = "THỜI KHOÁ BIỂU CÁ NHÂN GIÁO VIÊN"
    ws['A2'].font = Font(name='Times New Roman', size=15, bold=True)
    ws['A2'].alignment = Alignment(horizontal='center')

    ws.merge_cells('A3:F3')
    ws['A3'] = f"Giáo viên: {data['full_name']} | Bộ môn: {data['subject']} | Tổng số: {len(data['slots'])} tiết/tuần (Theo TKB toàn trường CHECK - chuẩn)"
    ws['A3'].font = Font(name='Times New Roman', size=12, italic=True)
    ws['A3'].alignment = Alignment(horizontal='center')

    headers = ["Buổi / Tiết", "Thứ Hai", "Thứ Ba", "Thứ Tư", "Thứ Năm", "Thứ Sáu"]
    for c_i, h in enumerate(headers, 1):
        cell = ws.cell(5, c_i, h)
        cell.font = Font(name='Times New Roman', size=13, bold=True)
        cell.alignment = Alignment(horizontal='center', vertical='center')
        cell.fill = PatternFill(start_color="D9E1F2", end_color="D9E1F2", fill_type="solid")
        cell.border = thin_border

    # Sáng Header
    ws.merge_cells('A6:F6')
    ws['A6'] = "BUỔI SÁNG"
    ws['A6'].font = Font(name='Times New Roman', size=13, bold=True)
    ws['A6'].alignment = Alignment(horizontal='center')
    ws['A6'].fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")

    for c_i in range(1, 7):
        ws.cell(6, c_i).border = thin_border

    for p_num in range(1, 6):
        r_idx = 6 + p_num
        ws.cell(r_idx, 1, f"Tiết {p_num}").font = Font(name='Times New Roman', size=13, bold=True)
        ws.cell(r_idx, 1).alignment = Alignment(horizontal='center', vertical='center')
        ws.cell(r_idx, 1).border = thin_border
        for col_i in range(2, 7):
            ws.cell(r_idx, col_i).border = thin_border
            ws.cell(r_idx, col_i).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    # Chiều Header
    ws.merge_cells('A12:F12')
    ws['A12'] = "BUỔI CHIỀU"
    ws['A12'].font = Font(name='Times New Roman', size=13, bold=True)
    ws['A12'].alignment = Alignment(horizontal='center')
    ws['A12'].fill = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")

    for c_i in range(1, 7):
        ws.cell(12, c_i).border = thin_border

    for p_num in range(1, 6):
        r_idx = 12 + p_num
        ws.cell(r_idx, 1, f"Tiết {p_num}").font = Font(name='Times New Roman', size=13, bold=True)
        ws.cell(r_idx, 1).alignment = Alignment(horizontal='center', vertical='center')
        ws.cell(r_idx, 1).border = thin_border
        for col_i in range(2, 7):
            ws.cell(r_idx, col_i).border = thin_border
            ws.cell(r_idx, col_i).alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)

    day_to_col = {'2': 2, '3': 3, '4': 4, '5': 5, '6': 6}

    for slot in data['slots']:
        day, sess, p_num, subject, cls_name = slot
        col_i = day_to_col[day]
        row_i = (6 + p_num) if sess == 'Sáng' else (12 + p_num)
        
        ws.cell(row_i, col_i, f"{subject}\n({cls_name})")

    ws.column_dimensions['A'].width = 16
    for c_let in ['B', 'C', 'D', 'E', 'F']:
        ws.column_dimensions[c_let].width = 22

    wb.save(save_path)
    print(f"Saved EXCEL: {save_path}")

def main():
    print("Loading Timetable data from 'TKB toàn trường CHECK - chuẩn.xlsx'...")
    teacher_data = load_teacher_data_from_excel(TKB_EXCEL_SOURCE)
    
    for t_key, data in teacher_data.items():
        fname_word1 = os.path.join(OUT_DIR_WORD, f"Thời khóa biểu - {data['full_name']}.docx")
        fname_word2 = os.path.join(OUT_DIR_MAU, f"Thời khóa biểu - {data['full_name']}.docx")
        fname_excel = os.path.join(OUT_DIR_WORD, f"Thời khóa biểu - {data['full_name']}.xlsx")
        
        create_docx_tkb(t_key, data, [fname_word1, fname_word2])
        create_excel_tkb(t_key, data, fname_excel)

if __name__ == '__main__':
    main()
